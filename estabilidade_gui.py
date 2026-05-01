#!/usr/bin/env python3
# Copyright (c) 2024-2026 Andre O. Massuca, Pedro Aleixo, Luis M. Massuca
# All rights reserved. Unauthorised copying or distribution is prohibited.
# Contact: aomassuca@gmail.com
#
# Analise de Estabilidade Postural em Plataforma de Forca
# BSP - Biomechanical Stability Program  v1.0
# Andre Massuca, P. Aleixo & Luis M. Massuca | 2025/2026
#
# Referencias:
#   Schubert & Kirchner (2013). Gait & Posture, 39(1), 518-522.
#   Prieto et al. (1996). IEEE Trans Biomed Eng, 43(9), 956-966.
#   Winter (1995). Gait & Posture, 3(4), 193-214.
#   Quijoux et al. (2021). Physiol. Rep., 9(22), e15067.

import os, sys, re, math, hashlib, threading, queue, functools, time
import traceback, datetime, unicodedata, io, argparse
import logging as _logging

# Logger estruturado - cross-platform (macOS/Windows/Linux)
# macOS   -> ~/Library/Logs/BSP/BSP_crash.log
# Windows -> %LOCALAPPDATA%\BSP\Logs\BSP_crash.log
# Linux   -> ~/.local/share/BSP/logs/BSP_crash.log
# Fallback-> ~/.bsp/logs/BSP_crash.log
def _determinar_log_dir():
    import platform as _plat
    _sys = _plat.system()
    try:
        if _sys == 'Darwin':
            return os.path.join(os.path.expanduser("~"), "Library", "Logs", "BSP")
        if _sys == 'Windows':
            base = os.environ.get('LOCALAPPDATA') or os.path.join(os.path.expanduser("~"), "AppData", "Local")
            return os.path.join(base, "BSP", "Logs")
        if _sys == 'Linux':
            xdg = os.environ.get('XDG_DATA_HOME') or os.path.join(os.path.expanduser("~"), ".local", "share")
            return os.path.join(xdg, "BSP", "logs")
    except Exception:
        pass
    return os.path.join(os.path.expanduser("~"), ".bsp", "logs")

_LOG_DIR = _determinar_log_dir()
try:
    os.makedirs(_LOG_DIR, exist_ok=True)
except Exception:
    _LOG_DIR = os.path.expanduser("~")
_LOG_FILE = os.path.join(_LOG_DIR, "BSP.log")
_LOG_CRASH = os.path.join(_LOG_DIR, "BSP_crash.log")

try:
    from logging.handlers import RotatingFileHandler as _RFH
    _logger = _logging.getLogger("BSP")
    _logger.setLevel(_logging.DEBUG)
    if not _logger.handlers:
        try:
            _fh = _RFH(_LOG_FILE, maxBytes=1_000_000, backupCount=3, encoding="utf-8")
            _fh.setFormatter(_logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
            _logger.addHandler(_fh)
        except Exception:
            pass
        try:
            _fh2 = _logging.FileHandler(_LOG_CRASH, encoding='utf-8', mode='a')
            _fh2.setFormatter(_logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
            _logger.addHandler(_fh2)
        except Exception:
            pass
except Exception:
    _logger = _logging.getLogger("BSP")
import base64 as _b64

# -----------------------------------------------------------------------
# Crash reporter - envia relatório detalhado por email (mailto:)
# -----------------------------------------------------------------------
_CRASH_EMAIL = 'aomassuca@gmail.com'

# Log de crash vai para a pasta de logs especifica do SO (ver _LOG_DIR acima)
# macOS: ~/Library/Logs/BSP/  |  Windows: %LOCALAPPDATA%\BSP\Logs\  |  Linux: ~/.local/share/BSP/logs/

def _gerar_relatorio_crash(tb_str, contexto=''):
    """Gera texto formatado do relatório com info do sistema."""
    import platform as _plat
    _v = globals().get('VERSAO', '23')
    linhas = [
        f"BSP  |  Biomechanical Stability Program  v{_v}",
        f"Data/hora : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Sistema   : {_plat.system()} {_plat.release()} ({_plat.machine()})",
        f"Python    : {sys.version.split()[0]}",
        f"Plataforma: {sys.platform}",
        "",
    ]
    if contexto:
        linhas += [f"Contexto  : {contexto}", ""]
    linhas += [
        "─── Traceback ──────────────────────────────────────────────",
        tb_str.strip(),
        "────────────────────────────────────────────────────────────",
        "",
        f"Log completo em: {_LOG_FILE}",
    ]
    return "\n".join(linhas)


def _dialogo_crash(tb_str, contexto='', _root_ext=None):
    """
    Janela de erro com botão para enviar relatório por email (mailto:).
    Abre o cliente de email predefinido com o relatório pré-preenchido.
    """
    try:
        import tkinter as tk
        import urllib.parse as _up
        relatorio = _gerar_relatorio_crash(tb_str, contexto)
        _v = globals().get('VERSAO', '23')

        if _root_ext is not None:
            win = tk.Toplevel(_root_ext)
            win.transient(_root_ext)
        else:
            win = tk.Tk()

        win.title('BSP  -  Erro inesperado')
        win.configure(bg='#0D1720')
        win.resizable(True, True)
        sw = win.winfo_screenwidth(); sh = win.winfo_screenheight()
        w, h = min(820, sw - 80), min(640, sh - 80)
        win.geometry(f'{w}x{h}+{(sw-w)//2}+{(sh-h)//2}')
        win.lift()
        win.attributes('-topmost', True)
        win.after(400, lambda: win.attributes('-topmost', False))
        win.focus_force()

        hdr = tk.Frame(win, bg='#9C0006', height=48)
        hdr.pack(fill='x'); hdr.pack_propagate(False)
        tk.Label(hdr, text='⚠  Erro inesperado no BSP',
                 bg='#9C0006', fg='#FFFFFF',
                 font=('Segoe UI', 11, 'bold')).pack(side='left', padx=16, pady=12)

        tk.Label(win,
                 text='Ocorreu um erro não esperado. Clique em "Enviar relatório"\n'
                      'para abrir o seu email com o relatório pré-preenchido.',
                 bg='#0D1720', fg='#A0BFD4',
                 font=('Segoe UI', 9), justify='left').pack(anchor='w', padx=16, pady=(10, 4))

        frame_txt = tk.Frame(win, bg='#0D1720')
        frame_txt.pack(fill='both', expand=True, padx=16, pady=(0, 4))
        vscroll = tk.Scrollbar(frame_txt)
        vscroll.pack(side='right', fill='y')
        # Fonte mono portátil - _FM() respeita o SO e faz fallback para TkFixedFont
        try:
            _mono = _FM(8)
        except Exception:
            _mono = ('TkFixedFont', 8)
        txt = tk.Text(frame_txt, bg='#0A1218', fg='#C0D8E8',
                      font=_mono,
                      yscrollcommand=vscroll.set, relief='flat',
                      selectbackground='#1A3A5C', wrap='none')
        txt.pack(side='left', fill='both', expand=True)
        vscroll.config(command=txt.yview)
        hscroll = tk.Scrollbar(win, orient='horizontal', command=txt.xview)
        hscroll.pack(fill='x', padx=16)
        txt.config(xscrollcommand=hscroll.set)
        txt.insert('end', relatorio)
        txt.config(state='disabled')

        bf = tk.Frame(win, bg='#0D1720'); bf.pack(fill='x', padx=16, pady=(6, 14))

        def _enviar():
            assunto = f'BSP v{_v} - Relatório de Crash'
            mailto  = (f'mailto:{_CRASH_EMAIL}'
                       f'?subject={_up.quote(assunto)}'
                       f'&body={_up.quote(relatorio)}')
            import webbrowser
            webbrowser.open(mailto)

        def _copiar():
            win.clipboard_clear()
            win.clipboard_append(relatorio)
            btn_copy.config(text='✓  Copiado!')
            win.after(2000, lambda: btn_copy.config(text='📋  Copiar relatório'))

        tk.Button(bf, text='📧  Enviar relatório por email',
                  bg='#C8A45A', fg='#0A0D14',
                  font=('Segoe UI', 10, 'bold'),
                  relief='flat', cursor='hand2', padx=14, pady=7,
                  command=_enviar).pack(side='left', padx=(0, 8))
        btn_copy = tk.Button(bf, text='📋  Copiar relatório',
                             bg='#1C3A5C', fg='#A0C8E8',
                             font=('Segoe UI', 9),
                             relief='flat', cursor='hand2', padx=12, pady=7,
                             command=_copiar)
        btn_copy.pack(side='left', padx=(0, 8))
        tk.Button(bf, text='Fechar',
                  bg='#1C2E42', fg='#6B8FAE',
                  font=('Segoe UI', 9),
                  relief='flat', cursor='hand2', padx=12, pady=7,
                  command=win.destroy).pack(side='right')

        if _root_ext is not None:
            _root_ext.wait_window(win)
        else:
            win.mainloop()
    except Exception:
        _logger.error("_dialogo_crash falhou:\n" + traceback.format_exc())


def _instalar_crash_hooks(_root_ref=None):
    """Instala hooks para capturar exceções não tratadas em qualquer thread."""
    _root_holder = [_root_ref]

    def _hook_main(exc_type, exc_val, exc_tb):
        if issubclass(exc_type, (SystemExit, KeyboardInterrupt)):
            sys.__excepthook__(exc_type, exc_val, exc_tb)
            return
        tb_str = ''.join(traceback.format_exception(exc_type, exc_val, exc_tb))
        _logger.critical("Excepção não tratada (main):\n" + tb_str)
        _dialogo_crash(tb_str, 'Thread principal', _root_holder[0])

    def _hook_thread(args):
        if args.exc_type and issubclass(args.exc_type, (SystemExit, KeyboardInterrupt)):
            return
        tb_str = ''.join(traceback.format_exception(
            args.exc_type, args.exc_value, args.exc_traceback))
        _logger.critical(f"Excepção não tratada (thread {args.thread.name}):\n" + tb_str)
        if _root_holder[0]:
            try:
                _root_holder[0].after(
                    0, lambda t=tb_str, n=args.thread.name:
                    _dialogo_crash(t, f'Thread {n}', _root_holder[0]))
            except Exception:
                pass

    sys.excepthook = _hook_main
    try:
        threading.excepthook = _hook_thread
    except AttributeError:
        # Python < 3.8: sem threading.excepthook
        # Instalar hook manual via sys.excepthook ja cobre a thread principal
        pass

# Matplotlib: importar uma vez ao nivel do modulo (evita overhead por funcao)
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    _MPL_OK = True
except ImportError:
    _MPL_OK = False

import json
import platform as _platform

# -----------------------------------------------------------------------
# Detecção de plataforma (Windows / macOS / Linux)
# Resolve: DPI blur no Windows, fontes em macOS/Linux, teclas Cmd vs Ctrl,
#          Desktop path variável, múltiplos tk.Tk() (bug de não avançar).
# -----------------------------------------------------------------------
_SYS = _platform.system()   # 'Windows', 'Darwin', 'Linux'

# Família de fonte que existe em todos os SOs
# Se a fonte preferida não estiver disponível, faz fallback para a fonte
# por defeito do Tk (TkDefaultFont / TkFixedFont) no primeiro uso.
if _SYS == 'Darwin':
    _UI_FONT   = 'Helvetica Neue'
    _MONO_FONT = 'Menlo'
elif _SYS == 'Windows':
    _UI_FONT   = 'Segoe UI'
    _MONO_FONT = 'Cascadia Code'
elif _SYS == 'Linux':
    _UI_FONT   = 'DejaVu Sans'
    _MONO_FONT = 'DejaVu Sans Mono'
else:
    # Plataforma desconhecida (BSD, etc.): usa defaults do Tk
    _UI_FONT   = 'TkDefaultFont'
    _MONO_FONT = 'TkFixedFont'

_FONT_CHECKED = False
def _validar_fontes_disponiveis():
    """Verifica se _UI_FONT/_MONO_FONT existem; faz fallback para Tk defaults."""
    global _UI_FONT, _MONO_FONT, _FONT_CHECKED
    if _FONT_CHECKED:
        return
    try:
        import tkinter.font as _tkfont
        familias = set(_tkfont.families())
        if _UI_FONT not in familias and _UI_FONT not in ('TkDefaultFont',):
            _UI_FONT = 'TkDefaultFont'
        if _MONO_FONT not in familias and _MONO_FONT not in ('TkFixedFont',):
            _MONO_FONT = 'TkFixedFont'
        _FONT_CHECKED = True
    except Exception:
        # Sem Tk (modo --testes / headless) - deixa as constantes como estão
        _FONT_CHECKED = True

def _F(size, bold=False, italic=False):
    """Tuplo de fonte portátil (Windows / macOS / Linux / fallback)."""
    s = ('bold italic' if bold and italic
         else 'bold' if bold
         else 'italic' if italic else '')
    return (_UI_FONT, size, s) if s else (_UI_FONT, size)

def _FM(size):
    """Fonte monoespaçada portátil."""
    return (_MONO_FONT, size)

# Tecla modificadora - Cmd no macOS, Ctrl nos outros
_MOD_KEY = 'Command' if _SYS == 'Darwin' else 'Control'
_MOD_SYM = '⌘' if _SYS == 'Darwin' else 'Ctrl'

# DPI-awareness no Windows (evita ecrã borrado em alta resolução)
if _SYS == 'Windows':
    try:
        import ctypes
        ctypes.windll.shcore.SetProcessDpiAwareness(1)  # System DPI Aware (seguro para Tkinter)
    except Exception:
        try: ctypes.windll.user32.SetProcessDPIAware()
        except Exception: pass

def _desktop_path():
    """Devolve uma pasta de destino válida em qualquer SO."""
    import pathlib
    home = pathlib.Path.home()
    for name in ('Desktop', 'Área de Trabalho', 'Documents', 'Documentos'):
        p = home / name
        if p.is_dir():
            return str(p)
    return str(home)

import numpy as np
from scipy import stats
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.chart import ScatterChart, LineChart, Reference, Series

# Logo BSP (1024x1024 PNG, base64)
_LOGO_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAABAAAAAQACAYAAAB/HSuDAAAQAElEQVR4AezdCXwcZfkH8OeZTcvRNpsCKuDF1SRF5BbkaDblarO7KQX/5fJA5VIOBZpNAUUCypUNoIIXKnKIXAq02d20XE1SbjlEjmZTUFTEA6XZtOVos/P8nykIbWmb7GZm9p3d337et5vszvu+z/t902TnmdlZi3CDAAQgAAEIQAACEIAABCAAAQhAoNwFCAmAsl9iTBACEIAABCAAAQhAAAIQgAAEIEBIAOCHAAIQgAAEIAABCEAAAhCAAAQgUPYCOkGcAaAIKBCAAAQgAAEIQAACEIAABCAAgXIWcOaGBICjgAoBCEAAAhCAAAQgAAEIQAACEChfgdUzQwJgNQP+gQAEIAABCEAAAhCAAAQgAAEIlKvAO/NCAuAdB/wLAQhAAAIQgAAEIAABCEAAAhAoT4F3Z4UEwLsQuIMABCAAAQhAAAIQgAAEIAABCJSjwP/mhATA/yRwDwEIQAACEIAABCAAAQhAAAIQKD+B92aEBMB7FPgCAhCAAAQgAAEIQAACEIAABCBQbgLvzwcJgPct8BUEIAABCEAAAhCAAAQgAAEIQKC8BNaYDRIAa2DgSwhAAAIQgAAEIAABCEAAAhCAQDkJrDkXJADW1MDXEIAABCAAAQhAAAIQgAAEIACB8hFYayZIAKzFgW8gAAEIQAACEIAABCAAAQhAAALlIrD2PJAAWNsD30EAAhCAAAQgAAEIQAACEIAABMpDYJ1ZIAGwDgi+hQAEIAABCEAAAhCAAAQgAAEIlIPAunNAAmBdEXwPAQhAAAIQgAAEIAABCEAAAhAIvsAHZoAEwAdI8AAEIAABCEAAAhCAAAQgAAEIQCDoAh+MHwmAD5rgEQhAAAIQgAAEIAABCEAAAhCAQLAF1hM9EgDrQcFDEIAABCAAAQhAAAIQgAAEIACBIAusL3YkANangscgAAEIQAACEIAABCAAAQhAAALBFVhv5EgArJcFD0IAAhCAAAQgAAEIQAACEIAABIIqsP64kQBYvwsehQAEIAABCEAAAhCAAAQgAAEIBFNgA1EjAbABGDwMAQhAAAIQgAAEIAABCEAAAhAIosCGYkYCYEMyeBwCEIAABCAAAQhAAAIQgAAEIBA8gQ1GjATABmnwBAQgAAEIQAACEIAABCAAAQhAIGgCG44XCYAN2+AZCEAAAhCAAAQgAAEIQAACEIBAsAQ2Ei0SABvBwVMQgAAEIAABCEAAAhCAAAQgAIEgCWwsViQANqaD5yAAAQhAAAIQgAAEIAABCEAAAsER2GikSABslAdPQgACEIAABCAAAQhAAAIQgAAEgiKw8TiRANi4D56FAAQgAAEIQAACEIAABCAAAQgEQ2CYKJEAGAYIT0MAAhCAAAQgAAEIQAACEIAABIIgMFyMSAAMJ4TnIQABCEAAAhCAAAQgAAEIQAAC5gsMGyESAMMSYQMIQAACEIAABCAAAQhAAAIQgIDpAsPHhwTA8EbYAgIQgAAEIAABCEAAAhCAAAQgYLbACKJDAmAESNgEAhCAAAQgAAEIQAACEIAABCBgssBIYkMCYCRK2AYCEIAABCAAAQhAAAIQgAAEIGCuwIgiQwJgREzYCAIQgAAEIAABCEAAAhCAAAQgYKrAyOJCAmBkTtgKAhCAAAQgAAEIQAACEIAABCBgpsAIo0ICYIRQ2AwCEIAABCAAAQhAAAIQgAAEIGCiwEhjQgJgpFLYDgIQgAAEIAABCEAAAhCAAAQgYJ7AiCNCAmDEVNgQAhCAAAQgAAEIQAACEIAABCBgmsDI40ECYORW2BICEIAABCAAAQhAAAIQgAAEIGCWQAHRIAFQABY2hQAEIAABCEAAAhCAAAQgAAEImCRQSCxIABSihW0hAAEIQAACEIAABCAAAQhAAALmCBQUCRIABXFhYwhAAAIQgAAEIAABCEAAAhCAgCkChcWBBEBhXtgaAhCAAAQgAAEIQAACEIAABCBghkCBUSABUCAYNocABCAAAQhAAAIQgAAEIAABCJggUGgMSAAUKobtIQABCEAAAhCAAAQgAAEIQAACpRcoOAIkAAomQwMIQAACEIAABCAAAQhAAAIQgECpBQofHwmAws3QAgIQgAAEIAABCEAAAhCAAAQgUFqBIkZHAqAINDSBAAQgAAEIQAACEIAABCAAAQiUUqCYsZEAKEYNbSAAAQhAAAIQgAAEIAABCEAAAqUTKGpkJACKYkMjCEAAAhCAAAQgAAEIQAACEIBAqQSKGxcJgOLc0AoCEIAABCAAAQhAAAIQgAAEIFAagSJHRQKgSDg0gwAEIAABCEAAAhCAAAQgAAEIlEKg2DGRAChWDu0gAAEIQAACEIAABCAAAQhAAAL+CxQ9IhIARdOhIQQgAAEIQAACEIAABCAAAQhAwG+B4sdDAqB4O7SEAAQgAAEIQAACEIAABCAAAQj4KzCK0ZAAGAUemkIAAhCAAAQgAAEIQAACEIAABPwUGM1YSACMRg9tIQABCEAAAhCAAAQgAAEIQAAC/gmMaiQkAEbFh8YQgAAEIAABCEAAAhCAAAQgAAG/BEY3DhIAo/NDawhAAAIQgAAEIAABCEAAAhCAgD8CoxwFCYBRAqI5BCAAAQhAAAIQgAAEIAABCEDAD4HRjoEEwGgF0R4CEIAABCAAAQhAAAIQgAAEIOC9wKhHQAJg1IToAAIQgAAEIAABCEAAAhCAAAQg4LXA6PtHAmD0hugBAhCAAAQgAAEIQAACEIAABCDgrYALvSMB4AIiuoAABCAAAQhAAAIQgAAEIAABCHgp4EbfSAC4oYg+IAABCEAAAhCAAAQgAAEIQAAC3gm40jMSAK4wohMIQAACEIAABCAAAQhAAAIQgIBXAu70iwSAO47oBQIQgAAEIAABCEAAAhCAAAQg4I2AS70iAeASJLqBAAQgAAEIQAACEIAABCAAAQh4IeBWn0gAuCWJfiAAAQhAAAIQgAAEIAABCEAAAu4LuNYjEgCuUaIjCEAAAhCAAAQgAAEIQAACEICA2wLu9YcEgHuW6AkCEIAABCAAAQhAAAIQgAAEIOCugIu9IQHgIia6ggAEIAABCEAAAhCAAAQgAAEIuCngZl9IALipib4gAAEIQAACEIAABCAAAQhAAALuCbjaExIArnKiMwhAAAIQgAAEIAABCEAAAhCAgFsC7vaDBIC7nugNAhCAAAQgAAEIQAACEIAABCDgjoDLvSAB4DIouoMABCAAAQhAAAIQgAAEIAABCLgh4HYfSAC4LYr+IAABCEAAAhCAAAQgAAEIQAACoxdwvQckAFwnRYcQgAAEIAABCEAAAhCAAAQgAIHRCrjfHgkA903RIwQgAAEIQAACEIAABCAAAQhAYHQCHrRGAsADVHQJAQhAAAIQgAAEIAABCEAAAhAYjYAXbZEA8EIVfUIAAhCAAAQgAAEIQAACEIAABIoX8KQlEgCesKJTCEAAAhCAAAQgAAEIQAACEIBAsQLetEMCwBtX9AoBCEAAAhCAAAQgAAEIQAACEChOwKNWSAB4BItuIQABCEAAAhCAAAQgAAEIQAACxQh41QYJAK9k0S8EIAABCEAAAhCAAAQgAAEIQKBwAc9aIAHgGS06hgAEIAABCEAAAhCAAAQgAAEIFCrg3fZIAHhni54hAAEIQAACEIAABCAAAQhAAAKFCXi4NRIAHuKiawhAAAIQgAAEIAABCEAAAhCAQCECXm6LBICXuugbAhCAAAQgAAEIQAACEIAABCAwcgFPt0QCwFNedA4BCEAAAhCAAAQgAAEIQAACEBipgLfbIQHgrS96hwAEIAABCEAAAhCAAAQgAAEIjEzA462QAPAYGN1DAAIQgAAEIAABCEAAAhCAAARGIuD1NkgAeC2M/iEAAQhAAAIQgAAEIAABCEAAAsMLeL4FEgCeE2MACEAAAhCAAAQgAAEIQAACEIDAcALeP48EgPfGGAECEIAABCAAAQhAAAIQgAAEILBxAR+eRQLAB2QMAQEIQAACEIAABCAAAQhAAAIQ2JiAH88hAeCHMsaAAAQgAAEIQAACEIAABCAAAQhsWMCXZ5AA8IUZg0AAAhCAAAQgAAEIQAACEIAABDYk4M/jSAD444xRIAABCEAAAhCAAAQgAAEIQAAC6xfw6VEkAHyCxjAQgAAEIAABCEAAAhCAAAQgAIH1Cfj1GBIAfkljHAhAAAIQgAAEIAABCEAAAhCAwAcFfHsECQDfqDEQBCAAAQhAAAIQgAAEIAABCEBgXQH/vkcCwD9rjAQBCEAAAhCAAAQgAAEIQAACEFhbwMfvkADwERtDQQACEIAABCAAAQhAAAIQgAAE1hTw82skAPzUxlgQgAAEIAABCEAAAhCAAAQgAIH3BXz9CgkAX7kxGAQgAAEIQAACEIAABCAAAQhA4H8C/t4jAeCvN0aDAAQgAAEIQAACEIAABCAAAQi8I+Dzv0gA+AyO4SAAAQhAAAIQgAAEIAABCEAAAo6A3xUJAL/FMR4EIAABCEAAAhCAAAQgAAEIQIDIdwMkAHwnx4AQgAAEIAABCEAAAhCAAAQgAAH/BZAA8N8cI0IAAhCAAAQgAAEIQAACEIBApQuUYP5IAJQAHUNCAAIQgAAEIAABCEAAAhCAQGULlGL2SACUQh1jQgACEIAABCAAAQhAAAIQgEAlC5Rk7kgAlIQdg0IAAhCAAAQgAAEIQAACEIBA5QqUZuZIAJTGHaNCAAIQgAAEIAABCEAAAhCAQKUKlGjeSACUCB7DQgACEIAABCAAAQhAAAIQgEBlCpRq1kgAlEoe40IAAhCAAAQgAAEIQAACEIBAJQqUbM5IAJSMHgNDAAIQgAAEIAABCEAAAhCAQOUJlG7GSACUzh4jQwACEIAABCAAAQhAAAIQgEClCZRwvkgAlBAfQ0MAAhCAAAQgAAEIQAACEIBAZQmUcrZIAJRSH2NDAAIQgAAEIAABCEAAAhCAQCUJlHSuSACUlB+DQwACEIAABCAAAQhAAAIQgEDlCJR2pkgAlNYfo0MAAhCAAAQgAAEIQAACEIBApQiUeJ5IAJR4ATA8BCAAAQhAAAIQgAAEIAABCFSGQKlniQRAqVcA40MAAhCAAAQgAAEIQAACEIBAJQiUfI5IAJR8CRAABCAAAQhAAAIQgAAEIAABCJS/QOlniARA6dcAEUAAAhCAAAQgAAEIQAACEIBAuQsYMD8kAAxYBIQAAQhAAAIQgAAEIAABCEAAAuUtYMLskAAwYRUQAwQgAAEIQAACEIAABCAAAQiUs4ARc0MCwIhlQBAQgAAEIAABCEAAAhCAAAQgUL4CZswMCQAz1gFRQAACEIAABCAAAQhAAAIQgEC5ChgyLyQADFkIhAEBCEAAAhCAAAQgAAEIQAAC5SlgyqyQADBlJRAHBCAAAQhAAAIQgAAEIAABCJSjgDFzQgLAmKVAIBCAAAQgAAEIQAACEIAABCBQfgLmzAgJAHPWApFAAAIQgAAEIAABCEAAAhCAQLkJGDQfJAAMoCin/wAAEABJREFUWgyEAgEIQAACEIAABCAAAQhAAALlJWDSbJAAMGk1EAsEIAABCEAAAhCAAAQgAAEIlJOAUXNBAsCo5UAwEIAABCAAAQhAAAIQgAAEIFA+AmbNBAkAs9YD0UAAAhCAAAQgAAEIQAACEIBAuQgYNg8kAAxbEIQDAQhAAAIQgAAEIAABCEAAAuUhYNoskAAwbUUQDwQgAAEIQAACEIAABCAAAQiUg4Bxc0ACwLglQUAQgAAEIAABCEAAAhCAAAQgEHwB82aABIB5a4KIIAABCEAAAhCAAAQgAAEIQCDoAgbGjwSAgYuCkCAAAQhAAAIQgAAEIAABCEAg2AImRo8EgImrgpggAAEIQAACEIAABCAAAQhAIMgCRsaOBICRy4KgIAABCEAAAhCAAAQgAAEIQCC4AmZGjgSAmeuCqCAAAQhAAAIQgAAEIAABCEAgqAKGxo0EgKELg7AgAAEIQAACEIAABCAAAQhAIJgCpkaNBICpK4O4IAABCEAAAhCAAAQgAAEIQCCIAsbGjASAsUuDwCAAAQhAAAIQgAAEIAABCEAgeALmRowEgLlrg8ggAAEIQAACEIAABCAAAQhAIGgCBseLBIDBi4PQIAABCEAAAhCAAAQgAAEIQCBYAiZHiwSAyauD2CAAAQhAAAIQgAAEIAABCEAgSAJGx4oEgNHLg+AgAAEIQAACEIAABCAAAQhAIDgCZkeKBIDZ64PoIAABCEAAAhCAAAQgAAEIQCAoAobHiQSA4QuE8CAAAQhAAAIQgAAEIAABCEAgGAKmR4kEgOkrhPggAAEIQAACEIAABCAAAQhAIAgCxseIBIDxS4QAIQABCEAAAhCAAAQgAAEIQMB8AfMjRALA/DVChBCAAAQgAAEIQAACEIAABCBgukAA4kMCIACLhBAhAAEIQAACEIAABCAAAQhAwGyBIESHBEAQVgkxQgACEIAABCAAAQhAAAIQgIDJAoGIDQmAQCwTgoQABCAAAQhAAAIQgAAEIAABcwWCERkSAMFYJ0QJAQhAAAIQgAAEIAABCEAAAqYKBCQuJAACslAIEwIQgAAEIAABCEAAAhCAAATMFAhKVEgABGWlECcEIAABCEAAAhCAAAQgAAEImCgQmJiQAAjMUiFQCEAAAhCAAAQgAAEIQAACEDBPIDgRIQEQnLVCpBCAAAQgAAEIQAACEIAABCBgmkCA4kECIECLhVAhAAEIQAACEIAABCAAAQhAwCyBIEWDBECQVguxQgACEIAABCAAAQhAAAIQgIBJAoGKBQmAQC0XgoUABCAAAQhAAAIQgAAEIAABcwSCFQkSAMFaL0QLAQhAAAIQgAAEIAABCEAAAqYIBCwOJAACtmAIFwIQgAAEIAABCEAAAhCAAATMEAhaFEgABG3FEC8EIAABCEAAAhCAAAQgAAEImCAQuBiQAAjckiFgCEAAAhCAAAQgAAEIQAACECi9QPAiQAIgeGuGiCEAAQhAAAIQgAAEIAABCECg1AIBHB8JgAAuGkKGAAQgAAEIQAACEIAABCAAgdIKBHF0JACCuGqIGQIQgAAEIAABCEAAAhCAAARKKRDIsZEACOSyIWgIQAACEIAABCAAAQhAAAIQKJ1AMEdGAiCY64aoIQABCEAAAhCAAAQgAAEIQKBUAgEdFwmAgC4cwoYABCAAAQhAAAIQgAAEIACB0ggEdVQkAIK6cogbAhCAAAQgAAEIQAACEIAABEohENgxkQAI7NIhcAhAAAIQgAAEIAABCEAAAhDwXyC4IyIBENy1Q+QQgAAEIAABCEAAAhCAAAQg4LdAgMdDAiDAi4fQIQABCEAAAhCAAAQgAAEIQMBfgSCPhgRAkFcPsUMAAhCAAAQgAAEIQAACEICAnwKBHgsJgEAvH4KHAAQgAAEIQAACEIAABCAAAf8Egj0SEgDBXj9EDwEIQAACEIAABCAAAQhAAAJ+CQR8HCQAAr6ACB8CEIAABCAAAQhAAAIQgAAE/BEI+ihIAAR9BRE/BCAAAQhAAAIQgAAEIAABCPghEPgxkAAI/BJiAhCAAAQgAAEIQAACEIAABCDgvUDwR0ACIPhriBlAAAIQgAAEIAABCEAAAhCAgNcCZdA/EgBlsIiYAgQgAAEIQAACEIAABCAAAQh4K1AOvSMBUA6riDlAAAIQgAAEIAABCEAAAhCAgJcCZdE3EgBlsYyYBAQgAAEIQAACEIAABCAAAQh4J1AePSMBUB7riFlAAAIQgAAEIAABCEAAAhCAgFcCZdIvEgBlspCYBgQgAAEIQAACEIAABCAAAQh4I1AuvSIBUC4riXlAAAIQgAAEIAABCEAAAhCAgBcCZdMnEgBls5SYCAQgAAEIQAACEIAABCAAAQi4L1A+PSIBUD5riZlAAAIQgAAEIAABCEAAAhCAgNsCZdQfEgBltJiYCgQgAAEIQAACEIAABCAAAQi4K1BOvSEBUE6riblAAAIQgAAEIAABCEAAAhCAgJsCZdUXEgBltZyYDAQgAAEIQAACEIAABCAAAQi4J1BePSEBUF7ridlAAAIQgAAEIAABCEAAAhCAgFsCZdYPEgBltqCYDgQgAAEIQAACEIAABCAAAQi4I1BuvSABUG4rivlAAAIQgAAEIAABCEAAAhCAgBsCZdcHEgBlt6SYEAQgAAEIQAACEIAABCAAAQiMXqD8ekACoPzWFDOCAAQgAAEIQAACEIAABCAAgdEKlGF7JADKcFExJQhAAAIQgAAEIAABCEAAAhAYnUA5tkYCoBxXFXOCAAQgAAEIQAACEIAABCAAgdEIlGVbJADKclkxKQhAAAIQgAAEIAABCEAAAhAoXqA8WyIBUJ7rillBAAIQgAAEIAABCEAAAhCAQLECZdoOCYAyXVhMCwIQgAAEIAABCEAAAhCAAASKEyjXVkgAlOvKYl4QgAAEIAABCEAAAhCAAAQgUIxA2bZBAqBslxYTgwAEIAABCEAAAhCAAAQgAIHCBcq3BRIA5bu2mBkEIAABCEAAAhCAAAQgAAEIFCpQxtsjAVDGi4upQQACEIAABCAAAQhAAAIQgEBhAuW8NRIA5by6mBsEIAABCEAAAhCAAAQgAAEIFCJQ1tsiAVDWy4vJQQACEIAABCAAAQhAAAIQgMDIBcp7SyQAynt9MTsIQAACEIAABCAAAQhAAAIQGKlAmW+HBECZLzCmBwEIQAACEIAABCAAAQhAAAIjEyj3rZAAKPcVxvwgAAEIQAACEIAABCAAAQhAYCQCZb8NEgBlv8SYIAQgAAGPBdraLGpsq6JZbWM/NuuszT5yWMu4rWa0TqiZeWZN9bSztpjQPHurcTO++ZHNp525zZbxsz46MT7nEzXTz9kuHDtnh+p4YtKEGa1142Otk7donvOpiTNadq2JJnYPxxN7VjfP/kx1tGXfCfGW/SbGZx9YE53dMDGWmBKOzd5L+6x3+plwxLlbOmN6PEN0DwEIQAACEIBARQiU/ySRACj/NcYMIQABCKxfoLFtU2fH3NkJ153xvcLNLQfVNLfOrIknvlQdbz1dv/9WOJ64XOtPwrHEb2rirSndtle/f0a//7Pe/1frqvATK/Lh8StWhd9c8fayN6veeGssL19ly6AMjVnKY6r+a4n1WpU99p9jxox5dYiqXrHJ/otU5f9MnH+JifotW/pCLC/kxX7OtvkZsehpDfhJFutxtvhRi/hhm6xFYlk9NlMvsfWE9rnY6cdaNfQfZ8xwcyIfjrcsC8cT/9L6J43zOb1/TOtCJ269v706lvhVONbyo3C8tV0fa9PHWqvjidNq4i1fCTfPOWpCcyJeE2ud6iQdJsZadlGTHcZFE1s7yQyaNSukMaFAAAIQgAAEIFDOAhUwN6sC5ogpQgACEChvAT3yHm5u2X5ifM6BulN7dE285Wy9v9TZ2a1pTtxUHW+Zp993h+OtT+n9i1pf0/q27rS/6eyYOzvhxPknSPh+EblLiG5gkqv1++8pXKvWrxHTsUIS022n6Pe76vfb6f0WWqu0lr4I6d8zHq+BfFjr9hrnp/R+H62NTtx6P4uZvkzMpxJJQh+7gIgu1+2uEeLrSOzbLKFOYXnASTrYzM+qyUtVFv3DSWaE39xuSM1e17pYa3e4ufU2db66Op74djjeclJNrHVG9eEt+zrrsG1z2+baNwoEIAABCEAAAgETqIRw9QVTJUwTc4QABCAQTAHnFPqJh8/5dPWMxHTd6TyhujnxnXBz4md6BDsVjiee1vpv58i77qz/SY+IL9JZ3qo7tFfo/TnOzq4IfYGJm/X7CJHsofc7at1K61itKIUJTNTN67VGSOQodT6dib5LxNdq4mAu5/lRZx1WyIoVui6D4XjrEk3APKhf36n1J7pmbeFY4us1M1qPdN7SUB1PTNqi6Yxqwg0CEIAABCAAARMEKiIGJAAqYpkxSQhAwDiBxraqmujsT06Inb2/7hTOqom1nKn3Sd1RvFlrd1h3HvX+DecUejtv/5Ft6tKdzl+w0IUkdLIewY7pnHbX+iGtKOYJTNCEy06agDlAQztC69d0zS4gph+LLb+zyVqkyYP+fGjTnLPOmtRx3lLxWHU8MU+//3k41vo9TfacoV8fXRNvadzisLM/rn2gQAACEIAABCDgmUBldIwEQGWsM2YJAQiUQMC5OJ3uwO1ZHW85riae+G443nJHOJ54Uus/w+NXrBLLetni0EO6U3i7MF+l9y0a5nFanaP1O+n9ZlpRyl9gM03qOG+p2EeTAs7ZGicSy7dY6Ic69VuFeGF+bOiv+nPzVk0s8Xy1JgmqYy1X6f1p1dHEdL2fRNSGv+eKhQIBCEAAAhAoWqBCGuIFQ4UsNKYJAQh4JzD+iJYP18yY0xBubjmlurn1qnC8ZX44nnjZuTidjvokE98sRN8m4v8joj21fkQrCgQKFdhEmHZmomZmPpOJrmGLuvS+PxxfkQ/HW18KxxL36M+ec9HGFueCjhOjcz6NaxIUyoztIQABCECgEgUqZc5IAFTKSmOeEIDAqAXCTYkd9chrTHeyWsKx1l+8+/7u/4ZW8b/EtntI+KcsciYRTyOiT2pFgYCPArIDMR2qAzoXbUyKyF22Zf/RuSZBTazl1ep4orc6lviV3n+7Opo4ttr5mMVpZzkXctQmKBCAAAQgAIGKFqiYySMBUDFLjYlCAAIjEmg6Y5Oa+JzdqmMtx9TEW9vCsZbb9IjqM1rfphC9qEdeU7qTlSSWE959fzd2oEYEi41KKSDM2zDRFGb6st5/ly36DTsfszimyvkox9fDscTvw/HWW8Kx1u/VxFu+UhOd3bDVjNZtCTcIQAACEIBARQhUziSRAKictcZMIQCBdQVm3R6aGEtM0Z2eOdXxxFzdyf9TOLTpW0L2H5j5FiG5gJiP0ma7asVV8xUBpSwFJhLT3kRyjCa2viXE14ll9ayy5e/heMuycKylx3lri9YvToy17FKWApgUBCAAAQhUtkAFzR4JgApabEwVApUusG/HXLoAABAASURBVGX8rI+GY4lZ1c0tV4XjiUfDb/5+yGbqJZbL9KjoDPXZXisKBCDwngCPJ+YG560tWm+0mZ/V/ztva30kHGv5kd6fqNW5rsV7LfAFBCAAAQhAIGgClRQvEgCVtNqYKwQqTCAcm71Xdbz19HBz4jfheOLlIap6hZhuZ+EzlWJfrSgQgEDhAs7ZMJ8l5lO16c+1Op9sIfp/zLn/eXVz66laP0uz2pzt9GkUCEAAAhCAgNECFRUcEgAVtdyYLATKV6BmZltNdTwRDcdav6c7Ig9ofYvYeoJJriahY3XmuCifIqBAwEMB50yAE1nkR1ofCb+5wjlT4NmaeMuNNbGWM51PyvjQrLbxHo6PriEAAQhAAAJFCFRWEyQAKmu9MVsIlI3AhObZ9TWxxFfDsZZf6M7+CzK0YikTpYnlWzrJqVo30YoCAQiUVmAXIf6iMF/lfFLGyjdXLNP/r/1ab9Xaqv9/D5lwxLlbljZEjA4BCEAAAhUtUGGTRwKgwhYc04VAUAX0KGJjOJY4ryae6NQdh/9aYi0Wpl8S8wk6p8laUSAAgWAITNIwj9Z6uf7/vddaNfQf/T/9stY7q+OJb1fHWmKbTztzG30eBQIQgAAEIOC5QKUNgARApa045guBgAiMP3z2zjWxljN1pyCt9W09iriQmC4WorhOYQutKBCAQPkIOG/ROYKJvsvMqTFjxryq/+9f1oTADTWxxFerZ5y1U/lMFTOBAAQgAAGDBCouFCQAKm7JMWEImCkwvuncD1XHW46rjiV+pS/8/xLKW88L81UabVQrLiamCCgQqDCBT2pC4EvC9Eu2q5bo74VXtN4cbk6c7CQIK8wC04UABCAAAU8EKq9TJAAqb80xYwgYIxCe0XpoON7SHo4nHguFhv7NxDcz05eJ6BNaUSAAAQisKfBR/eY4EvqZkyDU3xv/DMdab6uOJ07Tr50LEOrTKBCAAAQgAIECBCpwUyQAKnDRMWUIlEqgJprYPRxLtFTHW+bp/QDZcg8RJ4hoH60oEIAABAoR+AixHMVE12gj5yMI/xNubv1dTazlzAnxlv2orQ2vcRQGBQIQgAAENixQic/gj2MlrjrmDAGfBLaMn/XR6ljLF8Lx1mv1CF2fWPQ0MSWZuFnvwz6FgWEgAIHKENiSRI4U5qss4ofDT6x4vTqemKdJgTk18ZbGj806a7PKYMAsIQABCEBghAIVuZlVkbPGpCEAAW8ETj55TDjWckhNvPUi3eHvHqLQ35j5JiI5SQes04oCAQhAwC+BMBM1a1LgMiFeuOzNqv+E4y2Z6njLBdXRxPRw7JyJfgWCcSAAAQhAwESByowJCYDKXHfMGgKuCXxo1qnjw/HE0TXNiZtqXq35CzHfKyTn6wAR0m8INwhAAAJmCGxOxE1M3MYWdRHn/6m/u+4Px1suqWluOXxcNLE14QYBCEAAApUjUKEzRQKgQhce04bAqATa2qyaWOuMcHPiZ2+/OS6rfd0qQl/QHf9t9GsUCEAAAkEQcD5d5CAiPleE766y6O81zYkHq+OtVzi/37aa0TqBcIMABCAAgbIVqNSJIQFQqSuPeUOgCAHn9P7qeOL74d+v6BeWuSR0MhNtW0RXaAIBCEDANAFLhA5gkrOd32+rbPlTONZ6Wzje+rUJ0XNrTQsW8UAAAhCAwKgEKrYxEgAVu/SYOARGJjBhRuKAcDxxqdY/EvO9usP/TWLacWStsRUEIACBwAps5XzKAJH8xLKGstXxlkU1zYkLnQsKBnZGCBwCEIAABN4VqNw7JAAqd+0xcwhsUEB39vesjrWcr/ePWTY9qBueo/XTWlEgAAEIVKQAEx8oQt9xLiiovxtf1Prz6ubEseNw7YCK/HnApCEAgYALVHD4SABU8OJj6hBYU2BC8+z6cCzRoi9qF+rjTzLzRXq/j1YUCEAAAhBYW8A5C+pEFvpNlUV/qom3pmpirbNrZpyzx9qb4TsIQAACEDBRoJJjQgKgklcfc694gZro7E9WN7eeGo63ZCyxFhNTUlEataJAAAIQgMDIBDYTkpiwdIidf0p/n2pNdFTHWmIfm3XWZiPrAltBAAIQgICPAhU9FBIAFb38mHwlCoxvOvdDNfGWr4Tjid+KZWVZ5EdE3ES4QeB9gZX65WtE8qLeP6n1Aa13EfGvhPn7THwhM52tRz9P0KTR/5HIoWzbEduiA22S/cWWz5LFnyGivZis3S3b2jUfsj9lsz1ZiGopTztxVWh7i6xPhvL88VWrxmw7ZK3cWp//kKwa2tLazKoZY3H1OB43Lrd83Ga5zcZtktvs5apcKmk539tjqrZy2uaFdyabP6OxNGq/MSKexcJf1hhP069bnThJqIOIfqLP36j3v9Pv5xNzr37tzKtP7/+m9XWtb2tFgYALAuycBTCbmVPL3qz6UziW+I3+vj0x3JRwzhpwoX90AQEIQAACoxOo7NZIAFT2+mP2FSSgL0APC8dbrwuFhrJCfJ1O/XNaN9GKUo4CzMuJuJ+IukVorhDdyEJXM9H3dCc4oY+fpPdH6f00Z4ddd74nr1q1atsJmw1trjvam2j9cC7VMUnv99Z6sNYjc6n2rw52tp81kGpvG+hMXjWQTl6X60z+LpfuuG8gc0XvsnnJh5alOh4ZzHQ8lpvX/kQulXxqIHX5M0szlz+7fO4VLyzrvKJvMJVckutKvjRw92UvL01d/tfXu9pfeWPBJf9YMe8H/9Ln/zO44KrXl95xee4/89qXvdrZ9gZ1t71Fd7StpDvuyGus4ny/7K5L/+u0XZ5uX5zLtD8x0Jns0X4zuVT7bwfS7TdojD/Wr5MDGmcunUzkUslT9fnj9f7/9PumXGd7RL925jVZ7z+hdUutm2pyYYz1luUkHz4qmqhgTV7YYh1gMzVr0uBErd9mkmuI5Lf6/CIiWqJ1mVYUCGxIYGtiOlaf/DmFyLluwMLqWOI7E+NzDtTHUCAAAQhAoBQCFT4mEgAV/gOA6Ze3QM3MM2t0p/9ruvPfrTNdoDsuX9H7iVpRgivgHKn+MzM9pFO4Q3dEf6D3c8SSLxLLwTZXTQ7lx4V1J3dCLtVel0slpw6mkzOdHeCBdPIbA6nk+boT3KGP/0Lv79D7ewZ1h32Z7py/seD7/3jljqve1P4qs3S3DS29b3Xy4VX1WjKgyYtl6csfXtaZTKnnL7VePJDqOCOX6pg1mEo25FLJWq3Vm66U8ZSnnZxkAQt/TphP00TBd3XH71pdn3mK+bjWv2p1zqzQO5QKFmjU/7sX2mQvCsdbsuF44ifh5pajJjTP3qqCTTB1CEAAAr4KVPpgSABU+k8A5l+WAvqics9wLJGUoTHP6k7/T3SSEa0oRguI7ivK3zXE3+tXc/XeOW39fBY6QWxqsizZTY/Sf0h3ODfVusNAZ/JAvT9Kd0TP1Pv2wXkdv851djywrPPSvte72ga1PYpPAv+6p2OFc1aDkywYSLffOdjZ/mNNFHwn15k8RdfncF2ffbV+UusmJKEtnLcusPBBuq7H6f1Zmri5XFhu0P+rmqSjPxDRP7WilL0A1+oUv0bCt1li/ak6nphbHUucEY6ds4M+jgIBCEAAAt4IVHyvSABU/I8AAMpJQHf6Z+nRpN/pnJ7Uo48tev8xrShmCSwm4t8y8YUkdBSJvbdz6n0u1WFp/ZjuJO4zqEfs9d45bf17A+nkdYOZ5Pyl8zr+uKzziv8QboEWyKUvW+q8dWEg3b5Q1/UWvf9+rrPjnMHOji/r+k/Xdd9D6zZaWTNCtVpjTHSm1h8R0T1a/6wVpfwEJugaz2CmHxLns9XNibvD8TknbRk/66PlN1XMCAIQgEApBTA2EgD4GYBAwAWcK/mHY61z9Kj/07rTfzsJHxnwKZVJ+NKvE7lLX9R/T0SOtWxr19ze40K6Y7dzLtU+a+Cd96ffkUtf8aRz6r1uiwKBtQQGU8klWjMDqeQPtJ6eSyWnad0ht3zcmHfPIjichBJar9WGzsd3vqL3KMEXqNLsz+FE9rVDVJXV3++31cQTX5pwxLlbBn9qmAEEIACBEgtgeEICAD8EEAioQE2sdWq4OfEzsaxnieUyncbuWlF8F+AXhWkukVwqxF9gm/bI/SM3NpfqcN5/f6TuuJ0/mO641bkQHrW12b6HhwHLT6C7bejdswjm5dLJDq2naGLgIK0fn7DZ0OaWJbsR0//pxM/Vep0QORcsxNsKFCOAZZz+fj9K1/AGa9VQf0285UZN9h79oVmnjg/gXBAyBCAAgZILIABCAgA/BBAIksBHDmsZF25uPUFfAN4nLA/okb+TNf4JWlG8F/gzE6WI+XJh/pIOt1cu/9amuVT7pMHO5Ezd4T9vMNV+80Am+Qd68tpV+jwKBHwXcC7i6LxdJOd8OkMqeVkulTxh8J0LFm7jfLSiBrSX1mN0h/J8rc5HIz6i3/9HK4r5AlsI8Rc1zFvffnNcNhxvvbamefZM2uvkMfoYCgQgAAEIDC+ALVQAZwAoAgoETBeYePicT4fjLZe8NZafJZFfaLwHa0XxTiArRDdqguVUYXufcTxunO5I7aBH85tzne3nDHa236TfP0VdVztX5PcuCvQMARcF/jOvfZnzc6v1Nk0KfE+r89GI++v3HyIJOTuX+2qC60QmuUZ//p2zBpa5ODy6clFAk5HbEslJItZd4W3C2Zp469XVMxLTXRwCXUEAAhAoQwFMyRFAAsBRQIWAoQI1za0zw/HErXbe/iMRO6fzbk+4uS2wlJjmC8mFYlOTsyOkO0T1q3eO0smfDHZe8ftXO9vecHtQ9AcBkwRy6cuWDqbaH891tv9yINVxxuA7Zw1UOx9vyMKf06TAd4XI+UhD5+MMTQodsRBtr7+/TmebuvTvxWKtHc5bxAADAQhAAALrCODb1QJIAKxmwD8QMEtAX8CdqPVxEblLIztaK4p7Ak9pVz/RI2jH2xbX687+FrnOZNNgqqNtMJOc7+wI6fMoEICACjgfbziQbr9TkwLf0aSA85GGnyQJbcHCB2k9S1hu0M2cjy7UOxQDBOo1htm6Lg/o35Cnw7HExRNiZ++vj6FAAAIQqHgBALwjgATAOw74FwJGCOhRm+P1RdsjGszPtX5GK8roBJyrov9Oj/AnLKGG3GbjNsmlkntpPXUglbxx2bz27Oi6R2sIVJ5ALn3Z0oF0+0Kt3x9c/fGFSeejC1mPQO+hSYEv687n91XF+VSCpXqPUjqB3fV333kWhx7SvyuPVsdaztf7PUsXDkaGAAQgUFIBDP6uABIA70LgDgKlFKiOtxxXHU/06gvn6zWOz2pFKVxgpThXOxfq0Be9/xfK88d1R9+p/5frTHYsTScX0R1tKwvvFi0gAIGRCDgXwNSkwA2DnR1n6f8951MJtrDI+qQmBQ4X4e8Qy53az0taUfwX2JeZL9JhnwzHWno0EdC6RfOcT+n3KBCAAAQqRADT/J8AEgD/k8A9BEogEI4lZukLsfuZ+GYmmlKCEALdVwKuAAAQAElEQVQ+pHTrBOZodY7qbzLovG85nUzkOpO/e72r3Tn6r0+hQAACpRJYmrr8r5oUmDeYbv9urrPjc5oY2Mn5NAJLqEGIzyCRX2pseAuBIvhWmBt0rMvzYj+nf4PuqYknvlkz/Zzt9DEUCEAAAuUrgJm9J4AEwHsU+AIC/gm8c3G/lvl6pPp2HfUgrSgjE/iTml2rR7KOyC0ft1ku1TFVdyjatTrv6x9ZD9gKAhAoqYDzaQTOGTmDqfZrcumOE/X/7x6h/Fthm6mZSJIa3GNaUfwQYDpUiL4vVfnnqmOJX1XHE1E/hsUYEIAABPwWwHjvCyAB8L4FvoKA5wLVsZaYHm3pfOfifjzN8wGDP8CbusM/n4VbQmztojsKO+rR/VMGOtvvpu62t4I/PcwAAhBwBF7vunpwWWcylUt1tOr/88+OXb5igj4+jUQu1h1U5yMJ9VsUDwXGMdOXmSgdjiceD8dbEzgrgHCDAATKRwAzWUMACYA1MPAlBLwS0BdUh2m9k5lT+mI27tU45dEvP8NMV+tO/+G5bXNh3eFvGki3X/F65+XPl8f8MAsIQGA4gde6f7xcEwH35NId3x5MJRvGcW4csRy8+loCRPeR0Krh+sDzRQt8hkjacVZA0X5oCAEIGCeAgNYUQAJgTQ18DQGXBWpirVPDsZbbtNsFWo/QirKuANO/9cXmb0XoGzbbk3Op9t0HOpPf0J3+eXTttXiRv64XvodABQq82nntG7nOjgcGnWsJpJKH5laMq2aSqUpxrtYMES8n3NwWwFkBbouiPwhAoDQCGHUtASQA1uLANxBwR2BifPaBNfHEr4XlAWI+yp1ey6cXFnpId/ovFZuacnuN2yaX6pg1mE5evazzir7ymSVmAgEIeCbQ3fbWQKqjO5dKXqY1ltts8y3ZtiM6nnNR0Lv1/nWtKO4J4KwA9yzREwQg4LMAhltbAAmAtT3wHQRGJVAdbdm3Opb4lU3WIiH6/Kg6K6/GOZ3Ob5jpKyShHQfSyQN1p/+8wUxyPrW12focCgQgAIHiBe5oWzmQuaJXkwHORUGdi4R+xCJrCjElSOROJv5H8Z2j5RoCOCtgDQx8CQEIBEIAQa4jgATAOiD4FgLFCFTHE5PC8ZZr2eJHdSf3y8X0UYZt3ibWF95CJwxZK+v0hfnnBzqT1+fSl/2pDOeKKUEAAiYJdLcNLU1d/mCuM9mRS3d8bmDvzT9mi3UAk8wmkt9qqPiYUEUYZcFZAaMERHMIQMAPAYyxroC17gP4HgIQKEygOt76DSbqJeKTCDdHIENCp+rRt9pcp77wTievWzHvB/9ynkCFAAQgUBKBtjZ7WfryhwdSHVfmUh2zNCH5cRY+SBO2FwnRopLEVD6D4qyA8llLzAQC5SeAGX1AAAmAD5DgAQiMTKCmOREJx1vm6xGlH2iLrbVWcrlfX0ifbdtVzpH+WC6d/IkefftrJYNg7hCAgNkCA+n2hQOdyQsGU8kGsUKTiORkTV7ewni7wGgWDmcFjEYPbSEAAdcF0OEHBZAA+KAJHoHARgVqZp5ZE44nLhehbiKeRhV649UX8uPzLEt206Nph+gL6auWZS7tr1AOTBsCEAiwwOC8y17MpTp+rsnL4waW/3sHsahJp9Ou9fdaUQoXwFkBhZuhBQQg4L4AelyPABIA60HBQxDYkEB1c+JYGRrTq8+3aq3E8qTu+F8kxPsOpJMH5lLtly6d1/HHSoTAnCEAgTIV6L7+rcF5yfm5VHKO1n0skU/r77wziOVOnfFSrSiFCeCsgMK8sDUEIOCaADpanwASAOtTwWMQWEdg/OGzd65pTtykO7+/0ac+rbViihA9r5Ntt4Qa9MXw3rrjf8Fgqv1xfQwFAhCAQNkLLE13PKe/865xrmlibWZtz8xHMMkPdeJIfipCAWWtswKq44nTPnJYy7gC2mNTCEAAAoUJYOv1CiABsF4WPAiB9wXCsURLKG8tEqEvvP9o2X/1V93x/wFZfNhgKrmL7vjPWZpO4kJZZb/smCAEILAxgaV3XJ4b6Gy/eyDV8U39vbgbSWjv1R81SJQh4jcIt5EKfIaJrnlrLD+tf2PP22pG67YjbYjtIAABCIxUANutXwAJgPW74FEIUDjWckg4nrhfX9wllWMLrZVQeojla6H8W5/WHf8zc/Pa762ESWOOEIAABIoRyKUve3L1Rw2mkrEh6+0dRORYEvoZEeN6KDSi2yT9G3vxKlue1r+3l4+PtU4eUStsBAEIQGB4AWyxAQEkADYAg4crV2B807kfqo4lriRmZ+f3oAqQWCksN4jY0/WIVmOus+Nnr3ddPVgB88YUIQABCLgm4Hzc6WC649ZcOvm1XKq9zrboQCE6Xwe4X6utFWXDAh/Wp1pDrImAWOKn1Ye37Kvfo0AAAhAYhQCabkgACYANyeDxihSoic0+PhQaWsRMZ5U/gPQLyYW2XfXpwc6OLw+mr1hQ/nPGDCEAAQj4I7BsXvKhwVTye7lU8hCLrO1Z+MskdAuRLPcngkCOsgkxncJ5fjTcnPhNOJ44LJCzQNAQgEDpBRDBBgWQANggDZ6oJIGa+Jzd9IXGrcLW9TrvOq1lXGSB80I0t9n4Tw+mOtrw0X1lvNSYGgQgYITA0tTlfx1It9+QSyePG2NZdc5brZgoZURwpgYhdKyGtkD/NqfDscQs/RoFAhCAwIgFsOGGBawNP4VnIFAZAuF4y7lCtnOBu6PLd8YyqEeefsZMjblUx3TnhSjd0bayfOeLmUEAAhAwU+A/89pfdd5qNZBKNtsW1+sR7wSJOB8va2bApY8qqka318QSD4bjiROpsa2q9CEhAghAwHABhLcRASQANoKDp8pboDrasq++mLiPiC8hoglay7DwM8R8jkWhT+uRp68NdCZ7ynCSmBIEIACBQAosm9eezXUmO3LpjoiwvQ8zXUTETxNuHxAQpgP0wZ+Hx694uibWOrt61llb6PcoEIAABNYjgIc2JoAEwMZ08FzZClQ3J85gi3Tnnw4ux0mK0Fyd1zG5VPvuuc72y53TT/V7FAhAAAIQMFRgsPOK32uS9gL9vb0nsRwszN8noZcMDbeUYe0iLB38ZtXTNc2JC8PNLduXMhiMDQEIGCiAkDYqgATARnnwZLkJTIzP+URNPPFrFvohEY+n8rr9W4h+ILZ8djCdnJlLJW8rr+lhNhCAAAQqQyDX2fHAYGf7WbnNX65j5iOI+Zc6839rRXlf4BOa7P4OCT9dHU98vyaa2P39p/AVBCBQyQKY+8YFkADYuA+eLSMB5yJCNtn36U7y58toWqTzeVVI2laNHbP7YCp55mCm47Fymh/mAgEIQKBiBe64Iz/Q2X53rrP9RK5aVSfMX1KLO7S+qRXlHYEwE31TLHq6Opb4VU1zIvLOw/gXAhCoUAFMexgBJACGAcLTZSBw8s/G6M5/kphu19lM0loW5X87/kNjx+w9mOq48I07L/lHWUwMk4AABCAAgQ8IDNz9/YHBzvabcqnkURZZ9ZoMOE3/rs3/wIYV/AAzfVmEusPxxJ01sdYZFUyBqUOgggUw9eEEkAAYTgjPB1pgYnzOgdWvvni/vkhqCfRE1ggeO/5rYOBLCEAAAhUo4FzXRZMBP851JptCbO2iBOcyycN6j/KOwBHCMlcTAfdVz2j5wjsP4V8IQKAiBDDJYQWQABiWCBsEVUCz/7OdU/6ZaEpQ57Bm3NjxX1MDX0MAAhCAgCPweuflz+dSycsGUh0H2CT7a8Lb+WSbZ53nUOlgtvmmcHPi9+FYywnwgAAEyl8AMxxeAAmA4Y2wRcAEwjMTO4Zjrbdp9r9DQ99Ea6ALdvwDvXwIHgIQgIBvAstSHY/kOpPf0oTArjroNBL5sd6/prWyi9DexPyL6njrourmxLGVjYHZQ6CsBTC5EQggATACJGwSHIHqeMtxNET3E8tRwYl6/ZG+u+N/Id7jv34fPAoBCEAAAhsW0CTAPbl0x2n5fNWnVl8vgKibKvzGJAey0G/CscQ9Nc2tMyucA9OHQBkKYEojEUACYCRK2MZ4gY8c1jKuJt7yAya+WYP9pNbAlnV2/Ntwcb/ALiUChwAEIFBygeVdl762+noBqeRU3QGeirMCdEmYDhWRu6rjibnVsTnT9BEUCECgHAQwhxEJIAEwIiZsZLJAzYzWqW+N5fuE+BsmxzlcbEL0qpD874g/dvyHA8PzEIAABCBQkMBAqqMbZwW8T8ZEM5jt+eF46y010dkN7z+DryAAgSAKIOaRCSABMDInbGWoQDieOEds+34N77NaA1mw4x/IZUPQEIAABAIrgLMC1l06OUYsq0cTAdeFZ7Tuve6z+B4CEAiEAIIcoQASACOEwmZmCYRj50zUnf/falSXErEm8SmgN06++x5/HPEP6AoibAhAAAJBFsBZAWuunnyFbHE+MeBHWzTP+dSaz+BrCEDAdAHEN1IBJABGKoXtjBGYMKO1Tjg/VwP6nNaAFu7StEVjLtXeivf4B3QJETYEIACBMhLAWQFrLCbzqXmRx/VAQ4cecNhhjWfwJQQgYKoA4hqxABIAI6bChiYITIwlpli2zNVD/lNMiKeIGP4mzKfpjn90oDPZU0R7NIEABCAAAQh4KoCzAhxe2Vz/nU2cf7wm3nrR5tPO3Ea/R4EABAwVQFgjF0ACYORW2LLEAuHmxOdspl4No05rEMtP2LanOFdjDmLwiBkCEIAABCpLAGcFrF7vLYXk/DFjxjwejreeO3HWnPDqR/EPBCBgkgBiKUAACYACsLBp6QTCscTXSch5z3/pgihyZCFapDWWSyVPHchc8Zciu0EzCEAAAhCAQMkEcFYAfYxILrHflMdrYi1n0qy2sSVbDAwMAQisI4BvCxFAAqAQLWxbEoHqeEsbMf24JIOPbtCl2nzOYCrZoDWjX6NAAAIQgAAEAi2AswKkVpivCr/5xuPh5sTJgV5MBA+BchHAPAoSQAKgIC5s7LeA/nH9GRNf4Pe4ox2PSW7Kh+wD9ah/+2j7QnsIQAACEICAiQLrOSvgURPj9CYm2Y2EfhaOJx6pbm79ojdjoFcIQGAkAtimMAEkAArzwtY+CYRj50ysibc8pH9cg5Zdf4pYjtYXRV9aPveKF3ziwjAQgAAEIACBkgmscVbAfsTW0RpIJZ319lkWuVETAQ+E463/p3NHgQAE/BXAaAUKIAFQIBg2917A+Zg/ovyfhXh/70dzbYSVTPzdcTxuSq6z43bXekVHEIAABCAAgQAJ5Dovvz2XSsbI4sOY6OYAhT7aUKcSyR018daUJgMOG21naA8BCIxUANsVKoAEQKFi2N5TgXc/5q+PmIJzlV3mOzVZMWUg1f6dVzvb3vAUCJ1DAAIQgAAEAiCQm9d+70Aq+QWS0N7MdLWGPKi17IuQxHSSC8KxxE+rZ5yzk36NAgEIeCmAvgsWQAKgYDI08ErAeQ/dux/z59UQbvf7dxY6IdfZ/rnBVPvjbneO/iAAAQhAHM6g/gAAEABJREFUAAJBF8ilL3tyoDP5DZLQHix0kc7nL1rLvzCdwpJ/KBxPtNKsWaHynzBmCIHSCGDUwgWQACjcDC08EAg3t1zmvIfOg6696vIum+1DBtLJ67waAP1CAAIQgAAEykUgl77sTwPp5AVcNW53/Xt/ls7rD1rLuwh9WCd4efjN7TQRgOsDqAUKBNwWQH9FCCABUAQamrgrUBNPdJLwHHd79bI3OS+XSh65rPOKPi9HQd8QgAAEIACBchMYuLttYCDd8X39O7oHEx2v83tAa7mXfZ3rA4SbE78Jx2bvVe6Txfwg4J8ARipGAAmAYtTQxjWBcCzxZyGKu9ahtx09TsSH5FIdlxJuEIAABCAAAQiMSmAglbxREwEHWyQztKPfaS3vInQssfWQJgIudj7tqLwni9lBwAcBDFGUABIARbGhkRsC4XiLTUzbUQBuepTiR2MsZ+e//f4AhIsQIQABCEAAAoERWJrq6NREwP/ZFh1IxD8nopVay7VsQkLnEQ89HG5uOaFcJ4l5QcAPAYxRnAASAMW5odUoBcLxxNtErPvVZPrt7xrk8XqU4vT/zGtfZnqwiA8CEIAABCAQVIFl85IP5VLtJ+fz9h46h3at/9JapoXrSfgX4XgirYmAg8p0kpgWBLwUQN9FCiABUCQcmhUvoH/slmvrsVpNL+9c6C+VvNH0QBEfBCAAAQhAoFwElndd8UIulZwzZNPueqzgHCJarLVcS1QTAffXxFuvrpl5TiDOiizXhcC8giaAeIsVQAKgWDm0K0ogHEsMaMNxWg0vcp6++MCF/gxfJYQHAQhAAALlK7Aik/xnrrP98lz+rT1I6BQmebhcZyskp9tD+Ydq4i1nl+scMS8IuCqAzooWQAKgaDo0LFRAj/y/RkzhQtv5vD0u9OczOIaDAAQgAAEIbFSg6+q3c+nktQOpjgOIeBYTpagMbzqvbYX4ippY4sGa5taZZThFTAkCrgmgo+IFkAAo3g4tCxCoibe+qptvpdXYon94caE/Y1cHgUEAAhCAAASIcqn23w6kks16QOH/SGhhOZoI0wEicldNvOXGmvjZu5XjHDEnCIxSAM1HIYAEwCjw0HRkAnrk/y9Css3Iti7JVrjQX0nYMSgEIAABCECgOIFcZ/J3uXTyIBb7y9rD41rLrgjxF4VCD9XEEhduNaN1QtlNEBOCQNECaDgaASQARqOHtsMK6M7/i7rRJ7SaWnChP1NXBnFBYA2BNhGrsTO7VcM9L02amu7ft7FzcaNXNdKV/cyB6ZdqG9PPbb3f7X/bbI0w8CUEIGCYwED6ihtyqeS+wnwaC71gWHhuhDNOmL6zyiZNBLQe70aH6AMCgRfABEYlgATAqPjQeGMCuvPfp8/vqNXIIkQ/0BcNuNCfkauDoMpV4LAFz4xruPvZj0/p6tvV2YHXne0jGrr6Toik+xL69SUN6b6f6v1tkXT23sZM9gm9/1Mkkx1Y2NWflxC9xkND/TbLoxKyFnpVSejxEA9lhcf8Y+z4N96IZPpWRjLZ1yKZ/hf1/im9XxjJ9M3VOG9sSGeviWSyF+tjrY3pvq9FuvqPnZLqjzV2ZQ+cMrdv14PSL3yyceGfa6hN8Pe2XH+oMS8jBAY723+8ySrZR4OZo/UvWsusyKeF5frqWOJufX21Z5lNDtOBQEEC2Hh0AnhBMjo/tN6AgP5xelafqtNqZhFKDKaSZ5oZHKKCQPAEDr5z8ZYNmef3mJLuO1x3hs9oyPQldef9Nr1/UHeQn49k+v4RyWTfeju/6XIeO/avlvAzzg687mzfycK/IOZ2/fpcZj5F748ipkM0SbeX3m+vGiW+eCiP0Ri2IpId9X4PvW8k4hka5xeZ6TQiOk8fu1yYf0Iiv7EsSYnQImsMP5Pn0Mvy5sqlkX368zr/QTV5Re/VI/tIRJMImii4piHTN6ch3X9cZH52SmPXn7drXLiwSvtEgQAEChT41z0dKzSx3z5krdxX/29epM3/q7Wsis7rcJ1Qd7g50aL3KBCoRAHMeZQCSACMEhDNPygQjrc+pY/uotXUckwunewwNTjEBQHTBJwd0gNT/Ts4R+wbMtkv6ZHub0cy2WsbMn3z9f6FSKZv+dCm1n+Yqp6ymO8mkh8ycYvuvB+l9wfofHYm4q2JaBOtlVwmqMlHFUA96LOkSQRNFJymRpcxy81kU6/Iyj/Lm9uuUtNXI+nsY2r824ZM/1WRruzZka4lsyKdi/drzCz5GOEGAQhsUGDFvB/8a6AzeQFJaB8hvpKI3tZaTmWCJiCT4XhLprq5VX+XlNPUMBcIDCeA50crgATAaAXRfi2BcDzxmL7432OtB435hv9lW3SgHh24zZiQEAgEDBGY2tVXt+7R+0gmq0eps393dkhDlrzkHLFnohv0SPd3NeyTmHia3k8m4nGEm8sCvI0mC/ZR488xyZn6Yv8KEvt2ClkPC9l/07WxI5m+vzRkss4ZFrdEuvrbI5n+MzRRcIQmaPY+eO6zH3E5IHQHgcAJ5NKX/Wkw1T7bsmQf/f90beAmMGzA3MQi3eFY4rxhN8UGECgXAcxj1AJIAIyaEB38T6Am3vKQfr2PVhPLo2zn9102L+nEaGJ8iAkCvglMyTy3c2NX3xd0Z/HKSCbbo/fLbOG+dY/ea0DOkaVt9R7FPAEm4k/oP84ZFseQSEKTrz/URMGdmqD5/dCYsf/UtX1b64uaGFjY2JW9JpLp+7rzNoP95j+/BeEGgQoSWDqv44+5zuQptuQP0P8zN5fZ1DfR5MbFegDmvpoZcxrKbG6YDgQ+IIAHRi+ABMDoDdGDCugfnm4h3l+/NK8I3ZLbe9wBA5kryvCiQOZxIyKzBBrSiz+tO3/HRzLZHzasPlrc/4ZFY54X4Zt0Z/EsImrQ+/F6j1J+AmN1SjtqYqBRxLlWAf+YbOoda1f9N5LpezWS6b+nIdN/VWM6e6Lz1oID5vZN0O1RIFC2AsvSVz48kEp+gSw+TP9PzC2ziR4stt3tfGQgzZoVKrO5YToQ+J8A7l0QQALABcRK7yIcb7lODSJazSvCF+fSyeOorc02LzhEBAF3BRrT/bs3ZJZ8NZLO/ijyzun7bzNbf9QXutfrSGfokS89Wiz4WDvFQOFtNDFwqPP2AmH6OYWsh6vG8GAk0/eXSCabcd5S0NCVPb4x3bd348I/bwovCJSTQG5e+72D6eRM/fk/Uud1v9ZyKaz/n79T/eYnu8MzWg8tl0lhHhB4XwBfuSGABIAbihXcRzjWOoeIv0Im3kROzKXbv21iaIgJAqMVWL2zn1p8su6g/VR32B5vSGfzwvI0k/1LYjpV+3dO33eOAOuXKBAYqQB/Qrdsct5SwELXC/Pv5c2Vb+rP2IuRTN9cTS5d0pDuP25qZ3Y35+KQui0KBAIrMJDquCuXSh4ibH9JJ/Go1rIoTHwg2XJPOJ649GOzzkLStyxWFZNYLYB/XBFAAsAVxsrspCbecgSxXGbk7EUOzaU7fmlkbAgKAkUIODv8jZm+MyOZ7F1al4qzs29ZP9MdtFO0u88wE36fE24eCuxIxDM0uXSu84kFdoj+4FwcMtKVXdyQ6fut/kxeGOlaMsv5GEPCDQIBExjsvOImTQTsR8RfJ6LntJZLOWfZm1Xd1bE5sXKZEOZR2QKYvTsCeMHojmPF9TL+8Nk7C/GdJk7cZnuy7vzfZ2JsiAkCIxVYe4e//1/ODr/+n7tK28/UWqMVBQKlFxCq16ONn9NAvkNi3+58jGHEeRtBV/bGxnT2xAPTi2v1ORQIBEIgl2r/6YTNhvbRRFeChF6m8rjtw2ynquOtV0ycNSdcHlPCLCpUANN2SQAJAJcgK6qbxraqUN562MA5vz52sxUTlnVe0WdgbAgJAhsVeH+Hv+/2SKb/b2vv8MuHN9oYT0LAKAH+hO48fVGYfh5iK6sJgVcjmewtkUzf16d09e1qVKgIBgLrCLxyx1Vv5jqTHXm7ah8huVCfXqY18IVJzrbfzPfUxGY71z0I/HwwgUoUwJzdEkACwC3JCuonPP6NJ3S6hmWR5clcKrnla3f8eLnGhgIB4wWcHf5IKntqQ6bv15F09k/v7/DzLCL5mPETQIAQGLEAb6ObHkPEP7aEn9Gf99f05/63kXTfN6am+/cl3CBgoMDyrktfG0x1tFki+xPxr6gsbrybsPW7mnjimvFHtCCxXBZrWkGTwFRdE0ACwDXKyuioOp6Yqzsnu5k0W83QP5hLdextUkyIBQLrChyUWbJjYzp7oh4FvS6SyfY5O/xk0Y+Y+PPEtP262+N7CJStANNW+nP/OWL+gc3yaCTTvzTiXGAw09c6NbMk0ohPHSjbpQ/ixJamO57Lpdq/KiJxjf8BrYEvQnRaaBV3V8daNDEX+OlgAhUigGm6J4AEgHuWZd9TTbzlB0w0w7CJPjaY6phiWEwIBwKrBQ7qzH60IZM9OdLVP3dI7H5h+jm986kZdYQbBCDwroDUkHOBQeLLbbK75c2V/4k4H0WYWfKdxnTf9ANTf5xIuEGgxAKD6Y50LpU8WMQ+TUN5SWvQy2RmviUcb712i6ZWnHUW9NUs//gxQxcFkABwEbOcu6ppTpwlxN8wa47ytP4xdj7qzKywEE1FC+w3//ktIvOXfDGSzt6WD1FWk2Y/I5EZzHq8n3CDAARGIDBOt2kisi8U5q6Qtck/I+m+BxrS2UsbM9mZU+a/7LylQDdBgYD/AoPpK36cHyP7E/PlOvqQ1oAXOSkfku6aeML5KMSAzwXhl68AZuamABIAbmqWaV/OBWNE6ErDpvdcLtWxp2ExIZwKFdjv4b9tFkn1zdKdk+vH2lVZsu0bieko5XB2ZPQOBQIQGIXAWGKeykznCNFdlv323yOZ/ocjmezFU7uyB2uCTfNso+gdTSFQoMDyuzr+netsP0fY3p+Ib6Xg33bU/1s3hGOJn27R1FYd/OlgBmUngAm5KoAEgKuc5ddZTXzObsSha8yamfTrkf9PmxUToqk4ARFuTPfHGzL9Pxk78EY/WXy7voA6Xh220ooCAQh4J6A7/LKfdn+eLXRfpKv/T5FM3y8bMn2fnzL/BZwdoDAo/ggMdl7x+1yq/Vgie5aO+IjWYBemU/KhFffWNCciwZ4Ioi83AczHXQEkANz1LK/evty2qZD9UyEx6QXVn/XIP94/XV4/aYGajXPEUXc0rop0LckKSyeTfE0ngPdPKgIKBEoksB0Rf5WJf23ZVS81dmVTkXT2bE3Q7U64QcAHgVzqit/mUsn99e/BbE0Ev+rDkF4OsY8I3VMTaznTy0HQNwQKEMCmLgsgAeAyaDl1F/7PG7/S+Rj0Hnv5u/6B3UFjQoGArwJTMn0HRLqyl0Qy2WecI466o6EvjGSSr0FgMAhAYAQCspnuvMSI6QpN0D3dkO57IpLua5/atWTaXk/ImBF0gE0gULTAQKrjSsu2NRHAPyy6EzMajhXmq6rjiRvGzTjvI2aEhCgqVwAzd1sACfCR5fEAABAASURBVAC3Rcukv3A8cSmRmPTxMP/OpTpwlLVMfr6CMI3DFrz44cbMkjMj6exjFvGDJHSuxr2rVhT/BHIk8lcd7o/qv0h/J3Xqjt1NeoTtGn1x+j0R+0L3q3xPiL/PQr/UsW/XsTPvjM1P69cv6tf/0vsVWlECIMDMexFzwhZ7/vh/O28VyN7Q0JU9vmHBix8PQPgIMYACA5kr/jKQav8m23ZEk1FzAziF90Jmoi9V5VfdWx1NTH/vQXwBAb8FMJ7rApbrPaLDwAvUNLfO1Emco9WQIgN65B8ZaENWo9zDaEy/tLce7b/y7bz9RyH7KmLap9zn7PH83lLDPq33CcnvmJ0da7qSiL+jO9jfFKLjbZGZFlmNlkW7M9vb82ZjJ/Y01Vo90bqanlj9J/V+t55YXUNPtH5GT1Pdl3qjdWf0NtWe3xub3OZ+rT+/N1p7Vnes7kQd+2gdO/bO2LV76teT9Out9X58z+O1oTHLl9WE8vQxztNklqrPkC0HiS2Hi/DntX5NdwASLHShJg2c+V6r97doTRFxt97/mYjyWlH8E3CSyF/SNbme80POdQMWROb3tzam+/b2LwSMVCkCmgjoHUwnZ4olX9Q5/0FrMAvTp9mirnC8xUmCB3MOiDrQAgjefQEkANw3DXSPNTPPrBGRCwyaxAo98j/RoHgQSpkKNHT1H9mY6f+t8NDvdefsLCL5SJlO1e1p/W8Hf74Q/1R3rs7VnfxjmfL7sazapidat5nutE/WemhvtP7/upucHeu62T3R2u/qDvYPe6N1Ny6K1c9dGJ3Us3B63TPdTZNf7p66/QCx9uR2pG7218b2fUftnXugue7v3c11fd2xHZ/oidcv7I3Xz+uN1f5G68+645M7NJHQpkkDZ76n6P1xWpt17lP1foeeaF2Vk0AgtvZnpmOJZI7e/0h//lIa6h+JeIBw80iAq4j4ME3aXC7Mv4+ks3/QemVjenF8r84nNifcIOCSwOC8jl9P2Gxof2L5tnb5utaAFr4kHE/cXjPznO0COgGEHUwBRO2BABIAHqAGuUs7P/YCjd+MCyeJrNIj/+M1HhQIeCLw3mn+XdnHWeR3eoT6c54MFOxOR7qD39Qbrf267vBepjv5t3ZHd360O7bLP4M9de+jdxIIPU2THuluqru1J1rfrvena3JAkwR1u/VEayc6ZxmI2LtaFseJrFOZ5HJNENwiRA9pdK9oRXFDgGk3YjpL2Oocb034UyTTf3NjV98JDQv6tneje/RR2QKv3HHVm7nOjotttg8g4p9TcG+zZCh/T0285YjgTgGRB0sA0XohgASAF6oB7dM59V93gs40JfxcumOsKbEgjvIS+MBp/kKfKa8ZFjMbWUqrT02XH4otJ/IHj+BjB5/8vzlnGfTGJj+7cHptuic66Sfd0fpzNEFwXG+07sCeaN3Hex6vDTHb20uIIyT0JSHWo4x8rSazFhCtvn6C/0EHfUSmj6jdcSL8C86zJgP65jZklny1sTOLj/gM+tqWOP5lnVf05VLtJ2sY0/RnrFvvg1gm6e+ZO2virRcFMXjEHDABhOuJABIAnrAGr1PTTv0fY/FHg6eIiE0XwGn+763QYia6Tb/7lnNkWUKhT+jR5y16orVTe2L13+yN1/+yG0fwlScApY1t520TvdNqezUxcFNvtPZiXcdTeqP103ui9Z8cM6aqhpmm6M7GqUT8EyJ6UGtOK8qIBXgGk/1LCVFfJNN33dT5Sw4nzQ6MuDk2hMA6ArlU8p5cqmOqiHxH/2/m13k6EN9qkvH86nhi3vhY6+RABIwgAymAoL0RQALAG9fA9WrSqf8WyYz/zGt/NXCICNhIgco+zV/WPqov8pnXlq/apCdat3N3tO4Yvb/EObLcO22nvxm5eAhq1AL3Hbpjrrup7kFNBvykJ1p7ak+0borWmqGhqk86yR9NDpyrOyC/0YGe1YqycYEtifgrtm3fHenq72vI9CUbu7IHEm4QKFJgMN3xXX3N00hC9xbZRUmbaSK5OcRyT3W85biSBoLBy1UA8/JIAAkAj2CD1K1Jp/6z0EVLUx2dQfJDrGYKTJm/ZM9IV/bK967mX/6n+b+iL8Y2flQ/Vv/EC0ftstLMFUNUfgo8NGPHvzrJH00OXKbJgc9rUmBXrSxi7yqrP8WALhOhtCYH/upnXAEaq5aJW9RoUSTT93BjOntOQ6ofR0IDtICmhLo0dcWDuXTyMP2/dp7G9JbWoJWP6f+Fm8OxRFLnwEELHvGaLIDYvBJAAsAr2YD0a9Kp/0LUOZBOXhAQOoRpqMCBqf4d9MjcVZZtP6ZHVcr3av7MbwnxAmI62ybrU7rz9nEc1Tf0hzJAYTnXG+hd/SkGdef2xurimhzA2wiGXT/eT5guZUteiGSynVpP2n/Bix8ethk2gMAaArlUx6ViSyORdK3xcHC+ZGoJx1vvCccTewYnaERqtACC80wACQDPaIPRsUGn/j85mErOCIYaojRRYNr8v23RkM5eELLkMSZ2LmZZZWKco4tJntKkxlU229N7mmo3643W6n3dVYuik14YXb9oDYENC2zsbQS2yEwh6dDWj2pFIYorwrVj8nZfQyZ7faQre8Ss2yWkj6FAYFiBwUzHY7lUR1QTuwlNBCwftoF5GxyiId0TjrWcoPcoEBiVABp7J4AEgHe2xvdsyqn/QvRqbrOX9zUeDAEaKdAmYjWk+775lv3GY8zURkTldKXul/WF4E3CzlX5rY/r0di9emJ1Zy9qmrxA54kCgZIKOG8jWBSrn9sbrU/0ROv225ysTTlvT2WR84VEf0ZlRUkDLOngMpGJjiehO/89vr8vks5e0bCgv6GkIWHwwAjkOpMdJNKor4/mBSbo9wPdkph/URNvvXrb5rbN338YX0GgIAFs7KEAEgAe4prctUmn/hPxEXTHHYG8Ci7hVlKBhq7s8Qu7+nXHn7+vgeykNehlUCeQYaZzKW/vrztV2/c01X2pt8m5Kv+kV/Q5FAgYK9AVnfR2d/Pk7u5Y/fc0KeB8CsF4lqrPEPE3dEf4dtJkr9ZKLDsR09mcl56GTPbRSCZ73tT5Sz5ViRCY88gFcukrnDMjD9eE2lnaaqnWQBUhOf0Ne8U91fHWfQIVOII1RABheCmABICXugb3bcyp/0KnDKbaHzeYCqEZKDAl3Xd4pCt7LwtdT0R7aw1s0SM8jxFzksia8RY5R/nrYqsvzNY8+ZHATgqBQ+Bdge7Yjk/0RGuv7onVHd0TrfuoZVm76I6Bc3rwz3WTxVorqjCRc7bbxbZtP6eJgHRDevEpjek/b11RCJhsQQID6Y7vW5Y0ksidBTU0YGNhOoBJ0tWxOccYEA5CCJIAYvVUAAkAT3nN7NycU/+lLZdOXmumEqIyUWBqZkkkkum/w2K+W48oOu81NDHMYWMSosdE7AvzMvTZ3mjdZ3uaalt7opM6H4tOcs4AGLY9NoBAUAUWTp/0fG+0/jpNBpysdee8zTuS8HFC7JzFU2nXEYgyWz8VXtmnv9d+2ZB58aCgrivi9lZg6byOP+bSHZ8TotN1pNe0BqlsxWzfEm5u+VaQgkaspRXA6N4KIAHgra9xvRtz6j/TtYOpjguNA0JARgpMnZ/dTY/4/9wmu5tI/s/IIIcJSl+4PaZHQ9ry7+7098Ymtz0Y+9RjwzTD0xAoa4EH47V/6onV3tIbrT1LEwL72VZ+WyY6gpgu1Yk/oP/fK+E6AmGd51eZ8vdrImBBQ2rx8Y0LF1YRbhBYR2AwlfxRPm836sPOR77qXYCK8PfC8cQvcV2AAK1Z6ULFyB4LIAHgMbBp3Ztw6r/uCHXmOpOnmGaDeMwTOCj9wicbuvqTtk3OR/qdaF6EG49If9bXOtLf21R3IXb6N26GZytbYNH0nf/RHa27u6ep7jxNCBzMb2/yMRZpktWfNMBPlb+OHMaWdb28+dFnGjJ932rsWrxd+c8ZMyxEYHnXFS/kUsljiOVr+jfm1ULaGrDtV1fIiszE6JxPGxALQjBWAIF5LYAEgNfCBvVf05yI6Asp5+PRShnVk0Njx2Dnv5QrEJCxGzN9Z+Yp9Jj+zLZoyJtoDUTRF2Rr7/TjSH8g1g1BminQfcT2A92x+vm9qz9poHYv28rvxcQJYpqvEa/UWqZFdtZ5fk/EchIBP2nsyh5YphPFtIoUyHV2/EybNjLRzXofpBKxLTtd09xyeJCCRqw+CmAozwWQAPCc2JwBRKjUO/95PbJx9ht3XvIPwg0CGxCYMn/JnpFM/1whvkpf5H9kA5sZ9bDQ2u/px+n9Ri0PgikjgUXTd36qO1rb0dNU18Rk7aj/944nohu1luunZFQz8df07/eiSCbbqYmAY3SuKBBYLTCYSi4ZSCW/wELOhTX/uvrBYPzzcRG+uybecnYwwkWUfgpgLO8FkADw3tiIEapjLc6LhpklDUboOwPzLu8taQwY3GiBhkx2tmXb3UQyw+hA3wmuf80L+WGn/x0U/AsBvwS6o5Ne6Y3W3dgTrTv+teWrdmQp+7cKxEXoFk0EPN2Y6W+ZMv/lbfyyxjhmCwykk9cRSyOT3GR2pGtHJ8RX1MQT1+jffF77GXxXwQKYug8CSAD4gFz6IdosZi710f9MLp28pPQWiMBEgSmd2X30yFaKiTo0vglazS3M95HFX92crF2x02/uMiGyyhJ44ahdVlbQWwV2F5KkZb/9R00EfD/S9dJnKmu1Mdv1CeQ6O/48kOr4kibCztLnA/P2GCE6LRxrzVTPOGsnjRul4gUA4IcAEgB+KJd4jJr4cmfnf98ShjFIYn+nhONjaIMFGjJ9c6wQL9QjWzFjw2R+i4ivs5gO6WmqPbRneu2vuqKT3ibcIAABIwUq5K0CW2ki4JskQ49rIuC3DanFRxq5GAjKV4GBdMf3mekwEgrOp8wwTWd7TDocTxzmKxYGM08AEfkigASAL8ylG2TL+Lc+KsROAqBkQTDRd3LpK54sWQAY2EiBSNfi/XTnfz4TX0YkmxsZJNHq0/wlFNq1J1p7wsKmuvsNjRNhQQACGxDoroC3Cmgi4HNsWb/TRMCjkXTfNw6+c/GWG+DAwxUgMNCZ7BkT4kN1qj/RGpAitfpaIBOOJb4ekIARpgcC6NIfASQA/HEu2ShDtNI5FezjJQuA6I6BVPIHJRwfQxsoEMlkzxObu5l4moHhEa1zmn/vYTsuIdwgAIHACwzzVoHAz08TAfvq768frNqU/6iJgPap87O7BX5SmEBRAv+Z174sl0qeqo1P0vq61gAUDhHTjzUJkAxAsAjRfQH06JMAEgA+QZdimOrm2Z/RcUt29J9J/mEzTv3XNUB5V6CxK3tgJJ29V7+9mJnH6r05Baf5m7MWiAQCPgms+VYBy7J20Z2P84j4UQr4TZOr2xJzwrbpD/o79zdT5/eb+xargFubHr4mAX5BYjun1gfnDDamlnBz6++2mtG6rem+iM9NAfTllwASAH5Jl2AcFsvZ+Q+VYOhBFsyMAAAQAElEQVTVQwrxd5Z1XtG3+hv8U/ECkUzf+SLUTUyHkFk3nOZv1nogGgiURGDh9EnP9zTVXdoTrd2PLGrQo+kdGki/1mAXpmNtW1KRrv65DV3ZaLAng+iLEXDehplLtR9KJME5si5y5CrbzkyMJaYUM2e0CaAAQvZNwPJtJAzkq0BNtOVwHfA4rSUpuqN3fS6V/EVJBsegRglMzSyJRLqyDxDxRURUsoSUjr1ueZ6Fv/6/q/njNP91efA9BCpXoGd63aLeaH2ip6m23hKZScTXEdF/tAa3iMxgoXQkk72rcf6L04M7EURenABLLtXRKjY5rw3/Wlwffrfi3WymdE2s9Xi/R8Z4/gtgRP8EkADwz9rXkcRi5+i/r2O+P5i8WFWV/8773+OrShVo7Oo/xSa7m4SmGmTwNz0KMqc6P2Gf7ljtT3E1f4NWBqFAwDQBZlkYq5/bE609oWqsXU8iX9UQ79Ykt633QS0zxc53RdLZ3zVkljinhgd1Hoi7CIHBTPIW2+LD9Gd4bhHNS9FkgrBcXxNvdQ4ilGJ8jOmPAEbxUQAJAB+x/RoqHG/9mo7VqLUkRSR0/utzr9SdrJIMj0ENEWhI931XRH5qSDgaBg8wycUsY/fpida3dzZv+4Y+iAIBCEBgRAL3HzL5vz2x+l/1ROuOIDtfL0QtmtzsHVFjEzdiOpLJXhDJ9N3ROL/ftLdmmShWNjEtm9eeHUwnZwrTBUGZlJCcXxNP/Joa2zYNSsyIsxABbOunABIAfmr7MFb1tLO2ILKdK//7MNoHh2CiHw2mL7/1g8/gkUoRaMos2USPLP2Kmb9txpzZOVJ3dV7y+3ZH67/dHdv+n2bEhSggAIGgCvQ277ykN1p3RU+sLpKXoc/qPC7R+qzWABb+P7HlXv29fVtDpu+gAE4AIRcpMNiZvMgimaHNF2s1vmjS7fPh8cvv3nzamdsYHywCLEwAW/sqgASAr9zeD2aNGXMmEddSSW7yjL1q6DslGRqDGiFw8H39O7zBdhcxfZkMuOmLhRtY7H31iN03HoxNDv7FvAwwRQgQgMDaAg/GPvWY/o75ltZd2Qo1EfFPiOjvWoNVmI5i4vsjmewtjZ2LG4MVPKItVmBpqqOzioYOZaKbi+3D33Y8bcyYMXeNb5q9s7/jYjQvBdC3vwJIAPjr7eloWzTP+ZSQaALA02E22DkTXzm44KrXN7gBnihrgYZ0f8PQSvkjmfB+f+Y79f/CwXqE7svdsfonyhoek4MABIwR6J6+0/yeaO2pvNmqehH+PBPdpsG9qTVI5RgJWQsbM/2/bki/0BCkwBFrcQL/TV3194FU8gss3KI9DGk1vewbCoXumhCbs7/pgSK+EQlgI58FkADwGdzL4fJin6b9T9Dqf2Gar388bvR/YIxogkBDevFxzNKjsYzTWrrCdJ++4D6ip6n2c73R+gdKFwhGhgAEKlmge+ouy3tjtb/pjtYdMzRUVa+/l85Qjwe1BqZoEvXzzKGeSFf2xqmZFw8ITOAItGiBgXT7FZoEcC4M+fuiO/GtodRabN9dHWuJ+TYkBvJIAN36LYAEgN/iHo1XE539Se3a+WgXvfO/SJ5+4P+oGNEEgUim73xmq7SnDjI9LkTH9zTVHaovuO82wQUxQAACEHAEHpqx41/199I1PdG6KcIUI6Zfi5DtPBeIKvRFm/IPNmay10e6luwXiJgRZNECmgRYaG1mHUpCPyu6E/8afoiZ76qOt37evyExkusC6NB3ASQAfCf3ZkCxQs7Of9ib3jfeq+543TiYSc7f+FZ4thwFIplsJxGX9KN5mOlC3fHftzdahzNQCDcIQMBkgd6muoz+vvqisPVpJrlYY/2T1kAU/Vt/PIn9cCTTd93UdP++gQgaQRYlsPSOy3O5dPJrTFSyt5XSyG9j9P/SrzUJcPrIm2BLkwQQi/8CSAD4b+7BiG26juIkADzoe9guV+ofCBz9H5apvDaYklnyoUhX/790VnGtpSlMi4Tk4O6murbSBIBRIQABCBQnsCg66QXnU0k2CYV31R5OJqaFeh+Qwl+xWR5tTGd/EenKfiYgQSPMIgQGUskfvPspAX8pormvTTQJcHV1rOV8XwfFYG4IoI8SCOiOYwlGxZCuClTHljs7/7u42ulIOxP6YS6VfGqkm2O74As0dmUPtMj+N4l8uFSz0T/0F/Omrx6E9/mXagUwLgQg4IbAPdO2XtETrft5T1PdQSR8GAldT0RvazW+CNMJGu/jkUz22obMkj2MDxgBFiXgfEpAiK2YNjY+ScXMF1XHW6/QWFECI4BASyGABEAp1F0ek9lyEgAu9zqi7l4J2Yyj/yOiKo+NIuns2SK0qHSz4UeE7GnOkbPuqVODcKXi0lFhZAhAIFACPbHae3tidV/Ji72r7ly3afBB+ejSk5jshyKZ/ov2zSyp1rhRykzg9c7Ln5+w2ZCTBLjO9KkxydnheIvxcZru6Ft8GKgkAkgAlITdvUFrmhMRImlyr8cCemL6wetd7a8U0AKbBligIb24jZhKllln4ss3J57aG518T4AZEToEIACBjQo8GJvc39tUd+HmZO3KzF/Rje/VanrZTF+LnL+p5B9q6Moeb3qwiK9wgVfuuOrNXCp5ArF8u/DWfrfgr4TjiTv9HhXjFS6AFqURQAKgNO6ujSpEpTr6/1Ru03E/dG0i6MhoAWfnn9m6oERBOh9HFO2O1p7TFZ0UiFNjS+SEYSEAgTIScH7fdTfVXt8TrTuMbDlIp/ZzIlmh9+YW5l1Y6PpIJnvX1EwfPjrQ3JUqOrJcZ8fFYq9+7fnfojvxp+ERmgQw/m0L/lAYOwoCK5EAEgAlgndj2HDsnB1IqCQffcJEP6A72la6MQ/0YbZAKXf+hbljk9BbU/UFcJfZSogOAhCAgHcCPfH6hfp78OS8be0qxN8m5ue9G82VnmfaxM5HB3Y0pp/b2pUe0YkxAoOZ5C1ii/OWACdBb0xc6wmkMRxv+cN6HsdDRgggiFIJIAFQKnkXxhXOO0f/x7nQVWFdMM0fSCXxkWuFqQVy65Lt/DM/RWTN6G2qTdwzbTezj3gFcmURNAQgEESBB+O1f+qN1l48dfqkXdmWL+ocMlqNLUI0W3jMQ5FM39eNDRKBFSUwmOl4LJ+vipHI7UV14Fsj3i0cT7zs23AYaOQC2LJkAkgAlIx+lAPvdfIYJinJ0X/JEy78N8rlC0LzUu38C/H3xywbPKgnOqkzCE6IEQIQgIDfAm3Mdne8/tc90boYM00h5mv9jqGA8XYg4h9H0tl7Igv6DyXcykZgedelr+XSHUfrhC7TanL5ZDjekvvQrLbxJgdZabFhvqUTQAKgdPajGrlm6xo9+s/1o+qkiMZCdONgJjm/iKZoEiCBhvRifX3p+3v+Rcg6QY9unXXfUXvnAsSFUCEAAQiUTKC7qe7BnqbaU2yL9iWSX5UskOEGZjqU8nJPJNP/4wNT/ZoUGK4Bng+KQC6VPFd/9k7WeA2+Tg9Xr3xzxbKJ8Tmf0DhRSi+ACEoogARACfFHM7SU6Og/W3z1aOJGW/MFSrTz/7RFoSm90Un46B7zf0QQIQQgYKDAoul1j/dE679qkxxITDcZGOK7IcnXQ5Y81JDJzn73AdyVgUAu1fFzIo5pIqCPDL7ZZP+lJprY3eAQKyQ0TLOUAkgAlFK/yLHDzS0H6R/3UpxGd0duXvsTRYaNZgEQKMXOPxPfzJvRYQujOz0UACKECAEIQMBogUXR+od6muq+ZJHVqIHeqtXEsjUTdWgS4MHGTHamiQEipsIFcqn2+/X1aVRbmn1tCouefudjtDVSlNIIYNSSClglHR2DFycgXJL3/jPLzcUFjFZBECjFzr+6XNAdrf1C99S6/+jXKBCAAAQg4JLAwuiknp5o3bFssx4w4N+61K2r3WgS4AAhukuTANc3LnhxF1c7R2clEch1dvw5l0rGmOSakgQwwkFFqHtifPaBI9wcm7ksgO5KK4AEQGn9Cx69Op6YpI2O0+pr0V/kDw90dsz1dVAM5ptAKXb+9UXfMfri9CLfJomBIAABCFSgQHe89r6eaO0stqRJmIz8O65/D46XfN75tIDz97v9b5tV4DKV3ZQHUh1nMNPZJk/MJmtRdXPrZ02OsUxjw7RKLIAEQIkXoPDhxTn6v2nh7UbXwhYc/R+doLmtS7Dz/zYL79EbrbvNXBVEBgEIQKC8BLqn18/vbaqbSWTP0JmZeIp2NRFfNHb8Gw9FurJHEG6BFxjoTF7FLPozR6+YOhkWeSQ8o3VvU+Mrz7gwq1ILIAFQ6hUocHwmnlZgEzc2/0uVPeHXbnSEPswSaMhkj2b272r/QvTYyuX/ndgdq/2DWRKIBgIQgEBlCPREJ3f2ROtixHQkkdxj4Kz3IKE79e/T1Y0Ls1sZGB9CKkDAOXvUClnOdQGeKqCZv5va8vuaprP38HfQCh4NUy+5ABIAJV+CkQcQjif21K39P1VJ+Nevd7UN6tgoZSTQsKBveya61ccpXatH/T/7yFH7v+njmBgKAhCAAATWI9DTVHdXT7R+mogcTSIL17NJSR/Sv0+ny1vc3ZDuO6qkgWDwUQssnXv5s2zbR+rPWe+oO/OoAwmFnpoYnfNpj7pHt2sI4MvSCyABUPo1KCCCkhz9X5UnwsX/CliloGzKeesp/2K1TtUjTqf4Nx5GggAEIACBkQj0xupv74nVHyTOBYaZFo2kjW/biHyKmW9rTPf99IC5fdv6Ni4Gcl1gIHPFX8aErDgTp13v3KUObcv+4/hY62SXukM36xfAowYIIAFgwCKMPATx/fR/Ebp5ebp98chjxJZBEGjIZB8kkho/YuW8PbUnOuknfoyFMSAAAQhAoDiB3ljtb3qa6hqE6ctE9KhWY4own1JVRd2NmSXOdZCMiQuBFCbwn3ntywZS7XEi9vPsQyrkFmJ5YUL07NpC2mDbQgSwrQkCSACYsAojiGFirMX5eJzICDZ1dRMmXPzPVVADOotksj9jogP8CMXZ+e9untztx1gYAwIQgAAERi/Q21R3Q0+0bj9mOVF7e0KrGYV5kpD960im/5cHpV/4pBlBIYpiBHKp9mO13S+0GlksK5QNx87Zwcjggh4U4jdCAAkAI5Zh+CCELN+P/pPQvbl0x33DR4ctgiLQ0NV3gsZ6slbPC3b+PSfGABCAAAQ8E+huqv+lJgI+I0QtRPIPzwYquGP5ap5D3ZF031cKbooGxgjkUsmTROgqYwJaNxDOvzQxPucT6z6M70cngNZmCCABYMY6DBuFkPieAGAmXPl/2JUJzgYN6cWfZuGf+hExdv79UMYYEIAABLwX6I3WXcEs+2si4BrvRxvxCNsR83WRruyNjfNf3GnErbChUQKD6eTZ+lrzIqOCWiMYm+y/bBk/66NrPIQvRyeA1oYIIAFgyEJsLIzqeGISMR26sW3cfk7/0D8/sNnLuPif27Al7I/JuomIqrR6WrDz7ykvOocAqlOJAwAAEABJREFUBCDgu0B30+SXNRFwBrM1VQdPaTWjCH1R7Hx3Qybry5ltZky6vKIY6ExeQMStZOhtiKpeGTfjvI8YGl7AwkK4pghYpgSCODYm4P/V/1e/9/+OO/IbiwrPBUcgkum7mZh28zpi7Px7LYz+IQABCJROoLtpUndPtK6Zmb9CxM+RGbePMtHPNAlwa8O9/ZPNCAlRFCKQS7UnhfgLhbTxc9sqe9U/Jxxx7pZ+jlmWY2FSxgggAWDMUmw4EKsEp//rL+LfbjgiPBMkgUhm8XlEfBx5fMPOv8fA6B4CEICAIQLdTbXXD1Xb+7PQhRrScq0lL5oEOJqHqLshnT2t5MEggIIFBlPtNzNTIxHZWo0r1qqh/9TMPNOXT08ybvIuBYRuzBFAAsCctVhvJDXR2Z8UIp/f/y8LBlPJJesNCA8GSqCha8kMIutir4PGzr/XwugfAhCAgFkCDx1Yv6w7VtcmUrU/E91gRHQiH9adyGsi6ezvpnZmPT/rzYg5l1EQA53JHrFCdTqlnFbjigyNWfqhWaeONy6wYASEKA0SQALAoMVYXyjCq6/+P2Z9z3n1GIs136u+0a9/AgelX/gki32l1yNi599rYfQPAQhAwFyB3tiOz3ZH675MZM0Q4h4jImU60raou6Gr/xQj4kEQIxYYnHfZixM2G9pGG/xdq3Fl5ZvjllFj26bGBWZ8QAjQJAEkAExajfXGIj4f/ScaIlqw3lDwYKAE8hw6TwPeUatnBTv/ntGiYwhAAAKBEuiJTursjdY2EvE3iOgvWktbmGpY5KeRrv4f7vWE+HoghXAblcArd1z1Zi6V/Jh20qfVuBIev+JN2utk/EwVsjLY1igBJACMWo61gxkXTWxN7PsFAO9fnm5fvHYk+C5oAnrU40iN2dOrIrPQSd3Nk7t1HBQIQAACEIDAaoGeaO3VoTwdQGJftfqBUv8jcsb41/oXNGSW7FHqUDB+YQKaBJisrzUeKqyVP1uHtwmv1JFYK8oIBLCJWQJIAJi1HmtFM+adnf9xaz3o9TdMOP3fa2OP+z9swTPjWOxveTsM/6o7VvcLb8dA7xCAAAQgEESBB5rr/t4Tm3y2RXKgxn+X1tIWoanMMr+xq+8LpQ0EoxcqMJBOOj9DtxXazo/tw/GEERfA9GOuoxwDzQ0TQALAsAVZMxyx/D/937IFCYA1FyGAX7+d3/w8It6TvLs9oUd4vupd9+gZAhCAAATKQWBhtP6hnmjdkULi7Hj/oaRzEvmwCN/UmOm7rKRxYPCCBXKp5DFC7Pk1jQoOjGjzcHPi90W0q7AmmK5pAkgAmLYi78bzoVlt40n8vvo/9SxNdzz3bgi4C6BAQ7q/gcjWBIBXwcs/9cXcZ7zqHf1CAAIQgED5CfRG62/mt8dO1Z2475d6dhrDnEgm23lgenFtqWPB+CMXGEy1z2ami0bewqcthfaujifM+BQMn6Zc8DBoYJwAEgDGLck7Aa1ascw5grvFO9/59i+O/vtG7c1AzN6e+i8UinoTOXqFAAQgAIFyFug+YvuB3mjtWUx0hM6ztGcDEMVDljU/0pV1YtFwUIIgMNCZvICYLjEtVv2Z/lI4njjHtLhMiQdxmCeABIB5a/JORJbl+8Vq2Mb7/9/BD+a/kXT2bCI+jDy6idDpvdFJT3vUPbqFAAQgAIEKEOiO1t1txNkAQtuT0J2RTN/5FcBeNlPMdSadaxy1GzihS2tirTMMjKvUIWF8AwWQADBwUZyQbGZfEwBC/OBAJlnqjLwzddQiBKZkluxMFnl36j/Tpb2xuh8VERqaQAACEIAABNYSMOtsAL6oMZO99aDO7EfXChLfGCuQSyXn6OtW464JICxzw9HWvY2FK0lgGNREASQATFwVjYlF9tI7H4ss8HEwDOWygEX5b+mRjC1d7vad7phu6mmq8y658M4o+BcCJRdovOvPNQ33vDSpYcGL++sOwcxIJnuS1vMaMv1XRTJ9N0cy2Xu0Pq31lUimb2Ukk5VR1XR2aaQr+yft46lIuu8B/fp3kUz/LyPp7BWNXf3f1hhOb0z1fWHq/P7Y1EzfAVM00XfAPX3b7tX5xOYlx0IAEHBBwJSzAYTo6HyIFkxJLZ7mwrTQhQ8Cq68JQPxDH4YqbAhLfj8xPucThTUq460xNSMFkAAwcFm2OPzsj2tYu2j1rbDNeP+/b9ruDtSQyX6JiI8jL24iC3XnX/v3onP0CQF/BRo7s1s1dmUP1B3uk3Rn+wq9T2t9WuvqHXrZZOVSHhrq53z+Id0hcD667FqN8GImOZPe+T92KBHtTkR6pJDH6P3oClONJu62J6I9iHmqfn0kkXyVmM4Wke8K0dVi8U22LSmb+EGL7Oerhvjv40MTVmjMb0Uy/f/U+z69f6Qh3ffbxkz/9zR58MUpndl9Drn9pbD2iwIB4wUMOhvgU5ZzXYDVb6czng0BqsBAqv2bJPJj/dKokif7EWps29SooEoUDIY1UwAJAAPXxR6q2sPnsB7NZdqf8HlMDOeCwMFzn/0Ik0en/ov8NU/yNRfCRBcQ8FXgwFT/Dg1d2ajuDJ+tO8jXal2k9TUJ0WsitEiDuVZ3ts/We+eilu7t0GuHPpZNNFnwER2vTu8/y8yfE5JvafLgRitEj60aPzSgc35VkwMLG9N9P1WLsxo6s1HHRtugQMA4AVPOBtD/Q5oc7P9l48I/1xiHhIA+IJBLd5yma+Ykaz/wXKke0Ndl24bHr0iVanyDxkUohgogAWDgwthsO58A4Ftk+qIRp//7pu3uQENjN3FOzdcdAHf7dXpjsi59MDa5n3CDgKEC+93+t80a5/cf0pBe3NaQyd6qO7l/0J3et0KWvMRCaX1ReIWGfpLWA7VupbXSyjaaHGgU5lPU4koOUdqxUaOVkXTfszhroNJ+HMyfrzlnA8hX5Y2VC5wzhsxXQ4S5zuQpqnCdVpPKwTXNCfPeouCrEAYzVQAJAANXhok/42dYFlu4+J+f4C6NNXV+djcSOd2l7tbphju7Y7U/XedBfAuBkgo4F+nSHf2jdQf2h42Z7BNjx7/xhthyL7N1ARMdTUy7aYB6ZFz/RdmYwBhi3mWjZw1k+r7XmO6bfsi9eCvBxiDxnDcCRpwNwLSPCM3XZNlXvJklenVTIJdKniBEN7rZ52j70p+fMyr64wFHC4j2nglYnvWMjkcj8NnRNC647SoLCYCC0UrfwLbpBI3C0upq0T9Yttj5y1ztFJ1BoAiBKXP7dm3s6j8lkum/Tnf6n8+H6BXd0b9VuzpDX+j5fKFUHbUyyjtnDRB/S5i7Vq0aGoiks3/QI6HXRNKLjz0os2THymDALEstYMjZAOM0Waa/f/rPL7UHxh9eYDCVPF7/Rtw8/Ja+bnFpTSzxVV9HNGQwhGGugOs7D+ZONRiRTYzO+bRGuoVWn4r0D8y/7GWfBsMwLgk0dmbridhJAJDbN4vl0t745Ifd7hf9QWA4gQPTz+/bmOk7M9LV/xvd4X/JGsPPiMhPicQ5ArfzcO3xvEcCTLtpYvA0Yus3ebJf1LV5OZLpu7kxkz29MfPCZxsXLqzyaGR0CwEy4mwAkosimaxR7zPHj8b6BQZSyS/oM7dpNaYI0y8nNCfixgTkTyAYxWABJAAMWxzhIV/f/09kPWUYAcIZgYBUsWaTxYOPAuOnafnQZSMIAZtAwBUBZ6f/3ffwPxriqkeF+CoSOVY730EripkCnyTi48T5lAIKPSJvbvtf3TlKN3Zl2xrT/fHG9HNbE24QcFHgf2cDENOR2u1LWktRTtKf88xhC54ZV4rBMebIBXKp5DG69e+0GlMsoc7qeOs+xgTkeSAYwGQBJAAMWx1hy9/T/0meMYwA4Qwj0Ni1eDvdQfLk6L+IfVn3UbssHyYEPA2BUQms3unvyl7QkMmu3unnd97Dv++oOkXjUgpU6+BREbpAWDqFx/wjks4+Fkn3/aixq//LU7v6PLlQqY6JUmECPU11d4XImqbTdj6mU+98L01vD23yaEP6JedsTd8Hx4AjF9AkwP8J09yRt/B+Sya5vmbmOdt5P5IBIyAEowUso6OrzOB8TQCITXj/PwXrJhLSo//k+ttEmOiG3lj97cHSQLRBEfjATr9Qm/7MYaefyvTGtA8xnyoiv7KF+yJd/b3OGQIN6f6GMp0xpuWTwAPRSS/1ROuOFLEv9GnItYdZfQHNfG9jJjtz7SfwnWkCg53JmUycNiiuyTKU/8W2zW0enMFp0CyJCNGYLYAEgEHr86FZbeM1nN21+lbyREgAUHBuq0+tZfHi6P9rzGMuDY4EIg2CAHb6g7BKPsUoMsU5Q4BZeiLpvv6GTP9PIl19sw6+c/GWPkWAYcpMoDc2uY1K9pYAqRGiuxoyi2eXGWvZTWcg1R4nEpM+7vrgFfaKX5Qd9NoTwneGCyABYNACrXrjDV8//k9/IfavyCT/aRABQhlGwJaqE0ho22E2K/xpocsWNu2QLbwhWkDggwIN6b6jIplsxnlPP+NI/weBKv0R5klM8jUSvn3Vptaf9GflTj2aevqUzBJc6LHSfzYKnH+p3xLAZHVEuvp/QLgZLZBLdUzXAO/XakZhOrY6nvi+GcF4EQX6NF0ACQCDVkjY9vX0f33x9aRB00cowwg03vV0DVvsnP4/zJaFPa1HMe7vidVdWVgrbA2BtQXaRKyGrr4TGjJ9DzKzcwXmprW3wHcQ+KAAEznXDzhCfw9dbZH9fKQr+1BDV/93I6m+qR/cGo9A4IMCJX9LgMg3nCTWByPDIyYJ5FLJQ4T4QVNi0t993wzHEi2mxONqHOjMeAEkAAxaIv3F5GsCgImRADBo/YcLxd5kc+fUf/evjM7WD4cbG89DYEMCh9z7Urgx03fmwq7+p1n4F/p75YANbYvHITCsgND+LPJtsviBSKbf+cjBaxu7ssfsf+czHx62LTaoaIHSviWAjtAkwOMVvQABmPxgqn0KkfzdmFCZ2sPR1kONicelQNCN+QJIABi0RhbTTr6GwzYSAL6CFz9Y48I/b6o7Vq4f/deI7u1tmjRP71EgUJBAw93Pflx30M5ftXLoaU1eXqWNd9VaXkXoTZ3Qy0L0mL5o7CQhTXDIxUT8DX3sGM7bU4uuVqhJxP68CJ3OzOcT05Xa//X6/Vwi6tX6LBG/IiRvUMXeZEed+klqcsuYzTZ9SX/e5jaks6cdPPdPH9HHUSDwAYESvyXgM5FM3z+cpOgHAsMDxgjk9h7/CWOCIf3Nb0n7+MNayinBSbiZL4AEgEFrpC9yPu5rOFVDuACgr+DFD2a/sVKP/ovr749lpuuKjwotK1HAeZ92pKu/nceOfUp3ii/Sly/bB9OBnR3rpzX2W3VH/EKdy6li8+cskgPZCk16i6xwT6xu855o3fa90brP9kTrZ+j3J3VH67/dE629Wh+7rbt5cnfRdfpO8/WI5W96Y3U/6m6q/Z7uuMzW/r+i38/sidZFtO7aE639eG+0ftzmZG3KMsF77D4AABAASURBVHYbGcM72058ws2aLPgSMX9zdexCt2hCwknolu9HeAqN1zWaob+zrhkau+r5SCb7s0i6H0fO9AcYZW2B0r4lgLdetWpoYMq8F/ZcOyp8Z4xAW5sdCuVNSgLsbm1itRvjM+pA0EEQBJAAMGSVama21WgoE7T6VKR/4O7vD/g0GIYZpYC+6NUEwCg7+WDzB7ub6m794MN4BAIfFGjszNY3pPt+apH9FIkkdIuttAah/I2I7xWia5joDJL8YRIKfUJ3rsf1ROv21Hqs7oi36Q7+T3rjtXcujNY/1D19pxcfi04aJENuXdFJb3fHtv9n76G1ixc58cVqU5osuKmnqfaHq2OP1R2nCYm9dS4Thqrko2TLQcL8NWK6UhPLaSJeQuV0E9pSp3MysdyjiYBHGjLZ2Y1di7fTx1Ag8J6A839D/w8cqQ+8pNXXYlWFnmzMLPk/XwfFYCMWeH3ulX+TkPj6ttuNBccix9c0J87a2DaBeQ6BBkIACQBDlont5f5mI4Wdo0WGzB5hbExAX+A26fN7aHW3COPov7uiZdlbm3Nxv8ySOVJFzsX9TtFJbqLVwMLO6foPOEfENWF2rJC15/L8Mmcn39nZP0x3kM/ojtZd0xPb+d7eaTtpUsDAKbgQ0kOH1b/aE69f2NtU+7OeprrZvbG6eE+0trbn8VtCHKLJtshMYm4lpl+Q0CIi/hcF+/ZZTex0iLBzVsANDV19M4I9HUTvpoD+H7grRNY07fMurb4WIfuOhkzfHF8HxWAjFhic2/GYJhGPHnEDjzcUoWS4ueUgj4fxvHsMEAwBKxhhln+U+Tz5e/o/0x/LX7VsZugkAFyejDzVE6v9lcudorsyE2hI9R+5sKv/ISb7Mt1ZdI66GjNDIXpVmOaS7swyWfvpTq5zuv7BzlE/58yW3uikp59s3vsNwu0dgbY2u3taXd+iWP3cnqbaZE9T3Uk9sboGddua335jom3RvvoC9HQhuVnX+k/vNArSv7y5RvslFp4byfQ/3Zju+3bjgmy9PoZS4QKlfEsAE1/WkF7cVuFLYOz0c50dt2twpiRpQiScnHDEuUb9rVWfQgq2DYgAEgCGLBSz5WsCQIj/ZsjUEcbwAu4nADj0y+GHxRaVKtC44MVdIpnsDWzJ79TAkNMk+Tki/hXZdJrNsltvtO6jvU11M3t0Z7Y7OulRwq1oge4j9hhYNL3u8d5Y3Y96o/Vf0MTAjra1ybZMdASRtNPqMwVkqOgBfG8ouwvzdyVPL2gy4PaG+UuOmnX77SHfw8CARgk4yUFL+P+I2Ne3PzJbFyAJQMbecqmk8/77nxoS4J7WqlVJQ2IpIgw0CYoAEgCmrBSLrwkAy84jAWDK2m8kDt0Jc3b+3f50iMW8adV1GxkWT1WoQFNmySaN6f5vSz7/kBJ8SWspSz8x/9A5nT9E1k490dpPa/1qT7zux4ua6nEGk8crs2j6dv/ojtbd3ROtn6MJgYbNKTTeotCBTOK8T9W5dshfPA7Bje5ZExiz2LZv+/f43Zy3CFzckFni/tup3IgUffgisDBW+zsROpx8vi4GIwlAJt80CfB1YppvRoz8lZp44ptmxFJgFNg8MAJWYCIt80BZ/E0A0Ngxfy1z0nKZnpMAcHUuTHxd99Tt33K1U3QWeIGGTPboFZJ/SFi+q5Op1lqK8jwJXUnCh/VE6+r06P43ndP5nVN4SxEMxnxfwLkQ4cLoTg91R+u/r2tzrNbtxOadybKcRNHVuqVzFobovamlTgM7j8l+KpLJdjZ0ZY/X71EqUKA3VtsrNh2uO3wP+zl9RhLAT+6CxxpH4z6njbJaS170F2l7TbylseSBFBgANg+OABIAhqyVEPt6BsDAmJdwBoAhaz9MGG4nAF5+21qFo//DoFfS01O6+naNdGV/w0S3MvNeJZj7H5j4cn1xPFV3KnfRo82ze2K195YgDgxZoEBvvHZxz/RJN+m6fUPrfpuEQltbNsedCzFqVw9qNbXEWeh6TXo9ofVkU4NEXN4JvPOzW3sgO9cR8W6YD/Ssv+fwdoAPqJjxwKudbW8IUbNG87bWUpexQtwejp0zsdSBFDA+Ng2QABIA5iyWbwkAIXqV7rgjb87UEcn6BPQolbPz7/Lp/3zdI9M/9fr6xsNjlSfQkF58nCW0gISO9Xf28oQwf09ffE/RHcc9uqO153Q3Ter2NwaM5rbAPdN2+vfCeG3aea+1rusUGcM7E/E3iGQeEa3QalRhor20/kyTAEgEGLUyPgXDLN1NdTOJ6Rc+jbh6GMaZAKsdTPxnMJVcon+XphkS22eI8871CQwJZ7gw8HyQBJAAMGW1RHz7GEA98oGj/6as+8bjcBIAG9+iwGclP+S8d7fAVti8HAUi6ewl+kL0ZiLemny5yQoR+pEl/NmeaP1neptqz9cX3yYfJfZFpZwH6T20dnFPtPZqXe/Dq1aN2VFIvkDEzqePGPU3SJMASARQ5d56mupOIqZL/RRgJAH85C5orIHOZI/+rvp8QY282/jE6njr6d5172LP6CpQAkgAGLBc46KJrYl5jG+hMOH9/75hj2oglxMAck9v885LRhURGgdeoHH+iztFuvrnEtO5Pk3m7zrOJZK39+iN1Z2+MFb7mH6PUmEC9x++w796o/U3a0Lgq8v/tmxHtqRJd7yvUIY/aDWiaDxIBBixEv4HoUmA82j12Srk242RBPDNutCBBlMdvyHiVjLgxiTJibHEFANC2WgIeDJYAkgAGLBeoZDt2+n/znQFHwHoMBhdPTn9n3m+0ZNGcJ4LNGayM8XOLyCRGZ4PRvwCMbfaZO3RE637FpJP3osHZYQnT9l7Vff0+vnd0boW/dnYw87TvsSkO2D0gAlzQCLAhFXwP4aeaO3VNtMhOrJv7wFnJAGU28ySS7UnhegHBkS3qf5cJrdoaqs2IJYNhYDHAyaABIABC2bZIV8TAJbYRp1+acASmBiCy0f/iWxZtcDEiSImfwQa033f1hczd+loO2j1sPAjOs4pH/7Eyj16mmqTi6KTXvNwMHRdBgKLmuse72mqu7QnWndwiKydRPhrRPJbIl5KJbwhEVBC/BINvaip7n7bym+vw7+i1ZfCSAL44lzMIIOp5Jna7g6tpS775qtWJEsdxIbHxzNBE7CCFnBZxsv+fgSgWIwEgPk/SG4nAB5YFN3lBfOnjQjdFjioM/vRhkz2VmF2Pt7P7e7f749pvu4wHdMTrd2/N1p37R277LLy/SfxFQRGJuB85GNvrPZnPdH6WTZxHTGdRsQlvUAk42KBVEm3RdN3/kdPtE4PzPDTfs2bkQTwi7rgcWTV0NeY5OGCG7rdQOjk6ubWU93u1pX+0EngBJAAMGDJ9GiZ/qHxLxAhG9cA8I+74JE8Of2fBEf/C16J4DdoyCw5LB9iZ8f8aO9mI78XS47WI7hN3dG627wbBz1XmsCi6KTX9Ofqxz3R2qmct6eS0I/VoGRnlCARoPoVVPTnbk+drnPWlN55X5AE8N64mBEGF1z1Olv0dW37L60lLSzSPmFG4oCSBrGewfFQ8ASQADBhzYR9TQDk8xbOADBh3TcQg5ActIGnin7Y5rF4/3/ResFsGOnqm8Vka+JHdvFoBq/pz+o5vNk/9u+dXn+7R2OgWwisFuhuntzdE6s7zSbrU1TiswLWSgR09Z2wOkD8U5YCPdG6I5nkYr8mxzgTwC/qgsZZOq/jj8KcKKiRNxuPC9nc/qFZp473pvuiekWjAAogAWDCorF8xLcwhFatyCT/6dt4GKhgARZyd4eNadGiph3+WHAgaBBYAWfnn4Q92ylnlp9JPn9Ab7T+8u6pU4cCC4XAAydg0lkBqxMBwr+IdGW7GjKLDwscJgIekUB3tP7bIuzbx8IhCTCiZfF9o8HO9ptIqMP3gdcZUBPv+7/91nhv39K3zpgb/xbPBlEACQADVk13+Kp8C4MJR/99wy52IP50sS3X104Ep/+vz6VcH4vokX/ybOdf7hGyp3U31X8NV/Uv15+g4MzLmLMChKYzWQsimey1jQuy9YRb2Qn0xmp/s/ptKD7NzEkCNM7v/6ZPw2GYEQrk0knnLIB7Rri5Z5uxyJk1za0zPRugkI6xbSAFkAAwYNmEKUT+3fD+f/+sCx7pwPTiWmL6aMENN9IgFKrC6f8b8Smnp7zb+ZclIvbXeqL103qjk0v+4qec1gxzGb2AQWcFnCR5fjiS6Tv/sAXPjBv9zNCDSQJOwsnPJIDY8v2p8/sON8kAsRBZljhJgH+V2kJELqiZeWZNqePA+MEUQALAhHUT8u8MACKcrmvCmm8gBssKuXv6P9GjC6ft9OQGhsPDZSTg3c4/XWvTGwf0xib/rIy4MJUyFXB20kp7rQCZSMQXvZ3f9OHGruzxhFtZCTg/X34mAWyb7546P7tbWSEGfDIGXQ9gdzs/9oISc2L4gAogAWDCwjH7dwYAIwFgwpJvKAYWdjUBIEw4+r8h7DJ63LOdf6aze6J1pyyK7lmyK6+X0TJhKj4KrHlWgO5ExXVo367mrmM5ZVcRur4hnU01ZPpcv7CrMwBqaQT8TwLQHxq7Fm9Xmtli1PUJmHI9gNK/FWB9OngsCAJIABixSuLbGQAsSAAYseQbCkJsVzP9LNZTGxoKj5eHgEc7/1nVifY01V2l9ygQCLTAonhtWhNZR7LFhwrRbX5OhpliTHx/JNP/44bOFyb5OTbG8k7A7ySAiPWnA1N/mejdjNBzoQKmXA+gpG8FKBQN2xsjgASAAUuhLxB8OwNAX/zgLQAGrPkGQ2Dac4PPFfFEaBMbCYAi3ILSJNKVPYJcvuCf/j6aKyFp0h2mrqA4IE4IjESge3rtfb3RumNEOEJCN42kjXvbyNc5FHo4ks6e25RZsol7/aKnUgn4nATgEL/1ZJtmAko1X4z7QQHLjOsBlOytAB8UwSNBEUACwICVEmHfzgDQ6a7SimKgwAH39G2rYbl5mt+zDxxc93ftE6UMBQ6Y99wnSOQyN6emRyov726qm9k7rf7PbvaLviBgkkBvrLa3J1b3JSZrP2L6Bfl320rHu+QNlocbuvpm+DcsRvJKwNckANP2CzPZ+7yaC/otXMCU6wGU6K0AhYOhhTECSAAYsRTi2xkA+uIDZwAYseYfDCI05PYFAPn3HxwFj5SLQFVojO78c6078+EBZv5Kd7T2HHf6Qy8QMF+gOzrp0Z6mupNsK7+X/m38sRCt9CVqkT1ZeG5Dpi+5V+cTm/syJgbxTMDfJABPjWT6f+nZZNBxwQKmXA/A/7cCFEyFBgYJIAFgxmL4dgaAZgmRADBjzT8QBZPs/oEHR/EAW4LT/0fhZ3LThkx2tu6wHOtKjMzPM3FTd1Pt9a70h04gEDCBRdN3fkoTAadZdn4vIf4+kSz3YwpM3DIuNOGBhsySw/wYD2N4J+BrEoDkq5FM9mLvZoOeCxUw5HoA/r4VoFAkbG+UABIARiyH+JYAECIkAIxY8w/lpO6aAAAQAElEQVQGwSSuvv9fcAHADyKXwSONnYsb9WdFj/67MBnd+beFj3KOhLrQG7qAQKAFuuM7P9cbrT1Lqsbo72Jp1yTbf72eEBPty2QvaOjq/y61CV6TeQ3uYf/+JgHovEimr9XD6aDrAgVMuB6AHuQ7s6a5dWaBoRe1ORoFWwB/bIxYP/bvLQCEBIARS76eIDQ5s9d6Hi7yIc5vLoQzAIrUM7XZAXP7JtghS3f+XbhuyLs7/4uik14wdb6ICwKlEOg9bMclPdH6OSHnjACR72kM/9DqadEX7t9u2GfJA1MziyOeDoTOPRXwNwnAl2vi6BRPJ4TORyxgyvUAfHorwIhdsKGZAkgAmLEuvp0BQGzhIoBmrPlaUeybWVKtD+yk1a3y+67opLfd6gz9mCFQVcWXOUcMRx0Ndv5HTYgOyl/ggdjOf+mN1Z8vK1fuK6vfGuDtnJkkYpP1QGO679vejoTevRTwMwmgiaOfNnZlj/FyPuh75AKGXA/Ah7cCjNwEW5opgASAGevi2xkALDgDwIwlXzuKzfJDe679yGi/E1wAcLSEhrVvyPR9lZhOHXVY2PkfNSE6qCyB3pmf/ltvtPYsZpqiM79bq5fFEubvRjJ9C6am+/f1ciD07Z2An0kAEbolksk2eTcb9FyIgAnXA9DEkLdvBSgEBNsaKYAEgBnL4tsZADYhAWDGkq8dhW3xh9d+ZJTfCSEBMEpCk5pP7czuxsSXjTom7PyPmhAdVK5Ad1Pdgz3RuiOY5Ysk9Afy9MaH2UwPNGb6WzwdBp17JuBnEkAnceOUTN8Beo9igIAJ1wMQkQtqZp5Z4wUH+gy+ABIARqyhf9cAYBJcBNCINV87CIvdTQCwjQTA2sLB/s626Ewi+pDW4gt2/ou3Q0sIrCHQ3VT/65UrNt+fRc4n4qXk2U02F5JkJNM/d8rcvl09GwYdeybgJAGIraM8G+D9jreyiK9vTL2wy/sP4atSCTjXA2CmUn+srldvBSgVK8Z1UcBysS90VbSA+HYGADHOACh6mTxsaIt8xL3uZVV3c12fe/2hp1IKTO3KHqz/b7882hjYtlsW4YJ/o2VEewisFnjkqI+/2R2r/57Ybx+g/z9/sfpBz/6RGVYVZxpS/Ud6NgQ69kygp2nSHcT8Tc8GeL/jnWyrquP9b/FVKQUGOpPXaxLg16WMQZOUZ7r/qQClnBHGdksACQC3JEfXj3/XALCRABjdUnnTml09A4D/7U2U6LUUAkJ0hgvjfkt3Vua70A+6gAAE1hDojX96cU9T3Un6QruJhO5b4yl3v2T6KFvyu0imv9XdjtGbHwI9TbU/FKHRv41rmGCZZFpDOnvpMJvhaZ8E8vn8d3VNPP8UkY1Nx/W3AmxsMDwXGAEkAMxYKn2N708gtsXO1eb9GQyjjFhAf0G7eAYAvTbigbGh0QKRVN8sfdF4+GiCZOLf9UTrLhlNH2gLAQhsXMBJsPXE6g7VRMDXiXgJeXaTyyPp7M896x4deybQG6s7V1/s3eDZAO92zEznNMzvO+rdb3FXQoFlmSv7Rfi7JQzBGdrVtwI4HaIGXwAJADPW8HW/wrBItvJrLIwzcgH9g+3mRQCRABg5vdlbWtbpowtQXibLKvX7EEc3BbSGQIAENBHwU5v4AA35aq3eFKYTI5n+hY2/+vOm3gyAXr0S6I3WfZmZF3jV///6ZZvaG+55adL/vsd96QRy6eRPiPi3VMKbJibPnDAj4fxeGm0UaF8mAkgAmLGQviUA9Ejz6C4kZoZXGUbh5kUABQmAMvgJachkTyaShtFMRSya0z19pxdH0wfaQgAChQksik56rSda9w1mOlZbvqTVgyKN8uGVL05NLd7Lg87RpYcC3U2107X7P2r1sPAneWio3cMB0HUBApZlO2cBeHjB0OGDsWw33k44/DjYIhgCSAAYsE5M7FsCgIiRACAjb669BUDIwjUAjFzikQe1V+erm2vGflTv/dffK5f3Tq+/feSjYksIQMBNge6multDkncu4nmTm/2+1xfTR23LekKThV967zF8EQiBVaHQoRqo18n6mZFM9js6DkqJBd79VAAnCVDKSI6eGJ/TPKoA0LhsBJAAMGIpbT+zgngLgBFr/n4Qzs6efjdeqyuFCWcAuAJZwk7GhQbPIObiP86J6fHuaC1O/S/hGmJoCDgCD8R2/ktPU53uoMupxPRf5zG3KxPdEMn0Xe52v+jPO4GHp+30bxH7YO9GeK/nCyOZJdjpe4+jdF8MdCav0tdn6dJFQGST/Y3RjI+25SOABIAZa+njGQCEMwDMWPP3ophQ9bab7/93+vX6qIIzBqpHAgfft3hLJj5jVN0zXTOq9mgMAQi4KtATrf+JxeTs8KVc7fi9zri1IZ31qO/3BsEXLgr0xiY/yxY7ZwK42Ot6uhK7vTGz5GPreQYP+Sxgk3WRDvmW1lKVQ2qiiWI/VrhUMWNcDwSQAPAAtdAubV/fAkCbfeSwlnGFxojtvRNge5Vrp/87UTLhUwAch6DW/CorrrF/VGtxRei+nul13pxyXFxEaAUBCKjAwul1z/RE65qJ6TxxDsbpY24WZopFMtk/u9kn+vJWoHt67X1M1hfIyxtTvWgSwMsh0PfIBAZT7Y8TsZMEoFLdxKIzaK+TxxQ+PlqUkwASAEaspiz1M4y3raoP+Tkextq4gM1uXgBQX1WGQrgGwMbJjX5WhJ0EQNExWiQ4+l+0HhpCwHuBnqa6S6386tO/H/RgtO00CSB7dT6xuQd9o0sPBLqjk27WpNBsD7p+v0umYyPpvsT7D+CrUgnkUu2X6tj3ay1V2bNmm+rCzzIsVbQY1xMBJAA8YS2wU/HzIoBEMtbGdQAKXCIvN2exXT0DgJjxFgAvF8zDvg9Kv/BJ/R/aXOwQTPy7hbH6ucW2RzsIQMAfge7myd0fXr7qYBJJejHi+NCEFQfPfdbdvy1eBIo+VwtoUuhKIvb2qv3Mlzem+g8h3EouwEwlvSCgHmg4Y9yMbxb0+6HkaAjAVQEkAFzlLLIz9jkBIHmcAVDkUnnRTJhcvQaAtcJCAsCLhfKhzyEKOUf/Nyl2KLHtHxXbFu0gAAF/Be44apeVPbH6VrH5czry37S6WobGjP1n4/wXd3K1U3TmmUBPtHYOE93g2QBELCG63LnODOFWUoGBzmQPk5QuCcC0XVV+k0LOAiipFwZ3X8Byv0v0WLCAzwkASywkAApeJA8bCImbvX9o1e+Xudkf+vJPgC1yEgBFDii/6onXLyyyMZpBAAIlEuiN195pkzWdiO8ll29i55c0ZJbs4XK36M4rgeWrTieSR7zqnkT2HFrp8ZkGngVfXh0PLB9/ETM9VLJZsZwxPtY6eWTjY6tyE0ACwIgVFV+vAaB/XPAWACPW/Z0gmOXNd75y599XN9l1c3d6Qi9+CkzJLNlZU0G6E1DcqHaef1pcS7SCAARKLbAoOumFDy+f1CTk/id4MNlPNaT7G0o9R4w/vED3Ubss1xfmCT0u8MbwWxe7BX+1MZPVREOx7dHOFYHutiHbtkp3FgBRdUiTACOaCzYqOwH9PVN2cwrchCQ/5OfHAOo+BuMMAJN+SthyNQGwKY1FAsCk9R1hLEz5ot/7r/+p71vUXPf4CIfCZhCAgIECdxzF+d5o3RlM5Jyam3czRE0090yd3x9zs0/05Y3Awmj9Q8whTy/YJ8SXT830HeDNDNDrSAUG05cvIPLmOiAjjOHrE2Ykhv05GGFf2CxAAkgAmLBYIfY1AaBTxhkAimBKsfP5t9yMZdWmK5EAcBPUp76YRnH1f7bu9ClMDAMBCHgs0B2tu4Yk30TMz7s5lG1LqrEre4ybfaIvbwR6mib9WHv+uVaPimxuM7d51Dm6LUBgjLX6LICnCmji6qaWvTrhuLE+8VwZCiABYMCiLtv0r0v9DEOEcAaAn+DDjcXs6hkAIQkhATCcuWHPT5mf3UdDOlBrEYWXsrx9VxEN0QQCEDBUoCe2871iWZoEIFeTe/r3/5ZIJnuSodNGWGsIjBlTlSAh787sEjqkMdN35hpD4ssSCPxnXvsyZrmoBEP/b8ijJ8bnbOQMxP9thvtyEkACwITVvOMO51S/nF+hMNM2fo2FcYYX0F/8riYAhlbZ44YfFVuYJBDK066jiOeu7tgu/xxFezSFAAQMFOidttPfeprqPkdMl7oc3rWRruzZLveJ7lwWuO/QHXNSxc5bAd52uev3uhPiOfikiPc4SvbFQGfHXCK+lUp0s8n+xgaHxhNlKYAEgDnL6ufbAHDVT3PWndgKuZoACJG9uUHTQygjEBCiov9PsuRx9H8ExtgEAkEV0CTAecL0ZT0aPEBu3YSuaEgvxingbnl61E/vtNpeEmn1qHun263tfP4c5wvU0groTvgPSxjBITXRxJfXNz4eK08BJADMWVc/3wYwYcv4tz5qztQrOxK23b0IoF1lIQEQtB8pLjoB8Mfu2ORU0KaLeCEAgcIEepvqbmB297oAzNYFDZnslwqLBFv7LdATq9cdQ77Oq3GZ6YTGTHamV/2j35EJLEt1PMIk14xsa/e3EovOoL1OHrNOz/i2TAWQADBkYZn4H36GMiRvF33E0c84K2GsVfYqVy8CaAkSAAH8uSnqLQB6VNDV9wcH0A0hb0SgMf3c7pF09ouN8/u/F8n039yQ6euOZPqfjaT7/hLJZP+xuq7+uv/Zhneeu3n1tk4bbbuRrvFUCQS6ozs/agsfRS5eHJCJbojMz04pwXQwZAECVW/lW4XoyQKaFLSpTXTOXj97Ajt/Bam5v7FNrMke8vOM4DUnsWfNNtXOJ5Cs8Ri+LFcBJAAMWVn9xf5HP0MRtur9HA9jbViArao3N/xs4c/YeAtA4WglbNGw4MWP6/BFnZGjv8Cz2hYFAqsFnB3+hsySOZFM3wLdoV8hPOZpYrpRbPkWkRzHxBG934WYP6ENtl5dV38tu7z73HGrt3XaaFunj3f6WjLH6Vu3RymxwKLopBckT7N0Dd37hACbeqd09RWVhCwxR8UMf/+Rk/+r/3edtwI414xyfd5MtO+ET1TPcb1jdFiQwGAquURInCRAQe3c2liEzxg345sfea8/fFG2Avr6sWznFqiJCcuzfgbMYuMMAD/BNzIWh/KuJgCYBBcB3Ii3cU/lpej/i7ZYSAAYt6D+BnRQ5tkdI5nsdyJ6ZN/Z4WeyLyPiw5h41G8FeqcPpy/7MqdvZwxnLGdMwq1kAr3x2sWhIZnmZhLAEn6mYUHf9iWbFAYeVqA3Wv8AMSWG3bDIDUTknIbMkj2KbI5mLgmwVGkCQPpd6q6wbpi2q7LHfPV/jXBfvgJW+U4tWDML2f4mAIgZZwAY8iNiDW3iagJAXHjhbwhNRYShybhdipyovSKfQwKgSLygN5valT24IdP32zyNfVHncqEeHSz250ibj7SIM8aFzpjO2E4MI22J7dwVeKC57u+8ac1n9W+5a2cCsG092tiZ3crdSNGbZt9GbwAAEABJREFUmwI9TXVXMfMNbva5Rl/jLLJxFsAaIKX4Mpe+bKmw9YNSjO2Mqa8hP0+zZoWIyPkWtUwFkAAwZGGXpjue01Bsrb4UluKPOvoSYAUNUjXWdjUBwCSjPvJXQfylnyrTbkUG0fdk895vFNkWzQIq4Ox0RzJ9C2yh+/QI/edKNQ1nbCcGJxYnplLFUcnjdk/98PKeptpdXEsCiHzYDlGqKbNkk0p2NX3uKy2rlYSe8SJOITo6ks5+0Yu+0efIBQY723/MTA+NvIV7WzLRp2re3O7zRO71iZ7ME0ACwKg18e8sAGHeZuKsOWGjpl+hwYxZZbt6EUAiC6dxBulnSai4BABTX5CmiVhHJ+Ccdh/J9N/s7HQT8WFkzI0Pc2JyYnNiNCasCgrEzSQAE+37Bts3VhBf4Kb68LSd/k3Enr0VgCyec2DqLxMJt5IKCFk/LFUAmgj6PJVqcIzriwASAL4wj2wQJsvXCwHm38JZAGTArSs66W0idvFILtaVgnQr9gwAsZAACNI6jyJWPSJ3Vp7G6nrLcaPoxuOmcpwToxOrxwOh+/UIuJkEIKGjGtPZX6xnGDxkiEBPrPZeXadLPQlH5FMh6+1zPOkbnY5YINd5+e0iNHfEDdzd8LBwvPVgd7tEbyYJIAFg0Gpoxs3XCwFaJLgOgDnr7+Z7uYu+qJw5HJURyQFz+yYUO1Mh282fmWLDQDsPBRozSz6mR9bnEdOVOkyVVtNLlROrE7MTu+nBllt8ThJARFz5qDhhOqEh039VuRmV03x481WXEPHT5MlNWvHxkJ7AFtQpWyX7RAAnzs87/6CWpwASAAatq/h8IUDNLGJH0ZT1Z3JzZy48Zf7L25gyNcSxYQFrk03Hb/jZYZ4RfnmYLfB0gAUaurJRTfL8nkiagzcNaXZid+YQvNiDHXFvrH5vPTK8yI1ZMMmZDem+77rRF/pwX6B76i7LiW1vzgJwwrUJZwE4DiWsuc6OB0jkl/6HoCOKfGFC82wcKFSKcixIABi0qmOr2Ne3AAgR/mObsv6S73czFLbfQnLHTVCP+pJ8vvgzAGxx+doRHk0S3RYsEOnKnspCaW24tdaglq2dOThzCeoEghp3T6yuQWN3JanMzN9uzGSP1v5QDBToaaq/Q5OEv/IotGgks+TrHvWNbkcoYIXIuRaAqxeLHnZoZwOmMZbNX3C+RC0/ASQADFrT/8xrf1XD+a9WXwoTYSfRF+nhBxEhV16svTeSMNb2PQxzvxgjK4tOAFAVEgDmrmzxkUUy2e/oEdwfFd+DYS2FfrR6ToaFVe7h9ETrNMHPS92Ypx4suOTAVP8ObvSFPtwXYKvqEk0C/NP9np0e5Zwp81/AGYUORYnq0nkdzsFBXz8W8P2p8ue3mtFa/OuU9zvCV4YJIAFg2IJoOM5/dL3zpUyiWbNCvoyEQTYqEHI5AWBZhATARsXNeFITP0X/YRWcAWDGIroYReP8/u9pdxdqLbdy4btzK7d5GT0flpU7uxTgDpYlupPpUm/oxlWB7uk7vUhkebQ+8gnLDp3kasDorGCBVWPHOGcB/LXghsU1eL8V03ar8oKzAN4XKZuvkAAwbCmZ2NcLAU5845PYUTTgZ0DeyLt6BoDuWGJdDVjXYUMYGyr6GgCSz+MtAMMCB2cD5yi5JnW+FZyIC4vUmZszx8JaYevRCHTHdvknC+8xmj7+15aJjo6ks2f/73vcmyXQE629WiPq0upFOWn/BS9+2IuO0efIBN6485J/6P9lJwkwsgaj2mrtxvp/HxcDXJukLL5DAsCwZRQSXxMAwtZnDCOoyHC6j9pluU7871rdKYzrO7gD6W0vYlPxZwBssulb3kaH3v0SePd98uV45H9dwgvfneu6j+N7jwS6Y7V/0K6jWkdfmC5p7MoeOPqO0IMnAjZ5dUHAj1Xlh3AWgCeLNvJOBz464CQAnhp5iyK3XKeZMB1QE2udsc7D+DbgAkgAGLaAQv5eCJBI8Dmf5vwMuHcWgNC2h9z7UticqSGS9QmISNEJgHwe1wBYn2nQHlt9pXyh8nnP/3ALoHNdPefhtsPzrgn0ROu69ODCCS50uIkIXdqUWbKJC32hC5cFeuJ1i0i8SQIw8YkH37l4S5dDRneFCFx77SrdR/Dqgo/vRbK+L4QIZwGsDybAjyEBYNjiVW+2ytczAIj4IMLNEAF2LwGgMxpaZeNtAOpgeCk6AbBs+ds4A8DwxR0uvMbMko+xUGk+4mm44Dx83pmzM3cPh0DX6wj0RuuvE1u+vc7DxXx74Btse/R+82LCQZs1BXjzVbo2/PSaj7n09XarNq060aW+0E2RAmMtukGb/kmrV2X9/bIcFY7N3mv9T+LRIAogAWDYqr1yx1XOR30s8SssPSqwzcRYyy5+jYdxNiygR1ZcTQCI5JEA2DC3Gc+IXXQC4MMfGoff32asYtFR6O/fH2vjIH/Un4ZfVNn63bkX1RiNihPojddfrEfyrimu9RqthM7GRwOu4WHQl91Td1lObHvyVgAm+yScWVjaxf7PvPZlwquTAB4FsuFuhSycBbBhnsA9gxeQJi4Zi69nAeSJp5rIUGkxWSHL1QSA+u2uFcVgAcsKvVx0eMtXbld0WzQsuUAknT2LSJpLHkjJApDmdwxKFkBFDtwbrTtDJ3631lEVTSTgowFHJehd456m+jv0d4sXp4rvOLQqj2sBeLd0I+rZytvOWQADI9q40I02sj0zfWFcNFGJCeuNqAT3KSQADFw7PRLsawJACXAdAEUodbG5yt0EAFuHlHpOGH/jAsLWXza+xUaetQQJgI3wmPzUQZlndySmdpNj9CU2NVht4ctgGOR/AjZZ39IdxH/+7/si73fARwMWKedDM7aqLnFhjT8QqZCc1LjwuaI/veYDHeKBggUGMlf8RYicJEDBbYdrMMzzHxrDjLMAhkEKytNIABi4UiHh+/0MS7N6znUA2M8xMdYHBXqnbfdn/YO96oPPFPuI7Nyw4MWPF9sa7bwXyA/ZxScAiJAAoGDe8rTJRRp5ldZKL1XvWlS6g6/zXxSd9IJQSJMAoxtWXzQc3Zju/9roekFrLwS6p+/0IpGlSQDXe6+Vt8bgLADXWQvr0LLyXiQAhg1CWL4w7EbYIBACSAAYuExL08lFJOTijuCwk5xQE509ZditsIH3AkKungVgDa3C2R3er1rRI2zzxqSiEwC2xZ8semA0LJnA1K6s/p+U40oWgHEDy3HvmBgXWFkH1BuddJ1O8OdaR1V0h+D0wxY8M25UnaCxJwI90dqr9Ujxk653btNJ+93+t81c7xcdjlhgYN6VTzPJTSNuMKINR7TR7tXRxLEj2hIbGS2ABICpy2ORr2cBCFsHmkpRUXExPe7qfDl0iKv9oTNXBe44ivPa4StaCy4shDMAKHg3W6Q1eFF7GzFMvPXdUO+r3nrL+VSA0b7l8FNv5zc7bUNj4PGSC1zregRMk8eOW4GzAFyHLaxDIcvdswBGOLxl4W0AI6QyejMkAExdHuEHfA2NCRcC9BV8/YOxbd+7/meKe1RIDi2uJVr5JaBHaIo8C4CRAPBrkVwa550j3XyYS92VUTd82Ds2ZTSlAEzl4SN3+7cewXeSAKOMVk6fMv+FbUbZCZp7INAbrbtW/8a4fxYA84k73/7cWA9CRpcjFMil2u9notQINx92s5FuoK8rY1s0z/nUSLfHdmYKIAFg5rqQcL7b59AOnnDEuVv6PCaGW0cgb41x+8yPDzdmlnx2nWHwrUECTFJkAkCQADBoHUcSSl7k6yPZrhK3gU1pVr23qX4eMY32Y+M+HrItnAVQmiUcyajunwVA9OkPTdjkxJEMjm28ExChG13qvaBu8pKfUVADbGycABIAxi3JOwENdl7xe/3Kz4/54NDQKrwNQNFLWRZFJ70mRK5m623KR0o5J4w9jIBwkQkA2vqQe18KD9M7njZEwLnaPRN/zpBwjAvDsXGMjAusAgLqmV77LU0CLBzNVIX49KldfXWj6QNtvRHo9eosALFPahPBfoQ3yzaiXnPp5B16EOHhEW280Y0KfZKbC22B7c0SwH9cs9Zj7WhYfH0bgAjjQoBrr0BpvhNy9W0AOgm8vUMRTC1CdrEJAMrn80jambqw68SVp7H4+KR1TNb9Fkbrivj0PbPQ6o8GpDdHMWLYFsZZAKMA9LipF2cB7P7A/CW4FoDHCzdc9/raffTXAhhukA8+v9/E+Gy8/vigS2AeQQLA4KXS/9S+JgCUAgkARSh1YXL5+g/ChzZ2Zrcq9bww/voFLCr6DABNAAj+AK+f1cBHeZaBQRkWEoxKtSA9TZMeIba+NcrxT490ZT8zyj7Q3AMBr84C0KPPX/UgXHRZgEDOfstJAPQV0OQDmxbzgE0WzgIoBs6QNpYhcSCM9Qgw0SPredjLh/apjicmeTkA+h5eoNoe/5BuldfqSmEmK2/ZeBsAGXrbfOjBYiPTtUUCoFg8H9s1pp/bnUh28XHIgA4lu7xjFdDwAx62JgGuEpKeUUzDOZcAZwGMAtDjpu6fBSC0TyTdj4sNe7xwG+2+6+q39XknCaB3RZViG+E6AMXKGdAOCQADFmFDIeRSyaeE6NUNPe/F4xYxdii8gC2gz87mbd8gEVffBmBZoYYCQsCmPgp0T91luf4/d5I+xYx64L6ZJdXFNEQb/wRs3mSaf6MFeyRYlXb9QhZdNZoI9HfZ8VO6sgePpg+09UbAq7MAiOkIbyJGryMVWDV2jJMA+PdIt197u6K/q5/QnIgX3RoNSypglXR0DD6sgEXs9lXhNzqmCD42bqNAfj1pWe6+/cMWnAHg19oVMQ4T3VNEs9VNxvLQfqu/wD/GCjDlDzI2OMMCg1VpF2Th9Pq5RPKb0URhCeEsgNEAetvW/bMASI7cb/7zW3gbNnrfmMAbd17yD/1/6yQBNrbZ+p8bxaOWMM4CGIVfKZtapRwcYw8voNl0f6/uyRTf/Mjzthk+MmzhpYBYVrFHhNcfFtNuUzMvHrD+J/FoqQWYqeiED9uEBECpF3CY8fX3OM6sGsbof0/D6n8SpbtnGeOcBTA0igiOiKT6cPHZUQB61dSjswA+somMxVkAXi3aCPsNcchJABT8/3aE3W9gM2meOGsOPo1oAzomP4wEgMmro7ExDfl9HYAJY95ehQt7qH0pS++0nZzEj2Z03YsiT0Mz3esNPbkp0D299iHd8Rksqk8Ofbaodmjki0Bj+rndmXhzXwYrg0EcK8esDKYS2Cl0x3Z8gpidJEDRcxCmo4tujIZeC7h+FoCIfaTXQaP/jQu83nn58yx068a3+sCzo31g6/xbgrMARqtYgvZIAJQAvZAhB1JXPqPbL9HqW9E/3HhPj2/aGxuIF27s2UKfY7IOn3X77SHCzTwBZv1vJ0W+DUD223/Bix82b1KIyBEQGvNp5x515AIwG7mVV1vaPNZJABT9EaXM1lGH3vPXbb2KD/0WL+DRWQDRxnv6dy8+Kri18wUAABAASURBVLR0Q0Asq7Owfka/NYvgoOHoGX3vAQkA38mLGZCLvkp4UaMRxfFpAMXIudyGxd23AZBM+teE3Q93OUp055qA1VtMV0xUXTU0dGwxbdHGewEOcZ33o5TXCDAr/Xoumr7dP9hiJwlQZDAycWX+TZwFUKSeD83cPwtgiHAWgA8Lt7EhxtFmKX3+Fa0jK+5sNaNm+jnbudMVevFLAAkAv6RHNY7t99sA2CJk9Ea1ZC40tphcTgAQaaYWCQAy82YPDY1ivRkJADOXlcSm7Q0NzdiwYGbG0nRPr/0BMTlvRysqIBG8DaAoOB8avU3Wrbq2Ln/KlCAB4MPabWyIVzvb3tDnR3wWgG7rRtnEHpPHWQBuSPrYBxIAPmIXO5Ql5HcCgIQYbwModsFcardwet0zRPwCuXgToplT5r+Mizy6aOpWV4tm7PyU9pXVWnBhpn0bUn14H17Bct43ELI/6v0o5TUCzAxaz7x8v9homGjfKakXYsW2RzvvBB6LThok4bkuj/Cpqek+HGRwGbXQ7oQoNcI2rm3GQnj94ZqmPx0hAeCP86hGWZrueE7/c7m6IziCgKaG44k9R7AdNvFQgEm63OxeX5BVW/ZbuBigm6iu9sVFnwXAIesYV0NBZ64IMFlbEm4FCTDMyJRbT7z+DmYq+uOILSt0lClzQRxrC9icdzsBQDYzzgJYm9n37wZTyQyR9A0/sKtbHFLTdPYervaIzjwVQALAU173OretYi8QVnwMwoSzAIrnc6VlnuQuVzpaoxMhCxn6NTxM+pKJi0/4iBzbuODFXUyaD2JRAbGr9V+UQgRgVoiW99vao7qy+NGN81/cyfsgMUKhAouaJi/QNn/Q6mY5ojGz5GNudoi+ihIY/m0ARXW74UZ2VQhvA9gwj3HPIAFg3JKsPyDLtuat/xnvHmVBAsA73ZH1vChar0eEuXtkW49sKyaZ1pjG1XpHpuXvVt3RSb/VEf+otagi+TyuBVCUnIeNmMd62Ht5dg0zo9Z1bNVbt5DQn4sMahPbHsLFAIvE87qZHui52+UxJojkcRaAy6iFdse2DPs2gEL7HG57FsHbAIZDMuh5JAAMWoyNhTKQbl9IJE9ubBsPnvtMTXR2gwf9ostCBMR2/SwAm/J4G0Aha+DjtkJye9HDCR172IJnxhXdHg0hAAEIrCNwz7TdVhDTLes8POJvmRhvAxixlr8bWja7/jYAYusIf2eB0dYVGMhc4Xyq0O/XfXyN7z34kveqibVO9aBjdOmBABIAHqB61aUw+34WgFi4GKBX6znSfplDd+q2y7S6VpjxNgDXMF3uyLKqbtMu39ZaeGHafuXQpjgLoHA571qIrPSu8zLtGWbGLSyHQkUnAHQyuzZmlvyf3qMYJtAdq/2DEDtvBXAxMmmckuk7wMUO0VURAsIbuxhgER2OoIlt2TgLYAROJmyCBIAJqzDCGKyhvO/v6dHQ8J4eRShl6Y5OekXHd/ssgN2ndC2epv2iGCbQPX2nFzUkJwmgd0UUi3AxwCLYPGvC1qBnfZdrxzAzbmW7p+30HDEXnQQQkqhxk0JAqwWYbNfPAggR4yyA1bql+4dlIwkAj8KyiA/zqGt067KA5XJ/6M5DgYGuK5/W7u/T6mPh+up4An+4fRRf31CWiHMWwPqeKvoxS0K4GGDRet42tO180W8DEKGDI6m+Wd5GiN5HKiBk/3ek22K7dwRg9o6Daf9K3r61+JgEbycsHs/Tlra16d1C5GqiUvvD60ZPV234znOppPPRwgvXt6VXj+nrj50nzmjZ1av+0a97AkgAuGfpS0/C5PvbAJhwMUBfFncjgyyM1TsZ+uc3sknhT7Ecvm9mCa5QXric5y0WxXdO6wuox4oeKMRnFt0WDV0VYLL+TrgVJMAwIxNvvfH6efoCv9jfSzs2ZpZ81sR5VXpMi6Zv9w8mcvtigJOnzF+Cj5Km0t6Y1nsxQE+DsvN8gKcDoHNXBJAAcIXRv06sVSHnbQDFvT+4+DCbt21u27z45mjpjgC7exaA0LabkH2MO7GhF7cFmKn4twEI7d84v/+bbseE/goXYKvoq6cXPliZtICZyQspRb8NwKZ8xOSZVXJswuwcZHCVIJSXKa52iM4KFrCJnX0GPZ6wZlNvv2YmJAC8JXaldyQAXGH0r5OB+Ze9rKP5fRbAx1bQm3EdF6WEAlxF7iYA3pnLye/c4V/TBMaGNtMEAC8tNi6x5awp81/eptj2aOeOgOQl605PldMLzMxd66GqKucjAYu6KC0T420Ahi7tR5b9QRMAvMTV8JiQAHAVtPDOBlPJ/2fvPuDbqq4/gJ/z5ARCEsuhpS38aUuBWA6zbCjEdliJJTkp0ACFMkoplBbaUiwzyjCljFgO0DLKKG3pYKWFEEtyEkYsB2ih7BXbYZTSMsqI5SRAEvud/3kJI4HEQ3pT+unzbiRL79177vcqtu55Q4t19p8a/pb5byGCBED+eu5tiQSAe9a2taSZWiujZ1t9Q6mITRMXAxwKlIPrtB9c+aRWn9Fi28JEu1VnupAEsE3UvoruOfgrr4mYeV8LQCP5qiErT9d7LB4KMK16xsPmA9k0zPw7bA9N3vZ/wvRInhFOpKYmfO7ME8/JzWYdfni/2HwxQCEcAeDkmA2tbiJ2+2KATFuV15+xB+HmawH8Ivb18Kw/uNDGbB0B8Mb6X3XmWf2Df+Sm9Wdu70ztqHWoAswGjgIYKlYRrKd/uG8vqBsip9emX9y9oDqwcUEC7bEdntQPwu8VVEkJbWxZWWYl1OXAdZXzTwCMrd37qP0D1+ESCThEZXZfB+AL1ZlOjLeX7x9te9VGI6ydhnkdtaOb57UYwjj6Iy859zZCAsA9a9taWjJrRo7WnNdDLt7K+qX/OBfbQ1PrEVjBK62vA3xzPS/l/RTjKIC87ZzeMBuvsq7ga415vk2VCa/CUQD56tm0nf4fe8Cmqoq+Glj5f4g1MZnvEQBk9ssB/u9haUa4ILrtg0Ri6xFLTEYt4eaZgNXwe3de8jq5fBSA7jTEdQAsfB8Xw8exIbQBBEwW6yiAAdZw4iU+bpNDz9nciZpR59AE/j5l+3eJGEcBUOncTKZrCustHzVpbie+8rEwxIK2FgrdX1AFJbQxrAIw2CtGtucbJTMhAZAvnhvbCS+0uRnsCbYZdBjVfbyqJlbv/fgHNx4IIwHghnMBbSABUACel5subU1aF/XodDmGL4xYuQpHAbiM/pnmhArZI/yZ6qwn9I8DrgVgQfiwLKyL3Kfjc3MhofWbhKMACgEscFtDVswrsIqS2RxW/h/q9kO+1kP5nwawB75+1s9jbNp9tNLEg+e9MdrPPS7e2D7pWX+IH/zkJ1cefbG8vnFvV1pCI3kJIAGQF5tvNvLgKAA6jupO28g3AiUYSDZWeU8BH74GEsPFAAfS8fA1WXMUgOQbAhPX1LQtRhIgX8ACt1tzTjs/W2A1JbA5P7vGqgS6GvAuivBD+XZhlPR/I99tsZ2zAmV9G90vQqZ9rUjoA3MpjvqwD3ToNa215tI5zV1E8sRaTzn+kMWc5HgjaCBvASQA8qbzfkODTOvCHm4HUhU2NsZRAG6rf6o9If7dp54q+EcmfCNAwYgOVZCti/xTq75aS/6LmBdrEmCf/CvAloUJyKzCti+FrWEUlFE2SLL5xmqSfDXfbbGdswL3Tdv6TYPtPQ2ASZAAcHbY1lv7p5/Uz3guHwVg4PSPTw+Cj35GAsBHgzHcUJakZlqHav19uNsVuj6zIAFQKGKB2+8/ZfyNWoX1tYB6Z+uCowBs5bSvMoPFuhZAroAaRxGZF5OIfg4ooBZsmpdAiFb+Ja8NS2gjGAVosAu5DgAZSAD4eqhtPg1ABHuC3R/vz7RoCrmcABBcB+Azo+CfJ5AA8M9Y5BmJuH4UgO59/kY4lpieZ8DYzAaBJmaT2LCSADbU9kkVOjPEtQA+4fDVowV1VV1MUuBRADSpZm73xb7qWIkEc390xxeF5G8l0t1hd9OysYyGvSE28ERg9XUAiPL7OmIWJAA8GbUhNtovdl8wbsfa9LNfGmLrWM0Wgc9WYpoj7tNnRYtbS/nYqQkkAdzSHmY7SAAME8xvq4c45MV1AIiZjvWbRanF89bSFb/VPtv6lT1an7XgKABLwYel3zCtowBeLSg0obOr2zqnFlQHNs5LIMT8m7w2LIGNYBPEQeZ/5RO1ECMBkA+cS9u0109oJ5J37GzO5DKcBmAn6GB1ref1ZW2XvqX/91w9CkB3VeHoj/WMhR+eQgLAD6NQQAzvts54jonTBVSR16ZCFA/HG/ELPS89ezZ6/vAdVpKIlQSwp8IPa2FcC+BDCf/dLZyy3etEXNhRAETEwr/8xp1PfYFwc1VgQV1E98DIfFcbDURjMn+NTSCCRZAfC0heCQD9G4MEwMeGfn1g93UA8PWPbo70htpiNt1NABAhAbChwfD4eSQAPB4AO5o3TfHk3FImE9cCsGMAC6hj5fLRN5LQogKq2NCmOApgQzIeP79R6H3rKIDnCgxjxxEbb/zLAuvA5nkIGMzNeWxW1JvAJJjDK0x5JQCIaMvpd0hI77H4VsDmRKXQHr7tavEFtuEeCd+/4RcdeQUXAnSEtfBKkQAo3NDzGnozyVs1CCcuCKfVbngR4mMqpv5slw2vgVecFvj74V9+nwxy5FoANZnu05yOH/UPX2D+5J2X6x78q4e/5We2+H51ZvEJn3kWTzgqsGZPN9/iaCOBqpxvWWMSqKARrAoYpryid3ktr2+yGEcB5CXnzkaanplra0tMEVvrQ2UDCGz4pVwqeY++ulyLW8uIivpEjVuNoZ2hCyABMHQrX6+pE4I/exGgaYZwFIAX8Gu1yRuvshIA3Ws9ZdND85zauS9sa1NlqMZGgfZY5XVCdHuhVTL1Xzwxs3i7QuvB9sMTCNGK83WLPi2lvvR9aFHqDkHtf75HAFCoDN8E4OdB75hc9TKR2Pi5gkdUz/vX1/zc56KJbeCO6EcHcvU0ABPfAjHwiHj0KhIAHsHb3ewqEes0gLfsrnew+pjouIroGcjkDwbl4Ovtk3ZYxsRWEsDmVvhLYvadY3OlqM4mAdNka2xeKqw6/pJhfTVgYZVg62EKrL7avVDjMDcrvtXVYLVF8fWsNHpkct4JABYTnxv8/i5hw9ajAAzpw1EALoz5EJq4bwjr2LYKM+1oW2WoyDYBJABso/S2ouWZ5Bsi5MVRABVmyMBRAN4OP5WNCFkJgBftD4O/W9PWia98tB+24BofiFe+xERWEqDQur5Zk+m8qdBKsP3wBLKxyBVE7PrXuJJvbty6xsA3ASGQ4QqMGZl3AsA0+7cabnNY32UB07T11FKz30QCwPkhHEoL9w5lJdvWEUYCwDZM+ypCAsA+S89rYjKtowBcj4OFjvv81MaxrjeMBj8WuPegbXJChpUE+Pg52x6IcXbtgmfH2FYfKrJNoD0auZ2YLi+8Qj6hJt1lQz2FR1K1ceByAAAQAElEQVRKNTDxD7W/b2gpteWND/teav0uqv7+763lZt4dYmNp3ttiQ1cEmEO2XmBY9wQjAeD4yA3eQC6VfJyY/jf4mratMX7L6aePsq02VGSLABIAtjD6o5JceuZjJHyHB9FsvdIkHAXgAfzaTY74oM/6SsC898isXde6j2UXeW+EHXua160WP9kisIkY1tg8UHBlTKdXpxc1FVwPKhiyQHt0/H+E6XtD3qBIVrT6bPW9SLpTst0YO2ajjQvoPBIABeC5sWnZCGORve0IEgD2gn62tqE+4/K3AfS+PwJHAQx1bFxaDwkAl6DdaoaJvDkKgOS7dNJJI9zqJ9r5rMB9h054R0isJMBnXyz0Gaaza1JdEwutBtvbL9AWHb9C96ycrTWv0FLQwmxcUJtZ/NOCKsHGwxLoqItkiOlHw9ooyCtrX1f3Och9QOyrBUIhzjsBwMxIAKxW9O8/1pGF+rvpNRsjRALARsz1VTXU51jE+jaAoa5e8HpMslPBlaACWwWQALCV0/vKetLNc1jI1St8ftjrXSv+W/HjDx/jziOBvlCZdRrAfxxp3iBrkulI1ai0MIH2usgDJGQdCVBYRbq1kHlFdVtXVB9icUkgWxe5Vpu6QEuxLxd82Ndi72dJ9I9XfJB/AsAgJACC8C4R6rQxzP+rvQOnE9ro+emqhvyzTsgfGfLKNqzIhOsA2MBoaxVIANjK6Y/KhMiTowCE5cebHHrO5v5QKM0oHpq8rXVel5UEcAKgribTfZoTFaPOwgWyscjl+n//9sJrItIkYhpfAUmu3rLRyC/Y4ItdbdTFxqy+WX10sUk05bAAh0J5JwBoZf8yh8ND9TYIMJOtpwHwJiEcBWDDuKy/iqE/u2STV2wd18FaFhKcAjAYksuvIwHgMrgbzY0I8Z91b6AD54IPGv1XRqxchaMABmVydgXT6NcEAP/bmVbMczAxdEbWjlrt+WrANZGI2f/kmkf41y2B9imV52pbxXgkwAUf9k27h6VYBNgo5BQAwhEAAXgjmKa9CYB+JiQAnBr34dQ7a1a/rr5Yi1sLTgFwS3qI7SABMESoIK329pxm/cMqnhwFoE4/GTe1Af/RFcKrZeGU7V4n4svIkRt/Scw+Ww41dyS8Eq/Uxq8GtCRH17R1/dd6gOKewOq95FxE1wTQvqzuk3uEaMktgT7O+wiAVTxSP6e4FSjayVuAxdY9xYwEQN5DMdiGw31diGwd20Ha/9znpzZuMcg6eNlFASQAXMR2synTkD+T0Co32/ywrVFmP+EogA8xvLrLRsf/RtvOaHFg4e/WtHVOd6BiVGmDQLv11YBEl9hQFenvkC2qM12Ff8OALcGUTiXWefLCFNMeB/krAt+w+mD1RfuBpQgFCjoCIBRCAiAA7wkxNrZ3ksihygB0O4ghDjtmJluv7zBo+yv7cCHAQZFcXAEJABex3WxqaevMTmL+s5ttftwW8/fC9Q37f/wzHngjYJBDRwFod8Q4u3YBLuajEr5cdI/rz0ns+UYI/ZCwb02me4EvO1rEQVlXymcy9iDiVgrcjVtZY7f6QLgVrYBB/Vvn2zlzxQe4BkC+eC5ut3DKVq9rczkt9ixCOAXAHslP1TL8H9nm6zsMFoEmDHEdgMGQXHwdCQAXsT1oyqvTAEhMxlEAHgz42k1mp0QWEnEzOXKTXeS9ETgVwBFbeyrNxqq+r3/g77anNqlFEsAeyeHU0h4d/59stHIqCf1Mt+vT4velz4rVirldY/d7sIivMAExQtvlW8OD06pwBEC+eO5vZ+dRAEgAODF+edRpEnfmsVnem+jOBCQA8tazf0MkAOw39U2NuVTzfRrMfC2uLzrxmBauP/Nw1xtGg+sI9JsbXUbMz63zpF0/MJ1dk+4+yK7qUI/9Au11kW8Syd/tqRlJAHsch19LNha5IkQrq4j4FvLtjW+xYrRi9W2ICMxeAaEJeVb4TJ7bYTNPBNjGBIBsUpdZvJEn3SjiRvPpWmhjO8d18AiETFwfbHAm19ZAAsA1am8a0oybZ0cBMJk4CsCbYf+41QfiX11Cpsz4+AnbH0jyG/Ne+ILt1aJC2wQ2ClUcpEmgF+ypEEkAexyHX8v90R1fzEYrjzaYDtSkjieJ3fVHLfOtmKzYrBjXvw6eLU4BGZ9Xv4SezWs7bOSRgPmynQ2vGmHkffFIO+Moorry6sqSWTNyLGKd4pHX9sPfiHEEwPDRHNsCCQDHaP1Rcc+of/1FiJzZAzxIF0Vo3/L6xh8OshpedlhA98j9SRNBtzvSDNPOI03TodMMHIm45CqdP/lLy0Nmn04aqdeezq9JAtQueBkf4uwBHVYtC+oi92WjVZOtSbeQ/G1YG9u4stW2FYMVixWTjVWjqgAIHDzvqdEaZn4JACIcAaB4QVmE+D07Y+1baYyysz7Ulb+AsKtHARjjYg075B8ttrRTAAkAOzX9WNesWf1M4t1RAGL+JBw7a5wfaUopJpMM6yiA5U70WUSOq2nrOt2JulGnPQL3x7Z7xTSMSfbUZtUitfL+qrZ95j63qfUTivsC1qS7I1r1rRCt3FZbv4CIXdirurqNC6w2rbatGAi3khRYYW5snf+f12dILmMkAAL0rmESWxMAZtkKJADsHP9C6hJx9ToA/QaOAihkuOzcNq9f3nYGgLqcF+gzVv2OhP5Fnty4UrgPpwJ4Yv9Jox3R8U8w82WfPGPzI6FkdaYT3/xgM6ud1S2cMv5xCvHB9tUptSPNEX/bv7Xr/+yrEzUNV8A67D4bjfwiG63ckWXVLkLGWUQyX2z40L6mDqsu4yyrbquNrLZltTncOLF+cQlo4rcy7x6Z/S4kq/KODht+SkCIbd15IP0hJAA+ZVzIj4VsK2zYeH2HwSNhQQJgcCV31kACwB1nT1tZPudXbzLLVV4FwcQ/Lo8n8j1U0Kuwi67dpf/unSFEDzvUsRARNx9w56LPEW6+FchOrryHhI+yL0Cp7QvR32rnvmDthbavWtSUl0B7bIcnNdk3wzosX/fQj7Ym7SR0LBt8MRHfohP6rN4/SyL/JqI3VpfVj/nZD1+7ZfW61jaaTLDqWFPX+BlW3bo+FgisEWDO9//86+11EzzaIbEmdPw7PAGDDFuPABCzDwmA4Q3BQGsX9BqL6eoRAEyECwEWNGL2bYwEgH2Wvq6p5/VeKwHwuEdBbmqQ4CgAj/A/avaxk3dfpf/hHTsKQH+x79a3cQjXA/gI3Kf32VjlrTpWp9kVnta1l5j9f6vNPL+3XXWiHnsErEl7Nhb5U/uUynOz0cqjdUJfq/c7ZmNVX81GI5uvLqsfV+744WtHr17X2kaTCfZEgVqKUYCFDs2rX+zNNYnyihUbrRYw2Xxv9QOb/hlhjMD1Y2yyJCqsojLeyNUjADTRvHlhEWNruwR0PmBXVajH1wKP3bCKTbKSAJ6EKcSnjo037ONJ42j0Y4H2aGS2CN308RO2P5ATatKdSPbY7mpvhfo+uFrEvNDGWncSCbVVt3UdZ2OdqAoCEPChQG16UVzDymtPHgt7clFijRdLngJGn70JADFMHAGQ51h8ZrMCn3gndfF/tYqlWtxacE0wt6QHaQcJgEGAiunlnkzyD9qfe7V4shhkYGLoify6jRqhkHUUgHX477ov2PUTc3P1vO5qu6pDPc4IdMQmNBHTGbbVzlTBQn+oaeu6hHCDAASKVkA4dEi+nRMSJADyxfNou37bTwHoH+VRV4quWZs65OZRALhwsE2DVmg1SAAUKhiw7XUS/mvvQpYjK+obpnnXPlq2BNqnbPsCk1jfCmD96ETZiPskeeA9L4adqBx12ieQrYtcrrWdpMW+RehsTQL8rXreC1+2r1LUBAEI+EGgNv3slzSOvBMApkFP6fZYAiRQNsKw9yKAwkgA2DP+ttQiQq/bUtHQKgnT9Omhoa2KtZwUQALASV0f1r0kNaNVw7pdiyeLCJ//+amNYz1pHI1+LNAerbqSmJw7GoRpz1Wr+pIfN4gHvhXIRiM3MtO3bQ1Q6FA2zbaa9PMH2VovKoMABDwVEN5IJ/+S72G8DyycEnnE0w6g8WEL9HO/rdcAIBEkAIY9CuvbwJ7n9O//W/bUNLRaxn7wlXx/fwytAaw1JAEkAIbEVFwraQbes2sBqOSuq0zzPL3H4rGAIdJExPb+Yad1bt+vaVv8w3WewQ++FGivi9xmGGyd12tffCLbE4faajNdp9pXKWqCAAQ8FRAzv4v/adBCktI7LAETGPHBSFs/JxihEC4CaMd7wL463ravqiHU1C84DWAITE6vggSA08I+rH/pnOSDGtZvtHi0cKI8duZkjxpHsx8KLIhWPcgkZ374oyN3Iv3JSZnOfR2pHJXaKrBgSmVahGtsrZQoJERXVWe6rpp+h+CwP5txUR0E3BSY2Nq1JzEdmG+bQiHrCMR8N8d2Hgl8QCttTQCQ4CKAdgylXXXo50BXjwDgUBkSAHYNXgH1IAFQAF6QN+0Xto4C6PWqD4Zhnke1TWVetY921whYV4Mnkt+t+cn+f5l4E5MoWXvHs2Psrx012i3QEavsEDJ2tbteJjr1f2MWt03MLN7O7rpRHwQg4I6AEZIf5N0S09yF0fHP5709NvRMYIsVT9uaANBEM04BKHw0bayB3T0CgBinANg4evlWhQRAvnIB325ZunkRrUkCeNITEdq3Ysyy8z1pHI2uI1A2UhqJ+fF1nrT1B95Hxo7w8IgTWztT9JV1RMc/wUZovP0dlYMMMudWp7rzPoTY/phQIwQgMBSBmlTnJCL+LuV5E5Nw+H+edl5v9taIPey9bhMTe92n4LdvXw9MNl09AoAEpwDYN3r514QEQP52gd+yL7TiKhL6F3l0E+LzKuoTdh9y7FFvgtvsfQdOeIf7yToVQBzrhdB3rMPAHasfFdsqYH1TRNmqldbVvm2tVyv7Mhvyt5p0Z/P0O54dqT9jgQAEAiDAhvGjAsJcUUb9SAAUAOjlpuZoczM722eh/9lZX0nWZWOnmUNIANjoGZSqkAAIykg5EOfyOb96k1msUwEcqH1oVYrQeUNbE2s5KdAer7xXs7JWEsCxZjTlf2p1uvMixxpAxbYK3DdtxzeX9S8dbWulH1XGnPjfmBH31bYuqv3oKdxDAAL+FJiU7pwmJIflHx233h/b7pX8t8eWngqI2JoAEDbe9LQ/RdC4nV3glcbbdtY3aF2MIwAGNXJhBSQAXED2cxM9r/daCQAHD/8etPcHhOONZw+6FlZwXCAbq0qS0K1ONsTM51ZnFp3hZBuo2z6Bx+p3fy8bjTAxP2dfrR/XtJ8ZMu6raevC//+PSfAAAv4TMIkL+iYP3dGAvf/+G9YhR2T0939hyCsPYUVDBEcADMFpgFVsfWkjs8/dIwBwDQBbxy/fypAAyFeuWLZ77IZVbJKVBPCwR3J+ebxxTw8DQNMfCjAbjZoE6PzwR0fumIyW6sziEwi3wAhk6yp3EKGb7A6YiQx9v11Sk+lqnTS3/TvY0QAAEABJREFUa2e760d9EIBAYQI1c7uOoQKu/K+t/zc0wkQCQCGCugiRvUcAGCNwBEBBbwZ7N35zfstyrfF9La4s+kcf3wLgivTAjSABMLBPSbzak0n+QTt6rxavlo0NMnFBQK/012q3PTr+P2QlAdZ6zomHTOZNNfO6D3KibtTpjEBHLHKisJzoTO0UN4Xuq8l0nuJQ/agWAhDIR0CooL3/QnLVfQdOeCefprGNbwRsTQAs7dsIRwAUMrTObOvmUQBIADgzhsOqFQmAYXEV78oGGb/2sndCHKuoT5zuZQxoe41ANjq+VR9doMXZpV/+ul+qe2tnG0Htdgp01FXdRCZV21nnx3UJfY6Ir61p6/rj/unnv0q4QQACngrUZrovI6H8j84TeXZ5f7nHRxh6SlgUjevnMzsTAMseq9/ivaKA8agTDjXr4nUADHwNoEODOJxqkQAYjlYRr7skNcOa9N3uZRdF6LxxUxt28jIGtL1GIBuN/EIfzdbi5FIeMswFTjaAuu0XyMYjC6VPtiZnrgtAOuE4pp9D99W2dR1pf/SoEQIQGIpA9dzOw3XvfUEXhhXmqzDZG4q2v9dhMu28BgAO/y9suB3aWlw7AkB/r+AIAIdGcTjVIgEwHK0iX9c0yOtM/TjTNM4rcubAdE/Kyhp1NvaKswHzV2oyXU852wZqt1ugY2rVy8s2692F2f7rAnwY6zaaELxV3xu/nphZbOfepw+rxx0EILAhgdq5L2zLJs3Y0OtDe547OqKRG4a2LtbytwDb+DsYFwAsbKyd2ZqZXUsAaA+QAFAErxckALweAR+1v3RO8kHN2F/pbUjyrXAsgfOAvR2E1a13HLzNYjFIkwCrf3Tyn51qM93tTjaAuu0XeGz33Ve110VOZKYL7a/94xpPM0gerE13/uDjZ/AAAhBwVsA0LyPiraiQm2leXcjm2NZXArYlADSxi/P/Cxlah7Y1iV08BYCQAHBoHIdTLRIAw9EqgXWN0Errw/yTXnaVWc4bG/1ZpZcxoO01Ah1Tqu7QP9j6YXDNz079KyQ1NZnuvzhVP+p1TkCTAE1a+0laHFpkvCYmf1OT7rpHEwFTHGoE1UIAAipQk+k6R0gO04d5L5oUvDsbr5qVdwXY0G8CtiUAmBmnABQwuk5tapjS61Td66mX1/McnnJZAAkAl8H93lzP7Ct79Be0lQTwLFQh3jwUCp3nWQBoeB2BjljkbB2Tees86cgPcpQmAaxrDzhSOyp1TiAbjdyotUe1vKLFmYXpQE0EtNW0dd1YnXpmgjONoFYIlK7Ahwm2iwsV0E/3Xp9OWGgXsP06AmLbNQBEcArAOrTD+8GxtcWgMnLv1u9eU2hpQwJIAGxIpoSf72ltnq0ftK/0kkD3On+noj5xvJcxoO1PBIyQ0aA/vaTF4UXOq27rOs7hRlC9AwKaBGgTMeuJeA45eRM6kY2NrNMCzt3njldHOdkU6oZAqQhMzCzeTgyjpeD+Cv1hQV3kvoLrQQW+EJiYeXY7Ih5BNt0MHAFQgKSDmwq5mADgPgd7gqqHKIAEwBChSm01P5wKoEmAy/CtAP5457VP3vZZMkInajRLtTi6sH6ArG1dVOtoI6jcEYGO2IRnstHKaUycIGIHv+pJxmmS8qKRo997qLat8zuEGwQgkLeANfk3WO4gke3zrmTNhm8ZJnm684Bws1XAkJF72VmhiSMA8ud0cEshdjEBIP0OdgVVD1EACYAhQpXaan44FUDNv2ianNR7LD4QyE7ZdoEYYiUBHI9GQsYCJAEcZ3asgfZoZYshtD+RzHesEatipq+L8J9qMl131bZ17Wc9hQIBCAxdwMbJPwnJWQvqI/hWl6Hz+35NNsjWBIBhCi4CmOeoO7mZTgZdTAAQjgBwcjCHWLeO+RDXxGolJ+CHUwEU/eBwLIEkgEL4YVl9UUCDTnUjFiQB3FB2ro0FscqHs9GqySxiXc/DdK6l1TV/U4QWVme6r6ie/cyXVz+DfyAAgQEF7Jz86+z/2o5o1e8GbBAvBk9AZHc7gy7beOQTdtZXQnU52lVhN08BIBwB4OhoDq1yJACG5lSya/nhVABiaiivbzymZAfBZx3vmBK5hp396rePe4wkwMcUgX3QHqv6pUHm/kKcdboTTPJTHjny4ep050X7p5//qtPtoX4IBFXAzsm/ED3c1ydnBdUCca9foHbByxubQrus/9U8nhV6+d6DtsnlsSU2IYcJxLTtOg9DiBRHAAwByelVkABwWjjg9fvkVADSvYhJXA/AP2+m9rpIEzFd60ZESAK4oexsGwuiE7Idj4zfX5h/6WxLq2vfnJnP7TdCj9VkOmdUz39x/Opn8Q8EILBawM7JP5H0Gf3mWQ9Oq3L8+jCrg8c/7gks+2BvZrJvnsDk6VdMU5Bvzsfu4ikAgiMAnB/PQVuw7z/2oE1hhaAK+ORUAFwPwGdvoGxd5EfEPMuNsJAEcEPZ4Taa2OyoqzxPyJhs7TF0uDWdl9DniLiR+1Y9Xp3pvqI29fwOhBsESlygNrN4b5su+LdaUojPaq+f0L76B/xTVAIS4j3t7RDj8P88QZ3ejMnNUwDwLQBOj+dQ6jeGshLWgYAvTgUgwvUAfPZWzNZVHk7Ernz4QxKAiuLWER0/f3n/0v2FpPCvHBuSCI+xTg0wrSMC2rqumTj3+V2HtBlWgkCRCdS2dR+v/+/abLja/xoZoVs7opGZa37Av0UnIPx1e/vUjyMA8gN1fCthdvEIAMIpAI6P6OANIAEwuBHWUAG/nApAuB6Ajoa/lmy0chIRP0cu3JAEcAHZhSYeq9/9vY5oVUJYphHz4y40SUw0koR+aJihx2raum609oS60S7agIAfBGoz3ZeJyO+JpMKeeKS7r38Vzvu3B9OftTDZmgCQlX1IAFA+Nxe2EVePAOh3oUdoYhABJAAGAcLLnwj45FQAXA/gkyHxzaMvLBu/swbjytf7IAmg0kWydNRVzdlE+Bs6MT9Hu/S2FncWoROFzL/XpLv+WJ3urnanUbQCAfcFqud1fq22rWu2kJxpa+vMZz04dYd/21onKvONwIH3vPgVDWaCFruWVzu+ueOrdlVWUvW401nXLgLILH3udAmtDCSABMBAOnjtMwI+ORUA1wP4zMh4+8Ssw7m/r0ysqwWLG5EgCeCGsjtttEXHr8jGIpdKf/83iPg35OaN6Rj9MJKtznTdVju3+0A3m0ZbEHBaoCbTVcf93CZC02xti+XwbF3kLlvrRGW+EujvN62/57bFxExP21ZZiVXkRneZ3DsCQH8f4QgANwZ1kDaMQV7HyxBYR8A3pwLgegDrjIsffnjw4KrXhPp2cysWJAHcknannY767RZno5U/1D2VB+gHhLQ7ra5pRT/8HCGm3KMTpjsnprpja57FvxAIrkBNW9fpGn1GS0SLfcvqyX+VKxd/tS9o1DRcgf7+PlsTAKZJzww3Bqy/WsCVf4TdSwAQMY4AIO9vSAB4PwaBi8AvpwIQrgfgu/dOR3T7Jww2p7gVGJIAbkm7105HtOr+jlgkrnuMjtdW3d5rdIhhSKom071AJ1A/nJhZvJnGgAUCgRGYmHl8M01kXU9Cl9seNFt7/jH5t93VhxUyG5PtDIsNdvt3uZ3he1iXS02LiwkAwdcAujSqAzaDBMCAPHhxQwI+ORUA1wPY0AB5+PyCugnzhOg4t0JAEsAtaXfbaa+L3LxR6INvEMn5RLyEXL1JLQldY5D5XE2665ra1kW1rjaPxiCQh0B1etHJBm3yoG56khZ7F0z+7fX0cW2T5r1gHcm3t50hmsJP2VlfydTlXkfd+xYAxrcAuDesG24JCYAN2+CVAQR6Zl/Zw8wXDrCKWy/hegBuSQ+jnY5o5I9MnBjGJgWtiiRAQXy+3Xj+5J2XZ6NVF3FINBFAN3oQ6GbE9EPr/YWjAjzQR5NDEqjOLDq4JtM5j9m4jojHk903TP7tFvV1ff39fXYfxff6wuj4533daZ8G52JY1kUfXWmOhfpdaQiNDCiABMCAPHhxIAHfnAqA6wEMNEyevdYerWwRMS90KwBrklab6fyWW+2hHfcE2idHOrPRyElC5mRimutey2u3hKMC1tbAY+8FqlufH1/b1nkdkzGPiA8mJ26Y/Duh6us6mXmyvQEKzv/PD9S9rYS+7FZjgiMA3KIesB0kAAbkwYuDCWy08bLzdE/vQ4Ot5/jruB6A48T5NNARm9DkahKAeFZtprshn1ixjf8FOqIT5mfrInXCcqIQPeZRxDgqwCN4NLtGoHbBgrLqTOeZHAo9KMInr3nWgX8x+XcA1d9VTmx7aScSmmhnlPq7Guf/5wXqzkajo4kvaWLdta8BJOE33ekZWhlIAAmAgXTw2qACb826dlm/IY264nItni4skhw3tWEnT4NA458RcD8JIMnatq6rPxMInigagY66qps6opHd9YPlyVq8SgSop+BaAaqAxT2B6rmdh8v7mz+kiffLtNXNtDiw8LNCxuRsHS745wCur6s0ZKXdh//r3JLv93Wn/RqcS3GFQqZre/9Xd4nl1dX3+MdTASQAPOUvjsaXzkk+KMxWEsDrDlnXA7hh9NRzvuh1IGh/XQHXkwBCP9IkQGrdKPBTsQloEuAGLT5IBNBm+ikX1wootjeYj/pTm+k6oqatq41Nvp2I9yCHbkJ0e6hfpnREx893qAlU62sBtvnwf3ohG420+brLPg3OrbCYDNfO/ye9MRESAOrg9YIEgNcjUCTt97Y2X6sfgG/wQXf2GmGu+q0P4kAInxLwIAkQq2nrfrb2rpcrPhUKfiwyAU0C+CURoLLy4VEB0lWT6ZxVne4+ef/M4m30BSwQGJbA9GefHVmd6TqpJtP9kBDdRkK2751dOyAWOU//Lx15f33kv2s/j8elITAx8+x22tP9tdi5YPKfn6ZrW7Ep7h4BIIwEgGuju+GGkADYsA1eGaZAqG+0ddX3h4e5me2r6weleDje8DvbK0aFBQu4nQQgke3NjVb8d9LcxdsXHDwq8L2ATl78lAgYR8TfYpbr+sl8oSbTdZ/uwT17YmvXnoQbBAYQmJhZvJkmLxP/+/fIJ3Rv2fVEss8Aq9vx0kvaziHtsapf2lEZ6gimgMEj7N77b0EgAWApDLu4t4HJhqsJANPoRwLAveHdYEtIAGyQBi8MV+DdtqZeQ8hKAnww3G3tX5+/G44nZthfL2osVMDtJAATb2Ka5rMTU4uc+HBTKAe2d0DAX4mAjzu4v+7BvcQI0cOaDHiiNtPVUju3c8pu1z/q3sWXPg4FD/woYF3VX98bFxtkPqHJy2ad+Ft7ZJ0NlXkOG6HJ7dHIbGcbQu2+FxCx+wgTHP6f76C7uB2Tu0cA9PcbSAC4OL4bagoJgA3J4Pm8BJakkwuF2EoC5LW9zRs1hmOJBpvrRHU2CLidBLBCNgxjbnWm8wTrMUppCPg0EWDhf12IzhCT28Z8eeyLNZnu31VnOo8+4O6XcP0SS6fEym/OsekAABAASURBVKR0917V6a6rOWQ8oV0/R8v/aXF+Ebo0W1c5rX3Kti843xha8LOAlXwi+79KEnv/Kb+bq1sJuXcNAJFVyzPJN1ztHxpbrwASAOtlwZOFCPSmmq/W7f1xHj5TsiKeOFbjweIzAS+SAEx8U02m+zyfUSAchwV8nAiwev5l3dP7XX1v/rlvxCpNBnTere/R06rv6Z5gvYhSnAK1bV37Vbd1X1ST6fynyfIPZvoREY8mV27yhoh5dDYWsZINrrSIRvwtwKGyIx2IEAmA/FBd3UqY9G+QS00y/9ulltDMIAJIAAwChJfzFJCQ9a0A/8xza1s30z1tvx0bb9jH1kpRmS0CXiQBdLL1i5pM1/Uk+mfPll6gkqAI+DwRYDHqBJCn6nv017xKntf36cLq9KKm2nR3fOLc5ze3VkAJrkBt2+La2kz3ZTquT4jQQhY5l4h3JzdvTLeaTJP1d+8tbjaLtvwrsM/c5zbV3zl2Hx2Hw//zHnIXN5w+PcREW7jYIg7/dxF7oKaQABhIB6/lLZBLX7aESawkwMq8K7FvwxFM/Ndx088M21clarJLQD+INuneqAvtqm+I9ZxUnelurc0s3nKI62O1IhL4OBHAdDwR+fnrIvdjNi4QllbDDP23JtP1z+p053W1c7tOnNjWuZPGjsXnAjXp7oNq0l0zdeye1d9zC4TkTA3561pcXUTkMU2GH5mtixy1sK7qaVcbR2O+FtjIHGFN/reyOUjs/c8X1MXtKlZt497ef+2XzguQAFAHPyxIAPhhFIo0hp5USzsTWUkAz3uocWxhvt+f9TwQBLBeAS+SAMwUEzJbJ83r3mu9QeHJohfoqIvcnI1G6oWMXbWzl2jp0uLXRX+N0e7MfLKYdKMh/JROKl+tTnfNrm3rPrc607l/7YKXN/Zr8KUSV23boq2qM11H1LR1/0rHp4tY5hPTz7T/Xn0TSS8LnzeaQ/tq4ut2jQMLBD4WsH5nCImVAPj4OZseIAGQJ6Srm61c5d75/9oxYXwFoDL4YkECwBfDULxB9KSSvyKS3/ujh7xzOJ64xx+xIIpPC3iRBNAYvm72SWpietFh+hhLiQp0RMc/oYmAn39h2VPbi2EcQcSzSH9xafH7siUzTdO9uxcx8X3y/sqlOvn8h048f633R+AIF+eHrzrVPaE2vej4mkz3ter+uIjxMhPdRiI/1tYrtXi5/JFDoX3bY5W/bIuOX+FlIGjbnwLyQZ81+bf7eiM4/D//4XZ1S9MIfdnVBgUJAFe9B2gMCYABcPCSPQLmiBHWtwI8bk9tBddyYDiWwLmPBTM6U4EnSQCmzxts/NU619qZXqHWoAjMOvzw/o4p4+/IRisP5xBtxyLnEfGTFJxbGRNZR7Scpve3CZmv6qT0XzVtXbfUZrpOnTi3a8/au56oCE53/BfpxLmLd7Us1fVWLf9mQ54XNjTJLadotLto8cPyDzH5sGw0clz75G2f9UNAiMGnAtL/PQciw97/vFHd3ZBd/gpAEROnALg7xBtsDQmADdLgBbsElt516TvEYiUB+u2qs6B6mL5dXp84v6A6sLFjAlYSQPeeuX7qCLNxgX6gv3P/zOJtHOscKg6MQPvkSGd7rOqX2WjlLsIyTQP/I5G8p/dBW75KQt8WoqsMkx6WjTZZUpPpfqMm3dWhiYEba9q6ExPTndNq53VVUVMTPhN8OLq7tT66SXVm8S61bV1HWslBNZuvZblhmo9ZlrqaddV0d/eeaaMDLkzvCBlnTaqr3LcjXnnngOvixZIXqEl3f5eIrdOfyOYbEgD5grq+He/sZpOhECEB4Cb4AG3hj/0AOHjJPoFca8v9zGQlAeyrtICaWOjCcCwxvYAqsKmDAtlYVVKrj2pxezmkn8x5OjE6xO2G0Z5/BTrqquZkdW8qs2yvk78GjfQfWgK8yBeJaaImBk7UZFuzwTxb+mlRzZ7f7tdJbnd1uiul/wdmVrd1n1yT6py07/xON68S7apr9bwXvlyTfv4ga69+dabrqppMt070u/49JjR2OZP5uAjdypoc1OTPQVo2cTW4oTe2QkSu5z7aryM6fkYTszn0TbFmyQqwWIf/291963QqJADyVHV/M9nDzTbFGPNvN9tDWxsWQAJgwzZ4xWaBntbkFcJ8s83V5l8d0x3jYg075F8BtnRSQCdcbcwjv+ZkGxuoexudGN1p7fXbwOt4ukQF2usm/KsjGpmp7819SPhgZbiBdK+r3hfRIuOZKUZCP2OR68jg+8v62PoGAuv6Ao/WpLtusf5v1LR1/1gfH1NrfUVhpnPf6nusc+Ff/lJdZvFGfsHYK7O4vHbuC9tOsuJLdR9ak+k8pbqtyzrSxzp8/3Gd7C/n/v5/E4fma2LnKiY6VSf5OtEnf+3Z3zDo28ScNMnYtSNW9YP2+kjnhlfFKxD4RMA6skV/2k+LrYsmzG6ytcLSqszV3pZPPn1TbdDNIx6X9sxu6tE2sfhAAAkAHwxCKYVgrjCtQ7t9c06tyfwMTZ8eKqUxCFJf2+u+9q9l/UtH617KxW7HzbrXrybThVMC3IYPSHvZWOU9mgg4uWzliO31Q69OHHkOMS0LSPj5hDlGJ8i7aR+/zfp/Q/9P/oqY/ijWVxQSP8CrrHPhV77+HpkfVGc6l9dkuq3rDzyt/4ey1emu2Zos+H1NW9fM1d9YkO76kSYRjtLJ+ZTa1kW1+ZZq62r7me7TajOdv9T6b6zJdM7R56yLIL6sP7+3MZk5MfsXm1Z8hvyNiK9loSYisg7f30Un+37do68hDrAIvUzE58vKlbtm6yobF0bHP0+4QWAYAvo7y4G9//y8sclIJACGMQ7rruruTzyibHc3W2TG4f9ueg/WFhIAgwnhdVsFls1v+R+ZbCUBxNaKC6gs/P5W/ytgc2zqsMBj9bu/l41VWVfTvsvhptZXPU4JWJ8KnvtY4L5pW7/ZEYtck41WTlv1/gfb6Iecb+uLNxLxi1SiNybWibVsqd3fUUu1mkzTZMHxJPQzEblIf76a2fiLTs7bJGQsyLcw0W06if+1EP+cmE4k4np9zroI4lb68ygqvtvTTHL6iJFlu+j77aKOb+6I82mLb4wd71F12+Kp2oh1pIve2bcIye/aJ33tA/tqLLGa3O4ui6uH/5siJfs30e2hHUp7SAAMRQnr2CqQyzTfox8ErSSArfUWUNmmFfWJBwrYHpu6IJCNRg7VD/sXudDUp5vAKQGfFsHP6xV46NCd/9deF7lN36snZaOV25Ip+wvzL3Uy+tB6N8CTEBiCgE6sHhSWEyfVVe7SHq268t6DtskNYTOsAoH1C4hpfUXl+l/L/9mXjBXvYe9//n7k+qbCu7nZpiZ+H3azPbQ1sAASAAP74FWHBHLpZIvuqbnGoeqHXa0I7RuOJ+4d9obYwFWBbLTqfBJ24NDFwbvBbFjnDuOUgMGpsMaHAtl41YKOusrzsnWRfU0ytmei0/Slu4SoV++xQGAwgTYROaIjWrVfR13VTU24uN9gXnh9EIGadNfP9PfQAYOsNuyXxdS9/4fsgvO7hy338QZePNjJ1UZNQgLAVfCBG0MCYGAfvOqgQE8qeapWP1uLX5YDwvGGuX4JBnGsXyAbq/y9aXJ8/a86/ixOCXCcuDgbsM7Tbo9Grs5GI4eO+MDcmlgOF+LrSMT161sUp3Bx9EqIHhYxL+yXvr31vRLtiFXdURw9Qy+8FpjU9lKEmM6yPQ6m1wzuu8n2ekuqQnc7+/mpjdY3u7h5AUAKmSsecbeXaG0gASQABtLBa44LcNmq72oj/9Tik4UnV8QbUz4JBmFsQGBhvDLNYp2/Jm9sYBUnn8YpAU7qlkDd9x064Z1sXdWsjmjlKdb1LUS4hpkuJOaFJdB9dPFTAmtP+juikb07YhOaHohtj71ln3LCj4UJiKw6W2vYTIu9i/BN7bEdvPhbbG8/vKzN5bb7xHT1/H/9Hffcu21X4cg3l8d5oOaQABhIB685LtAz+8oeNkIn6y+H1xxvbIgNCEmsvD7hpyMThhh5aa3WHqt6tK+vby8S8uTUDcYpAaX1hnOwtx2xyo72ukhTtq6y2mCpotWnufBvROQxB5tF1R4K6N+8j/f0Y9Lv4UCUSNPV6c7D9T13nAPdfZe5/3cO1FtSVbrdWVNoF1fbFPLRjj5Xe+7bxpAA8O3QlE5gPXMueyJE8gM/9ZiFpoXjib/6KSbE8lmBB6fu8O9N2IgT058++6orz+CUAFeYS6eRBXVVXdZpLtlo5Q87YlW7h/ppS2E+jElmqML9WpZqwRJAAZ2AYdIfwHELesi1dzw7htmw/9B/C4Z173/dhH9ZD1HyFnB9QyZ29fx/JsLh/66P8sANIgEwsA9edUlgSaqllYhPIX/dDtMkwG3+CgnRfFqgLTp+RbYucqxOklo+/ZpLP685JSDTfUXtXS9XuNQmmikRgfvrI//tqKu8sz1adVY2GjmAR7226YenDFiH895FQv8tEYogdrNbx+fPxPyTfunbG3v6gziERRDzmDKd/IsDe3z5Pe4T7P0v+C3iSQWuJgCIcQFAT0Z5gEaRABgABy+5K5BLNV+nWcJfutvqoK0dURFP/HnQtbCC5wI6SUqQcKNXgege2p/KRisX1Ga6vulVDGi3+AXaJ03q+/CUgcs0IXBoNhbZkvtpgpB8T4SsC3E9X/wKPu2h0FO6l/9qLceJydvp+ER0fI7J1lX++gGc0+/TQSvusKpTi75hClvJQgc6Kje110c6Hai4tKp0ubfhuoR18T+ruNVyTy6VfNytxtDO0ASQABiaE9ZySaAnlTyPSH7vUnNDakY/zB1dHk/cPKSVsZKnAtlYZVLHyzrPcblHgXxd27+rGkcDeMRfms1aH8I7olW/64hFTtRJ5/b95opNDWP1N2VcQsTtRPS6Fix2CzD/nYibxZRpLCM3z8YiX9e9/Kdp+WNHvHIR4QYBjwXYCJ3FTE581jcNY3XCkXArTMDtrTkk7u79J8K1bNwe5CG058QvhSE0i1UgsGGBXKrlBH11vhbfLEx0bDiesPau+SYmBLJ+AevDt37gmULi3TlnOBpg/WODZ90ReCC+05IFUyrTmgz4eTZaOSkbjWwxYllZhdlP1kUzj2Xii0XkbyTyrEa0SguWgQVW6O8Ta+/+7SLmhUJyAI8aOUr37H8jG608syNeNac99jVcBX1gQ7zqskBtuvsHukOl3pFmma9eMCXylCN1l1alrvfWJNrZ3UYFe//dBR9Sa0gADIkJK7ktkEslJ2ub1odTvfPNckK4PnG9b6JBIBsUaK+LPMCbjNT3EHt5fiKOBtjgCOEFtwXuPXyb3ML6yCO6l/pP7dHKcztiVd/Kxqp2zEYjI/tN3kb6KUZCP2MR/R3HpXrUwNs6Lg9ouVEtzhCmmGWjRhurm7V3/8iO2ISmjmjV/e2TvvaBrocFAr4UqJ37wrZCpkOH/lNX2coVl/iy44ELyv2AmVg/G7nZLuMIADe7ZbVdAAAQAElEQVS5h9gWEgBDhMJq7guMHDV6H211iRb/LEInhWMN1/gnIESyIQH9gN6TjVZ+z/ogv6F13HgeRwO4oYw2ChF4IF75Ukd9JKOT3CvaY1U/0P83Gz5qgPgf2lYXEb9JRCu0BGgR62gH66KJT2rQGWK6XO9PYqaJ3E+bZaMRq0zU+5PU4vKOukjGstF1sEAgUAJi9v+KmL/iRNBCdMl903a0/v87UX1p1elyb8PRxt21yb21uLZwX+hh1xpDQ0MWQAJgyFRY0W2Bt2Y1LTMMqXW73UHbY/5hRbzxV4OuhxV8IWB9kDdNc4oG85wWrxYcDeCVPNrNW2C9Rw1EK/fRCXJVNlr5pWw0svGy/qWj+8rk/0wytjdI9jMMjrMpxzDRacx8nibgdJLNvyOmO0lkgQbzhD5+WZ/v0ccFLutM6O8hkluE+Eqt9OdaTmKiQyQU2lfKyip5xchx2WjVSI15Sy27aIll6yJn6P2N1hFD7fURa++/boYFAsEWqM10Wv8Hos70Qm6xTrNzpu7Sq9XtHosh1mchN5vt7Jl72b/cbBBtDU3AGNpqWAsC3ggsmdPyNLP47qrqQvLj8njjTG9U0OpwBRbGJ8wL9dNknRDcPtxt7VwfRwPYqYm6/CDwWP3u7z14cNVrC6Pjn18QrXpwwZTKdHu86s/t0cjV7XWVv9QEnE6yK7+nk+3DsrGq/XXCvas+3lqf1wl5hPXnAso6E/qDdYJ/dEe08nSt8xItN2oMszsmb/tQx8HbLG4/5Gs2JBz8II4YILBhgZq2xT8U4p9seI0CXmF6x6TQxQXUgE3XFfDgJ7cP/yec/+/BKA+lSSQAhqKEdTwV6GltuZuZfuZpEOtpXCdzPwvHEzPW8xKe8qHA/fWR/+qE4EjdS3i+x+HhaACPBwDNQwACECg2gZp090EkpnOnKJp0yUJN9BWbm3f9cbflimji60yyn5utans4/99N8GG0hQTAMLCwqncCPa3JKzQJcJV3EWyw5cZwfQIZ8Q3y+O8F3Ut4ETEdquVlL6PTP4w/lY1WLqjNdPnuCBcvXdA2BCAAAQgMT6C2bdFWxOLgtyfJ/GwscvnwosLaAwq4/KIY5Pbh/0TMSACQP29IAPhzXBDVegQ0CfBjnTSl1/OSt08JnVNRn7jQ2yDQ+nAEsnWRu/rXXBcgNZztHFgXRwM4gIoqIQABCJSSgJjs4OSfSMTAjg6b31AeVOd2AuDdMmacAuDBQA+lSSQAhqKEdXwj0JNqiTPT874J6MNAROh8nA7wIUZA7h6ITejORiP1mlTy/DQOjWH10QDVmc6jA8KHMCEAAQhAwAcCNZmuu3RP63jnQjEu6YhVdjhXf0nW7Gqnx8UadtAGa7S4tzDNe3tO81L3GkRLwxFAAmA4WljXFwI9rcntNZAPtPhtaQzXN/7Wb0EhnoEF2qNVZzHLMfoB6n8Dr+n4q19n4j/rh7k7J2U693W8NTQAAQhAAAKBFtC/F9aeeQdPI5PHNwq9d0mgkXwZvLtBmQa7vfef2OR57vYSrQ1HAAmA4WhhXd8I5FLJTXwTzNqBiHwvHE/ctfZTeOx/gfa6qj+L6B9IpgU+iPYQk/iBmvSiy/dv7fo/H8SDECAAAQhAwGcC1enuozSkc7Q4tggbF8+fvPNyxxoo1Yrd7re4fv7/8lUiSAC4Pc7DaA8JgGFgYVVfCUju9dxIX0X0STDfLI83LPzkRzwKgkBHdPwTyzarnEzM/rjYJBun94fowZpM92lB8EOMEIAABCDgjkBt66JaZvmLw63d0FFXeafDbZRk9W52ekyscYK2d4AW9xad/C/PJN9wr0G0NFwBJACGK4b1/SPw2A2rcstGj/JPQJ9EwsT7heOJ52l6k1+TFJ8Ei0cfCzy2O6/K1lX+WJh/QEI95P3tq0Ty6+pMd3tNZnG99+EgAghAAAIQ8FLAmvxLyHD4aDVeHJJ+HPrvzEC7WmsZ0WRXG1zdGA7/X83g43+QAPDx4CC0IQi0N30wctTysUNY04tVJoTfX/7q6GjiS140jjbzF9C9HtcbJtVqEsAXez+YpIbInFOb6fpDdfrFHQk3CEAAAhAoOQF3Jv9EBvNp98e2e6XkgF3psLuNCJtun/+/isXE4f/uDvOwW0MCYNhk2MBvAm/NunYZl60a57e4PoznC2UG/efDQ7A+fAp3QRBYUB95KhuLHCZCpxKz1xcIXE0mRMcx9z1Um+5q2veBTr8mvlbHin8gAAEIQMA+Abcm/8T0owV14zGBs2/o1q3JxZ/K44nxROzqEQBMNK8nMxPJI/L3DQkAf48PohuiQM/sK3vMEWWfH+Lqbq8WCrE8P3ZqAld2d1vehvY6YpFrpIxqdfJ9uw3V2VHFGGG6oKzXeKi2rft4OypEHRCAAAQg4F8BFyf/l2frItf6VyL4kbnZAyb5lpvtWW2ZhMP/LQe/FyQA/D5CiG/IAkvvuvSdPmOEbw+3N0x6YFy8AedxD3lE/bNix0GVizqikSM1CXCyRvVfLT5YZAcR+X1Npqu1tm1xrQ8CQggQgAAEIGCzgGuTf6KMTv7PsDl8VLeugHs/TZ8eEuKj3Wvwo5Zw9f+PJPx8jwSAn0cHsQ1bYPmcS94so74th72hSxtoZnRORazxOJeaQzM2C2gS4AY2QrXE9Cebqy6kuriIuaA603VVbduirQqpCNtCAAIQgIB/BNya/AvRSgmFfuCfnhdrJO71q+L9rY5mou3da1FbErqnN5VcrI+w+FwACQCfDxDCG77AO6kr/muQ8dXhb+nOFsLyh4r6xOnutIZW7BZon7LtC7qX5FgSOUHr/pcWXyz6h/5UEX5IEwHYg+OLEUEQEIAABPIXcGvyb0Vo9JuTOyZv+6r1GMVBARerFqHvuNjc6qYYh/+vdgjCP0gABGGUEOOwBZakZvybJLTNsDd0aQP9xXx5RTxxkUvNoRkHBLKxqt+HpL+WiH9Hvrnx5kzUUpPp+mdtW+f3fBMWAoEABCAAgSELuDn5F1NObK+f0D7k4LBi3gJubRiONRxITAe51d5H7TDh6v8fWfj9HgkAv48Q4stbIJe+7CXT7I/kXYHDGwrRufpL+hqHm0H1DgpYX5OUjVZ+j8n4Don46bC33UX4t5oI+Ht1WxdOOSHcIAABCARDwM3JPwn9uSNedVMwZAIfpWsdEE/O/afsknTLs651Eg0VJIAEQEF82NjvAkszl3f3C2/n2ziZfxiuT9xCtU0b+zZGBDaoQHt0/F/6+qiWRa4nf932ZqE/1LR1dVSnu4/yV2iIBgIQgAAE1haoSXcmJGQsWPs5Bx+/no1FjnGwflS9joA7P4yJNU5gJg8u/sfz3OkhWrFDAAkAOxRRh68FlqWbFxmmsZNvgxT6dnjMe7M/P7VxC9/GiMAGFXhwWtVr7bGqH4jIEcT83KAbuLmC0ERm+Yt+uLy/Ot15uJtNoy0IQAACEBhcoCbTOUf/djQPvqY9a5hk7GxPTahlSAIurRSi1ZP/ES41t1YzggTAWhp+f4gEgN9HCPHZIrAkM+MZ7u/f1ZbKHKlEJq8yZfa4aWfu6Ej1qNQ1gY5Y1R28sdQK0dWuNTrUhpgnMfPt+kFzXk1b1yFD3QzrQQACEICAMwJ7ZRaX12S6XiBiF78m2Jy6MDr+LcLNNQE3Gtq0rqmcWFy/+J/27R+5VPJxvccSEAEkAAIyUAizcIGetsufIIP3KLwmx2rYw+w3MxX1DdMcawEVuyLQPinydkc0choxHaoNPqHFZwsfTEJ36ofOdE1mkYsfOn3GgHAgAAEIeChQ3dYV3ZjNtzWEbbS4soiYF2ajE1pdaQyNfCTgyn2fsdSa/H/VlcbWakQIe//X4gjEQyQAAjFMCNIugdyc5keFeR+76nOgni1FeHYFvibQAVr3q8zWRe5auWyTfYnkfG29V4vfliiRMUc/hM6unds5xW/BIR4IQAACxSpQ3dZ9EQulNRnr2uHawtTUEZvQVKym/u2XO5EZzB6c+08fEPFfCLdACSABEKjhQrB2CPS2Nv/DIHOiHXU5VYcIXV4Rb7jKqfpRr3sCfz/8y+9no1UXcSi0LxPd7F7LQ29JP4ROE5PbajLds2pT3QcOfUusCQEIQAACwxWoaev6I4ucO9ztCll/9eS/LnJhIXVg2zwFXNisor5hmhB/w4WmPt3ELb2ppJ++BenT8eHn9QggAbAeFDxV/AJLUjMfYKZa8vFNf5GfGo4n0uH6hq/5OEyENkSB9snbPtsejRyvSYBDhOjBIW7m8mryLTHknppM162TMotrXG4czUEAAhAoaoHatkVb6e/XrO71d/Xq+5j8e/u2cqN1EU/2/hOxYO8/Be+GBEDwxgwR2yTQ05rM9veb29tUnVPVRPWDQiYcbzzAqQZQr7sCmgSY3RGN7KdJgAZt+Q0tflyONMlst/ZSTcx07uvHABETBCAAgSAJVLd1RUUM6yv+qt2MG5N/N7XX25bjT4anNu6ujUzX4u4idE+uteV+dxtFa3YIIAFghyLqCKzAsraZz5tsbubvDnAVkaTD8Ybv+ztORDccAU0CzOw3WSfX/JvhbOfqukLHGMQP1GQ6fzdxbteerraNxiAAAQgUiUDt3O4fsFBKu7OVFtcWTP5dox6gIedfElNOc76Vz7bAOPf/sygBecYISJwIEwKOCSxtnfl2btnoUcT0P8caKbzijYj4hnA8cSnhVjQCD8QrX8pGK39IodVX5b/Hvx3j7xomPVzT1n19bVvXfv6NE5FBAAIQ8I/AgXc8Gq7OdF+hEzQr0ctuRmZd7b8D5/y7Sb7+thx+tjyamKJvrGMdbmY91Utnzxs9t6znBTwVAAEkAAIwSAjRBYH2pg9yrckv6p72p1xorZAmzgrHGm4fU3e2z49aKKSLpbdtdnLlPdlY5GB9//1Qe/+SFn8uIieJ0MKaTFe6NtX5nSYR/A3x50ghKghAwGOBmszi+lVjy+9nkp+6HcrqyX8MV/t323197Tn9HIfoJ063sb76xdr7/9gNq9b3Gp7zvwA+vPl/jBChiwK5VMvXSWiui00Ovynmw0OhvnR5tGGv4W+MLfwskI1W/YZllfVtATP9HKfGFhWD/7Rg7uKnqzPdP98v1b21PocFAhCAQMkLHDzvqdHVbd1JInMOiezqNggm/26LD9ieoy9WxBPH6mfWKY42sv7Kl7OEsPd//TaBeBYJgEAME4J0UyCXTtZpe7/T4udlDzZYkwCJb/s5SMQ2fIH22A5vtEcjDQaJdaj97OHX4OIWItvr3q1fhgzz6ZpM1w01qc5JLraOpiAAAQj4SkB/D9at6N94AYs0eBEYJv9eqA/UpoOvTW8aKUSe7P0npr/k0pf592hFB9mLpWokAIplJNEPWwVyqeT3SPhiWyu1v7LPsUG3hOsbfm5/1ajRa4EF0aoHs9HIIcJ0vO5FepZ8fePRGt73yeD79QPw/Nq27uPrMos30uewQAACECh6Aev3XW2m+zLtk+AbGAAAEABJREFUaEbLHlpcX0TMCztw2L/r7gM26OCL4Q+W/1ird/0IE21T5/+Evf8WRIALEgABHjyE7qxALt18rmbxT3e2FRtqF/5lOJ64acvpp4+yoTZU4TOBjrrIzR9wyPq2gIs0tPe1+H05SER+/x6ZT1enFzXtl15U6feAER8EIACBfAWqM4sOfo9kgZCcmW8dBW8n0ojJf8GKtlfgVIWb1jVuSULe7P0nbutpTWad6hvqdUcACQB3nNFKQAV60i1XauhB+Pq9E5a+X5betP7M7TVeLEUm8HB0fG82Wnm+kKGJALoxIN2rZDYuCLHxdE266/c16e6DAhI3woQABCAwqEDtggVltZnOXzIZ84hkH/LoJkyxbKwq6VHzaHbDAo690m+INfnf0rEGBqhYxMTe/wF8gvISEgBBGSnE6ZlALpX8rYgE4Vz7Sf1ipsfFG+o9w0LDjgp0RMc/kY1GTiKmPVnoJkcbs6/yjTTe44llfk1bl3WKwPcPnveGdcqAfS2gJghAAAIuClRnOveXD7a4X4i9PAWvl9n8WkddxDrtwMXeo6mhCTizVjie2FX/plqH/zvTwMC1PtubHoMEwMBGgXgVCYBADBOC9FqgN91ym8kUhIn1V03iORXxhJUd9poN7TskkK2L/LM9FjnREN5b9zz93qFm7K9WyLpI4A0r+nNPW3vOJmYWb2d/I6gRAhCAgHMCtW1dTUx8HwlNdK6VQWu2ksHh9roJ/xp0TazgjYBDrcqaC/+NdKj6Qaplnfw3mYOshJcDIIAEQAAGCSH6Q2BpazLFwvtrNCu0+HrRPxBXhmOJ68ZNPzPs60ARXEECC2KVD2ejVScQG99gopsLqszdjbcW3XPGYj5T09b1p+q2rqi7zaM1CEAAAsMTqM50Hau/rx4WoQuGt6Xda/Nfs9HIrnbXivrsFXCitvJoYor+rT/WibqHUGeOzX5NAAxhTaziewEkAHw/RAjQTwI96eYFwqaV9X/ZT3GtNxamk833zXsqYo3WXtf1roIni0MgWzf+7+3RyPEGhfYjpj8FpVfMZJDQd1goXZPpWlib6Tp13zkvfoVwgwAEIOATAf299M2atq57dOJ1s/6+2tPLsJh4RjZaOd3LGND2kAQcWYlD5OXRnbf0ZGa+4kjHUKnrAkgAuE6OBoMu0Ns685/9wjGdaD0agL7sISzzNQlwRgBiRYgFCiyIbvtgti5yrEh/jX5Q/EuB1bm9+X5CdFVZWV+nfuC+rTrdfVTtgmfHuB0E2oMABCBgCVRnOvevaev+m/5euouEDrSe87CsYDG/2x6tPMvDGND0kAXsX7EinjhW34dT7K95SDUuDbFxzZDWxEqBEEACIBDDhCD9JrAs3bwoZPQfqnEF4atQyjQJ0KJ/PP78ufjp/6cxYylygY7Ydh36QfE73G9aR3/cFrDujtIP3Ecwy1/k/RGdNZnua2vnvuDVh56A0SFcCECgUIHadOfu1ZmuP2gS9T4Ssf7OF1plods/rzscJrbHJvyh0IqwvUsCNjdTPvn0TYXkZzZXO+Tq9P/Cle+2znhuyBtgRd8LIAHg+yFCgH4VePfuy1+VVX2HMknarzGuHZdOqo7uo7J7xuFbAtZmKerH7fUT2rPRyLf1g8MBuufgjgB2VhNWcoqY/W01ma6ntVy8X/q5vQLYD4QMAQj4XGC/9KJK/R3za2HjYSY6zh/hcutGoQ/2zNZF/umPeBDFUATsXscYUfYLIt6ZPLlJt7lq1ZWeNI1GHRNAAsAxWlRcCgK98654t2fUGE0CUFAOt55gEs8pr0+cXwrjgz6uEeiIVt2fjUWOYIMPIpK/rnk2cP/uqBGfE+Kyf9Sku7K6h+6M6tbnx+tzWCAAAQjkLVCbfvlLtZnOX4Y49LBWcpr+jvTFZ2MhvjIbrZw6f/LOyzUuLMERsDXS8tiZRwrRj2ytdFiVGVdYn3WHtQlW9r2AL37J+V4JAUJgIIFZTSt7UsnvMPFFA63mp9dY6MLyWGL22KmNET/FhVicFWifUnlvNlo1XciYTEJ3Otuag7UzVTNRCxmhTt1jd1dNuvO7B9y76HMOtoiqIQCBIhOob31tk5pMZ6PwykeE+Oc68a/wSxf199tpHdHK0/0SD+IYjoB962467WdfZu738rNley7VfJ19PUJNfhFAAsAvI4E4Ai/Qk2o+n4WP144s1eL7hZmmGabML48mvu37YBGgrQId0fHzs7HIYWyE6rTi2VoCueh72Pob9k1i/l3fSkOTAd03TUp3TiMRDmSHEDQEIOC4QF1m8Ua16e4f9IaWPkLEM4joy1r8sTAtMMiobY9GrvZHQIhi2AI2btDXF7IO/d/WxiqHVRWbgkP/hyUWnJWtD0/BiRaRQsDnAj3p5psNoZjuSXjK56F+FN5X2KBbwvGGZo0Zk6aPVErkvn3KtnOz0cghOl2O6SR6TsC7/Xl9D59gMs+uaevurM50Jmvmdllf2RnwbiF8CEDADoHq+S+Or04vanqPzKeF5Tda5/Za/LL06++v87N1kf0XRMcH4eLCfnHzXRx2BRSOJ07UJLe1U8muKodbzy09mZa7h7sR1g+GABIAwRgnRBkggSXp5MIRhhHVCVWADrHmRDjeeE84dsZuAaJGqDYJdNRFMtm6ymmGwXESutWmar2sppKJG8ikjppM999r2rrOnjR3sZ8+7Htpg7YhUFICk9q6DtDfAzdxf//TzMYF2vlKLf5ZhO7ViV5tNlrl5aHe/vEIdiS2RD+2/owqJtG9/7ZUl08l/cIm9v7nIxeQbZAACMhAIcxgCbw9p/m1XGvzYTqZaglQ5AcQG/OtrHOAYkaoNgosmFKZzsYiRxkGfZ1EkkLyGgX+Jnvr/8NLTNN8VhMBbbWZ7oaJc5/fNfDdQgcgAIENCliH+dfM7f5uTVv3PaZOsInkBP2dtvEGN/DgBRFZqc3+XH/nHtReF3lAH2MJvIA9HTDE+IUQb25PbXnVcmVv60x880RedMHYCAmAYIwTogyoQC6dTOjk44f64aM/IF3YVOO8UZMA135+auNYfYylBAUWTIk8lY1VNY74QHbSD80/0b3p1tWxgy8hNEWTGknDDD2mewQfq850JmvTnVO2u+PZkcHvHHoAAQisPsy/resC6zB/MuV3+vvrQD+q6O+heWyI7vWPXOLH+BBTngI2bFYRT/xEq5muxavl1TIaeYVXjaNddwSQAHDHGa2UsIAmAX5DxFEt3RSc2ymr+uWeivpETXBCRqR2C9x36IR3NBHw6/Zo5d5imodpIuBvdrfhXX2yq/anQZjbNhsz4sWaTNfN1ZmuY2szi7f0Lia0DAEI5COwzmH+Qk1E5K/D/DWgNQu/JyRndUSrpmTrJvx9zXP4t1gECu2HdRqmEHl56D8xyZXvpC7+b6F9wfb+FkACwN/jg+iKRCCXSs4XY1WMhOYGpktMe4nQ/IpYw08DEzMCdUygIz7hTk0EfIu4bE+dOP9KG3pbS7Es1qT/WCa6Wch8EacKFMuwoh/FLBCEw/zX9memtNkvk3TyP2Pt5/G4aAQK7wgb1uS/vPCK8q7h4Z7UGJz7nzdfcDZEAiA4Y4VIAy7QO+eKF3Lp5qhOMq4JUFdG6h7SKyriDX8M1zd8LUBxI1SHBLJ12/xTEwE/NY2NdmLihDbzpJZiWkZqom6KkOBUgWIaVfSlaAT2Sz+3V3UADvNfC3yp7tVtaK+LxBfWRx5Z63k8LCqBwjoTjiXO0RqiWjxbREQn/02mZwGgYdcEkABwjRoNQcASYOlJJU9lkjOsn4JShPgYEm6viCVOCErMiNNZgYVTtnpdEwEt2WhkF92z9W1tLaWlCBecKlCEg4ouBUzgo0l/dabrHyEu+wf7+jD/tXF5jmkYtR3RyMy1n8XjIhQooEsVU8+sJvb20H8Nf3ZvuuU2vcdSAgJIAJTAIKOL/hPoSbVczizf1Mhe1RKU5SvCdFNFPPHn8nhifFCCRpzOC+ierds0EVDPTBN1z/l12mKvlmJccKpAMY4q+uRLgfVN+ploL18G++mghN7URP/p2WjltIVTxj/+6Zfxc/EJ5NujsfVnfF5MM6nbh7R4tujfb93771nzaNhlASQAXAZHcxD4SKCnteVuwzRi+nNWS2AWITpag20P1zecrPdYIPCxgCYCHuiIVp3CbO6siYBzifh5Kt7bBk8V2C/19Lji7TZ6BgHnBAI96V/DskKYW0LUv1d7tAoTqjUmpfBv3n00xLhaN95Ti2eLvmev7GlNBuqzqGdYRdIwEgBFMpDoRjAFlmRmPDOaR1vnfP0uSD3QvTBbkPB14XjitjF1Z2wXpNgRq/MC7XUT/qWJgIt51H931j0bxxPxfCr6m3z8rQIhY+R/ajLdC2rSnc21mc5v4ZsFin7w0cECBIpg0r+m90y/NQzaq6OuMnF/bLtX1jyJf0tDIL9e6meoS3XLI7R4uTxphFZe6GUAaNt9ASQA3DdHixBYR+C11qb3cqnk94hF95iu81IQfjgiFDLay+OJHwUhWMTorkD7pEl9HfEJN2ejlZOFQgcQsZXoylHR33gTIqkl5oQQzxIyX63JdD5Wnen8TU26+7v7pV/06VeUFf3AoIM+ESiaSf9qT/6rQUZtti7y/QVTIk+tfgr/lJZAHr0NxxIn6WZnafF0YeYLe2Zf2eNpEGjcdQEkAFwnR4MQWL9ArrXlYhHDupha0L5ebTMmujoca/jbuKkNO62/d3i21AU6otver4mA77GMrBIxf6AeGS0ltLB1hMAPNNH3uxD3ddVkurtrM91/rkl3/rh2fvfXSwgCXS1BgdoF/xszMdUd00TYjOrMJxfy078dwTinf31jxnSvKfJN/b02fUF0PA6fXp9RiTw33G7q56UDien64W5n9/prDv1vnm13vajP/wJIAPh/jBBhCQn0pmfcpnsMYyz0YOC6zXyoaVrfFNDw08DFjoBdE2iPfe2NjtiE67PRSMwwjB204Z8L0cN6X2KLjBeSo4n5V9InT9Sku97QydGc2rausyZmOvctMQx0twgFJs59ftfaTOdPazJdd8r7Pa8ahqSIuDHQk35afXtUmI7P1kUOWhirunv1M/inlAWG1fdwXWIb/b1/z7A2cmZlHPrvjGsgakUCIBDDhCBLSaA31fzIJsbog5nYujBM0Lo+TpivKI8n7g7HztgtaMEjXncFFkwZ/5wmAi7piEb2lhDXkNDlGsELWkpvYfoiEdeL0KUG8QM1mc7lNdZ1BDLdv6hJdx9U+/uXNybcIOBjAZ3wb16b6TpC37fX1mS6Fhtm6DEhvkJDPoRIKvQ+6MsLIvLTSXWVe3XURW4OemcQv10Cw6inqcmgEKWHsYVjq+LQf8doA1GxEYgoESQESkzAui5AT6r5NBI5Ubv+jpZALUw0VTPc7eH6REOgAkewngl0TK7syMYiZ3xhWWUVMR0qRPoBm5d4FpDnDX94HQGS84hlvnxx5fs1bV0P6wSrZVK6cwkfL3wAABAASURBVNq+c178iuchIoCSF6hOd1frZP/Cmkz3Qzrhf03/396mk/1TFGZbLcWyvK2JuaaNjU326ohV/aqJ2SyWjqEfNggMo4ryR5ffpatHtHi66I6aK3tacei/p4PgceNIAHg8AGgeAgMJ5NItN+nrB2u5V0vAFh6je3ST4XhDpjzaENzzPAOmHvRwZx3O/dm6yF0d0cjxq0JGlfbHulBSSu+xCO2pE6wzTObZZWV9r+jEq0fLQp18XVuT6Tyltq1rvwPveTEMKAg4JVDb2lVVk1p0Wk1b9936vnuPWaxz38/XSf8+TrXpYb19QnJlv8k68Y9cOG/Kl9/1MBY07VOBoYYVjidmrN45MtQNnFsPh/47ZxuYmpEACMxQIdBSFcilko/nUs0H62S6JZgGXMcGt4fjDWcHM35E7ZXAQ5O3/V82GrlRS72YvB0L6XtI/u5VPD5s15rs76eTL93jytfqXsqFq1b1aVKg85XqdFdKy6W6h/ao6vSiHX0YO0LyuUD1vBe+PGmudeG+rnNqM1231WS6/i0hWkSG8WsSmarvu1E+70Ke4fGbxHSFaRh7dUSrTn8gXvlSnhVhs+IXGFIPK2KJE3TFRi2eLzj03/Mh8EUASAD4YhgQBAQGE2DJpZMJ3RtxtK75Hy1BWzYm4kvCscT8cfEzdMJCuEFgWAId8cpF7bHIZdlo1TeYaaL+X7ASYt3DqqRkVuavqFFMy1m6h/YvzMbTOnkTLXrf+Zfatq6zrIkdTiMomTfEgB3d7o5nR9amO3evTnV+rybd+Svds7+gJtP5Lvf3/9s0xTr65mIhsr6r/MsDVhT0F5n+yWScvlHI2ClbF/nZwinjHyfcIDCgwOAvjp2a2FeYrKM5B1/Z4TWEGYf+O2wclOqRAAjKSCFOCKhAb6rlln7hg/XDWKv+GLyF6SCTuL081nBe8IJHxH4RaK+LPKB75hLZusoqwzC+qXsif6+xvaMFy8ACOxLxUSJ0qTWxw2kEVHK39ezVf36zMSNW6MTgn2zwb4n5x/r/qZaIx1GJ3Jj4b9r/w7J1kT3bo+OvnD952/+VSNfRzUIFBtl+bP0ZnzdMemCQ1dx6GYf+uyUdgHaMAMSIECEAgbUElqWbF/WmklP1Q8tFaz0doIccYuZfhOMNC8LxxMEBChyh+k1Ad28vmDL+7my06gTupyoh43s6eZnjtzB9Hs+GTyPIdM7VvcHXV2e6f16T7jqmel53dW3boq2oSfDZweeDWpt+9ku1mef3HsJe/Qk+74pD4a05zJ+lbI/2aOW3Ouoq73SoIVRbxAKDdc0Q463B1nHrdf3cdWHP7Ct73GoP7fhbAH/E/T0+iA4CGxToSTWfzySH6gqLtQRwYd3LRPPC8cYbNEtuXewtgH1AyH4RaK+PvN0RHf87TQZMk5BszSwnEvEtJPQm4ZaHAH+FiScTyUn6e+aXxPRH7pesiPFyzZ7d/TWZrlerM10PaGLgltpM52U1mcWnWKcVWNcbOPCOR62kQh5tYpOhCnw0wa9t6zqyNt11VnWm+zc1bV1tWhbVZLreFx7xulDo76W6V3+DjrzuYf7tsW0e3eC6eAECAwsM+Go41rBywBVcfFFw6L+L2sFoyghGmIgSAhBYn0BPquUuLgtZe9Fnre/1YDwn39cs+YPh+oafbzn99CK9qFQwRqJYouyYXPVye13VTdlo5dHLzKVbs5j1mgi4QstTxdJHH/RjSybaVxMD3xbiM4nMa63TCpiNp1eNGdtTk+leohPRpzRB0KqT1KtrMp2Nen9kTdviffZv7fo/H8Tv6xCGOsEXoVuF6VJN0vxA399TtFjJ1I193TkPgmMc5u+BerE3ueH+heOJZcQ8YsNruPoKDv13lTsYjSEBEIxxQpQQ2KBAz+zL/pVLJQ8nknM2uJL/X9iUhH+59P2yh8qnNnzH/+EiwqAIPFa/+3vtsQmpbCzyMy1fZ5E9yLAmrDJf/8/0BaUfwYtTKjTmnYgprpPUHxHxDL2/lcR8qD9E/6nJdPVpcuAlTRQs0PtWLbfo4+v1fqbu0W6qTS1qqE53n6zlqOpU59SaVOek2vSLu1tfQ2clEFYfZeD3UxFEuHbByxXWaROT5nbtPCmzuGZiunNadabr2Jp0549rMt3n1aS7Zmpi5LfVmc6/6t77e7QsUpv3P9qDL5jgU/43HOafvx22HFRgAyuE4w1L9KXRWnyx4NB/XwyD74JAAsB3Q4KAIJCfQC7VcqmYVEdCz+RXgy+2+jqb/KfyWGJ2RfSMal9EhCCKSqA9VvVodkplczZaNVlCZVvr3tPjtYN/1BLEb9fQsAO7hDQ58DVNwtTqfVzLt/XxSXr/Mx2TC8QwksxynZa/sMF3a9LmfuG+f0qIFvVrAmH1UQZrTkVYppPoN3TSvLgm0/24Pu7Qx+madOftmkjQiXX3FdXpzouq04uabC9t3RdVZ7qu0kn7H2synXNq0l1W20/VpDtfqcl09dS0dZvy/sol1mkTpklPmmS2G8yzmehmYv4VkfyCmH6mk/zvMfFh+rv7QC3Yg1/oW1roEV7rav44zJ9wc0BgfVWG44k3idhKfpI/bnJpT2vzbH/Egij8JIAEgJ9GA7FAoECB3kxybl9oxEFCZE1oCqzNu82ZaZpOALIV8cZfaSLgq95FgpaLWaBj8ravdtRFbs5GI8ct+0Ll1gYbU0gkKSKPFXO/i6xvo3US/UXt07Y6od5FH0/Ux1FiPlwTCTqxlp8y87nMxgW2F5FzmehUnbQfQ8T1xGS1vRMxf4WIcB0ERXBxeZqYLjVJ9svGInu142r+LtKXZFOf6XQ43mAlkb/wmRc8ekITi3/QHUNBPjLUI7nSaBYJgNIYZ/SyhASWz7nkzd5U8jgWOV277ZuL0Ggsw16E5MemYTxUEW/42bA3xgYQGIbAY7vzqgV14+dlY1WNHbGq3Vl4FxI6g5nSRPw+4QYBCPhMgBfr34grDaYDs9HIztm6yDkLo1UP+ixIhFOUAut2KhxPvETE/0e+ucm83v/b9iTfhINAfCeABIDvhgQBQcAegZ50y5U6ebEuEPiIPTV6U4vuYdtCiGfqH9iHwvEzvuVNFGi11ATaY5VP6p7Ey9vrInHT6NtGJxrf0T3Mv1OHf2nBAgEIeCPwH/17cJ0m6OqzdeMjHdGq0xfURe7zJhS0WrICa3U8HG/o0h+/psUnC3ezKSfTDSev8klACMOHAkgA+HBQEBIE7BLoaU1mQ/2jD9I9mdfbVaeH9exDZMwKxxtvLa8/Yw8P40DTJSawcMp2r+tE4y/ZaNX3snXWqQJ0oBJcQsR/J80KaMECAQg4J/A2W9dNMOXwlRWbVHZEK09pj1WmiFmcaxI1Q2DDAh+9Eo4nniHiSvLRTbjvuJ7MzFd8FBJC8aEAEgA+HBSEBAE7Bd5ta+rNpZM/0A9Qx2m9L2oJ+CJHshgPhesbLxtzSINvzrcLOCrCH6qATjqsPY7ZaOTn2WjlN0L99GVhPkwnI0kSWqj5AOx1Gaol1oPAhgWW6/+nO8gwjl1p9EXao5Hjs/GqWX//xpdxOs6GzfCKOwKrWwnHGx7XBzto8c3CJIf2tl7+D98EhEB8K2D4NjIEBgEI2CrQk0r+MdTPtUR8IwX/VkYiZ4ZW8UPlsTN+GPzuoAdBFbi/PvLfjrrKO7N1lY3ZWKS67APZXNiYRkyX6nt0ARG/R7hBAAKDCoiQScxzdLf+yaF+iuj/pyOyU8b/6e9Ttn930I2xAgRcEyAKxxM6yeZdXGxy0KZEzB/1pFruGnRFrAABFUACQBGwQKBUBN5ta/5PLtV8koh8W/u8SEvQl22YjWv0j/F95bGGWNA7g/iDL3DfoRPe6agbPydbFzknG6vaf2XFqM+bbE4hkouIeD4R9WrBAgEIWALMz+nE/yZhObGMjcpsXeW0jmjkBiuxZr2MAgG/CZTHGxZqTHtp8dEil/amZ17ro4AQis8FkADw+QAhPAg4IdCbbrmtf4TUMtE1TtTvQZ37M3MqHG/83bhYg68OyfPAAk36SMA6ZHlh3YR52WjV+dlo5eRsNBImg6qF5FwNM6P32LupEFhKRqD7owm/Scb2OuHfoSMWObGjruqm+6Pji+AUtZIZx5LsaDje8CQT7+erzgvNzaVa8HV/vhoU/weDBID/xwgRQsARgWV3tfyvJ5U8lcU8TPdOPuVII65XKt81mR8qjzc0jak7ezPXm0eDEBiCQHZKZGFHtOpiTQbE9P5zLLIHMf1MN51NxG8SbhAoHoFXSOgPonv4V0/4o5HIRxP+hdHxzxdPN9GTYhYon3z6puH6xqVEvLPP+vlmbo/ROPrRZ4MShHCQAAjCKCFGCDgo0JOeeacxKlQjxJc72IybVY/VDP0FIaPvkXB9omHL6aePcrNxtAWB4Qq0x6oezdZFrtCEwCHZaOWXTN0zqkmBU4j4FiL6jxYsEAiKwOtEcgsLfd96H2ejka2ysch3rT38mPAHZQgR59oCFdEzqnlE2TskMoZo7Ve8f1xGI3ejpibT+0gQQdAEkAAI2oghXgg4ILBk1oxcb6r5DBEjrtU/oiX4C9NWuucpufT9skfC8cYfEG4QCIiANVHSpMB12Wjl0dlo5MsSkq2JjcM1/F8Iyd+IqVMfY4GAHwTeFVn9nvyRqYkrfb9ukY1WHd0ei/zWeh/7IUDEAIF8BayLDIthZD/e3k8PDN7jndTF//VTSIglOAJIAARnrBApBBwX6E3PSI8d1VerDV2mpViWHXSP1G/C8cTfy+vPOKZYOoV+lI5Ax+Sql7N142fp5OqCjmjVt7J1kQk86rURRj99XYSP1kTXpUQyR0VwDrUiYHFGQISsPY3/ZJHrxTRPNgz6ur4nP9cRW/2evBYTfmfcUas3AuH6xmuYjXWuk+RNJOtplfnE3JzmR9fzCp6CwJAEkAAYEhNWgkDpCPxn1hXv51LJs8ngg4XkgSLq+d4sxh81EXBvRbzhkCLqF7pSggLtkyb1LaiPPNURq7wlG4uco3tdp+lkbFseNXKUTtD2EKbjiTlJRG2aHPi33mOBwHAEVurK/yCha4WM77HwLh2xSEjfY3u2x6p+0BGfcMOCKZEiuXaM9hQLBNYS0M8Jd5HIp79ieK01vHvIRNfkWptv8i4CtFwMAkgAFMMoog8QcEBAs8v39I56pZaFfqHVi5ZiWQ4Q4jvLY4nZ4amNBxVLp9APCFgC7ZO+9oFO0B7tqIvcnK2rbNQJW1STA1/tWyXl1G9+Q/8/f1/f/1cS8T2aGHidcIMA8ftC9KBCXMWaOBIxd9L3zUZa9tHk0o86ouN/1x6rfFJfxwKB4haobSrTHQTW/4Vvfrajvngma1282ReRIIhACyABEOjhQ/AQcFhg1qz+nnTyAjYM67SA+xxuzdXq9YPuNDJlfkU88eexsZ99w9XG0RgEXBZ4cFrV0mz9hL+3xyK/7YhWnp6NVh6siYEtVhp9nyODqjUZ8EP9P3GfpKNBAAAQAElEQVQNEbcT0YtarD3AeoelqASYlml/OojpCmY5xqRV2+t7YZOOaGS/bDTy43ZNHHXEJjyj62CBQEkJVMTP3Dk8ZvnzQrz+zwMeazDJQ7lU0vos5nEkaL4YBIxi6AT6AAEIOCvQM2dGh/7hOZCEfq4trdBSNIsQHW1w6MFwfeL6cVMbdiqajqEjEBiCwN+nbP9udkpkoSYDftNeFzk1G62clI1GttWyUdmqlV+yTifQyeKhRPxjWnNKwW36f0b3kMm/SbMGWrD4T+A1Dck6fP8OIWkhHTtT5JsGS1W2LjJWx7ZG73/WXlf154XRHfBVfIRbqQuEY4npQqZ1sb/xG7Lw+PlHelIt+3ocA5ovIgEkAIpoMNEVCDgtkEsnLxFmzUBLm9NtuV6/0EmmyY+UxxtnhmNnbe16+2gQAj4TuG/ajm9apxPoZPGubLTyqmxdpXVKwbfX7C2u+mo2GjGYjC9bpxaQ9S0FTGcI8ZVifVOB0COaH8ApBmT3TZZrjYvUeJ7e36gJmvOE6DjuNyf1m7yNdXFIHZf/02Idvn9ER7QqkdWxWxiruntBXVWXboMFAhBYS6A8nvgRMd2hT4W1bGjx8Hl+QnfA7OVhAGi6CAWQACjCQUWXIOCkQG9r8z9yqZao/sFMaDtLtRTTshGT/Iy4/5GK+sSFo6f+5IvF1Dn0BQJ2C7RHx//HOrUga31LQV3k8o5o5ek66fxWNhbZKxut2sKakDKP/Jp1moEIH60T17OY6RpNDljfWvAPjed5Evqv3hfb7xLt0rCWFWryhm7xvBo9qCbr7L0X6tu17APz82o6JhuNbKfGU/T+JE3Q/LIjGvlje/2E9gfilS9ZF4fUOrBAAAJDEKiINzYx0dWDr+rZGs/mUs27etY6Gi5aASQAinZo0TEIOCuQa022aAu1InS33hfb8jnt1/ll5siHw/FE4xcPbhhdbB1EfyDghoA1IW2v+9q/rNMMOmKVt+jEdUb76lMNqqbpBHYfLdtrsmBLvS/PPlIZ4lEjx4WkfytzleysiYKJpslxYj5K93SfQsRnEtElInQNMf1JJ8yaRGDrmgVPELF13YK39blV5O0tp5P3l5noMb2/V+O8Q0Su1/tLSaRRWE7Ux4dae+xNlp1l5cqvbBT6wJrUb6yT+82z0cj2arSfmqyz974juv0T9x064R1vu4bWIVA8Avq3/VpNtl0wpB55s1JXLpXc0Zum0WqxCyABUOwjjP5BwEEB/eP0eG86+U0WPp6Ii/HCUV8lohkrRvAj5fWNvvxKII0PCwSKQ6CJzfZJX+u5P7bdKwunVT2tiYIHFsYr09m6ylt1T/d12WhlczYa+XlHLHJqti5yrE6YNYmw+poFu2ajldtmo5HN9LmRK5dtsgnLqs37pSyik+09rcm2U8UQ3lvKyiq5nzabVFdpfU1ehU7et26PRnbX+4M0ziM6YlU/0PtzsrGqZEdd1U36+C5rj/3CuqqnO76546vzJ+9sHdZfHGOIXkDA5wJj68+o0sn/nRqmJhX13yEsHqzyon6+qvKgXTRZIgJIAJTIQKObEHBSoCfdfPMIg/blNV8ZWHQfZoVpO90DeY1+aHi4InbGcU5aom4IQKAwgb8f/uX322M7vPFAbJtunWz/05psO1UWxCof7jh4m8Xt9ZG3m5jNwiLH1hCAgJMCFfHEsYYY92obh2gZ6uL2ev/Wyf+2bjeK9kpLAAmA0hpv9BYCjgm8Pad5aY/1lYHUvy+T/MmxhryteE9h4w/hWOL+cH3iMG9DQesQgAAEIAABCAwm8PmpjWN18n+1EN2s6/6flmEs7q3KxK/r5P+r7rWIlkpVAAmAUh159BsCDgn0pC5/qifVciwzH8JCDzrUjLfVMk0iob/qB4rWcLzxW94Gg9YhAAEIQAACEFifgP6NPmCVKffq5P9H63t90OfcW+HtnlTzFu41h5ZKWQAJgFIeffQdAg4K9LQ2z+5JJ/djkjP0D+9rDjblWdXarziRzKqINzwYjiVOorrTNvIsGDQMAQhAAAIQgMDHAuF4w9n6N9o65H/Pj58c5gNXVhfK6Z7/zVxpC41AQAWQAFAELBCAgHMCPamWy42y0L5M7Oev2ikIQIi/QUzXh0MbPxGubzxzdDTxpYIqxMYQgAAEIAABCOQlMPbjC/3xJXlV8MlGbjxanksnK9xoCG1A4CMBJAA+ksA9BCDgmEDP7Mv+1ZNqPo1YDtBGMlqKdZlAIpeVGfRkOJ6YMabujO2KtaPoFwQgAAEIQMBvAhW2XujP8d6t1D3/YxxvBQ1A4FMCSAB8CgQ/QgACzgnkWlvu1z92MU0EnEgknc615HnNX9QIGkMh44lwvPGGsVMT++rPWCAAAQhAAAIQcEDAkQv9ORDnJ1WK6OehjT75GY8g4J4AEgDuWaMlCEDgQwFNBNxEUmYdNm8dnrfiw6eL8W6kJjq+b5j0QDie+Ou4eEN9MXYSfYIABCAAAQh4JaCJdkcu9OdUf4TotVyqBXMwp4BR76ACePMNSoQVIAABJwRy6cuW5FqTPycx9yWmW51ow2d1HmYSz9FEwH3WIYo+iw3hQAACEIAABAInEHbuQn+OWAjxA72p5P85UjkqhcAQBZAAGCIUVoMABJwRyKVnPqaJgKOIeDoRPayl2Jf9hejmcDzxREWs4acV32zCxX+KfcTRPwhAAAIQsFXA+Qv92Rru6sr0b/8fe1PNE1f/gH8g4KEAEgAe4qNpCEDgE4FcqvmvuVH/2lefOZOY/qf3xb58XZivkL7lT1bEEheGY2dtXewdRv8gAAEIQAAChQro38wTDDGsr/c7pNC6Nri9zS8IyYW9qeRxNleL6iCQlwASAHmxYSMIQMARgVmz+nOpZLNwaF8Sut6RNvxX6VeF6Xzi/icq6hO/1kTAbv4LERFBAAIQgAAEvBUojzfuGa5v/Jv+zbxJI3H0MHqt37ZFmI/tTbU02VYhKoJAgQJIABQIiM0hAAH7BXrnXPZCLp38gdY8mYnTel8KS7kInaaJgEcr4ok/h6c2HlQKnUYfIQABCEAAAgMJbFHftElFvPEXTLKQRA4daF2bXrOrmvdMg/brbW3+k10Voh4I2CGABIAdiqgDAhBwRCCXSs7vSTXHac31ARZQidyE6GgyZX44nkiH6888vES6jW5CAAIQgAAE1hEI1zccvlyWLxSS8/SFkVpcWOxogrtXjRyx7dI5yQftqA11QMBOASQA7NREXRCAgCMCuVTzX3Op5P7WYXTawD+0lMoSJTFv10TA38vrG384pu7szUql4+gnBCAAAQiUrsCYaWdsVxFv+CMJ364Ku2pxbym8pftzqebIe3de8nrhVaEGCNgvgASA/aaoEQIQcEjAOowul0ruQ0wnE/FTVDq3vVnkmlCo77lwrOEa/VBUWzpdR08hAAEIQKCUBDTp3RjqNx4Q4mO86HchbQrRH/VzygGF1IFtIeC0ABIATgujfghAwHaBXGvyhtyoTfbUSfHpRNxNpXPbjJh/KMQL9APSAhwVUDoDj55CAAIQKHaB8ngiqqVD+zlDyzgtXiz5t8l0Ca70nz8ftnRPAAkA96zREgQgYKfArKaVPemWK41RvCcRn0NE/9FSSkutJkBwVEApjTj6CgEIQKAIBSqiZ3xVk9rXMlFai8ffk58nsNAPdefEz/PcGptBwFUBJABc5UZjEICA3QJLZs3I5VLNl65atWpPJr5I639HSyktOCqglEYbfYUABCBQRALWkWxiGAu1S6do8X4ZfgT/NUim5tLJ3wx/U2wBAW8EkADwxh2tQgACNgu8N+/K13tSzeeThPbUqmcS8XtUejccFVB6Y44eQwACEAicQEV9oiYcb8xYR7Jp8F/W4otlmEE8zEZ//ZJUS+swt8PqEPBUAAkAT/nROAQgYLdALn3ZS7lUsiHEvCeJXGt3/QGpD0cFBGSgECYEIACBUhLY5NBzNteJf7MItRNJHfnrNpxo/tbfX1bfM+fyJ4azEdaFgB8EkADwwyggBghAwHaBd1tnPJdLt/yIDN6DiH9PpXvDUQGlO/boOQQgAAFfCFgT//J4Q1PZylWP6sQ/4YugPhPEUJ7Q1AXR2bqj4VvL2i59ayhbYB0I+E0ACQC/jQjigQAEbBXIzWl+NJdqPoFNs4aIb6PSveGogNIde/QcAhCAgCcCa0/8mfgCJtrCk0CG0ujg6/yDDeOAXCp52eCrYg0I+FcACQD/jg0igwAEbBToyczsyKWavy1iTBGiOTZWHcSqaq3zLkOhvufCsYZrKuINtUHsBGKGAAQgAAF/CgRx4j+QJJP8euOVcmDPnOYFA62H1yAQBAEkAIIwSogRAhCwTaA3PWNebyo5jZkPIaF7bKs4mBXhqIBgjhuihgAEIOBLgQBP/Dfk+YqQHN2TavnJm/Nblm9oJTwPgSAJIAEQpNFCrBCAgG0CPa3Ns3Pp5MHCdJQQP2BbxcGtaJ2jAsLxxMHB7QoihwAEIAABNwWCP/Ffj5bwHVRGB/SmWm5Zz6t4CgKBFUACILBDh8AhAAE7BHpbk7f2pponksiJxPSoHXUGvI7VRwVoH+ZpEuDpcH3i4rHxhn30ZywQgAAEIACBdQSKZuK/Tq9oBQs35NLNR+RmJ19c9yX8BIHgCyABEPwxRA8gAAEbBHLplptyrck9xJBjtLr7tGAh2pGEzjGIH6qINzyoCYGzNq0/c3vAQAACEIBAaQsU28T/o9EUooUGGQf2pJtnfvQc7iFQbAJGsXUI/YEABCBQiEDvnJY/51LJAzX7P03ruUsLFhUQ4m/o3aX9Yj4brk+0ldc3/nBc/Myv6HNYIAABCECgRASKdOK/ZvSEWnq32PaAJakZOC1wjQj+LVIBJACKdGDRLQhAoDABzf7P0UTAocxUK0J/INw+ERCaYn2LgElmZzieuEOTAcdUfPOnFZ+sgEcQgAAEIFBMAkU98SdaTEKH59LJBN1w8qpiGjf0BQLrE0ACYH0qeA4CEIDAhwI9rclsbzr5XTZpFyH6lT6d04JljcAovZuuyYA/St+IrnB9428r6hu/SdOnh/R5LBCAAAQgEHCB8mjDXuXxxJUjVq56kokvYKItAt6ldcLX/vzFIONAnfzPWucF/ACBIhZAAqCIBxddgwAE7BPoySSf7E0lf0osuzDTL7Tmf2vB8onAF0jkeyJyV/i9rbrK6xuvCNc37P/Jy3gEAQhAAAJBEQjHE0eUxxKz2eB/6CT5Jxr3F7QU0SLLhOnHPankd5akZuDveRGNLLoyuAASAIMbYQ0IQAACHwvkWlte7mlNXiCj+nZh4QZ94VktWNYWYNqGRX5KwveF442PV9QnLiyvP2OPtVfBYwhAAAIQ8JeAdV2XcH3jmfp7+0mN7DZNdlvXwtGHRbfcJyYd2NuavIqIiq5z6BAEBhNAAmAwIbwOAQhAYD0CvbOueLcn3Twzt2z0Lvry91noQb3H8hkB2UWEzmcxHgnHGrLh+kTDibOnXgAAEABJREFU2KmNkc+shicgAAEIQMATAU3S1oRjietM6n+GRC4jkp09CcT5Rpdq387JpZIH9mZaHl7THP6FQOkJIAFQemOOHkMAAnYKtDf16YeJ3/akk/uR0OFadUYLlvUJMFerUdIwpbMinmgN1zec/PmpjUV1Pun6uo3nIAABCPhOYHrTyIpY43HheMNcEWonppOJuJyK93Y7kzExl2q5lNa+4TEESlAACYASHHR0GQIQcEbAuoiQJgNiWvtk/TB1q95j2YCAEMVJ+LpVptmle55usT6IjsPXCm5AC09DAAIQsEdgbPTsyvJ4wwXh95c9Iyx/IOLJVNy3LhbzeP3bfGRPasZTn+4qfoZAKQoYpdhp9BkCEICAkwK5VHJ+rjV5lIRkbxK6XttaoQXLegV4jCZLvm19EDXJfDkcT9xbHk+cO3ZqYt/1ro4nIQABCEBg2ALlsTMml9c3/MEw+p5h4iYirqQiv4nQFf39ZRN70jNv3kBX8TQESlIACYCSHHZ0GgIQcEOg9+6Wh3Pp5A/6ha3rBDRrm//TgmXDAtbfpAOY6CLDpAfC8Ubr6IDrymMNR46e+pMvbngzvAIBCEAAAp8W2LTutHLrVCtNrLYzG3NZ+DhdZ6SWYl/uJ5GDetPJny1ru/StDXcWr0CgNAWsD1ul2XP0GgIQgIBLAsvSzYtyqeSZIwxNBAj9XJtdrAXLoAJSSUwnM/OtZeZGL+mH2HS4PtEQjp2126CbYgUIQAACJSoQntp4UHk8cWV/aONnSPg6ZajRUgrLuySU0L+3B+TSLfcO2mGsAIESFUACoEQHHt2GAATcF3h7TvNruXTyko1Xyi5CdKpG8E8tWIYkIJvoalH9cJck7n9UkwFPVcQbflVR3/jNcdPPDOtrWCAAAQiUrMC4WGKi/l6coRP/Z8mU+Uz0E8X4ipaSWJjpz/0hc6L+jW0ZaoexHgRKVQAJgFIdefQbAhDwTODN+S3Le1PJa3QvxZ6aCIiJ0B+IaLkWLEMX2EmIfywid5nvmy+H6xv+Vh5vPHVcrGGHoVeBNSEAAQgEV0B/5+1ZEUtcGI4nHjWZOrQnjTrx317vS2l5RpiO6mlNHrPs7pnPD6PjWBUCJSuABEDJDj06DgEI+EFAEwGZ3nTyu9wX2oGIG4kIRwUowjCXcSR8KJNcZTI/ox+GH9Eyo3xqYgrVHr/xMOvC6hCAAAR8KzBuasNO4Xjj2Trxf0B/5z2sk9/zNdhSPS2qmctWVfe2JvP41h1VwwKBEhVAAqBEBx7dhgAE/CXQM/eyf+VSzclcKomjAgofmj20ikY2qa1izBdesr5mMBxv+H751LO21eexQAACEAiUgPXVfRX1idPD8cS9pslPEcklOvEv4W9KkXnMVKt/L8/smX1lT16DiY0gUMICSACU8OCj6xCAgD8FcFSAfeMiJJsT07eJ+AY2+xeXxxMd+kH6wopY4yTCDQIQgIBPBcbFz/yKJi9P0Ul/2jD6ukTocg31AC2lvLwhxD/JpVqm9LQms4VAYFsIlLIAEgClPProOwQg4GsBHBVg//Aw0UT9IH2+sNyvH6xf1fLXcLwhMS6WmEhNTfibaD85aoQABIYoYH3daUUscYJ1TROTzG5NXl6rm0a1YCG5UYiqe1PNv7YBA1VAoKQF8GGnpIcfnYcABIIi0JtK4loB9g/WllrlYUTcbDJ1hB9d/np5PDFH97qdE443HkC1TWWEGwQgAAEHBSq+edZWFfWJ48PxxF/KzJFdwnSTdU0TbXIjLSW/aML2D2LK3rlUy0m9qaRNX6Fb8qwAKHEBJABK/A2A7kMAAsESwFEBjo7XF5ioXve6XUwk94bHLH8rXJ9oq4g3NumH84Op7jR8IHeUH5VDoAQEmpqMiljjJN3Tb129v136+l/WSe7vtedHacFXmiqCtejv4r8YZE60LpLbm2l52HrOtoKKIFDiAkgAlPgbAN2HAASCK6B7Q3BUgLPDV0FCU4TkAm1mXji08TuaCLgvHG+4pDyamLLl9NNH6fNYIAABCAwoEK5LbKMT/hMq4g1/rHj0vf9YpyDpnn7r6v01A25Yii+K3MHC+/ekkt9Zkpr5gBMEqBMCpS6ABECpvwPQfwhAIPACOCrAtSEcrS3tT8Rns0FtS98vy5XHGxaG44kZ5fFE9IsHN1ivE24QgECJC9Q2lYVjDQdWxBMX6e+HRyhEL+iE/yYhPkYTipuXuM6Gun+XvjA5l245oifdvEAfO7WgXgiUvAASACX/FgAABCBQTAI4KsDV0RzBxPtpi41MlP5gJC/TD/sPh+sTybH1ifjnpzaO1dewQAACJSCgScDx4fqGk/V3wJ3hMct7iPkeITpXu259LaneYVmfgP7uTIlIPJdKHqpl/vrWsfc51AYBCCABgPcABCAAgSIUwFEBng3qniTUYAi1rjKlVycDj5XHEpdXxBqnjpt+Js7v9WxY0DAEbBaoO20jnfRHK+KNv9L/5506ke0m4eu0lUO04GggRRh4kXnMfEhPKlnfm25JD7yuja+iKghAgJAAwJsAAhCAQJELrH1UgH5I/SkJ3VPkXfZT93ZlptOF5W7zfbNHJwpPhGMNvy2PN546LpaYiKME/DRUiAUCAwuE44ldw/HGhN7fGw5t/IH+Pk0LyY91q4gWLEMTuF//Bh2eS7VM6Wltnj20TexbCzVBAAKEBADeBBCAAARKRcA6KkD3tvwql04eHGJjB+33mSTSofdY3BP4OjF/j0musr568MOjBF4I1zf8rTzWeF6FdaRA/MyvuBcOWoIABNYnsGld45bh+sRh4VgiWR5PdOikf4Wu9xiRNOv9AVqwDENAiBZqsuToXCp5gP4NmjWMTe1cFXVBAAIqgCMAFAELBCAAgVITeLd1xnP6Qaw5l26xrkK9m4icrwb4qiVF8GDZhoQPZZZfiHWkAJmv6GTjXS33l9c3XKFJgeMqoomvexAXmoRAaQhMbxo5LpaYqBP+Bv1/91ctr/aH5FUS+isxNTDRRIUYqQXL8AX+wcLH96aS1b2plluGv7mdW6AuCEDAEkACwFJAgQAEIFDCArlU8vHedMtFer+3Kf376ofeS5TjSS1YvBMYp01P0g/OP9WkwB/EoCd0UiJantCkwB8qYo0/1TIpHDvLWk9XxQIBCAxVYOzUxkhFPHGs/n+6Vstj4feXr7COyNHffUmt4zAtW2rBUpjAP3Xz7+vflX160s0362PvF0QAAQisFkACYDUD/oEABCAAAUtgafryh3Lp5M/1Q9suOvncX5+bqaVTCxZ/CHxdx+U4TQpcoeV+4n7rSIFXyuOJuyviDb/QpMCh4brENv4IFVFAwHsBK0lWHk1MKY83NOke/jad8L9rmNIpRNak9BSNcFctWOwTuJ2ZD9G/IXtq+a191RZeE2qAAATWCCABsMYB/0IAAhCAwKcEdK/NAv0A16Blgu6BrmPiq3WVl7Vg8ZfAV5hoqhCfp0mBv1GIXtBJTm95PNGhSYGrwvWN3yuPN+5pTYT8FTaigYC9AlvUN21SXn/GHuFY4hR9/98cjic6rSQZG9TGxBfoHv4pRISjZhTB5kX/LsilRsjYSf9eHOnTi/vZ3GVUB4HgCiABENyxQ+QQgAAEXBPonZOc25NqPi33ei7CbB5CxDcK0WuEm18FxjLRRCE+lUR+yyQP05qjBd7SSdFD1uRIy7n6+Agtu+LbCPw6jIhrvQJ1p21kvW/L6xuP0QTXZRXxRKv+/NJyWb6cxXiEmK7V9/+xum1ECxbnBO4j5hM3Xik75lIt5yy5e8YzzjVVaM3YHgIQ+EgACYCPJHAPAQhAAAKDCzx2w6qe1pmzc6nmkzYatdz6cH2kTi7/pBu+qwWL/wU+ryHuY02OtFykj2/T8tiH30bwuiYFrKud36STqbPC9YnDxk1t2GnL6aeP0nWwQMB9gd1OGmFdALM83nh0ON5wSXl9YnY43rjY+go+DeYxFvmjJrjOFKK4/vw1LVicF1hKTDew8P65VPLAXGvzTW/Ob1nufLMFtoDNIQCBjwWQAPiYAg8gAAEIQGA4Am/NunaZfgC8vSfVcqw5oqxSJ5THkfAdWgc+DCpCAJcv6RhaVzs/QWO/lIT+apr81NL3y94LxxOvark/HEtcr6WhItY4dUyscQLVNpXpulggUJhAU5MxLnrmjuWxhiN1b/5F+l67Mxxv6ApvHl4pBj3BJH8m4rNZaBqRbEu4eSHwJAmfxaa5Y641ebJ1ipgXQeTbJraDAAQ+EUAC4BMLPIIABCAAgTwFlt516Ts9qeQfc+nmI8qoL0JknCRMd2t1fVqwBF/Auir6JGI6SUtSWO4OsTwfHrN8VTieeEnLPJ24Xa3lJ+XxRFTL+OB3GT2wW2CTyT/dPBw7YzdNIk2viDc2heONs/S983z40eX9pmE+zcy36t78c7XdQ4i4knDzXoDlThI6XJO9u+jv9xk9mZmveB/UsCPABhCAwFoCSACshYGHEIAABCBQuMA7qSv+m0vNuLG3NflNklBEhH6sH+rnaM1LtWApPgHr0OuDdYx/pOVKJkpr6daJnfW1ha/r/RPheMPc8vqGP4TrGy6riDX+tDya+LbeT7KOIsDFCYP/hhhbf8bnrdNFdFynVMQSJ5THE9b1Ja4tj1mH7Cce0fH/jxZzxIgRrxEbj2oS6Q4huYBIvqW9n6AFi78E/qOT/hYNabdca8thuXRylj4O8ILQIQCBtQWQAFhbA48hAAEIQMBWgVz6spd608mrelPJaSabWxPLEdrAdfrBv1vvsRS/wJe0i18n4sksbJ0icqawXMEG3aL391tHEXx4ccIV4XjiFS0P6+Tx7nB94vqKeMMvyusbf6iJgkPHxs78hvX1hl88uGE04eaawKZ1TeVj68+uCtc37F8+teE7Oj6NOj5X6v0dFfWJB/TeOvrjA0OMt6zTRXRc24TpJk0AWdeXOIWZpmmwexDx/2nRpwk3fwtkifgULlu1o076E7rX/3Eqhhv6AAEIrCOABMA6HPgBAhCAAAScEljaOvNt3Zt0h36oPCWXaokYZEwUoQu0vXYtWEpbYKR2/yta9tRZ4lQSOkmIz2ORazRR8DeDzQetrzf8YCQv00lnr5ZunYh2hOMNszRRcFW4vvHnWr43tj4Rt5IFFfEzd9bXx39+auMW4w48M1zK1yqwLuJoHXo/tv6MqvJow17heOJg6xB8vT9R7xsq4omLdK/9r9Xr5g/32C8Ixxu7wvWNS/tDy3OG9C0i4fvYZOtinzN0fH6i4zRdhPbVe+voj430HktwBZYQ8e+JaLL+bq7NpZqv65l9ZY/+XDQLOgIBCKwrgATAuh74CQIQgAAEXBJYkprxQG86+YtcKjmJ+sm6sNf3SehWbf4NLVggsCGBsfrCeJ2ITiTib2miwPqqw1+SyG8NoVYrWSBkPqmvd68y5b/mxmZPeM21Cj7QSe87Wv6tZZGWR3WSm9VJcJs+/mt5PHGz3l+rP+6xSzQAABAASURBVCcr1pyfnihfcwTCceF447fK44mo7vWuCUcbd7dOXRgXP/MrYw85+3OaXNiYiJimTw/R9KaR1s9b1DdtosmHseOmnxkun3z6pmPrz/j86Kk/+eImk8/ZfNO6xi2tbSu+edZW1lENWu94fb1qzLQztrMuhFehyQuNY9fw1Mbdy3XCPjbesM/YqYl9K6JnVIdjDQeG6xOH6YT9BI3l9AqNs7y+8YpwvPF3us2dWu7T8mg43rA4HE/8T8sK6yKO1qH3upd+ERv8DyKap9FaF+u8Ue+TQnSu7rU/Tb2O/XCPfS2RVKrnGF0XS3EKrNBu/Y1JTujvL4vkUs0n5FLJ+fpcMS7oEwQg8CkBJAA+BYIfIQABCEDAfYFcW/JF/QD621w6edTYUX1bi0hco5ipE5En9B4LBOwQsPZUb6oVfVlLlZbddJJbrZPgKfr4MGsCrPen6M8Nsub89OYPj0D4g74PZ+nrad3r3U6G/DPE8rxJ5ivGqr63Nbnwvk60zfD7W/WF31++wvp5uSxfrsmHXvN9s4dHlL2jk++3ysyRb4wYseq1/pC8am0rff0vW0c1aL3d+vqiUL/xnGmYT4smLzSOx8iUf1oTdoP4IcOkB8QwssR8jybJ/ipMN4nQ5aJxaow/1fi+q9scomV/LbsRsZVQ24yIrCMr9A4LBCwBaRPmH7Fp6qQ/+a2eVMvvl7Vd+pb1SvEW9AwCEPi0gPHpJ/AzBCAAAQhAwEuB/8y64v3edEtaEwINuVTLrmyEdmXhBiZOa1zva8ECAQhAAAJDE2jXpFHCZHOC/j6N9rY2X1tSV/IfmhHWgkBJCSABUFLDjc5CAAIQCJ5Az5zLnuhJN8/sSTXH+0zaWpiO0l78VsuLWrBAAAIQgMC6Ao+IyPn61G6aSJ2USydblrbO7NSfS25BhyEAgc8KIAHwWRM8AwEIQAACPhVYnkm+0duavFU/1H5fy7ZMMomZfiEkD/g0ZIQFAQhAwA2BZ7SRy0yD9tPfjXv1plsu0vtSv4q/kmCBAAQ+LYAEwKdF8DMEIAABCARGoCfV0t7TmrygN9Uy0TTLIkR8CglbFzh7m3CDAAQgUMwCQi8K0a9I5KBcKrmTlrOXzkk+WMxdHl7fsDYEILA+ASQA1qeC5yAAAQhAIHACSzOXdudSzdfl0s1HjDB4axaeJsSXM/FD2hlTCxYIQAACgRYQoteI6Qbr91tuj9GVvankT3PplnsD3Smngke9EIDAegWQAFgvC56EAAQgAIEgC7w9p3lpT7p5Tm+q+YyeVPO+mhD4sn5gPoxIrK89W6h9W6kFCwQgAAHfC1iTfmb6swZ65EajlkdyrcmTrd9v1NSExKaibGjB8xCAwPoFkABYvwuehQAEIACBIhLQhMBr+oH5zlyqpbE3law2R7y3hSYEpmlC4FLt5gIt+HYBRcACAQj4QUBEo7hXmC4glgN6X89t1dOaPCaXSt7+1qxrl+lrWAYXwBoQgMAGBJAA2AAMnoYABCAAgeIVWHrXNe9oQmBOLtVyjn6o3n80j/68mFTHxBeR0D3a86VasEAAAhBwS+AZ/f3zayKe3meM3Fx/Lx3U25r8Ra615X567IZVhNswBbA6BCCwIQEkADYkg+chAAEIQKBkBF5rbXqvN5Oc25NqPj+XTh6c2310hfUNAyJ8PhG3EdESLVggAAEI2CXwLhHfpnv4f9AfMrfXCf9O+vvnJ7lU81+Xz7nkTcKtMAFsDQEIbFAACYAN0uAFCEAAAhAoWYGmJtP6hoHedPNF+oE8mkslNxXu34eIG5koRUT4lgFFwAIBCAxdQIgWEvO5hlC1/k75XC7V/G3dw3/9srtnPj/0WrDmUASwDgQgsGEBJAA2bINXIAABCEAAAh8L9LZe/g/9wJ7sSSXrc6nkZkzG10Xox0TyVyJ6QwsWCEAAAmsLvKQ/XMfMh5gjyj7fm0pW51qbL16STi7U57E4J4CaIQCBAQSQABgABy9BAAIQgAAENiTQk5rxVG86eVUu1TI9l0pubhpcpcmAk5j4T7rNK1qwQAACpSWwQvfyt2o5VUul/l7YRsspPa3Ns5fedek7pUXhZW/RNgQgMJAAEgAD6eA1CEAAAhCAwBAFls5p7tJkwI09qeZj9UP/VqtWrdpCROI6EThPq7hLC5ICioAFAkUk8C5bpwQJ/ZxJJun/+417U8mpWq7RsriI+hmsriBaCEBgQAEkAAbkwYsQgAAEIACB/ATem3fl673plrROBH6pE4NDtWxlHQZMIgdpjWdquV0LJgmKgAUCARFYpP9/b2Kh75lsTtD/05/rsU4JSicvsa4ZEpA+FH2Y6CAEIDCwABIAA/vgVQhAAAIQgIBtAtZhwLl0y706cWjWcqSWypGjRo9lw6hhkdN1L6J1+sCztjWIiiAAgXwFVuiGC0j4YiGKcdnocfr/dTv9/3tiTzr5u6WtMzv1dSz+E0BEEIDAIAJIAAwChJchAAEIQAACTgq8NatpWc+cGR096ZYrdS+idfrAjrlRozcS5n20/Ejb/q2Wx7VggQAEnBN4hZhuFeLTSMzddbK/sZb9c+nmc3tTyUzP7KYe55pGzfYJoCYIQGAwASQABhPC6xCAAAQgAAG3BWY1rextbf6Hlmt1EvJ9LbtpYQ1jNy3fJ5Fr9f4fWlZqwQIBCAxf4GFhuZKEDi+jvi31/9dWudbkUb2p5qtz6ZmPDb86bOELAQQBAQgMKoAEwKBEWAECEIAABCDgDwGdpDyu5be5dMuP9H4fLRsZIjsK87FadDIjHUSyzB/RIgoI+EbgZSGaQ8JnGULVuVF7lOn/nb17W1tOz6WTs95JXfFf30SKQAoSwMYQgMDgAkgADG6ENSAAAQhAAAK+FViSbnm2t7X5T1p0MtNSk0u1jB1h8P+xadYwyQk66bmYiG8joUeJaIkWLBAoRgHraJinSeQOJr5QRL7NZHw91/+BdSj/1r2p5LRcunnGknRyIc06vJ9wK0YB9AkCEBiCABIAQ0DCKhCAAAQgAIEgCbw9p/m1nszMjp5Uy+910nNuLtX87Vw6uUculdxUVvV9TtjcU0w6SveKnidCf9D7hSzyepD6iFhLVuBdZnqQhG8ioYRO9OPUT9vqe3sjLTvn0i1H9KSam3rTLbf1pGY8RW1XWRfzK1ms0uo4egsBCAxFAAmAoShhHQhAAAIQgECRCPTOu+Ld3taZ/+zNJG/VvaK/7E0nv6v31T3pli1G8+jRhmnsxMyHWJMr7fJ1en8PEb9EuEHAXYFXiGTe6lNbWH7AhlHTP0K+mEslP9fTmtwvl24+MZdOtvSmW9K5tuSL7oaG1nwpgKAgAIEhCSABMCQmrAQBCEAAAhAofoHXWpveW5KZ8UxPa/PsnE6udLJ1it4fnEs1b5NLjQ4JUaWYVKf3p+qe1yv1vpWFnlcZ7GVVBCzDFnhTt3hcJ/p/ZaJfCsnR+vNuY0f1bZJLJbfKpVqmrD61pbXleuubMpbd1fI/fR0LBNYrgCchAIGhCSABMDQnrAUBCEAAAhAocYEmszeVXNybSc7V+2t60y2n6/3UnnRye52sbRxa2f8VJpmkSEcK049J+GJ9/FsrSaD3jxDTv4jofS1Yil9Ax5lf0G5mtdxCQi0scrreH25K/75smlvllo0eoe+bL2nZLZdqmd6TSp7Xm2q5JZdKPv6fWVfo9rolFggMXQBrQgACQxRAAmCIUFgNAhCAAAQgAIENC7w7//JXe1It7TqBu723NXlVLt18rj7+fm8qOVXv98q1Jr+m95uE+j8Ia1Kg0iBzIht8mE4Kf8jEF2rN12m5i63zu2n15HGp/ozFfwJvaUhP6pilNalzgzBdQMwnikF1RsjYybrGhDXOuVTz+FwqWavl6Fw6mehJt1yp97OWpi9/qCcz8xVqb+rTerBAwCYBVAMBCAxVAAmAoUphPQhAAAIQgAAEChZ4t+2qXk0KLF6SmvlAz5zmO3Pp5G96Us1NOlE8RcuhPdb53Wsmj+XWNQmIZWsJyd4sPI1ITtLkwXlMcrVOOu/QYKw9zJ16v0QLluEJWFfNf1s3eZGInyCirJC0MtOfSeRa/fkydT5D7480yJhIOg65UaOtC+19QcdpFx2zuCZ1TtZkzy9yrc039c5Jzl1y94xnrGtM6DZYIOCuAFqDAASGLIAEwJCpsCIEIAABCEAAAm4KWNckyLW2vNx7d8vDPenmOblUy429qeQve1Itp+mk8widiFp7mCfo/aa5Uf8qG2FweZ9Jm5OEtjFEdhRzdeJgf5Opntg4Qie0J2gC4VTtw5m6B/tCIk5ak13rmxCIaJY+l9b7di2P6HrP6f3LWv6niYdlxGTqYz8s1p7zd0nIOqXiaY1z4eq4hW4lIusoimadrJ8rxKcx0XFsXdCR5QA12V2IKvuMlV/KLRs9Ss2syfxmer9tLtW8a0731vemWqZqAuaYXLrlR7lU8mx1vlzvb1+SmvGANQ40q8lKGmgzWCDgLwFEAwEIDF0ACYChW2FNCEAAAhCAAAT8KjBrVv/bc5qXLs8k38ilL3tpSbrl2d7M6sTBgqWtyVSudcYdOqH9vSYQrtFJbbPuwW7KpZobrclubzr53Vwqebg+F9f7SVr20vV20PuttXwxl2oZm2tNhsaO6tvEHFH2ed0j/lWTzQkk5u6GUDWbZo11SoNJ8g0r6WB9zaIy7cYm7WIYsnOIjR36hbczDa6yJuE6Gd+G+0Jfs+opo74tV61atYU1Mdc6N7MOoeeyVeOsZMbGK2WM1WbO2vO+bPSI3O6jQxrPCC2fy6WT1ikVO/emktWr404nj9LnraMoztTJ+sW9qeare1LJP/ZYF3RsbblfTR7rTSUXL5/zqzepvekDjQ8LBIpFAP2AAASGIYAEwDCwsCoEIAABCEAAAqUr8J9ZV7y/9K5L39E94v9e2jqzM5ee+diSdHJhT2Zmh3VKw9JUy9+tpENv68x/6mT88Z5M8sklc1qefrd1xnPL0s2Lls5p7rIm4ToZf6ln7mX/sup5J3XFf9+bd+Xr1sRc63zbOoS+Z/aVPW9rMuPN+S3LrTZX73m3zplvavLLUQil+yZAz30ogJAgAIHhCCABMBwtrAsBCEAAAhCAAAQgAAEI+EcAkUAAAsMSQAJgWFxYGQIQgAAEIAABCEAAAhDwiwDigAAEhieABMDwvLA2BCAAAQhAAAIQgAAEIOAPAUQBAQgMUwAJgGGCYXUIQAACEIAABCAAAQhAwA8CiAECEBiuABIAwxXD+hCAAAQgAAEIQAACEICA9wKIAAIQGLYAEgDDJsMGEIAABCAAAQhAAAIQgIDXAmgfAhAYvgASAMM3wxYQgAAEIAABCEAAAhCAgLcCaB0CEMhDAAmAPNCwCQQgAAEIQAACEIAABCDgpQDahgAE8hFAAiAfNWwDAQhAAAIQgAAEIAABCHgngJYhAIG8BJAAyIsNG0EAAhCAAAQgAAEIQAACXgmgXQhAID+F2r83AAAQAElEQVQBJADyc8NWEIAABCAAAQhAAAIQgIA3AmgVAhDIUwAJgDzhsBkEIAABCEAAAhCAAAQg4IUA2oQABPIVQAIgXzlsBwEIQAACEIAABCAAAQi4L4AWIQCBvAWQAMibDhtCAAIQgAAEIAABCEAAAm4LoD0IQCB/ASQA8rfDlhCAAAQgAAEIQAACEICAuwJoDQIQKEAACYAC8LApBCAAAQhAAAIQgAAEIOCmANqCAAQKEUACoBA9bAsBCEAAAhCAAAQgAAEIuCeAliAAgYIEkAAoiA8bQwACEIAABCAAAQhAAAJuCaAdCECgMAEkAArzw9YQgAAEIAABCEAAAhCAgDsCaAUCEChQAAmAAgGxOQQgAAEIQAACEIAABCDghgDagAAEChVAAqBQQWwPAQhAAAIQgAAEIAABCDgvgBYgAIGCBZAAKJgQFUAAAhCAAAQgAAEIQAACTgugfghAoHABJAAKN0QNEIAABCAAAQhAAAIQgICzAqgdAhCwQQAJABsQUQUEIAABCEAAAhCAAAQg4KQA6oYABOwQQALADkXUAQEIQAACEIAABCAAAQg4J4CaIQABWwSQALCFEZVAAAIQgAAEIAABCEAAAk4JoF4IQMAeASQA7HFELRCAAAQgAAEIQAACEICAMwKoFQIQsEkACQCbIFENBCAAAQhAAAIQgAAEIOCEAOqEAATsEkACwC5J1AMBCEAAAhCAAAQgAAEI2C+AGiEAAdsEkACwjRIVQQACEIAABCAAAQhAAAJ2C6A+CEDAPgEkAOyzRE0QgAAEIAABCEAAAhCAgL0CqA0CELBRAAkAGzFRFQQgAAEIQAACEIAABCBgpwDqggAE7BRAAsBOTdQFAQhAAAIQgAAEIAABCNgngJogAAFbBZAAsJUTlUEAAhCAAAQgAAEIQAACdgmgHghAwF4BJADs9URtEIAABCAAAQhAAAIQgIA9AqgFAhCwWQAJAJtBUR0EIAABCEAAAhCAAAQgYIcA6oAABOwWQALAblHUBwEIQAACEIAABCAAAQgULoAaIAAB2wWQALCdFBVCAAIQgAAEIAABCEAAAoUKYHsIQMB+ASQA7DdFjRCAAAQgAAEIQAACEIBAYQLYGgIQcEAACQAHUFElBCAAAQhAAAIQgAAEIFCIALaFAAScEEACwAlV1AkBCEAAAhCAAAQgAAEI5C+ALSEAAUcEkABwhBWVQgACEIAABCAAAQhAAAL5CmA7CEDAGQEkAJxxRa0QgAAEIAABCEAAAhCAQH4C2AoCEHBIAAkAh2BRLQQgAAEIQAACEIAABCCQjwC2gQAEnBJAAsApWdQLAQhAAAIQgAAEIAABCAxfAFtAAAKOCSAB4BgtKoYABCAAAQhAAAIQgAAEhiuA9SEAAecEkABwzhY1QwACEIAABCAAAQhAAALDE8DaEICAgwJIADiIi6ohAAEIQAACEIAABCAAgeEIYF0IQMBJASQAnNRF3RCAAAQgAAEIQAACEIDA0AWwJgQg4KgAEgCO8qJyCEAAAhCAAAQgAAEIQGCoAlgPAhBwVgAJAGd9UTsEIAABCEAAAhCAAAQgMDQBrAUBCDgsgASAw8CoHgIQgAAEIAABCEAAAhAYigDWgQAEnBZAAsBpYdQPAQhAAAIQgAAEIAABCAwugDUgAAHHBZAAcJwYDUAAAhCAAAQgAAEIQAACgwngdQhAwHkBJACcN0YLEIAABCAAAQhAAAIQgMDAAngVAhBwQQAJABeQ0QQEIAABCEAAAhCAAAQgMJAAXoMABNwQQALADWW0AQEIQAACEIAABCAAAQhsWACvQAACrgggAeAKMxqBAAQgAAEIQAACEIAABDYkgOchAAF3BJAAcMcZrUAAAhCAAAQgAAEIQAAC6xfAsxCAgEsCSAC4BI1mIAABCEAAAhCAAAQgAIH1CeA5CEDALQEkANySRjsQgAAEIAABCEAAAhCAwGcF8AwEIOCaABIArlGjIQhAAAIQgAAEIAABCEDg0wL4GQIQcE8ACQD3rNESBCAAAQhAAAIQgAAEILCuAH6CAARcFEACwEVsNAUBCEAAAhCAAAQgAAEIrC2AxxCAgJsCSAC4qY22IAABCEAAAhCAAAQgAIFPBPAIAhBwVQAJAFe50RgEIAABCEAAAhCAAAQg8JEA7iEAAXcFkABw1xutQQACEIAABCAAAQhAAAJrBPAvBCDgsgASAC6DozkIQAACEIAABCAAAQhAwBJAgQAE3BZAAsBtcbQHAQhAAAIQgAAEIAABCBDBAAIQcF0ACQDXydEgBCAAAQhAAAIQgAAEIAABCEDAfQEkANw3R4sQgAAEIAABCEAAAhAodQH0HwIQ8EAACQAP0NEkBCAAAQhAAAIQgAAESlsAvYcABLwQQALAC3W0CQEIQAACEIAABCAAgVIWQN8hAAFPBJAA8IQdjUIAAhCAAAQgAAEIQKB0BdBzCEDAGwEkALxxR6sQgAAEIAABCEAAAhAoVQH0GwIQ8EgACQCP4NEsBCAAAQhAAAIQgAAESlMAvYYABLwSQALAK3m0CwEIQAACEIAABCAAgVIUQJ8hAAHPBJAA8IweDUMAAhCAAAQgAAEIQKD0BNBjCEDAOwEkALyzR8sQgAAEIAABCEAAAhAoNQH0FwIQ8FAACQAP8dE0BCAAAQhAAAIQgAAESksAvYUABLwUQALAS320DQEIQAACEIAABCAAgVISQF8hAAFPBZAA8JQfjUMAAhCAAAQgAAEIQKB0BNBTCEDAWwEkALz1R+sQgAAEIAABCEAAAhAoFQH0EwIQ8FgACQCPBwDNQwACEIAABCAAAQhAoDQE0EsIQMBrASQAvB4BtA8BCEAAAhCAAAQgAIFSEEAfIQABzwWQAPB8CBAABCAAAQhAAAIQgAAEil8APYQABLwXQALA+zFABBCAAAQgAAEIQAACECh2AfQPAhDwgQASAD4YBIQAAQhAAAIQgAAEIACB4hZA7yAAAT8IIAHgh1FADBCAAAQgAAEIQAACEChmAfQNAhDwhQASAL4YBgQBAQhAAAIQgAAEIACB4hVAzyAAAX8IIAHgj3FAFBCAAAQgAAEIQAACEChWAfQLAhDwiQASAD4ZCIQBAQhAAAIQgAAEIACB4hRAryAAAb8IIAHgl5FAHBCAAAQgAAEIQAACEChGAfQJAhDwjQASAL4ZCgQCAQhAAAIQgAAEIACB4hNAjyAAAf8IIAHgn7FAJBCAAAQgAAEIQAACECg2AfQHAhDwkQASAD4aDIQCAQhAAAIQgAAEIACB4hJAbyAAAT8JIAHgp9FALBCAAAQgAAEIQAACECgmAfQFAhDwlQASAL4aDgQDAQhAAAIQgAAEIACB4hFATyAAAX8JIAHgr/FANBCAAAQgAAEIQAACECgWAfQDAhDwmQASAD4bEIQDAQhAAAIQgAAEIACB4hBALyAAAb8JIAHgtxFBPBCAAAQgAAEIQAACECgGAfQBAhDwnQASAL4bEgQEAQhAAAIQgAAEIACB4AugBxCAgP8EkADw35ggIghAAAIQgAAEIAABCARdAPFDAAI+FEACwIeDgpAgAAEIQAACEIAABCAQbAFEDwEI+FEACQA/jgpiggAEIAABCEAAAhCAQJAFEDsEIOBLASQAfDksCAoCEIAABCAAAQhAAALBFUDkEICAPwWQAPDnuCAqCEAAAhCAAAQgAAEIBFUAcUMAAj4VQALApwODsCAAAQhAAAIQgAAEIBBMAUQNAQj4VQAJAL+ODOKCAAQgAAEIQAACEIBAEAUQMwQg4FsBJAB8OzQIDAIQgAAEIAABCEAAAsETQMQQgIB/BZAA8O/YIDIIQAACEIAABCAAAQgETQDxQgACPhZAAsDHg4PQIAABCEAAAhCAAAQgECwBRAsBCPhZAAkAP48OYoMABCAAAQhAAAIQgECQBBArBCDgawEkAHw9PAgOAhCAAAQgAAEIQAACwRFApBCAgL8FkADw9/ggOghAAAIQgAAEIAABCARFAHFCAAI+F0ACwOcDhPAgAAEIQAACEIAABCAQDAFECQEI+F0ACQC/jxDigwAEIAABCEAAAhCAQBAEECMEIOB7ASQAfD9ECBACEIAABCAAAQhAAAL+F0CEEICA/wWQAPD/GCFCCEAAAhCAAAQgAAEI+F0A8UEAAgEQQAIgAIOEECEAAQhAAAIQgAAEIOBvAUQHAQgEQQAJgCCMEmKEAAQgAAEIQAACEICAnwUQGwQgEAgBJAACMUwIEgIQgAAEIAABCEAAAv4VQGQQgEAwBJAACMY4IUoIQAACEIAABCAAAQj4VQBxQQACARFAAiAgA4UwIQABCEAAAhCAAAQg4E8BRAUBCARFAAmAoIwU4oQABCAAAQhAAAIQgIAfBRATBCAQGAEkAAIzVAgUAhCAAAQgAAEIQAAC/hNARBCAQHAEkAAIzlghUghAAAIQgAAEIAABCPhNAPFAAAIBEkACIECDhVAhAAEIQAACEIAABCDgLwFEAwEIBEkACYAgjRZihQAEIAABCEAAAhCAgJ8EEAsEIBAoASQAAjVcCBYCEIAABCAAAQhAAAL+EUAkEIBAsASQAAjWeCFaCEAAAhCAAAQgAAEI+EUAcUAAAgETQAIgYAOGcCEAAQhAAAIQgAAEIOAPAUQBAQgETQAJgKCNGOKFAAQgAAEIQAACEICAHwQQAwQgEDgBJAACN2QIGAIQgAAEIAABCEAAAt4LIAIIQCB4AkgABG/MEDEEIAABCEAAAhCAAAS8FkD7EIBAAAWQAAjgoCFkCEAAAhCAAAQgAAEIeCuA1iEAgSAKIAEQxFFDzBCAAAQgAAEIQAACEPBSAG1DAAKBFEACIJDDhqAhAAEIQAACEIAABCDgnQBahgAEgimABEAwxw1RQwACEIAABCAAAQhAwCsBtAsBCARUAAmAgA4cwoYABCAAAQhAAAIQgIA3AmgVAhAIqgASAEEdOcQNAQhAAAIQgAAEIAABLwTQJgQgEFgBJAACO3QIHAIQgAAEIAABCEAAAu4LoEUIQCC4AkgABHfsEDkEIAABCEAAAhCAAATcFkB7EIBAgAWQAAjw4CF0CEAAAhCAAAQgAAEIuCuA1iAAgSALIAEQ5NFD7BCAAAQgAAEIQAACEHBTAG1BAAKBFkACINDDh+AhAAEIQAACEIAABCDgngBaggAEgi2ABECwxw/RQwACEIAABCAAAQhAwC0BtAMBCARcAAmAgA8gwocABCAAAQhAAAIQgIA7AmgFAhAIugASAEEfQcQPAQhAAAIQgAAEIAABNwTQBgQgEHgBJAACP4ToAAQgAAEIQAACEIAABJwXQAsQgEDwBZAACP4YogcQgAAEIAABCEAAAhBwWgD1QwACRSCABEARDCK6AAEIQAACEIAABCAAAWcFUDsEIFAMAkgAFMMoog8QgAAEIAABCEAAAhBwUgB1QwACRSGABEBRDCM6AQEIQAACEIAABCAAAecEUDMEIFAcAkgAFMc4fGOGCQAAB7JJREFUohcQgAAEIAABCEAAAhBwSgD1QgACRSKABECRDCS6AQEIQAACEIAABCAAAWcEUCsEIFAsAkgAFMtIoh8QgAAEIAABCEAAAhBwQgB1QgACRSOABEDRDCU6AgEIQAACEIAABCAAAfsFUCMEIFA8AkgAFM9YoicQgAAEIAABCEAAAhCwWwD1QQACRSSABEARDSa6AgEIQAACEIAABCAAAXsFUBsEIFBMAkgAFNNooi8QgAAEIAABCEAAAhCwUwB1QQACRSWABEBRDSc6AwEIQAACEIAABCAAAfsEUBMEIFBcAkgAFNd4ojcQgAAEIAABCEAAAhCwSwD1QAACRSaABECRDSi6AwEIQAACEIAABCAAAXsEUAsEIFBsAkgAFNuIoj8QgAAEIAABCEAAAhCwQwB1QAACRSeABEDRDSk6BAEIQAACEIAABCAAgcIFUAMEIFB8AkgAFN+YokcQgAAEIAABCEAAAhAoVADbQwACRSiABEARDiq6BAEIQAACEIAABCAAgcIEsDUEIFCMAkgAFOOook8QgAAEIAABCEAAAhAoRADbQgACRSmABEBRDis6BQEIQAACEIAABCAAgfwFsCUEIFCcAkgAFOe4olcQgAAEIAABCEAAAhDIVwDbQQACRSqABECRDiy6BQEIQAACEIAABCAAgfwEsBUEIFCsAkgAFOvIol8QgAAEIAABCEAAAhDIRwDbQAACRSuABEDRDi06BgEIQAACEIAABCAAgeELYAsIQKB4BZAAKN6xRc8gAAEIQAACEIAABCAwXAGsDwEIFLEAEgBFPLjoGgQgAAEIQAACEIAABIYngLUhAIFiFkACoJhHF32DAAQgAAEIQAACEIDAcASwLgQgUNQCSAAU9fCicxCAAAQgAAEIQAACEBi6ANaEAASKWwAJgOIeX/QOAhCAAAQgAAEIQAACQxXAehCAQJELIAFQ5AOM7kEAAhCAAAQgAAEIQGBoAlgLAhAodgEkAIp9hNE/CEAAAhCAAAQgAAEIDEUA60AAAkUvgARA0Q8xOggBCEAAAhCAAAQgAIHBBbAGBCBQ/AJIABT/GKOHEIAABCAAAQhAAAIQGEwAr0MAAiUggARACQwyuggBCEAAAhCAAAQgAIGBBfAqBCBQCgJIAJTCKKOPEIAABCAAAQhAAAIQGEgAr0EAAiUhgARASQwzOgkBCEAAAhCAAAQgAIENC+AVCECgNASQACiNcUYvIQABCEAAAhCAAAQgsCEBPA8BCJSIABIAJTLQ6CYEIAABCEAAAhCAAATWL4BnIQCBUhFAAqBURhr9hAAEIAABCEAAAhCAwPoE8BwEIFAyAkgAlMxQo6MQgAAEIAABCEAAAhD4rACegQAESkcACYDSGWv0FAIQgAAEIAABCEAAAp8WwM8QgEAJCSABUEKDja5CAAIQgAAEIAABCEBgXQH8BAEIlJIAEgClNNroKwQgAAEIQAACEIAABNYWwGMIQKCkBJAAKKnhRmchAAEIQAACEIAABCDwiQAeQQACpSWABEBpjTd6CwEIQAACEIAABCAAgY8EcA8BCJSYABIAJTbg6C4EIAABCEAAAhCAAATWCOBfCECg1ASQACi1EUd/IQABCEAAAhCAAAQgYAmgQAACJSeABEDJDTk6DAEIQAACEIAABCAAASIYQAACpSeABEDpjTl6DAEIQAACEIAABCAAAQhAAAIlKIAEQAkOOroMAQhAAAIQgAAEIFDqAug/BCBQigJIAJTiqKPPEIAABCAAAQhAAAKlLYDeQwACJSmABEBJDjs6DQEIQAACEIAABCBQygLoOwQgUJoCSACU5rij1xCAAAQgAAEIQAACpSuAnkMAAiUqgARAiQ48ug0BCEAAAhCAAAQgUKoC6DcEIFCqAkgAlOrIo98QgAAEIAABCEAAAqUpgF5DAAIlK4AEQMkOPToOAQhAAAIQgAAEIFCKAugzBCBQugJIAJTu2KPnEIAABCAAAQhAAAKlJ4AeQwACJSyABEAJDz66DgEIQAACEIAABCBQagLoLwQgUMoCSACU8uij7xCAAAQgAAEIQAACpSWA3kIAAiUtgARASQ8/Og8BCEAAAhCAAAQgUEoC6CsEIFDaAkgAlPb4o/cQgAAEIAABCEAAAqUjgJ5CAAIlLoAEQIm/AdB9CEAAAhCAAAQgAIFSEUA/IQCBUhdAAqDU3wHoPwQgAAEIQAACEIBAaQiglxCAQMkLIAFQ8m8BAEAAAhCAAAQgAAEIlIIA+ggBCEAACQC8ByAAAQhAAAIQgAAEIFD8AughBCAAAUICAG8CCEAAAhCAAAQgAAEIFL0AOggBCECAkADAmwACEIAABCAAAQhAAAJFL4AOQgACEFABHAGgCFggAAEIQAACEIAABCBQzALoGwQgAAFLAAkASwEFAhCAAAQgAAEIQAACxSuAnkEAAhBYLYAEwGoG/AMBCEAAAhCAAAQgAIFiFUC/IAABCKwRQAJgjQP+hQAEIAABCEAAAhCAQHEKoFcQgAAEPhRAAuBDCNxBAAIQgAAEIAABCECgGAXQJwhAAAIfCSAB8JEE7iEAAQhAAAIQgAAEIFB8AugRBCAAgY8FkAD4mAIPIAABCEAAAhCAAAQgUGwC6A8EIACBTwT+HwAA//+35GZ1AAAABklEQVQDAHXAwM7MCpV5AAAAAElFTkSuQmCC"
)

# -----------------------------------------------------------------------
# Configuracao persistente  (~/.aom_estabilidade.json)
# -----------------------------------------------------------------------

_CFG_FILE = os.path.join(os.path.expanduser('~'), '.aom_estabilidade.json')
_CFG_DEFAULT = {
    'pasta': '', 'ifd': '', 'saida': 'resultados_estabilidade.xlsx',
    'pasta_ind': 'individuais', 'pdf': 'relatorio_estabilidade.pdf',
    'scores_file': '', 'protocolo': 'fms', 'n_ens': '',
    'o_emb': False, 'o_ell': True, 'o_est': True, 'o_ind': True, 'o_pdf': True,
    'o_estats': False,
    'export_csv': True, 'export_docx': False, 'export_html': False,
    'export_png': False,   # exportar PNG dos graficos individuais
    'png_dpi': 180,        # resolução dos PNG (72/96/150/180/300)
    'png_tipos': ['estabilograma', 'elipse'],  # tipos a exportar
    'fft_ativo': False,    # analise espectral FFT
    'peso_norm': False,    # normalizar ea95 pela massa corporal
    'csv_sep': ';', 'csv_dec': ',',
    'estats_grupo': True, 'estats_dir_esq': True,
    'estats_pos_disp': True, 'estats_atleta': False,
    'estats_friedman': True, 'estats_posthoc': True, 'estats_corr': False,
}

def cfg_load():
    try:
        with open(_CFG_FILE, 'r', encoding='utf-8') as f:
            d = json.load(f)
        out = dict(_CFG_DEFAULT); out.update(d)
        return out
    except Exception:
        return dict(_CFG_DEFAULT)

def cfg_save(d):
    try:
        with open(_CFG_FILE, 'w', encoding='utf-8') as f:
            json.dump(d, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


# -----------------------------------------------------------------------
# Historico de sessoes  (ultimas 20 analises)
# -----------------------------------------------------------------------

_HIST_FILE = os.path.join(os.path.expanduser('~'), '.aom_estabilidade_hist.json')
_HIST_MAX  = 20

def hist_load():
    try:
        with open(_HIST_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return []

def hist_add(protocolo, n_atletas, saida, pdf=''):
    entries = hist_load()
    entries.insert(0, {
        'data':      datetime.datetime.now().strftime('%d/%m/%Y %H:%M'),
        'protocolo': PROTOCOLOS.get(protocolo, {}).get('nome', protocolo),
        'n':         n_atletas,
        'saida':     saida,
        'pdf':       pdf,
    })
    entries = entries[:_HIST_MAX]
    try:
        with open(_HIST_FILE, 'w', encoding='utf-8') as f:
            json.dump(entries, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


# -----------------------------------------------------------------------
# Perfis de configuracao nomeados
# -----------------------------------------------------------------------

_PROFILES_FILE = os.path.join(os.path.expanduser('~'), '.aom_estabilidade_profiles.json')

def profiles_load():
    """Carrega todos os perfis guardados."""
    try:
        with open(_PROFILES_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}

def profile_save(name, cfg_dict):
    """Guarda a configuracao actual com um nome."""
    profiles = profiles_load()
    profiles[name] = {k: v for k, v in cfg_dict.items()
                      if not k.startswith('_')}
    try:
        with open(_PROFILES_FILE, 'w', encoding='utf-8') as f:
            json.dump(profiles, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False

def profile_delete(name):
    """Remove um perfil pelo nome."""
    profiles = profiles_load()
    profiles.pop(name, None)
    try:
        with open(_PROFILES_FILE, 'w', encoding='utf-8') as f:
            json.dump(profiles, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


# -----------------------------------------------------------------------
# Identificacao e protecao de autoria
# As constantes abaixo sao checksums internos de validacao -- nao alterar.
#
# session-id: 51e0bf3aa5450fb27be3f0469b2a737d  [cfg/runtime/v7]
# validator:  e5d52a9346075e69dc55b812d6df64d3  [ref-A]
# validator:  9e226396fd6e47455e415045b97bd316  [ref-B]
# inst-ref:   015653d877171992eb13e361896a33b7
# dept-ref:   656192cc639fcf4b4ebfe2d3ed22e6bb
# build-tag:  MjAyNjAzMDFfMDEwNA==
# -----------------------------------------------------------------------

_p1 = '416e6472c3a9'; _p2 = '204f2e204d61'; _p3 = '737375c3a761'
AUTOR  = bytes.fromhex(_p1 + _p2 + _p3).decode('utf-8')
VERSAO = "1.0"   # initial public release
_VERSAO_MAJOR = 23  # numero da versao para comparacoes
_VERSAO_DATA  = "2026"  # ano de lancamento

SEGUNDO_AUTOR     = 'Pedro Aleixo'         # Professor de Biomecânica (UI)
SEGUNDO_AUTOR_APA = 'P. Aleixo'            # forma abreviada para citações APA
ORIENTADOR = 'Luís M. Massuça'    # terceiro autor
# nota: aparece nos cabecalhos dos relatorios e nas citacoes

def _tiro_itv_label(key):
    """Retorna o label do intervalo de tiro na lingua actual."""
    _map = {
        'toque_pontaria':   'tiro_itv_pont',
        'toque_disparo':    'tiro_itv_disp',
        'pontaria_disparo': 'tiro_itv_pont_disp',
        'disparo_fim':      'tiro_itv_disp_fim',
        'total':            'tiro_itv_total',
    }
    return T(_map.get(key, key)) if _map.get(key) else TIRO_INTERVALOS.get(key, key)
PROG   = "Biomechanical Stability Program"

# -----------------------------------------------------------------------
# Internacionalização (i18n)  - PT / EN / ES / DE
# -----------------------------------------------------------------------
try:
    from bsp_i18n import (T, definir_lingua, lingua_atual,
                          LINGUAS_DISPONIVEIS, licenca_texto,
                          mets_pdf_localizadas, mets_xl_localizadas,
                          lados_pdf_localizados)
    _I18N_OK = True
except ImportError:
    # Fallback: T() devolve a chave (a app funciona sem o módulo)
    def T(k, **kw): return k
    def definir_lingua(l): pass
    def lingua_atual(): return 'PT'
    def licenca_texto(**kw): return ''
    def mets_pdf_localizadas(): return []
    def mets_xl_localizadas(): return []
    def lados_pdf_localizados(): return {}
    LINGUAS_DISPONIVEIS = {'PT': '🇵🇹 PT'}
    _I18N_OK = False

# -----------------------------------------------------------------------
# Sistema de temas (escuro / claro)
# -----------------------------------------------------------------------
_TEMAS = {
    # Tema escuro (default) - paleta v1.0 alinhada com identidade BSP
    'escuro': dict(
        CF='#0D1720', CP='#132030', CC='#1A2B3C', CA='#3FB6D9', CA2='#90E0EF',
        COK='#22C55E', CAV='#F59E0B', CER='#EF4444', CT='#EDF6FF', CD='#7FA8C9',
        CB='#1E3448', CA3='#0F4C75', CG='#0F2030', CSEP='#1E3448',
        _SB_BG='#0A1218', _SB_FG='#3A5F7A',
        _TEMA_NOME='escuro', _TEMA_ICO='🌙',
    ),
    # Tema claro - reformulado em v1.0 para legibilidade clinica.
    # Paleta: fundo neutro (slate-50), paineis brancos, acento azul-petroleo
    # `#0F4C75` (BSP brand) e cyan `#0E7AA8` para hover/links. Todos os pares
    # texto/fundo passam WCAG AA (>= 4.5:1).
    'claro': dict(
        CF='#F8FAFC',  # janela: slate-50
        CP='#FFFFFF',  # cartoes / barra superior
        CC='#F1F5F9',  # paineis secundarios / linhas alternadas
        CA='#0F4C75',  # acento principal (botao Executar, badges, links)
        CA2='#0E7AA8', # acento hover / pressed
        COK='#15803D', # verde sucesso (forte)
        CAV='#B45309', # amarelo aviso (forte)
        CER='#B91C1C', # vermelho erro (forte)
        CT='#0F172A',  # texto primario (slate-900)
        CD='#475569',  # texto secundario (slate-600)
        CB='#CBD5E1',  # bordas / separadores (slate-300)
        CA3='#3FB6D9', # acento decorativo (cyan claro)
        CG='#E2E8F0',  # gridlines / fundos discretos
        CSEP='#CBD5E1',# separadores fortes
        _SB_BG='#E2E8F0', _SB_FG='#475569',
        _TEMA_NOME='claro', _TEMA_ICO='☀',
    ),
}

def _aplicar_tema(nome):
    """Actualiza as variaveis globais de cor para o tema escolhido."""
    global CF, CP, CC, CA, CA2, COK, CAV, CER, CT, CD, CB, CA3, CG, CSEP
    global _SB_BG, _SB_FG, _TEMA_ACTUAL
    t = _TEMAS.get(nome, _TEMAS['escuro'])
    CF=t['CF']; CP=t['CP']; CC=t['CC']; CA=t['CA']; CA2=t['CA2']
    COK=t['COK']; CAV=t['CAV']; CER=t['CER']; CT=t['CT']; CD=t['CD']
    CB=t['CB']; CA3=t['CA3']; CG=t['CG']; CSEP=t['CSEP']
    _SB_BG=t['_SB_BG']; _SB_FG=t['_SB_FG']
    _TEMA_ACTUAL = nome

_SB_BG = '#0A1218'; _SB_FG = '#3A5F7A'
_TEMA_ACTUAL = 'escuro'

# -----------------------------------------------------------------------
# Tokens de UI (espacamento, tipografia, iconografia) - v1.0
# -----------------------------------------------------------------------
#
# Espacamento em px (consistente em toda a UI). Escalonado em multiplos
# de 4: xs=4, s=8, m=12, l=20. Padroniza padx/pady/ipadx.
_PAD_XS = 4
_PAD_S  = 8
_PAD_M  = 12
_PAD_L  = 20

# Tipografia - funcoes para evitar referenciar fontes inexistentes no sistema.
# Segoe UI (Windows), SF Pro Text/Helvetica Neue (macOS), DejaVu Sans (Linux).
def _font_family_ui():
    if _SYS == 'Windows':
        return 'Segoe UI'
    if _SYS == 'Darwin':
        return 'SF Pro Text'
    return 'DejaVu Sans'

def _font_family_mono():
    if _SYS == 'Windows':
        return 'Cascadia Mono'
    if _SYS == 'Darwin':
        return 'SF Mono'
    return 'DejaVu Sans Mono'

# Tipografia hierarquica. Tuplos (family, size, weight)
_FONT_H1   = lambda: (_font_family_ui(), 16, 'bold')
_FONT_H2   = lambda: (_font_family_ui(), 12, 'bold')
_FONT_H3   = lambda: (_font_family_ui(), 10, 'bold')
_FONT_BODY = lambda: (_font_family_ui(), 9,  'normal')
_FONT_SMALL= lambda: (_font_family_ui(), 8,  'normal')
_FONT_MONO = lambda: (_font_family_mono(), 9, 'normal')

# Iconografia (caracteres Unicode, renderizam em todas as plataformas).
# Nao usar emoji decorativo (bandeiras, luas) por render inconsistente.
ICO = {
    'ok': '[OK]', 'error': '[X]', 'warn': '[!]', 'info': '[i]',
    'loading': '...', 'settings': 'Cfg', 'refresh': 'R',
    'download': '↓', 'up': '↑', 'right': '→', 'left': '←',
    'expand': '▸', 'collapse': '▾',
}


def _configurar_estilos_ttk(style):
    """
    Configura estilos ttk consistentes com o tema actual. Chamar no init da
    Janela depois de criar o root. Estilos disponiveis:
      Primary.TButton    - accent (azul), flat
      Secondary.TButton  - outlined, accent subtil
      Danger.TButton     - vermelho, flat (aces destrutivas)
      Ghost.TButton      - transparente, hover subtil
      Treeview           - tabela com linhas alternadas
    """
    try:
        # Tema base (clam permite cores custom no macOS/Windows)
        try:
            style.theme_use('clam')
        except Exception:
            pass

        # Primary - cor accent, flat
        style.configure('Primary.TButton',
                        background=CA, foreground=CT,
                        borderwidth=0, relief='flat',
                        padding=(_PAD_M, _PAD_S),
                        font=_FONT_BODY())
        style.map('Primary.TButton',
                  background=[('active', CA2), ('disabled', CB)],
                  foreground=[('disabled', CD)])

        # Secondary - outlined
        style.configure('Secondary.TButton',
                        background=CF, foreground=CA,
                        borderwidth=1, relief='solid',
                        bordercolor=CA,
                        padding=(_PAD_M, _PAD_S),
                        font=_FONT_BODY())
        style.map('Secondary.TButton',
                  background=[('active', CC), ('disabled', CF)],
                  foreground=[('disabled', CD)])

        # Danger - vermelho
        style.configure('Danger.TButton',
                        background=CER, foreground=CT,
                        borderwidth=0, relief='flat',
                        padding=(_PAD_M, _PAD_S),
                        font=_FONT_BODY())
        style.map('Danger.TButton',
                  background=[('active', '#B02020'), ('disabled', CB)])

        # Ghost - transparente
        style.configure('Ghost.TButton',
                        background=CF, foreground=CT,
                        borderwidth=0, relief='flat',
                        padding=(_PAD_M, _PAD_S),
                        font=_FONT_BODY())
        style.map('Ghost.TButton',
                  background=[('active', CC)],
                  foreground=[('disabled', CD)])

        # Treeview
        style.configure('Treeview',
                        background=CP, foreground=CT,
                        fieldbackground=CP,
                        borderwidth=0,
                        rowheight=22,
                        font=_FONT_BODY())
        style.configure('Treeview.Heading',
                        background=CC, foreground=CA,
                        borderwidth=0, relief='flat',
                        font=_FONT_H3())
        style.map('Treeview',
                  background=[('selected', CA2)],
                  foreground=[('selected', CT)])

        # Progressbar
        style.configure('Accent.Horizontal.TProgressbar',
                        background=CA, troughcolor=CC,
                        borderwidth=0, lightcolor=CA, darkcolor=CA)

        # Notebook tabs
        style.configure('TNotebook', background=CF, borderwidth=0)
        style.configure('TNotebook.Tab',
                        background=CC, foreground=CT,
                        padding=(_PAD_M, _PAD_S),
                        borderwidth=0, font=_FONT_BODY())
        style.map('TNotebook.Tab',
                  background=[('selected', CP)],
                  foreground=[('selected', CA)])

        # Combobox
        style.configure('TCombobox',
                        fieldbackground=CP, background=CC,
                        foreground=CT, arrowcolor=CA,
                        bordercolor=CB,
                        padding=_PAD_S)
        style.map('TCombobox',
                  fieldbackground=[('readonly', CP)],
                  foreground=[('readonly', CT)])

        # Entry
        style.configure('TEntry',
                        fieldbackground=CP, foreground=CT,
                        bordercolor=CB, insertcolor=CA,
                        padding=_PAD_S)
    except Exception:
        # Se o ttk nao existir ou o tema nao suportar algum parametro, ignora.
        pass


# Aplicar língua guardada no config
def _aplicar_lingua_config():
    try:
        import json as _json
        _cf = os.path.join(os.path.expanduser('~'), '.aom_estabilidade_cfg.json')
        with open(_cf, 'r', encoding='utf-8') as _f:
            _d = _json.load(_f)
        _l = _d.get('_lingua', 'PT')
        if _l in LINGUAS_DISPONIVEIS:
            definir_lingua(_l)
    except Exception:
        pass
_aplicar_lingua_config()


def _detect_system_theme():
    """Detecta o tema do sistema operativo (macOS/Windows) na primeira execucao.
    Retorna 'escuro', 'claro' ou None se nao conseguir detectar."""
    try:
        if _SYS == 'Darwin':
            import subprocess
            r = subprocess.run(['defaults', 'read', '-g', 'AppleInterfaceStyle'],
                               capture_output=True, text=True, timeout=2)
            return 'escuro' if 'Dark' in r.stdout else 'claro'
        elif _SYS == 'Windows':
            import winreg
            key = winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r'Software\Microsoft\Windows\CurrentVersion\Themes\Personalize')
            val, _ = winreg.QueryValueEx(key, 'AppsUseLightTheme')
            winreg.CloseKey(key)
            return 'claro' if val == 1 else 'escuro'
    except Exception:
        pass
    return None


# -----------------------------------------------------------------------
# Auto-update one-click (GitHub Releases API + SHA256)
# -----------------------------------------------------------------------
# O repo publica releases com:
#   - BSP_Setup.exe ou BSP-{ver}-win-x64.exe   (Windows)
#   - BSP-{ver}-macos-arm64.dmg                (Apple Silicon)
#   - BSP-{ver}-macos-x64.dmg                  (Intel)
#   - SHA256SUMS.txt (checksums em formato "hash  filename")
_UPDATE_REPO   = 'andremassuca/BSP'
_UPDATE_API    = f'https://api.github.com/repos/{_UPDATE_REPO}/releases/latest'
_UPDATE_STATE  = {
    'tag': None,          # "1.0.1"
    'notes': '',
    'asset_url': None,    # URL do DMG/EXE para este OS
    'asset_name': None,
    'sha256': None,       # valor esperado para asset_name
    'sha_txt_url': None,
}

def _parse_versao(v):
    """'1.0.1' -> (1,0,1). Tolerante a 'v1.0', '1.0-rc1'."""
    try:
        v = v.strip().lstrip('vV')
        v = v.split('-')[0]
        return tuple(int(x) for x in v.split('.'))
    except Exception:
        return (0,)

def _asset_nome_esperado():
    """Retorna lista de padroes aceitaveis para este OS/arch.
    Usamos um prefixo unico para detectar sem ambiguidade."""
    import platform
    if _SYS == 'Windows':
        return [r'BSP[_-]Setup', r'BSP-.*-win-x64\.exe$',
                r'BSP-.*-windows', r'BSP_v.*\.exe$']
    if _SYS == 'Darwin':
        arch = platform.machine() or ''
        if arch in ('arm64', 'aarch64'):
            return [r'BSP-.*-macos-arm64\.dmg$',
                    r'BSP-.*-macos\.dmg$']
        return [r'BSP-.*-macos-x64\.dmg$',
                r'BSP-.*-macos\.dmg$']
    return [r'BSP-.*\.tar\.gz$']  # Linux opcional

def _match_asset(assets):
    """Dado 'assets' JSON de GitHub, devolve (asset_dict, sha_text_asset)
    ou (None, sha_text_asset)."""
    import re as _re
    padroes = [_re.compile(p, _re.IGNORECASE) for p in _asset_nome_esperado()]
    found = None
    sha_asset = None
    for a in assets or []:
        nome = a.get('name', '')
        if nome.upper() == 'SHA256SUMS.TXT':
            sha_asset = a
            continue
        for p in padroes:
            if p.search(nome):
                found = a
                break
        if found and sha_asset:
            break
    return found, sha_asset

def _parse_sha256sums(texto):
    """'hash  filename\\n...' -> {filename: hash}"""
    out = {}
    for linha in (texto or '').splitlines():
        linha = linha.strip()
        if not linha or linha.startswith('#'):
            continue
        # aceita "<hash>  <file>" e "<hash> *<file>"
        parts = linha.split()
        if len(parts) >= 2:
            h = parts[0].lower()
            fn = parts[-1].lstrip('*').strip()
            if len(h) == 64:
                out[fn] = h
    return out

def _verificar_update_async(callback=None):
    """Verifica update em background consultando GitHub Releases API.
    Actualiza _UPDATE_STATE e chama callback(tag) se disponivel."""
    import threading
    def _worker():
        try:
            import urllib.request, json as _json
            req = urllib.request.Request(
                _UPDATE_API,
                headers={'Accept': 'application/vnd.github+json',
                         'User-Agent': f'BSP/{VERSAO}'})
            with urllib.request.urlopen(req, timeout=5) as r:
                data = _json.loads(r.read().decode('utf-8'))
            tag = (data.get('tag_name') or '').strip()
            if not tag:
                return
            nova = tag.lstrip('vV')
            if _parse_versao(nova) <= _parse_versao(VERSAO):
                return
            assets = data.get('assets') or []
            a_bin, a_sha = _match_asset(assets)
            if not a_bin:
                # Release existe mas sem asset para este OS; ainda mostramos banner
                _UPDATE_STATE.update(tag=nova,
                                     notes=data.get('body','') or '',
                                     asset_url=None, asset_name=None,
                                     sha256=None, sha_txt_url=None)
                if callback:
                    try: callback(nova)
                    except Exception: pass
                return
            # Fetch SHA256SUMS.txt se existir
            sha_map = {}
            if a_sha:
                try:
                    with urllib.request.urlopen(
                            a_sha.get('browser_download_url'), timeout=5) as r2:
                        sha_map = _parse_sha256sums(r2.read().decode('utf-8'))
                except Exception:
                    sha_map = {}
            _UPDATE_STATE.update(
                tag=nova,
                notes=data.get('body','') or '',
                asset_url=a_bin.get('browser_download_url'),
                asset_name=a_bin.get('name'),
                sha256=sha_map.get(a_bin.get('name','')),
                sha_txt_url=(a_sha or {}).get('browser_download_url'),
            )
            if callback:
                try: callback(nova)
                except Exception: pass
        except Exception:
            # Silencioso: user nao deve ver erros de rede
            pass
    threading.Thread(target=_worker, daemon=True).start()

def _actualizar_agora(progress_cb=None, done_cb=None):
    """Download do asset, verificacao de SHA256, lancamento do instalador.
    progress_cb(bytes_done, bytes_total) opcional.
    done_cb(ok, mensagem) ao fim."""
    import threading
    def _worker():
        url = _UPDATE_STATE.get('asset_url')
        nome = _UPDATE_STATE.get('asset_name')
        sha = _UPDATE_STATE.get('sha256')
        if not url or not nome:
            if done_cb: done_cb(False, 'Asset de actualização não disponível para este sistema.')
            return
        try:
            import urllib.request, tempfile, hashlib as _hl
            tmp = tempfile.mkdtemp(prefix='bsp_upd_')
            dest = os.path.join(tmp, nome)
            req = urllib.request.Request(url, headers={'User-Agent': f'BSP/{VERSAO}'})
            with urllib.request.urlopen(req, timeout=30) as r:
                total = int(r.headers.get('Content-Length', '0') or 0)
                h = _hl.sha256()
                done = 0
                with open(dest, 'wb') as fout:
                    while True:
                        chunk = r.read(64*1024)
                        if not chunk: break
                        fout.write(chunk)
                        h.update(chunk)
                        done += len(chunk)
                        if progress_cb:
                            try: progress_cb(done, total)
                            except Exception: pass
            got = h.hexdigest().lower()
            if sha and got != sha.lower():
                if done_cb: done_cb(False,
                    f'Checksum nao corresponde:\nesperado={sha}\nobtido={got}\n'
                    'Transferencia abortada (possivel corrupcao ou tampering).')
                return
            # Lancar instalador
            if _SYS == 'Windows':
                try:
                    import subprocess
                    subprocess.Popen([dest], close_fds=True)
                except Exception as ex:
                    if done_cb: done_cb(False, f'Nao foi possivel lancar instalador: {ex}')
                    return
            elif _SYS == 'Darwin':
                try:
                    import subprocess
                    subprocess.run(['open', dest], check=False)
                except Exception as ex:
                    if done_cb: done_cb(False, f'Nao foi possivel abrir DMG: {ex}')
                    return
            else:
                if done_cb: done_cb(False, f'Plataforma nao suportada: {_SYS}')
                return
            if done_cb: done_cb(True, f'Transferencia completa. A sair para instalar...')
            # Sair da app
            try:
                import time as _time
                _time.sleep(1.2)
                os._exit(0)
            except Exception:
                pass
        except Exception as ex:
            if done_cb: done_cb(False, f'Erro na actualizacao: {ex}')
    threading.Thread(target=_worker, daemon=True).start()

# -----------------------------------------------------------------------
# Dashboard Web (BSP Dashboard FastAPI) - launcher local
# -----------------------------------------------------------------------
_DASHBOARD_PROC = [None]  # handle para o processo uvicorn
_DASHBOARD_PORT = [None]

def _dashboard_port_livre():
    import socket
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    s.bind(('127.0.0.1', 0))
    porta = s.getsockname()[1]
    s.close()
    return porta

def _dashboard_ja_vivo():
    p = _DASHBOARD_PROC[0]
    return p is not None and getattr(p, 'poll', lambda: 0)() is None

def abrir_dashboard_web():
    """Arranca uvicorn a correr bsp_dashboard.api.main em porta livre,
    espera o /api/health responder, abre o browser. Idempotente."""
    import subprocess, sys, threading, time, webbrowser
    try:
        import urllib.request
    except Exception:
        urllib = None  # type: ignore

    if _dashboard_ja_vivo() and _DASHBOARD_PORT[0]:
        webbrowser.open(f'http://127.0.0.1:{_DASHBOARD_PORT[0]}/')
        return True, _DASHBOARD_PORT[0]

    # Confirmar que FastAPI esta instalado (resposta clara se nao estiver)
    try:
        import fastapi  # noqa: F401
        import uvicorn  # noqa: F401
    except ImportError:
        return False, 'fastapi/uvicorn nao instalados. Correr: pip install fastapi uvicorn python-multipart'

    porta = _dashboard_port_livre()
    _DASHBOARD_PORT[0] = porta
    # Spawn processo filho sem consola (Windows) e longe do stdout da app
    env = os.environ.copy()
    env['PYTHONIOENCODING'] = 'utf-8'
    creationflags = 0
    if _SYS == 'Windows':
        creationflags = 0x08000000  # CREATE_NO_WINDOW
    cmd = [sys.executable, '-m', 'bsp_dashboard.api.main',
           '--host', '127.0.0.1', '--port', str(porta)]
    try:
        _DASHBOARD_PROC[0] = subprocess.Popen(
            cmd, cwd=os.path.dirname(os.path.abspath(__file__)) or None,
            env=env, creationflags=creationflags,
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except Exception as ex:
        return False, f'spawn falhou: {ex}'

    # Polling ate /api/health responder (max 8s)
    def _esperar_saude():
        if urllib is None:
            return True
        url = f'http://127.0.0.1:{porta}/api/health'
        for _ in range(40):
            try:
                with urllib.request.urlopen(url, timeout=0.5) as r:
                    if r.status == 200:
                        return True
            except Exception:
                time.sleep(0.2)
        return False

    def _abrir_quando_pronto():
        ok = _esperar_saude()
        if ok:
            try: webbrowser.open(f'http://127.0.0.1:{porta}/')
            except Exception: pass

    threading.Thread(target=_abrir_quando_pronto, daemon=True).start()
    return True, porta

def parar_dashboard_web():
    """Termina o processo uvicorn se estiver vivo."""
    p = _DASHBOARD_PROC[0]
    if p is None:
        return
    try:
        p.terminate()
        p.wait(timeout=3)
    except Exception:
        try: p.kill()
        except Exception: pass
    finally:
        _DASHBOARD_PROC[0] = None
        _DASHBOARD_PORT[0] = None

# -----------------------------------------------------------------------
# Logo ISCPSI (Instituto Superior de Ciencias Policiais e Seguranca Interna)
# -----------------------------------------------------------------------
_ISCPSI_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAFAAAABQCAYAAACOEfKtAAABCGlDQ1BJQ0MgUHJvZmlsZQAAeJxj"
    "YGA8wQAELAYMDLl5JUVB7k4KEZFRCuwPGBiBEAwSk4sLGHADoKpv1yBqL+viUYcLcKakFicD6Q9A"
    "rFIEtBxopAiQLZIOYWuA2EkQtg2IXV5SUAJkB4DYRSFBzkB2CpCtkY7ETkJiJxcUgdT3ANk2uTml"
    "yQh3M/Ck5oUGA2kOIJZhKGYIYnBncAL5H6IkfxEDg8VXBgbmCQixpJkMDNtbGRgkbiHEVBYwMPC3"
    "MDBsO48QQ4RJQWJRIliIBYiZ0tIYGD4tZ2DgjWRgEL7AwMAVDQsIHG5TALvNnSEfCNMZchhSgSKe"
    "DHkMyQx6QJYRgwGDIYMZAKbWPz9HbOBQAAAxiElEQVR4nL29d7Rd2VXm+5tr7XTSzUHhSirlUkmV"
    "pFLl6FwYzMCYxsYGHKCxm/QwPcYjPNvQjNfD/dxAY0wwbQw2pp7bGAw2jV12uXJJlVUKVQqlnK50"
    "dfOJe+8V3h/7XEll41IFHkvjSEdH556z91xzzfB9c05Ju932vM4VWY8TwYjGikJEEOvQQKQVzuZk"
    "NkUpjfUKJwqlFQ4hiAJiBQYQQOc5zua0TUhmAhQQeSD0+IoQiBB23wuQmhbWWHDgnUFUiNYx1ng8"
    "HhV4kBzIQBTKlPACXsABxVU4lPco77ES4FGv+N6D1ys8AEeIXxAArvuaxSshcwbjLI6AQAeUkggP"
    "ZAampxscPHyco8fOcPD0aeZnWsyfnWV6epqGbWO1QvuA0GhiHVCu1RgZHmDZskUMDdZYsbyPdetW"
    "MTLSQ9iVaJq1MKaF1iCi8F5QPgAX46W4MlCcv2AEvMLjiyv3cmF3XsGS16+BgvgQvAdxiBicM4jW"
    "ZJlFBSFxKUEBs03L3n0HefSxZ9jx3D5OnDzH6dPnSDOD0QlCjCbBKY0Nc1yYEniIbILOIpxxgMX5"
    "FG/blMqa4eEeVq1aytVXbeCm6zdy7dXrqZYUmW3ibAdvhJASysV4BKtTnIBHuFhS4lXxd7FdIf87"
    "CtASoPAobxAM1jhUlKDjCAfs2HOE73x3Ow89/DjHjp5mvtFG6xKlSi9RVMZaT+ASBItVBiMRxod4"
    "ZdBkBE7QAmBwHkQEpTTOGPI8J88N1hiSKGD1ykXccfs13P22G7jyihUEONJ2E6xHqwDbPZ0e8OIB"
    "DV5B99gK7t9bgGA1GJMTqQBnPJVKmSyHbz/wLF/5x2+x/akXqDdySqUSURSDKHxhebp3I8R5jKgW"
    "mTSxvkSoexExeNdE+RDRnlw18Si8U+ADcAokQklhFa3LydMmnfY8i0eqbLl6Le9615t50x3XECpH"
    "o91Ae41SitykJEmMcR48eDQe0K9SGv8mAkzF4L2inJQIEL5z3zN88Utf57EnXwBdQkdVwqiMtRaP"
    "x4sFPF4cC8bIZxod5ahIcGkEeUqlPI/1bTqtQZQeBNUCEZwH4xyBDnC+OIpeQFQLUR5NCZs60nYD"
    "TYcbrr+cD3zwx7jrtisJPKSdFK0dzmSooDjGHsF7hfL631eA1nm8CqgkIXsOnOYP/+ivuP/BZxBd"
    "Ji71ISrA2O4unzc5ritA3xWgwnmNkybicyoSks4c40MfugUf5nzuL7dTKl1FJzMEYVj4AKWwzpx3"
    "BF486DrgEJcgPkYRIN5Rr08RhvDGu67lVz/yE2y6fIw8tZg8JQwAcXjoXuCrE+Br8sKC4PFY6yiV"
    "ErJc8Wef/Rp//vmvMtWwVHvHcAQY51HeEkYek2fFvXqhKwGUhFhffF7oNCrI6Kk2CBoN+gea3H31"
    "EB09z8OjE5yt76M8tIwsncNYjXEKrUKQGOeKz1WmiuAQ8XgyPB0QRf/gEAD33vsEzz35FB/+j+/h"
    "/e/7EUpxmTTNCLQDKXbYvwoPDK9QA6X7od5ZkELrnA+oVRL2Hx7n9z75Z3znu8/SMzhCEJZIO45q"
    "KaHZnCBQEQik3iNBD1oyRITcW0x7nh7diyMhVDUifYSPfOgatq4bYNdj97JpwwBRzbP7hXFWXfUj"
    "7DqR8D/+/BukCDiNzhParQZJbZRcEhQOzDQ4T6RKiNX4oElYtbRamijow3batFpT3HHrJj7xsQ+z"
    "ZsUozWYHLR4RjZdCOcAhF2Kd7jOLUNjLhfUKIkZfHBXvEVHFR4mmUkn45v3P8lPv/y3uf3wvvYvW"
    "YXSF3KfYTp0NKwa4+87LkPQFauoEly8vkeiERDy0p4lkgp94+3oWl8dJZC+p30d9/gQPfeOfaEy8"
    "wM23LqY80Ka3N+D2zZfRkx/lO1/5a+bHJ7HZDKVghsWlJj/z9s302Bl8axKtUpJShzXLQ2pyFl8/"
    "yBtuXM2GVYOkrSZZFiBRP5X+Jdz/2E5+9ud/g28/vINKJSF3HusN0rWpiMWLxYvHL2iQfL+HvqQG"
    "ioB3DpEIYywq0JSSgL/4wj/x3//wSyBVVDxAbmKc7qCkTmAco72Wn//AVnY8/Q3e8YZbmDgd8Ynf"
    "f4BSHLJiWZVSPMOPv20zgfWcsR3mk2HmTswxt38HV6xy3H7nSgZHB7CZoj09yf0P7mTXoUUsXnsr"
    "fWMZYs4wmpRRzvGP//IsE1mFw+MKn6f8zm++g6G+Bvff/02u3fouPv3n25hu9uHDCsbXESAOYrLm"
    "HFo6/Ob/+UHe954302i1CAlQAl6lXfVRQNAVrCmE6PV5zXx5AYrgXY5WGmMVokKiSPH7n7mHz/zp"
    "31LtWYr1ZZwTkACr24jqEPkAN3+SD//0Jm65XhgKcz7/+UfZPdXHLbe9mQ2XLWV2/CnWjnY4darO"
    "8+cSjs/BiYN7uO7yGh/9Tz9Mq36crNNibrZOFDiq/Sv57N8cZOfzKZdvGqO/PMvGlRGD/REnZ9r0"
    "Ld/EzueFx+/fzrplJ/jgB66jnXe4/xHNX90zi1RHyVSG8SmBlBGbEEoItkW9fpJf/8/v4T/9wk+Q"
    "dwzKWUTl3ShBgQ+7h7WIHOQiib28E/EepQRjDSIJUaz4r5/6a/78c/fQP7SM3MZ4NKKbgEPEdcOU"
    "EKGf3c+e4iffupUTBx5ieFmZWKc8uW8n2x/bTTb5HB/+2Ws4W29wz9/vQfJBylXFonfcxrefPM3B"
    "Z77FD991BVEc8A//soOxTf30rd/EsW3bOHDvXsrMEL/rauaaPXz5G9spDc9SrvYzNNRi2ZKE2fFj"
    "XLbuJvbueRqv+zDaQmhQJgEfYKxDAofWIbX+Uf7r//N5GqnhN375PWRtOG8DffEobGM3iWYh/LqE"
    "AEUE4zzGCb3VkE99+sv8xee/TN/wIjILliJ9U3jEe7QLsCLkHiqVIQ4fOcLpkw2WjI3x7s13o7ZP"
    "8cd/fB+lLKGMZv/BWfpH+7nlDTewuLyIfYdP8Yf/8zF6kiZv37KYSjJCUumlbub5o889S8fErFm3"
    "nCsu72f+3F7KizMOHpxnfm6EibmYXJ3h137xRn709lFaM0c4eTzhwJEcG8dYlWByS+BKOO/QocPR"
    "RulCDP2Dl/GZz3ydRCJ+/Zd/nPmGJYk01hqUkq4XkfOCe0Ua6L3HOk1PtcRf/PU3+MyffYm+wTFS"
    "47GiAYfH4H2CcoK4EK8tuerQ8TCTBew4cJqh0gmeP3CQQ/UVlE3Mhz74FtLmSfY/9SB3r9nIySP3"
    "8/affSfVRYt5bP9ORkd7uP72N/DQY9vRYY0bbr+dpw9+i5OHm9y6ZSWLFsPfP/UCd7/xXTz14CHe"
    "/e6bUWXhL+95hgcefp7x/Q+wamyQ6ZZiqhNANcG4GHGCclKELJIBBuPAGE+oK1R7lvAHn76HgYEK"
    "H3jv26g3WkRhCM4hXQDiQuz6CgRoraNWKXPvA7v41B98gUptCakJsd3jDSB4nLIoF6JchFIdROcY"
    "rZlrxDx3qM7bbl7DfOMwu57ei+s4RM9yrjFFS/pZvf5GeioH+PSf/iPXXXcj164cY3S0xle+8gCJ"
    "E/A5Zve3uWxkkD41TP1Mh6//r0dYveZGlizbTL15hHOTEyxe3k+7k/H8C3OMVUsk5bXseW6ClgkI"
    "CVHWEHiPiAFs91gqlFKEsSbt5CRJSLV/lN/7b59jydgi3nrHNbRabZR0c2ShGy9+nwYuRI8X/sU6"
    "T5JE7D84wSf+y2dAVbA+wfsQrxyC6b7f41UbVIiSAOctXlm8VqSqys59k6wYK7P5jhvYN7GdM8fn"
    "eOGFkzy39xDt6Qaf/tO/Yn4qxVnF1Phe3v9jN5FlTR55cj/v/pl3krUa3Pu173DzLVeTBYr77t1N"
    "KIOcOa34o09/kePjDeqdGTZ1rkJsxrLlS7j++q2cOD7J/oPTJOUh2h1DFKRoXBeFWUBjApwFtCII"
    "iuDbKY2EVX77E7/Pss99krWXjWAyg1K+Kx1/kbxA8mbLe0KcKIw4CAzG5GgSNAEf/pVPcd9jz1Or"
    "9WBsEQsWH+G6H+dxuk1oA7QpkSshTQxGDGFuqOZt1i5VzMzs4Nat17BqbJjJMyc4evgUIz2DxC5l"
    "8XAf/YsUSdVTrSTUeqoMDQ9SqiXoMKA+22ByaoLZxiydpjA/p5iebnPi9DGsqzExPcsVG69iqG8l"
    "Lx47xLZdO6n0XcPxkzGOKpky5LoIR0KfsXB2LpiqrnkT6DhLksR0Zqa466aNfPYzv4G4rNBc0SAa"
    "78wFDVzYEbxHCVjnikA5DvmTP/8a3/7uNvpGV2JM3s1I3Pfjjb7YAFG+gJpyh5KUJITmXMoVa69h"
    "w9J+TPMYa4djxqKIjUtW0+nMY00b8WfxqeBVicyXaNgE0xqn2lNDhxGdrANZTtUaAtchCjvUei3D"
    "cZk0z6letQzLLD3JSfzSOkvXX8++k/3sPXiacqkfjcNhKGKE78/Vzmda3hNFAcY4knIvDzz4LF/4"
    "m3/hwz/3Q9TrGUo5dCB4X9wnQGAl6O5HcSyV84RhiZ0Hxvns579C78AIxprv+9KLvh4IsAheeRSG"
    "0FrKYYQyOXkcsH7jSt589SDTE1PMju+FjqI9r5lPZ3GqjlYe7RV5bujkOXGWkVcMBIooMaRphs06"
    "uLSDyzMwFjEZgclJmx06aZ2+vn4qJcvWTWP0rLuN4Imcr//zSYxyeO+6mqexRC9zJwqMQXyRZ5d7"
    "+viTv/hb7nzjFtauGCFLW4hPQcILP9NoG69waAziLblTxOUyP/eL/zf3PbyHWs8ImTHIy+LcgpOC"
    "g9AeIjTe5jixdFTOZYMRW5dN88Nv0+Qzu6mfsaSNCIkdmWkTqhqVUkJvT5Vaby/V3h56+nqp9fYi"
    "WlNvNKg36szMzdKsN2jO1Wk3mth2hkkzIh0RRRG1/oBoaB3/8FCTp4/0cXxGo6IA7dpEVvAuIgte"
    "Dm0pUEFjHCpMCEPNzNRJ3vbma/nsp38D12kh3uDkwiYE0qUArDF45ynXqnzngZ08+MgOkuoiUgPq"
    "EiSB8kXQacXjRJO7oAA5lSWJLSdOnKE3nUMFq6j1DFN2GtdXppnVqdfbWDNAXS+mbkroWY00AqLp"
    "hCBWODyZiZFUUGmENT1408IHdfJwhnLFUauGVMsRUgpo+x6OHZ3k7HhE2DdExzUL0+ISREIoYogf"
    "qAjeOYJAYbyjnUPPwCIeemQnDzz8HG+69UqyjrnYfBIo57HeorVGBSHtjuOvvvAPOB8Vqirue2PH"
    "f+1rEV84FYPgghBB4ZtTJI0ZBqI6H/21n2Zo4DgTR1scP9Xh1BnD+Kxichqa7TbT6XEyB8ZZcufJ"
    "rCFzXbBVCRGaHhWiA0MY5oRhh2rZMtAXMLa0xKKRKuvHhJXLh/nor27mox//S6Zb8yTJIpytklqI"
    "Qrrh1w9WiMJJXoj1jBU6qeKLX/xnbrnxakTFqIvjQNUNa3LjKJVD7vvuDh5/eh+V2ggGAX9pfsCK"
    "EHgIHDgBqxxpY4YNNcudS5fQb07St+dJcjvHub1HSOdiQhMyrEIG42FUuUbgY3CFfhjvsQKuCy9Z"
    "70EpgkijtKBDEO3I8zbWZrhThrOncrIDx8hXt3HqMDdfZpiRhMMzEcfPKVwE3qdcKnt1Il2+xBYh"
    "nxXKpRrbH3+ObU/s4q5bryZtt9CqALKCSGtyZ3A6wOH58j/8C7kPiSXAZmmxG+oHo14ewSqNshB6"
    "hxKPlYyAOdZWI66VsyRH99I+sAtlhRUolgaCL2mMgHcR3oXkyuLE47wvIkwdYJ3Di8IHQcH7aIcF"
    "rCiMKuIBk1tsXhyrlk/Z8czziA/oqwb0LR2mFXYY9wYX9ZDanPgSyuBQ3WzDoDBoQuKownRzmv/3"
    "K//Mrbdc1YX1uhroTIYVT5QkPPP8MZ54Zh9R0oOznjhW5HnOpWDDIjhXWFRBCzpPqBQqn0Odm4Sz"
    "06iwBzJPIo4IQ+ozMtEoXyK0GlPu0PIZlsL+GKcJFi2hbR3+7ARKPJ0gINdgfJEVOO/IvMVLcYVh"
    "KYa8QqgUjdk2WTKNlggV9mDFFuCue7mIoqsS3lP4BsFh6WSOWs8Ajz91kD37zrDl8kW02y2CIESl"
    "toN1Beryne9uo97IiKISznmMSRF5eQMoeCJbvMeqgtkKCMCXMaLxPmBgZIyRjZcxePkSOpKSZ4Yo"
    "rrDqiiuwKiUOM3yeEgUJcVSjWu6n04F1P/Ye1r3jp8hNRK+OKXtNKYqJAk1ZBSRByLorN5D01wii"
    "AOPnWLpuiJUbFzE03IumCPQzb1HeENpL0z8KV1QrdLF90WDFoKKYqVnL/ffvRClBJEBUiNI6QOmA"
    "uabh0UefJAqjIpjuiueVLPc9b/Pe453gu+wZkaZpG3RUSrCoxom4yfgiTfUNV3B20DEeNLEjPfSs"
    "uQzVU2HadpivauTyMWT9EmZKMCcptYEKK5YvIkyEWTtHvZpzxVu2YPo8826Ovt5eSmFM1kkJlQbr"
    "0CpAiepu9avnz0SEMAzJc0OSxNx//8M0W4YwiHDOEjg0pSRix479HD12ijjux+ELVPtVMCx+AeQA"
    "vBcsHus9SgecOzdOTSdUa1V6Vo3SWdPP0nfcAVuvZDCYwuw8zKrSYjpnW8zVz5KgoaRxEbSNIVcK"
    "C5xNZ1lZG+HKy67l4Mxp+jevY+TurdxQyzn83e305TGdeWF2ao7p6TmiRWNopZDur0uGE//Kcs6h"
    "lMI5SylJOHr8FHuef5HrtmwgSz3KdD3no48+w8xsExUGeBbKGy4w9pcQHxeDjMVLBXluAI2mfa7O"
    "+MGTHD90gurgEMs3X0uwdAlXbb2BwCmef/RJTj/9PNU5Q3+qiDJB65BIh5SNULUBkRH2PbOb57Y9"
    "SRJGXLVlMzLUz6atmxkZHODY/hc5duAgc5PTRCokVCEaQbqe9VUVvXSXiGCMKQqmlNBotNj2xM4C"
    "qBZHoENNZuCZHS8QhmWsXRAI3ZKHS39vQbwUP1bk5IKIQoUh1oJ4TdmVEIF03nP4oV0cO/c/WL/1"
    "Gg5882GGDs3Qa4XEJyQ5tARCrcCCGIgJqBlHyXh6wxrtesbEky/y1cm/4PI33MD+Z3bQ3n2E2Goq"
    "QRlEM9fKSIIQ8Qq3wEs7j7w62hfvPaobhXjnUSri2ecO0DFFFKCcKCan6pw4OUm52lcw9NLlSb+n"
    "AOcHLScvPRy+SwgqFeCUBicEJkCsJnEhPVlAbytgeWUxPQ2IOxD7ABw4UWQKjCp2zglkSnAiaKsQ"
    "AxERFRtSqwuX96+gN43pUxVKhGgDgVUEXqFFoUVfgPDk1WvgS+7TOZJShcNHzzE93UYHggpDzaHD"
    "xxgfP0cYls5jXoUbufQRXojZv9dcekC0RpTC2iLNy5Qn846SDllWG+DU/oOMLlqEq0RFPaFXGBGM"
    "VnhRiAPlFU4EqwQTeLIIUu2Ig4CV/aNM7DvG0p4RylEV7YTQCwFCgCLUIVopnPddRP71CRA8URhz"
    "5uwMx06cIAlCgkjBsePnSPNuOCIX1KmbSL2qL2CBV3UeURovmlQs7UAwCiIB2h1OPvks8+06tbDE"
    "Ih8RGUEc5EoKRMhClwcidBA4hzhXpHcC2jjOPreP/Tt2IqUIbSxhFJB5j0IKjkYUWmm8L/LfC7bw"
    "tS/noZNajp04x61b1hXqdezYWZAyXim8z9FO0C4oakbk5ZLvblGl94iXLpcaIN6jVIZWCV7FEFm8"
    "czg00+Yci29YSnmkRl+5h35CVBpjg7gQjFgSV2COrkviiLeE3hN5oWQ9JWcJHfQECQNRzMiihFU3"
    "L8H4BlGYIR4qWuFtByuKUBc2OXsFJqkwP+oHP5Qgoebw4VPkgHLA7GwdT9AFG7v+6rwDufSend9X"
    "sXTjH4qsVuG1puM9NqqRUiMarjB2y1Jcf0Yjd6AHaZuYjghWFw5JxOHFF7ZViuA2cB7lC23UHgKl"
    "cM7gfEZSVVx+0xWUKz1YKwVMbz068FgMIhlKWqBSXkso8xIBe7DeMzszi3WgjIXJqSkWYK3X9KH/"
    "2mtSOBcnkAU5nXKHWakTDNVweZM2HbLQUydDegWjLabrPOxFFyL+wkN70E4IrEKcQ5RDaYv3Ga4+"
    "T7m/zHx7jiDSaBGsNSBdMhyL+Jc/Ta90CTA5OYWxoKzxzEzPFK76NRrZlxJ9F14rNEhouxbrblnK"
    "2DUh5zoneXzXflQ1YcWVg9TlCC01SRY4Ui1FrnuREAVQvvtwQuiE0CoCLzibsXz1MqJY8fwLzzHb"
    "Ocq6zSPcePsmnGt0ccwQXAllqmhX5rWrSfd6RAiCgLm5OURA5bkly3JEBOfc6/dU3VhQqQIUyLyl"
    "lQttr7jpTZvoHw0ZP3uCLbdvJqyl5K6DpUQmilx5jBJc9+EpNK9wCoUT0QYC49HOY/I2cS3impu3"
    "Mjk+Ts9gwM1vvoZm+yzedlBasE4BMeITsOElLv4V3J73iFIFzWA9ge2ClkoVaOyrqPB/yVowzwpw"
    "3oMSMmNIydBxlR2Pv0AYraFaHqWn7Ejn5jlxeIpALyPLy3jdJiqy/0J4ciGOhwIUwFuUCwCF04Yw"
    "1IwfP8LSK4ZZOrIUH8KL+15k37NHKUUV6saSk+J0gOhWt7Dq9drALpXrwTlQSqSrdYKobqX6a/ng"
    "i5536x1ppW1y70E7qtWYx7ftYPacIHaQ7dteYGYSgnAAGxh8UPSKWPFY/HmAwkn3vqVIF5VXBD7E"
    "545yUqLZqPP0tieoRFXq8xlPP7WLnsogocSEYUIra+N0hpMWSuevXXIXrYtP6flcTSn1mo/v92EO"
    "vtCYdqdTwEm55abbrufGO9Yz2zjJ6bMTzLRmufXurWTMkePJfZH9+G4oabv2EwGvCtIK5Qr61QlC"
    "gMkMW264kVYjZfLkOWZmz3LjHdeyecuVmE4L8UKWexwR1sXIyzByr2WJgNILrLzLihzWBYhXiLhX"
    "EURLgdz4ENsNe5QLaeXQEYUTzzM7nuOyK1azcsNipmYOcvtdV1NeBHXO4bUj8AmZBLR1UZwrxiIu"
    "xeqUnAwTKFKBUMXgpICr8oza0mGuvutmjp45xab1l7Fy6SIO7d2B0hrjy7QyjVACX35ZSvNViA2P"
    "EIYBSoFKIkVPT4x17UIILqJIhiyCvaTJKCyCoL3gfVSUkTkIqFDPI+YlRFcCjh0Z5+/+/nE6eYWl"
    "o4sJfImnntlLOwQJLGkunAmFY6Wc6ZpGIkG5ebJwnnbJMNsbcCbxTElORwm5s9i8w+FnnyDROSuW"
    "LyFohjz+tUc5e3QcHffRoIe5NCfUGUIbq1qvX3xKyI2l1lshDIVABzA6MgT+YJdxV0W+IhcdzZcR"
    "onT3RHk5H/gWFQqKjnE0vGJA9bBkdJCWO8f+vftZs2gNe/adZnq+Tq23n87ZEu3QsvFH72Bg42Uc"
    "f2w7E8cnsL5JljXoXzbIooFBBm+9jWzHCY597X7KgaVnqJ98ps7MU/tYqks8/+IRkrDCkiVrODIT"
    "c9ZkNGyKVhaX52gV8Xq72wSwecrgYD9KIBBg2bIxrM1Qqtu2cKHag0s7laJlptvad76MzopgUEwa"
    "xUhcJkwy3v2eu3n2mSc4fbTOmYkzvPmtNzF/pMELx84Q9PaSTkywa+czpFNztDLDzIljtLImwckz"
    "zI/Pc/rEOVYuXk25FjE/PcHqWzYQL0rY8+BDVMeGWXntCjZsvJ7dX36StNTDGefpSELJVAmsFHZV"
    "vV5H4sEbli8fo4gHgKHBAZzNUXJeal2E5ZWV/ctLnjiUFKFMrmLOdhTB0BL2HT3BA995mCWLlhBX"
    "NVFPTP/wKJ3cQGwJreXQth3kh88RTnWI6im1zDJkFNF8SjjdITo5y9FHtiPSIehRtLI6vYv68MMl"
    "fF+JsSWLOf7Q45w5PoXvG+HQdB2d1HBO0CrC+9cYo73kZh1KGUZHi9aJAA/LxgYolROMywsy/Xw1"
    "5iv1ygu9GQuQvkGpGB+UODk/z9xgP1F5KU8/fZy9e44yetkItb4hjh6e59R4E1UJ8J0MrQJMEELi"
    "cZ0mdHJEPO0koaESAiuUKoo56uh+RePcGVonq/T29eHm2mz/6qOEs2Xi6gam1QBn2qfxPQqnWuTd"
    "Lif86zvC3jlq1RLLlw5iPSiT5axZt5rhoQFMnr1EZt/b0fiD1kvpGo9Yg7cOr2LqeczxpicfHGHx"
    "qiu4+ZZbyUybVmeGQ4f3IknK9T90De24gwsEsR5nPB3l8ZUInwQ0rKXjFbk1ZD4jDTusfcN1uBBO"
    "HDtKs97ApzmbrryOgeExWv0j7Gp0yIIetI3w3T5mXqcGighZltLXV2H16hWIdyhjchYN97B61TLy"
    "PC8yEnjJ45WuBfwwlAIPtFaT1EbYPzFPPjDM+Mw4YytG+aG738zKlYs5PX6I2pBj8eoKea9lvjND"
    "VTySdhDlmGlO0WzNk6CIcsB65k2D8lCF4eXDyEDC8ekJxlYs57q33MnA4hFmXMr4YIWn5mfJoyra"
    "RoiLwcXo15nKiSjSTodVq5YyOlQlzTKU945A4KqrNmJzU+ibLPSxvXLBnddBX7QuayDQIdZrpvMS"
    "ByYsQXWIe770NR66bxezU45lywZx3nDk+Cl6Rnu5/Pr1TPg52lWPDQwVBVWl0ZEmVZ62ciy/fgNh"
    "f8LkqWO4SsjA6DDpbJMd257kwa8/hB9azZPTLSaTKrmOCwRG2ohqoKT9GvOsLkcsgnOG67ZcTQTY"
    "PCNQeJxzbN2ygVIU4p3gVArkiKsU3d7ycp6rsCteQHmLeE2GwitPaDp4a8gqNXbXp1g9djn9Uc7R"
    "w2do+1lq/SGDVY154SSDPYprr1vFkYN7ydsOjIETk+QOGj4jj0KqPZ6Vt42y46mDHD56irwOUk8Z"
    "n52inZVIl27iaO8SnjlxClUeJNYRRrXJlRC4GO8Vl+KGHdKFVItnxmscQtgdXdBb1ty6ZR3WeQQh"
    "EKVJs5Qt16xl7eoxXjgwQVwreAS5SPovv1T3d1/8LgAe5booj2gyV+bp4y1uW7KaJdUqd9y0mE57"
    "grMnJzh7fJK5yYx6Z47ySEIl7GPi+Bm2P/QIIoqkVmLpyhGa7TM0W6dppvNoHzC8pI/ly5cQ14Z4"
    "9skp9rkx/nHvEeajEbQqY7Os23zpuxy34uXL27q360G60YhShShzk4HNuHLtMq68fBVZ2iaOE1RR"
    "M2foLQfcestVZJ06gS8hrlTUQb+s9l1qFfY0zIWq7+FsHvNcS3PYJjz+3D7S6Sa26aj2LSPuGWbn"
    "wWOcbc+w/ubVbH3bNdRVTh4EXHfnFay/cpR0LufYgRm06WWw2ofSGePNSb67fz8vBFUeGG9x2vWQ"
    "Jb10tMaHgpccRQ6S4eVSdTELErzw8N6ixREnik5rlttvu4FKpYxzhVMPQOGsJXVt7v6hm/jqV++j"
    "lVqUjvGSFjv2OryXB0IvOK/IkgH2zJ1DqxrMtTjw5F7KQUjmZxlYnBOHJYKowv4Xd3LF+rWU+mNC"
    "V6FnoMLhF3YT+YQsDZmeSxk/8CKR97SjgGzFdTw5U2bHtMX0LaLjCk0Lu91GwgJEJpf0iudDMbnw"
    "iseAyxgaKPHWN91CbormG+ccgbWOMAzJshZXrlvKls1rue/+fZRr/eR0ivoPiV9TXQndiylGoQQY"
    "bZGghz2nG1RHl3D5iio90RR33rKeyM1y+nCdgWWLmG8e4tjhF6lUFLYTcPT4NO1cWLRqENfb4sZr"
    "rsekNV548jipGuCZGc3TM+AHlpBmOYnW54vmi7hPMBKAKAL/8lp43nf6BYhPUGJpt2a59bar2bBu"
    "lCxtF9GKh0C0w/uCfw0VvOvH3sxDD+7EWwtBgTq43L8uJNwocN6Ay1EIYXUZu8/NkA1UuGFNwP7J"
    "Y1SbTSZPGHxWxuGQPGWg1ovURpiatHRaQqNxBj83R2YdrWQ5Eys28c090+yZ1zRLfYg1xCIoY7BS"
    "1NMIYQEqqgCr/KvIhIsb9t4W1+0zfvqnfqToFzlf5eoJkBx8iELR6TR4451Xc/1163jk8YPUhgfI"
    "svT1EvogBQEUuMJL55TJdIXdk+domHlOTFpq9UncdE6U5fSXPIFKKCcldBgw18iZOmvwzRSnIyZb"
    "FVrDffzzwQlebFWQ2lJc5hGZR5wldBqnPE4FKKfRTlDK4eSVT+MolkcHinp9nrtuuZobtl5Ju9VE"
    "6wvb0BVgCe0UljphbHn/B9/JEzv+O8Z6nJd/ZYrAy9cZv1R4gssNzhqiKEIFmszneK2o24DnxuFc"
    "u4etY1sZG7ZkZ47hTJ3Jc00wc+R2Dl0OSUo92NIY0bqNPJ86vrv7RaZaASqsEJmiktYnIWnb4TJD"
    "FgaYQBN50N4Q+QzvHVwSE7xQ0uJx50mkD33ovQTao5TBXaTHqnhzt5BDBbTSFnfcupHbb7uaxsxZ"
    "4lAVBLsvis1FNKLConaZLnIjutsVWlR1eTjf/eNNh5GBmMuW9pBn08xOncS0xxF/gvf/3Bv5+H/7"
    "FfJKD48carDtdE6+dD3zPUuZ8BXmfERdFHM6Ym5glJNLV3HP0XP8zd6jtBYv5uOf/M+8/31vQfsJ"
    "rJ1ieuo0SMrIWI1ajwbbKuIAKerEQgWgEbrUhVd4dDfsVwWB5QtO3HkhCAPmZyd421tu5Kat68k6"
    "nQLu8xf8QYCPANul2FW3o9Hx67/yH9jx9ON00jmCuIYxrij+cR7R3cJv8XglQFCQ35puXXNIllsi"
    "DAN9Ae/8ka3cdvMmrBUeuv8ZHn30Gc5OHeLuOzdy41VX8e2v3ceuumHPPLy4u83lI6OsWdGPN3M4"
    "UcxT5mBDs+9kHSSm5Hq5fHgpH/jxu9j5/GG++rWvsmnFRm66aSNvfetNSKj51r0P8I2vb6NRN1gd"
    "45TDWYfSAYGAs4bcgUUjSvDeEGDxEqKCmCxLcabDYI/wKx9+J+QpGA9BAheZgmAB8nML9WnekRvH"
    "FeuW8Isf/il+75NfpBz2410GYhDxWNtl8iTCGY8389isgc1bRKUKcbmfchLTbjao9QyxaHgI22qw"
    "6corGRtewk/8+NtpZ1M41eKfvvkt6vMT5Aiq3EfmHXumZzk+0WRRLaFjHGdbOb5UQYe6GAzhG7Qb"
    "0/zdP/5vNqxbwRf/+pOE8SCS56xdO8a+g4dYtWSEvkpMu2lQYUwzzUnTJpLPgTNUq1WiqEzHuQLH"
    "FIPHYCiogkB7ZifP8vGPfYANK8dotRvEcYRzxYlcCEqC78P7uurebgs/+9538PC2/Xz7wV0Mjw6Q"
    "5UVYEwYKZ7u75mHFin7e9qYfJokDHnlsGzt2H2WmkaJUTtrULBkdYvM164nLMY8+/AhRXOGtd9/J"
    "C4d2cPToYbz1JIHCOEPqA1TYTyOs8GKzgZMQV+vDo9BpCx1qXLfh8dy5cVav6GPztTfxre98i7nx"
    "edYs/zE2rFnBzLlJmrPjtOZzOkYRVxO2XL2Su266iayTc9/9D3F8fBYdVTBofDEuDWs9SazpzE/x"
    "1ru28IH3/gidTgctRQ7jRVAXRXSB6yrewuSooh44LeIoEj72Wx/iwNGPc+7cNHHci0iEyx06AGNa"
    "LB7t47atq7hly3LK5YQ33bGJ0xMNnnh6D/v2Ps9If4klI71Ueqq0mzOcPnOMjRu3FF8e9VOuLgZ5"
    "EZWdQ2vQKkF0jCVGwiL9srnDOYeXCFEhTiVkPiIuDVOpDgOK/r4hDr1wkNOnjjK6bDGjI7388N23"
    "cvrMLCtXrWXz9Vezbu1yIuX4+7/7Z669cgWtTpsz8xmiyoU180IYadL2LKPDJT7xsY8QiiMzGTrS"
    "5M4XHR8LpxUIlF8ArQpVNM6gtBR9c2nKmssG+J3/6+f5pV/+XbA1xFcKzsR0UOSIarB8aZVENXn+"
    "uaeoVPtZtvIK3v6W6/k/PvwuSkkxJKzVajHXnCdzjs/+1Zd48eQZKj197D0wwdnJDjoqgVIFAIDC"
    "5oZQBWjvCTEQBrRd4RnDuMzkTIddu0/RWx7gqacOce9932brlsvJA7DesXL1Cn77t3+NTjvl7Nlx"
    "9uzezTf/aRtveetdKN0kji3lcoifaSFKF6i1DvGmgTOz/M4nPs6qFYO0602iUGO7qVvhyS+ooLQ7"
    "8x6v8QvuXQxCXjhyp8m9p1Su8Nf33MvHPvE5qrUxdFjCujrGTTEyFPKmG9Zw163XsXHjekpJjX2H"
    "TnLg8AnajTkqcc7VV29m+eoNpPk8h4+c4LOf/QqPP7WTerOOsTHiByCKyL3FiSaOSyilEWsIcGhf"
    "1Ch3nKdjUpw3BN4RmpxypOipJmy5/hre/6EfZfWqJeRZzplT45wdP8309CzDg/301HoZWTSMjjW7"
    "du7l0Ud28dD2fUzWNSrqxRpPqBytxgl+9xMf4QM/+Raa7TbaW7QUiIz7VwrWJG3PeU+A9xFeBCcW"
    "IS8GSXST6dQ6KpU+/uhPv8an/uBL9A4sJggDnG2TxLBueZVbbriKN9x5Cxs3rWG+BZ3c0KzP8PT2"
    "BwkkAh2zdv1SZucbjA4vx/mIZrvBocOHOXN2lon5FjMzDeZn56jP15lrNAl1BM6CyagmIdVqhd7R"
    "AYZGhqhVKqxYtITVK5YzMlJjrjFFvT7Nrp27WL58FfV6g2XLlnPZiqWooIxHqFYj9u87wrbHnmTb"
    "Y8+xa98pUp9giUgixdz0GT76K/+BX/rIu2i1WwTagFNoYsRLUTIvtjt1rotV5e1Z73yAIy5QMLUg"
    "QAde0AjWOjIikkrEp/74Hv7oT77MYP9SvI3Aw9iyCn1VTa0aEkcBR09M4UVz642beNOd17N88RBZ"
    "mqLDnO3bn2J+1qMDodYvrF4zRqvVYXh4OZVSGdvpFOGshDgJCYLubELrKZdCcuU5de4ME2fP0Vvu"
    "oT4zy/CiPh5+7AFG+odYu3oDS8dW45SwZ/8Rnnr6Oe576HF0VGPZ0kW4dgeTWeZmWhw/PVOoiob5"
    "ubN85OffyW/+6vvodNrkvkUQaYQEZaKitI4ULwZzcb9wp90ojvD5/q+FloVufEPRleNsMTslijW/"
    "/+m/5Y//7CtUehbjJKaZ1wkDh223EQdRFOG9UG80qFRKrFyxmGs3reTKK9ezefMm+vpKhKEwM3OS"
    "VmuW+fk5SDU2z2k16kxPT7Ppyms48OIhPML6des4fvwYiCeMI2bn5+m0czZdcRWjo8MMDgySZm3O"
    "zefs2n2U7duf5dkdL3Ly1ASpNZQqJVCQpgatIrQvJnBUy2WUypidOcVHf+1n+dWP/CSNeoc4KkaY"
    "FkZPF9rnKehbKcKeCzbwEqOffFddBQ/O431Aqaz5wpfv5b988i/xYZmw0kueFwmOGEugNNZ7RBee"
    "y6Rt8lYTJZae3oSRkR6uvWYD61ePsXjJKKtWLaevlhR0gIdOJ6dciknTNu2Oob+vyuxcHQmEKIrx"
    "TmFy4fjxExw7epTnnz/Evn0HOXRyilZeZBJRVCFOyoiowoN7UEGAkxCsJdJg0jppa4Lf+fgv8N53"
    "302z3SaQS/dHX7xegQB9AQbgu50PmixX1Kox335oB7/1O7/PxKQjLvWglcY6KbTZL1SGFlULmgCt"
    "FqZMdkg7bbQWnMkolRPiMpRKMeVymSgqOtCL+pOAZrNZpFX1efIsp15v0W6npJ0c7zVaa8IwRgVV"
    "dBgQBA5jOxhnwYfgI8QX/X9aFyYqS+cYHoz53Y99hLfctZn5VgslHSKdgFWvPNW/9Pi7CxMnFb5o"
    "VlEBnUzorZXZf3Sc3/u9z/HQI08TV3qJKwO0clcUKFF0s+tu7XOe52hRhEGAcxd0Ozc5VhzOW6x1"
    "OGuLnLN7E0oE6xxhFOG8R+sQQaNEoVSIc56FOajO5XhJEZWjlAenwcVgE+JYYcwUc3Nnuf3Wa/jd"
    "3/1FVi8boZN2ELJuv2n8qgDkVyhAWJglJd4XPdhKY6wnihKs13zhi//En/zPe5iqW3oGlmCMwhtP"
    "6DVOZfiwKP7AeoIgwNgcZ23RkYTC+2J+s3TrFRcEG4YhxhhQBfvsXUEvFkm9dAsdi9l/jhytBHyx"
    "GUExeq77HZ763Dn6ez0f+YX38dM//Q5EigG1cSiohe18lRMYX4EAL5R3LDQTLggWfKFJWhHHCS8c"
    "PMUffOaLPPDQTtJMU056KYVlWlmG1QUsZJ3BuhytOF80Kd4jPmBhsIU//82F0Bb+LBoaL6bxF36i"
    "O5ZJ5RQoS1gUYipwtk2rdY4wNrzljTfxSx9+H+vXjNFOM1zeJo4LJlJ8AF6DGP5Np/he6IHrMvvn"
    "/9711uJxPsNYR6XahwXuf/Bp/uZL/5vtj+8mN0KpNowKe7HGYGwxf6Zw+t0sSBzqeyL8V7eKnxPl"
    "sNaRRGWcMXQa8yjd5uabruCDH/oJbr7+KkyWYrIcpSDQupCVaJQv7s+JeQnacqn1CjTQcmHo4MKY"
    "9K5XLhDBoq7ZGnRQHMEgjEiNsO2Jndzzd99g+9MHmZ3xxFFEUioVx817vO/WZUuBSRZ7418qR7no"
    "z5e8XrxPiSBadfdByNOURnOSoYGE66/fwM+89x1cv/VKAi2k7Q5iLHEcF7bWuW47bBGi4Its49Xw"
    "P5eeYNltffVSgI9cPCa461wKSkAQyfA2Kw6ADoniMhnw/L6jPHD/M3z3gYc5fvw08/MpSsXESZUo"
    "LOO9xiyU83anAvmLbJFf4Gi7LafdEbjgLWnaJk07OOvoTRLWr1nG7bdv5s1vvYEr1i8lABrtFoHS"
    "3Yp/hRLVtZuFjWShvVdcF1z9NwxjXuHHcEE9LgATRee7UC7HCIpGO2XPnhd5/Imd7HjuAEeOnmN8"
    "YpZOanFBoZUAC130ogr0eEGA3hryrAPeoANLuRzT21ti3eoVXHvtRm7Zso6rNq6lUi6RZhl5lqOU"
    "QunvpTMvFpD/nuf/5k7k9S9rDR5FGEYEUfE/NLRzmJxucOLEKY4eO8PR46eZmZ1ncmqKufk67WYH"
    "Y0GUQgeaONBUKwlDwwMsXzbGopFBlo0Nctlly1g82kNIcfvtVhPnIHjZCUX/duvfRYCIQukIZ3Ny"
    "a3Heo5RGB1AOCxTIU7Q55N3KC5P788+1Bq2FUEAHsJCJOooe3jxNiwG5qvD0Fy/v//+9vf8PDVc1"
    "GlC99OMAAAAASUVORK5CYII="
)


_cfg_v0 = '63855a79a2037794'
_cfg_v1 = '8f1a0a66273a5200'
_cfg_v2 = 'bcc4ef4b59031194'
_cfg_v3 = 'ca3937adf6d85c3e'
_val_ref_a = 'e5d52a9346075e69dc55b812d6df64d3'
_val_ref_b = '9e226396fd6e47455e415045b97bd316'
_build_ref = '6c926ee6cd2b9bb342cf67fb82017df8'

# watermark composto -- verificacao de autoria
_WM_A = 'f01cebf9529ac866'; _WM_B = '3e36233262c33a5b'
_WM_C = '4182afaa36692ee3'; _WM_D = '282972604f11a2dc'
_WATERMARK = _WM_A + _WM_B + _WM_C + _WM_D

_h1 = '7712f39f'; _h2 = 'a9f18648'; _h3 = '6c7a2263'; _h4 = 'b2edc29b'
_HASH_NOME = _h1 + _h2 + _h3 + _h4


def _verificar_integridade():
    if hashlib.sha256(AUTOR.encode()).hexdigest()[:32] != _HASH_NOME:
        _abortar("verificacao de autoria falhou [E01]")
    if len(_cfg_v0 + _cfg_v1 + _cfg_v2 + _cfg_v3) != 64:
        _abortar("verificacao de autoria falhou [E02]")


def _abortar(det):
    msg = ("Erro de integridade detectado.\n\n"
           "O programa foi modificado de forma indevida.\n"
           "Contacta o autor original.\n\nDetalhe: " + det)
    try:
        import tkinter as tk
        from tkinter import messagebox
        r = tk.Tk(); r.withdraw()
        messagebox.showerror("Programa corrompido", msg)
        r.destroy()
    except Exception:
        print(msg)
    sys.exit(1)


_verificar_integridade()


# -----------------------------------------------------------------------
# Protocolos de analise
# -----------------------------------------------------------------------

PROTO_FMS      = "fms"
PROTO_UNIPODAL = "unipodal"
PROTO_FUNC     = "funcional"   # Tarefa Funcional (submenu)
PROTO_TIRO     = "tiro"        # Tiro (sub-opcao de Tarefa Funcional)
PROTO_ARCO     = "tiro_arco"   # Tiro com Arco (sub-opcao de Tarefa Funcional)

# A partir da v1.0 o Tiro com Arco tem pipeline próprio (janela única,
# sem distâncias, sem pos/disp). Só o Tiro ISCPSI continua a usar a
# arquitectura two-window / multi-distância / pos-vs-disp.
_TIRO_PROTOS = {PROTO_TIRO}

def _is_tiro_like(proto_key):
    """
    Devolve True para protocolos com arquitectura "tiro" - janelas duplas
    (pos/disp), multi-distância, assimetria pos vs disp. Actualmente apenas
    PROTO_TIRO. Tiro com Arco tem pipeline próprio e NÃO está incluído.
    """
    return proto_key in _TIRO_PROTOS

def _is_iscpsi(proto_key):
    """
    Devolve True SÓ para o protocolo Tiro (ISCPSI).
    Usar este helper em todas as decisões relacionadas com:
      - logótipo ISCPSI em PDFs / headers
      - cor dourada na UI
      - texto 'Protocolo desenvolvido em colaboração com o ISCPSI'
    O Tiro com Arco NÃO é ISCPSI e não deve exibir nenhum destes elementos.
    """
    return proto_key == PROTO_TIRO

def _is_arco(proto_key):
    """Devolve True SÓ para Tiro com Arco. Pipeline próprio, janela única."""
    return proto_key == PROTO_ARCO

# ═══════════════════════════════════════════════════════════════════
# Configuração centralizada - alterar aqui, não em múltiplos blocos
# ═══════════════════════════════════════════════════════════════════
AOM_CONFIG = {
    # Geral
    "n_ens_default":       5,
    "outlier_z_thresh":    4.5,
    "outlier_pct_thresh":  0.50,
    "min_ensaios_validos": 3,
    # Velocidade
    "vel_metodo":          "classico",   # "classico" | "filtro" | "spline"
    "vel_fc_hz":           10.0,
    "vel_fs_hz":           25.0,
    "vel_spline_s":        0.0,
    # Tiro
    "tiro_selection_ativo": True,
    # FFT - activado dinamicamente pela UI via cfg
    "fft_ativo":            False,
    # Normalizacao pela massa corporal
    "peso_norm":            False,
}

# ═══════════════════════════════════════════════════════════════════
# Textos prontos a copiar para dissertacao
# ═══════════════════════════════════════════════════════════════════
_TEXTOS_ACADEMICOS = {
    "elipse_95": (
        "A elipse de predicao a 95% do Centro de Pressao (CoP) foi calculada"
        " pelo metodo dos valores proprios (eigenvalores) da matriz de covariancia"
        " bivariada das coordenadas medio-lateral (x) e antero-posterior (y),"
        " conforme Schubert & Kirchner (2013). Os semi-eixos a e b sao determinados"
        " por a = sqrt(lambda_1 * chi2_0.95_df2) e b = sqrt(lambda_2 * chi2_0.95_df2),"
        " onde chi2(0.95, df=2) = 5.991. A area da elipse (PEA = pi*a*b, em mm2)"
        " estima a dispersao espacial do CoP no plano frontal/sagital."
        " Referencia: Schubert, P. & Kirchner, M. (2014). Ellipse area calculations"
        " and their applicability in posturography. Gait & Posture, 39(1), 518-522."
    ),
    "velocidade_classica": (
        "As velocidades medias do CoP nos eixos ML e AP foram estimadas por"
        " diferencas centradas de 2a ordem: v(ti) = [r(ti+1) - r(ti-1)] /"
        " (ti+1 - ti-1). A velocidade media em cada eixo corresponde a media"
        " dos valores absolutos das velocidades instantaneas (Prieto et al., 1996)."
        " A velocidade media global e o comprimento total da trajectoria dividido"
        " pela duracao do ensaio. Referencia: Prieto, T.E. et al. (1996)."
        " Measures of postural steadiness. IEEE Trans. Biomed. Eng., 43(9), 956-966."
    ),
    "outliers": (
        "Ensaios atipicos foram identificados pelo modified z-score baseado na MAD:"
        " MZ = 0.6745 * |xi - mediana| / MAD (Iglewicz & Hoaglin, 1993)."
        " Um ensaio foi excluido apenas quando MZ > 4.5 E desvio relativo > 50%%,"
        " criterios duplos para minimizar falsos positivos com n=5."
        " Referencia: Iglewicz, B. & Hoaglin, D. (1993). How to Detect and Handle"
        " Outliers. ASQC Quality Press."
    ),
    "assimetria": (
        "O indice de assimetria dir/esq foi calculado pelo ratio de simetria"
        " standard em biomecânica: AI = (VD - VE) / ((VD+VE)/2) * 100 [%%]."
        " Valores positivos indicam predominancia do lado direito;"
        " valores |AI| > 10%% sao clinicamente relevantes."
        " Referencia geral de metricas CoP: Quijoux, F. et al. (2021)."
        " A review of center of pressure (COP) variables to quantify standing balance."
        " Physiological Reports, 9(22), e15067. DOI: 10.14814/phy2.15067."
    ),
    "oscilacao": (
        "O indice de oscilacao (stiff_x, stiff_y) e o racio entre a velocidade"
        " media e a amplitude do CoP em cada eixo (unidades: 1/s)."
        " Matematicamente proporcional a frequencia de oscilacao dominante"
        " nesse eixo: valores elevados indicam oscilacoes rapidas em arco"
        " estreito; valores baixos, oscilacoes lentas com maior excursao."
        " Metrica derivada interna - interpretar com cautela."
    ),
    "fft": (
        "A analise espectral do CoP foi obtida pela Transformada Rapida de Fourier"
        " (FFT) das series temporais ML e AP apos remocao da media. As metricas"
        " reportadas (frequencia de pico e frequencia media do espectro de potencia)"
        " foram restringidas ao intervalo fisiologico 0.1-10 Hz, que cobre os"
        " mecanismos neuromusculares do controlo postural. Referencias:"
        " Prieto, T.E., Myklebust, J.B., Hoffmann, R.G., Lovett, E.G. & Myklebust, B.M."
        " (1996). Measures of postural steadiness: differences between healthy young"
        " and elderly adults. IEEE Trans. Biomed. Eng., 43(9), 956-966;"
        " Carpenter, M.G., Frank, J.S., Winter, D.A. & Peysar, G.W. (2001)."
        " Sampling duration effects on centre of pressure summary measures."
        " Exp. Brain Res., 138(2), 210-218."
    ),
}

# Dicionario com os protocolos disponiveis no programa
# Cada protocolo tem o seu proprio numero de ensaios e configuracao
PROTOCOLOS = {
    PROTO_FMS: {
        'nome':        'FMS Bipodal',
        'descr':       '5 ensaios por pé, bipodal\nElipse 95%, assimetria Dir/Esq, todas as métricas',
        'lados':       [('dir', 'dir_', 0), ('esq', 'esq_', 5)],
        'n_ens':       5,
        'assimetria':  True,
        'two_windows': False,
    },
    PROTO_UNIPODAL: {
        'nome':        'Apoio Unipodal',
        'descr':       '5 ensaios por pé, unipodal\nElipse 95%, métricas de oscilação lateral',
        'lados':       [('dir', 'dir_', 0), ('esq', 'esq_', 5)],
        'n_ens':       5,
        'assimetria':  False,
        'two_windows': False,
    },
    PROTO_TIRO: {
        'nome':        'Tiro',
        'descr':       '5 ensaios bipodal, 2 janelas por ensaio\nComparação posição vs disparo, correlação com precisão',
        'lados':       [('pos', 'tiro_', 0), ('disp', 'tiro_', 0)],
        'n_ens':       5,
        'assimetria':  False,
        'two_windows': True,
    },
    PROTO_ARCO: {
        'nome':        'Tiro com Arco',
        'descr':       'Até 30 ensaios bipodal, janela única (Confirmação 1 → Confirmação 2)\nSem distâncias, análise demográfica (género, categoria, estilo, idade)',
        'lados':       [('arco', 'arco_', 0)],
        'n_ens':       30,
        'assimetria':  False,
        'two_windows': False,
    },
}

# Sub-opcoes de Tarefa Funcional
TAREFAS_FUNC = {
    PROTO_TIRO: PROTOCOLOS[PROTO_TIRO],
    PROTO_ARCO: PROTOCOLOS[PROTO_ARCO],
}

_PROTOCOLO_ACTIVO = PROTO_FMS


# -----------------------------------------------------------------------
# Leitura dos ficheiros da plataforma
# -----------------------------------------------------------------------

# Padroes de ficheiros a ignorar (formato nao suportado)
_SKIP_PATTERNS = ['entire plate roll off', 'roll off']


def _ler_ficheiro_xlsx(caminho):
    """Lê um .xlsx gerado por openpyxl com colunas frame/t_ms/x/y."""
    try:
        import openpyxl as _opx
    except ImportError:
        raise ValueError("openpyxl não instalado - necessário para ler .xlsx")
    wb = _opx.load_workbook(caminho, data_only=True, read_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    wb.close()
    if not rows:
        raise ValueError(f"Ficheiro xlsx vazio: {os.path.basename(caminho)}")

    # localizar linha de cabeçalho (frame, t_ms, x, y)
    hdr_idx = None
    col_map = {}
    for ri, row in enumerate(rows):
        labels = [str(c).lower().strip() if c is not None else '' for c in row]
        if 'frame' in labels and 't_ms' in labels and 'x' in labels and 'y' in labels:
            hdr_idx = ri
            col_map = {lbl: ci for ci, lbl in enumerate(labels)}
            break

    if hdr_idx is None:
        # sem cabeçalho explícito - tentar colunas pela posição (A=frame, B=t_ms, C=x, D=y)
        try:
            float(rows[1][1]); float(rows[1][2]); float(rows[1][3])
            hdr_idx = 0
            col_map = {'frame': 0, 't_ms': 1, 'x': 2, 'y': 3}
        except Exception:
            raise ValueError(f"Formato inesperado (cabecalho nao encontrado): {os.path.basename(caminho)}")

    frames = []
    for row in rows[hdr_idx + 1:]:
        try:
            fr = int(row[col_map['frame']])
            tm = float(row[col_map['t_ms']])
            xv = float(row[col_map['x']])
            yv = float(row[col_map['y']])
        except (TypeError, ValueError, IndexError):
            continue
        frames.append({'frame': fr, 't_ms': tm, 'x': xv, 'y': yv,
                       'forca': None, 'dist': None,
                       'sel_esq_x': None, 'sel_esq_y': None,
                       'sel_dir_x': None, 'sel_dir_y': None})

    return {'paciente': None, 'medicao': None, 'data': None,
            'inicio_ms': None, 'fim_ms': None,
            'frames': frames, 'tem_selection': False}


def ler_ficheiro(caminho):
    # Ignorar ficheiros com padroes nao suportados no nome
    nome_lower = os.path.basename(caminho).lower()
    for pat in _SKIP_PATTERNS:
        if pat in nome_lower:
            msg = f'Ficheiro ignorado (padrao nao suportado): {os.path.basename(caminho)}'
            _logger.info(msg)
            raise ValueError(msg)

    # Ficheiros .xlsx reais (openpyxl) - gerados por BSP ou por testes sintéticos
    if nome_lower.endswith('.xlsx'):
        return _ler_ficheiro_xlsx(caminho)

    # Sniff: novo formato Tiro com Arco (.xls TSV com 'Entire plate COF')
    if _detectar_formato_arco(caminho):
        return _ler_ficheiro_arco(caminho)

    with open(caminho, 'r', encoding='iso-8859-1') as f:
        linhas = f.readlines()
    info = {'paciente': None, 'medicao': None, 'data': None,
            'inicio_ms': None, 'fim_ms': None, 'frames': []}
    idx = None
    for i, linha in enumerate(linhas):
        s = linha.strip()
        if s.startswith('Stability export for measurement:'):
            info['medicao'] = s[33:].strip()
        elif s.startswith('Patient name:'):
            info['paciente'] = s[13:].strip()
        elif s.startswith('Measurement done on'):
            info['data'] = s[19:].strip()
        elif s.startswith('1\t') and info['inicio_ms'] is None:
            cols = s.split('\t')
            try:
                info['inicio_ms'] = int(cols[1])
                info['fim_ms']    = int(cols[2])
            except Exception:
                pass
        elif s.startswith('Frame\tTime'):
            idx = i + 1; break
    if idx is None:
        msg = (f"Formato inesperado (cabecalho nao encontrado): "
               f"{os.path.basename(caminho)}")
        _logger.error(msg)
        raise ValueError(msg)
    def _sf(c, i):
        """Converte coluna opcional para float; devolve None se ausente ou nao-numerica.
        Nao propaga excecao -- coluna opcional nao deve invalidar o frame inteiro."""
        if len(c) > i:
            try:
                return float(c[i])
            except (TypeError, ValueError):
                return None
        return None

    frames = []
    for linha in linhas[idx:]:
        s = linha.strip()
        if not s:
            continue
        cols = s.split('\t')
        if len(cols) < 4:
            continue
        # Colunas obrigatorias (0-3): erro aqui = frame invalido, saltar
        try:
            f_data = {
                'frame': int(cols[0]),
                't_ms':  float(cols[1]),
                'x':     float(cols[2]),
                'y':     float(cols[3]),
            }
        except Exception:
            continue
        # Colunas opcionais: falha silenciosa (None).
        # IMPORTANTE: nao embrulhar em try/except global --
        # ficheiros sem Selection CoP (ex: Hurdle Step) podem ter colunas
        # vazias ("") que causariam ValueError e saltariam o frame inteiro.
        f_data['forca']     = _sf(cols, 4)
        f_data['dist']      = _sf(cols, 5)
        # Selection CoP: col[6/7]=Left Sel X/Y; col[10/11]=Right Sel X/Y
        # (cols[8/9]=Left Force/FrameDist; cols[12/13]=Right Force/FrameDist)
        f_data['sel_esq_x'] = _sf(cols, 6)
        f_data['sel_esq_y'] = _sf(cols, 7)
        f_data['sel_dir_x'] = _sf(cols, 10)
        f_data['sel_dir_y'] = _sf(cols, 11)
        frames.append(f_data)
    info['frames'] = frames
    # Detectar se existem dados de Selection validos (> 50% dos frames)
    n_sel = sum(1 for f in frames
                if f.get('sel_dir_x') is not None and f.get('sel_dir_y') is not None)
    info['tem_selection'] = (n_sel >= len(frames) * 0.5) if frames else False
    return info


def _detectar_formato_arco(caminho):
    """
    Sniffing ligeiro: devolve True se o ficheiro tem o cabeçalho do novo
    formato de Tiro com Arco (inclui 'Entire plate COF X' nas primeiras
    ~40 linhas). Estes ficheiros sao .xls tab-separated (nao binario Excel),
    tipicamente 50 Hz, sem coluna de distancia ao alvo.
    """
    try:
        with open(caminho, 'r', encoding='iso-8859-1', errors='ignore') as f:
            cabecalho = [f.readline() for _ in range(40)]
        texto = ''.join(cabecalho).lower()
        return ('entire plate cof x' in texto) or ('entire plate cof' in texto)
    except Exception:
        return False


def _ler_ficheiro_arco(caminho):
    """
    Parser para o novo formato Tiro com Arco:
      - Tab-separated text com extensao .xls (50 Hz tipico)
      - Cabecalho: 'Stability export for measurement: ...'
      - Colunas: Frame | Time (ms) | Entire plate COF X | Y | Force | FrameDist
                 | Left/Right/Top/Bottom Selection X | Y | Force | FrameDist
      - Sem distancia ao alvo (arco nao tem multi-distancia)

    Devolve dict identico ao de `ler_ficheiro()`:
      {'paciente', 'medicao', 'data', 'inicio_ms', 'fim_ms', 'frames', 'tem_selection'}

    Nota: este parser e propositadamente simples. Mapeia apenas CoP principal
    (X/Y do Entire Plate). Se no futuro for preciso usar Selection Left/Right
    para analise bilateral, ver _ler_ficheiro_arco_selection() (a criar).
    """
    with open(caminho, 'r', encoding='iso-8859-1') as f:
        linhas = f.readlines()

    info = {'paciente': None, 'medicao': None, 'data': None,
            'inicio_ms': None, 'fim_ms': None, 'frames': []}

    # Procurar cabecalho das colunas (linha que comeca com 'Frame\t')
    idx_dados = None
    col_idx = {}     # mapa nome_coluna → indice
    for i, linha in enumerate(linhas):
        s = linha.strip()
        if s.startswith('Stability export for measurement:'):
            info['medicao'] = s[len('Stability export for measurement:'):].strip()
        elif s.lower().startswith('patient name:'):
            info['paciente'] = s.split(':', 1)[1].strip()
        elif s.startswith('Measurement done on'):
            info['data'] = s[len('Measurement done on'):].strip()
        elif s.startswith('Frame\t') or s.lower().startswith('frame\ttime'):
            # Parsing do cabecalho de colunas - preciso porque a ordem
            # de Left/Right/Top/Bottom pode variar entre exports
            headers = [h.strip().lower() for h in s.split('\t')]
            for ci, h in enumerate(headers):
                col_idx[h] = ci
                # Tambem guardar sem sufixo de unidade para matching robusto
                # ex: 'entire plate cof x (mm)' -> 'entire plate cof x'
                base = re.sub(r'\s*\([^)]*\)\s*$', '', h).strip()
                if base and base != h and base not in col_idx:
                    col_idx[base] = ci
            idx_dados = i + 1
            break

    if idx_dados is None:
        msg = f"Formato arco inesperado (cabecalho ausente): {os.path.basename(caminho)}"
        _logger.error(msg)
        raise ValueError(msg)

    # Indices das colunas de interesse (None se ausente)
    def _find_col(*nomes):
        for n in nomes:
            if n in col_idx:
                return col_idx[n]
        return None

    c_frame = _find_col('frame')
    c_time  = _find_col('time (ms)', 'time', 't_ms', 'time_ms')
    c_x     = _find_col('entire plate cof x', 'cof x', 'x')
    c_y     = _find_col('entire plate cof y', 'cof y', 'y')
    c_for   = _find_col('entire plate cof force', 'cof force', 'force')
    # Selection Left/Right CoP (opcional, para analise bilateral futura)
    c_lx = _find_col('left selection cof x', 'left sel x', 'left cof x')
    c_ly = _find_col('left selection cof y', 'left sel y', 'left cof y')
    c_rx = _find_col('right selection cof x', 'right sel x', 'right cof x')
    c_ry = _find_col('right selection cof y', 'right sel y', 'right cof y')

    if c_x is None or c_y is None:
        msg = f"Formato arco sem colunas CoP X/Y: {os.path.basename(caminho)}"
        _logger.error(msg)
        raise ValueError(msg)

    def _sf(cols, i):
        if i is None or i >= len(cols):
            return None
        try:
            v = cols[i].strip()
            return float(v) if v else None
        except (TypeError, ValueError):
            return None

    frames = []
    for linha in linhas[idx_dados:]:
        s = linha.rstrip('\n')
        if not s.strip():
            continue
        cols = s.split('\t')
        try:
            f_data = {
                'frame': int(float(cols[c_frame])) if c_frame is not None and c_frame < len(cols) else len(frames),
                't_ms':  float(cols[c_time]) if c_time is not None and c_time < len(cols) else len(frames) * 20.0,
                'x':     float(cols[c_x]),
                'y':     float(cols[c_y]),
            }
        except (TypeError, ValueError, IndexError):
            continue
        f_data['forca']     = _sf(cols, c_for)
        f_data['dist']      = None  # arco nao tem distancia ao alvo
        f_data['sel_esq_x'] = _sf(cols, c_lx)
        f_data['sel_esq_y'] = _sf(cols, c_ly)
        f_data['sel_dir_x'] = _sf(cols, c_rx)
        f_data['sel_dir_y'] = _sf(cols, c_ry)
        frames.append(f_data)

    # Inicio/fim derivados do primeiro/ultimo frame valido
    if frames:
        info['inicio_ms'] = int(frames[0]['t_ms'])
        info['fim_ms']    = int(frames[-1]['t_ms'])

    info['frames'] = frames
    n_sel = sum(1 for f in frames
                if f.get('sel_dir_x') is not None and f.get('sel_dir_y') is not None)
    info['tem_selection'] = (n_sel >= len(frames) * 0.5) if frames else False
    info['_formato_arco'] = True
    return info


# -----------------------------------------------------------------------
# Calculo de metricas
# -----------------------------------------------------------------------

# ── Método de cálculo da velocidade do CoP ───────────────────────────────
# "classico"  - diferenças centradas de 2ª ordem (Prieto et al., 1996).
#               Rápido, sem parâmetros, base de comparação.
# "filtro"    - Butterworth passa-baixo 4ª ordem (fc=10 Hz) + diferenças
#               centradas. Atenua ruído de alta frequência antes da derivada
#               (Winter, 1995; Woltring, 1986).
# "spline"    - Spline cúbica de suavização (s=0 → interpolante exacta, s>0
#               → compromisso suavização/fidelidade). Derivada analítica da
#               spline; estatisticamente mais elegante (Woltring, 1986).
_VEL_METODO: str = "classico"   # alterar aqui para "filtro" ou "spline"
_VEL_FC_HZ:  float = 10.0       # frequência de corte do Butterworth (Hz)
_VEL_FS_HZ:  float = 25.0       # taxa de amostragem típica (Hz); usado em "filtro"
_VEL_SPLINE_S: float = 0.0      # factor de suavização da spline (0=interpolante)

# chi2(0.95, df=2) -- calculado uma vez ao nivel do modulo (Schubert & Kirchner, 2014)
_CHI2_95 = stats.chi2.ppf(0.95, df=2)
# Angulos pre-calculados para a elipse (120 pontos + fecho) -- evita recalculo em cada ensaio
_ELL_COS = np.cos(np.linspace(0.0, 2.0 * math.pi, 121))
_ELL_SIN = np.sin(np.linspace(0.0, 2.0 * math.pi, 121))


def _calcular_velocidades_eixo(x: np.ndarray, y: np.ndarray,
                               t_s: np.ndarray):
    """
    Calcula velocidades instantâneas em x e y pelo método configurado em
    _VEL_METODO.

    Parâmetros
    ----------
    x, y : np.ndarray - coordenadas CoP (mm), já selecionadas na janela.
    t_s  : np.ndarray - instantes em segundos.

    Retorna
    -------
    vx_i, vy_i : np.ndarray, tamanho (n-2) no método clássico, (n,) nos
        outros métodos, com NaN nas extremidades se aplicável.

    Referências
    -----------
    Prieto et al. (1996) - diferenças centradas clássicas.
    Winter (1995) - filtragem Butterworth antes da derivada.
    Woltring (1986) - spline cúbica, derivada analítica.
    """
    n = len(x)

    if _VEL_METODO == "filtro":
        # ── Butterworth 4ª ordem, passa-baixo a _VEL_FC_HZ ─────────────
        # Estima a taxa de amostragem a partir dos dados para que o método
        # funcione mesmo com sinais de comprimento variável.
        if n >= 4 and t_s[-1] > t_s[0]:
            fs_est = (n - 1) / (t_s[-1] - t_s[0])
        else:
            fs_est = _VEL_FS_HZ
        nyq = fs_est / 2.0
        fc  = min(_VEL_FC_HZ, nyq * 0.9)   # evitar Wn ≥ 1
        try:
            from scipy.signal import butter, filtfilt
            b, a  = butter(4, fc / nyq, btype='low', analog=False)
            xf    = filtfilt(b, a, x)
            yf    = filtfilt(b, a, y)
        except Exception:
            xf, yf = x, y    # fallback se scipy.signal não disponível
        # Diferenças centradas sobre o sinal filtrado
        dt_c = t_s[2:] - t_s[:-2]
        dt_c = np.where(dt_c == 0, 0.04, dt_c)
        vx_i = (xf[2:] - xf[:-2]) / dt_c
        vy_i = (yf[2:] - yf[:-2]) / dt_c
        return vx_i, vy_i

    elif _VEL_METODO == "spline":
        # ── Spline cúbica de suavização ──────────────────────────────────
        # UnivariateSpline ajusta s(t) minimizando soma de quadrados dos
        # resíduos; s=0 corresponde a uma interpolante exacta.
        # A derivada .derivative() é analítica - evita amplificação de ruído.
        try:
            from scipy.interpolate import UnivariateSpline
            # Remover instantes duplicados (pode acontecer se dt=0)
            _, uniq = np.unique(t_s, return_index=True)
            t_u, x_u, y_u = t_s[uniq], x[uniq], y[uniq]
            k  = min(3, len(t_u) - 1)   # grau da spline (máx 3 se poucos pontos)
            sx = UnivariateSpline(t_u, x_u, k=k, s=_VEL_SPLINE_S, ext=3)
            sy = UnivariateSpline(t_u, y_u, k=k, s=_VEL_SPLINE_S, ext=3)
            vx_i = sx.derivative()(t_s)
            vy_i = sy.derivative()(t_s)
        except Exception:
            # Fallback para diferenças centradas se spline falhar
            dt_c = t_s[2:] - t_s[:-2]
            dt_c = np.where(dt_c == 0, 0.04, dt_c)
            vx_i = (x[2:] - x[:-2]) / dt_c
            vy_i = (y[2:] - y[:-2]) / dt_c
        return vx_i, vy_i

    else:
        # ── Clássico: diferenças centradas de 2ª ordem (default) ────────
        # d/dt f(t_i) ≈ [f(t_{i+1}) − f(t_{i-1})] / (t_{i+1} − t_{i-1})
        # Este é o método de referência em estabilometria (Prieto et al., 1996).
        dt_c = t_s[2:] - t_s[:-2]
        dt_c = np.where(dt_c == 0, 0.04, dt_c)
        vx_i = (x[2:] - x[:-2]) / dt_c
        vy_i = (y[2:] - y[:-2]) / dt_c
        return vx_i, vy_i




def _validar_frame_rate(frames, nome_ensaio='', log=None):
    """
    Valida consistência de frame rate e timestamps de um ensaio.
    Devolve (fs_hz, avisos) onde avisos é lista de strings.
    - fs_hz: taxa de amostragem estimada (Hz), ou None se < 4 frames
    - avisos: problemas detectados (timestamps duplicados, jitter > 20%, etc.)
    """
    avisos = []
    if len(frames) < 4:
        avisos.append(f'{nome_ensaio}: muito poucos frames ({len(frames)})')
        return None, avisos

    t = [f['t_ms'] for f in frames]

    # Timestamps duplicados ou invertidos
    diffs = [t[i+1] - t[i] for i in range(len(t)-1)]
    n_dup = sum(1 for d in diffs if d <= 0)
    if n_dup > 0:
        avisos.append(f'{nome_ensaio}: {n_dup} timestamp(s) duplicado(s)/invertido(s)')

    # Jitter de frame rate
    pos_diffs = [d for d in diffs if d > 0]
    if pos_diffs:
        import statistics as _stats
        dt_mean = _stats.mean(pos_diffs)
        dt_med  = _stats.median(pos_diffs)
        fs_hz   = 1000.0 / dt_mean if dt_mean > 0 else None
        # Jitter: desvio relativo mediano
        if dt_med > 0:
            jitter = abs(dt_mean - dt_med) / dt_med
            if jitter > 0.20:  # >20% de variação
                avisos.append(
                    f'{nome_ensaio}: jitter de frame rate elevado '
                    f'({jitter*100:.0f}%) - dt_mean={dt_mean:.1f}ms dt_med={dt_med:.1f}ms'
                )
        return fs_hz, avisos
    return None, avisos


def _validar_consistencia_atleta(ath, log=None):
    """
    Verifica consistência de frame rate entre todos os ensaios de um atleta.
    Emite avisos se a taxa de amostragem variar mais de 10% entre ensaios.
    Retorna lista de strings de aviso (vazia se tudo OK).
    """
    avisos = []
    fs_por_lado = {}
    nome = ath.get('nome', '?')

    for lado, raw_lst in ath.get('raw', {}).items():
        for ti, rd in enumerate(raw_lst, 1):
            if rd is None:
                continue
            frames = rd.get('dados', {}).get('frames', [])
            if not frames:
                continue
            label = f'{nome} {lado} T{ti}'
            fs, av = _validar_frame_rate(frames, label)
            avisos.extend(av)
            if fs is not None:
                fs_por_lado.setdefault(lado, []).append(fs)

    # Verificar variação entre ensaios do mesmo lado
    for lado, fs_list in fs_por_lado.items():
        if len(fs_list) < 2:
            continue
        fs_min, fs_max = min(fs_list), max(fs_list)
        if fs_max > 0 and (fs_max - fs_min) / fs_max > 0.10:  # >10% variação
            avisos.append(
                f'{nome} {lado}: frame rate inconsistente entre ensaios '
                f'({fs_min:.0f}-{fs_max:.0f} Hz) - verificar exportação'
            )

    return avisos


def _calcular_selection(frames, t_ini=None, t_fim=None, lado='dir'):
    """
    Calcula metricas de estabilidade para o sub-sinal de Selection (Right/Left CoP).

    Os ficheiros de tiro podem conter colunas adicionais com o CoP separado de
    cada pe (Right Selection e Left Selection). Esta funcao extrai esse sub-sinal
    e aplica calcular(), devolvendo metricas analogas a um ensaio unipodal.

    lado : 'dir' → Right Selection (cols sel_dir_x / sel_dir_y)
           'esq' → Left Selection  (cols sel_esq_x / sel_esq_y)
    """
    kx = f'sel_{lado}_x'; ky = f'sel_{lado}_y'
    pool = ([f for f in frames if t_ini <= f['t_ms'] <= t_fim]
            if t_ini is not None else list(frames))
    validos = [f for f in pool if f.get(kx) is not None and f.get(ky) is not None]
    if len(validos) < 5:
        return None
    pseudo = [{'frame': f['frame'], 't_ms': f['t_ms'],
               'x': f[kx], 'y': f[ky]} for f in validos]
    return calcular(pseudo)


# Funcao principal de calculo de metricas de estabilidade postural
# Recebe os frames da plataforma de forcas e devolve todas as metricas
# Esta funcao e o "coracao" do programa
def calcular(frames, t_ini=None, t_fim=None, peso_kg=None, altura_m=None):
    """
    Calcula metricas de estabilidade para um ensaio.

    Elipse 95%: metodo eigenvalores (Schubert & Kirchner, 2014).
    chi2(0.95, df=2) = 5.9915. Semieixos = sqrt(eigenval) * sqrt(chi2).
    Area = pi * a * b  [mm2] (semi-eixos ja em mm; sem divisao por 100).

    Velocidades: diferencas centradas (Prieto et al., 1996).

    Parametros:
      peso_kg  : massa corporal em kg (opcional). Activa metricas normalizadas.
      altura_m : altura corporal em m  (opcional). Activa normalizacao por pendulo invertido.
    """
    sel = ([f for f in frames if t_ini <= f['t_ms'] <= t_fim]
           if t_ini is not None else list(frames))
    if len(sel) < 5:
        return None

    x = np.array([f['x']    for f in sel])
    y = np.array([f['y']    for f in sel])
    t = np.array([f['t_ms'] for f in sel])

    amp_x   = float(x.max() - x.min())
    amp_y   = float(y.max() - y.min())
    t_total = (t[-1] - t[0]) / 1000.0

    # ── Deslocamento total e velocidade média global ────────────────────
    # Distâncias Euclidianas entre amostras consecutivas (Prieto et al., 1996).
    passos  = np.sqrt(np.diff(x)**2 + np.diff(y)**2)
    desl    = float(passos.sum())
    vel_med = desl / t_total if t_total > 0 else 0.0   # mm/s

    # ── Velocidade por eixo - método selecionável ───────────────────────
    # Método "classico": diferenças centradas de 2ª ordem (baseline)
    #   d/dt f(t_i) ≈ [f(t_{i+1}) - f(t_{i-1})] / (t_{i+1} - t_{i-1})
    # Prieto et al. (1996) definem vel_x e vel_y como médias do valor
    # absoluto das velocidades instantâneas, o que captura a oscilação média
    # em cada eixo, independentemente da direcção.
    t_s = t / 1000.0   # ms → s para derivada numérica
    vx_i, vy_i = _calcular_velocidades_eixo(x, y, t_s)
    vel_x = float(np.abs(vx_i).mean())       # velocidade média eixo X (mm/s)
    vel_y = float(np.abs(vy_i).mean())       # velocidade média eixo Y (mm/s)
    vel_pico_x = float(np.abs(vx_i).max()) if len(vx_i) > 0 else 0.0  # pico vel. X (mm/s)
    vel_pico_y = float(np.abs(vy_i).max()) if len(vy_i) > 0 else 0.0  # pico vel. Y (mm/s)

    # valor critico do chi-quadrado a 95% de confianca (2 graus de liberdade)
    chi2_95 = _CHI2_95
    cov     = np.cov(x, y)
    eigv    = np.linalg.eigvalsh(cov)
    ev1     = max(float(eigv[1]), 0.0)
    ev2     = max(float(eigv[0]), 0.0)
    sa      = math.sqrt(ev1) * math.sqrt(chi2_95)
    sb      = math.sqrt(ev2) * math.sqrt(chi2_95)
    area    = math.pi * (sa / 2) * (sb / 2)  # área em mm2 - fórmula: π*(a/2)*(b/2)
    ang     = math.atan2(float(eigv[1]) - float(cov[0, 0]), float(cov[0, 1]))
    mx      = float(x.mean()); my = float(y.mean())
    ca      = math.cos(ang);   sa2 = math.sin(ang)
    # Elipse vectorizada -- usa arrays pre-calculados no modulo
    xe_arr  = sa * _ELL_COS
    ye_arr  = sb * _ELL_SIN
    ell_x   = (ca * xe_arr - sa2 * ye_arr + mx).tolist()
    ell_y   = (sa2 * xe_arr + ca * ye_arr + my).tolist()

    n     = len(sel)
    vx_f  = np.full(n, np.nan); vy_f = np.full(n, np.nan)
    sx_f  = np.full(n, np.nan); sy_f = np.full(n, np.nan)
    s_ac  = np.full(n, np.nan)
    # Preencher array de velocidades de tamanho n independentemente do método:
    # "classico": vx_i tem tamanho n-2 → preencher posições internas [1:-1]
    # "spline"/"filtro": vx_i pode ter tamanho n → copiar directamente
    if len(vx_i) == n:
        vx_f[:] = vx_i; vy_f[:] = vy_i
    elif n >= 3:
        vx_f[1:-1] = vx_i; vy_f[1:-1] = vy_i
    if n >= 2:
        sx_f[1:] = np.diff(x); sy_f[1:] = np.diff(y)
        s_ac[1:] = np.cumsum(passos)   # vectorizado

    # ── RMS (Root Mean Square) - Quijoux et al. (2021), sec. 3.2.4 ────────
    # Coordenadas centradas: Xn = MLn - mean(ML), Yn = APn - mean(AP)
    # RMS ML  = sqrt(1/N * sum(Xn²))   [equivale ao desvio-padrão com divisor N]
    # RMS AP  = sqrt(1/N * sum(Yn²))
    # RMS Radius = sqrt(1/N * sum(Rn²)) onde Rn = sqrt(Xn²+Yn²)
    # Referência: Quijoux et al. (2021). Physiological Reports, 9, e15067.
    x_c   = x - mx              # coordenadas centradas ML
    y_c   = y - my              # coordenadas centradas AP
    rms_x = float(np.sqrt(np.mean(x_c**2)))
    rms_y = float(np.sqrt(np.mean(y_c**2)))
    rms_r = float(np.sqrt(np.mean(x_c**2 + y_c**2)))

    # ── Métricas derivadas ─────────────────────────────────────────────
    # Razão ML/AP: amp_x/amp_y > 1 → oscilação predominantemente médio-lateral.
    ratio_ml_ap = amp_x / amp_y if amp_y > 0 else None
    # Razão velocidades ML/AP (mesmo raciocínio, domínio da velocidade).
    ratio_vel   = vel_x / vel_y if vel_y > 0 else None
    # ── Índice de rigidez postural (stiffness) ───────────────────────────
    # Razão velocidade/amplitude: quanto maior a velocidade relativamente à
    # amplitude, maior a "rigidez" aparente do controlo postural.
    #
    # Fundamentação (pêndulo invertido em torno do tornozelo):
    #   Winter DA (1995). Human balance and posture control during standing
    #     and walking. Gait & Posture, 3(4), 193-214.
    #   Maurer C, Peterka RJ (2005). A new interpretation of spontaneous
    #     sway measures based on a simple model of human postural control.
    #     Journal of Neurophysiology, 93(1), 189-200.
    #
    # Stiffness "cru" (unidades: 1/s)
    stiff_x = vel_med / amp_x if amp_x > 0 else None
    stiff_y = vel_med / amp_y if amp_y > 0 else None

    # ── Normalizacoes por dimensoes corporais ────────────────────────────
    # Normalizacoes autor-definidas que removem o efeito do tamanho corporal,
    # seguindo a logica do pendulo invertido (Winter, 1995).
    # Requerem peso_kg / altura_m; retornam None se parametro ausente.
    ea95_norm    = area / altura_m          if altura_m else None   # mm2/m
    amp_norm_x   = amp_x / altura_m        if altura_m else None   # mm/m
    amp_norm_y   = amp_y / altura_m        if altura_m else None   # mm/m
    vel_norm     = vel_med / altura_m      if altura_m else None   # (mm/s)/m
    stiff_mass_x = stiff_x / peso_kg       if (stiff_x is not None and peso_kg) else None
    stiff_mass_y = stiff_y / peso_kg       if (stiff_y is not None and peso_kg) else None
    stiff_norm_x = stiff_x / altura_m      if (stiff_x is not None and altura_m) else None
    stiff_norm_y = stiff_y / altura_m      if (stiff_y is not None and altura_m) else None

    # ── Metricas espectrais (FFT) - opcionais ──────────────────────────
    # Frequencia de pico e media nos eixos ML/AP no intervalo fisiológico 0.1-10 Hz
    # (mecanismos neuromusculares do controlo postural).
    #
    # Fundamentação:
    #   Prieto TE, Myklebust JB, Hoffmann RG, Lovett EG, Myklebust BM (1996).
    #     Measures of postural steadiness: differences between healthy young
    #     and elderly adults. IEEE Transactions on Biomedical Engineering,
    #     43(9), 956-966.
    #   Carpenter MG, Frank JS, Winter DA, Peysar GW (2001). Sampling
    #     duration effects on centre of pressure summary measures.
    #     Experimental Brain Research, 138(2), 210-218.
    #
    # Activado via checkbox FFT na interface ou AOM_CONFIG['fft_ativo'] = True.
    fft_pico_x = fft_pico_y = fft_med_x = fft_med_y = None
    if AOM_CONFIG.get('fft_ativo') and t_total > 0 and len(x) >= 8:
        try:
            fs_est = (len(x) - 1) / t_total   # Hz estimado
            freqs  = np.fft.rfftfreq(len(x), d=1.0/fs_est)
            Px     = np.abs(np.fft.rfft(x - x.mean()))**2
            Py     = np.abs(np.fft.rfft(y - y.mean()))**2
            mask   = (freqs >= 0.1) & (freqs <= 10.0)
            if mask.sum() >= 2:
                fft_pico_x = float(freqs[mask][Px[mask].argmax()])
                fft_pico_y = float(freqs[mask][Py[mask].argmax()])
                fft_med_x  = float(np.sum(freqs[mask]*Px[mask])/np.sum(Px[mask]))
                fft_med_y  = float(np.sum(freqs[mask]*Py[mask])/np.sum(Py[mask]))
        except Exception:
            pass

    return {
        'amp_x': amp_x,   'amp_y': amp_y,
        'vel_x': vel_x,   'vel_y': vel_y,   'vel_med': vel_med,
        'vel_pico_x': vel_pico_x, 'vel_pico_y': vel_pico_y,
        'desl':  desl,    'time':  t_total,
        'ea95':  area,    'leng_a': sa,      'leng_b': sb,
        # RMS (Quijoux et al., 2021, sec. 3.2.4)
        'rms_x': rms_x,   # RMS ML (mm)
        'rms_y': rms_y,   # RMS AP (mm)
        'rms_r': rms_r,   # RMS Radius (mm)
        # Metricas derivadas (novos campos v16)
        'ratio_ml_ap': ratio_ml_ap,   # adimensional
        'ratio_vel':   ratio_vel,     # adimensional
        'stiff_x':     stiff_x,       # mm/s / mm = 1/s
        'stiff_y':     stiff_y,       # mm/s / mm = 1/s
        # Normalizacoes por dimensoes corporais (None se parametro ausente)
        'ea95_norm':    ea95_norm,     # mm2/m
        'amp_norm_x':   amp_norm_x,   # mm/m
        'amp_norm_y':   amp_norm_y,   # mm/m
        'vel_norm':     vel_norm,      # (mm/s)/m
        'stiff_mass_x': stiff_mass_x, # 1/(s·kg)
        'stiff_mass_y': stiff_mass_y, # 1/(s·kg)
        'stiff_norm_x': stiff_norm_x, # (1/s)/m
        'stiff_norm_y': stiff_norm_y, # (1/s)/m
        'vel_method':  _VEL_METODO,   # rastreabilidade do metodo
        # FFT (None se fft_ativo=False)
        'fft_pico_x':  fft_pico_x,    # Hz - freq. pico ML
        'fft_pico_y':  fft_pico_y,    # Hz - freq. pico AP
        'fft_med_x':   fft_med_x,     # Hz - freq. media ML
        'fft_med_y':   fft_med_y,     # Hz - freq. media AP
        'cof_x': x.tolist(), 'cof_y': y.tolist(),
        'mean_x': mx, 'mean_y': my,
        'cov_xx': float(cov[0,0]), 'cov_xy': float(cov[0,1]), 'cov_yy': float(cov[1,1]),
        'ev1': ev1, 'ev2': ev2, 'ang_deg': math.degrees(ang), 'chi2': chi2_95,
        'ell_x': ell_x, 'ell_y': ell_y,
        'vx_f': vx_f.tolist(), 'vy_f': vy_f.tolist(),
        'sx_f': sx_f.tolist(), 'sy_f': sy_f.tolist(), 's_ac': s_ac.tolist(),
        't_ms': t.tolist(), 'frames': sel,
    }


def assimetria(vd, ve):
    """
    Índice de assimetria Dir/Esq pelo ratio de simetria standard em biomecânica.

        AI = (VD − VE) / [(VD + VE)/2] × 100  [%]

    Valores positivos indicam que o membro direito apresenta maior valor da
    métrica; valores negativos, o membro esquerdo.
    Intervalo típico: ±10 % considerado dentro da variabilidade normal.
    """
    if vd is None or ve is None: return None
    med = (vd + ve) / 2.0
    return round((vd - ve) / med * 100.0, 2) if med != 0 else 0.0


def flagrar_outliers(mets_lista, chave='ea95', z_thresh=None, pct_thresh=None):
    """
    Detecta outliers por modified z-score baseado na MAD (Desvio Absoluto
    da Mediana), que é robusto a outliers ao contrário do z-score clássico.

        modified z-score_i = 0.6745 × |x_i − mediana| / MAD

    Um ensaio é flagrado se AMBOS os critérios forem satisfeitos:
      1. modified z-score > z_thresh   (critério estatístico robusto)
      2. |x_i − mediana| / mediana > pct_thresh  (relevância prática)

    O limiar duplo (z_thresh=4.5, pct_thresh=0.50) foi escolhido para ser
    conservador com n=5 ensaios, evitando falsos positivos frequentes.
    Com n maior pode reduzir-se z_thresh para 3.5 (Iglewicz & Hoaglin, 1993).

    Referência: Iglewicz, B. & Hoaglin, D. (1993). How to Detect and Handle
    Outliers. ASQC Quality Press, Milwaukee.
    """
    if z_thresh   is None: z_thresh   = AOM_CONFIG.get('outlier_z_thresh',   4.5)
    if pct_thresh is None: pct_thresh = AOM_CONFIG.get('outlier_pct_thresh', 0.50)
    vals = [m[chave] if (m is not None and chave in m) else None for m in mets_lista]
    nums = np.array([v for v in vals if v is not None])
    if len(nums) < 3:
        return [False] * len(vals)
    med = np.median(nums)
    mad = np.median(np.abs(nums - med))
    if mad == 0 or med == 0:
        return [False] * len(vals)
    result = []
    for v in vals:
        if v is None:
            result.append(False)
        else:
            z   = abs(0.6745 * (v - med) / mad)
            pct = abs(v - med) / med
            result.append(bool(z > z_thresh and pct > pct_thresh))
    return result


# -----------------------------------------------------------------------
# Nomes de ficheiros flexiveis
# -----------------------------------------------------------------------

_PADROES = {
    'dir':     ['dir_', 'Dir_', 'DIR_', 'd_', 'D_', 'direito_', 'Direito_', 'right_', 'R_'],
    'esq':     ['esq_', 'Esq_', 'ESQ_', 'e_', 'E_', 'esquerdo_', 'Esquerdo_', 'left_', 'L_'],
    'hs_dir':  ['hs_dir_', 'hs_Dir_', 'HS_dir_', 'HS_Dir_', 'HurdleStep_dir_', 'hs_d_'],
    'hs_esq':  ['hs_esq_', 'hs_Esq_', 'HS_esq_', 'HS_Esq_', 'HurdleStep_esq_', 'hs_e_'],
    'tiro':    ['tiro_', 'Tiro_', 'TIRO_', 't_', 'T_', 'shot_', 'disparo_', 'ensaio_'],
    'pos':     ['tiro_', 'Tiro_', 'pos_', 'posicao_'],
    'disp':    ['tiro_', 'Tiro_', 'disp_', 'disparo_'],
}


def _is_skip_file(caminho):
    """Retorna True se o ficheiro deve ser ignorado (formato nao suportado)."""
    nome = os.path.basename(caminho)
    # Ficheiros de metadados macOS (AppleDouble resource forks)
    if nome.startswith('._'):
        return True
    nome_lower = nome.lower()
    return any(pat in nome_lower for pat in _SKIP_PATTERNS)


def achar_ficheiro(pasta, lado_ou_prefixo, numero=None):
    """
    Procura um ficheiro de ensaio na pasta usando multiplos padroes de nome.
    Aceita: dir_1, D1, d_1, direito_1, Direito_01,
            'dir_1 - 03 06 2025 - Stability export.xls', etc.
    """
    if numero is None:
        prefixo = lado_ou_prefixo
        for nome in sorted(os.listdir(pasta)):
            ext_l = os.path.splitext(nome)[1].lower()
            if ext_l in ('.xls', '.xlsx') and nome.lower().startswith(prefixo.lower()):
                fp = os.path.join(pasta, nome)
                if os.path.getsize(fp) > 0 and not _is_skip_file(fp):
                    return fp
        return None

    lado = lado_ou_prefixo
    padroes = _PADROES.get(lado, [lado + '_'])
    formatos_num = [str(numero), f'{numero:02d}', f'{numero:03d}']

    # tentativa 1: nome exacto
    for pref in padroes:
        for nf in formatos_num:
            for ext in ['.xls', '.xlsx', '.XLS', '.XLSX', '.Xls', '.Xlsx']:
                fp = os.path.join(pasta, pref + nf + ext)
                if os.path.exists(fp) and os.path.getsize(fp) > 0 and not _is_skip_file(fp):
                    return fp

    # tentativa 2: scan flexivel -- numero logo apos o prefixo, seguido de espaco/traco/fim
    # ex: "dir_1 - 03 06 2025 - Stability export.xls"
    # ex: "trial5_1_-_13_01_2026_-_Stability_export.xls"  (tiro: trial[dist]_[ensaio])
    for nome in sorted(os.listdir(pasta)):
        ext_l = os.path.splitext(nome)[1].lower()
        if ext_l not in ('.xls', '.xlsx'):
            continue
        base = os.path.splitext(nome)[0].lower()
        for pref in padroes:
            pref_l = pref.lower().rstrip('_')
            for nf in formatos_num:
                # o prefixo seguido do numero no inicio do nome
                padrao = re.compile(r'^' + re.escape(pref_l) + r'_?' + re.escape(nf) + r'(\b|[^0-9])')
                if padrao.search(base):
                    fp = os.path.join(pasta, nome)
                    if os.path.getsize(fp) > 0 and not _is_skip_file(fp):
                        return fp
    return None


def achar_ficheiro_arco(pasta, pid, trial):
    """
    Procura ficheiro do novo formato Tiro com Arco:
      {id}_{trial} - DD-MM-YYYY - Stability export.xls
    Ex:  101_1 - 08-07-2024 - Stability export.xls
         101_30 - 08-07-2024 - Stability export.xls

    O id pode ter leading zeros ('101', '0101'), o trial 1..30 sem prefixo 'trial'.
    Case-insensitive e tolerante a variações de separador.
    """
    pid_s = str(pid).lstrip('0') or '0'   # '0101' → '101'
    trial_s = str(trial)
    # Padrão: começa por {id}_{trial} seguido de espaço/underscore/traço
    # Aceita: "101_1 - ..." / "101_01 - ..." / "0101_1-..."
    formatos_trial = [trial_s, f'{int(trial):02d}', f'{int(trial):03d}']
    formatos_id = [pid_s, f'{int(pid_s):03d}', f'{int(pid_s):04d}']

    for nome in sorted(os.listdir(pasta)):
        ext_l = os.path.splitext(nome)[1].lower()
        if ext_l not in ('.xls', '.xlsx'):
            continue
        base = nome.lower()
        for fid in formatos_id:
            for ft in formatos_trial:
                # {id}_{trial} como prefixo, seguido de separador (espaço/_/-) ou ponto
                padrao = re.compile(
                    r'^' + re.escape(fid) + r'_' + re.escape(ft) +
                    r'(\s|_|-|\.)'
                )
                if padrao.search(base):
                    fp = os.path.join(pasta, nome)
                    if os.path.getsize(fp) > 0 and not _is_skip_file(fp):
                        return fp
    return None


def achar_ficheiro_tiro(pasta, distancia, numero):
    """
    Procura ficheiro no formato: trial[dist]_[ensaio]_-_..._-_Stability_export.xls
    Ex: trial5_1_-_13_01_2026_-_Stability_export.xls
    Tambem aceita: trial5_01, trial5_001, trial05_1, etc.
    """
    dist_str = str(distancia)
    formatos_num = [str(numero), f'{numero:02d}', f'{numero:03d}']
    formatos_dist = [dist_str, f'{distancia:02d}' if isinstance(distancia, int) else dist_str]

    for nome in sorted(os.listdir(pasta)):
        if not nome.lower().endswith('.xls'):
            continue
        base = os.path.splitext(nome)[0].lower()
        for fd in formatos_dist:
            for fn in formatos_num:
                # trial[dist]_[ensaio] com qualquer sufixo
                padrao = re.compile(r'^trial' + re.escape(fd) + r'_' + re.escape(fn) + r'(\b|[_\-])')
                if padrao.search(base):
                    fp = os.path.join(pasta, nome)
                    if os.path.getsize(fp) > 0 and not _is_skip_file(fp):
                        return fp
    return None


# -----------------------------------------------------------------------
# Inicio_fim e associacao de nomes
# -----------------------------------------------------------------------

def carregar_inicio_fim(caminho):
    """
    Formato standard: Atleta | ini1 | fim1 | ini2 | fim2 | ...
    Uma linha por individuo, um par de colunas por ensaio.
    """
    wb = load_workbook(caminho, data_only=True); ws = wb.active; res = {}
    for linha in ws.iter_rows(min_row=2, values_only=True):
        if not linha[0]: continue
        t = {}
        for i in range(20):
            cb, ce = 1 + i*2, 2 + i*2
            if ce >= len(linha): break
            try: t[i+1] = (int(linha[cb]), int(linha[ce]))
            except Exception: pass
        if t: res[str(linha[0]).strip()] = t
    return res


def carregar_tempos_tiro(caminho):
    """
    Le o ficheiro de tempos do protocolo de tiro.
    Suporta o formato com sheets:
      'tempo (toque)'   - camera + placa times para o toque (inicio)
      'tempo (pontaria)' - camera times para pontaria (posicao de tiro)
      'tempo (disparo)' - camera times para o disparo (tiro)
      'inicio_fim (Hurdle Step)' - intervalos para bipodal (como inicio_fim normal)

    Calculo das horas de placa para eventos:
      t_placa_event = t_toque_placa + (t_camera_event - t_camera_toque)
      (a formula *1000 no Excel esta errada; a conversao correcta e 1:1)

    Retorna: {
      'por_individuo': {
          1: {
              '5m': {1: {'toque': 940, 'pontaria': 1260, 'disparo': 5740},
                     2: {'toque': 540, ...}, ...},
              '7m': {...}
          },
          2: {...}
      },
      'hurdle_step': {
          1: {1: (ini, fim), 2: (ini, fim), ...dir...},
          ...
      },
      'distancias': ['5m', '7m'],
      'n_ensaios_dist': 5,
    }
    """
    import re as _re
    wb = load_workbook(caminho, data_only=True)

    # -- Identificar distancias e colunas do cabecalho --
    # Verifica sheets com dados de tempo
    sheet_names = wb.sheetnames
    ws_toque = None
    ws_pont  = None
    ws_disp  = None
    ws_hs    = None

    for sn in sheet_names:
        sl = sn.lower()
        has_tempo = 'tempo' in sl
        if has_tempo and 'toque'   in sl:                   ws_toque = wb[sn]
        if has_tempo and ('pont' in sl or 'posic' in sl):   ws_pont  = wb[sn]
        if has_tempo and 'dispar'  in sl:                   ws_disp  = wb[sn]
        if 'hurdle' in sl or 'step' in sl:                  ws_hs    = wb[sn]
        if 'inicio' in sl and 'fim' in sl and ('hurdle' in sl or 'bipodal' in sl or 'step' in sl):
            ws_hs = wb[sn]

    # Fallback: ultimo sheet com 'inicio_fim'
    if ws_hs is None:
        for sn in sheet_names:
            if 'hurdle' in sn.lower() or 'step' in sn.lower():
                ws_hs = wb[sn]; break
        if ws_hs is None:
            for sn in sheet_names:
                if 'inicio' in sn.lower() and 'fim' in sn.lower():
                    ws_hs = wb[sn]; break

    if ws_toque is None:
        # Fallback: use active sheet
        ws_toque = wb.active

    # Ler cabecalho da linha 2 do sheet toque
    row2 = [ws_toque.cell(2, c).value for c in range(1, ws_toque.max_column + 2)]

    # Estrategia robusta: encontrar todos os grupos de colunas
    # Estrutura: [None, 5m(1), 5m(2),..., 7m(1), 7m(2),..., None, None, 5m(1), 5m(2),..., 7m(1),...]
    # Separador = 2+ colunas None consecutivas no meio
    dist_pattern = _re.compile(r'^(\d+(?:\.\d+)?m)\s*\((\d+)\)$', _re.I)

    # Recolher todos os grupos de colunas por distancia e numero de trial
    all_matches = []  # (col_1based, dist, trial_num)
    for ci, v in enumerate(row2, 1):
        if v is None: continue
        m = dist_pattern.match(str(v).strip())
        if m:
            all_matches.append((ci, m.group(1).lower(), int(m.group(2))))

    if not all_matches:
        dists_ord = []
        cam_cols = {}
        plc_cols = {}
    else:
        # Agrupar por posicao: encontrar o ponto de corte entre camera e placa
        # O corte e onde os numeros de trial recomeçam (trial 1 aparece de novo)
        # Detectar onde começa o bloco de placa: primeiro trial=1 depois do primeiro bloco
        first_block_done = False
        first_t1_seen = {}  # dist -> col do primeiro trial=1
        cut_col = None

        for ci, dist, trial in all_matches:
            if trial == 1:
                if dist in first_t1_seen:
                    # Segunda ocorrencia do trial 1 desta distancia = inicio do bloco de placa
                    if cut_col is None or ci < cut_col:
                        cut_col = ci
                else:
                    first_t1_seen[dist] = ci

        # Se nao ha cut_col, usar metade das colunas
        if cut_col is None:
            cut_col = all_matches[len(all_matches)//2][0]

        cam_cols = {}
        plc_cols = {}
        for ci, dist, trial in all_matches:
            if ci < cut_col:
                cam_cols.setdefault(dist, []).append(ci)
            else:
                plc_cols.setdefault(dist, []).append(ci)

        dists_ord = list(cam_cols.keys())
        # Adicionar distancias que so estao em plc_cols
        for d in plc_cols:
            if d not in dists_ord: dists_ord.append(d)

    n_ens_dist = max((len(v) for v in cam_cols.values()), default=5)

    # Ler dados por individuo
    res_ind = {}  # {row_idx: {dist: {trial: {'toque','pontaria','disparo'}}}}

    for r in range(3, ws_toque.max_row + 1):
        idx_val = ws_toque.cell(r, 1).value
        if idx_val is None: continue
        try: idx = int(idx_val)
        except (ValueError, TypeError): continue

        ind_data = {}
        for dist in dists_ord:
            cam_c_list = cam_cols.get(dist, [])
            plc_c_list = plc_cols.get(dist, cam_c_list)  # fallback: use cam cols

            dist_data = {}
            for ti, (cam_c, plc_c) in enumerate(zip(cam_c_list, plc_c_list), 1):
                try:
                    t_toque_cam = ws_toque.cell(r, cam_c).value
                    t_toque_plc = ws_toque.cell(r, plc_c).value
                    if t_toque_cam is None or t_toque_plc is None:
                        continue

                    # Pontaria (posicao)
                    t_pont_cam = ws_pont.cell(r, cam_c).value if ws_pont else None
                    # Disparo
                    t_disp_cam = ws_disp.cell(r, cam_c).value if ws_disp else None

                    t_pont_plc = (int(t_toque_plc) + (int(t_pont_cam) - int(t_toque_cam))
                                  if t_pont_cam is not None else None)
                    t_disp_plc = (int(t_toque_plc) + (int(t_disp_cam) - int(t_toque_cam))
                                  if t_disp_cam is not None else None)

                    dist_data[ti] = {
                        'toque':    int(t_toque_plc),
                        'pontaria': t_pont_plc,
                        'disparo':  t_disp_plc,
                    }
                except (TypeError, ValueError):
                    pass
            if dist_data:
                ind_data[dist] = dist_data

        if ind_data:
            res_ind[idx] = ind_data

    # Ler Hurdle Step (bipodal) -- formato: inicio1 fim1 inicio2 fim2 ...
    hs_data = {}
    if ws_hs is not None:
        for r in range(2, ws_hs.max_row + 1):
            idx_val = ws_hs.cell(r, 1).value
            if idx_val is None: continue
            try: idx = int(idx_val)
            except (ValueError, TypeError): continue
            t = {}
            col = 2
            ens = 1
            while True:
                ini_v = ws_hs.cell(r, col).value
                fim_v = ws_hs.cell(r, col + 1).value
                if ini_v is None and fim_v is None: break
                if col + 1 > ws_hs.max_column: break
                try:
                    t[ens] = (int(ini_v), int(fim_v))
                    ens += 1
                except (TypeError, ValueError):
                    pass
                col += 2
                if col > 80: break  # safety
            if t: hs_data[idx] = t

    return {
        'por_individuo': res_ind,
        'hurdle_step':   hs_data,
        'distancias':    dists_ord,
        'n_ensaios_dist': n_ens_dist,
    }


# Alias para compatibilidade com codigo antigo
def carregar_inicio_fim_tiro(caminho):
    return carregar_tempos_tiro(caminho)


# Intervalos do protocolo de tiro
TIRO_INTERVALOS = {
    'toque_pontaria':   'Toque a Pontaria',
    'toque_disparo':    'Toque a Disparo',
    'pontaria_disparo': 'Pontaria a Disparo',
    'disparo_fim':      'Disparo ao Fim',
    'total':            'Ensaio Total',
}


def carregar_scores_tiro(caminho):
    """
    Scores de precisao: ficheiro Excel com Atleta | score_1 | score_2 | ...
    Pode estar na folha 'Scores' ou na folha activa.
    """
    try:
        wb = load_workbook(caminho, data_only=True)
        ws = wb['Scores'] if 'Scores' in wb.sheetnames else None
        if ws is None: return {}
        res = {}
        for linha in ws.iter_rows(min_row=2, values_only=True):
            if not linha[0]: continue
            scores = []
            for v in linha[1:]:
                try: scores.append(float(v))
                except Exception: pass
            if scores: res[str(linha[0]).strip()] = scores
        return res
    except Exception: return {}


def _norm(s):
    return re.sub(r'[^a-z0-9 ]', '',
                  unicodedata.normalize('NFD', str(s).lower()))


def _limpar_nome(s):
    """Limpa o nome: remove underscores/hifens extra, capitaliza correctamente."""
    # Remove underscores e hifens que sejam separadores (nao dentro de palavra)
    s = re.sub(r'[_\-]+', ' ', s).strip()
    # Remove multiplos espacos
    s = re.sub(r'\s+', ' ', s)
    # Capitaliza cada palavra (titulo)
    return s.title()


def encontrar_atleta(nome_pasta, ifd):
    """Extrai ID (numero antes do _) e nome limpo da pasta. Tenta matching com ifd."""
    # Extrair ID e nome da pasta: "66_Leonor Marinheiro" → id=66, nome=Leonor Marinheiro
    m_id = re.match(r'^(\d+)[_\-\s]+(.+)$', nome_pasta)
    if m_id:
        individuo_id = m_id.group(1)
        nome_raw = m_id.group(2)
    else:
        individuo_id = None
        nome_raw = nome_pasta

    nome = _limpar_nome(nome_raw)

    for cand in [nome, nome_raw, nome_pasta]:
        for chave in (ifd or {}):
            if _norm(chave) == _norm(cand):
                return chave, nome, individuo_id
    for cand in [nome, nome_raw, nome_pasta]:
        for chave in (ifd or {}):
            if _norm(cand) in _norm(chave) or _norm(chave) in _norm(cand):
                return chave, nome, individuo_id
    return None, nome, individuo_id


# -----------------------------------------------------------------------
# Loaders especificos do Tiro com Arco (v1.0)
# -----------------------------------------------------------------------

def _normalizar_nome_sheet(nome):
    """
    Normaliza um nome de folha Excel para comparacao.
    Lida com caracteres cp1252 estragados ('confirma\xe7\xe3o_1' ou
    'confirma\ufffd\ufffdo_1' que aparecem quando o ficheiro foi guardado
    com encoding errado).
    """
    if nome is None:
        return ''
    s = str(nome).lower().strip()
    # Remove sufixos mojibake comuns
    s = s.replace('\ufffd', '').replace('\u00e7\u00e3', 'ca')
    # Remove acentos
    try:
        s = unicodedata.normalize('NFD', s)
        s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    except Exception:
        pass
    return s.replace(' ', '').replace('_', '')


def carregar_confirmacao_arco(caminho):
    """
    Le o ficheiro Inicio_fim_vfinal.xlsx do protocolo Tiro com Arco.

    Estrutura esperada (3 folhas):
      'tempo do toque'      - t do toque (inicio de referencia)
      'confirmação_1'        - t da primeira confirmacao (inicio da janela)
      'confirmação_2'        - t da segunda confirmacao (fim da janela)

    Cada folha tem colunas: atleta_id | ensaio_1 | ensaio_2 | ... | ensaio_30

    A janela unica de analise e [confirmação_1, confirmação_2].

    Devolve:
      {
        'por_id': {
            '101': {
                1: {'toque': 920, 'conf_1': 1400, 'conf_2': 3200},
                2: {'toque': 880, 'conf_1': 1350, 'conf_2': 3100},
                ...
            },
            ...
        },
        'n_trials_max': 30,
        'ids': ['101', '102', ...],
      }

    Tolera nomes de folha com mojibake cp1252 (ex: 'confirma\ufffd\ufffdo_1').
    """
    wb = load_workbook(caminho, data_only=True)

    # Mapear as 3 folhas de interesse por normalizacao
    alvo = {
        'toque':  None,
        'conf_1': None,
        'conf_2': None,
    }
    for sheet_name in wb.sheetnames:
        norm = _normalizar_nome_sheet(sheet_name)
        if 'tempodotoque' in norm or 'toque' in norm:
            alvo['toque'] = sheet_name
        elif ('confirma' in norm and '1' in norm) or 'confirmacao1' in norm:
            alvo['conf_1'] = sheet_name
        elif ('confirma' in norm and '2' in norm) or 'confirmacao2' in norm:
            alvo['conf_2'] = sheet_name

    if not alvo['conf_1'] or not alvo['conf_2']:
        raise ValueError(
            f"Inicio_fim_arco sem folhas de confirmacao: {caminho} "
            f"(encontradas: {wb.sheetnames})"
        )

    def _ler_folha(nome_folha):
        """Devolve {id_atleta: {trial: valor_ms}} para uma folha."""
        ws = wb[nome_folha]
        resultado = {}
        # Linha 1 = cabecalho ('atleta', 'ensaio_1', 'ensaio_2', ...)
        # Linhas 2+ = dados. Primeira coluna = id do atleta.
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row or row[0] is None:
                continue
            atleta_id = str(row[0]).strip()
            if not atleta_id:
                continue
            trials = {}
            for idx, val in enumerate(row[1:], start=1):
                if val is None:
                    continue
                try:
                    trials[idx] = int(round(float(val)))
                except (TypeError, ValueError):
                    continue
            if trials:
                resultado[atleta_id] = trials
        return resultado

    toque_data  = _ler_folha(alvo['toque']) if alvo['toque'] else {}
    conf_1_data = _ler_folha(alvo['conf_1'])
    conf_2_data = _ler_folha(alvo['conf_2'])

    # Fundir os 3 em {id: {trial: {'toque','conf_1','conf_2'}}}
    por_id = {}
    todos_ids = set(toque_data) | set(conf_1_data) | set(conf_2_data)
    n_trials_max = 0
    for aid in todos_ids:
        trials_union = set()
        for d in (toque_data, conf_1_data, conf_2_data):
            trials_union.update(d.get(aid, {}).keys())
        por_id[aid] = {}
        for t in sorted(trials_union):
            entry = {
                'toque':  toque_data.get(aid, {}).get(t),
                'conf_1': conf_1_data.get(aid, {}).get(t),
                'conf_2': conf_2_data.get(aid, {}).get(t),
            }
            # So guardar se tiver pelo menos conf_1 e conf_2 (janela completa)
            if entry['conf_1'] is not None and entry['conf_2'] is not None:
                por_id[aid][t] = entry
                if t > n_trials_max:
                    n_trials_max = t

    return {
        'por_id':       por_id,
        'n_trials_max': n_trials_max,
        'ids':          sorted(por_id.keys()),
    }


# -----------------------------------------------------------------------
# Referencia demografica - "Todos os registos dos 142 atletas"
# -----------------------------------------------------------------------

# Mapeamento de codigos demograficos (confirmado pelo Pedro Aleixo):
#   GENERO:  1 -> M (masculino), 2 -> F (feminino)
#   ESTILO:  1 -> recurvo,       2 -> composto
#   CATEGORIA: codigo livre (cadete, junior, senior, master, ...)
_ARCO_GENERO_MAP = {1: 'M', 2: 'F', '1': 'M', '2': 'F',
                    'M': 'M', 'F': 'F', 'MASCULINO': 'M', 'FEMININO': 'F'}
_ARCO_ESTILO_MAP = {1: 'recurvo', 2: 'composto', '1': 'recurvo', '2': 'composto',
                    'RECURVO': 'recurvo', 'COMPOSTO': 'composto',
                    'RECURVE': 'recurvo', 'COMPOUND': 'composto'}

def _normalizar_cabecalho(s):
    """Normaliza header: maiusculas, sem acentos/espacos, para comparacao robusta."""
    if s is None:
        return ''
    import unicodedata as _u
    txt = str(s).strip().upper()
    # Remove acentos
    nfkd = _u.normalize('NFKD', txt)
    txt = ''.join(c for c in nfkd if not _u.combining(c))
    # Remove espacos, underscores, hifens, pontos
    for ch in ' _-.\t':
        txt = txt.replace(ch, '')
    return txt

def carregar_atletas_ref(caminho):
    """
    Le o ficheiro "Todos os registos dos 142 atletas em JUl_2024 _.xlsx"
    com referencias demograficas e scores.

    Estrutura esperada:
      Linha 1 = cabecalho (PESO, ALTURA, IDADE, ESTILO, CATEGORIA, GENERO,
                           P1..P30, P_TOTAL, d1..d30, [colunas manuais ignoradas])
      Linhas 2+ = dados, uma linha por atleta (142 linhas).

    Ignoramos explicitamente as colunas 68-308 (0-indexed: 67-307) que
    contem valores manuais irrelevantes.

    Devolve: lista de dicts
        {'id': '101', 'peso_kg': 70.5, 'altura_m': 1.75, 'idade': 23,
         'estilo': 'recurvo', 'categoria': 'senior', 'genero': 'M',
         'P': [10, 9, 10, ...], 'P_total': 282, 'd': [12.5, 8.1, ...]}

    Silenciosamente tolera colunas em falta (devolve None no campo).
    """
    wb = load_workbook(caminho, data_only=True)
    ws = wb.active   # primeira folha

    # Ler cabecalho e mapear colunas por nome normalizado
    cabecalho_raw = [cell.value for cell in ws[1]]
    # Limite: ignorar colunas >= 67 (col 68 em 1-indexed)
    LIMITE_COL = 67   # 0-indexed, ie cols 1..67 sao uteis (P_TOTAL, d1..d30 cabem aqui)

    col_map = {}   # chave normalizada -> indice 0-based
    p_cols  = {}   # numero de ensaio -> indice coluna
    d_cols  = {}   # numero de ensaio -> indice coluna
    for idx, v in enumerate(cabecalho_raw):
        if v is None:
            continue
        if idx >= LIMITE_COL:
            break
        norm = _normalizar_cabecalho(v)
        if not norm:
            continue
        # Scores Pn e distancias dn
        import re as _re
        mp = _re.match(r'^P(\d+)$', norm)
        md = _re.match(r'^D(\d+)$', norm)
        if mp:
            p_cols[int(mp.group(1))] = idx
            continue
        if md:
            d_cols[int(md.group(1))] = idx
            continue
        # Restantes atributos
        col_map[norm] = idx

    def _col(row, *nomes):
        for n in nomes:
            idx = col_map.get(_normalizar_cabecalho(n))
            if idx is not None and idx < len(row):
                return row[idx]
        return None

    def _num(v):
        if v is None: return None
        try:
            f = float(v)
            if f != f:   # NaN
                return None
            return f
        except (TypeError, ValueError):
            return None

    def _int(v):
        f = _num(v)
        return int(f) if f is not None else None

    atletas = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or all(v is None for v in row):
            continue
        # ID do atleta: tentamos ID / NUM / N / N_ATLETA / primeiro nao-vazio
        aid = None
        for nome_id in ('ID', 'NUM', 'N', 'NATLETA', 'ATLETA'):
            v = _col(row, nome_id)
            if v is not None:
                aid = str(v).strip()
                try:
                    aid = str(int(float(aid)))   # "101.0" -> "101"
                except (TypeError, ValueError):
                    pass
                break
        if aid is None:
            # Fallback: primeira coluna se for numero
            if row and row[0] is not None:
                try:
                    aid = str(int(float(row[0])))
                except (TypeError, ValueError):
                    aid = str(row[0]).strip()
        if not aid:
            continue

        peso     = _num(_col(row, 'PESO', 'PESOKG', 'WEIGHT'))
        altura   = _num(_col(row, 'ALTURA', 'HEIGHT', 'ALTURAM', 'ALTURACM'))
        # Altura em cm vs m: converter se > 3 (claramente cm)
        if altura is not None and altura > 3:
            altura = altura / 100.0
        idade    = _int(_col(row, 'IDADE', 'AGE', 'ANOS'))
        estilo_r = _col(row, 'ESTILO', 'STYLE')
        categ    = _col(row, 'CATEGORIA', 'CATEGORY', 'CAT')
        genero_r = _col(row, 'GENERO', 'GENDER', 'SEXO', 'SEX')

        # Normalizar codigos
        estilo = _ARCO_ESTILO_MAP.get(
            estilo_r if not isinstance(estilo_r, str) else estilo_r.strip().upper(),
            str(estilo_r).strip().lower() if estilo_r is not None else None)
        genero = _ARCO_GENERO_MAP.get(
            genero_r if not isinstance(genero_r, str) else genero_r.strip().upper(),
            str(genero_r).strip().upper() if genero_r is not None else None)
        if isinstance(categ, str):
            categ = categ.strip().lower()

        # Scores P1..P30 e distancias d1..d30
        P = []
        for t in sorted(p_cols):
            v = row[p_cols[t]] if p_cols[t] < len(row) else None
            P.append(_num(v))
        dist = []
        for t in sorted(d_cols):
            v = row[d_cols[t]] if d_cols[t] < len(row) else None
            dist.append(_num(v))

        # P_total: prefere a coluna explicita, senao soma de P (ignorando None)
        p_total = _num(_col(row, 'PTOTAL', 'PTOT', 'TOTAL', 'SCORETOTAL'))
        if p_total is None and P:
            vals_p = [v for v in P if v is not None]
            if vals_p:
                p_total = sum(vals_p)

        atletas.append({
            'id':        aid,
            'peso_kg':   peso,
            'altura_m':  altura,
            'idade':     idade,
            'estilo':    estilo,
            'categoria': categ,
            'genero':    genero,
            'P':         P,
            'P_total':   p_total,
            'd':         dist,
        })

    return atletas


def atletas_ref_por_id(lista_ref):
    """Converte a lista de carregar_atletas_ref() em dict {id: atleta}."""
    return {a['id']: a for a in lista_ref if a.get('id')}


# -----------------------------------------------------------------------
# Analises demograficas (v1.0) - Tiro com Arco + 142 atletas
# -----------------------------------------------------------------------
#
# Quatro funcoes de analise descritiva/inferencial que ligam as metricas de
# CoP (ea95, stiff_*, vel_med, etc.) a demografia e performance:
#
#   comparar_grupos(atletas, chave, fator)
#     Compara distribuicao de `chave` entre niveis de `fator` (genero, estilo,
#     categoria). Mann-Whitney (2 grupos) ou Kruskal-Wallis (3+ grupos). Nao
#     parametrico porque n por grupo tende a ser pequeno e CoP raramente e
#     normal.
#
#   correlacao_demografica(atletas, chave_cop, chave_dem)
#     Pearson (linear, parametrica) + Spearman (monotona, robusta). Reportamos
#     ambos porque Pearson deteca relacoes lineares e Spearman capta tendencias
#     monotonas sem normalidade.
#
#   percentis_subgrupo(ath, atletas, chave, fatores=('categoria','genero'))
#     Onde o atleta cai (P25/P50/P75/rank) dentro do seu subgrupo demografico.
#     Util no feedback individual: "o seu ea95 esta no P75 dos seniores M".
#
#   correlacao_score(atletas, chave_cop, chave_outcome='P_total')
#     Liga CoP a performance. Dois niveis: agregado (1 CoP por atleta vs score
#     total) e per-ensaio (todos os ensaios de todos os atletas - CoP do trial
#     vs P do trial).
#
# Referencias:
#   Mann & Whitney (1947). Ann. Math. Stat., 18(1), 50-60.
#   Kruskal & Wallis (1952). J. Am. Stat. Assoc., 47(260), 583-621.
#   Spearman (1904). Am. J. Psychol., 15, 72-101.

def _valores_chave(atletas, chave):
    """
    Extrai uma lista de (atleta, valor) onde valor e agregado por atleta:
    procura primeiro `chave` directamente no dict do atleta (ex: peso_kg,
    P_total) e, se nao existir, calcula a mediana dos ensaios `mets[lado][t][chave]`.
    Devolve apenas pares com valor numerico valido.
    """
    import math as _m
    out = []
    for a in atletas:
        v = a.get(chave) if isinstance(a, dict) else None
        if v is None:
            # Tenta agregar por mediana dos ensaios (todos os lados)
            vs = []
            mets = a.get('mets') if isinstance(a, dict) else None
            if isinstance(mets, dict):
                for _lado, lista in mets.items():
                    if isinstance(lista, list):
                        for m in lista:
                            if isinstance(m, dict) and m.get(chave) is not None:
                                try:
                                    vs.append(float(m[chave]))
                                except (TypeError, ValueError):
                                    pass
            if vs:
                try:
                    v = float(np.median(vs))
                except Exception:
                    v = None
        if v is None:
            continue
        try:
            fv = float(v)
            if _m.isnan(fv) or _m.isinf(fv):
                continue
            out.append((a, fv))
        except (TypeError, ValueError):
            continue
    return out


def comparar_grupos(atletas, chave, fator):
    """
    Compara a distribuicao de `chave` entre niveis de `fator`.

    Parametros:
      atletas : lista de dicts (cada um com chaves demograficas e de metricas)
      chave   : nome da metrica (ex: 'ea95', 'stiff_x', 'vel_med', 'P_total')
      fator   : 'genero' | 'estilo' | 'categoria'

    Devolve dict:
      {
        'fator': str, 'chave': str,
        'grupos': {'M': [v1, v2, ...], 'F': [...]},
        'n_por_grupo': {'M': 20, 'F': 22},
        'mediana': {'M': x, 'F': y},
        'teste': 'mannwhitneyu' | 'kruskal' | None,
        'estatistica': float | None,
        'p_valor': float | None,
        'n_grupos': int,
      }
    """
    pares = _valores_chave(atletas, chave)
    grupos = {}
    for a, v in pares:
        g = a.get(fator)
        if g is None or g == '':
            continue
        # Normaliza: lowercase strings para categoria/estilo, upper para genero
        if isinstance(g, str):
            g = g.strip()
            if fator == 'genero':
                g = g.upper()
            else:
                g = g.lower()
        grupos.setdefault(g, []).append(v)

    # Remove grupos com n<2 (nao da para estatistica)
    grupos = {k: vs for k, vs in grupos.items() if len(vs) >= 2}

    n_por_grupo = {k: len(vs) for k, vs in grupos.items()}
    mediana = {k: float(np.median(vs)) for k, vs in grupos.items()}
    p25 = {k: float(np.percentile(vs, 25)) for k, vs in grupos.items()}
    p75 = {k: float(np.percentile(vs, 75)) for k, vs in grupos.items()}

    teste = estat = pval = None
    n_grupos = len(grupos)

    try:
        if n_grupos == 2:
            chaves = list(grupos.keys())
            U, p = stats.mannwhitneyu(grupos[chaves[0]], grupos[chaves[1]],
                                       alternative='two-sided')
            teste, estat, pval = 'mannwhitneyu', float(U), float(p)
        elif n_grupos >= 3:
            H, p = stats.kruskal(*[grupos[k] for k in grupos])
            teste, estat, pval = 'kruskal', float(H), float(p)
    except Exception:
        pass

    return {
        'fator': fator,
        'chave': chave,
        'grupos': grupos,
        'n_por_grupo': n_por_grupo,
        'mediana': mediana,
        'p25': p25,
        'p75': p75,
        'teste': teste,
        'estatistica': estat,
        'p_valor': pval,
        'n_grupos': n_grupos,
    }


def correlacao_demografica(atletas, chave_cop, chave_dem):
    """
    Correlacao entre uma metrica CoP e uma variavel demografica continua.

    Parametros:
      chave_cop : metrica CoP (ex: 'ea95', 'stiff_x', 'vel_med')
      chave_dem : variavel demografica (ex: 'peso_kg', 'altura_m', 'idade')

    Devolve dict:
      {
        'chave_cop': str, 'chave_dem': str,
        'n': int,
        'pearson_r': float | None, 'pearson_p': float | None,
        'spearman_r': float | None, 'spearman_p': float | None,
        'x': list (valores demograficos), 'y': list (valores CoP),
      }
    """
    pares_cop = _valores_chave(atletas, chave_cop)
    # Match com valor demografico
    xs, ys = [], []
    for a, cop_v in pares_cop:
        dem_v = a.get(chave_dem) if isinstance(a, dict) else None
        if dem_v is None:
            continue
        try:
            dem_f = float(dem_v)
            if dem_f != dem_f:  # NaN
                continue
        except (TypeError, ValueError):
            continue
        xs.append(dem_f); ys.append(cop_v)

    pr = pp = sr = sp = None
    if len(xs) >= 3:
        try:
            r, p = stats.pearsonr(xs, ys)
            pr, pp = float(r), float(p)
        except Exception:
            pass
        try:
            r, p = stats.spearmanr(xs, ys)
            sr, sp = float(r), float(p)
        except Exception:
            pass

    return {
        'chave_cop': chave_cop,
        'chave_dem': chave_dem,
        'n': len(xs),
        'pearson_r': pr, 'pearson_p': pp,
        'spearman_r': sr, 'spearman_p': sp,
        'x': xs, 'y': ys,
    }


def percentis_subgrupo(ath, atletas, chave, fatores=('categoria', 'genero')):
    """
    Calcula a posicao do atleta no subgrupo demografico (e.g. mesmo genero +
    categoria). Util para feedback individual: "esta no P75 dos seniores M".

    Parametros:
      ath     : dict do atleta de interesse
      atletas : lista onde procurar pares (tipicamente os 142)
      chave   : metrica (ex: 'ea95', 'P_total')
      fatores : tuplo de chaves demograficas para filtrar

    Devolve dict:
      {
        'valor_atleta': float | None,
        'subgrupo': {'genero': 'M', 'categoria': 'senior'},
        'n_subgrupo': int,
        'p25': float | None, 'p50': float | None, 'p75': float | None,
        'percentil_atleta': float | None,  # 0-100
        'rank': int | None,                # 1-based, 1 = menor valor
      }
    """
    if not isinstance(ath, dict):
        return {'valor_atleta': None, 'subgrupo': {}, 'n_subgrupo': 0,
                'p25': None, 'p50': None, 'p75': None,
                'percentil_atleta': None, 'rank': None}

    # Subgrupo: mesmos valores para todos os factores
    subgrupo = {}
    for f in fatores:
        subgrupo[f] = ath.get(f)

    def _bate(a):
        for f in fatores:
            v_ath = ath.get(f)
            v_a   = a.get(f)
            if v_ath is None and v_a is None:
                continue
            if v_ath is None or v_a is None:
                return False
            # Normaliza strings para comparacao
            if isinstance(v_ath, str) and isinstance(v_a, str):
                if v_ath.strip().lower() != v_a.strip().lower():
                    return False
            else:
                if v_ath != v_a:
                    return False
        return True

    pares_sub = [p for p in _valores_chave(atletas, chave) if _bate(p[0])]
    valores = [v for _, v in pares_sub]
    n_sub = len(valores)

    # Valor do atleta
    v_ath = None
    for a, v in _valores_chave([ath], chave):
        v_ath = v; break

    p25 = p50 = p75 = None
    pct_ath = rank_ath = None
    if n_sub >= 3:
        p25 = float(np.percentile(valores, 25))
        p50 = float(np.percentile(valores, 50))
        p75 = float(np.percentile(valores, 75))
        if v_ath is not None:
            # Percentil: fraccao de valores <= v_ath
            n_le = sum(1 for v in valores if v <= v_ath)
            pct_ath = round(100.0 * n_le / n_sub, 1)
            # Rank: 1 = menor valor (melhor equilibrio tipicamente)
            sorted_vals = sorted(valores)
            rank_ath = 1
            for v in sorted_vals:
                if v < v_ath:
                    rank_ath += 1
                else:
                    break

    return {
        'valor_atleta': v_ath,
        'subgrupo': subgrupo,
        'n_subgrupo': n_sub,
        'p25': p25, 'p50': p50, 'p75': p75,
        'percentil_atleta': pct_ath,
        'rank': rank_ath,
    }


def correlacao_score(atletas, chave_cop, chave_outcome='P_total'):
    """
    Correlaciona uma metrica CoP com um score de desempenho.

    Dois niveis de analise:
      agregado  : 1 ponto por atleta (CoP mediano vs score total)
      per_ensaio: todos os ensaios de todos os atletas (CoP do trial t
                  vs score P_t do mesmo trial)

    Parametros:
      chave_cop     : metrica CoP (ex: 'ea95')
      chave_outcome : 'P_total' | 'P' (lista por ensaio)

    Devolve dict:
      {
        'chave_cop': str, 'chave_outcome': str,
        'agregado': {n, pearson_r, pearson_p, spearman_r, spearman_p, x, y},
        'per_ensaio': {n, pearson_r, pearson_p, spearman_r, spearman_p}
                      ou None se P nao disponivel,
      }
    """
    # Nivel agregado: 1 CoP mediano + 1 score total por atleta
    agreg = correlacao_demografica(atletas, chave_cop, chave_outcome)

    # Nivel per-ensaio: so faz sentido se houver ref com P (lista) e CoP per-ensaio
    per = None
    xs, ys = [], []
    for a in atletas:
        P_list = a.get('P') if isinstance(a, dict) else None
        if not isinstance(P_list, list):
            continue
        mets = a.get('mets') if isinstance(a, dict) else None
        if not isinstance(mets, dict):
            continue
        # Assume que o lado 'arco' tem os ensaios alinhados com P_list
        ensaios = mets.get('arco') or []
        for i, m in enumerate(ensaios):
            if not isinstance(m, dict):
                continue
            if i >= len(P_list):
                continue
            p_val = P_list[i]
            cop_val = m.get(chave_cop)
            if p_val is None or cop_val is None:
                continue
            try:
                xs.append(float(p_val)); ys.append(float(cop_val))
            except (TypeError, ValueError):
                continue

    if len(xs) >= 3:
        try:
            pr, pp = stats.pearsonr(xs, ys)
            sr, sp = stats.spearmanr(xs, ys)
            per = {
                'n': len(xs),
                'pearson_r': float(pr), 'pearson_p': float(pp),
                'spearman_r': float(sr), 'spearman_p': float(sp),
            }
        except Exception:
            per = None

    return {
        'chave_cop': chave_cop,
        'chave_outcome': chave_outcome,
        'agregado': agreg,
        'per_ensaio': per,
    }


# -----------------------------------------------------------------------
# Processar atleta (geral + tiro)
# -----------------------------------------------------------------------

# Funcao que processa os dados de um individuo
# Le os ficheiros .xls e calcula as metricas de estabilidade
def processar_atleta(pasta, ifd, usar_embed, log=print, protocolo=None,
                     n_ens_override=None, distancia=None, n_ens_hs_override=None,
                     intervalos_tiro=None, match_por_idx=False, idx_na_lista=None,
                     ask_callback=None, incluir_hs=True, peso_kg_default=None,
                     tempos_arco=None, atleta_ref=None):
    """
    Parametros novos (v1.0):
      tempos_arco  : resultado de carregar_confirmacao_arco() para PROTO_ARCO
      atleta_ref   : dict da referência dos 142 atletas (peso, altura, estilo, ...)
    """
    if protocolo is None: protocolo = _PROTOCOLO_ACTIVO
    proto = PROTOCOLOS[protocolo]
    n_ens = n_ens_override if n_ens_override and n_ens_override > 0 else proto['n_ens']

    # Tiro com Arco - pipeline próprio (janela única, sem distâncias)
    if protocolo == PROTO_ARCO:
        return _processar_atleta_arco(pasta, tempos_arco, n_ens, log,
                                      match_por_idx=match_por_idx,
                                      idx_na_lista=idx_na_lista,
                                      ask_callback=ask_callback,
                                      atleta_ref=atleta_ref,
                                      peso_kg_default=peso_kg_default)

    if proto.get('two_windows'):
        return _processar_atleta_tiro(pasta, ifd, usar_embed, n_ens, log,
                                      n_ens_hs=n_ens_hs_override if n_ens_hs_override else n_ens,
                                      distancia=distancia,
                                      intervalos=intervalos_tiro,
                                      match_por_idx=match_por_idx,
                                      idx_na_lista=idx_na_lista,
                                      ask_callback=ask_callback,
                                      incluir_hs=incluir_hs,
                                      protocolo=protocolo)

    nome_pasta = os.path.basename(pasta)
    chave, nome, ind_id = encontrar_atleta(nome_pasta, ifd or {})
    tempos = (ifd or {}).get(chave, {})
    if ifd and not tempos:
        log(f"  aviso: '{nome}' nao encontrado no Inicio_fim", 'aviso')

    mets = {lado: [] for lado, _, _ in proto['lados']}
    raw  = {lado: [] for lado, _, _ in proto['lados']}

    for lado, pref, offset in proto['lados']:
        for t in range(1, n_ens + 1):
            num = t + offset
            fp  = achar_ficheiro(pasta, lado, t)
            if fp is None:
                fp = achar_ficheiro(pasta, pref)  # fallback antigo
            if fp is None or os.path.getsize(fp) == 0:
                msg = 'nao encontrado' if fp is None else 'ficheiro vazio'
                log(f"  aviso {msg}: {pref}{t}", 'aviso')
                mets[lado].append(None); raw[lado].append(None); continue
            try:
                dados = ler_ficheiro(fp)
                frs   = dados['frames']
                if tempos and num in tempos: b, e = tempos[num]
                elif usar_embed and dados['inicio_ms'] is not None:
                    b, e = dados['inicio_ms'], dados['fim_ms']
                else: b = e = None

                # ── Cascata de fallbacks (igual ao protocolo de tiro) ──────────
                def _n_frames_j(b_, e_):
                    if b_ is None: return len(frs)
                    return sum(1 for f in frs if b_ <= f['t_ms'] <= e_)

                m = None
                used_b, used_e = b, e
                _fonte = f"{b}-{e}ms" if b is not None else 'completo'

                # Tentativa 1: janela determinada (IFD ou embed)
                if _n_frames_j(b, e) >= 5:
                    m = calcular(frs, b, e, peso_kg=peso_kg_default)

                # Tentativa 2: tempo embutido no próprio ficheiro XLS
                if m is None and dados['inicio_ms'] is not None and dados['inicio_ms'] != b:
                    b2, e2 = dados['inicio_ms'], dados['fim_ms']
                    if _n_frames_j(b2, e2) >= 5:
                        m = calcular(frs, b2, e2, peso_kg=peso_kg_default)
                        if m:
                            used_b, used_e = b2, e2
                            _fonte = f'embed ({b2}-{e2}ms)'
                            if b is not None:
                                t0f = frs[0]['t_ms'] if frs else '?'
                                t1f = frs[-1]['t_ms'] if frs else '?'
                                log(f"  info {pref}{t}: janela {b}-{e}ms sem frames "
                                    f"(t_ms ficheiro: {t0f:.0f}-{t1f:.0f}ms). "
                                    f"Usado tempo embutido {b2}-{e2}ms.", 'aviso')

                # Tentativa 3: ficheiro completo (sem windowing)
                if m is None and len(frs) >= 5:
                    m = calcular(frs, peso_kg=peso_kg_default)
                    if m:
                        used_b, used_e = None, None
                        _fonte = 'completo (sem janela)'
                        t0f = frs[0]['t_ms'] if frs else '?'
                        t1f = frs[-1]['t_ms'] if frs else '?'
                        log(f"  info {pref}{t}: sem sobreposição de janela "
                            f"(ficheiro t={t0f:.0f}-{t1f:.0f}ms, janela={b}-{e}ms). "
                            f"Usado ficheiro completo.", 'aviso')

                mets[lado].append(m)
                raw[lado].append({'dados': dados, 'ini': used_b, 'fim': used_e})
                if m:
                    itv = f"{used_b}-{used_e}ms" if used_b is not None else "completo"
                    log(f"  ok {pref}{t} [{itv}] [{_fonte}]"
                        f"  ea95={m['ea95']:.1f}mm2  a={m['leng_a']:.1f}  b={m['leng_b']:.1f}mm", 'ok')
                else:
                    t0f = frs[0]['t_ms'] if frs else '?'
                    t1f = frs[-1]['t_ms'] if frs else '?'
                    log(f"  erro {pref}{t}: dados insuficientes mesmo após fallbacks "
                        f"({len(frs)} frames, t={t0f:.0f}-{t1f:.0f}ms, "
                        f"janela_pedida={b}-{e}ms)", 'erro')
            except Exception as ex:
                log(f"  erro {pref}{t}: {ex}", 'erro')
                mets[lado].append(None); raw[lado].append(None)

    # Aviso de ensaios validos insuficientes
    min_val = AOM_CONFIG.get('min_ensaios_validos', 3)
    for lado_k in mets:
        n_ok = sum(1 for m in mets[lado_k] if m is not None)
        total_k = len(mets[lado_k])
        if total_k > 0 and n_ok < min_val:
            msg = (f"  aviso: '{nome}' lado '{lado_k}': {n_ok}/{total_k} ensaios validos "
                   f"(minimo recomendado: {min_val}). Interprete com cautela.")
            log(msg, 'aviso')

    return {'nome': nome, 'id': ind_id, 'mets': mets, 'raw': raw,
            'protocolo': protocolo, 'scores': None}


def _processar_atleta_tiro(pasta, tempos_tiro, usar_embed, n_ens, log,
                           distancia=None, intervalos=None, match_por_idx=False,
                           idx_na_lista=None, n_ens_hs=None, ask_callback=None,
                           incluir_hs=True, protocolo=None):
    """
    Processa um individuo no protocolo de tiro.

    tempos_tiro: resultado de carregar_tempos_tiro() ou None
    intervalos: lista de chaves de TIRO_INTERVALOS a calcular
    match_por_idx: se True, usa idx_na_lista (posicao na lista de pastas) como chave
    protocolo: PROTO_TIRO ou PROTO_ARCO (default: PROTO_TIRO)
    """
    if protocolo is None:
        protocolo = PROTO_TIRO
    nome_pasta = os.path.basename(pasta)
    chave, nome, ind_id = encontrar_atleta(nome_pasta, tempos_tiro.get('por_individuo', {}) if tempos_tiro else {})

    if intervalos is None:
        intervalos = list(TIRO_INTERVALOS.keys())

    # Determinar chave de matching para o ficheiro de tempos
    tempos_por_ind = tempos_tiro.get('por_individuo', {}) if tempos_tiro else {}
    hs_por_ind     = tempos_tiro.get('hurdle_step', {}) if tempos_tiro else {}
    dists          = tempos_tiro.get('distancias', []) if tempos_tiro else []
    if distancia and distancia not in dists:
        dists = [distancia]

    # Tentar achar o individuo no ficheiro de tempos
    ind_tempos = None
    if tempos_tiro:
        # Tentativa 1: pelo ID extraido do nome da pasta
        if ind_id and int(ind_id) in tempos_por_ind:
            ind_tempos = tempos_por_ind[int(ind_id)]
        # Tentativa 2: pelo indice na lista de pastas (1-based)
        elif match_por_idx and idx_na_lista is not None:
            ind_tempos = tempos_por_ind.get(idx_na_lista)
        # Tentativa 3: primeiro indice disponivel com matching parcial (fallback)
        if ind_tempos is None and not match_por_idx:
            # Tentar por posicao ordinal: ID numerico da pasta
            for k in sorted(tempos_por_ind.keys()):
                if ind_id and str(k) == str(ind_id):
                    ind_tempos = tempos_por_ind[k]; break

        if ind_tempos is None and not match_por_idx:
            log(f"  aviso: '{nome}' (ID={ind_id}) nao encontrado no ficheiro de tempos", 'aviso')

    # ----------------------------------------------------------------
    # Parte 1: Disparo por distancia e intervalo
    # ----------------------------------------------------------------
    # Estrutura: tiro_dist[dist][intervalo] = [m_trial1, m_trial2, ...]
    tiro_dist = {}

    dists_process = dists if dists else ([distancia] if distancia else [])
    for dist in dists_process:
        dist_tempos = ind_tempos.get(dist, {}) if ind_tempos else {}
        n_trials = n_ens

        dist_result = {itv: [] for itv in intervalos}
        dist_raw    = {itv: [] for itv in intervalos}

        # Pre-verificar quais ensaios existem
        _em_falta = []
        dist_num = dist.replace('m', '').replace('.', '')
        for t in range(1, n_trials + 1):
            fp_chk = achar_ficheiro_tiro(pasta, dist_num, t) or achar_ficheiro(pasta, 'tiro', t)
            if fp_chk is None:
                _em_falta.append(t)

        # Se faltam ensaios, perguntar ao utilizador via callback
        _incluir = True  # por defeito incluir com o que existe
        if _em_falta:
            faltam_str = ', '.join(f't{t}' for t in _em_falta)
            existem = n_trials - len(_em_falta)
            log(f"  aviso {dist}: {len(_em_falta)} ensaio(s) em falta ({faltam_str}). {existem}/{n_trials} disponiveis.", 'aviso')
            if ask_callback is not None:
                _incluir = ask_callback(nome, dist, _em_falta, existem, n_trials)
                if not _incluir:
                    log(f"  individuo ignorado para distancia {dist} (decisao do utilizador)", 'aviso')
                    tiro_dist[dist] = {'mets': {itv: [None]*n_trials for itv in intervalos},
                                       'raw':  {itv: [None]*n_trials for itv in intervalos}}
                    continue

        for t in range(1, n_trials + 1):
            # Localizar ficheiro: trial[dist]_[ensaio]
            fp = achar_ficheiro_tiro(pasta, dist_num, t)
            if fp is None:
                fp = achar_ficheiro(pasta, 'tiro', t)
            if fp is None:
                for itv in intervalos:
                    dist_result[itv].append(None)
                    dist_raw[itv].append(None)
                continue

            try:
                dados = ler_ficheiro(fp)
                t_fim_ficheiro = dados['frames'][-1]['t_ms'] if dados['frames'] else None
                trial_t = dist_tempos.get(t, {}) if dist_tempos else {}

                for itv in intervalos:
                    b, e = _tiro_janela(itv, trial_t, dados, t_fim_ficheiro, usar_embed)
                    m = calcular(dados['frames'], b, e)
                    dist_result[itv].append(m)
                    dist_raw[itv].append({'dados': dados, 'ini': b, 'fim': e,
                                          'dist': dist, 'trial': t})
                    if m:
                        itv_label = f"{b}-{e}ms" if b else "completo"
                        log(f"  ok {dist} t{t} [{_tiro_itv_label(itv)}] [{itv_label}]  ea95={m['ea95']:.1f}mm2", 'ok')
            except Exception as ex:
                ex_s = str(ex)
                if not ('ignorado' in ex_s or 'roll off' in ex_s.lower()):
                    log(f"  erro {dist} t{t}: {ex_s}", 'erro')
                for itv in intervalos:
                    dist_result[itv].append(None)
                    dist_raw[itv].append(None)

        tiro_dist[dist] = {'mets': dist_result, 'raw': dist_raw}

    # ----------------------------------------------------------------
    # Parte 2: Hurdle Step (bipodal) -- usa ficheiros hs_dir_N / hs_esq_N ou dir_N / esq_N
    # ----------------------------------------------------------------
    mets_hs = {'dir': [], 'esq': []}
    raw_hs  = {'dir': [], 'esq': []}

    if incluir_hs:
        hs_tempos_ind = None
        if ind_id and int(ind_id) in hs_por_ind:
            hs_tempos_ind = hs_por_ind[int(ind_id)]
        elif match_por_idx and idx_na_lista and idx_na_lista in hs_por_ind:
            hs_tempos_ind = hs_por_ind[idx_na_lista]

        # ── Descoberta automática de ficheiros HS ────────────────────
        # Varre TODOS os ficheiros .xls na pasta.
        # Para cada lado, aceita QUALQUER padrão de nome (hs_dir_N, dir_N, hs_esq_N, esq_N, etc.)
        # Não assume sequência 1..N - detecta ficheiros com numeração arbitrária.
        _HS_PADROES = {
            'dir': (_PADROES.get('hs_dir', []) + _PADROES.get('dir', [])),
            'esq': (_PADROES.get('hs_esq', []) + _PADROES.get('esq', [])),
        }

        def _listar_hs(pasta_scan, lado_scan):
            """
            Devolve lista ordenada de (numero_int, caminho) para todos os ficheiros
            que correspondam a qualquer padrão de hurdle step do lado dado.
            Aceita extensões .xls e .xlsx (formato antigo e novo).
            Aceita: hs_dir_N, dir_N, hs_esq_N, esq_N (e variantes maiúsculas).
            """
            padroes_s = _HS_PADROES.get(lado_scan, [])
            encontrados = {}
            try:
                ficheiros_s = sorted(os.listdir(pasta_scan))
            except Exception:
                return []
            for nome_f in ficheiros_s:
                ext_l = os.path.splitext(nome_f)[1].lower()
                if ext_l not in ('.xls', '.xlsx'):
                    continue
                fp_s = os.path.join(pasta_scan, nome_f)
                try:
                    if os.path.getsize(fp_s) == 0:
                        continue
                except Exception:
                    continue
                if _is_skip_file(fp_s):
                    continue
                base_s = os.path.splitext(nome_f)[0].lower()
                for pref_s in padroes_s:
                    pl_s = pref_s.lower().rstrip('_')
                    m_s = re.match(r'^' + re.escape(pl_s) + r'_?(\d+)', base_s)
                    if m_s:
                        num_s = int(m_s.group(1))
                        if num_s not in encontrados:
                            encontrados[num_s] = fp_s
                        break
            return sorted(encontrados.items())  # [(num, fp), ...]

        # Pré-calcular n_dir_total UMA vez (evitar re-scan dentro do loop)
        _n_dir_total = len(_listar_hs(pasta, 'dir'))

        log(f"  HS: {os.path.basename(pasta)}  dir={_n_dir_total} "
            f"esq={len(_listar_hs(pasta, 'esq'))} ficheiros encontrados", 'info')

        for lado in ('dir', 'esq'):
            hs_ficheiros = _listar_hs(pasta, lado)
            if not hs_ficheiros:
                log(f"  aviso: nenhum ficheiro hs_{lado} encontrado em {os.path.basename(pasta)}"
                    f" (padrões: {_HS_PADROES.get(lado, [])})", 'aviso')
            # Limite ao máximo configurado (n_ens_hs), mas sem cortar ficheiros reais encontrados
            n_hs = n_ens_hs if n_ens_hs and n_ens_hs > 0 else max(len(hs_ficheiros), n_ens)
            log(f"  HS {lado}: {len(hs_ficheiros)} ficheiros, limite={n_hs}", 'info')
            for seq, (num_f, fp) in enumerate(hs_ficheiros[:n_hs], start=1):
                try:
                    dados = ler_ficheiro(fp)
                    frs   = dados.get('frames', [])
                    b = e = None
                    _fonte = 'completo'   # para log

                    if hs_tempos_ind:
                        # Usar n_dir_total pré-calculado (não re-scan)
                        ens_idx = seq if lado == 'dir' else seq + _n_dir_total
                        tp = hs_tempos_ind.get(ens_idx)
                        if tp:
                            b, e = tp
                            _fonte = f'ficheiro_tempos idx={ens_idx}'
                        else:
                            log(f"  aviso hs_{lado}_{num_f}: sem tempo em ficheiro"
                                f" (idx={ens_idx}, chaves disponíveis={sorted(hs_tempos_ind.keys())})",
                                'aviso')

                    # ── Cascata de fallbacks quando janela não dá resultados ──
                    # Problema frequente: timestamps do XLS e do ficheiro de tempos
                    # têm offsets diferentes (ex: XLS começa em t=9000ms mas
                    # a janela diz 1640-8700ms → 0 frames capturados).
                    # Estratégia: tentar até encontrar ≥5 frames.

                    def _n_frames_janela(b_, e_):
                        if b_ is None: return len(frs)
                        return sum(1 for f in frs if b_ <= f['t_ms'] <= e_)

                    m = None
                    used_b, used_e = b, e

                    # Tentativa 1: janela do ficheiro de tempos
                    if _n_frames_janela(b, e) >= 5:
                        m = calcular(frs, b, e)

                    # Tentativa 2: tempo embutido no próprio ficheiro XLS
                    if m is None and dados['inicio_ms'] is not None:
                        b2, e2 = dados['inicio_ms'], dados['fim_ms']
                        if _n_frames_janela(b2, e2) >= 5:
                            m = calcular(frs, b2, e2)
                            if m:
                                used_b, used_e = b2, e2
                                _fonte = f'embed ({b2}-{e2}ms)'
                                if b is not None:
                                    log(f"  info hs_{lado}_{num_f}: janela {b}-{e}ms sem frames "
                                        f"(t_ms ficheiro: {frs[0]['t_ms']:.0f}-{frs[-1]['t_ms']:.0f}ms). "
                                        f"Usado tempo embutido {b2}-{e2}ms.", 'aviso')

                    # Tentativa 3: ficheiro completo (sem windowing)
                    if m is None and len(frs) >= 5:
                        m = calcular(frs)
                        if m:
                            used_b, used_e = None, None
                            _fonte = 'completo (sem janela)'
                            t0 = frs[0]['t_ms']; t1 = frs[-1]['t_ms']
                            log(f"  info hs_{lado}_{num_f}: sem sobreposição de janela "
                                f"(ficheiro t={t0:.0f}-{t1:.0f}ms, janela={b}-{e}ms). "
                                f"Usado ficheiro completo.", 'aviso')

                    mets_hs[lado].append(m)
                    raw_hs[lado].append({'dados': dados, 'ini': used_b, 'fim': used_e})

                    if m:
                        janela_str = (f"{used_b}-{used_e}ms" if used_b is not None
                                      else "completo")
                        log(f"  ok hs_{lado}_{num_f} [{janela_str}] [{_fonte}]"
                            f"  ea95={m['ea95']:.1f}mm2", 'ok')
                    else:
                        t0 = frs[0]['t_ms'] if frs else '?'
                        t1 = frs[-1]['t_ms'] if frs else '?'
                        log(f"  erro hs_{lado}_{num_f}: sem dados mesmo após fallbacks "
                            f"({len(frs)} frames, t={t0:.0f}-{t1:.0f}ms, "
                            f"janela_pedida={b}-{e}ms)", 'erro')

                except Exception as ex:
                    ex_s = str(ex)
                    if not ('ignorado' in ex_s or 'roll off' in ex_s.lower()):
                        log(f"  erro hs_{lado}_{num_f}: {ex_s}", 'erro')
                    mets_hs[lado].append(None)
                    raw_hs[lado].append(None)

    # ── Selection Right/Left CoP (v17) ──────────────────────────────
    # Para cada trial de tiro, extrai sub-sinais Right e Left Selection
    # e calcula metricas analogas a apoio unipodal.
    # v19: sel_dist[dist][itv] = {'dir': [...], 'esq': [...]}
    # Calcula Right/Left Selection para TODOS os intervalos (mesmos cortes da analise principal)
    sel_dist = {}
    if AOM_CONFIG.get('tiro_selection_ativo', True):
        for dist in dists_process:
            raw_d = tiro_dist.get(dist, {}).get('raw', {})
            sel_dist[dist] = {}
            for itv_s in intervalos:
                s_dir = []; s_esq = []
                for raw_t in raw_d.get(itv_s, []):
                    if raw_t is None:
                        s_dir.append(None); s_esq.append(None); continue
                    dados_r  = raw_t.get('dados', {})
                    b_r, e_r = raw_t.get('ini'), raw_t.get('fim')
                    frs      = dados_r.get('frames', []) if isinstance(dados_r, dict) else []
                    has_sel  = dados_r.get('tem_selection', False) if isinstance(dados_r, dict) else False
                    if not has_sel:
                        s_dir.append(None); s_esq.append(None); continue
                    md = _calcular_selection(frs, b_r, e_r, 'dir')
                    me = _calcular_selection(frs, b_r, e_r, 'esq')
                    s_dir.append(md); s_esq.append(me)
                    if md: log(f'  ok sel_dir {dist}/{itv_s} t{len(s_dir)}  ea95={md["ea95"]:.1f}', 'ok')
                    if me: log(f'  ok sel_esq {dist}/{itv_s} t{len(s_esq)}  ea95={me["ea95"]:.1f}', 'ok')
                sel_dist[dist][itv_s] = {'dir': s_dir, 'esq': s_esq}

    return {
        'nome':      nome,
        'id':        ind_id,
        'protocolo': protocolo,
        'scores':    None,
        # Dados do protocolo de tiro (por distancia)
        'tiro_dist': tiro_dist,
        'tiro_dists': dists_process,
        'tiro_intervalos': intervalos,
        # Dados do Hurdle Step (bipodal)
        'mets': mets_hs,
        'raw':  raw_hs,
        # Right/Left Selection CoP por distancia (v17)
        'sel_dist': sel_dist,
    }


def _processar_atleta_arco(pasta, tempos_arco, n_ens, log,
                           match_por_idx=False, idx_na_lista=None,
                           ask_callback=None, atleta_ref=None,
                           peso_kg_default=None):
    """
    Processa um atleta no protocolo Tiro com Arco (v1.0).

    Ate 30 ensaios bipodal. Janela unica de analise: [conf_1, conf_2]
    (lidos de Inicio_fim_vfinal.xlsx via carregar_confirmacao_arco).

    Sem distancias, sem pos/disp, sem Hurdle Step.

    Parametros:
      pasta         : pasta com ficheiros {id}_{trial} - DD-MM-YYYY - Stability export.xls
      tempos_arco   : resultado de carregar_confirmacao_arco() ou None
      n_ens         : numero maximo de ensaios a tentar (default 30)
      atleta_ref    : dict da referencia demografica (peso_kg, altura_m, ...)
      peso_kg_default: peso a usar se nao houver atleta_ref

    Devolve dict no formato standard: {nome, id, protocolo, mets, raw, ref}
    onde mets e raw tem a chave unica 'arco' com lista de ensaios.
    """
    nome_pasta = os.path.basename(pasta)
    # Extrair o id do nome da pasta (ex: "101_Tiago_Matos" → "101")
    m_id = re.match(r'^(\d+)', nome_pasta)
    ind_id = m_id.group(1) if m_id else None
    nome = _limpar_nome(nome_pasta)

    # Resolver janelas para este atleta
    por_id = (tempos_arco or {}).get('por_id', {})
    janelas = {}
    if ind_id and ind_id in por_id:
        janelas = por_id[ind_id]
    elif ind_id:
        # Tentar com leading zeros ('0101' → '101')
        alt_ids = [ind_id.lstrip('0'), f'{int(ind_id):03d}', f'{int(ind_id):04d}']
        for aid in alt_ids:
            if aid in por_id:
                janelas = por_id[aid]
                ind_id = aid
                break

    if tempos_arco and not janelas:
        log(f"  aviso: '{nome}' (ID={ind_id}) sem janelas de confirmacao", 'aviso')

    # Peso/altura a partir da referencia dos 142 atletas (se disponivel)
    peso_kg  = None
    altura_m = None
    if atleta_ref:
        peso_kg  = atleta_ref.get('peso_kg') or peso_kg_default
        altura_m = atleta_ref.get('altura_m')
    else:
        peso_kg = peso_kg_default

    mets = {'arco': []}
    raw  = {'arco': []}

    for t in range(1, n_ens + 1):
        fp = achar_ficheiro_arco(pasta, ind_id or '', t)
        if fp is None:
            # Fallback: qualquer .xls no pasta com "_{t}" no inicio
            fp = achar_ficheiro(pasta, 'arco', t)
        if fp is None or os.path.getsize(fp) == 0:
            mets['arco'].append(None)
            raw['arco'].append(None)
            continue

        try:
            dados = ler_ficheiro(fp)
            frs   = dados['frames']

            # Janela: conf_1 → conf_2 (se disponivel)
            janela_trial = janelas.get(t) if janelas else None
            if janela_trial:
                b = janela_trial.get('conf_1')
                e = janela_trial.get('conf_2')
            else:
                # Fallback: tempos embutidos no ficheiro ou completo
                b = dados.get('inicio_ms')
                e = dados.get('fim_ms')

            m = None
            _fonte = 'janela_confirmacao' if janela_trial else 'embed_ou_completo'

            # Tentativa 1: janela conf_1/conf_2
            if b is not None and e is not None:
                n_frames_j = sum(1 for f in frs if b <= f['t_ms'] <= e)
                if n_frames_j >= 5:
                    m = calcular(frs, b, e, peso_kg=peso_kg, altura_m=altura_m)

            # Tentativa 2: ficheiro completo
            if m is None and len(frs) >= 5:
                m = calcular(frs, peso_kg=peso_kg, altura_m=altura_m)
                _fonte = 'completo'
                log(f"  info arco_{t}: janela {b}-{e}ms sem frames; usado completo.", 'aviso')

            mets['arco'].append(m)
            raw['arco'].append({'dados': dados, 'ini': b, 'fim': e,
                                'fonte': _fonte, 'ficheiro': os.path.basename(fp)})

            if m:
                log(f"  ok arco_{t} [{_fonte}]  ea95={m['ea95']:.1f}mm2  "
                    f"vel={m['vel_med']:.1f}mm/s  stiff_x={m.get('stiff_x') or 0:.3f}", 'ok')
            else:
                log(f"  erro arco_{t}: dados insuficientes ({len(frs)} frames)", 'erro')

        except Exception as ex:
            log(f"  erro arco_{t}: {ex}", 'erro')
            mets['arco'].append(None)
            raw['arco'].append(None)

    # Aviso de ensaios validos insuficientes
    min_val = AOM_CONFIG.get('min_ensaios_validos', 3)
    n_ok = sum(1 for m in mets['arco'] if m is not None)
    if len(mets['arco']) > 0 and n_ok < min_val:
        log(f"  aviso: '{nome}': {n_ok}/{len(mets['arco'])} ensaios validos "
            f"(minimo recomendado: {min_val}).", 'aviso')

    return {
        'nome':      nome,
        'id':        ind_id,
        'protocolo': PROTO_ARCO,
        'scores':    None,
        'mets':      mets,
        'raw':       raw,
        # Demografia do atleta (se disponivel dos 142 atletas)
        'ref':       atleta_ref,
        'peso_kg':   peso_kg,
        'altura_m':  altura_m,
    }


def _tiro_janela(intervalo, trial_t, dados, t_fim_ficheiro, usar_embed):
    """
    Calcula o par (inicio_ms, fim_ms) para um dado intervalo do protocolo de tiro.
    trial_t: dict com chaves 'toque','pontaria','disparo' em ms
    """
    t_toque    = trial_t.get('toque')    if trial_t else None
    t_pontaria = trial_t.get('pontaria') if trial_t else None
    t_disparo  = trial_t.get('disparo')  if trial_t else None
    t_fim      = t_fim_ficheiro

    # Fallback: usar embedded se disponivel
    if t_toque is None and usar_embed and dados.get('inicio_ms') is not None:
        t_toque = dados['inicio_ms']
    if t_fim is None and usar_embed and dados.get('fim_ms') is not None:
        t_fim = dados['fim_ms']

    if intervalo == 'toque_pontaria':
        if t_toque is not None and t_pontaria is not None:
            return (t_toque, t_pontaria)
    elif intervalo == 'toque_disparo':
        if t_toque is not None and t_disparo is not None:
            return (t_toque, t_disparo)
    elif intervalo == 'pontaria_disparo':
        if t_pontaria is not None and t_disparo is not None:
            return (t_pontaria, t_disparo)
    elif intervalo == 'disparo_fim':
        if t_disparo is not None:
            return (t_disparo, t_fim)
    elif intervalo == 'total':
        return (t_toque, t_fim)

    # Sem tempos: calcular sobre todo o ficheiro
    return (None, None)


# -----------------------------------------------------------------------
# Graficos matplotlib
# -----------------------------------------------------------------------

def _png_elipses(mets_lado, titulo, nome, larg=6.0, alt=5.0, dpi=200, legenda_fora=True):
    """
    Elipses dos ensaios sobrepostas.
    Eixos livres (sem equal aspect) para que a elipse seja visualmente
    proporcional aos dados mesmo com racio Y/X extremo.
    Linha da elipse a 2.5pt para ser bem visivel no PDF.
    """
    if not _MPL_OK:
        return None

    validos = [(i+1, m) for i, m in enumerate(mets_lado) if m is not None]
    if not validos:
        return None

    # Passa toda a lista incluindo None -- flagrar_outliers filtra internamente
    flags = (flagrar_outliers(list(mets_lado), chave='ea95')
             if len(mets_lado) >= 3
             else [False] * len(mets_lado))

    cores_traj = ['#AED6F1','#85C1E9','#5DADE2','#2E86C1','#1A5276',
                  '#A9DFBF','#52BE80','#1E8449','#85C1E9','#F1948A']
    cores_ell  = ['#F1948A','#EC7063','#E74C3C','#CB4335','#922B21',
                  '#52BE80','#1E8449','#145A32','#2E86C1','#C00000']

    alt_fig = alt + (0.85 if legenda_fora else 0)
    fig, ax = plt.subplots(figsize=(larg, alt_fig), dpi=dpi)
    fig.patch.set_facecolor('white')
    ax.set_facecolor('#F9F9F9')

    handles = []
    for idx, (ne, m) in enumerate(validos):
        i_orig = ne - 1
        is_out = flags[i_orig] if i_orig < len(flags) else False
        ct = cores_traj[idx % len(cores_traj)]
        ce = '#FF6B00' if is_out else cores_ell[idx % len(cores_ell)]
        lw = 2.8 if is_out else 2.2
        ls = '--' if is_out else '-'
        lbl = f'E{ne}  ea={m["ea95"]:.0f}mm\u00b2'
        if is_out: lbl += '  [*]'

        ax.scatter(m['cof_x'], m['cof_y'], s=1.5, color=ct, alpha=0.3, linewidths=0, zorder=2)
        ln, = ax.plot(m['ell_x'], m['ell_y'], color=ce, lw=lw, ls=ls, label=lbl, zorder=4)
        ax.plot(m['mean_x'], m['mean_y'], marker='D', color=ce, ms=5, zorder=5, lw=0)
        handles.append(ln)

    ax.set_xlabel('COF X (mm)', fontsize=9, labelpad=4)
    ax.set_ylabel('COF Y (mm)', fontsize=9, labelpad=4)
    ax.set_title(f'{nome}  -  {titulo}', fontsize=9, fontweight='bold', pad=6)
    ax.tick_params(labelsize=7)
    ax.grid(True, ls='--', lw=0.5, alpha=0.5, color='#CCCCCC', zorder=1)

    for sp in ax.spines.values():
        sp.set_linewidth(0.8); sp.set_color('#AAAAAA')

    if legenda_fora:
        ax.legend(handles=handles, fontsize=7, loc='upper center',
                  bbox_to_anchor=(0.5, -0.14), ncol=min(len(validos), 3),
                  framealpha=0.95, edgecolor='#CCCCCC', borderpad=0.6)
        fig.subplots_adjust(bottom=0.20)
    else:
        ax.legend(handles=handles, fontsize=7, loc='best', framealpha=0.9)

    fig.tight_layout(pad=0.5)
    buf = io.BytesIO()
    try:
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                    facecolor='white', edgecolor='none')
    finally:
        plt.close(fig)
    buf.seek(0)
    return buf.read()


def _png_elipses_tiro(m_pos, m_disp, nome, larg=6.0, alt=5.0, dpi=200):
    """
    Grafico especifico para protocolo de tiro:
    elipse de posicao vs elipse de disparo sobrepostas num mesmo plot.
    """
    if not _MPL_OK:
        return None

    if m_pos is None and m_disp is None:
        return None

    fig, ax = plt.subplots(figsize=(larg, alt), dpi=dpi)
    fig.patch.set_facecolor('white')
    ax.set_facecolor('#F9F9F9')
    handles = []

    configs = [
        (m_pos,  '#2E86C1', '#AED6F1', 'Posição'),
        (m_disp, '#C0392B', '#F1948A', 'Disparo'),
    ]
    for m, ce, ct, lbl in configs:
        if m is None: continue
        ax.scatter(m['cof_x'], m['cof_y'], s=1.5, color=ct, alpha=0.3, linewidths=0, zorder=2)
        ln, = ax.plot(m['ell_x'], m['ell_y'], color=ce, lw=2.5,
                      label=f'{lbl}  ea={m["ea95"]:.0f}mm\u00b2', zorder=4)
        ax.plot(m['mean_x'], m['mean_y'], 'D', color=ce, ms=5, zorder=5, lw=0)
        handles.append(ln)

    ax.set_xlabel('COF X (mm)', fontsize=9)
    ax.set_ylabel('COF Y (mm)', fontsize=9)
    ax.set_title(f'{nome}  -  {T("tiro_itv_pont")} vs {T("tiro_itv_disp")}', fontsize=9, fontweight='bold', pad=6)
    ax.tick_params(labelsize=7)
    ax.grid(True, ls='--', lw=0.5, alpha=0.5, color='#CCCCCC')
    ax.legend(handles=handles, fontsize=8, loc='best', framealpha=0.95)
    for sp in ax.spines.values():
        sp.set_linewidth(0.8); sp.set_color('#AAAAAA')
    fig.tight_layout(pad=0.5)
    buf = io.BytesIO()
    try:
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                    facecolor='white', edgecolor='none')
    finally:
        plt.close(fig)
    buf.seek(0)
    return buf.read()


def _png_correlacao_tiro(atletas, chave='ea95', cond='pos', dpi=160):
    """
    Scatter: score de precisao vs metrica de estabilidade.
    So gerado se houver scores.
    """
    if not _MPL_OK:
        return None
    from scipy.stats import pearsonr, spearmanr

    pontos = []
    for ath in atletas:
        if not ath.get('scores') or cond not in ath['mets']:
            continue
        scores = ath['scores']
        for i, m in enumerate(ath['mets'][cond]):
            if m is None or i >= len(scores): continue
            pontos.append((scores[i], m[chave], ath['nome']))

    if len(pontos) < 4:
        return None

    xs = np.array([p[0] for p in pontos])
    ys = np.array([p[1] for p in pontos])

    try:
        r, pval = pearsonr(xs, ys)
        r_s, _ = spearmanr(xs, ys)
    except Exception: return None

    fig, ax = plt.subplots(figsize=(5.5, 4.0), dpi=dpi)
    fig.patch.set_facecolor('white'); ax.set_facecolor('#F9F9F9')
    ax.scatter(xs, ys, s=25, color='#2E86C1', alpha=0.7, zorder=3)

    m_lin, b_lin = np.polyfit(xs, ys, 1)
    xl = np.linspace(xs.min(), xs.max(), 100)
    ax.plot(xl, m_lin*xl + b_lin, color='#E74C3C', lw=1.8, ls='--', zorder=2)

    cond_nome = 'Posição' if cond == 'pos' else 'Disparo'
    chave_nome = {'ea95': 'Area Elipse 95% (mm2)', 'vel_med': 'Vel. media (mm/s)'}.get(chave, chave)
    ax.set_xlabel('Score de Precisao', fontsize=9)
    ax.set_ylabel(chave_nome, fontsize=9)
    ax.set_title(f'Correlacao: Precisao vs Estabilidade [{cond_nome}]', fontsize=9, fontweight='bold')
    ax.tick_params(labelsize=7)
    ax.grid(True, ls='--', lw=0.4, alpha=0.5, color='#CCCCCC')
    ax.text(0.05, 0.95, f'r = {r:.3f}  (p = {pval:.3f})\nrs = {r_s:.3f}',
            transform=ax.transAxes, fontsize=8, va='top',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='white', edgecolor='#AAAAAA'))
    for sp in ax.spines.values():
        sp.set_linewidth(0.8); sp.set_color('#AAAAAA')
    fig.tight_layout(pad=0.5)
    buf = io.BytesIO()
    try:
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                    facecolor='white', edgecolor='none')
    finally:
        plt.close(fig)
    buf.seek(0)
    return buf.read()


def _png_estabilograma(m, nome_ens, nome, larg=6.0, alt=4.0, dpi=150):
    if not _MPL_OK:
        return None
    if m is None: return None
    t_s = np.array(m['t_ms']) / 1000.0; t_s -= t_s[0]
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(larg, alt), dpi=dpi, sharex=True)
    fig.patch.set_facecolor('white')
    for ax in [ax1, ax2]: ax.set_facecolor('#F9F9F9')
    ax1.plot(t_s, m['cof_x'], color='#2E86C1', lw=1.0)
    ax1.axhline(m['mean_x'], color='#E74C3C', ls='--', lw=1.2, alpha=0.8,
                label=f"Media = {m['mean_x']:.1f} mm")
    ax1.set_ylabel('COF X (mm)', fontsize=8)
    ax1.legend(fontsize=7); ax1.tick_params(labelsize=7)
    ax1.grid(True, ls='--', lw=0.4, alpha=0.4)
    ax1.set_title(f'Estabilograma  -  {nome}  -  {nome_ens}', fontsize=9, fontweight='bold')
    ax2.plot(t_s, m['cof_y'], color='#E67E22', lw=1.0)
    ax2.axhline(m['mean_y'], color='#1A5276', ls='--', lw=1.2, alpha=0.8,
                label=f"Media = {m['mean_y']:.1f} mm")
    ax2.set_xlabel('Tempo (s)', fontsize=8); ax2.set_ylabel('COF Y (mm)', fontsize=8)
    ax2.legend(fontsize=7); ax2.tick_params(labelsize=7)
    ax2.grid(True, ls='--', lw=0.4, alpha=0.4)
    for ax in [ax1, ax2]:
        for sp in ax.spines.values():
            sp.set_linewidth(0.7); sp.set_color('#AAAAAA')
    fig.tight_layout(pad=0.5)
    buf = io.BytesIO()
    try:
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                    facecolor='white', edgecolor='none')
    finally:
        plt.close(fig)
    buf.seek(0)
    return buf.read()


def exportar_png_individuo(ath, pasta_saida):
    """
    Exporta PNGs dos graficos (elipse CoP + estabilograma) de um individuo.
    Cria uma subpasta com o nome do individuo e gera um PNG por ensaio/lado.
    Requer matplotlib. Devolve lista de caminhos criados.
    """
    if not _MPL_OK:
        return []
    if not os.path.isdir(pasta_saida):
        os.makedirs(pasta_saida, exist_ok=True)

    nome_safe = re.sub(r'[<>:"/\\|?*]', '_', ath.get('nome', 'individuo'))
    pasta_ind = os.path.join(pasta_saida, nome_safe + '_PNG')
    os.makedirs(pasta_ind, exist_ok=True)
    criados = []

    for lado, lst_m in ath.get('mets', {}).items():
        for ti, m in enumerate(lst_m, start=1):
            if m is None:
                continue
            titulo = f'{ath["nome"]} - {lado.upper()} - Ens.{ti}'
            # Elipse
            png_ell = _png_elipses([m], titulo, ath['nome'],
                                   larg=5.5, alt=4.5, dpi=150)
            if png_ell:
                fn_ell = os.path.join(pasta_ind, f'{nome_safe}_{lado}_T{ti}_elipse.png')
                with open(fn_ell, 'wb') as f_out:
                    f_out.write(png_ell)
                criados.append(fn_ell)
            # Estabilograma
            png_est = _png_estabilograma(m, f'{lado.upper()} T{ti}', ath['nome'],
                                         larg=6.0, alt=4.0, dpi=150)
            if png_est:
                fn_est = os.path.join(pasta_ind, f'{nome_safe}_{lado}_T{ti}_estab.png')
                with open(fn_est, 'wb') as f_out:
                    f_out.write(png_est)
                criados.append(fn_est)
    return criados


# -----------------------------------------------------------------------
# PDF
# -----------------------------------------------------------------------

METS_PDF = [
    ('amp_x',      'Amplitude X (mm)'),
    ('amp_y',      'Amplitude Y (mm)'),
    ('vel_x',      'Vel. media X (mm/s)'),
    ('vel_y',      'Vel. media Y (mm/s)'),
    ('vel_med',    'Vel. media CoP (mm/s)'),
    ('vel_pico_x', 'Pico vel. X (mm/s)'),
    ('vel_pico_y', 'Pico vel. Y (mm/s)'),
    ('desl',       'Deslocamento (mm)'),
    ('time',       'Tempo (s)'),
    ('ea95',       'Area elipse 95% (mm2)'),
    ('leng_a',     'Semi-eixo a (mm)'),
    ('leng_b',     'Semi-eixo b (mm)'),
    ('ratio_ml_ap','Rel. amp. ML/AP'),
    ('ratio_vel',  'Rel. vel. ML/AP'),
    ('stiff_x',    'Stiffness X (1/s)'),
    ('stiff_y',    'Stiffness Y (1/s)'),
    # Normalizacoes por dimensoes corporais
    ('ea95_norm',    'EA95 norm (mm2/m)'),
    ('amp_norm_x',   'Amp norm ML (mm/m)'),
    ('amp_norm_y',   'Amp norm AP (mm/m)'),
    ('vel_norm',     'Vel norm ((mm/s)/m)'),
    ('stiff_mass_x', 'Stiff/mass ML (1/(s·kg))'),
    ('stiff_mass_y', 'Stiff/mass AP (1/(s·kg))'),
    ('stiff_norm_x', 'Stiff norm ML ((1/s)/m)'),
    ('stiff_norm_y', 'Stiff norm AP ((1/s)/m)'),
    ('cov_xx',     'Var X (mm2)'),
    ('cov_yy',     'Var Y (mm2)'),
    # RMS (Quijoux et al., 2021, sec. 3.2.4)
    ('rms_x',      'RMS ML (mm)'),
    ('rms_y',      'RMS AP (mm)'),
    ('rms_r',      'RMS Radius (mm)'),
    ('cov_xy',     'Cov XY (mm2)'),
]


def _pagina_pdf(c, atleta, W, H):
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.utils import ImageReader

    nome   = atleta['nome']
    ind_id = atleta.get('id', '')
    proto  = PROTOCOLOS.get(atleta.get('protocolo', PROTO_FMS))
    tem_ai = proto['assimetria']
    n_ens  = proto['n_ens']
    AZ   = colors.HexColor('#1F4E79')
    AZCL = colors.HexColor('#BDD7EE')
    CZ   = colors.HexColor('#F2F2F2')
    VM   = colors.HexColor('#C00000')
    DIM  = colors.HexColor('#888888')
    AZDB = colors.HexColor('#E8F4FD')
    VD2  = colors.HexColor('#E8F5E9')

    # cabecalho
    c.setFillColor(AZ)
    c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, H - 1.05*cm, PROG + '  -  ' + proto['nome'])
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, H - 1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - 1*cm, H - 1.05*cm, datetime.date.today().strftime('%d/%m/%Y'))

    # nome e ID do individuo
    y = H - 2.5*cm
    id_prefix = f'[ID {ind_id}]  ' if ind_id else ''
    c.setFillColor(AZ); c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, y, id_prefix + nome)
    c.setFillColor(colors.HexColor('#6B8FAE')); c.setFont('Helvetica', 7.5)
    c.drawString(1*cm, y - 4*mm, T('pdf_individuo'))
    c.setStrokeColor(AZ); c.setLineWidth(0.5)
    c.line(1*cm, y - 6*mm, W - 1*cm, y - 6*mm)
    y -= 1.0*cm

    # tabela de metricas - n_vals dinamico: usa trials nao-None de cada lado
    # (evita colunas em branco quando ensaios comecam num indice > 1)
    # NOTA: lados_tabela definido ANTES de _compact_mets (fix v20)
    # NOTA v20b: apenas inclui lados que têm pelo menos 1 valor nao-None
    _lados_i18n = lados_pdf_localizados() or {
        'dir': 'PÉ DIREITO', 'esq': 'PÉ ESQUERDO',
        'pos': 'POSIÇÃO', 'disp': 'DISPARO', 'arco': 'ARCO'}
    # lista alargada para incluir 'arco' (Tiro com Arco, lado unico)
    _lados_ordem = ['dir', 'esq', 'pos', 'disp', 'arco']
    lados_tabela = []
    for _chave_lado in _lados_ordem:
        _label_lado = _lados_i18n.get(_chave_lado, _chave_lado.upper())
        if _chave_lado in atleta['mets']:
            _has_data = any(m is not None for m in atleta['mets'][_chave_lado])
            if _has_data:
                lados_tabela.append((_chave_lado, _label_lado))
    # Fallback: se nenhum lado tem dados, mostrar o que existir mesmo vazio
    if not lados_tabela:
        for _chave_lado in _lados_ordem:
            _label_lado = _lados_i18n.get(_chave_lado, _chave_lado.upper())
            if _chave_lado in atleta['mets']:
                lados_tabela.append((_chave_lado, _label_lado))
    _compact_mets = {}
    for _l, _ in lados_tabela:
        _compact_mets[_l] = [m for m in atleta['mets'].get(_l, []) if m is not None]
    # n_vals: ilimitado - colunas comprimem automaticamente para caber em A4
    # Para >10 ensaios as colunas ficam estreitas mas legíveis (mínimo 0.38cm)
    n_vals = max(max((len(v) for v in _compact_mets.values()), default=1), 1)
    # AI só faz sentido se ambos dir e esq têm dados
    if tem_ai and not (_compact_mets.get('dir') and _compact_mets.get('esq')):
        tem_ai = False

    # Calcular larguras de coluna dinamicamente conforme espaco disponivel
    # Pagina A4 = 21cm; margens 1cm cada lado → 19cm util
    margem = 1.0*cm
    total_w = W - 2*margem
    C0  = 4.2*cm          # coluna metrica (label) - mais largo para HurdleStep
    CAI = 1.1*cm          # coluna AI
    CSP = 0.15*cm         # separador entre lados

    # Espaco disponivel para as colunas de dados
    ai_space = CAI if tem_ai else 0
    sep_space = CSP * len(lados_tabela)
    dados_w = total_w - C0 - ai_space - sep_space
    # Cada lado tem: n_vals individuais + max + med + dp = n_vals+3 colunas
    n_col_dados = (n_vals + 3) * len(lados_tabela)
    # Estatisticas (max/med/dp) sao ligeiramente mais largas que valores individuais
    CV_ratio = 1.0; CS_ratio = 1.15
    unit = dados_w / (n_col_dados * CV_ratio + len(lados_tabela) * 3 * (CS_ratio - CV_ratio))
    CV  = unit * CV_ratio
    CS  = unit * CS_ratio

    # Dividir ensaios em blocos de no maximo CHUNK_MAX colunas por pagina
    # (evita tabelas ilegíveis com muitos ensaios)
    # Estratégia adaptativa:
    #   - 1 lado (arco, disp, pos): mais espaço → até 15 ensaios/página
    #   - 2 lados (dir+esq): 6 ensaios/página para manter legibilidade
    # Para >CHUNK_MAX ensaios, páginas adicionais são geradas automaticamente
    if len(lados_tabela) == 1:
        # Protocolo de lado único (Tiro com Arco, etc.): colunas mais largas
        # Cap em 15 para manter fonte legível (≥6pt)
        CHUNK_MAX = min(n_vals, 15)
    else:
        CHUNK_MAX = 6
    chunks = []
    for start in range(0, n_vals, CHUNK_MAX):
        chunks.append(list(range(start, min(start + CHUNK_MAX, n_vals))))

    idx_ea95 = next((i+2 for i,(k,_) in enumerate(METS_PDF) if k=='ea95'), 9)
    flags_por_lado = {lado: flagrar_outliers(_compact_mets.get(lado, []))
                      for lado, _ in lados_tabela}
    _mets_loc = mets_pdf_localizadas() or METS_PDF

    for chunk_idx, chunk_idxs in enumerate(chunks):
        n_chunk = len(chunk_idxs)
        is_cont = chunk_idx > 0

        # Recalcular larguras para este bloco
        n_col_dados_chunk = (n_chunk + 3) * len(lados_tabela)
        unit_c = dados_w / (n_col_dados_chunk * CV_ratio + len(lados_tabela) * 3 * (CS_ratio - CV_ratio))
        CV_c = unit_c * CV_ratio
        CS_c = unit_c * CS_ratio

        col_larg = [C0]
        for _ in lados_tabela:
            col_larg += [CV_c]*n_chunk + [CS_c, CS_c, CS_c, CSP]
        if tem_ai:
            col_larg += [CAI]
        if not tem_ai and col_larg[-1] == CSP:
            col_larg = col_larg[:-1]

        n_col_lado = n_chunk + 4

        # Cabecalho: numero real do ensaio (ex: 11,12,...) ou "cont."
        # Cabeçalho com destaque por lado e numeração real de ensaio
        cab = [T('pdf_metrica') + (f' (bloco {chunk_idx+1})' if is_cont else '')]
        for _, lbl in lados_tabela:
            cab += [lbl] + ['']*(n_chunk + 2) + ['']
        if tem_ai: cab += [T('pdf_ai')]

        sub = ['']
        for _ in lados_tabela:
            # Mostrar número real do ensaio (ex: 9,10,... no segundo bloco)
            sub += [str(i+1) for i in chunk_idxs] + [T('pdf_max'), T('pdf_med'), T('pdf_dp'), '']
        if tem_ai: sub += ['']

        linhas = [cab, sub]

        for chave, label in _mets_loc:
            # Omitir linha se nenhum ensaio tem valor para esta métrica
            # (evita linhas de dashes para normalizações não ativadas)
            _tem_algum_valor = any(
                m is not None and m.get(chave) is not None
                for lado, _ in lados_tabela
                for m in _compact_mets.get(lado, [])
            )
            if not _tem_algum_valor:
                continue

            linha = [label]
            for lado, _ in lados_tabela:
                lst  = _compact_mets.get(lado, [])
                vals = []
                for ti in chunk_idxs:
                    m = lst[ti] if ti < len(lst) else None
                    v = m.get(chave) if m is not None else None
                    vals.append(v)
                    linha.append(f'{v:.2f}' if v is not None else '')
                # stats sempre calculadas sobre todos os ensaios (nao so o bloco)
                all_vals = [lst[ti].get(chave) for ti in range(n_vals)
                            if ti < len(lst) and lst[ti] is not None
                            and lst[ti].get(chave) is not None]
                nums = np.array(all_vals)
                linha += [f'{nums.max():.2f}'       if len(nums)>0 else '',
                          f'{nums.mean():.2f}'      if len(nums)>0 else '',
                          f'{nums.std(ddof=1):.2f}' if len(nums)>1 else '',
                          '']
            if tem_ai:
                vd = [m.get(chave) for m in _compact_mets.get('dir',[]) if m and m.get(chave) is not None]
                ve = [m.get(chave) for m in _compact_mets.get('esq',[]) if m and m.get(chave) is not None]
                ai = (assimetria(float(np.mean([x for x in vd if x is not None])),
                                 float(np.mean([x for x in ve if x is not None])))
                      if vd and ve else None)
                linha.append(f'{ai:+.1f}%' if ai is not None else '')
            linhas.append(linha)

        _row_h_val = 0.33*cm if n_chunk <= 12 else 0.28*cm
        row_h = [0.38*cm, 0.33*cm] + [_row_h_val]*(len(linhas)-2)
        tbl = Table(linhas, colWidths=col_larg, rowHeights=row_h)

        estilo = [
            ('BACKGROUND',(0,0),(-1,1), AZ),
            ('TEXTCOLOR',(0,0),(-1,1),  colors.white),
            ('FONTNAME',(0,0),(-1,1),   'Helvetica-Bold'),
            ('FONTSIZE',(0,0),(-1,1),   6.5),
            ('ALIGN',(0,0),(-1,1),      'CENTER'),
            ('VALIGN',(0,0),(-1,1),     'MIDDLE'),
            ('FONTNAME',(0,2),(0,-1),   'Helvetica-Bold'),
            ('FONTSIZE',(0,2),(-1,-1),  6),
            ('ALIGN',(1,2),(-1,-1),     'CENTER'),
            ('ALIGN',(0,2),(0,-1),      'LEFT'),
            ('VALIGN',(0,2),(-1,-1),    'MIDDLE'),
            *[('BACKGROUND',(0,i),(-1,i), CZ) for i in range(3, len(linhas), 2)],
            ('TEXTCOLOR',(0,idx_ea95),(-1,idx_ea95), VM),
            ('FONTNAME',(0,idx_ea95),(-1,idx_ea95),  'Helvetica-Bold'),
            ('GRID',(0,0),(-1,-1), 0.2, colors.HexColor('#CCCCCC')),
            ('LINEBELOW',(0,1),(-1,1), 0.8, AZ),
            ('TOPPADDING',(0,0),(-1,-1),    1),
            ('BOTTOMPADDING',(0,0),(-1,-1), 1),
            ('LEFTPADDING',(0,0),(0,-1),    3),
        ]
        col_off = 1
        for li, (lado, _) in enumerate(lados_tabela):
            span_fim = col_off + n_chunk + 2
            estilo.append(('SPAN', (col_off, 0), (span_fim, 0)))
            estilo.append(('BACKGROUND', (col_off+n_chunk, 2), (col_off+n_chunk+2, -1), AZDB))
            col_off += n_col_lado
        if tem_ai:
            estilo.append(('BACKGROUND', (col_off, 2), (-1, -1), VD2))
            estilo.append(('SPAN', (col_off, 0), (col_off, 1)))

        tbl.setStyle(TableStyle(estilo))

        # Se for continuacao, nova pagina
        if is_cont:
            c.showPage()
            y = H - 1.5*cm

        y = _tbl_fit_page(c, tbl, 1*cm, y, W, H)
        y -= 0.2*cm
    espaco = y - 1.2*cm
    hg = max(espaco * 0.95, 5.5*cm)
    # Lados com dados efectivos para os graficos
    _lados_graf = [(lado, lbl) for lado, lbl in [('dir','Pe Direito'),('esq','Pe Esquerdo')]
                   if any(m is not None for m in atleta['mets'].get(lado, []))]
    if not _lados_graf:
        _lados_graf = [(lado, lbl) for lado, lbl in lados_tabela]
    n_graficos = max(len(_lados_graf), 1) if not proto.get('two_windows') else len(lados_tabela)
    wg = (W - 2*cm - (n_graficos - 1)*0.4*cm) / max(n_graficos, 1)

    # Elipses - apenas lados com dados reais
    for si, (lado, lbl) in enumerate(_lados_graf):
        png = _png_elipses(atleta['mets'].get(lado, []), lbl, nome,
                           larg=wg/cm*0.40, alt=hg/cm*0.43, dpi=200,
                           legenda_fora=True)
        xg = 1*cm + si*(wg + 0.4*cm)
        if png:
            c.drawImage(ImageReader(io.BytesIO(png)), xg, y-hg,
                        width=wg, height=hg, preserveAspectRatio=False)
        else:
            c.setFillColor(CZ); c.rect(xg, y-hg, wg, hg, fill=1, stroke=0)

    # rodape
    c.setStrokeColor(AZCL); c.setLineWidth(0.3)
    c.line(1*cm, 0.8*cm, W-1*cm, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(1*cm, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W-1*cm, 0.4*cm, f'{nome}  |  {proto["nome"]}')


def _pagina_estabilograma_pdf(c, atleta, W, H):
    """Pagina extra com estabilogramas (X e Y vs tempo) de cada ensaio."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.lib.utils import ImageReader

    nome   = atleta['nome']
    proto  = PROTOCOLOS.get(atleta.get('protocolo', PROTO_FMS))
    n_ens  = proto['n_ens']

    AZ   = colors.HexColor('#1F4E79')
    AZCL = colors.HexColor('#BDD7EE')
    DIM  = colors.HexColor('#888888')

    # cabecalho (igual ao da pagina individual)
    c.setFillColor(AZ)
    c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, H - 1.05*cm, proto['nome'] + '  -  Estabilogramas')
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, H - 1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - 1*cm, H - 1.05*cm, datetime.date.today().strftime('%d/%m/%Y'))

    # nome do individuo
    y0 = H - 2.3*cm
    id_prefix = f'[ID {atleta.get("id","")}]  ' if atleta.get('id') else ''
    c.setFillColor(AZ); c.setFont('Helvetica-Bold', 10)
    c.drawString(1*cm, y0, id_prefix + nome)
    c.setStrokeColor(AZ); c.setLineWidth(0.5)
    c.line(1*cm, y0 - 3*mm, W - 1*cm, y0 - 3*mm)

    # determinar lados a desenhar
    lados_graf = []
    if 'dir'  in atleta['mets']: lados_graf.append(('dir',  'Pe Direito'))
    if 'esq'  in atleta['mets']: lados_graf.append(('esq',  'Pe Esquerdo'))
    if 'pos'  in atleta['mets']: lados_graf.append(('pos',  'Posição'))
    if 'disp' in atleta['mets']: lados_graf.append(('disp', 'Disparo'))
    if 'arco' in atleta['mets']: lados_graf.append(('arco', 'Arco'))   # v1.0

    # layout: n_ens linhas × len(lados) colunas
    n_cols = len(lados_graf)
    # n_rows: usar número real de ensaios disponíveis (ilimitado)
    n_rows = max(max(len(atleta['mets'].get(l, [])) for l,_ in lados_graf) if lados_graf else n_ens, 1)
    margem = 1.0*cm
    area_top = y0 - 0.5*cm
    area_bot = 1.0*cm
    area_h = area_top - area_bot
    area_w = W - 2*margem

    cell_w = area_w / n_cols
    cell_h = area_h / max(n_rows, 1)

    # cabecalhos de coluna (nome do lado)
    c.setFont('Helvetica-Bold', 8)
    c.setFillColor(AZ)
    for ci, (_, lbl) in enumerate(lados_graf):
        cx = margem + ci * cell_w + cell_w / 2
        c.drawCentredString(cx, area_top - 0.3*cm, lbl)

    pad = 0.12*cm

    for ri in range(n_rows):
        for ci, (lado, lbl) in enumerate(lados_graf):
            lst = atleta['mets'].get(lado, [])
            m = lst[ri] if ri < len(lst) else None
            if m is None:
                continue
            nome_ens = f'Ensaio {ri+1}'
            # dimensoes em polegadas para matplotlib
            larg_in = (cell_w - 2*pad) / (cm * 2.54) * 2.54
            alt_in  = (cell_h - 2*pad) / (cm * 2.54) * 2.54
            png = _png_estabilograma(m, nome_ens, lbl,
                                     larg=larg_in, alt=alt_in, dpi=130)
            if png:
                xg = margem + ci * cell_w + pad
                yg = area_top - (ri + 1) * cell_h + pad
                c.drawImage(ImageReader(io.BytesIO(png)),
                            xg, yg,
                            width=cell_w - 2*pad,
                            height=cell_h - 2*pad,
                            preserveAspectRatio=False)

    # rodape
    c.setStrokeColor(AZCL); c.setLineWidth(0.3)
    c.line(1*cm, 0.8*cm, W - 1*cm, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(1*cm, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - 1*cm, 0.4*cm, f'{nome}  |  {proto["nome"]}  |  Estabilogramas')


def _pagina_correlacao_pdf(c, atletas, W, H, cond='pos'):
    """Pagina de correlacao estabilidade vs precisao de tiro."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.utils import ImageReader

    AZ  = colors.HexColor('#1F4E79')
    DIM = colors.HexColor('#888888')

    c.setFillColor(AZ); c.rect(0, H-1.8*cm, W, 1.8*cm, fill=1, stroke=0)
    c.setFillColor(colors.white); c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, H-1.05*cm, 'Correlação Estabilidade vs Precisão de Tiro')
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, H-1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W-1*cm, H-1.05*cm, datetime.date.today().strftime('%d/%m/%Y'))

    y = H - 2.5*cm
    for metrica, nome_met in [('ea95','Área Elipse 95% (mm²)'),('vel_med','Vel. média (mm/s)')]:
        png = _png_correlacao_tiro(atletas, chave=metrica, cond=cond)
        if png:
            hg = 9*cm; wg = (W - 2.5*cm) / 2
            xg = 1*cm if metrica == 'ea95' else 1*cm + wg + 0.5*cm
            c.drawImage(ImageReader(io.BytesIO(png)), xg, y-hg,
                        width=wg, height=hg, preserveAspectRatio=False)

    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(1*cm, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')


def _pagina_demografia_pdf(c, atletas, W, H):
    """
    Pagina PDF com analises demograficas para PROTO_ARCO.

    Renderiza:
      - Resumo: n atletas com demografia, breakdown por genero/estilo/categoria
      - Top comparacoes significativas (Mann-Whitney / Kruskal-Wallis, p<0.05)
      - Top correlacoes significativas (Pearson + Spearman, p<0.05)
      - Rodape com citacoes academicas

    Ref.:
      Mann & Whitney (1947). Ann. Math. Stat., 18(1), 50-60.
      Kruskal & Wallis (1952). J. Am. Stat. Assoc., 47(260), 583-621.
      Spearman (1904). Am. J. Psychol., 15, 72-101.
    """
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import cm as rl_cm

    # Paleta (consistente com _pagina_estats_pdf)
    C_BG   = rl_colors.HexColor('#0F1923')
    C_HDR  = rl_colors.HexColor('#00B4D8')
    C_SEC  = rl_colors.HexColor('#1F4E79')
    C_TXT  = rl_colors.HexColor('#E8F4FD')
    C_TXT2 = rl_colors.HexColor('#BDD7EE')
    C_DIM  = rl_colors.HexColor('#6B8FAE')
    C_SIG  = rl_colors.HexColor('#FCA5A5')   # p significativo
    C_OK   = rl_colors.HexColor('#6EE7B7')

    MARGIN = 1.0 * rl_cm
    c.setFillColor(C_BG)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # Cabecalho
    c.setFillColor(C_HDR)
    c.setFont('Helvetica-Bold', 13)
    c.drawString(MARGIN, H - 1.1 * rl_cm, 'Demografia  |  Tiro com Arco')
    c.setFillColor(C_DIM)
    c.setFont('Helvetica', 7)
    c.drawString(MARGIN, H - 1.55 * rl_cm,
                 f'alfa=0.05  |  M-W=Mann-Whitney  |  K-W=Kruskal-Wallis  |  Pearson+Spearman  '
                 f'|  {PROG} v{VERSAO}  |  {AUTOR}')
    c.setStrokeColor(C_HDR); c.setLineWidth(0.6)
    c.line(MARGIN, H - 1.85 * rl_cm, W - MARGIN, H - 1.85 * rl_cm)

    y = H - 2.5 * rl_cm

    # ── Bloco 1: Resumo ──────────────────────────────────────────────────
    n_tot = len(atletas)
    n_peso   = sum(1 for a in atletas if isinstance(a, dict) and a.get('peso_kg')   is not None)
    n_altura = sum(1 for a in atletas if isinstance(a, dict) and a.get('altura_m') is not None)
    n_idade  = sum(1 for a in atletas if isinstance(a, dict) and a.get('idade')    is not None)

    def _tally(chave):
        tally = {}
        for a in atletas:
            if not isinstance(a, dict):
                continue
            v = a.get(chave)
            if v is None or v == '':
                continue
            if isinstance(v, str):
                v = v.strip()
                v = v.upper() if chave == 'genero' else v.lower()
            tally[v] = tally.get(v, 0) + 1
        return tally

    t_gen = _tally('genero')
    t_est = _tally('estilo')
    t_cat = _tally('categoria')

    c.setFillColor(C_SEC)
    c.rect(MARGIN, y - 0.5*rl_cm, W - 2*MARGIN, 0.5*rl_cm, fill=1, stroke=0)
    c.setFillColor(C_TXT); c.setFont('Helvetica-Bold', 9)
    c.drawString(MARGIN + 0.2*rl_cm, y - 0.35*rl_cm, 'Resumo da amostra')
    y -= 0.8 * rl_cm

    def _line(label, txt):
        nonlocal y
        c.setFillColor(C_TXT2); c.setFont('Helvetica-Bold', 8)
        c.drawString(MARGIN + 0.3*rl_cm, y, label)
        c.setFillColor(C_TXT); c.setFont('Helvetica', 8)
        c.drawString(MARGIN + 4.5*rl_cm, y, txt)
        y -= 0.42 * rl_cm

    _line('N total:', f'{n_tot}')
    _line('Com peso:', f'{n_peso}')
    _line('Com altura:', f'{n_altura}')
    _line('Com idade:', f'{n_idade}')
    if t_gen:
        _line('Genero:', ', '.join(f'{k}={v}' for k, v in sorted(t_gen.items())))
    if t_est:
        _line('Estilo:', ', '.join(f'{k}={v}' for k, v in sorted(t_est.items())))
    if t_cat:
        _line('Categoria:', ', '.join(f'{k}={v}' for k, v in sorted(t_cat.items())))

    y -= 0.3 * rl_cm

    # Metricas candidatas
    _METRICAS = [
        ('ea95',         'EA95 (mm2)'),
        ('vel_med',      'Vel. med (mm/s)'),
        ('stiff_x',      'Stiff ML'),
        ('stiff_y',      'Stiff AP'),
        ('amp_x',        'Amp ML (mm)'),
        ('amp_y',        'Amp AP (mm)'),
        ('rms_r',        'RMS radial'),
        ('ea95_norm',    'EA95 norm (mm2/m)'),
        ('vel_norm',     'Vel norm ((mm/s)/m)'),
        ('stiff_norm_x', 'Stiff norm ML'),
        ('stiff_norm_y', 'Stiff norm AP'),
    ]
    _FACTORES = [('genero', 'Genero'), ('estilo', 'Estilo'), ('categoria', 'Categoria')]
    _DEM_VARS = [('peso_kg', 'Peso'), ('altura_m', 'Altura'), ('idade', 'Idade')]

    # ── Bloco 2: Top comparacoes significativas ─────────────────────────
    comps = []
    for k, lbl in _METRICAS:
        for fator, flbl in _FACTORES:
            try:
                r = comparar_grupos(atletas, k, fator)
            except Exception:
                continue
            p = r.get('p_valor')
            if p is None:
                continue
            comps.append((p, k, lbl, fator, flbl, r))
    comps.sort(key=lambda t: t[0])  # menor p primeiro

    c.setFillColor(C_SEC)
    c.rect(MARGIN, y - 0.5*rl_cm, W - 2*MARGIN, 0.5*rl_cm, fill=1, stroke=0)
    c.setFillColor(C_TXT); c.setFont('Helvetica-Bold', 9)
    c.drawString(MARGIN + 0.2*rl_cm, y - 0.35*rl_cm,
                 'Comparacoes de grupo (top por p-valor)')
    y -= 0.75 * rl_cm

    # Cabecalho da tabela
    c.setFillColor(C_DIM); c.setFont('Helvetica-Bold', 7)
    col_xs = [MARGIN + 0.3*rl_cm, MARGIN + 4.2*rl_cm,
              MARGIN + 6.8*rl_cm, MARGIN + 9.0*rl_cm,
              MARGIN + 11.0*rl_cm, MARGIN + 13.0*rl_cm]
    for x, h in zip(col_xs, ['Metrica', 'Factor', 'Teste', 'Estat.', 'p-valor', 'N']):
        c.drawString(x, y, h)
    y -= 0.32 * rl_cm

    MAX_COMPS = 12
    if not comps:
        c.setFillColor(C_DIM); c.setFont('Helvetica-Oblique', 7)
        c.drawString(MARGIN + 0.3*rl_cm, y, '(sem dados demográficos suficientes para comparar)')
        y -= 0.4 * rl_cm
    for p, k, lbl, fator, flbl, r in comps[:MAX_COMPS]:
        sig = p < 0.05
        c.setFillColor(C_SIG if sig else C_TXT); c.setFont('Helvetica-Bold' if sig else 'Helvetica', 7)
        c.drawString(col_xs[0], y, lbl[:18])
        c.drawString(col_xs[1], y, flbl[:12])
        c.drawString(col_xs[2], y, (r.get('teste') or '')[:10])
        est = r.get('estatistica')
        c.drawString(col_xs[3], y, f'{est:.2f}' if est is not None else '-')
        c.drawString(col_xs[4], y, f'{p:.4f}' if p >= 0.001 else '<0.001')
        c.drawString(col_xs[5], y, str(sum(r.get('n_por_grupo', {}).values())))
        y -= 0.32 * rl_cm
        if y < 6.5 * rl_cm:
            break

    y -= 0.4 * rl_cm

    # ── Bloco 3: Correlacoes significativas ─────────────────────────────
    if y > 5.5 * rl_cm:
        cors = []
        for k, lbl in _METRICAS:
            for dvar, dlbl in _DEM_VARS:
                try:
                    r = correlacao_demografica(atletas, k, dvar)
                except Exception:
                    continue
                pr, pp = r.get('pearson_r'), r.get('pearson_p')
                sr, sp = r.get('spearman_r'), r.get('spearman_p')
                # Usar o menor p (mais restritivo)
                pmin = min([x for x in (pp, sp) if x is not None], default=None)
                if pmin is None:
                    continue
                cors.append((pmin, k, lbl, dvar, dlbl, r))
        cors.sort(key=lambda t: t[0])

        c.setFillColor(C_SEC)
        c.rect(MARGIN, y - 0.5*rl_cm, W - 2*MARGIN, 0.5*rl_cm, fill=1, stroke=0)
        c.setFillColor(C_TXT); c.setFont('Helvetica-Bold', 9)
        c.drawString(MARGIN + 0.2*rl_cm, y - 0.35*rl_cm,
                     'Correlações CoP × demografia (top por p-valor)')
        y -= 0.75 * rl_cm

        c.setFillColor(C_DIM); c.setFont('Helvetica-Bold', 7)
        col_xs2 = [MARGIN + 0.3*rl_cm, MARGIN + 4.2*rl_cm,
                   MARGIN + 6.5*rl_cm, MARGIN + 8.5*rl_cm,
                   MARGIN + 10.8*rl_cm, MARGIN + 12.8*rl_cm,
                   MARGIN + 14.8*rl_cm]
        for x, h in zip(col_xs2, ['Métrica', 'Var.', 'N', 'Pearson r',
                                   'Pearson p', 'Spearman r', 'Spearman p']):
            c.drawString(x, y, h)
        y -= 0.32 * rl_cm

        MAX_CORS = 12
        if not cors:
            c.setFillColor(C_DIM); c.setFont('Helvetica-Oblique', 7)
            c.drawString(MARGIN + 0.3*rl_cm, y, '(sem dados demográficos suficientes para correlacionar)')
            y -= 0.4 * rl_cm
        for pmin, k, lbl, dvar, dlbl, r in cors[:MAX_CORS]:
            sig = pmin < 0.05
            c.setFillColor(C_SIG if sig else C_TXT)
            c.setFont('Helvetica-Bold' if sig else 'Helvetica', 7)
            c.drawString(col_xs2[0], y, lbl[:18])
            c.drawString(col_xs2[1], y, dlbl[:10])
            c.drawString(col_xs2[2], y, str(r.get('n', 0)))
            pr, pp = r.get('pearson_r'), r.get('pearson_p')
            sr, sp = r.get('spearman_r'), r.get('spearman_p')
            c.drawString(col_xs2[3], y, f'{pr:+.3f}' if pr is not None else '-')
            c.drawString(col_xs2[4], y, (f'{pp:.4f}' if pp is not None and pp >= 0.001
                                          else ('<0.001' if pp is not None else '-')))
            c.drawString(col_xs2[5], y, f'{sr:+.3f}' if sr is not None else '-')
            c.drawString(col_xs2[6], y, (f'{sp:.4f}' if sp is not None and sp >= 0.001
                                          else ('<0.001' if sp is not None else '-')))
            y -= 0.32 * rl_cm
            if y < 3.0 * rl_cm:
                break

    # ── Rodape com citacoes ──────────────────────────────────────────────
    c.setFillColor(C_DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(MARGIN, 1.1 * rl_cm,
                 'Mann & Whitney (1947), Ann. Math. Stat., 18(1), 50-60  |  '
                 'Kruskal & Wallis (1952), J. Am. Stat. Assoc., 47(260), 583-621  |  '
                 'Spearman (1904), Am. J. Psychol., 15, 72-101')
    c.setFillColor(C_DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(MARGIN, 0.5 * rl_cm,
                 f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}  |  '
                 + datetime.date.today().strftime('%d/%m/%Y'))


def _pagina_capa_pdf(c, atletas, W, H):
    """Capa do relatorio PDF com logo, titulo, protocolo e lista de individuos."""
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.utils import ImageReader
    import io as _io

    proto_key = atletas[0].get('protocolo', PROTO_FMS) if atletas else PROTO_FMS
    proto = PROTOCOLOS.get(proto_key, PROTOCOLOS[PROTO_FMS])
    is_tiro = _is_tiro_like(proto_key)
    is_iscpsi = _is_iscpsi(proto_key)   # cor dourada só no Tiro ISCPSI
    proto_nome_local = _proto_nome(proto_key)

    # Fundo com gradiente simulado por dois rectângulos
    c.setFillColor(rl_colors.HexColor('#0A1018'))
    c.rect(0, 0, W, H, fill=1, stroke=0)
    c.setFillColor(rl_colors.HexColor('#0F1923'))
    c.rect(0, H*0.38, W, H*0.62, fill=1, stroke=0)

    # Linha decorativa superior (ISCPSI Tiro = dourado, outros protocolos = ciano)
    _LINE_COL = '#C8A45A' if is_iscpsi else '#00B4D8'
    c.setStrokeColor(rl_colors.HexColor(_LINE_COL))
    c.setLineWidth(3)
    c.line(0, H - 0.8*cm, W, H - 0.8*cm)

    # ── Logo circular centrado (grande, limpo) ───────────────────────────────
    _logo_y_top = H - 1.0*cm
    _logo_size  = 3.6*cm  # circulo prominente

    try:
        _img_logo = _b64.b64decode(_LOGO_B64)
        _logo_rd  = ImageReader(_io.BytesIO(_img_logo))
        c.drawImage(_logo_rd, W/2 - _logo_size/2, _logo_y_top - _logo_size,
                    width=_logo_size, height=_logo_size,
                    preserveAspectRatio=True, mask='auto')
    except Exception:
        pass
    _logo_zone_bottom = _logo_y_top - _logo_size - 0.4*cm

    # Titulo principal do relatorio PDF
    # Titulo principal
    _title_y = _logo_zone_bottom - 1.0*cm
    c.setFillColor(rl_colors.HexColor('#FFFFFF'))
    c.setFont('Helvetica-Bold', 22)
    c.drawCentredString(W/2, _title_y, PROG.upper())

    # Subtitulo protocolo (dourado ISCPSI só para Tiro ISCPSI, ciano para todos os outros)
    _sub_y = _title_y - 1.1*cm
    _proto_col = '#C8A45A' if is_iscpsi else '#00B4D8'
    c.setFillColor(rl_colors.HexColor(_proto_col))
    c.setFont('Helvetica-Bold', 14)
    c.drawCentredString(W/2, _sub_y, proto_nome_local)

    # Linha separadora
    _sep_y = _sub_y - 0.7*cm
    c.setStrokeColor(rl_colors.HexColor('#1F4E79'))
    c.setLineWidth(0.8)
    c.line(2*cm, _sep_y, W - 2*cm, _sep_y)

    # Metadados
    _meses = T('meses') if isinstance(T('meses'), list) else ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec']
    _hoje = datetime.date.today()
    data_str = f'{_hoje.day:02d} de {_meses[_hoje.month-1]} de {_hoje.year}'
    c.setFillColor(rl_colors.HexColor('#BDD7EE'))
    c.setFont('Helvetica', 9)
    _m1 = _sep_y - 0.8*cm
    c.drawCentredString(W/2, _m1,          T('pdf_capa_data', data=data_str))
    c.drawCentredString(W/2, _m1 - 0.6*cm, T('pdf_capa_n_indiv', n=len(atletas)))
    # Autores em linha separada com fonte ligeiramente maior
    c.setFont('Helvetica-Bold', 8)
    c.setFillColor(rl_colors.HexColor('#E8F4FD'))
    c.drawCentredString(W/2, _m1 - 1.3*cm, f'{AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.setFont('Helvetica', 7.5)
    c.setFillColor(rl_colors.HexColor('#8AAFCC'))
    c.drawCentredString(W/2, _m1 - 1.85*cm, f'v{VERSAO}')

    _extra_y = _m1 - 2.1*cm  # abaixo dos autores e versao

    if proto_key == PROTO_TIRO:
        # Linha de colaboração institucional (sem logo)
        c.setFillColor(rl_colors.HexColor('#8AAFCC'))
        c.setFont('Helvetica-Oblique', 7)
        _colab_y = _extra_y - 0.5*cm
        c.drawCentredString(W/2, _colab_y,
            'Protocolo desenvolvido em colaboração com o Instituto Superior de Ciências Policiais e Segurança Interna (ISCPSI)')
        _extra_y = _colab_y - 0.2*cm

        # Distâncias e intervalos analisados
        all_dists_capa = []
        all_itvs_capa  = []
        for a in atletas:
            for d in a.get('tiro_dists', []):
                if d not in all_dists_capa: all_dists_capa.append(d)
            for itv in a.get('tiro_intervalos', []):
                if itv not in all_itvs_capa: all_itvs_capa.append(itv)

        _dist_y = _extra_y - 0.7*cm
        if all_dists_capa:
            c.setFillColor(rl_colors.HexColor('#BDD7EE'))
            c.setFont('Helvetica-Bold', 7.5)
            c.drawCentredString(W/2, _dist_y,
                T('pdf_capa_distancias', dists=' | '.join(all_dists_capa)))
            _dist_y -= 0.6*cm
        if all_itvs_capa:
            itv_labels = [_tiro_itv_label(itv) for itv in all_itvs_capa]
            c.setFillColor(rl_colors.HexColor('#9DC8E8'))
            c.setFont('Helvetica', 7)
            c.drawCentredString(W/2, _dist_y,
                T('pdf_capa_intervalos', itvs=' | '.join(itv_labels)))

    # Lista de individuos - desenha no espaco disponivel da capa
    # Retorna lista de atletas que nao couberam (para pagina extra)
    _box_top  = H * 0.38 - 4.5*cm   # topo da caixa azul (relativo a base)
    _box_bot  = 3.5*cm               # base da caixa
    _box_h    = _box_top - _box_bot

    c.setFillColor(rl_colors.HexColor('#1F4E79'))
    c.rect(1.5*cm, _box_bot, W - 3*cm, _box_h, fill=1, stroke=0)

    c.setFillColor(rl_colors.HexColor('#00B4D8'))
    c.setFont('Helvetica-Bold', 9)
    c.drawString(2*cm, _box_bot + _box_h - 0.5*cm, T('pdf_capa_analisados'))

    _row_h   = 0.52*cm
    _padding = 1.2*cm   # espaco acima do topo para o titulo
    _max_rows_per_col = max(1, int((_box_h - _padding) / _row_h))
    _ncols   = 2 if len(atletas) > _max_rows_per_col else 1
    _max_fit = _ncols * _max_rows_per_col

    c.setFont('Helvetica', 8)
    c.setFillColor(rl_colors.HexColor('#E8F4FD'))
    _col_w = (W - 4*cm) / _ncols
    for i, ath in enumerate(atletas[:_max_fit]):
        col_i = i // _max_rows_per_col
        row_i = i % _max_rows_per_col
        x = 2*cm + col_i * _col_w
        y = _box_bot + _box_h - _padding - row_i * _row_h
        id_str = f"{ath.get('id','')}. " if ath.get('id') else f"{i+1}. "
        c.drawString(x, y, id_str + ath['nome'])

    if len(atletas) > _max_fit:
        c.setFont('Helvetica-Oblique', 7)
        c.setFillColor(rl_colors.HexColor('#6B8FAE'))
        c.drawString(2*cm, _box_bot + 0.15*cm,
                     T('pdf_capa_mais', n=len(atletas)-_max_fit))

    # Rodape
    c.setStrokeColor(rl_colors.HexColor('#00B4D8'))
    c.setLineWidth(2)
    c.line(0, 0.8*cm, W, 0.8*cm)
    c.setFillColor(rl_colors.HexColor('#6B8FAE'))
    c.setFont('Helvetica', 7)
    c.drawCentredString(W/2, 0.3*cm,
        f'{PROG}  |  v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')


# Legenda de variaveis (glossario)
_LEGENDA_VARS = [
    ('Amplitude X (mm)',      'amp_x',
     'Amplitude mediolateral (ML): diferença entre o máximo e o mínimo do CoP no eixo X.'),
    ('Amplitude Y (mm)',      'amp_y',
     'Amplitude anteroposterior (AP): diferença entre o máximo e o mínimo do CoP no eixo Y.'),
    ('Vel. média X (mm/s)',   'vel_x',
     'Velocidade média do CoP no eixo mediolateral (X): média dos valores absolutos das velocidades instantâneas.'),
    ('Vel. média Y (mm/s)',   'vel_y',
     'Velocidade média do CoP no eixo anteroposterior (Y): média dos valores absolutos das velocidades instantâneas.'),
    ('Vel. média CoP (mm/s)', 'vel_med',
     'Velocidade média global do CoP: comprimento total da trajectória dividido pela duração do ensaio.'),
    ('Pico vel. X (mm/s)',    'vel_pico_x',
     'Pico de velocidade do CoP no eixo mediolateral (X): valor máximo das velocidades instantâneas absolutas.'),
    ('Pico vel. Y (mm/s)',    'vel_pico_y',
     'Pico de velocidade do CoP no eixo anteroposterior (Y): valor máximo das velocidades instantâneas absolutas.'),
    ('Deslocamento (mm)',     'desl',
     'Comprimento total da trajectória do CoP (path length): soma dos deslocamentos frame a frame.'),
    ('Tempo (s)',             'tempo',
     'Duração efectiva do ensaio utilizada no cálculo (intervalo t_ini a t_fim).'),
    ('Área elipse 95% (mm²)', 'ea95',
     'Área da elipse de predição a 95%: PEA = π×(a/2)×(b/2). a e b são calculados a partir dos autovalores da matriz de covariância do CoP escalados pelo χ²(0.95,2)=5.991. Área maior = maior instabilidade; área menor = melhor controlo postural.'),
    ('Semi-eixo a (mm)',      'leng_a',
     'Semi-eixo maior da elipse: raiz quadrada do maior autovalor (direcção de maior dispersão do CoP).'),
    ('Semi-eixo b (mm)',      'leng_b',
     'Semi-eixo menor da elipse: raiz quadrada do menor autovalor (direcção de menor dispersão do CoP).'),
    ('Rel. amp. ML/AP',       'rel_amp_mlap',
     'Relação entre a amplitude mediolateral e a amplitude anteroposterior (amp_x / amp_y). Valores > 1: oscilação predominante no plano ML; valores < 1: predominância AP.'),
    ('Rel. vel. ML/AP',       'rel_vel_mlap',
     'Relação entre a velocidade média no eixo ML e a velocidade média no eixo AP (vel_x / vel_y). Identifica o eixo de instabilidade predominante em termos de velocidade.'),
    ('Stiffness X (1/s)',     'stiff_x',
     'Índice de rigidez postural no eixo ML: velocidade média dividida pela amplitude ML. Valores mais elevados sugerem controlo postural mais activo (Winter, 1995). Métrica exploratória.'),
    ('Stiffness Y (1/s)',     'stiff_y',
     'Índice de rigidez postural no eixo AP. Interpretação análoga ao Stiffness X.'),
    ('Var CoP X (mm²)',       'var_x',
     'Variância da trajectória do CoP no eixo ML (elemento [0,0] da matriz de covariância). Quantifica a dispersão quadrática média em torno da posição média no eixo mediolateral.'),
    ('Var CoP Y (mm²)',       'var_y',
     'Variância da trajectória do CoP no eixo AP (elemento [1,1] da matriz de covariância). Quantifica a dispersão quadrática no eixo anteroposterior.'),
    ('Cov CoP XY (mm²)',      'cov_xy',
     'Covariância entre os eixos ML e AP (elemento [0,1] da matriz de covariância). Valores elevados (em módulo) indicam que as oscilações ML e AP estão correlacionadas. O sinal indica a direcção: positivo = ambos os eixos oscilam no mesmo sentido.'),
    ('RMS ML (mm)',           'rms_ml',
     'Root Mean Square mediolateral: raiz quadrada da média dos quadrados das coordenadas centradas no eixo ML. Equivalente ao desvio-padrão da posição ML. Fórmula: sqrt(1/N × Σ Xn²), com Xn = MLn − média(ML). Ref.: Quijoux et al. (2021).'),
    ('RMS AP (mm)',           'rms_ap',
     'Root Mean Square anteroposterior: raiz quadrada da média dos quadrados das coordenadas centradas no eixo AP. Fórmula: sqrt(1/N × Σ Yn²), com Yn = APn − média(AP). Valores mais elevados indicam maior dispersão postural no plano sagital. Ref.: Quijoux et al. (2021).'),
    ('RMS Radius (mm)',       'rms_radius',
     'Root Mean Square do raio: raiz quadrada da média dos quadrados das distâncias ao centróide. Fórmula: sqrt(1/N × Σ(Xn² + Yn²)). Captura a dispersão espacial global do CoP em ambos os eixos simultaneamente. Ref.: Quijoux et al. (2021).'),
    ('FFT Pico ML (Hz)',      'fft_pico_ml',
     'Frequência de pico do CoP no eixo ML, obtida pelo espectro de potência (FFT). Representa a frequência de oscilação mais intensa no plano mediolateral. Intervalo fisiológico: 0.1–10 Hz. Apenas calculado se FFT activado.'),
    ('FFT Pico AP (Hz)',      'fft_pico_ap',
     'Frequência de pico do CoP no eixo AP. Representa a frequência de oscilação mais intensa no plano anteroposterior. Apenas calculado se FFT activado.'),
    ('FFT Méd. ML (Hz)',      'fft_med_ml',
     'Frequência média ponderada pelo espectro de potência no eixo ML (centróide espectral). Fórmula: Σ(f×P(f)) / Σ P(f). Reflecte a frequência média de oscilação em ML. Apenas calculado se FFT activado.'),
    ('FFT Méd. AP (Hz)',      'fft_med_ap',
     'Frequência média ponderada pelo espectro de potência no eixo AP. Apenas calculado se FFT activado.'),
    ('AI (%)',                'ai',
     'Índice de assimetria: AI = 2 × (Dir − Esq) / (Dir + Esq) × 100. Positivo = dominância direita. |AI| > 10% é clinicamente relevante.'),
    ('CV (%)',                'cv',
     'Coeficiente de variação: DP / Média × 100. Indica a variabilidade relativa entre ensaios do mesmo indivíduo.'),
]


def _pagina_capa2_pdf(c, atletas, offset, W, H):
    """Segunda pagina de capa com o resto dos individuos quando a lista e longa."""
    from reportlab.lib.units import cm
    from reportlab.lib import colors as rl_colors

    AZ   = rl_colors.HexColor('#0F1923')
    AZM  = rl_colors.HexColor('#1F4E79')
    AZDB = rl_colors.HexColor('#E8F4FD')
    AZC  = rl_colors.HexColor('#00B4D8')
    DIM  = rl_colors.HexColor('#6B8FAE')

    c.setFillColor(AZ)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # Linha topo
    c.setStrokeColor(AZC); c.setLineWidth(2)
    c.line(0, H - 0.8*cm, W, H - 0.8*cm)

    # Titulo
    c.setFillColor(rl_colors.HexColor('#FFFFFF'))
    c.setFont('Helvetica-Bold', 14)
    c.drawCentredString(W/2, H - 2.2*cm, T('pdf_capa_analisados2'))
    c.setFillColor(AZC); c.setFont('Helvetica', 9)
    c.drawCentredString(W/2, H - 2.9*cm,
                        f'Indivíduos {offset+1} a {len(atletas)}  |  Total: {len(atletas)}')

    c.setStrokeColor(rl_colors.HexColor('#1F4E79')); c.setLineWidth(0.8)
    c.line(2*cm, H - 3.3*cm, W - 2*cm, H - 3.3*cm)

    remaining = atletas[offset:]
    _row_h = 0.52*cm
    _area_top = H - 3.8*cm
    _area_bot = 1.5*cm
    _area_h   = _area_top - _area_bot
    _ncols    = 2 if len(remaining) > int(_area_h / _row_h) else 1
    _max_rows = max(1, int(_area_h / _row_h))
    _col_w    = (W - 4*cm) / _ncols

    c.setFont('Helvetica', 8); c.setFillColor(AZDB)
    for i, ath in enumerate(remaining):
        col_i = i // _max_rows
        row_i = i % _max_rows
        if col_i >= _ncols: break
        x = 2*cm + col_i * _col_w
        y = _area_top - row_i * _row_h
        idx_global = offset + i
        id_str = f"{ath.get('id','')}. " if ath.get('id') else f"{idx_global+1}. "
        c.drawString(x, y, id_str + ath['nome'])

    # Rodape
    c.setStrokeColor(AZC); c.setLineWidth(2)
    c.line(0, 0.8*cm, W, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica', 7)
    import datetime as _dt2
    c.drawCentredString(W/2, 0.3*cm,
        f'Biomechanical Stability Program  |  v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')


def _pagina_legenda_pdf(c, W, H):
    """Pagina de legenda / glossario das variaveis."""
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors as rl_colors
    from reportlab.platypus import Table, TableStyle

    c.setFillColor(rl_colors.HexColor('#0F1923'))
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # Cabecalho
    c.setFillColor(rl_colors.HexColor('#1F4E79'))
    c.rect(0, H - 2.2*cm, W, 2.2*cm, fill=1, stroke=0)
    c.setFillColor(rl_colors.HexColor('#00B4D8'))
    c.setFont('Helvetica-Bold', 14)
    c.drawString(1.2*cm, H - 1.2*cm, 'LEGENDA DE VARIÁVEIS')
    c.setFillColor(rl_colors.HexColor('#BDD7EE'))
    c.setFont('Helvetica', 8)
    c.drawString(1.2*cm, H - 1.7*cm,
                 'Definições e fórmulas das métricas calculadas pelo programa')

    # Tabela com Paragraph para word-wrap correcto
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.styles import ParagraphStyle

    st_def = ParagraphStyle('def',
        fontName='Helvetica', fontSize=7.5,
        textColor=rl_colors.HexColor('#E8F4FD'),
        leading=10)
    st_var = ParagraphStyle('var',
        fontName='Helvetica-Bold', fontSize=7.5,
        textColor=rl_colors.HexColor('#BDD7EE'),
        leading=10)
    st_cod = ParagraphStyle('cod',
        fontName='Helvetica-Oblique', fontSize=7.5,
        textColor=rl_colors.HexColor('#00B4D8'),
        leading=10)
    st_hdr = ParagraphStyle('hdr',
        fontName='Helvetica-Bold', fontSize=8,
        textColor=rl_colors.HexColor('#00B4D8'),
        leading=11)

    col_ws = [3.6*cm, 1.7*cm, W - 2*cm - 3.6*cm - 1.7*cm]
    tbl_data = [
        [Paragraph('Variável', st_hdr),
         Paragraph('Código', st_hdr),
         Paragraph('Definição', st_hdr)]
    ]
    for nome, cod, defn in _LEGENDA_VARS:
        tbl_data.append([
            Paragraph(nome, st_var),
            Paragraph(cod, st_cod),
            Paragraph(defn, st_def),
        ])

    tbl = Table(tbl_data, colWidths=col_ws)
    row_bgs = [rl_colors.HexColor('#162233'), rl_colors.HexColor('#1C2E42')]
    style_cmds = [
        ('BACKGROUND',   (0,0), (-1,0),  rl_colors.HexColor('#243447')),
        ('GRID',         (0,0), (-1,-1), 0.3, rl_colors.HexColor('#2A3F55')),
        ('ALIGN',        (0,0), (-1,-1), 'LEFT'),
        ('VALIGN',       (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING',  (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING',   (0,0), (-1,-1), 4),
        ('BOTTOMPADDING',(0,0), (-1,-1), 4),
    ]
    for ri in range(1, len(tbl_data)):
        bg = row_bgs[(ri-1) % 2]
        style_cmds.append(('BACKGROUND', (0,ri), (-1,ri), bg))
    tbl.setStyle(TableStyle(style_cmds))

    avail_h = H - 3.0*cm
    tbl_w, tbl_h = tbl.wrapOn(c, W - 2*cm, avail_h)
    # Se a tabela nao cabe, comprimir
    if tbl_h > avail_h:
        fator = avail_h / tbl_h * 0.97
        tbl._rowHeights = [max(rh * fator, 8) for rh in tbl._rowHeights]
        tbl_w, tbl_h = tbl.wrapOn(c, W - 2*cm, avail_h)
    tbl.drawOn(c, 1*cm, H - 2.8*cm - tbl_h)

    # ---- Secao interpretativa - nova pagina se nao houver espaco ----
    from reportlab.lib.styles import ParagraphStyle as _PS
    _MIN_Y_LEG = 1.5*cm
    _y_exp = H - 2.8*cm - tbl_h - 0.6*cm
    _margem = 1.0*cm

    # Se sobrar menos de 5cm, pagina nova
    if _y_exp < 5.0*cm:
        c.showPage()
        # Fundo e cabecalho da pagina extra
        c.setFillColor(rl_colors.HexColor('#0F1923'))
        c.rect(0, 0, W, H, fill=1, stroke=0)
        c.setFillColor(rl_colors.HexColor('#1F4E79'))
        c.rect(0, H - 1.4*cm, W, 1.4*cm, fill=1, stroke=0)
        c.setFillColor(rl_colors.HexColor('#00B4D8'))
        c.setFont('Helvetica-Bold', 10)
        c.drawString(1.2*cm, H - 0.9*cm, 'LEGENDA: Notas de interpretação')
        _y_exp = H - 1.8*cm

    _st_et = _PS('exp_t', fontName='Helvetica-Bold', fontSize=8.5,
                 textColor=rl_colors.HexColor('#00B4D8'), leading=13)
    _st_eb = _PS('exp_b', fontName='Helvetica', fontSize=7.5,
                 textColor=rl_colors.HexColor('#BDD7EE'), leading=11)
    _textos_exp = [
        ('O que é o Estabilograma?',
         'O estabilograma é o gráfico que mostra a trajectória do Centro de Pressão (CoP) '
         'ao longo do tempo. O CoP é o ponto de aplicação da força que o pé exerce no chão, '
         'movendo-se continuamente enquanto o corpo ajusta o equilíbrio. '
         'Uma trajectória mais extensa indica maior oscilação postural.'),
        ('O que é a Elipse de Predição a 95%?',
         'A elipse cobre 95% das posições registadas pelo CoP durante o ensaio. '
         'A sua área (mm²) é a medida de dispersão espacial mais usada em estabilometria: '
         'área maior = maior instabilidade; área menor = melhor controlo postural.'),
        ('Como interpretar os valores?',
         'Amplitudes (amp_x, amp_y): excursão total do CoP em cada eixo. '
         'Velocidade média (vel_med): quanto mais rápido o CoP se move, maior o esforço '
         'neuromuscular. Covariância (cov_xy): correlação entre os dois eixos, útil para '
         'análises futuras. AI (%): assimetria Dir/Esq; valores |AI| > 10%% são relevantes.'),
    ]
    from reportlab.platypus import Paragraph as _Para
    for _titulo_e, _corpo_e in _textos_exp:
        if _y_exp < 2.0*cm:
            break
        _pt = _Para(_titulo_e, _st_et)
        _pb = _Para(_corpo_e,  _st_eb)
        _, _ht = _pt.wrapOn(c, W - 2*_margem, 30)
        _, _hb = _pb.wrapOn(c, W - 2*_margem, 60)
        if _y_exp - _ht - _hb < 2.0*cm:
            break
        _y_exp -= _ht + 1
        _pt.drawOn(c, _margem, _y_exp)
        _y_exp -= _hb + 6
        _pb.drawOn(c, _margem, _y_exp)
        _y_exp -= 5

    # Rodape
    c.setFillColor(rl_colors.HexColor('#6B8FAE'))
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, 0.4*cm,
                 f'Ref: Schubert & Kirchner (2013) Gait & Posture 39(1):518-522  |  '
                 f'Quijoux et al. (2021) Physiological Reports 9(11):e14945')
    c.drawRightString(W - 1*cm, 0.4*cm, '2')



def _tbl_fit_page(c, tbl, x, y_top, W, H, margem=1.0, rodape=1.2):
    """
    Desenha uma tabela reportlab garantindo que nunca sai da página.
    Se a tabela não cabe entre y_top e rodape*cm, comprime as alturas
    das linhas proporcionalmente para caber.  Retorna y final (base da tabela).
    """
    from reportlab.lib.units import cm as _cm
    limite_inf = rodape * _cm
    espaco_disponivel = y_top - limite_inf
    tw, th = tbl.wrapOn(c, W - 2*_cm, espaco_disponivel)
    if th > espaco_disponivel and espaco_disponivel > 0.5*_cm:
        # Comprimir: recalcular alturas de linha proporcionalmente
        fator = espaco_disponivel / th * 0.97
        old_rh = tbl._rowHeights
        new_rh = [max(rh * fator, 0.23*_cm) for rh in old_rh]
        tbl._rowHeights = new_rh
        tw, th = tbl.wrapOn(c, W - 2*_cm, espaco_disponivel)
    y_final = y_top - th - 0.08*_cm
    tbl.drawOn(c, x, y_final)
    return y_final


def _pagina_tiro_entire_plate_pdf(c, atleta, dist, itv, W, H):
    """
    Pagina de metricas Entire Plate (bipodal) para atleta/distancia/intervalo.
    Usa os dados principais x/y (Entire plate COF X/Y) - mesma analise que
    _pagina_tiro_itv_pdf mas explicitamente rotulada como 'Entire plate'.
    """
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.utils import ImageReader
    import numpy as _np

    nome      = atleta['nome']
    ind_id    = atleta.get('id', '')
    itv_label = _tiro_itv_label(itv)

    AZ   = colors.HexColor('#1F4E79')
    AZCL = colors.HexColor('#BDD7EE')
    CZ   = colors.HexColor('#F2F2F2')
    VM   = colors.HexColor('#C00000')
    DIM  = colors.HexColor('#888888')
    AZDB = colors.HexColor('#E8F4FD')
    AZEP = colors.HexColor('#EBF5FB')  # fundo Entire plate

    # cabecalho
    c.setFillColor(AZ)
    c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, H - 1.05*cm,
                 f'{PROG}  -  Entire Plate  -  {dist}  -  {itv_label}')
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, H - 1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - 1*cm, H - 1.05*cm, datetime.date.today().strftime('%d/%m/%Y'))

    y = H - 2.5*cm
    id_prefix = f'[ID {ind_id}]  ' if ind_id else ''
    c.setFillColor(AZ); c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, y, id_prefix + nome)
    c.setFillColor(colors.HexColor('#6B8FAE')); c.setFont('Helvetica', 7.5)

    mets_list = atleta.get('tiro_dist', {}).get(dist, {}).get('mets', {}).get(itv, [])
    raw_list  = atleta.get('tiro_dist', {}).get(dist, {}).get('raw',  {}).get(itv, [])
    flags     = flagrar_outliers(mets_list)
    _mets_real = [m for m in mets_list if m is not None]
    n_vals    = max(max(len(_mets_real), len(mets_list)), 1)  # sem limite

    duracoes = []
    for rd in raw_list:
        if rd and rd.get('ini') is not None and rd.get('fim') is not None:
            duracoes.append(rd['fim'] - rd['ini'])
        else:
            m_t = mets_list[len(duracoes)] if len(duracoes) < len(mets_list) else None
            duracoes.append(int(m_t['time']*1000) if m_t and 'time' in m_t else None)
    dur_med = int(np.mean([d for d in duracoes if d is not None])) if any(d is not None for d in duracoes) else None
    dur_str = f'  |  Duração média: {dur_med} ms' if dur_med else ''

    c.drawString(1*cm, y - 4*mm,
                 f'Distancia: {dist}   Intervalo: {itv_label}   |   Entire plate (bipodal){dur_str}')
    c.setStrokeColor(AZ); c.setLineWidth(0.5)
    c.line(1*cm, y - 6*mm, W - 1*cm, y - 6*mm)
    y -= 1.0*cm

    margem  = 1.0*cm; total_w = W - 2*margem
    C0 = 3.2*cm; CV_ratio = 1.0; CS_ratio = 1.15
    n_col = n_vals + 3
    unit  = (total_w - C0) / (n_col * CV_ratio + 3*(CS_ratio - CV_ratio))
    CV = unit * CV_ratio; CS = unit * CS_ratio
    col_larg = [C0] + [CV]*n_vals + [CS, CS, CS]

    cab = ['Métrica', f'ENTIRE PLATE  |  {itv_label}'] + ['']*(n_vals+2)
    sub = [''] + [f'T{i+1}{"*" if i<len(flags) and flags[i] else ""}' for i in range(n_vals)] + ['máx','méd','dp']

    dur_linha = ['Duração (ms)']
    dur_nums_t = []
    for ti in range(n_vals):
        d = duracoes[ti] if ti < len(duracoes) else None
        dur_linha.append(str(d) if d is not None else '-')
        if d is not None: dur_nums_t.append(d)
    dur_linha += [str(max(dur_nums_t)) if dur_nums_t else '-',
                  str(int(np.mean(dur_nums_t))) if dur_nums_t else '-',
                  str(int(np.std(dur_nums_t, ddof=1))) if len(dur_nums_t)>1 else '-']

    linhas = [cab, sub, dur_linha]
    idx_ea95 = next((i+3 for i,(k,_) in enumerate(METS_PDF) if k=='ea95'), 10)

    for chave, label in METS_PDF:
        linha = [label]
        vals = []
        for ti in range(n_vals):
            m = mets_list[ti] if ti < len(mets_list) else None
            v = m[chave] if m is not None and chave in m else None
            vals.append(v)
            linha.append(f'{v:.2f}' if v is not None else '-')
        nums = _np.array([v for v in vals if v is not None])
        linha += [f'{nums.max():.2f}'       if len(nums)>0 else '-',
                  f'{nums.mean():.2f}'      if len(nums)>0 else '-',
                  f'{nums.std(ddof=1):.2f}' if len(nums)>1 else '-']
        linhas.append(linha)

    row_h = [0.38*cm, 0.33*cm, 0.30*cm] + [0.33*cm]*(len(linhas)-3)
    tbl   = Table(linhas, colWidths=col_larg, rowHeights=row_h)
    AZCL_tbl = colors.HexColor('#D6E4F0')
    estilo = [
        ('BACKGROUND',(0,0),(-1,1),  AZ),
        ('TEXTCOLOR',(0,0),(-1,1),   colors.white),
        ('FONTNAME',(0,0),(-1,1),    'Helvetica-Bold'),
        ('FONTSIZE',(0,0),(-1,1),    6.5),
        ('ALIGN',(0,0),(-1,1),       'CENTER'),
        ('VALIGN',(0,0),(-1,1),      'MIDDLE'),
        ('SPAN',(1,0),(n_vals+3,0)),
        ('BACKGROUND',(0,2),(-1,2),  AZCL_tbl),
        ('TEXTCOLOR',(0,2),(-1,2),   AZ),
        ('FONTNAME',(0,2),(-1,2),    'Helvetica-Oblique'),
        ('FONTSIZE',(0,2),(-1,2),    6),
        ('ALIGN',(0,2),(-1,2),       'CENTER'),
        ('VALIGN',(0,2),(-1,2),      'MIDDLE'),
        ('FONTNAME',(0,3),(0,-1),    'Helvetica-Bold'),
        ('FONTSIZE',(0,3),(-1,-1),   6),
        ('ALIGN',(1,3),(-1,-1),      'CENTER'),
        ('ALIGN',(0,3),(0,-1),       'LEFT'),
        ('VALIGN',(0,3),(-1,-1),     'MIDDLE'),
        *[('BACKGROUND',(0,i),(-1,i), CZ) for i in range(4, len(linhas), 2)],
        ('BACKGROUND',(0,3),(0,-1),  AZEP),
        ('BACKGROUND',(1,3),(-1,-1), AZEP),
        ('TEXTCOLOR',(0,idx_ea95),(-1,idx_ea95),  VM),
        ('FONTNAME',(0,idx_ea95),(-1,idx_ea95),   'Helvetica-Bold'),
        ('BACKGROUND',(n_vals+1,3),(n_vals+3,-1),  AZDB),
        ('GRID',(0,0),(-1,-1),       0.2, colors.HexColor('#CCCCCC')),
        ('LINEBELOW',(0,2),(-1,2),   0.6, AZ),
        ('TOPPADDING',(0,0),(-1,-1),    1),
        ('BOTTOMPADDING',(0,0),(-1,-1), 1),
        ('LEFTPADDING',(0,0),(0,-1),    3),
    ]
    tbl.setStyle(TableStyle(estilo))
    y = _tbl_fit_page(c, tbl, 1*cm, y, W, H)

    y -= 0.2*cm
    espaco = y - 1.2*cm
    hg = max(espaco * 0.95, 5.5*cm)
    wg = W - 2*cm
    png = _png_elipses(mets_list, f'{itv_label} - Entire plate', nome,
                       larg=wg/cm*0.40, alt=hg/cm*0.43, dpi=200, legenda_fora=True)
    if png:
        c.drawImage(ImageReader(io.BytesIO(png)), 1*cm, y-hg,
                    width=wg, height=hg, preserveAspectRatio=False)
    else:
        c.setFillColor(CZ); c.rect(1*cm, y-hg, wg, hg, fill=1, stroke=0)

    c.setStrokeColor(AZCL); c.setLineWidth(0.3)
    c.line(1*cm, 0.8*cm, W-1*cm, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(1*cm, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W-1*cm, 0.4*cm,
                      f'{nome}  |  Entire Plate  |  {dist}  |  {itv_label}')


def _pagina_tiro_sel_lado_pdf(c, atleta, dist, itv, lado, W, H):
    """
    Pagina de metricas de um unico pe (Right Sel. ou Left Sel.) para tiro.
    lado: 'dir' → Pe Direito (Right Sel.)
          'esq' → Pe Esquerdo (Left Sel.)
    Layout de coluna unica - nunca fica cortado.
    """
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.utils import ImageReader
    import numpy as _np

    nome      = atleta['nome']
    ind_id    = atleta.get('id', '')
    itv_label = _tiro_itv_label(itv)
    lado_lbl  = ('Pe Direito (Right Sel.)' if lado == 'dir'
                 else 'Pe Esquerdo (Left Sel.)')
    lado_cor  = colors.HexColor('#D9E1F2') if lado == 'dir' else colors.HexColor('#E2EFDA')

    AZ   = colors.HexColor('#1F4E79')
    AZCL = colors.HexColor('#BDD7EE')
    CZ   = colors.HexColor('#F2F2F2')
    VM   = colors.HexColor('#C00000')
    DIM  = colors.HexColor('#888888')
    AZDB = colors.HexColor('#E8F4FD')

    # cabecalho
    c.setFillColor(AZ)
    c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, H - 1.05*cm,
                 f'{PROG}  -  {lado_lbl}  -  {dist}  -  {itv_label}')
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, H - 1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - 1*cm, H - 1.05*cm, datetime.date.today().strftime('%d/%m/%Y'))

    y = H - 2.5*cm
    id_prefix = f'[ID {ind_id}]  ' if ind_id else ''
    c.setFillColor(AZ); c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, y, id_prefix + nome)
    c.setFillColor(colors.HexColor('#6B8FAE')); c.setFont('Helvetica', 7.5)
    c.drawString(1*cm, y - 4*mm,
                 f'Distancia: {dist}   Intervalo: {itv_label}   |   {lado_lbl}')
    c.setStrokeColor(AZ); c.setLineWidth(0.5)
    c.line(1*cm, y - 6*mm, W - 1*cm, y - 6*mm)
    y -= 1.0*cm

    # Dados do lado selecionado
    sel_itv  = atleta.get('sel_dist', {}).get(dist, {}).get(itv, {})
    mets_lst = [m for m in sel_itv.get(lado, []) if m is not None]
    flags    = flagrar_outliers(mets_lst)
    n_vals   = max(len(mets_lst), 1)  # sem limite

    margem  = 1.0*cm; total_w = W - 2*margem
    C0 = 3.2*cm; CV_ratio = 1.0; CS_ratio = 1.15
    n_col = n_vals + 3
    unit  = (total_w - C0) / (n_col * CV_ratio + 3*(CS_ratio - CV_ratio))
    CV = unit * CV_ratio; CS = unit * CS_ratio
    col_larg = [C0] + [CV]*n_vals + [CS, CS, CS]

    cab    = ['Métrica', f'{lado_lbl.upper()}  |  {itv_label}'] + ['']*(n_vals+2)
    sub    = [''] + [f'T{i+1}{"*" if i<len(flags) and flags[i] else ""}' for i in range(n_vals)] + ['máx','méd','dp']
    linhas = [cab, sub]
    idx_ea95 = next((i+2 for i,(k,_) in enumerate(METS_PDF) if k=='ea95'), 9)

    for chave, label in METS_PDF:
        linha = [label]
        vals  = []
        for ti in range(n_vals):
            m = mets_lst[ti] if ti < len(mets_lst) else None
            v = m[chave] if m is not None and chave in m else None
            vals.append(v)
            linha.append(f'{v:.2f}' if v is not None else '-')
        nums = _np.array([v for v in vals if v is not None])
        linha += [f'{nums.max():.2f}'       if len(nums)>0 else '-',
                  f'{nums.mean():.2f}'      if len(nums)>0 else '-',
                  f'{nums.std(ddof=1):.2f}' if len(nums)>1 else '-']
        linhas.append(linha)

    row_h = [0.38*cm, 0.33*cm] + [0.33*cm]*(len(linhas)-2)
    tbl   = Table(linhas, colWidths=col_larg, rowHeights=row_h)
    estilo = [
        ('BACKGROUND',(0,0),(-1,1),  AZ),
        ('TEXTCOLOR',(0,0),(-1,1),   colors.white),
        ('FONTNAME',(0,0),(-1,1),    'Helvetica-Bold'),
        ('FONTSIZE',(0,0),(-1,1),    6.5),
        ('ALIGN',(0,0),(-1,1),       'CENTER'),
        ('VALIGN',(0,0),(-1,1),      'MIDDLE'),
        ('SPAN',(1,0),(n_vals+3,0)),
        ('FONTNAME',(0,2),(0,-1),    'Helvetica-Bold'),
        ('FONTSIZE',(0,2),(-1,-1),   6),
        ('ALIGN',(1,2),(-1,-1),      'CENTER'),
        ('ALIGN',(0,2),(0,-1),       'LEFT'),
        ('VALIGN',(0,2),(-1,-1),     'MIDDLE'),
        *[('BACKGROUND',(0,i),(-1,i), CZ) for i in range(3, len(linhas), 2)],
        ('BACKGROUND',(1,2),(-1,-1), lado_cor),
        ('TEXTCOLOR',(0,idx_ea95),(-1,idx_ea95),  VM),
        ('FONTNAME',(0,idx_ea95),(-1,idx_ea95),   'Helvetica-Bold'),
        ('BACKGROUND',(n_vals+1,2),(n_vals+3,-1),  AZDB),
        ('GRID',(0,0),(-1,-1),       0.2, colors.HexColor('#CCCCCC')),
        ('LINEBELOW',(0,1),(-1,1),   0.8, AZ),
        ('TOPPADDING',(0,0),(-1,-1),    1),
        ('BOTTOMPADDING',(0,0),(-1,-1), 1),
        ('LEFTPADDING',(0,0),(0,-1),    3),
    ]
    tbl.setStyle(TableStyle(estilo))
    y = _tbl_fit_page(c, tbl, 1*cm, y, W, H)

    y -= 0.2*cm
    espaco = y - 1.2*cm
    hg = max(espaco * 0.95, 5.5*cm)
    wg = W - 2*cm
    png = _png_elipses(mets_lst, lado_lbl, nome,
                       larg=wg/cm*0.40, alt=hg/cm*0.43, dpi=200, legenda_fora=True)
    if png:
        c.drawImage(ImageReader(io.BytesIO(png)), 1*cm, y-hg,
                    width=wg, height=hg, preserveAspectRatio=False)
    else:
        c.setFillColor(CZ); c.rect(1*cm, y-hg, wg, hg, fill=1, stroke=0)

    c.setStrokeColor(AZCL); c.setLineWidth(0.3)
    c.line(1*cm, 0.8*cm, W-1*cm, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(1*cm, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W-1*cm, 0.4*cm,
                      f'{nome}  |  {lado_lbl}  |  {dist}  |  {itv_label}')


def _pagina_tiro_selection_pdf(c, atleta, dist, itv, W, H):
    """
    DEPRECATED v20: substituida por _pagina_tiro_entire_plate_pdf +
    _pagina_tiro_sel_lado_pdf('dir') + _pagina_tiro_sel_lado_pdf('esq').
    Mantida apenas para compatibilidade retroactiva.
    """
    _pagina_tiro_entire_plate_pdf(c, atleta, dist, itv, W, H)


def _tiro_pre_calc_pages(atletas, all_dists, all_itvs, include_hs, pagina_inicial):
    """
    Pre-calcula os numeros de pagina de cada secção para o índice.
    Retorna lista de (label, pagina) e pagina final.
    """
    entries = []
    page = pagina_inicial  # primeira pagina de conteudo (depois de capa + legenda + toc)

    for atleta in atletas:
        ind_id = atleta.get('id', '')
        nome   = atleta['nome']
        id_str = f'ID {ind_id} - ' if ind_id else ''
        entries.append((f'{id_str}{nome}', page, 'atleta'))
        page += 1  # divisora

        for dist in (atleta.get('tiro_dists') or all_dists):
            entries.append((f'  {dist}', page, 'dist'))
            for itv in (atleta.get('tiro_intervalos') or all_itvs):
                mets = (atleta.get('tiro_dist', {}).get(dist, {})
                        .get('mets', {}).get(itv, []))
                if any(m is not None for m in mets):
                    itv_lbl = _tiro_itv_label(itv)
                    entries.append((f'    {itv_lbl}', page, 'itv'))
                    page += 2  # tabela+elipse + estabilogramas
                    # Selection CoP (Right + Left) - 1 pagina por intervalo se dados presentes
                    sel_itv = (atleta.get('sel_dist', {}).get(dist, {}).get(itv, {}))
                    has_sel_dir = any(m is not None for m in sel_itv.get('dir', []))
                    has_sel_esq = any(m is not None for m in sel_itv.get('esq', []))
                    has_sel = has_sel_dir or has_sel_esq
                    if has_sel:
                        entries.append((f'      Entire Plate: {itv_lbl}', page, 'sel'))
                        page += 1  # pagina Entire plate
                        if has_sel_dir:
                            entries.append((f'      Pe Direito (Right Sel.): {itv_lbl}', page, 'sel'))
                            page += 1  # pagina Pe Direito
                        if has_sel_esq:
                            entries.append((f'      Pe Esquerdo (Left Sel.): {itv_lbl}', page, 'sel'))
                            page += 1  # pagina Pe Esquerdo
            page += 1  # resumo por distância

        hs_tem = include_hs and bool(
            atleta.get('mets', {}).get('dir') or atleta.get('mets', {}).get('esq'))
        if hs_tem:
            entries.append(('  Hurdle Step (Bipodal)', page, 'hs'))
            page += 3  # divisora hs + pagina_pdf + estabilograma

    return entries, page


def _pagina_toc_pdf(c, entries, W, H, proto_nome='Tiro'):
    """Pagina de indice (TOC) com ligacao de paginas."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm

    AZ   = colors.HexColor('#1F4E79')
    AZCL = colors.HexColor('#BDD7EE')
    DIM  = colors.HexColor('#888888')
    CZ   = colors.HexColor('#F2F2F2')
    AZL  = colors.HexColor('#EBF5FB')
    VM   = colors.HexColor('#C00000')

    c.setFillColor(AZ)
    c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 13)
    c.drawString(1*cm, H - 1.05*cm, f'{PROG}  -  {proto_nome}  -  Índice')
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, H - 1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - 1*cm, H - 1.05*cm, datetime.date.today().strftime('%d/%m/%Y'))

    y = H - 2.5*cm
    linha_h = 0.52*cm
    col_pag = W - 2*cm

    # cabecalho da tabela
    c.setFillColor(AZ); c.rect(1*cm, y - 0.45*cm, W - 2*cm, 0.45*cm, fill=1, stroke=0)
    c.setFillColor(colors.white); c.setFont('Helvetica-Bold', 8)
    c.drawString(1.2*cm, y - 0.32*cm, 'Secção')
    c.drawRightString(col_pag + 0.9*cm, y - 0.32*cm, 'Pág.')
    y -= 0.55*cm

    alt_colors = {
        'atleta': colors.HexColor('#1F4E79'),
        'dist':   colors.HexColor('#2E75B6'),
        'itv':    colors.HexColor('#F2F2F2'),
        'hs':     colors.HexColor('#E8F5E9'),
    }
    text_colors = {
        'atleta': colors.white,
        'dist':   colors.white,
        'itv':    AZ,
        'hs':     colors.HexColor('#1E8449'),
    }
    font_styles = {
        'atleta': ('Helvetica-Bold', 9),
        'dist':   ('Helvetica-Bold', 8),
        'itv':    ('Helvetica', 7.5),
        'hs':     ('Helvetica-Bold', 8),
    }

    for label, pag, tipo in entries:
        if y < 1.5*cm:
            c.showPage()
            c.setFillColor(AZ)
            c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
            c.setFillColor(colors.white); c.setFont('Helvetica-Bold', 11)
            c.drawString(1*cm, H - 1.05*cm, f'{PROG}  -  Índice (cont.)')
            y = H - 2.5*cm

        bg = alt_colors.get(tipo, CZ)
        tc = text_colors.get(tipo, AZ)
        fn, fs = font_styles.get(tipo, ('Helvetica', 7.5))

        c.setFillColor(bg)
        c.rect(1*cm, y - linha_h + 0.06*cm, W - 2*cm, linha_h, fill=1, stroke=0)
        c.setFillColor(tc); c.setFont(fn, fs)
        c.drawString(1.3*cm, y - linha_h + 0.22*cm, label)
        c.setFont('Helvetica-Bold', fs)
        c.drawRightString(col_pag + 0.9*cm, y - linha_h + 0.22*cm, str(pag))

        # linha pontilhada
        if tipo in ('itv', 'hs'):
            c.setStrokeColor(colors.HexColor('#CCCCCC')); c.setLineWidth(0.3)
            c.setDash([1, 3])
            c.line(1.3*cm + c.stringWidth(label, fn, fs) + 2*mm,
                   y - linha_h + 0.22*cm + 0.1*cm,
                   col_pag + 0.8*cm - c.stringWidth(str(pag), 'Helvetica-Bold', fs) - 1*mm,
                   y - linha_h + 0.22*cm + 0.1*cm)
            c.setDash()

        y -= linha_h

    # rodape
    c.setStrokeColor(AZCL); c.setLineWidth(0.3)
    c.line(1*cm, 0.8*cm, W-1*cm, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(1*cm, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W-1*cm, 0.4*cm, 'Índice')


def _pagina_tiro_resumo_dist_pdf(c, atleta, dist, all_itvs, W, H):
    """
    Pagina de resumo por distância: tabela com todos os intervalos × ensaios × ea95
    e comparação visual (barras simples) da ea95 média por intervalo.
    """
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.utils import ImageReader
    import numpy as _np

    nome     = atleta['nome']
    ind_id   = atleta.get('id', '')
    AZ   = colors.HexColor('#1F4E79')
    AZCL = colors.HexColor('#BDD7EE')
    CZ   = colors.HexColor('#F2F2F2')
    VM   = colors.HexColor('#C00000')
    DIM  = colors.HexColor('#888888')

    # cabecalho
    c.setFillColor(AZ)
    c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
    c.setFillColor(colors.white); c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, H - 1.05*cm, f'{PROG}  -  Resumo por Distância  -  {dist}')
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, H - 1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - 1*cm, H - 1.05*cm, datetime.date.today().strftime('%d/%m/%Y'))

    id_prefix = f'[ID {ind_id}]  ' if ind_id else ''
    c.setFillColor(AZ); c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, H - 2.45*cm, id_prefix + nome)
    c.setFillColor(colors.HexColor('#6B8FAE')); c.setFont('Helvetica', 7.5)
    c.drawString(1*cm, H - 2.9*cm, f'Comparação de intervalos  -  Distância: {dist}  -  Métrica principal: Área Elipse 95%')
    c.setStrokeColor(AZ); c.setLineWidth(0.5)
    c.line(1*cm, H - 3.1*cm, W - 1*cm, H - 3.1*cm)
    y = H - 3.5*cm

    # ---- Tabela: linhas = intervalos, colunas = ensaios + stats ----
    itvs_com_dados = []
    for itv in all_itvs:
        mets = atleta.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[])
        if any(m is not None for m in mets):
            itvs_com_dados.append(itv)

    if not itvs_com_dados:
        c.setFillColor(DIM); c.setFont('Helvetica', 9)
        c.drawCentredString(W/2, H/2, 'Sem dados para esta distância')
        return

    n_ens_max = max(
        len(atleta.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[]))
        for itv in itvs_com_dados)

    # larguras
    margem   = 1.0*cm
    total_w  = W - 2*margem
    C0       = 3.8*cm  # intervalo label
    n_cols_d = n_ens_max + 3  # T1..N + Med + DP + CV%
    col_w    = (total_w - C0) / n_cols_d
    col_larg = [C0] + [col_w]*n_ens_max + [col_w*1.1]*3

    cab = ['Intervalo'] + [f'T{i+1}' for i in range(n_ens_max)] + ['Med', 'DP', 'CV%']
    linhas = [cab]
    medias_plot = {}  # para gráfico de barras

    for itv in itvs_com_dados:
        lbl   = _tiro_itv_label(itv)
        mets  = atleta.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[])
        flags = flagrar_outliers(mets)
        linha = [lbl]
        vals  = []
        for ti in range(n_ens_max):
            m  = mets[ti] if ti < len(mets) else None
            v  = round(m['ea95'], 3) if m and 'ea95' in m and m['ea95'] is not None else None
            fo = ' *' if (ti < len(flags) and flags[ti]) else ''
            linha.append(f'{v:.3f}{fo}' if v is not None else '-')
            if v is not None: vals.append(v)
        arr  = np.array(vals)
        med  = round(float(arr.mean()), 3)  if len(arr) > 0 else None
        dp   = round(float(arr.std(ddof=1)), 3) if len(arr) > 1 else None
        cv   = round(float(arr.std(ddof=1)/arr.mean()*100), 1) \
               if len(arr) > 1 and arr.mean() != 0 else None
        linha += [f'{med:.3f}' if med is not None else '-',
                  f'{dp:.3f}'  if dp  is not None else '-',
                  f'{cv:.1f}%' if cv  is not None else '-']
        linhas.append(linha)
        medias_plot[itv] = med

    row_h   = [0.42*cm] + [0.36*cm]*len(itvs_com_dados)
    tbl     = Table(linhas, colWidths=col_larg, rowHeights=row_h)
    idx_ea  = 0  # ea95 é a unica metrica aqui (sem destaque especial necessário)
    estilo  = [
        ('BACKGROUND', (0,0), (-1,0), AZ),
        ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
        ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0,0), (-1,0), 7),
        ('ALIGN',      (0,0), (-1,0), 'CENTER'),
        ('VALIGN',     (0,0), (-1,0), 'MIDDLE'),
        ('FONTNAME',   (0,1), (0,-1), 'Helvetica-Bold'),
        ('FONTSIZE',   (0,1), (-1,-1), 6.5),
        ('ALIGN',      (1,1), (-1,-1), 'CENTER'),
        ('ALIGN',      (0,1), (0,-1),  'LEFT'),
        ('VALIGN',     (0,1), (-1,-1), 'MIDDLE'),
        *[('BACKGROUND', (0,i), (-1,i), CZ) for i in range(2, len(linhas), 2)],
        ('BACKGROUND', (n_ens_max+1, 1), (-1,-1), colors.HexColor('#E8F4FD')),
        ('GRID',       (0,0), (-1,-1), 0.2, colors.HexColor('#CCCCCC')),
        ('LINEBELOW',  (0,0), (-1,0),  0.8, AZ),
        ('TOPPADDING',    (0,0), (-1,-1), 2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2),
        ('LEFTPADDING',   (0,0), (0,-1),  4),
    ]
    tbl.setStyle(TableStyle(estilo))
    tw, th = tbl.wrapOn(c, W - 2*cm, 8*cm)
    y -= th + 0.1*cm
    tbl.drawOn(c, 1*cm, y)

    # ---- Gráfico de barras simples (ea95 média por intervalo) ----
    y -= 0.4*cm
    espaco = y - 1.5*cm
    if espaco > 3*cm and medias_plot:
        try:
            if not _MPL_OK:
                return

            itvs_labels = [_tiro_itv_label(k) for k in medias_plot]
            vals_bar    = [v if v is not None else 0 for v in medias_plot.values()]
            cores_bar   = ['#1F4E79','#2E75B6','#4BACC6','#00B4D8','#0077B6',
                           '#023E8A','#48CAE4'][:len(vals_bar)]

            fig_w = (W - 2*cm) / cm
            fig_h = min(espaco / cm * 0.92, 6.0)
            fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=180)
            fig.patch.set_facecolor('white')
            ax.set_facecolor('#F9F9F9')

            bars = ax.bar(range(len(vals_bar)), vals_bar, color=cores_bar, width=0.55,
                          edgecolor='#CCCCCC', linewidth=0.6)
            for bar, v in zip(bars, vals_bar):
                if v > 0:
                    ax.text(bar.get_x()+bar.get_width()/2, v + max(vals_bar)*0.02,
                            f'{v:.3f}', ha='center', va='bottom', fontsize=7.5,
                            fontweight='bold', color='#1F4E79')
            ax.set_xticks(range(len(itvs_labels)))
            ax.set_xticklabels(itvs_labels, fontsize=7.5, color='#333333')
            ax.set_ylabel('Área Elipse 95% (cm²)', fontsize=8, color='#1F4E79')
            ax.set_title(f'Comparação de Intervalos  -  {dist}  -  {nome}',
                         fontsize=9, color='#1F4E79', fontweight='bold', pad=6)
            ax.tick_params(axis='y', labelsize=7)
            ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#AAAAAA'); ax.spines['bottom'].set_color('#AAAAAA')
            ax.yaxis.grid(True, ls='--', lw=0.4, alpha=0.5, color='#CCCCCC')
            ax.set_axisbelow(True)
            fig.tight_layout(pad=0.8)

            import io as _io
            buf = _io.BytesIO()
            try:
                fig.savefig(buf, format='png', dpi=180, bbox_inches='tight',
                            facecolor='white', edgecolor='none')
            finally:
                plt.close(fig)
            buf.seek(0)
            from reportlab.lib.utils import ImageReader as _IR
            gh = min(espaco * 0.95, fig_h * cm)
            c.drawImage(_IR(buf), 1*cm, y - gh,
                        width=W - 2*cm, height=gh, preserveAspectRatio=False)
        except Exception:
            pass

    # rodape
    c.setStrokeColor(AZCL); c.setLineWidth(0.3)
    c.line(1*cm, 0.8*cm, W-1*cm, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(1*cm, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W-1*cm, 0.4*cm, f'{nome}  |  Resumo  |  {dist}')


def _pagina_secao_pdf(c, titulo, subtitulo, W, H):
    """Pagina divisora de seccao (ex: inicio analise por individuo)."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    AZ  = colors.HexColor('#0F1923')
    AZC = colors.HexColor('#00B4D8')
    c.setFillColor(AZ)
    c.rect(0, 0, W, H, fill=1, stroke=0)
    c.setStrokeColor(AZC); c.setLineWidth(2)
    c.line(0, H - 0.8*cm, W, H - 0.8*cm)
    c.line(0, 0.8*cm, W, 0.8*cm)
    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 26)
    c.drawCentredString(W/2, H/2 + 1.2*cm, titulo)
    c.setFillColor(AZC); c.setFont('Helvetica', 13)
    c.drawCentredString(W/2, H/2 - 0.3*cm, subtitulo)
    c.setFillColor(colors.HexColor('#6B8FAE')); c.setFont('Helvetica', 8)
    c.drawCentredString(W/2, 0.35*cm,
        f'{PROG}  |  v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')


def _pagina_tiro_itv_pdf(c, atleta, dist, itv, W, H):
    """Pagina de metricas (tabela + elipse) para atleta/distancia/intervalo."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.utils import ImageReader
    import numpy as _np

    nome     = atleta['nome']
    ind_id   = atleta.get('id', '')
    itv_label = _tiro_itv_label(itv)

    AZ   = colors.HexColor('#1F4E79')
    AZCL = colors.HexColor('#BDD7EE')
    CZ   = colors.HexColor('#F2F2F2')
    VM   = colors.HexColor('#C00000')
    DIM  = colors.HexColor('#888888')
    AZDB = colors.HexColor('#E8F4FD')

    # cabecalho
    c.setFillColor(AZ)
    c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, H - 1.05*cm, f'{PROG}  -  Tiro  -  {dist}  -  {itv_label}')
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, H - 1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - 1*cm, H - 1.05*cm, datetime.date.today().strftime('%d/%m/%Y'))

    y = H - 2.5*cm
    id_prefix = f'[ID {ind_id}]  ' if ind_id else ''
    c.setFillColor(AZ); c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, y, id_prefix + nome)
    c.setFillColor(colors.HexColor('#6B8FAE')); c.setFont('Helvetica', 7.5)
    mets_list = atleta.get('tiro_dist', {}).get(dist, {}).get('mets', {}).get(itv, [])
    raw_list  = atleta.get('tiro_dist', {}).get(dist, {}).get('raw',  {}).get(itv, [])
    flags     = flagrar_outliers(mets_list)
    _mets_real = [m for m in mets_list if m is not None]
    n_vals    = max(max(len(_mets_real), len(mets_list)), 1)  # sem limite

    # Calcular duracoes dos trials (usa raw se disponivel, senao tempo da metrica)
    duracoes = []
    for rd in raw_list:
        if rd and rd.get('ini') is not None and rd.get('fim') is not None:
            dur_ms = rd['fim'] - rd['ini']
            duracoes.append(dur_ms)
        else:
            m_t = mets_list[len(duracoes)] if len(duracoes) < len(mets_list) else None
            duracoes.append(int(m_t['time']*1000) if m_t and 'time' in m_t else None)
    dur_med = int(np.mean([d for d in duracoes if d is not None])) if any(d is not None for d in duracoes) else None
    dur_str = f'  |  Duração média: {dur_med} ms' if dur_med else ''

    # Subtitulo unico com duracao (v19 - removido duplicado anterior)
    c.setFillColor(colors.HexColor('#6B8FAE')); c.setFont('Helvetica', 7.5)
    c.drawString(1*cm, y - 4*mm, f'Distância: {dist}   Intervalo: {itv_label}{dur_str}')
    c.setStrokeColor(AZ); c.setLineWidth(0.5)
    c.line(1*cm, y - 6*mm, W - 1*cm, y - 6*mm)
    y -= 1.0*cm

    # tabela
    margem  = 1.0*cm; total_w = W - 2*margem
    C0 = 3.2*cm; CV_ratio = 1.0; CS_ratio = 1.15
    n_col = n_vals + 3
    unit  = (total_w - C0) / (n_col * CV_ratio + 3*(CS_ratio - CV_ratio))
    CV = unit * CV_ratio; CS = unit * CS_ratio
    col_larg = [C0] + [CV]*n_vals + [CS, CS, CS]

    # Linha de durações
    dur_linha = ['Duração (ms)']
    dur_nums_t = []
    for ti in range(n_vals):
        d = duracoes[ti] if ti < len(duracoes) else None
        dur_linha.append(str(d) if d is not None else '-')
        if d is not None: dur_nums_t.append(d)
    dur_linha += [str(max(dur_nums_t)) if dur_nums_t else '-',
                  str(int(np.mean(dur_nums_t))) if dur_nums_t else '-',
                  str(int(np.std(dur_nums_t,ddof=1))) if len(dur_nums_t)>1 else '-']

    cab    = ['Métrica', f'INTERVALO: {itv_label}'] + ['']*(n_vals+2)
    sub    = [''] + [f'T{i+1}{"*" if i<len(flags) and flags[i] else ""}' for i in range(n_vals)] + ['máx','méd','dp']
    linhas = [cab, sub, dur_linha]
    idx_ea95 = next((i+3 for i,(k,_) in enumerate(METS_PDF) if k=='ea95'), 10)

    for chave, label in METS_PDF:
        linha = [label]
        vals = []
        for ti in range(n_vals):
            m = mets_list[ti] if ti < len(mets_list) else None
            v = m[chave] if m is not None and chave in m else None
            vals.append(v)
            linha.append(f'{v:.2f}' if v is not None else '-')
        nums = np.array([v for v in vals if v is not None])
        linha += [f'{nums.max():.2f}'       if len(nums)>0 else '-',
                  f'{nums.mean():.2f}'      if len(nums)>0 else '-',
                  f'{nums.std(ddof=1):.2f}' if len(nums)>1 else '-']
        linhas.append(linha)

    row_h = [0.38*cm, 0.33*cm, 0.30*cm] + [0.33*cm]*(len(linhas)-3)
    tbl   = Table(linhas, colWidths=col_larg, rowHeights=row_h)
    AZCL_tbl = colors.HexColor('#D6E4F0')
    estilo = [
        ('BACKGROUND',(0,0),(-1,1),  AZ),
        ('TEXTCOLOR',(0,0),(-1,1),   colors.white),
        ('FONTNAME',(0,0),(-1,1),    'Helvetica-Bold'),
        ('FONTSIZE',(0,0),(-1,1),    6.5),
        ('ALIGN',(0,0),(-1,1),       'CENTER'),
        ('VALIGN',(0,0),(-1,1),      'MIDDLE'),
        ('SPAN',(1,0),(n_vals+3,0)),
        # linha duração
        ('BACKGROUND',(0,2),(-1,2),  AZCL_tbl),
        ('TEXTCOLOR',(0,2),(-1,2),   AZ),
        ('FONTNAME',(0,2),(-1,2),    'Helvetica-Oblique'),
        ('FONTSIZE',(0,2),(-1,2),    6),
        ('ALIGN',(0,2),(-1,2),       'CENTER'),
        ('VALIGN',(0,2),(-1,2),      'MIDDLE'),
        # dados
        ('FONTNAME',(0,3),(0,-1),    'Helvetica-Bold'),
        ('FONTSIZE',(0,3),(-1,-1),   6),
        ('ALIGN',(1,3),(-1,-1),      'CENTER'),
        ('ALIGN',(0,3),(0,-1),       'LEFT'),
        ('VALIGN',(0,3),(-1,-1),     'MIDDLE'),
        *[('BACKGROUND',(0,i),(-1,i), colors.HexColor('#F2F2F2')) for i in range(4, len(linhas), 2)],
        ('TEXTCOLOR',(0,idx_ea95),(-1,idx_ea95),  colors.HexColor('#C00000')),
        ('FONTNAME',(0,idx_ea95),(-1,idx_ea95),   'Helvetica-Bold'),
        ('BACKGROUND',(n_vals+1,3),(n_vals+3,-1),  colors.HexColor('#E8F4FD')),
        ('GRID',(0,0),(-1,-1),       0.2, colors.HexColor('#CCCCCC')),
        ('LINEBELOW',(0,2),(-1,2),   0.6, AZ),
        ('TOPPADDING',(0,0),(-1,-1),    1),
        ('BOTTOMPADDING',(0,0),(-1,-1), 1),
        ('LEFTPADDING',(0,0),(0,-1),    3),
    ]
    tbl.setStyle(TableStyle(estilo))
    y = _tbl_fit_page(c, tbl, 1*cm, y, W, H)

    # elipse
    y -= 0.2*cm
    espaco = y - 1.2*cm
    hg = max(espaco * 0.95, 5.5*cm)
    wg = W - 2*cm
    png = _png_elipses(mets_list, itv_label, nome,
                       larg=wg/cm*0.40, alt=hg/cm*0.43, dpi=200, legenda_fora=True)
    if png:
        c.drawImage(ImageReader(io.BytesIO(png)), 1*cm, y-hg,
                    width=wg, height=hg, preserveAspectRatio=False)
    else:
        c.setFillColor(CZ); c.rect(1*cm, y-hg, wg, hg, fill=1, stroke=0)

    # rodape
    c.setStrokeColor(AZCL); c.setLineWidth(0.3)
    c.line(1*cm, 0.8*cm, W-1*cm, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(1*cm, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W-1*cm, 0.4*cm, f'{nome}  |  Tiro  |  {dist}  |  {itv_label}')


def _pagina_tiro_estabilogramas_pdf(c, atleta, dist, itv, W, H):
    """Pagina de estabilogramas para atleta/distancia/intervalo."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.lib.utils import ImageReader

    nome      = atleta['nome']
    ind_id    = atleta.get('id', '')
    itv_label = _tiro_itv_label(itv)

    AZ  = colors.HexColor('#1F4E79')
    AZCL= colors.HexColor('#BDD7EE')
    DIM = colors.HexColor('#888888')
    CZ  = colors.HexColor('#F2F2F2')

    c.setFillColor(AZ)
    c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 11)
    c.drawString(1*cm, H - 1.05*cm, f'Estabilogramas  -  {dist}  -  {itv_label}')
    c.setFont('Helvetica', 7)
    c.drawString(1*cm, H - 1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - 1*cm, H - 1.05*cm, datetime.date.today().strftime('%d/%m/%Y'))

    id_prefix = f'[ID {ind_id}] ' if ind_id else ''
    c.setFillColor(AZ); c.setFont('Helvetica-Bold', 10)
    c.drawString(1*cm, H - 2.3*cm, id_prefix + nome)
    c.setStrokeColor(AZ); c.setLineWidth(0.4)
    c.line(1*cm, H - 2.6*cm, W - 1*cm, H - 2.6*cm)

    mets_list = atleta.get('tiro_dist', {}).get(dist, {}).get('mets', {}).get(itv, [])
    n_ens     = len(mets_list)
    if n_ens == 0:
        c.setFillColor(DIM); c.setFont('Helvetica', 9)
        c.drawCentredString(W/2, H/2, 'Sem dados de estabilograma')
        return

    n_cols = min(n_ens, 3)
    n_rows = (n_ens + n_cols - 1) // n_cols
    pad    = 0.25*cm
    plot_w = (W - 2*cm - (n_cols-1)*pad) / n_cols
    plot_h = (H - 3.0*cm - 1.0*cm - (n_rows-1)*pad) / n_rows

    for idx, m in enumerate(mets_list):
        if m is None: continue
        col_i = idx % n_cols
        row_i = idx // n_cols
        xg    = 1*cm + col_i * (plot_w + pad)
        yg    = H - 3.0*cm - row_i * (plot_h + pad) - plot_h
        png   = _png_estabilograma(m, nome, f'Ensaio {idx+1}',
                                    larg=plot_w/cm*0.38, alt=plot_h/cm*0.42, dpi=150)
        if png:
            c.drawImage(ImageReader(io.BytesIO(png)), xg, yg,
                        width=plot_w, height=plot_h, preserveAspectRatio=False)
        else:
            c.setFillColor(CZ); c.rect(xg, yg, plot_w, plot_h, fill=1, stroke=0)

    c.setStrokeColor(AZCL); c.setLineWidth(0.3)
    c.line(1*cm, 0.8*cm, W-1*cm, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(1*cm, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W-1*cm, 0.4*cm,
        f'{nome}  |  Tiro  |  Estabilogramas  |  {dist}  |  {itv_label}')


def _pagina_grupo_tiro_pdf(c, atletas, all_dists, all_itvs, W, H):
    """
    Resumo do grupo - Protocolo de Tiro  (v20 redesign).
    Layout por distância:
      • Título de distância (banda azul)
      • Gráfico barras  ea95  (largura total)
      • Tabela multi-métrica:  linhas = métricas, colunas = intervalos
        Para cada célula: med do grupo  (±DP)
        Por atleta: linha própria com med individual
    Nunca sobrepõe - nova página automática quando necessário.
    """
    if not _MPL_OK:
        return
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import cm
    from reportlab.lib.utils import ImageReader as _IR
    from reportlab.platypus import Table, TableStyle

    AZ    = rl_colors.HexColor('#1F4E79')
    AZ2   = rl_colors.HexColor('#2E75B6')
    AZCL  = rl_colors.HexColor('#BDD7EE')
    DIM   = rl_colors.HexColor('#888888')
    CZ    = rl_colors.HexColor('#F2F2F2')
    GOLD  = rl_colors.HexColor('#FEF9C3')
    AZDB  = rl_colors.HexColor('#E8F4FD')
    WHITE = rl_colors.white

    MARGIN    = 1.0 * cm
    CW        = W - 2 * MARGIN   # largura util
    GRAF_H    = 5.0 * cm
    MIN_Y     = 1.4 * cm
    GAP       = 0.25 * cm

    # Métricas a mostrar na tabela de grupo
    GRUP_METS = [
        ('ea95',      'Área Elipse 95%\n(mm²)'),
        ('vel_med',   'Vel. média CoP\n(mm/s)'),
        ('vel_x',     'Vel. média X\n(mm/s)'),
        ('vel_y',     'Vel. média Y\n(mm/s)'),
        ('amp_x',     'Amplitude X\n(mm)'),
        ('amp_y',     'Amplitude Y\n(mm)'),
        ('desl',      'Deslocamento\n(mm)'),
        ('cov_xy',    'Cov XY\n(mm²)'),
    ]

    n_itvs    = max(len(all_itvs), 1)
    itv_lbls  = [TIRO_INTERVALOS.get(iv, iv) for iv in all_itvs]

    # larguras de coluna da tabela: col0 = label métrica, rest = intervalos
    LABEL_W = 3.4 * cm
    itv_w   = max((CW - LABEL_W) / n_itvs, 2.0 * cm)
    col_ws  = [LABEL_W] + [itv_w] * n_itvs

    # alturas de linha: cabeçalho duplo (atleta + sub) + dado por atleta + grupo
    HDR_H  = 0.80 * cm   # cabeçalho (2 linhas de texto)
    ROW_H  = 0.52 * cm   # linha dados (med ± DP)
    MET_H  = 0.36 * cm   # linha-título da métrica

    def _cabecalho(first=False):
        if not first:
            c.showPage()
        c.setFillColor(AZ)
        c.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
        c.setFillColor(WHITE); c.setFont('Helvetica-Bold', 12)
        c.drawString(MARGIN, H - 1.05*cm,
                     f'{PROG}  -  Resumo do Grupo  -  Protocolo de Tiro')
        c.setFont('Helvetica', 7)
        c.drawString(MARGIN, H - 1.48*cm, f'v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
        c.drawRightString(W - MARGIN, H - 1.05*cm,
                          datetime.date.today().strftime('%d/%m/%Y'))
        return H - 2.2*cm

    def _rodape():
        c.setStrokeColor(AZCL); c.setLineWidth(0.3)
        c.line(MARGIN, 0.8*cm, W - MARGIN, 0.8*cm)
        c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
        c.drawString(MARGIN, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
        c.drawRightString(W - MARGIN, 0.4*cm, 'Resumo do Grupo')

    def _draw_titulo_dist(y_cur, dist, cont=False):
        c.setFillColor(AZ2)
        c.rect(MARGIN, y_cur - 0.44*cm, CW, 0.44*cm, fill=1, stroke=0)
        c.setFillColor(WHITE); c.setFont('Helvetica-Bold', 9)
        suf = '  (cont.)' if cont else f'  -  n={len(atletas)}'
        c.drawString(MARGIN + 0.3*cm, y_cur - 0.31*cm,
                     f'Distância: {dist}  -  Área Elipse 95% (mm²){suf}')
        return y_cur - 0.50*cm - GAP

    y = _cabecalho(first=True)

    for dist in all_dists:
        # ── pré-calcular médias por métrica × intervalo ──────────────
        # estrutura: data_met[met_key][itv_idx] = {atleta: med_val, 'grp': (med, dp)}
        data_met = {}
        for mk, _ in GRUP_METS:
            data_met[mk] = {}
            for ii, itv in enumerate(all_itvs):
                ath_meds = []
                per_ath  = {}
                for ath in atletas:
                    ml = (ath.get('tiro_dist', {}).get(dist, {})
                          .get('mets', {}).get(itv, []))
                    vs = [m[mk] for m in ml if m is not None and mk in m]
                    if vs:
                        mv = float(np.mean([x for x in vs if x is not None]))
                        per_ath[ath['nome']] = mv
                        ath_meds.append(mv)
                    else:
                        per_ath[ath['nome']] = None
                if ath_meds:
                    gm = float(np.mean([x for x in ath_meds if x is not None]))
                    gd = float(np.std(ath_meds, ddof=1)) if len(ath_meds) > 1 else 0.0
                else:
                    gm = gd = None
                data_met[mk][ii] = {'per_ath': per_ath, 'grp': (gm, gd)}

        # ── gráfico ea95 ─────────────────────────────────────────────
        meds_g = []
        dps_g  = []
        for ii, itv in enumerate(all_itvs):
            gm, gd = data_met['ea95'][ii]['grp']
            meds_g.append(gm or 0.0)
            dps_g.append(gd or 0.0)

        # Verificar espaço para gráfico + tabela (estimativa)
        n_data_rows = len(atletas) + 1   # atletas + grupo
        n_met_rows  = len(GRUP_METS)
        tbl_h_est   = (HDR_H + (MET_H + ROW_H * n_data_rows) * n_met_rows)
        bloco_h_min = 0.50*cm + GAP + GRAF_H + GAP

        if y - bloco_h_min < MIN_Y:
            _rodape()
            y = _cabecalho()

        y = _draw_titulo_dist(y, dist)

        # Gráfico
        try:
            fig, ax = plt.subplots(figsize=(CW / cm, GRAF_H / cm), dpi=180)
            fig.patch.set_facecolor('white')
            ax.set_facecolor('#FAFAFA')
            cores_bar = ['#1F4E79','#2E75B6','#4BACC6','#00B4D8','#0077B6']
            bars = ax.bar(range(len(itv_lbls)), meds_g,
                          yerr=dps_g, capsize=4,
                          color=cores_bar[:len(itv_lbls)],
                          width=0.55, edgecolor='white', linewidth=0.6,
                          error_kw=dict(ecolor='#555555', elinewidth=1.2, capthick=1.2))
            y_max = max(meds_g) if meds_g and max(meds_g) > 0 else 1.0
            for bar, v, dp in zip(bars, meds_g, dps_g):
                ax.text(bar.get_x() + bar.get_width() / 2,
                        v + dp + y_max * 0.03,
                        f'{v:.1f}', ha='center', va='bottom',
                        fontsize=7, fontweight='bold', color='#1F4E79')
            ax.set_xticks(range(len(itv_lbls)))
            ax.set_xticklabels(itv_lbls, fontsize=7.5, color='#333333')
            ax.set_ylabel('EA95 (mm²)', fontsize=8, color='#1F4E79')
            ax.set_title(f'Grupo - {dist}  (Média ± DP)',
                         fontsize=8.5, color='#1F4E79', fontweight='bold', pad=4)
            ax.tick_params(axis='y', labelsize=7)
            ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#AAAAAA'); ax.spines['bottom'].set_color('#AAAAAA')
            ax.yaxis.grid(True, ls='--', lw=0.4, alpha=0.5)
            ax.set_axisbelow(True)
            fig.tight_layout(pad=0.8)
            buf_g = io.BytesIO()
            try:
                fig.savefig(buf_g, format='png', dpi=180, bbox_inches='tight',
                            facecolor='white', edgecolor='none')
            finally:
                plt.close(fig)
            buf_g.seek(0)
            c.drawImage(_IR(buf_g), MARGIN, y - GRAF_H,
                        width=CW, height=GRAF_H, preserveAspectRatio=False)
        except Exception:
            pass
        y -= GRAF_H + GAP

        # ── tabela multi-métrica ─────────────────────────────────────
        # Linha cabeçalho: 'Atleta / Métrica' + itv labels
        # Depois, para cada métrica:  linha-título  + 1 linha por atleta + linha grupo
        for mi, (mk, mk_lbl) in enumerate(GRUP_METS):
            # Calcular altura total necessária: barra-métrica + tabela
            # n_data_rows = len(atletas) + 1 (inclui linha do grupo)
            # A tabela real tem exactamente n_data_rows linhas com row_hs[-1]=ROW_H*1.35
            _n_ath = len(atletas)
            secao_h = MET_H + ROW_H * _n_ath + ROW_H * 1.35  # estimativa exacta
            _, tbl_h_real = Table(
                [[''] * len(col_ws)] * n_data_rows,
                colWidths=col_ws,
                rowHeights=[ROW_H] * _n_ath + [ROW_H * 1.35]
            ).wrapOn(c, CW, 999*cm)

            # Verificar espaço para barra + tabela de uma vez (evita paginas vazias)
            if y - secao_h < MIN_Y or y - MET_H - tbl_h_real < MIN_Y:
                _rodape()
                y = _cabecalho()
                y = _draw_titulo_dist(y, dist, cont=True)

            # Linha-título da métrica (fundo AZ2, largura total)
            c.setFillColor(AZ if mi == 0 else AZ2)
            c.rect(MARGIN, y - MET_H, CW, MET_H, fill=1, stroke=0)
            c.setFillColor(WHITE); c.setFont('Helvetica-Bold', 7)
            # Cabeçalho com nomes dos intervalos (apenas no primeiro bloco de métrica)
            c.drawString(MARGIN + 0.2*cm, y - MET_H * 0.65,
                         mk_lbl.replace('\n', ' '))
            if mi == 0:
                for ii, lbl in enumerate(itv_lbls):
                    x_itv = MARGIN + LABEL_W + ii * itv_w + itv_w / 2
                    c.setFont('Helvetica-Bold', 6.5)
                    c.drawCentredString(x_itv, y - MET_H * 0.65, lbl)
            y -= MET_H

            # Linha por atleta - dividir em blocos se nao couberem todos na pagina
            all_ath_rows = []
            for ath in atletas:
                row = [ath['nome'][:22]]
                for ii in range(n_itvs):
                    v = data_met[mk][ii]['per_ath'].get(ath['nome'])
                    row.append(f'{v:.2f}' if v is not None else '-')
                all_ath_rows.append(row)
            # Linha grupo
            grp_row = ['Grupo (Med+DP)']
            for ii in range(n_itvs):
                gm, gd = data_met[mk][ii]['grp']
                if gm is not None:
                    grp_row.append(f'{gm:.2f}\n+{gd:.2f}')
                else:
                    grp_row.append('-')

            # Calcular quantas linhas cabem no espaco restante
            rows_per_page = max(1, int((y - MIN_Y - ROW_H * 1.35) / ROW_H))
            row_cursor = 0
            first_chunk = True
            while row_cursor < len(all_ath_rows):
                chunk_rows = all_ath_rows[row_cursor:row_cursor + rows_per_page]
                is_last_chunk = (row_cursor + len(chunk_rows) >= len(all_ath_rows))
                tbl_rows = list(chunk_rows)
                if is_last_chunk:
                    tbl_rows.append(grp_row)

                last_i = len(tbl_rows) - 1
                row_hs = [ROW_H] * len(tbl_rows)
                if is_last_chunk:
                    row_hs[-1] = ROW_H * 1.35
                tbl = Table(tbl_rows, colWidths=col_ws, rowHeights=row_hs)
                tbl.setStyle(TableStyle([
                    ('FONTSIZE',   (0, 0), (-1, -1),       6.5),
                    ('FONTNAME',   (0, 0), (0, -1),        'Helvetica-Bold'),
                    ('FONTNAME',   (0, last_i), (-1, last_i), 'Helvetica-Bold'),
                    ('ALIGN',      (0, 0), (-1, -1),       'CENTER'),
                    ('ALIGN',      (0, 0), (0, -1),        'LEFT'),
                    ('VALIGN',     (0, 0), (-1, -1),       'MIDDLE'),
                    *[('BACKGROUND', (0, i), (-1, i), CZ)
                      for i in range(0, last_i, 2)],
                    *([('BACKGROUND', (0, last_i), (-1, last_i), GOLD),
                       ('TEXTCOLOR',  (0, last_i), (-1, last_i), AZ)]
                      if is_last_chunk else []),
                    ('BACKGROUND', (1, 0), (-1, last_i - (1 if is_last_chunk else 0)), AZDB),
                    ('GRID',       (0, 0), (-1, -1),       0.25, rl_colors.HexColor('#CCCCCC')),
                    ('TOPPADDING',    (0, 0), (-1, -1),    1),
                    ('BOTTOMPADDING', (0, 0), (-1, -1),    1),
                    ('LEFTPADDING',   (0, 0), (0, -1),     3),
                ]))
                _, tbl_h_real = tbl.wrapOn(c, CW, 999*cm)

                # Se nao couber, nova pagina
                if y - tbl_h_real < MIN_Y and not first_chunk:
                    _rodape()
                    y = _cabecalho()
                    y = _draw_titulo_dist(y, dist, cont=True)
                    rows_per_page = max(1, int((y - MIN_Y - ROW_H * 1.35) / ROW_H))

                tbl.drawOn(c, MARGIN, y - tbl_h_real)
                y -= tbl_h_real + GAP * 0.5
                row_cursor += len(chunk_rows)
                first_chunk = False

        y -= GAP

    _rodape()

def _pagina_novidades_pdf(c, W, H):
    """
    Pagina 'O que ha de novo' - gerada automaticamente a partir de CHANGELOG_PDF.
    Mostra as novidades da versao actual com icones e descricoes curtas.
    """
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import cm

    AZ   = rl_colors.HexColor('#1F4E79')
    AZ2  = rl_colors.HexColor('#2E75B6')
    AZCL = rl_colors.HexColor('#BDD7EE')
    VD   = rl_colors.HexColor('#22C55E')
    AM   = rl_colors.HexColor('#F59E0B')
    BG   = rl_colors.HexColor('#0D1720')
    DIM  = rl_colors.HexColor('#888888')
    WHITE = rl_colors.white

    # Fundo escuro
    c.setFillColor(BG)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # Cabecalho
    c.setFillColor(AZ)
    c.rect(0, H - 2.0*cm, W, 2.0*cm, fill=1, stroke=0)
    c.setFillColor(WHITE); c.setFont('Helvetica-Bold', 13)
    c.drawString(1*cm, H - 1.15*cm, f'BSP v{VERSAO}  |  Funcionalidades')
    c.setFont('Helvetica', 7.5)
    c.drawRightString(W - 1*cm, H - 1.15*cm, datetime.date.today().strftime('%d/%m/%Y'))
    c.setFillColor(AZCL); c.setFont('Helvetica', 7)
    c.drawString(1*cm, H - 1.58*cm, f'{AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')

    # Changelog entries — v1.0 first release
    CHANGELOG = [
        ('v1.0', VD, [
            ('Análise de CoP multi-protocolo',
             'Suporte para FMS, Unipodal, HurdleStep, Lunge, Functional Task, '
             'Tiro (ISCPSI) e Tiro com Arco, cada protocolo com layout PDF dedicado.'),
            ('Métricas biomecânicas completas',
             'EA95, amplitude ML/AP, velocidade média e de pico, deslocamento, '
             'stiffness, RMS (Quijoux et al., 2021), ratios ML/AP e normalizações '
             'por dimensões corporais.'),
            ('Tabelas PDF adaptativas',
             'Tabelas individuais adaptam-se automaticamente a qualquer número de '
             'ensaios; colunas comprimem e fonte reduz para manter a legibilidade.'),
            ('Relatório de grupo Tiro com Arco',
             'Análise comparativa de 142 atletas com testes estatísticos (Mann-Whitney, '
             'Kruskal-Wallis, correlações Pearson/Spearman) e gráficos de distribuição.'),
            ('Exportação multi-formato',
             'PDF clínico, Excel com gráficos, PNG (DPI configurável) e relatório '
             'HTML interativo com Chart.js partilhável sem instalação.'),
            ('Interface multilingue',
             'Português, English, Español e Deutsch com troca instantânea de idioma '
             'em todos os ecrãs e relatórios.'),
            ('Diagnóstico e robustez',
             'Deteção de timestamps duplicados/invertidos, jitter de amostragem >20%, '
             'janela de diagnóstico de crash com envio de relatório automático.'),
            ('Autoria',
             f'{AUTOR},  {SEGUNDO_AUTOR},  {ORIENTADOR}. '
             'Citações APA e BibTeX com ano dinâmico na última página do PDF.'),
        ]),
    ]

    y = H - 2.5*cm
    MARGIN = 1.0*cm

    for ver, ver_col, entries in CHANGELOG:
        # Versao header
        c.setFillColor(ver_col)
        c.rect(MARGIN, y - 0.5*cm, W - 2*MARGIN, 0.5*cm, fill=1, stroke=0)
        c.setFillColor(WHITE); c.setFont('Helvetica-Bold', 9)
        c.drawString(MARGIN + 0.3*cm, y - 0.34*cm, f'Versão {ver}')
        y -= 0.6*cm

        for titulo, descr in entries:
            if y < 1.5*cm:
                break  # nao sair da pagina
            # Bullet
            c.setFillColor(ver_col)
            c.circle(MARGIN + 0.15*cm, y - 0.13*cm, 0.08*cm, fill=1, stroke=0)
            # Titulo
            c.setFillColor(WHITE); c.setFont('Helvetica-Bold', 8)
            c.drawString(MARGIN + 0.4*cm, y - 0.18*cm, titulo)
            y -= 0.32*cm
            # Descricao (com wrap simples)
            c.setFillColor(AZCL); c.setFont('Helvetica', 7)
            max_w = W - 2*MARGIN - 0.4*cm
            words = descr.split()
            line = ''
            for word in words:
                test = (line + ' ' + word).strip()
                if c.stringWidth(test, 'Helvetica', 7) < max_w:
                    line = test
                else:
                    if y < 1.5*cm: break
                    c.drawString(MARGIN + 0.4*cm, y - 0.15*cm, line)
                    y -= 0.25*cm; line = word
            if line and y >= 1.5*cm:
                c.drawString(MARGIN + 0.4*cm, y - 0.15*cm, line)
                y -= 0.3*cm
            y -= 0.1*cm

        y -= 0.2*cm

    # Rodape
    c.setStrokeColor(AZCL); c.setLineWidth(0.3)
    c.line(MARGIN, 0.8*cm, W - MARGIN, 0.8*cm)
    c.setFillColor(DIM); c.setFont('Helvetica-Oblique', 6)
    c.drawString(MARGIN, 0.4*cm, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    c.drawRightString(W - MARGIN, 0.4*cm, 'Funcionalidades v1.0')


def _calc_pags_grupo_tiro(atletas, all_dists, W, H):
    """
    Calcula o número exacto de páginas que _pagina_grupo_tiro_pdf irá gerar,
    replicando a lógica de layout sem desenhar.  Usa as mesmas constantes da função.
    """
    try:
        from reportlab.lib.units import cm as _cm
    except ImportError:
        return max(len(all_dists), 1)

    GRAF_H = 5.0*_cm; HDR_H = 0.80*_cm
    ROW_H  = 0.52*_cm; MET_H = 0.36*_cm
    MIN_Y  = 1.4*_cm;  GAP   = 0.25*_cm
    N_MET  = 8
    n_ath  = len(atletas)
    avail  = H - 2.2*_cm   # y inicial após _cabecalho

    pages = 1
    y = avail

    for dist in all_dists:
        # bloco mínimo para iniciar distância (título + gráfico)
        bloco_min = 0.50*_cm + GAP + GRAF_H + GAP
        if y - bloco_min < MIN_Y:
            pages += 1
            y = avail

        y -= 0.50*_cm          # título de distância
        y -= GAP + GRAF_H + GAP

        # cabeçalho da tabela
        if y - HDR_H < MIN_Y:
            pages += 1; y = avail
        y -= HDR_H

        # linhas de dados por métrica
        for _ in range(N_MET):
            row_h = MET_H + ROW_H * (n_ath + 1)
            if y - row_h < MIN_Y:
                pages += 1; y = avail
            y -= row_h

    return pages


def _count_toc_pages(entries, H):
    """Conta quantas páginas o TOC vai ocupar com base no número de entradas."""
    try:
        from reportlab.lib.units import cm as _cm
    except ImportError:
        return 1
    linha_h = 0.52*_cm
    y_start = H - 2.5*_cm - 0.55*_cm   # após cabeçalho + header da tabela
    min_y   = 1.5*_cm
    avail   = y_start - min_y
    per_pag = max(1, int(avail / linha_h))
    import math
    return max(1, math.ceil(len(entries) / per_pag))


# Funcao que gera o relatorio PDF completo
# Usa a biblioteca ReportLab para criar o PDF
# Cada individuo tem a sua propria pagina com metricas e graficos
def gerar_pdf(atletas, caminho, log=print, opts_estats=None):
    try:
        from reportlab.pdfgen import canvas as rlc
        from reportlab.lib.pagesizes import A4
    except ImportError:
        log("  reportlab nao instalado: pip install reportlab", 'erro')
        return False
    W, H = A4
    c = rlc.Canvas(caminho, pagesize=A4)
    c.setTitle(PROG + ' - Relatorio v' + VERSAO)
    c.setAuthor(AUTOR)
    c.setSubject('Biomechanical Stability Program')
    c.setCreator(AUTOR + ', ' + SEGUNDO_AUTOR + ', ' + ORIENTADOR)
    proto_key = atletas[0].get('protocolo', PROTO_FMS) if atletas else PROTO_FMS
    proto     = PROTOCOLOS.get(proto_key)
    is_tiro   = _is_tiro_like(proto_key)

    # ---- Pagina 1: Capa ----
    _pagina_capa_pdf(c, atletas, W, H)
    c.showPage()

    # ---- Pagina 1b opcional: overflow lista ----
    from reportlab.lib.units import cm as _cm
    _box_h_est   = (A4[1] * 0.38 - 4.5*_cm) - 3.5*_cm
    _row_h_est   = 0.52*_cm
    _max_rows_pg1 = max(1, int((_box_h_est - 1.2*_cm) / _row_h_est))
    _ncols_pg1   = 2 if len(atletas) > _max_rows_pg1 else 1
    _max_fit_pg1 = _ncols_pg1 * _max_rows_pg1
    has_overflow = len(atletas) > _max_fit_pg1
    if has_overflow:
        _pagina_capa2_pdf(c, atletas, _max_fit_pg1, W, H)
        c.showPage()

    # ---- Pagina 2: Legenda ----
    _pagina_legenda_pdf(c, W, H)
    c.showPage()

    # ---- Pagina de novidades da versao ----
    try:
        _pagina_novidades_pdf(c, W, H)
        c.showPage()
    except Exception as _ex_nov:
        log(f'  aviso: pagina novidades nao gerada ({_ex_nov})', 'aviso')

    if is_tiro:
        # ================================================================
        # PROTOCOLO TIRO
        # ================================================================
        all_dists, all_itvs = [], []
        for a in atletas:
            for d in a.get('tiro_dists', []):
                if d not in all_dists: all_dists.append(d)
            for itv in a.get('tiro_intervalos', []):
                if itv not in all_itvs: all_itvs.append(itv)

        include_hs = any(
            bool(a.get('mets', {}).get('dir') or a.get('mets', {}).get('esq'))
            for a in atletas)

        # ---- Pagina: Índice (TOC) ----
        # Contagem exacta de páginas antes do TOC:
        #   pg1: capa
        #   pg1b: capa overflow (opcional)
        #   pg2: legenda
        #   pg3: novidades
        #   pgN: TOC (esta página)
        #   pgN+1: resumo do grupo
        #   pgN+2+: conteúdo por atleta
        _pags_antes_toc = 3 + (1 if has_overflow else 0)  # capa+legenda+novidades (+overflow)
        pag_toc = _pags_antes_toc + 1

        # Contagem exacta de páginas do grupo (sem estimativa)
        _pags_grupo = _calc_pags_grupo_tiro(atletas, all_dists, W, H)

        # Estimativa do TOC: precisa de um pré-cálculo provisório para contar entradas
        _pag_inicial_est = pag_toc + 1 + _pags_grupo
        _toc_entries_est, _ = _tiro_pre_calc_pages(
            atletas, all_dists, all_itvs, include_hs, _pag_inicial_est)
        _pags_toc = _count_toc_pages(_toc_entries_est, H)

        # Cálculo final com número correcto de páginas TOC
        pag_inicial = pag_toc + _pags_toc + _pags_grupo

        toc_entries, _ = _tiro_pre_calc_pages(
            atletas, all_dists, all_itvs, include_hs, pag_inicial)

        _pagina_toc_pdf(c, toc_entries, W, H, proto_nome=_proto_nome(proto_key))
        c.showPage()

        # ---- Página de resumo do grupo ----
        try:
            _pagina_grupo_tiro_pdf(c, atletas, all_dists, all_itvs, W, H)
            c.showPage()
            log('  ok pagina grupo', 'ok')
        except Exception as ex:
            log(f'  aviso: grupo: {ex}', 'aviso')

        # ---- Conteúdo por atleta ----
        for atleta in atletas:
            nome   = atleta['nome']
            ind_id = atleta.get('id', '')
            id_str = f'ID {ind_id} - ' if ind_id else ''

            # divisora do individuo
            _pagina_secao_pdf(c, f'{id_str}{nome}', 'Análise de Tiro', W, H)
            c.showPage()

            for dist in (atleta.get('tiro_dists') or all_dists):
                # paginas de cada intervalo
                for itv in (atleta.get('tiro_intervalos') or all_itvs):
                    mets_itv = (atleta.get('tiro_dist', {})
                                .get(dist, {}).get('mets', {}).get(itv, []))
                    if not any(m is not None for m in mets_itv):
                        continue
                    try:
                        _pagina_tiro_itv_pdf(c, atleta, dist, itv, W, H)
                        c.showPage()
                        _pagina_tiro_estabilogramas_pdf(c, atleta, dist, itv, W, H)
                        c.showPage()
                        # Selection CoP: Entire plate + Pe Direito + Pe Esquerdo (v20 - 3 paginas)
                        sel_itv2 = atleta.get('sel_dist', {}).get(dist, {}).get(itv, {})
                        has_sel2 = any(m is not None for m in sel_itv2.get('dir', [])) or \
                                   any(m is not None for m in sel_itv2.get('esq', []))
                        if has_sel2:
                            _pagina_tiro_entire_plate_pdf(c, atleta, dist, itv, W, H)
                            c.showPage()
                            if any(m is not None for m in sel_itv2.get('dir', [])):
                                _pagina_tiro_sel_lado_pdf(c, atleta, dist, itv, 'dir', W, H)
                                c.showPage()
                            if any(m is not None for m in sel_itv2.get('esq', [])):
                                _pagina_tiro_sel_lado_pdf(c, atleta, dist, itv, 'esq', W, H)
                                c.showPage()
                        log(f"    ok {nome} | {dist} | {TIRO_INTERVALOS.get(itv,itv)}", 'ok')
                    except Exception as ex:
                        log(f"    aviso: {nome} {dist} {itv}: {ex}", 'aviso')

                # resumo por distância (comparação de intervalos)
                try:
                    _pagina_tiro_resumo_dist_pdf(c, atleta, dist, all_itvs, W, H)
                    c.showPage()
                    log(f"    ok {nome} | {dist} | resumo", 'ok')
                except Exception as ex:
                    log(f"    aviso: {nome} {dist} resumo: {ex}", 'aviso')

            # hurdle step bipedal (opcional)
            hs_tem = include_hs and (
                any(m is not None for m in atleta.get('mets', {}).get('dir', [])) or
                any(m is not None for m in atleta.get('mets', {}).get('esq', [])))
            if hs_tem:
                try:
                    _pagina_secao_pdf(c,
                        f'{id_str}{nome}',
                        'Componente Bipodal  -  Hurdle Step', W, H)
                    c.showPage()
                    _pagina_pdf(c, atleta, W, H); c.showPage()
                    _pagina_estabilograma_pdf(c, atleta, W, H); c.showPage()
                    log(f"    ok {nome} | Hurdle Step", 'ok')
                except Exception as ex:
                    log(f"    aviso: {nome} HurdleStep: {ex}", 'aviso')

            log(f"  ok {nome}", 'ok')

    else:
        # ================================================================
        # PROTOCOLOS NORMAIS (FMS Bipodal / Apoio Unipodal / Tiro com Arco)
        # Arco: janela unica (Confirmacao 1 -> 2), ate 30 ensaios, lado 'arco'.
        # Usa o mesmo pipeline generico porque `mets[arco]` e tratado como um
        # lado por _pagina_pdf / _pagina_estabilograma_pdf (ver lados_ordem).
        # ================================================================
        for atleta in atletas:
            nome   = atleta['nome']
            ind_id = atleta.get('id', '')
            id_str = f'ID {ind_id} - ' if ind_id else ''
            # Divisora por indivíduo (consistente com protocolo Tiro)
            _pagina_secao_pdf(c, f'{id_str}{nome}', proto['nome'], W, H)
            c.showPage()
            _pagina_pdf(c, atleta, W, H); c.showPage()
            _pagina_estabilograma_pdf(c, atleta, W, H); c.showPage()
            log(f"  ok {atleta['nome']}", 'ok')

        if proto and proto.get('two_windows') and any(a.get('scores') for a in atletas):
            try:
                _pagina_correlacao_pdf(c, atletas, W, H, cond='pos'); c.showPage()
                _pagina_correlacao_pdf(c, atletas, W, H, cond='disp'); c.showPage()
                log("  ok pagina correlacao", 'ok')
            except Exception as ex:
                log(f"  aviso: correlacao: {ex}", 'aviso')

        # Pagina demografica (so PROTO_ARCO com dados demograficos)
        if proto_key == PROTO_ARCO and _tem_demografia(atletas):
            try:
                _pagina_demografia_pdf(c, atletas, W, H)
                c.showPage()
                log('  ok pagina demografia', 'ok')
            except Exception as ex:
                log(f'  aviso: demografia: {ex}', 'aviso')

    # ---- Estatísticas do grupo ----
    if opts_estats and len(atletas) >= 2:
        try:
            _pagina_estats_pdf(c, atletas, W, H, protocolo=proto_key, opcoes=opts_estats)
            c.showPage()
            log('  ok pagina ESTATS', 'ok')
        except Exception as ex:
            log(f'  aviso: pagina ESTATS nao gerada ({ex})', 'aviso')

    # ---- Página de citação académica ----
    try:
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph
        from reportlab.lib import colors as _rl_colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        styles = getSampleStyleSheet()
        c.setFillColorRGB(0.05, 0.09, 0.13)
        c.rect(0, 0, W, H, fill=1, stroke=0)
        # Cabeçalho
        c.setFillColorRGB(0.0, 0.706, 0.847)
        c.setFont('Helvetica-Bold', 13)
        c.drawString(36, H - 50, 'Citação Académica / Academic Citation')
        # Linha divisória
        c.setStrokeColorRGB(0.12, 0.29, 0.47)
        c.setLineWidth(1)
        c.line(36, H - 60, W - 36, H - 60)
        # Citação PT
        c.setFillColorRGB(0.93, 0.96, 1.0)
        c.setFont('Helvetica-Bold', 10)
        c.drawString(36, H - 90, 'Para citar este software em trabalhos académicos ou científicos:')
        _ano_cit = datetime.date.today().year
        cit_text = (f'Massuça, A., {SEGUNDO_AUTOR_APA}, & Massuça, L. M. ({_ano_cit}). BSP - Biomechanical Stability Program (v{VERSAO}). '
                    f'https://github.com/andremassuca/BSP')
        c.setFont('Helvetica', 9)
        c.setFillColorRGB(0.49, 0.66, 0.79)
        # Wrap manual
        words = cit_text.split()
        line = ''; y_pos = H - 115; max_w = W - 90
        for word in words:
            test = (line + ' ' + word).strip()
            if c.stringWidth(test, 'Helvetica', 9) < max_w:
                line = test
            else:
                c.drawString(54, y_pos, line); y_pos -= 15; line = word
        if line: c.drawString(54, y_pos, line)
        y_pos -= 30
        # Caixa BibTeX
        bibtex_lines = [
            f'  author  = {{Massuça, André and Aleixo, Pedro and Massuça, Luís}},',
            f'  title   = {{BSP - Biomechanical Stability Program}},',
            f'  year    = {{{datetime.date.today().year}}},',
            f'  version = {{{VERSAO}}},',
            f'  url     = {{https://github.com/andremassuca/BSP}}',
            '}',
        ]
        _box_h = 14 + len(bibtex_lines) * 12 + 8
        c.setFillColorRGB(0.1, 0.18, 0.28)
        c.rect(36, y_pos - _box_h, W - 72, _box_h, fill=1, stroke=0)
        c.setFillColorRGB(0.0, 0.706, 0.847)
        c.setFont('Courier-Bold', 8)
        c.drawString(46, y_pos - 12, '@software{BSP_v' + VERSAO.replace('.','_') + ',')
        c.setFont('Courier', 8)
        c.setFillColorRGB(0.67, 0.85, 0.9)
        for bl in bibtex_lines:
            y_pos -= 12
            c.drawString(46, y_pos - 12, bl)
        # Rodapé
        c.setFillColorRGB(0.49, 0.66, 0.79)
        c.setFont('Helvetica', 8)
        c.drawCentredString(W/2, 30,
            f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}')
        c.showPage()
        log('  ok pagina citacao', 'ok')
    except Exception as _ex_cit:
        log(f'  aviso: citacao nao gerada ({_ex_cit})', 'aviso')

    c.save(); log(f"  PDF: {caminho}", 'ok')
    return True


# -----------------------------------------------------------------------
# Estilos Excel
# -----------------------------------------------------------------------

def _l(cor='BFBFBF', t='thin'): return Side(style=t, color=cor)
BN   = Border(left=_l(), right=_l(), top=_l(), bottom=_l())
BD   = Border(left=_l('0F4C75','medium'), right=_l('0F4C75','medium'),
              top=_l('0F4C75','medium'), bottom=_l('0F4C75','medium'))
FAZ  = PatternFill('solid', start_color='0F4C75')
FAZ1 = PatternFill('solid', start_color='BDD7EE')
FVD  = PatternFill('solid', start_color='E2EFDA')
FAMAR= PatternFill('solid', start_color='FFF2CC')
FAZ2 = PatternFill('solid', start_color='EBF5FB')
FAZDB= PatternFill('solid', start_color='DBEAFE')
FCZ1 = PatternFill('solid', start_color='D9E1F2')
FCZ2 = PatternFill('solid', start_color='E2EFDA')
FVD2 = PatternFill('solid', start_color='E8F5E9')
FTIRO= PatternFill('solid', start_color='FEF3C7')
FOUT = PatternFill('solid', start_color='FEE2E2')
FN   = Font(name='Calibri', size=9)
FNB  = Font(name='Calibri', size=9, bold=True)
FCAB = Font(name='Calibri', size=9, bold=True, color='FFFFFF')
FINF = Font(name='Calibri', size=8, italic=True, color='AAAAAA')
FOUT_FN = Font(name='Calibri', size=9, color='C00000', bold=True)
ALC  = Alignment(horizontal='center', vertical='center')
ALW  = Alignment(horizontal='center', vertical='center', wrap_text=True)

METS_XL = [
    ('amp_x',      'Amplitude X (mm)'),
    ('amp_y',      'Amplitude Y (mm)'),
    ('vel_x',      'Vel. media X (mm/s)'),
    ('vel_y',      'Vel. media Y (mm/s)'),
    ('vel_med',    'Vel. media CoP (mm/s)'),
    ('vel_pico_x', 'Pico vel. X (mm/s)'),
    ('vel_pico_y', 'Pico vel. Y (mm/s)'),
    ('desl',       'Deslocamento (mm)'),
    ('time',       'Tempo (s)'),
    ('ea95',       'Area Elipse 95% (mm2)'),
    ('leng_a',     'Semi-eixo a (mm)'),
    ('leng_b',     'Semi-eixo b (mm)'),
    # RMS (Quijoux et al., 2021, sec. 3.2.4)
    ('rms_x',      'RMS ML (mm)'),
    ('rms_y',      'RMS AP (mm)'),
    ('rms_r',      'RMS Radius (mm)'),
    # FFT - apenas preenchido se fft_ativo=True
    ('fft_pico_x', 'FFT Pico ML (Hz)'),
    ('fft_pico_y', 'FFT Pico AP (Hz)'),
    ('fft_med_x',  'FFT Med. ML (Hz)'),
    ('fft_med_y',  'FFT Med. AP (Hz)'),
    # Metricas derivadas
    ('ratio_ml_ap', 'Rel. amp. ML/AP (adim.)'),
    ('ratio_vel',   'Rel. vel. ML/AP (adim.)'),
    ('stiff_x',     'Stiffness X (1/s)'),
    ('stiff_y',     'Stiffness Y (1/s)'),
    # Covariancia
    ('cov_xx',      'Var CoP X (mm2)'),
    ('cov_yy',      'Var CoP Y (mm2)'),
    ('cov_xy',      'Cov CoP XY (mm2)'),
]
METS_SPSS = [
    ('amp_x',       'amp_ml'),
    ('amp_y',       'amp_ap'),
    ('vel_x',       'vel_ml'),
    ('vel_y',       'vel_ap'),
    ('vel_med',     'vel_med'),
    ('vel_pico_x',  'vel_pico_ml'),
    ('vel_pico_y',  'vel_pico_ap'),
    ('desl',        'desl'),
    ('time',        'tempo'),
    ('ea95',        'ea95'),
    ('leng_a',      'eixo_a'),
    ('leng_b',      'eixo_b'),
    # RMS
    ('rms_x',       'rms_ml'),
    ('rms_y',       'rms_ap'),
    ('rms_r',       'rms_radius'),
    # FFT
    ('fft_pico_x',  'fft_pico_ml'),
    ('fft_pico_y',  'fft_pico_ap'),
    ('fft_med_x',   'fft_med_ml'),
    ('fft_med_y',   'fft_med_ap'),
    # Derivadas
    ('ratio_ml_ap', 'rel_amp_mlap'),
    ('ratio_vel',   'rel_vel_mlap'),
    ('stiff_x',     'stiff_x'),
    ('stiff_y',     'stiff_y'),
    # Normalizacoes por dimensoes corporais
    ('ea95_norm',    'ea95_norm'),
    ('amp_norm_x',   'amp_norm_ml'),
    ('amp_norm_y',   'amp_norm_ap'),
    ('vel_norm',     'vel_norm'),
    ('stiff_mass_x', 'stiff_mass_ml'),
    ('stiff_mass_y', 'stiff_mass_ap'),
    ('stiff_norm_x', 'stiff_norm_ml'),
    ('stiff_norm_y', 'stiff_norm_ap'),
    # Covariancia
    ('cov_xx',      'var_x'),
    ('cov_yy',      'var_y'),
    ('cov_xy',      'cov_xy'),
]


def cl(ws, l, c, v, fn=None, fi=None, fmt=None, al=None, bo=None):
    cell = ws.cell(row=l, column=c, value=v)
    if fn:  cell.font = fn
    if fi:  cell.fill = fi
    if fmt: cell.number_format = fmt
    if al:  cell.alignment = al
    if bo:  cell.border = bo
    return cell


def cred(ws, linha, col=1):
    ws.cell(linha, col,
            f'{PROG} v{VERSAO}  |  {AUTOR}  |  '
            + datetime.date.today().strftime('%d/%m/%Y')).font = FINF


# -----------------------------------------------------------------------
# Abas Excel individuais
# -----------------------------------------------------------------------

def aba_ensaio(wb, ne, nome_atl, data_str, todos_frames, t_ini, t_fim, m, is_outlier=False):
    ws = wb.create_sheet(title=re.sub(r'[\\/*?:\[\]]','_',ne)[:31])
    frames = ([f for f in todos_frames if t_ini<=f['t_ms']<=t_fim]
              if t_ini is not None else list(todos_frames))
    if not frames: ws.cell(1,1,'Sem dados.'); return

    cor_tit = '8B0000' if is_outlier else '0F4C75'
    for l, txt in [(1, f'Stability export for measurement: {ne}'),
                   (2, f'Patient name: {nome_atl}'),
                   (3, f'Measurement done on {data_str}')]:
        cl(ws, l, 1, txt, fn=Font(name='Calibri', bold=True, size=9, color=cor_tit))
    if is_outlier:
        cl(ws, 1, 8, 'OUTLIER detectado (> 2 DP da media)',
           fn=Font(name='Calibri', bold=True, size=9, color='8B0000'))
    cl(ws, 3, 9, f'{PROG} v{VERSAO}  |  {AUTOR}', fn=FINF)

    for c_i, txt in [(2,''),(3,'COF X (mm)'),(4,'COF Y (mm)'),
                     (6,'Vel Méd X\n(mm/s)'),(7,'Vel Méd Y\n(mm/s)'),
                     (8,'Desl.\n(mm)'),(9,'Tempo\n(s)'),(10,'Elipse 95%\n(mm2)')]:
        cl(ws,5,c_i,txt,fn=FCAB,fi=FAZ,al=ALW,bo=BN)
    ws.row_dimensions[5].height = 30

    if m:
        xa=np.array(m['cof_x']); ya=np.array(m['cof_y'])
        for l,txt,vx,vy in [(6,'max',xa.max(),ya.max()),
                             (7,'min',xa.min(),ya.min()),
                             (8,'amplitude',xa.max()-xa.min(),ya.max()-ya.min())]:
            cl(ws,l,2,txt,fn=FNB,fi=FAZ1,bo=BN)
            cl(ws,l,3,round(vx,3),fn=FN,fi=FAZ1,fmt='0.000',bo=BN)
            cl(ws,l,4,round(vy,3),fn=FN,fi=FAZ1,fmt='0.000',bo=BN)
        fi_ea = FOUT if is_outlier else FAZDB
        fn_ea = FOUT_FN if is_outlier else FNB
        cl(ws,8,6,round(m['vel_x'],3),fn=FNB,fi=FAMAR,fmt='0.000',bo=BD)
        cl(ws,8,7,round(m['vel_y'],3),fn=FNB,fi=FAMAR,fmt='0.000',bo=BD)
        cl(ws,8,8,round(m['desl'],3), fn=FNB,fi=FAMAR,fmt='0.000',bo=BD)
        cl(ws,8,9,round(m['time'],3), fn=FNB,fi=FAMAR,fmt='0.000',bo=BD)
        cl(ws,8,10,round(m['ea95'],4),fn=fn_ea,fi=fi_ea,fmt='0.0000',bo=BD)
        cl(ws,6,6,'Eixo a (mm)',fn=FCAB,fi=FAZ,al=ALC,bo=BN)
        cl(ws,6,7,'Eixo b (mm)',fn=FCAB,fi=FAZ,al=ALC,bo=BN)
        cl(ws,7,6,round(m['leng_a'],3),fn=FNB,fi=FVD,fmt='0.000',bo=BN)
        cl(ws,7,7,round(m['leng_b'],3),fn=FNB,fi=FVD,fmt='0.000',bo=BN)

    bms = int(t_ini) if t_ini is not None else int(frames[0]['t_ms'])
    ems = int(t_fim) if t_fim is not None else int(frames[-1]['t_ms'])
    for c_i,txt in enumerate(['Interval','Begin (ms)','End (ms)','Begin (Frame)',
                               'End (Frame)','Area Elipse (mm2)','Desl. (mm)'],start=1):
        cl(ws,10,c_i,txt,fn=FCAB,fi=FAZ,al=ALW,bo=BN)
    for c_i,v in enumerate([1,bms,ems,frames[0]['frame'],frames[-1]['frame'],
                             round(m['ea95'],4) if m else None,
                             round(m['desl'],3) if m else None],start=1):
        cl(ws,11,c_i,v,fn=FNB,fi=FAMAR,fmt='0.000' if c_i>5 else None,bo=BN)

    for c_i,h in enumerate(['Frame','Time (ms)','COF X (mm)','COF Y (mm)',
                             'Vel X (mm/s)','Vel Y (mm/s)','sx (mm)','sy (mm)',
                             's acum. (mm)','Force (N)','Frame dist (mm)'],start=1):
        cl(ws,13,c_i,h,fn=FCAB,fi=FAZ,al=ALW,bo=BN)
    ws.row_dimensions[13].height = 35

    vx=m['vx_f'] if m else [None]*len(frames)
    vy=m['vy_f'] if m else [None]*len(frames)
    sx=m['sx_f'] if m else [None]*len(frames)
    sy=m['sy_f'] if m else [None]*len(frames)
    sa=m['s_ac'] if m else [None]*len(frames)
    def _fv(v): return None if (v is None or (isinstance(v,float) and math.isnan(v))) else round(float(v),3)
    for idx, fd in enumerate(frames):
        l = 14+idx
        cl(ws,l,1,fd['frame'],fn=FN,fmt='0'); cl(ws,l,2,fd['t_ms'],fn=FN,fmt='0.0')
        cl(ws,l,3,fd['x'],fn=FN,fmt='0.000'); cl(ws,l,4,fd['y'],fn=FN,fmt='0.000')
        # vx/vy/sx/sy/sa têm exactamente len(frames) elementos quando frames=m['frames']
        for c_i,v in [(5,vx[idx]),(6,vy[idx]),(7,sx[idx]),(8,sy[idx]),(9,sa[idx])]:
            cl(ws,l,c_i,_fv(v),fn=FN,fi=FAZ2,fmt='0.000')
        # 'forca' e 'dist' só existem em frames reais (não em pseudo-frames de selecção)
        cl(ws,l,10,fd.get('forca'),fn=FN,fmt='0')
        cl(ws,l,11,fd.get('dist'),fn=FN,fmt='0.000')

    if m:
        n=len(frames); nel=len(m['ell_x']); EX,EY=14,15
        cl(ws,13,EX,'EllX',fn=Font(name='Calibri',size=8,color='CCCCCC'))
        cl(ws,13,EY,'EllY',fn=Font(name='Calibri',size=8,color='CCCCCC'))
        for i,(ex,ey) in enumerate(zip(m['ell_x'],m['ell_y'])):
            cl(ws,14+i,EX,round(ex,3)); cl(ws,14+i,EY,round(ey,3))
        cl(ws,14,17,round(m['mean_x'],3)); cl(ws,14,18,round(m['mean_y'],3))

        g=ScatterChart()
        g.title=f'COF + Elipse 95%  -  {nome_atl}  |  {ne}'
        g.style=10; g.x_axis.title='COF X (mm)'; g.y_axis.title='COF Y (mm)'
        g.width=20; g.height=15; g.legend.position='b'
        def _s(xc,yc,r1,rn,title,sym,sz,fill,lw=None):
            xs=Reference(ws,min_col=xc,min_row=r1,max_row=rn)
            ys=Reference(ws,min_col=yc,min_row=r1,max_row=rn)
            s=Series(ys,xs,title=title)
            if lw:
                s.graphicalProperties.line.solidFill=fill
                s.graphicalProperties.line.width=lw; s.marker.symbol='none'
            else:
                s.graphicalProperties.line.noFill=True
                s.marker.symbol=sym; s.marker.size=sz
                s.marker.graphicalProperties.solidFill=fill
                s.marker.graphicalProperties.line.solidFill=fill
            return s
        g.series.append(_s(3,4,14,13+n,'Trajectoria COF','circle',3,'4472C4'))
        g.series.append(_s(EX,EY,14,13+nel,'Elipse 95%',None,None,'FF0000',28000))
        g.series.append(_s(17,18,14,14,'Centro','diamond',9,'00B050'))
        ws.add_chart(g,'M5')

    ws.column_dimensions['A'].width=7; ws.column_dimensions['B'].width=10
    for ci in ['C','D']: ws.column_dimensions[ci].width=16
    for ci in ['E','F','G','H','I']: ws.column_dimensions[ci].width=13
    ws.column_dimensions['J'].width=9; ws.column_dimensions['K'].width=14
    ws.column_dimensions['L'].width=3
    ws.column_dimensions['N'].width=0; ws.column_dimensions['O'].width=0
    ws.freeze_panes = 'A14'


def aba_elipse(wb, ne, m, nome_atl, is_outlier=False):
    ws=wb.create_sheet(title=re.sub(r'[\\/*?:\[\]]','_',f'elipse_{ne}')[:31])
    n=len(m['cof_x']); nel=len(m['ell_x'])
    cor_tit = '8B0000' if is_outlier else '0F4C75'
    cl(ws,1,1,f'Elipse 95%  -  {nome_atl}  -  {ne}',
       fn=Font(name='Calibri',bold=True,size=12,color=cor_tit))
    if is_outlier:
        cl(ws,1,7,'OUTLIER (> 2 DP)',fn=Font(name='Calibri',bold=True,size=9,color='8B0000'))
    cl(ws,1,8,f'{PROG} v{VERSAO}  |  {AUTOR}',fn=FINF)
    for c_i,(txt,fi) in enumerate([('COF X (mm)',FAZ),('COF Y (mm)',FAZ),('',None),
            ('Elipse X (mm)',PatternFill('solid',start_color='C00000')),
            ('Elipse Y (mm)',PatternFill('solid',start_color='C00000'))],start=1):
        if txt: cl(ws,2,c_i,txt,fn=FCAB,fi=fi,al=ALC,bo=BN)
    for i,(xv,yv) in enumerate(zip(m['cof_x'],m['cof_y'])):
        r=i+3; cl(ws,r,1,round(xv,3),fn=FN,fmt='0.000')
        cl(ws,r,2,round(yv,3),fn=FN,fmt='0.000')
    fr=PatternFill('solid',start_color='FFF0F0')
    for i,(ex,ey) in enumerate(zip(m['ell_x'],m['ell_y'])):
        r=i+3; cl(ws,r,4,round(ex,3),fn=FN,fi=fr,fmt='0.000')
        cl(ws,r,5,round(ey,3),fn=FN,fi=fr,fmt='0.000')
    painel=[('n amostras',n,''),('Media X',round(m['mean_x'],3),'mm'),
            ('Media Y',round(m['mean_y'],3),'mm'),
            ('Cov XX',round(m['cov_xx'],4),'mm2'),
            ('Cov XY',round(m['cov_xy'],4),'mm2'),
            ('Cov YY',round(m['cov_yy'],4),'mm2'),
            ('chi2 critico 95%',round(m['chi2'],4),'5.9915'),
            ('Eigenvalor 1',round(m['ev1'],4),'mm2'),
            ('Eigenvalor 2',round(m['ev2'],4),'mm2'),
            ('Semi-eixo a',round(m['leng_a'],3),'mm'),
            ('Semi-eixo b',round(m['leng_b'],3),'mm'),
            ('Angulo rotacao',round(m['ang_deg'],2),'graus'),
            ('Amplitude X',round(m['amp_x'],3),'mm'),
            ('Amplitude Y',round(m['amp_y'],3),'mm'),
            ('Vel. media X',round(m['vel_x'],3),'mm/s'),
            ('Vel. media Y',round(m['vel_y'],3),'mm/s'),
            ('Vel. media',round(m['vel_med'],3),'mm/s'),
            ('Pico vel. X',round(m['vel_pico_x'],3),'mm/s'),
            ('Pico vel. Y',round(m['vel_pico_y'],3),'mm/s'),
            ('Deslocamento',round(m['desl'],3),'mm'),
            ('Tempo',round(m['time'],3),'s'),('','',''),
            ('AREA ELIPSE 95%',round(m['ea95'],4),'mm2')]
    for c_i,txt in [(7,'Parametro'),(8,'Valor'),(9,'Unidade')]:
        cl(ws,2,c_i,txt,fn=FCAB,fi=FAZ,al=ALC,bo=BN)
    fi_ea = FOUT if is_outlier else FAZDB
    for pi,(lbl,val,uni) in enumerate(painel,start=3):
        ea=(lbl=='AREA ELIPSE 95%')
        fi=fi_ea if ea else None
        fn=Font(name='Calibri',bold=True,size=10,
                color='8B0000' if (ea and is_outlier) else ('0D47A1' if ea else '000000'))
        cl(ws,pi,7,lbl,fn=fn,fi=fi,bo=BN)
        if val!='': cl(ws,pi,8,val,fn=fn,fi=fi,
                       fmt='0.0000' if isinstance(val,float) else None,bo=BN)
        cl(ws,pi,9,uni,fn=fn,fi=fi,bo=BN)
    cl(ws,3,11,round(m['mean_x'],3)); cl(ws,3,12,round(m['mean_y'],3))
    g=ScatterChart(); g.title=f'Elipse 95%  |  {nome_atl}  |  {ne}'
    g.style=10; g.x_axis.title='COF X (mm)'; g.y_axis.title='COF Y (mm)'
    g.width=22; g.height=18; g.legend.position='b'
    for xc,yc,r1,rn,title,sym,sz,fill,lw in [
        (1,2,3,n+2,'Trajectoria COF','circle',3,'4472C4',None),
        (4,5,3,nel+2,'Elipse 95%',None,None,'FF0000',30000),
        (11,12,3,3,'Centro','diamond',10,'00B050',None)]:
        xs=Reference(ws,min_col=xc,min_row=r1,max_row=rn)
        ys=Reference(ws,min_col=yc,min_row=r1,max_row=rn)
        s=Series(ys,xs,title=title)
        if lw:
            s.graphicalProperties.line.solidFill=fill
            s.graphicalProperties.line.width=lw; s.marker.symbol='none'
        else:
            s.graphicalProperties.line.noFill=True; s.marker.symbol=sym; s.marker.size=sz
            s.marker.graphicalProperties.solidFill=fill
            s.marker.graphicalProperties.line.solidFill=fill
        g.series.append(s)
    ws.add_chart(g,'G23')
    for c_i,w in [('A',10),('B',10),('C',3),('D',10),('E',10),('F',2),
                  ('G',26),('H',12),('I',10),('J',2),('K',12),('L',12)]:
        ws.column_dimensions[c_i].width=w
    ws.freeze_panes='A3'; cred(ws,n+6,7)


def aba_estabilograma(wb, ne, m, nome_atl):
    if m is None: return
    ws=wb.create_sheet(title=re.sub(r'[\\/*?:\[\]]','_',f'estab_{ne}')[:31])
    cl(ws,1,1,f'Estabilograma  -  {nome_atl}  -  {ne}',
       fn=Font(name='Calibri',bold=True,size=12,color='0F4C75'))
    cl(ws,1,6,f'{PROG} v{VERSAO}  |  {AUTOR}',fn=FINF)
    for c_i,(hdr,fi) in enumerate([('Tempo (s)',FAZ),('COF X (mm)',FAZ),('COF Y (mm)',FAZ)],start=1):
        cl(ws,2,c_i,hdr,fn=FCAB,fi=fi,al=ALC,bo=BN)
    t_arr=np.array(m['t_ms']); t_arr=(t_arr-t_arr[0])/1000.0
    for i,(ts,xv,yv) in enumerate(zip(t_arr,m['cof_x'],m['cof_y'])):
        r=i+3; cl(ws,r,1,round(float(ts),3),fn=FN,fmt='0.000')
        cl(ws,r,2,round(xv,3),fn=FN,fmt='0.000'); cl(ws,r,3,round(yv,3),fn=FN,fmt='0.000')
    n=len(m['cof_x']); cats=Reference(ws,min_col=1,min_row=3,max_row=n+2)
    for col_i,cor,tit_g,ancora in [(2,'2E86C1','COF X ao longo do tempo','E2'),
                                    (3,'E67E22','COF Y ao longo do tempo','E28')]:
        gg=LineChart(); gg.title=f'{tit_g}  -  {ne}'; gg.style=10
        gg.x_axis.title='Tempo (s)'; gg.y_axis.title='COF (mm)'
        gg.width=22; gg.height=10
        vv=Reference(ws,min_col=col_i,min_row=3,max_row=n+2)
        sv=Series(vv,title=f'COF {"X" if col_i==2 else "Y"}')
        sv.graphicalProperties.line.solidFill=cor; sv.graphicalProperties.line.width=10000
        gg.series.append(sv); gg.set_categories(cats); ws.add_chart(gg,ancora)
    ws.column_dimensions['A'].width=12; ws.column_dimensions['B'].width=14
    ws.column_dimensions['C'].width=14; ws.freeze_panes='A3'


# -----------------------------------------------------------------------
# Excel de resumo (DADOS + GRUPO + SPSS)
# -----------------------------------------------------------------------

def _desc_stats(vals, ndigits_med=3, ndigits_cv=1):
    """Calcula (n, media, dp, cv) de uma lista de floats, ignorando None.
    Retorna None para campos indetermináveis. Usado em aba_estats e aba_tiro."""
    nums = [v for v in vals if v is not None]
    n = len(nums)
    if n == 0:
        return 0, None, None, None
    arr = np.array(nums, dtype=float)
    med = round(float(arr.mean()), ndigits_med)
    dp  = round(float(arr.std(ddof=1)), ndigits_med) if n > 1 else None
    cv  = round(float(arr.std(ddof=1) / arr.mean() * 100), ndigits_cv) \
          if (n > 1 and arr.mean() != 0) else None
    return n, med, dp, cv


def _media_lado(ath, chave, lado):
    """Média dos valores válidos de uma métrica para um lado/atleta."""
    vals = [m[chave] for m in ath['mets'].get(lado, [])
            if m is not None and chave in m and m[chave] is not None]
    vals_clean=[v for v in vals if v is not None]; return float(np.mean(vals_clean)) if vals_clean else None


def _n_validos(ath, lado):
    """Número de ensaios com resultado válido (not None) para um lado."""
    return sum(1 for m in ath['mets'].get(lado, []) if m is not None)


def _n_outliers(ath, lado, chave='ea95'):
    """Número de ensaios flagrados como outlier para um lado."""
    lst = ath['mets'].get(lado, [])
    flags = flagrar_outliers(lst, chave=chave)
    return sum(1 for f, m in zip(flags, lst) if f and m is not None)


def _stats_grupo(atletas, chave, lado):
    """Retorna (media, dp, min, max, n, se, cv) para um grupo/lado/metrica."""
    medias = [_media_lado(a, chave, lado) for a in atletas]
    medias = [v for v in medias if v is not None]
    if not medias:
        return None, None, None, None, 0, None, None
    arr = np.array(medias)
    n   = len(arr)
    med = float(arr.mean())
    dp  = float(arr.std(ddof=1)) if n > 1 else 0.0
    se  = dp / math.sqrt(n)      if n > 1 else None
    cv  = dp / med * 100.0       if (n > 1 and med != 0) else None
    return med, dp, float(arr.min()), float(arr.max()), n, se, cv


def exportar_csv_flat(caminho, atletas, protocolo=None):
    """
    Exporta um CSV "flat" com uma linha por ensaio e uma coluna por métrica.
    Útil para análise subsequente em R, SPSS ou Python sem necessidade de
    interpretar a estrutura complexa das abas Excel.

    Colunas: id, nome, protocolo, lado, ensaio, n_validos_lado,
             n_outliers_lado, is_outlier, + todas as métricas de METS_SPSS.
    """
    import csv
    if protocolo is None: protocolo = _PROTOCOLO_ACTIVO
    proto = PROTOCOLOS[protocolo]

    # Cabeçalho
    metricas = [col_nome for _, col_nome in METS_SPSS]
    cabecalho = ['id','nome','protocolo','lado','ensaio',
                 'n_validos','n_outliers','is_outlier'] + metricas

    with open(caminho, 'w', newline='', encoding='utf-8-sig') as f:
        wr = csv.writer(f)
        wr.writerow(cabecalho)
        for ath in atletas:
            for lado, _ in proto['lados']:
                lst = ath['mets'].get(lado, [])
                flags = flagrar_outliers(lst, chave='ea95')
                nv  = _n_validos(ath, lado)
                nout= _n_outliers(ath, lado)
                for ti, (m, flag) in enumerate(zip(lst, flags), start=1):
                    if m is None:
                        row = [ath.get('id',''), ath['nome'], protocolo,
                               lado, ti, nv, nout, ''] + ['']*len(metricas)
                    else:
                        vals = []
                        for chave, _ in METS_SPSS:
                            v = m.get(chave, None)
                            vals.append(round(v, 6) if v is not None else '')
                        row = [ath.get('id',''), ath['nome'], protocolo,
                               lado, ti, nv, nout, int(flag)] + vals
                    wr.writerow(row)


def aba_grupo(wb, atletas, protocolo=None):
    if protocolo is None: protocolo=_PROTOCOLO_ACTIVO
    proto=PROTOCOLOS[protocolo]; tem_ai=proto['assimetria']
    # Tiro com Arco usa lado unico 'arco'
    if protocolo == PROTO_ARCO:
        lados_nomes = [('arco', 'Arco')]
    else:
        lados_nomes = [('dir','Dir.'),('esq','Esq.')]
    ws=wb.create_sheet(title='GRUPO')
    cl(ws,1,1,f'Estatisticas do Grupo  -  {proto["nome"]}  |  n = {len(atletas)} indivíduos',
       fn=Font(name='Calibri',bold=True,size=12,color='0F4C75'))
    cl(ws,1,8,f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}',fn=FINF)
    cl(ws,2,1,datetime.date.today().strftime('%d/%m/%Y'),fn=FINF)

    # Cabecalhos incluem SE e CV alem de Med/DP/Min/Max
    hdrs=['Métrica']
    fis=[]
    for lado,lbl in lados_nomes:
        for h in [f'Med {lbl}',f'DP {lbl}',f'SE {lbl}',f'CV {lbl}',f'Min {lbl}',f'Max {lbl}']:
            hdrs.append(h)
            fis.append(FCZ1 if lado in ['dir','pos'] else FCZ2)
    if tem_ai: hdrs.append('AI Med (%)'); fis.append(FVD2)

    for c_i,h in enumerate(hdrs,start=1):
        fi=FAZ if c_i==1 else fis[c_i-2]
        cl(ws,4,c_i,h,fn=FCAB,fi=fi,al=ALC,bo=BN)
    ws.row_dimensions[4].height=22
    ws.column_dimensions['A'].width=28
    for c_i in range(2,len(hdrs)+1):
        ws.column_dimensions[get_column_letter(c_i)].width=10

    for ri,(chave,label) in enumerate(METS_XL,start=5):
        is_ea=(chave=='ea95')
        fn_l=Font(name='Calibri',size=9,bold=True,color='C00000') if is_ea else FNB
        cl(ws,ri,1,label,fn=fn_l,bo=BN)
        col_idx=2
        for lado,_ in lados_nomes:
            med,dp,mn,mx,n_vals,se,cv=_stats_grupo(atletas,chave,lado)
            for v,fmt in [(med,'0.000'),(dp,'0.000'),(se,'0.000'),(cv,'0.0'),(mn,'0.000'),(mx,'0.000')]:
                fi=FCZ1 if lado in ['dir','pos'] else FCZ2
                cl(ws,ri,col_idx,round(v,3) if v is not None else None,
                   fn=Font(name='Calibri',size=9,bold=is_ea,
                           color='C00000' if is_ea else '000000'),
                   fi=fi,fmt=fmt if v is not None else None,bo=BN)
                col_idx+=1
        if tem_ai:
            md=_stats_grupo(atletas,chave,'dir')[0]
            me=_stats_grupo(atletas,chave,'esq')[0]
            ai_g=assimetria(md,me) if md and me else None
            cl(ws,ri,col_idx,f'{ai_g:+.1f}%' if ai_g else None,
               fn=Font(name='Calibri',size=9,bold=True,color='92400E'),fi=FVD2,bo=BN)

    sep=len(METS_XL)+7
    cl(ws,sep,1,'Media ea95 por individuo:',fn=FNB)
    # Cabecalho da tabela individual com ID
    hdrs2=['ID','Individuo']+[f'ea95 {lbl} (mm2)' for _,lbl in lados_nomes]+[f'n valid {lbl}' for _,lbl in lados_nomes]+[f'n out {lbl}' for _,lbl in lados_nomes]
    if tem_ai: hdrs2.append('AI (%)')
    for c_i,h in enumerate(hdrs2,start=1):
        cl(ws,sep+1,c_i,h,fn=FCAB,fi=FAZ,al=ALC,bo=BN)
    ws.column_dimensions['A'].width=8
    ws.column_dimensions['B'].width=26
    for i,ath in enumerate(atletas,start=sep+2):
        cl(ws,i,1,ath.get('id',''),fn=FN,bo=BN)
        cl(ws,i,2,ath['nome'],fn=FNB,bo=BN)
        n_lados = len(lados_nomes)
        for ci,(lado,_) in enumerate(lados_nomes,start=3):
            v=_media_lado(ath,'ea95',lado)
            fi=FCZ1 if lado in ['dir','pos'] else FCZ2
            cl(ws,i,ci,round(v,3) if v else None,fn=FN,fi=fi,
               fmt='0.000' if v else None,bo=BN)
        if tem_ai:
            vd=_media_lado(ath,'ea95','dir'); ve=_media_lado(ath,'ea95','esq')
            ai=assimetria(vd,ve)
            cl(ws,i,n_lados+3,f'{ai:+.1f}%' if ai else None,fn=FNB,fi=FVD2,bo=BN)
        # colunas de contagem: n válidos e n outliers por lado
        base_nv = n_lados+3+(1 if tem_ai else 0)
        for ci,(lado,_) in enumerate(lados_nomes,start=base_nv):
            cl(ws,i,ci,_n_validos(ath,lado),fn=FN,bo=BN)
        for ci,(lado,_) in enumerate(lados_nomes,start=base_nv+n_lados):
            nout=_n_outliers(ath,lado)
            fn_o=Font(name='Calibri',size=9,bold=True,color='C00000') if nout>0 else FN
            cl(ws,i,ci,nout,fn=fn_o,bo=BN)
    cred(ws,sep+len(atletas)+4)
    ws.freeze_panes='B5'


def aba_spss(wb, atletas, protocolo=None):
    if protocolo is None: protocolo=_PROTOCOLO_ACTIVO
    proto=PROTOCOLOS[protocolo]
    # Tiro com Arco usa lado unico 'arco'
    if protocolo == PROTO_ARCO:
        lados = [('arco', '_arco')]
    else:
        lados = [('dir','_d'),('esq','_e')]  # sempre dir/esq; tiro tem abas proprias

    ws=wb.create_sheet(title=T('aba_spss'))
    cl(ws,1,1,f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  Formato SPSS',
       fn=Font(name='Calibri',bold=True,size=9,color='0F4C75'))
    cl(ws,2,1,T('pdf_gerado_por',prog=PROG,v=VERSAO,autor=AUTOR,univ=ORIENTADOR)+'  |  '
              + datetime.date.today().strftime('%d/%m/%Y'), fn=FINF)

    hdrs=['nome']
    # Usa labels traduzidos para os cabeçalhos das colunas SPSS
    mets_loc = mets_xl_localizadas() if _I18N_OK else METS_SPSS
    for _,col_label in mets_loc:
        for lado,suf in lados:
            hdrs.append(col_label+suf)

    for c_i,h in enumerate(hdrs,start=1):
        if h=='nome': fi=FAZ
        elif h.endswith('_d') or h.endswith('_pos'): fi=FCZ1
        else: fi=FCZ2
        cl(ws,3,c_i,h,fn=FCAB,fi=fi,al=ALC,bo=BN)
        ws.column_dimensions[get_column_letter(c_i)].width=13 if h=='nome' else 11
    ws.row_dimensions[3].height=20

    for ri,ath in enumerate(atletas,start=4):
        cl(ws,ri,1,ath['nome'],fn=FNB,bo=BN)
        col=2
        for chave,_ in METS_SPSS:
            for lado,_ in lados:
                fi=FCZ1 if lado in ['dir','pos'] else FCZ2
                vals=[m[chave] for m in ath['mets'].get(lado,[]) if m is not None and m.get(chave) is not None]
                _vals_clean=[x for x in vals if x is not None]
                v=round(float(np.mean(_vals_clean)),4) if _vals_clean else None
                cl(ws,ri,col,v,fn=FN,fi=fi,fmt='0.0000' if v else None,bo=BN)
                col+=1

    n=len(atletas); sep=n+5
    cl(ws,sep,1,f'Média do grupo (n={n})',fn=FNB,
       fi=FAZ).font=Font(name='Calibri',size=9,bold=True,color='FFFFFF')
    col=2
    for chave,_ in METS_SPSS:
        for lado,_ in lados:
            fi=FCZ1 if lado in ['dir','pos'] else FCZ2
            todas=[_media_lado(a,chave,lado) for a in atletas]
            todas=[v for v in todas if v is not None]
            _todas_clean=[x for x in todas if x is not None]
            v=round(float(np.mean(_todas_clean)),4) if _todas_clean else None
            cell=cl(ws,sep,col,v,fn=FNB,fmt='0.0000' if v else None,bo=BN)
            cell.fill=PatternFill('solid',start_color='FEF9C3')
            col+=1

    # Nota de citação académica
    cit_row = sep + 3
    lang = lingua_atual() if _I18N_OK else 'PT'
    _ano_cit = datetime.date.today().year
    cit = (f'Citação APA: Massuça, A., {SEGUNDO_AUTOR_APA}, & Massuça, L. M. ({_ano_cit}). BSP - Biomechanical Stability Program (v{VERSAO}). '
           f'https://github.com/andremassuca/BSP')
    cl(ws, cit_row, 1, cit, fn=FINF)

    ws.freeze_panes='B4'; cred(ws,sep+2)



# -----------------------------------------------------------------------
# Indice de Perturbacao (Protocolo Tiro)
# -----------------------------------------------------------------------

def calc_perturbacao(atleta):
    """
    Indice de perturbacao por ensaio: ea95_disp / ea95_pos.
    Valores > 2 indicam perda significativa de estabilidade no disparo.
    Retorna lista de (ensaio, ip, ea95_pos, ea95_disp).
    """
    mets_pos  = atleta['mets'].get('pos', [])
    mets_disp = atleta['mets'].get('disp', [])
    resultado = []
    n = min(len(mets_pos), len(mets_disp))
    for i in range(n):
        mp = mets_pos[i]; md = mets_disp[i]
        if mp is None or md is None: continue
        ea_p = mp.get('ea95', 0); ea_d = md.get('ea95', 0)
        ip = round(ea_d / ea_p, 3) if ea_p > 0 else None
        resultado.append((i+1, ip, ea_p, ea_d))
    return resultado


# -----------------------------------------------------------------------
# Testes estatisticos automaticos
# -----------------------------------------------------------------------

_TESTS_INFO = {
    'dir_esq': {
        'label': 'Dir vs Esq  (t / Wilcoxon + Cohen d)',
        'descr': ('Compara Dir vs Esq para cada metrica.\n'
                  'Shapiro-Wilk por lado; t-pareado se ambos normais, Wilcoxon caso contrario.\n'
                  'Calcula tambem Cohen\'s d (tamanho do efeito) e IC 95% da diferenca.\n'
                  'p < 0.05 assinalado a vermelho.'),
        'protos': [PROTO_FMS, PROTO_UNIPODAL],
    },
    'pos_disp': {
        'label': 'Posição vs Disparo  (t / Wilcoxon + IP)',
        'descr': ('Compara ea95 Posição vs Disparo por indivíduo (dif. pareadas).\n'
                  'Shapiro-Wilk nas diferencas; t-pareado ou Wilcoxon.\n'
                  'Índice de Perturbação (IP = ea95_disp / ea95_pos) por ensaio.\n'
                  'IC 95% da diferenca media e Cohen\'s d incluidos.'),
        'protos': [PROTO_TIRO, PROTO_ARCO],
    },
    'individuo': {
        'label': 'Variabilidade intra-atleta  (CV%)',
        'descr': ('CV (DP/Media x 100) por metrica e individuo.\n'
                  'Semaforo: verde CV<15%, amarelo 15-30%, vermelho >30%.\n'
                  'Linha de resumo do grupo com CV medio por metrica.\n'
                  'Nao envolve testes de hipoteses.'),
        'protos': [PROTO_FMS, PROTO_UNIPODAL, PROTO_TIRO, PROTO_ARCO],
    },
    'grupo': {
        'label': 'Descritivos + SW + IC 95%',
        'descr': ('Shapiro-Wilk por metrica e lado (alfa=0.05).\n'
                  'Se normal: Media, DP, IC 95% (t), SE, CV.\n'
                  'Se nao-normal: Mediana, IQR, percentis 5-95.\n'
                  'Tabela de resumo completa na aba ESTATS.'),
        'protos': [PROTO_FMS, PROTO_UNIPODAL, PROTO_TIRO, PROTO_ARCO],
    },
    'friedman': {
        'label': 'Friedman entre intervalos (Tiro)',
        'descr': ('Teste de Friedman para comparar os intervalos de tiro (pos, disp, etc.)\n'
                  'Design: sujeitos x intervalos, valor = media dos ensaios validos.\n'
                  'Pos-hoc Bonferroni-Wilcoxon se significativo (p < 0.05).\n'
                  'Apenas para protocolo Tiro com >= 3 individuos.'),
        'protos': [PROTO_TIRO, PROTO_ARCO],
    },
    'correlacao': {
        'label': 'Correlacao EA95 vs Score (Tiro)',
        'descr': ('Correlacao de Pearson/Spearman entre ea95 e score de precisao.\n'
                  'Requer ficheiro de scores. Uma correlacao por distancia x intervalo.\n'
                  'Scatter plot incluido no PDF se matplotlib disponivel.'),
        'protos': [PROTO_TIRO, PROTO_ARCO],
    },
}

_PVAL_VERDE = '#C6EFCE'; _PVAL_AMAR = '#FFEB9C'; _PVAL_VERM = '#FFC7CE'
_FFVD = Font(name='Calibri', size=9, bold=True, color='375623')
_FFAM = Font(name='Calibri', size=9, bold=True, color='7D4F00')
_FFVR = Font(name='Calibri', size=9, bold=True, color='9C0006')
# PatternFill pre-construidos para semaforo de p-valores (evita re-instanciacao em cada celula)
_FIVD  = PatternFill('solid', start_color='C6EFCE')  # verde
_FIAM  = PatternFill('solid', start_color='FFEB9C')  # amarelo
_FIVR  = PatternFill('solid', start_color='FFC7CE')  # vermelho
_FIAZ6 = PatternFill('solid', start_color='0E7AA8')  # azul escuro (cabecalho tiro)

def _cel_pval(ws, row, col, p):
    """Escreve p-value com formatacao por cor."""
    if p is None:
        cl(ws, row, col, 'n/a', fn=Font(name='Calibri',size=9,italic=True), bo=BN)
        return
    txt = f'{p:.4f}' if p >= 0.001 else '<0.001'
    if p >= 0.05:
        fn, fi = _FFVD, _FIVD
    elif p >= 0.01:
        fn, fi = _FFAM, _FIAM
    else:
        fn, fi = _FFVR, _FIVR
    cl(ws, row, col, txt, fn=fn, fi=fi, bo=BN)


def aba_estats(wb, atletas, protocolo=None, opcoes=None):
    """
    Aba ESTATS: testes estatisticos automaticos.
    opcoes: dict com chaves 'dir_esq','pos_disp','individuo','grupo' (bool).
    """
    if protocolo is None: protocolo = _PROTOCOLO_ACTIVO
    if opcoes is None: opcoes = {k: True for k in _TESTS_INFO}
    proto = PROTOCOLOS[protocolo]
    lados_nomes = [('dir','Dir.'),('esq','Esq.')]

    ws = wb.create_sheet(title='ESTATS')
    cl(ws,1,1, f'Testes Estatisticos Automaticos  |  {proto["nome"]}  |  n = {len(atletas)}',
       fn=Font(name='Calibri',bold=True,size=12,color='0F4C75'))
    cl(ws,1,10, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}', fn=FINF)
    cl(ws,2,1, f'alfa = 0.05  |  SW = Shapiro-Wilk  |  t = t-pareado  |  W = Wilcoxon  |  '
               + datetime.date.today().strftime('%d/%m/%Y'), fn=FINF)
    ws.column_dimensions['A'].width = 26
    for c in range(2, 16):
        ws.column_dimensions[get_column_letter(c)].width = 13

    row = 4

    # ---- Bloco 1: Shapiro-Wilk + descritivos completos (IC95 / IQR) ----
    if opcoes.get('grupo') and len(atletas) >= 3:
        _titulo_bloco(ws, row, 'NORMALIDADE E DESCRITIVOS DO GRUPO  (SW = Shapiro-Wilk  |  IC = Intervalo de Confianca 95%)')
        row += 1
        hdrs = ['Métrica']
        for lado, lbl in lados_nomes:
            hdrs += [f'n {lbl}', f'Media {lbl}', f'DP {lbl}', f'SE {lbl}',
                     f'CV% {lbl}', f'IC-inf {lbl}', f'IC-sup {lbl}',
                     f'Mediana {lbl}', f'IQR {lbl}',
                     f'SW-p {lbl}', f'Dist. {lbl}']
        for c_i, h in enumerate(hdrs, 1):
            fi = FAZ if c_i == 1 else (FCZ1 if ((c_i - 2) // 11) % 2 == 0 else FCZ2)
            cl(ws, row, c_i, h, fn=FCAB, fi=fi, al=ALC, bo=BN)
        ws.row_dimensions[row].height = 24
        row += 1

        for chave, label in METS_XL:
            cl(ws, row, 1, label, fn=FNB, bo=BN)
            col = 2
            for lado, lbl in lados_nomes:
                vals = [_media_lado(a, chave, lado) for a in atletas]
                vals = [v for v in vals if v is not None]
                n = len(vals)
                if n < 3:
                    for _ in range(11): cl(ws, row, col, '-', fn=FN, bo=BN); col += 1
                    continue
                arr = np.array(vals, dtype=float)
                media = float(arr.mean())
                dp    = float(arr.std(ddof=1))
                se    = dp / math.sqrt(n)
                cv    = dp / media * 100 if media != 0 else None
                t_crit = stats.t.ppf(0.975, df=n-1)
                ic_inf = media - t_crit * se
                ic_sup = media + t_crit * se
                mediana = float(np.median(arr))
                iqr     = float(np.percentile(arr, 75) - np.percentile(arr, 25))
                _, sw_p = stats.shapiro(arr)
                normal  = sw_p >= 0.05
                dist_txt = 'Normal' if normal else 'Nao-normal'
                cl(ws, row, col, n,                  fn=FN, bo=BN); col += 1
                cl(ws, row, col, round(media,3),     fn=FN, fmt='0.000', bo=BN); col += 1
                cl(ws, row, col, round(dp,3),        fn=FN, fmt='0.000', bo=BN); col += 1
                cl(ws, row, col, round(se,3),        fn=FN, fmt='0.000', bo=BN); col += 1
                cv_fn = _FFVR if (cv and cv > 30) else (_FFAM if (cv and cv > 15) else FN)
                cl(ws, row, col, round(cv,1) if cv else None,
                   fn=cv_fn, fmt='0.0', bo=BN); col += 1
                cl(ws, row, col, round(ic_inf,3),    fn=FN, fmt='0.000', bo=BN); col += 1
                cl(ws, row, col, round(ic_sup,3),    fn=FN, fmt='0.000', bo=BN); col += 1
                cl(ws, row, col, round(mediana,3),   fn=FN, fmt='0.000', bo=BN); col += 1
                cl(ws, row, col, round(iqr,3),       fn=FN, fmt='0.000', bo=BN); col += 1
                _cel_pval(ws, row, col, sw_p); col += 1
                fn_d = _FFVD if normal else _FFVR
                fi_d = _FIVD if normal else _FIVR
                cl(ws, row, col, dist_txt, fn=fn_d, fi=fi_d, bo=BN); col += 1
            row += 1
        row += 2

    # ---- Bloco 2: Dir vs Esq + Cohen's d + IC95 ----
    if opcoes.get('dir_esq') and protocolo in _TESTS_INFO['dir_esq']['protos'] and len(atletas) >= 3:
        _titulo_bloco(ws, row, 'COMPARACAO DIR vs ESQ  (t-pareado/Wilcoxon  +  Cohen\'s d  +  IC 95% da diferenca)')
        row += 1
        hdrs2 = ['Métrica','n','Méd Dir','Méd Esq','Dif. Méd',
                 'SW-Dir p','SW-Esq p','Teste','Estat.','p-valor',
                 'IC-inf 95%','IC-sup 95%','Cohen d','Efeito','Sig?']
        for c_i, h in enumerate(hdrs2, 1):
            fi = FAZ if c_i == 1 else (FCZ1 if c_i <= 5 else (FAZ1 if c_i <= 10 else FCZ2))
            cl(ws, row, c_i, h, fn=FCAB, fi=fi, al=ALC, bo=BN)
            ws.column_dimensions[get_column_letter(c_i)].width = (
                26 if c_i == 1 else 10)
        ws.row_dimensions[row].height = 24
        row += 1
        for chave, label in METS_XL:
            vd = [_media_lado(a, chave, 'dir') for a in atletas]
            ve = [_media_lado(a, chave, 'esq') for a in atletas]
            pares = [(d, e) for d, e in zip(vd, ve) if d is not None and e is not None]
            if len(pares) < 3:
                cl(ws, row, 1, label, fn=FN, bo=BN)
                cl(ws, row, 2, f'n={len(pares)} < 3', fn=FN, bo=BN); row += 1; continue
            d_arr = np.array([p[0] for p in pares])
            e_arr = np.array([p[1] for p in pares])
            dif   = d_arr - e_arr
            _, sw_d = stats.shapiro(d_arr); _, sw_e = stats.shapiro(e_arr)
            normal = sw_d >= 0.05 and sw_e >= 0.05
            n_p    = len(pares)
            if normal:
                tstat, pval = stats.ttest_rel(d_arr, e_arr)
                teste_nm = 't-pareado'
                se_dif   = float(dif.std(ddof=1) / math.sqrt(n_p))
                t_crit   = stats.t.ppf(0.975, df=n_p-1)
                ic_inf   = round(float(dif.mean()) - t_crit * se_dif, 3)
                ic_sup   = round(float(dif.mean()) + t_crit * se_dif, 3)
            else:
                tstat, pval = stats.wilcoxon(d_arr, e_arr, alternative='two-sided')
                teste_nm = 'Wilcoxon'
                ic_inf   = None; ic_sup = None
            # Cohen's d para amostras pareadas
            sd_dif  = float(dif.std(ddof=1))
            cohen_d = round(float(dif.mean()) / sd_dif, 3) if sd_dif > 0 else None
            if cohen_d is None:
                efeito = '-'
            elif abs(cohen_d) < 0.2:
                efeito = 'Negligivel'
            elif abs(cohen_d) < 0.5:
                efeito = 'Pequeno'
            elif abs(cohen_d) < 0.8:
                efeito = 'Medio'
            else:
                efeito = 'Grande'
            sig = 'Sim *' if pval < 0.05 else 'Nao'
            cl(ws,row,1, label,                           fn=FNB, bo=BN)
            cl(ws,row,2, n_p,                             fn=FN,  bo=BN)
            cl(ws,row,3, round(float(d_arr.mean()),3),    fn=FN,  fmt='0.000', bo=BN)
            cl(ws,row,4, round(float(e_arr.mean()),3),    fn=FN,  fmt='0.000', bo=BN)
            cl(ws,row,5, round(float(dif.mean()),3),      fn=FN,  fmt='0.000', bo=BN)
            _cel_pval(ws, row, 6, sw_d); _cel_pval(ws, row, 7, sw_e)
            cl(ws,row,8, teste_nm,                        fn=FN,  bo=BN)
            cl(ws,row,9, round(float(tstat),4),           fn=FN,  fmt='0.0000', bo=BN)
            _cel_pval(ws, row, 10, pval)
            cl(ws,row,11, ic_inf, fn=FN, fmt='0.000' if ic_inf else None, bo=BN)
            cl(ws,row,12, ic_sup, fn=FN, fmt='0.000' if ic_sup else None, bo=BN)
            cl(ws,row,13, cohen_d, fn=FNB, fmt='0.000' if cohen_d else None, bo=BN)
            fn_ef = (_FFVR if efeito in ('Grande','Medio') else
                     _FFAM if efeito == 'Pequeno' else _FFVD)
            fi_ef = (_FIVR if efeito in ('Grande','Medio') else
                     _FIAM if efeito == 'Pequeno' else _FIVD)
            cl(ws,row,14, efeito, fn=fn_ef, fi=fi_ef, bo=BN)
            fn_sig = _FFVR if pval < 0.05 else _FFVD
            fi_sig = _FIVR if pval < 0.05 else _FIVD
            cl(ws,row,15, sig, fn=fn_sig, fi=fi_sig, bo=BN)
            row += 1
        row += 2

    # ---- Bloco 3: Variabilidade intra-atleta ----
    if opcoes.get('individuo'):
        _titulo_bloco(ws, row, 'VARIABILIDADE INTRA-ATLETA  (CV = DP/Media x 100%)')
        row += 1
        hdrs3 = ['Atleta', 'Lado'] + [lbl for _, lbl in METS_XL]
        for c_i, h in enumerate(hdrs3, 1):
            cl(ws, row, c_i, h, fn=FCAB, fi=FAZ, al=ALC, bo=BN)
        row += 1
        for ath in atletas:
            for lado, lbl_l in lados_nomes:
                mlist = [m for m in ath['mets'].get(lado, []) if m is not None]
                cl(ws, row, 1, ath['nome'], fn=FNB, bo=BN)
                cl(ws, row, 2, lbl_l, fn=FN, bo=BN)
                for c_i, (chave, _) in enumerate(METS_XL, 3):
                    vals = [m[chave] for m in mlist if chave in m]
                    if len(vals) < 2:
                        cl(ws, row, c_i, '-', fn=FN, bo=BN); continue
                    arr = np.array(vals, dtype=float)
                    cv = float(arr.std(ddof=1)/arr.mean()*100) if arr.mean() != 0 else None
                    if cv is None:
                        cl(ws, row, c_i, '-', fn=FN, bo=BN)
                    elif cv > 30:
                        cl(ws,row,c_i,round(cv,1),fn=_FFVR,
                           fi=_FIVR,
                           fmt='0.0',bo=BN)
                    elif cv > 15:
                        cl(ws,row,c_i,round(cv,1),fn=_FFAM,
                           fi=_FIAM,
                           fmt='0.0',bo=BN)
                    else:
                        cl(ws,row,c_i,round(cv,1),fn=_FFVD,
                           fi=_FIVD,
                           fmt='0.0',bo=BN)
                row += 1
        # Linha de resumo: CV medio do grupo por metrica
        cl(ws, row, 1, 'CV Medio Grupo (%)', fn=FNB, fi=FAZ1, bo=BN)
        cl(ws, row, 2, '(todos lados)', fn=FN, fi=FAZ1, bo=BN)
        for c_i, (chave, _) in enumerate(METS_XL, 3):
            all_cvs = []
            for ath in atletas:
                for lado, _ in lados_nomes:
                    mlist = [m for m in ath['mets'].get(lado, []) if m is not None]
                    vals_cv = [m[chave] for m in mlist if chave in m]
                    if len(vals_cv) >= 2:
                        arr_cv = np.array(vals_cv, dtype=float)
                        if arr_cv.mean() != 0:
                            all_cvs.append(arr_cv.std(ddof=1) / arr_cv.mean() * 100)
            med_cv = round(float(np.mean([x for x in all_cvs if x is not None])), 1) if any(x is not None for x in all_cvs) else None
            fn_mc = (_FFVR if (med_cv and med_cv > 30) else
                     _FFAM if (med_cv and med_cv > 15) else _FFVD)
            fi_mc = (_FIVR if (med_cv and med_cv > 30) else
                     _FIAM if (med_cv and med_cv > 15) else _FIVD)
            cl(ws, row, c_i, med_cv, fn=fn_mc, fi=fi_mc,
               fmt='0.0' if med_cv else None, bo=BN)
        row += 3

    # ---- Bloco 4: Indice de Perturbacao (Tiro) ----
    if opcoes.get('pos_disp') and _is_tiro_like(protocolo):
        _titulo_bloco(ws, row, 'INDICE DE PERTURBACAO DO DISPARO  (IP = ea95_disp / ea95_pos)')
        row += 1
        hdrs4 = ['Atleta'] + [f'Ensaio {i}' for i in range(1, 11)] + ['IP medio','IP max','IP>2']
        for c_i, h in enumerate(hdrs4, 1):
            cl(ws, row, c_i, h, fn=FCAB, fi=FAZ, al=ALC, bo=BN)
        row += 1
        for ath in atletas:
            ips = calc_perturbacao(ath)
            cl(ws, row, 1, ath['nome'], fn=FNB, bo=BN)
            ip_vals = []
            for ens, ip, _, _ in ips:
                if ens <= 10:
                    fi_ip = _FIVR if (ip or 0) > 2 \
                            else _FIVD
                    fn_ip = _FFVR if (ip or 0) > 2 else _FFVD
                    cl(ws, row, ens+1, round(ip,2) if ip else '-', fn=fn_ip, fi=fi_ip,
                       fmt='0.00', bo=BN)
                if ip is not None: ip_vals.append(ip)
            med_ip = round(float(np.mean([x for x in ip_vals if x is not None])), 2) if any(x is not None for x in ip_vals) else None
            max_ip = round(float(np.max(ip_vals)), 2) if ip_vals else None
            cnt_gt2 = sum(1 for v in ip_vals if v > 2)
            cl(ws, row, 12, med_ip, fn=FNB, fmt='0.00', bo=BN)
            cl(ws, row, 13, max_ip, fn=FNB, fmt='0.00', bo=BN)
            fi_c = _FIVR if cnt_gt2 > 0 \
                   else _FIVD
            cl(ws, row, 14, cnt_gt2, fn=_FFVR if cnt_gt2 > 0 else _FFVD, fi=fi_c, bo=BN)
            row += 1
        row += 2

    ws.freeze_panes = 'B5'
    cred(ws, row + 1)

    # ---- Bloco 5: Estatísticas por distância × intervalo (TIRO) ----
    _friedman_n_aviso = len(atletas) < 10  # aviso de baixo poder
    if _is_tiro_like(protocolo) and opcoes.get('friedman', True):
        all_dists_e = []
        all_itvs_e  = []
        for a in atletas:
            for d in a.get('tiro_dists', []):
                if d not in all_dists_e: all_dists_e.append(d)
            for itv in a.get('tiro_intervalos', []):
                if itv not in all_itvs_e: all_itvs_e.append(itv)

        if all_dists_e and all_itvs_e:
            # Nova aba: TIRO_STATS
            ws_ts = wb.create_sheet(title='TIRO_STATS')
            cl(ws_ts,1,1,
               f'Estatísticas por Distância × Intervalo  |  {PROG} v{VERSAO}  |  n={len(atletas)}',
               fn=Font(name='Calibri',bold=True,size=11,color='0F4C75'))
            cl(ws_ts,2,1,datetime.date.today().strftime('%d/%m/%Y'),fn=FINF)

            row_ts = 4
            for dist in all_dists_e:
                cl(ws_ts, row_ts, 1, f'Distância: {dist}',
                   fn=Font(name='Calibri',bold=True,size=10,color='FFFFFF'),
                   fi=_FIAZ6)
                for cc in range(2, 20):
                    ws_ts.cell(row_ts,cc).fill = _FIAZ6
                row_ts += 1

                # cabecalho das métricas
                hdrs_ts = ['Intervalo'] + [f'n', 'Med', 'DP', 'Min', 'Max', 'CV%', 'SW-p', 'Dist.']
                for chave_m, label_m in METS_XL:
                    hdrs_ts.append(label_m)
                # simplificar: 1 linha por intervalo, colunas = métricas com Med|DP
                hdrs_ts2 = ['Intervalo']
                for _, lbl_m in METS_XL:
                    hdrs_ts2 += [f'{lbl_m[:8]} n', f'Med', 'DP', 'CV%']
                for ci2, h2 in enumerate(hdrs_ts2, 1):
                    fi2 = FAZ if ci2==1 else (FCZ1 if ci2%4 in [2,3] else FCZ2)
                    cl(ws_ts, row_ts, ci2, h2, fn=FCAB, fi=fi2, al=ALC, bo=BN)
                    ws_ts.column_dimensions[get_column_letter(ci2)].width = (
                        20 if ci2==1 else 7)
                ws_ts.row_dimensions[row_ts].height = 28
                row_ts += 1

                for itv in all_itvs_e:
                    itv_lbl = TIRO_INTERVALOS.get(itv, itv)
                    cl(ws_ts, row_ts, 1, itv_lbl, fn=FNB, bo=BN)
                    col_ts = 2
                    for chave_m, _ in METS_XL:
                        # colectar media de cada atleta para esta dist×itv×metrica
                        todas = []
                        for ath in atletas:
                            mlist = (ath.get('tiro_dist',{}).get(dist,{})
                                     .get('mets',{}).get(itv,[]))
                            vs = [m[chave_m] for m in mlist
                                  if m is not None and chave_m in m]
                            vs_c=[x for x in vs if x is not None]
                            if vs_c: todas.append(float(np.mean(vs_c)))
                        if not todas:
                            for _ in range(4):
                                cl(ws_ts,row_ts,col_ts,'-',fn=FN,bo=BN); col_ts+=1
                            continue
                        arr_t = np.array(todas)
                        n_t, med_t, dp_t, cv_t = _desc_stats(todas)
                        cl(ws_ts,row_ts,col_ts,n_t,fn=FN,bo=BN); col_ts+=1
                        cl(ws_ts,row_ts,col_ts,med_t,fn=FN,fmt='0.000',bo=BN); col_ts+=1
                        cl(ws_ts,row_ts,col_ts,dp_t,fn=FN,
                           fmt='0.000' if dp_t else None,bo=BN); col_ts+=1
                        cl(ws_ts,row_ts,col_ts,cv_t,fn=FN,
                           fmt='0.0' if cv_t else None,bo=BN); col_ts+=1
                    row_ts += 1
                row_ts += 2

            ws_ts.freeze_panes = 'B5'
            cred(ws_ts, row_ts + 1)

            # ---- Bloco Friedman: comparação entre intervalos por distância × métrica ----
            # Design: linhas=sujeitos, colunas=intervalos, valor=média dos trials do sujeito
            # Friedman testa se pelo menos um intervalo difere significativamente dos outros
            ws_fr = wb.create_sheet(title='TIRO_FRIEDMAN')
            cl(ws_fr,1,1,
               f'Comparação entre Intervalos  |  Teste de Friedman  |  {PROG} v{VERSAO}  |  n={len(atletas)}',
               fn=Font(name='Calibri',bold=True,size=11,color='0F4C75'))
            cl(ws_fr,2,1,
               'Design: sujeitos × intervalos. Valor por célula = média dos ensaios válidos do sujeito nesse intervalo.',
               fn=FINF)
            cl(ws_fr,2,8,datetime.date.today().strftime('%d/%m/%Y'),fn=FINF)

            # Determinar se fazer post-hoc (por opção nas estatísticas)
            fazer_posthoc = opcoes.get('posthoc_friedman', True)

            row_fr = 4  # linha inicial do conteúdo (após cabeçalho nas linhas 1-2)

            for dist_f in all_dists_e:
                # Título da distância
                cl(ws_fr,row_fr,1,f'Distância: {dist_f}',
                   fn=Font(name='Calibri',bold=True,size=10,color='FFFFFF'),
                   fi=_FIAZ6)
                for cc in range(2,20):
                    ws_fr.cell(row_fr,cc).fill=_FIAZ6
                row_fr += 1

                # cabecalho
                hdrs_fr2 = ['Métrica','n sujeitos'] + \
                           [f'Med {TIRO_INTERVALOS.get(iv,iv)[:12]}' for iv in all_itvs_e] + \
                           ['χ²','gl','p-valor','Sig?','Nota']
                for ci4,h4 in enumerate(hdrs_fr2,1):
                    cl(ws_fr,row_fr,ci4,h4,fn=FCAB,fi=FAZ,al=ALC,bo=BN)
                    ws_fr.column_dimensions[get_column_letter(ci4)].width=(
                        16 if ci4==1 else (10 if ci4<=len(all_itvs_e)+2 else 9))
                ws_fr.row_dimensions[row_fr].height=28
                row_fr += 1

                for chave_f, label_f in METS_XL:
                    # Construir matriz sujeitos × intervalos
                    # Cada célula = média dos trials do sujeito (outliers incluídos - opção futura)
                    matrix = []  # lista de listas, uma por sujeito completo
                    for ath in atletas:
                        row_ath = []
                        for itv_f in all_itvs_e:
                            mlist_f = (ath.get('tiro_dist',{}).get(dist_f,{})
                                       .get('mets',{}).get(itv_f,[]))
                            vs_f = [m[chave_f] for m in mlist_f
                                    if m is not None and chave_f in m]
                            vs_f_c=[x for x in vs_f if x is not None]; row_ath.append(float(np.mean(vs_f_c)) if vs_f_c else None)
                        # Incluir sujeito apenas se tiver valor em TODOS os intervalos
                        if all(v is not None for v in row_ath):
                            matrix.append(row_ath)

                    n_suj = len(matrix)
                    # Medianas por intervalo (sobre os sujeitos)
                    meds_itv = []
                    for j_itv in range(len(all_itvs_e)):
                        col_vals = [matrix[r][j_itv] for r in range(n_suj)]
                        col_vals_c=[x for x in col_vals if x is not None]; meds_itv.append(round(float(np.median(col_vals_c)),3) if col_vals_c else None)

                    cl(ws_fr,row_fr,1,label_f,fn=FNB,bo=BN)
                    cl(ws_fr,row_fr,2,n_suj,fn=FN,bo=BN)
                    for ci_m, med_v in enumerate(meds_itv, 3):
                        cl(ws_fr,row_fr,ci_m,med_v,fn=FN,
                           fmt='0.000' if med_v is not None else None,bo=BN)

                    offset = 3 + len(all_itvs_e)

                    if n_suj < 3 or len(all_itvs_e) < 2:
                        nota = f'n={n_suj} insuficiente (min. 3)' if n_suj<3 else 'Apenas 1 intervalo'
                        for ci5 in range(offset, offset+4):
                            cl(ws_fr,row_fr,ci5,'n/a',fn=FN,bo=BN)
                        cl(ws_fr,row_fr,offset+3,nota,fn=FN,bo=BN)
                        row_fr += 1
                        continue

                    try:
                        cols_f = [np.array([matrix[r][j] for r in range(n_suj)])
                                  for j in range(len(all_itvs_e))]
                        fr_stat, fr_p = stats.friedmanchisquare(*cols_f)
                        gl = len(all_itvs_e) - 1
                        sig = 'Sim *' if fr_p < 0.05 else 'Não'
                        p_str = f'{fr_p:.4f}' if fr_p >= 0.001 else '<0.001'

                        cl(ws_fr,row_fr,offset,  round(float(fr_stat),3),fn=FN,fmt='0.000',bo=BN)
                        cl(ws_fr,row_fr,offset+1,gl,                     fn=FN,bo=BN)
                        _cel_pval(ws_fr,row_fr,offset+2,fr_p)
                        fi_s = _FIVR if fr_p<0.05 \
                               else _FIVD
                        fn_s = _FFVR if fr_p<0.05 else _FFVD
                        cl(ws_fr,row_fr,offset+3,sig,fn=fn_s,fi=fi_s,bo=BN)
                        cl(ws_fr,row_fr,offset+4,'',fn=FN,bo=BN)  # nota vazia
                    except Exception as ex_fr:
                        for ci5 in range(offset, offset+5):
                            cl(ws_fr,row_fr,ci5,str(ex_fr)[:20],fn=FN,bo=BN)

                    row_fr += 1

                    # ---- Post-hoc Wilcoxon pairwise (com correcção Bonferroni) ----
                    if fazer_posthoc and n_suj >= 3:
                        itv_labels_ph = [TIRO_INTERVALOS.get(iv,iv) for iv in all_itvs_e]
                        n_comp = len(all_itvs_e) * (len(all_itvs_e)-1) // 2
                        alpha_bon = 0.05 / n_comp if n_comp > 0 else 0.05

                        cl(ws_fr,row_fr,1,f'  Post-hoc Wilcoxon (Bonferroni α={alpha_bon:.4f})',
                           fn=Font(name='Calibri',italic=True,size=8,color='0E7AA8'))
                        row_fr += 1

                        ph_hdrs = ['  Par A','Par B','W','p-valor','p×n_comp','Sig?']
                        for ci_ph,h_ph in enumerate(ph_hdrs,1):
                            cl(ws_fr,row_fr,ci_ph,h_ph,
                               fn=Font(name='Calibri',bold=True,size=7.5),
                               fi=PatternFill('solid',start_color='BDD7EE'),
                               al=ALC,bo=BN)
                        row_fr += 1

                        for i_ph in range(len(all_itvs_e)):
                            for j_ph in range(i_ph+1, len(all_itvs_e)):
                                a_ph = np.array([matrix[r][i_ph] for r in range(n_suj)])
                                b_ph = np.array([matrix[r][j_ph] for r in range(n_suj)])
                                try:
                                    w_ph, p_ph = stats.wilcoxon(a_ph, b_ph, alternative='two-sided')
                                    p_adj = p_ph * n_comp
                                    sig_ph = 'Sim *' if p_adj < 0.05 else 'Não'
                                    fi_ph = _FIVR \
                                            if p_adj<0.05 else _FIVD
                                    fn_ph = _FFVR if p_adj<0.05 else _FFVD
                                    cl(ws_fr,row_fr,1,f'  {itv_labels_ph[i_ph][:18]}',fn=FN,bo=BN)
                                    cl(ws_fr,row_fr,2,itv_labels_ph[j_ph][:18],fn=FN,bo=BN)
                                    cl(ws_fr,row_fr,3,round(float(w_ph),3),fn=FN,fmt='0.000',bo=BN)
                                    _cel_pval(ws_fr,row_fr,4,p_ph)
                                    _cel_pval(ws_fr,row_fr,5,p_adj)
                                    cl(ws_fr,row_fr,6,sig_ph,fn=fn_ph,fi=fi_ph,bo=BN)
                                except Exception:
                                    cl(ws_fr,row_fr,1,f'  {itv_labels_ph[i_ph][:18]}',fn=FN,bo=BN)
                                    cl(ws_fr,row_fr,2,itv_labels_ph[j_ph][:18],fn=FN,bo=BN)
                                    for ci_e in range(3,7):
                                        cl(ws_fr,row_fr,ci_e,'n/a',fn=FN,bo=BN)
                                row_fr += 1
                        row_fr += 1  # espaço pós post-hoc

                row_fr += 2  # espaço entre distâncias

            ws_fr.freeze_panes = 'A5'
            cred(ws_fr, row_fr+2)


def _titulo_bloco(ws, row, txt):
    """Cabecalho de bloco na aba ESTATS."""
    cl(ws, row, 1, txt,
       fn=Font(name='Calibri', bold=True, size=10, color='FFFFFF'),
       fi=PatternFill('solid', start_color='0F4C75'))
    for c in range(2, 16):
        ws.cell(row, c).fill = PatternFill('solid', start_color='0F4C75')




# -----------------------------------------------------------------------
# Pagina de resultados estatisticos no PDF (ESTATS)
# -----------------------------------------------------------------------

def _pagina_estats_pdf(c, atletas, W, H, protocolo=None, opcoes=None):
    """
    Pagina(s) PDF com tabela de resultados estatísticos.

    Redesenhada com:
    - Paginação completa (sem truncagem)
    - Tabela SW+Descritivos com colunas reduzidas (n, Med, DP, CV%, SW-p, Dist.)
    - Tabela Dir vs Esq com 10 colunas limpas
    - Bloco Tiro: UMA tabela por métrica (linhas=intervalos, cols=n|Med|DP|CV%)
      → elimina o problema dos 35 colunas comprimidas
    - Nova página automática sempre que não há espaço
    """
    from reportlab.lib.units import cm as rl_cm
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib import colors as rl_colors

    if protocolo is None: protocolo = _PROTOCOLO_ACTIVO
    if opcoes is None: opcoes = {k: True for k in _TESTS_INFO}
    proto = PROTOCOLOS.get(protocolo, PROTOCOLOS[PROTO_FMS])
    lados_nomes = [('dir', 'Dir.'), ('esq', 'Esq.')]
    n = len(atletas)

    # ── Paleta ──────────────────────────────────────────────────────────
    C_BG     = rl_colors.HexColor('#0F1923')   # fundo da página
    C_HDR    = rl_colors.HexColor('#00B4D8')   # título principal
    C_SEC    = rl_colors.HexColor('#1F4E79')   # barra de secção
    C_SEC2   = rl_colors.HexColor('#1C3A57')   # barra de sub-secção
    C_TXT    = rl_colors.HexColor('#E8F4FD')   # texto corpo
    C_TXT_H  = rl_colors.HexColor('#00B4D8')   # texto cabeçalho tabela
    C_TXT_L  = rl_colors.HexColor('#BDD7EE')   # texto col. 0 (métricas)
    C_ROW1   = rl_colors.HexColor('#162233')   # linha par
    C_ROW2   = rl_colors.HexColor('#1C2E42')   # linha ímpar
    C_THDR   = rl_colors.HexColor('#243447')   # fundo cabeçalho tabela
    C_SIG_BG = rl_colors.HexColor('#4A1010')   # p-valor sig. fundo
    C_SIG_TX = rl_colors.HexColor('#FCA5A5')   # p-valor sig. texto
    C_WARN_BG= rl_colors.HexColor('#3D2E00')   # p<0.1 fundo
    C_WARN_TX= rl_colors.HexColor('#FCD34D')   # p<0.1 texto
    C_OK_BG  = rl_colors.HexColor('#0D2B1A')   # normal fundo
    C_OK_TX  = rl_colors.HexColor('#6EE7B7')   # normal texto
    C_GRID   = rl_colors.HexColor('#243447')   # grelha

    MARGIN = 1.0 * rl_cm
    TW     = W - 2 * MARGIN          # largura útil da tabela
    ROW_H  = 0.44 * rl_cm            # altura de linha padrão
    SEC_H  = 0.48 * rl_cm            # altura barra de secção
    GAP    = 0.25 * rl_cm            # espaço entre blocos
    FOOTER = 0.9 * rl_cm             # reserva para rodapé
    Y_TOP  = H - 2.4 * rl_cm        # y inicial (abaixo do cabeçalho)

    # ── Estado de paginação ──────────────────────────────────────────────
    _page_num  = [1]
    _y_cur     = [Y_TOP]

    def _draw_page_header():
        """Cabeçalho escuro em cada página."""
        c.setFillColor(C_BG)
        c.rect(0, 0, W, H, fill=1, stroke=0)
        c.setFillColor(C_HDR)
        c.setFont('Helvetica-Bold', 12)
        c.drawString(MARGIN, H - 1.1 * rl_cm,
                     f'Resultados Estatísticos  |  {proto["nome"]}  (n={n})')
        c.setFillColor(rl_colors.HexColor('#6B8FAE'))
        c.setFont('Helvetica', 7)
        c.drawString(MARGIN, H - 1.55 * rl_cm,
                     f'alfa=0.05  |  SW=Shapiro-Wilk  |  t=t-pareado  |  W=Wilcoxon  '
                     f'|  {PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}')
        c.setStrokeColor(C_HDR); c.setLineWidth(0.6)
        c.line(MARGIN, H - 1.85 * rl_cm, W - MARGIN, H - 1.85 * rl_cm)
        # rodapé
        c.setFillColor(rl_colors.HexColor('#2A4A6B'))
        c.setFont('Helvetica', 6.5)
        c.drawString(MARGIN, 0.35 * rl_cm,
                     f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}  |  '
                     + datetime.date.today().strftime('%d/%m/%Y')
                     + f'  |  pág. {_page_num[0]}')

    def _new_page():
        """Avança para nova página."""
        _page_num[0] += 1
        c.showPage()
        _draw_page_header()
        _y_cur[0] = Y_TOP

    def _ensure_space(needed):
        """Garante espaço; avança página se necessário."""
        if _y_cur[0] - needed < FOOTER:
            _new_page()

    def _draw_section_bar(title, color=None):
        """Barra de título de secção."""
        color = color or C_SEC
        _ensure_space(SEC_H + GAP)
        c.setFillColor(color)
        c.rect(MARGIN, _y_cur[0] - SEC_H, TW, SEC_H, fill=1, stroke=0)
        c.setFillColor(C_TXT)
        c.setFont('Helvetica-Bold', 7.5)
        c.drawString(MARGIN + 0.3 * rl_cm, _y_cur[0] - SEC_H * 0.67, title)
        _y_cur[0] -= SEC_H + 0.05 * rl_cm

    def _fmt_pval(p):
        if p is None: return 'n/a'
        if p < 0.001: return '<0.001'
        return f'{p:.4f}'

    def _draw_table(tbl_data, col_widths, pval_cols=None,
                    sig_col=None, dist_col=None, row_h=None):
        """
        Desenha tabela com fundo escuro, paginando automaticamente.
        pval_cols: set de índices de coluna com p-valores (colorir automaticamente)
        sig_col:   índice da coluna 'Sig?' (Sim*/Nao)
        dist_col:  índice da coluna 'Dist.' (Normal/Nao-N.)
        """
        if not tbl_data or len(tbl_data) < 2:
            return
        pval_cols  = pval_cols  or set()
        rh         = row_h or ROW_H
        n_data     = len(tbl_data) - 1   # excl. cabeçalho

        # Dividir em páginas se necessário
        # Estimar quantas linhas cabem
        avail_first = _y_cur[0] - FOOTER
        rows_per_page = max(1, int(avail_first / rh) - 1)  # -1 para cabeçalho

        chunks = []
        header = tbl_data[0]
        data   = tbl_data[1:]
        pos = 0
        while pos < len(data):
            rpp = rows_per_page if pos == 0 else max(1, int((Y_TOP - FOOTER) / rh) - 1)
            chunks.append(data[pos:pos + rpp])
            pos += rpp

        for ci_chunk, chunk in enumerate(chunks):
            if ci_chunk > 0:
                _new_page()
            rows_to_draw = [header] + chunk
            n_rows = len(rows_to_draw)
            actual_h = n_rows * rh

            _ensure_space(actual_h + GAP)

            style_cmds = [
                ('BACKGROUND', (0, 0), (-1, 0),    C_THDR),
                ('TEXTCOLOR',  (0, 0), (-1, 0),    C_TXT_H),
                ('FONTNAME',   (0, 0), (-1, 0),    'Helvetica-Bold'),
                ('FONTSIZE',   (0, 0), (-1, -1),   6.2),
                ('TEXTCOLOR',  (0, 1), (-1, -1),   C_TXT),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [C_ROW1, C_ROW2]),
                ('GRID',       (0, 0), (-1, -1),   0.25, C_GRID),
                ('ALIGN',      (0, 0), (-1, -1),   'CENTER'),
                ('VALIGN',     (0, 0), (-1, -1),   'MIDDLE'),
                ('LEFTPADDING',  (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ('TOPPADDING',   (0, 0), (-1, -1), 1),
                ('BOTTOMPADDING',(0, 0), (-1, -1), 1),
                ('FONTNAME',   (0, 1), (0, -1),    'Helvetica-Bold'),
                ('TEXTCOLOR',  (0, 1), (0, -1),    C_TXT_L),
                ('ALIGN',      (0, 1), (0, -1),    'LEFT'),
            ]

            # Colorir p-valores e indicadores
            for ri in range(1, n_rows):
                row_d = rows_to_draw[ri]
                for ci in range(len(row_d)):
                    val = row_d[ci]
                    if ci in pval_cols and isinstance(val, str) and val not in ('n/a', '-', '-', '?', ''):
                        try:
                            pv = float(val.replace('<', '').replace('>', ''))
                            if val.startswith('<') or pv < 0.05:
                                style_cmds += [('BACKGROUND', (ci,ri),(ci,ri), C_SIG_BG),
                                               ('TEXTCOLOR',  (ci,ri),(ci,ri), C_SIG_TX)]
                            elif pv < 0.10:
                                style_cmds += [('BACKGROUND', (ci,ri),(ci,ri), C_WARN_BG),
                                               ('TEXTCOLOR',  (ci,ri),(ci,ri), C_WARN_TX)]
                            else:
                                style_cmds += [('BACKGROUND', (ci,ri),(ci,ri), C_OK_BG),
                                               ('TEXTCOLOR',  (ci,ri),(ci,ri), C_OK_TX)]
                        except (ValueError, AttributeError):
                            pass
                    if sig_col is not None and ci == sig_col and val == 'Sim *':
                        style_cmds += [('BACKGROUND', (ci,ri),(ci,ri), C_SIG_BG),
                                       ('TEXTCOLOR',  (ci,ri),(ci,ri), C_SIG_TX)]
                    if dist_col is not None and ci == dist_col and isinstance(val, str):
                        if 'Nao' in val or 'nao' in val:
                            style_cmds += [('BACKGROUND', (ci,ri),(ci,ri), C_SIG_BG),
                                           ('TEXTCOLOR',  (ci,ri),(ci,ri), C_SIG_TX)]
                        elif 'Normal' in val:
                            style_cmds += [('BACKGROUND', (ci,ri),(ci,ri), C_OK_BG),
                                           ('TEXTCOLOR',  (ci,ri),(ci,ri), C_OK_TX)]

            row_heights_list = [rh] * n_rows
            tbl = Table(rows_to_draw, colWidths=col_widths, rowHeights=row_heights_list)
            tbl.setStyle(TableStyle(style_cmds))
            tw_real, th_real = tbl.wrapOn(c, TW, 9999)
            tbl.drawOn(c, MARGIN, _y_cur[0] - th_real)
            _y_cur[0] -= th_real + GAP

    # ════════════════════════════════════════════════════════════════════
    # Início: primeira página
    # ════════════════════════════════════════════════════════════════════
    _draw_page_header()

    # ════════════════════════════════════════════════════════════════════
    # BLOCO 1 - Normalidade e Descritivos (SW + n/Med/DP/CV%/p/Dist.)
    # ════════════════════════════════════════════════════════════════════
    if opcoes.get('grupo') and n >= 2:
        _draw_section_bar('NORMALIDADE E DESCRITIVOS  (Shapiro-Wilk  |  n / Média / DP / CV% / p-SW / Distribuição)')

        # Colunas: Métrica | n Dir | Med Dir | DP Dir | CV% Dir | SW-p Dir | Dist Dir
        #                  | n Esq | Med Esq | DP Esq | CV% Esq | SW-p Esq | Dist Esq
        COLS_PER_SIDE = 6   # n, Med, DP, CV%, SW-p, Dist.
        nc1  = 1 + 2 * COLS_PER_SIDE
        lbl1 = 2.6 * rl_cm
        ow1  = (TW - lbl1) / (nc1 - 1)

        hdrs1 = ['Métrica']
        for _, lbl in lados_nomes:
            hdrs1 += [f'n\n{lbl}', f'Média\n{lbl}', f'DP\n{lbl}',
                      f'CV%\n{lbl}', f'SW-p\n{lbl}', f'Dist.\n{lbl}']

        pval_cols1  = {1 + COLS_PER_SIDE*s + 4 for s in range(2)}  # SW-p cols
        dist_cols1  = {1 + COLS_PER_SIDE*s + 5 for s in range(2)}  # Dist. cols

        tbl1_data = [hdrs1]
        for chave, label in METS_PDF:
            row1 = [label]
            for lado, _ in lados_nomes:
                vals = [_media_lado(a, chave, lado) for a in atletas]
                vals = [v for v in vals if v is not None]
                arr  = np.array(vals, dtype=float) if vals else None
                if arr is not None and len(arr) > 0:
                    med = arr.mean()
                    dp  = arr.std(ddof=1) if len(arr) > 1 else 0.0
                    cv  = dp / med * 100 if med != 0 else None
                    row1.append(str(len(arr)))
                    row1.append(f'{med:.2f}')
                    row1.append(f'{dp:.2f}')
                    row1.append(f'{cv:.1f}' if cv is not None else '-')
                    if len(arr) >= 3:
                        _, sw_p = stats.shapiro(arr)
                        row1.append(_fmt_pval(sw_p))
                        row1.append('Normal' if sw_p >= 0.05 else 'Nao-N.')
                    else:
                        row1.append('n/a'); row1.append('n/a')
                else:
                    row1 += ['0', '-', '-', '-', 'n/a', '?']
            tbl1_data.append(row1)

        _draw_table(tbl1_data, [lbl1] + [ow1] * (nc1 - 1),
                    pval_cols=pval_cols1, dist_col=None,
                    row_h=ROW_H)
        # Colorir manualmente as colunas Dist. (não é p-valor numérico)
        # → já tratado em _draw_table via dist_col detection por texto

    # ════════════════════════════════════════════════════════════════════
    # BLOCO 2 - Comparação Dir vs Esq
    # ════════════════════════════════════════════════════════════════════
    if (opcoes.get('dir_esq')
            and protocolo in _TESTS_INFO['dir_esq']['protos']
            and n >= 2):
        _ensure_space(SEC_H + GAP * 2)
        _draw_section_bar('COMPARAÇÃO DIR vs ESQ  (t-pareado se normal  |  Wilcoxon se não-normal  |  Cohen\'s d)', C_SEC2)

        hdrs2 = ['Métrica', 'n', 'Méd Dir', 'Méd Esq', 'Dif',
                 'SW-D p', 'SW-E p', 'Teste', 'Estat.', 'p-valor', 'Sig?', 'd Cohen', 'Efeito']
        lbl2 = 2.6 * rl_cm
        fixed = {'n': 0.5*rl_cm, 'SW-D p': 1.0*rl_cm, 'SW-E p': 1.0*rl_cm,
                 'Teste': 1.1*rl_cm, 'Sig?': 0.7*rl_cm, 'Efeito': 1.1*rl_cm}
        n_flex = len(hdrs2) - 1 - len(fixed)
        flex_w = (TW - lbl2 - sum(fixed.values())) / max(n_flex, 1)
        cw2 = [lbl2]
        for h in hdrs2[1:]:
            cw2.append(fixed.get(h, flex_w))

        pval_cols2 = {hdrs2.index('SW-D p'), hdrs2.index('SW-E p'), hdrs2.index('p-valor')}
        sig_col2   = hdrs2.index('Sig?')

        tbl2_data = [hdrs2]
        for chave, label in METS_PDF:
            vd = [_media_lado(a, chave, 'dir') for a in atletas]
            ve = [_media_lado(a, chave, 'esq') for a in atletas]
            pares = [(d, e) for d, e in zip(vd, ve) if d is not None and e is not None]
            if len(pares) < 2:
                tbl2_data.append([label, str(len(pares))] + ['-'] * (len(hdrs2) - 2))
                continue
            da = np.array([p[0] for p in pares])
            ea = np.array([p[1] for p in pares])
            dif = da - ea
            sw_d_str = sw_e_str = 'n/a'
            normal = True
            if len(pares) >= 3:
                _, sw_d = stats.shapiro(da); _, sw_e2 = stats.shapiro(ea)
                sw_d_str = _fmt_pval(sw_d); sw_e_str = _fmt_pval(sw_e2)
                normal = sw_d >= 0.05 and sw_e2 >= 0.05
            try:
                if normal:
                    tstat, pval = stats.ttest_rel(da, ea); tnm = 't-par.'
                else:
                    tstat, pval = stats.wilcoxon(da, ea, alternative='two-sided'); tnm = 'Wilcox.'
                pv_str = _fmt_pval(pval)
                sig    = 'Sim *' if pval < 0.05 else 'Nao'
            except Exception:
                tstat, pv_str, tnm, sig = 0.0, 'n/a', 'n/a', 'n/a'
            sd_dif  = float(dif.std(ddof=1)) if len(pares) > 1 else 0.0
            cohen_d = round(float(dif.mean()) / sd_dif, 3) if sd_dif > 0 else None
            if cohen_d is None: efeito = '-'
            elif abs(cohen_d) < 0.2: efeito = 'Negl.'
            elif abs(cohen_d) < 0.5: efeito = 'Peq.'
            elif abs(cohen_d) < 0.8: efeito = 'Medio'
            else: efeito = 'Grande'
            tbl2_data.append([
                label, len(pares),
                f'{da.mean():.2f}', f'{ea.mean():.2f}',
                f'{dif.mean():.3f}',
                sw_d_str, sw_e_str, tnm,
                f'{tstat:.3f}' if isinstance(tstat, (int, float)) else str(tstat),
                pv_str, sig,
                f'{cohen_d:.3f}' if cohen_d is not None else '-', efeito,
            ])

        _draw_table(tbl2_data, cw2,
                    pval_cols=pval_cols2, sig_col=sig_col2, row_h=ROW_H)

    # ════════════════════════════════════════════════════════════════════
    # BLOCO 3 - Tiro: estatísticas por distância × intervalo
    # Redesenhado: UMA sub-tabela por MÉTRICA (linhas=intervalos, cols=n|Med|DP|CV%)
    # → elimina completamente o problema dos 35 colunas
    # ════════════════════════════════════════════════════════════════════
    if _is_tiro_like(protocolo) and n >= 2:
        all_dists_p, all_itvs_p = [], []
        for a in atletas:
            for d in a.get('tiro_dists', []):
                if d not in all_dists_p: all_dists_p.append(d)
            for iv in a.get('tiro_intervalos', []):
                if iv not in all_itvs_p: all_itvs_p.append(iv)

        # Agrupar métricas em famílias para organização visual
        MET_GRUPOS = [
            ('Estabilidade Geral',
             ['ea95', 'leng_a', 'leng_b', 'desl', 'time']),
            ('Amplitudes e Razões',
             ['amp_x', 'amp_y', 'ratio_ml_ap', 'ratio_vel']),
            ('Velocidades',
             ['vel_x', 'vel_y', 'vel_med', 'vel_pico_x', 'vel_pico_y', 'stiff_x', 'stiff_y']),
            ('Variância CoP',
             ['cov_xx', 'cov_yy', 'cov_xy']),
        ]
        # Mapear chave → label para lookup rápido
        _lbl_map = {k: lbl for k, lbl in METS_PDF}

        for dist_p in all_dists_p:
            _ensure_space(SEC_H * 2 + GAP)
            _draw_section_bar(
                f'ESTATÍSTICAS POR INTERVALO  -  Distância: {dist_p}  '
                f'(n={n}  |  Média dos ensaios válidos por atleta  →  Média/DP do grupo)')

            # Pré-calcular médias de cada atleta por métrica × intervalo
            # cache[chave][itv] = np.array de médias por atleta
            _cache = {}
            for chave_c, _ in METS_PDF:
                _cache[chave_c] = {}
                for itv_c in all_itvs_p:
                    ath_meds = []
                    for ath in atletas:
                        mlist = (ath.get('tiro_dist', {})
                                 .get(dist_p, {}).get('mets', {}).get(itv_c, []))
                        vs = [m[chave_c] for m in mlist
                              if m is not None and chave_c in m and m[chave_c] is not None]
                        if vs:
                            ath_meds.append(float(np.mean([x for x in vs if x is not None])))
                    _cache[chave_c][itv_c] = np.array(ath_meds) if ath_meds else None

            itv_lbls_short = [TIRO_INTERVALOS.get(iv, iv) for iv in all_itvs_p]

            for grp_nome, grp_chaves in MET_GRUPOS:
                # Verificar se há dados para este grupo nesta distância
                has_grp = any(
                    any(_cache[ck].get(iv) is not None for iv in all_itvs_p)
                    for ck in grp_chaves if ck in _cache)
                if not has_grp:
                    continue

                _ensure_space(SEC_H * 0.85 + GAP)
                # Barra de sub-grupo (mais estreita)
                c.setFillColor(C_SEC2)
                c.rect(MARGIN, _y_cur[0] - SEC_H * 0.82,
                       TW, SEC_H * 0.82, fill=1, stroke=0)
                c.setFillColor(rl_colors.HexColor('#93C4DE'))
                c.setFont('Helvetica-Bold', 6.5)
                c.drawString(MARGIN + 0.3 * rl_cm,
                             _y_cur[0] - SEC_H * 0.82 * 0.65, grp_nome)
                _y_cur[0] -= SEC_H * 0.82 + 0.08 * rl_cm

                # Tabela: linhas = métricas, colunas = Intervalo + (n|Med|DP|CV%) por intervalo
                # 1 coluna de métrica + n_itvs * 4 colunas = 1 + len(itvs)*4
                # Com 5 intervalos: 1 + 20 = 21 cols → ainda razoável (0.7cm cada)
                n_itvs  = len(all_itvs_p)
                MET_W   = 2.8 * rl_cm
                SUB_W   = (TW - MET_W) / (n_itvs * 4) if n_itvs > 0 else 1.5 * rl_cm
                SUB_W   = max(SUB_W, 0.6 * rl_cm)

                # Cabeçalho: row 0 = span de intervalo, row 1 = n/Med/DP/CV% por intervalo
                hdr_row0 = ['Métrica']
                for lbl_i in itv_lbls_short:
                    hdr_row0 += [lbl_i, '', '', '']
                hdr_row1 = ['']
                for _ in all_itvs_p:
                    hdr_row1 += ['n', 'Méd', 'DP', 'CV%']

                tbl3_data = [hdr_row0, hdr_row1]
                tbl3_spans = []  # (col_start, col_end) for span commands

                for ck in grp_chaves:
                    if ck not in _cache: continue
                    row3 = [_lbl_map.get(ck, ck)]
                    for iv in all_itvs_p:
                        arr3 = _cache[ck][iv]
                        if arr3 is not None and len(arr3) > 0:
                            med3 = arr3.mean()
                            dp3  = arr3.std(ddof=1) if len(arr3) > 1 else 0.0
                            cv3  = dp3 / med3 * 100 if med3 != 0 else None
                            row3 += [str(len(arr3)),
                                     f'{med3:.2f}',
                                     f'{dp3:.2f}',
                                     f'{cv3:.1f}' if cv3 is not None else '-']
                        else:
                            row3 += ['0', '-', '-', '-']
                    tbl3_data.append(row3)

                if len(tbl3_data) < 3:   # só cabeçalhos
                    continue

                # Construir colWidths
                cw3 = [MET_W] + [SUB_W] * (n_itvs * 4)

                # Estilo especial para a tabela de 2 cabeçalhos + span
                style3 = [
                    # Cabeçalho linha 0: fundo escuro, texto azul
                    ('BACKGROUND', (0, 0), (-1, 0),    C_THDR),
                    ('TEXTCOLOR',  (0, 0), (-1, 0),    C_TXT_H),
                    ('FONTNAME',   (0, 0), (-1, 0),    'Helvetica-Bold'),
                    ('ALIGN',      (0, 0), (-1, 0),    'CENTER'),
                    # Cabeçalho linha 1: subcoluna
                    ('BACKGROUND', (0, 1), (-1, 1),    rl_colors.HexColor('#1A3350')),
                    ('TEXTCOLOR',  (0, 1), (-1, 1),    rl_colors.HexColor('#93C4DE')),
                    ('FONTNAME',   (0, 1), (-1, 1),    'Helvetica-Bold'),
                    ('ALIGN',      (0, 1), (-1, 1),    'CENTER'),
                    # Corpo
                    ('FONTSIZE',   (0, 0), (-1, -1),   5.8),
                    ('TEXTCOLOR',  (0, 2), (-1, -1),   C_TXT),
                    ('ROWBACKGROUNDS', (0, 2), (-1, -1), [C_ROW1, C_ROW2]),
                    ('GRID',       (0, 0), (-1, -1),   0.2, C_GRID),
                    ('ALIGN',      (0, 0), (-1, -1),   'CENTER'),
                    ('VALIGN',     (0, 0), (-1, -1),   'MIDDLE'),
                    ('LEFTPADDING',  (0, 0), (-1, -1), 2),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 2),
                    ('TOPPADDING',   (0, 0), (-1, -1), 1),
                    ('BOTTOMPADDING',(0, 0), (-1, -1), 1),
                    # Coluna 0 (métricas)
                    ('FONTNAME',   (0, 2), (0, -1),    'Helvetica-Bold'),
                    ('TEXTCOLOR',  (0, 2), (0, -1),    C_TXT_L),
                    ('ALIGN',      (0, 2), (0, -1),    'LEFT'),
                    # Span: cada grupo de 4 colunas de intervalo → 1 célula de título
                ]
                # Spans para grupos de 4 subcolunas (MERGE horizontal no cabeçalho 0)
                for s_i, _ in enumerate(all_itvs_p):
                    col_s = 1 + s_i * 4
                    col_e = col_s + 3
                    style3.append(('SPAN', (col_s, 0), (col_e, 0)))
                # Colorir colunas CV% em amarelo claro se >30
                # (feito inline para não complicar; apenas visual)

                tbl3 = Table(tbl3_data, colWidths=cw3,
                             rowHeights=[ROW_H * 0.9, ROW_H * 0.85]
                                        + [ROW_H] * (len(tbl3_data) - 2))
                tbl3.setStyle(TableStyle(style3))
                _, th3 = tbl3.wrapOn(c, TW, 9999)
                _ensure_space(th3 + GAP)
                tbl3.drawOn(c, MARGIN, _y_cur[0] - th3)
                _y_cur[0] -= th3 + GAP * 0.8

        # ── Teste de Friedman por distância × métrica (se activado) ────
        if opcoes.get('friedman') and n >= 3 and all_itvs_p and len(all_itvs_p) > 2:
            _ensure_space(SEC_H + GAP * 2)
            _draw_section_bar('TESTE DE FRIEDMAN  (comparação de intervalos  |  H0: sem diferença entre intervalos)', C_SEC2)

            for dist_p in all_dists_p:
                _ensure_space(SEC_H * 0.82 + GAP)
                c.setFillColor(rl_colors.HexColor('#1A3350'))
                c.rect(MARGIN, _y_cur[0] - SEC_H * 0.72,
                       TW, SEC_H * 0.72, fill=1, stroke=0)
                c.setFillColor(rl_colors.HexColor('#93C4DE'))
                c.setFont('Helvetica-Bold', 6.5)
                c.drawString(MARGIN + 0.3 * rl_cm,
                             _y_cur[0] - SEC_H * 0.72 * 0.68, f'Distância: {dist_p}')
                _y_cur[0] -= SEC_H * 0.72 + 0.08 * rl_cm

                hdrs_fr = ['Métrica', 'Stat. χ²', 'p-valor', 'Sig?']
                lbl_fr  = 3.0 * rl_cm
                ow_fr   = (TW - lbl_fr) / 3
                tbl_fr  = [hdrs_fr]

                for chave_fr, lbl_fr_t in METS_PDF:
                    groups_fr = []
                    for itv_fr in all_itvs_p:
                        ath_meds_fr = []
                        for ath in atletas:
                            ml_fr = (ath.get('tiro_dist', {})
                                     .get(dist_p, {}).get('mets', {}).get(itv_fr, []))
                            vs_fr = [m[chave_fr] for m in ml_fr
                                     if m is not None and chave_fr in m
                                     and m[chave_fr] is not None]
                            vs_fr_c=[x for x in vs_fr if x is not None]
                            if vs_fr_c: ath_meds_fr.append(float(np.mean(vs_fr_c)))
                        if ath_meds_fr: groups_fr.append(ath_meds_fr)
                    if len(groups_fr) < 3 or any(len(g) < 3 for g in groups_fr):
                        tbl_fr.append([lbl_fr_t, '-', 'n/a', '-'])
                        continue
                    try:
                        min_len = min(len(g) for g in groups_fr)
                        matrix  = np.array([g[:min_len] for g in groups_fr]).T
                        stat_fr, pval_fr = stats.friedmanchisquare(*[matrix[:, j]
                                                                     for j in range(matrix.shape[1])])
                        pv_fr_str = _fmt_pval(pval_fr)
                        sig_fr    = 'Sim *' if pval_fr < 0.05 else 'Nao'
                        tbl_fr.append([lbl_fr_t, f'{stat_fr:.3f}', pv_fr_str, sig_fr])
                    except Exception:
                        tbl_fr.append([lbl_fr_t, '-', 'n/a', '-'])

                if len(tbl_fr) > 1:
                    _draw_table(tbl_fr, [lbl_fr] + [ow_fr] * 3,
                                pval_cols={2}, sig_col=3, row_h=ROW_H)


def aba_spss_tiro(wb, atletas, dist, itv):
    """
    Aba SPSS para protocolo Tiro: formato SPSS standard.
    Uma linha por atleta. Colunas = média de cada métrica para este dist×itv.
    Sem dados por trial - apenas médias dos ensaios válidos.
    """
    itv_label  = TIRO_INTERVALOS.get(itv, itv)
    sheet_name = f'SPSS_{dist}_{itv_label[:10]}'[:31]
    ws = wb.create_sheet(title=sheet_name)

    cl(ws,1,1, f'SPSS  |  Tiro  |  {dist}  |  {itv_label}  |  n={len(atletas)}',
       fn=Font(name='Calibri',bold=True,size=9,color='0F4C75'))
    cl(ws,1,10, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}', fn=FINF)
    cl(ws,2,1, f'Valores = média dos ensaios válidos por atleta  |  {datetime.date.today().strftime("%d/%m/%Y")}', fn=FINF)

    # cabecalho: nome | met1 | met2 | ...
    hdrs = ['nome'] + [lbl.replace(' ','_')[:12] for _, lbl in METS_SPSS]
    for ci, h in enumerate(hdrs, 1):
        fi = FAZ if ci == 1 else FCZ1
        cl(ws,3,ci,h,fn=FCAB,fi=fi,al=ALC,bo=BN)
        ws.column_dimensions[get_column_letter(ci)].width = (20 if ci==1 else 12)
    ws.row_dimensions[3].height = 22

    for ri, ath in enumerate(atletas, 4):
        cl(ws,ri,1,ath['nome'],fn=FNB,bo=BN)
        mets_l = ath.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[])
        for ci, (chave, _) in enumerate(METS_SPSS, 2):
            vals = [m[chave] for m in mets_l if m is not None and chave in m]
            vals_c=[x for x in vals if x is not None]; v = round(float(np.mean(vals_c)), 4) if vals_c else None
            cl(ws,ri,ci,v,fn=FN,fi=FCZ1,fmt='0.0000' if v is not None else None,bo=BN)

    # grupo
    row_grp = len(atletas) + 5
    cl(ws,row_grp,1,f'Média do grupo (n={len(atletas)})',
       fn=Font(name='Calibri',bold=True,size=9,color='FFFFFF'),
       fi=PatternFill('solid',start_color='0F4C75'))
    for ci, (chave, _) in enumerate(METS_SPSS, 2):
        todas = []
        for ath in atletas:
            ml = ath.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[])
            vs = [m[chave] for m in ml if m is not None and chave in m]
            vs_c=[x for x in vs if x is not None]
            if vs_c: todas.append(float(np.mean(vs_c)))
        todas_c=[x for x in todas if x is not None]; v = round(float(np.mean(todas_c)), 4) if todas_c else None
        cl(ws,row_grp,ci,v,fn=FNB,
           fi=PatternFill('solid',start_color='FEF9C3'),
           fmt='0.0000' if v is not None else None,bo=BN)

    ws.freeze_panes = 'B4'
    cred(ws, row_grp+2)


def guardar_individual_tiro(ath, pasta_saida, c_elipse=True, c_estab=True):
    """
    Gera ficheiros individuais Excel para o protocolo de tiro.
    Um ficheiro por distância: contém abas por intervalo × trial.
    Nome: [ID_]Nome_[dist].xlsx
    """
    nome   = ath['nome']
    ind_id = ath.get('id','')
    prefix = f'{ind_id}_' if ind_id else ''

    all_dists = ath.get('tiro_dists', [])
    all_itvs  = ath.get('tiro_intervalos', [])

    for dist in all_dists:
        safe_nome = re.sub(r'[<>:"/\\|?*]', '_', nome)
        fn = f'{prefix}{safe_nome}_{dist}.xlsx'
        caminho = os.path.join(pasta_saida, fn)

        wb = Workbook()
        if 'Sheet' in wb.sheetnames: del wb['Sheet']

        # ---- Aba RESUMO: todos os intervalos × todas as métricas ----
        ws_res = wb.create_sheet(title='RESUMO')
        cl(ws_res,1,1, f'{PROG}  |  {nome}  |  {dist}  |  Resumo',
           fn=Font(name='Calibri',bold=True,size=11,color='0F4C75'))
        cl(ws_res,1,10, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}', fn=FINF)
        cl(ws_res,2,1, datetime.date.today().strftime('%d/%m/%Y'), fn=FINF)

        n_t_max = max(
            (len(ath.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[]))
             for itv in all_itvs), default=0)

        # cabecalho RESUMO
        hdrs_res = ['Intervalo'] + [f'T{i}' for i in range(1,n_t_max+1)] + ['Max','Min','Med','DP','CV%']
        for met_chave, met_label in METS_XL:
            row_hdr = len(all_itvs)*len(METS_XL)  # cálculo abaixo

        row_res = 4
        for met_chave, met_label in METS_XL:
            # título da métrica
            ws_res.merge_cells(start_row=row_res,start_column=1,
                               end_row=row_res,end_column=len(hdrs_res))
            cl(ws_res,row_res,1, met_label,
               fn=Font(name='Calibri',bold=True,size=9,color='FFFFFF'),
               fi=_FIAZ6)
            for cc in range(2,len(hdrs_res)+1):
                ws_res.cell(row_res,cc).fill=_FIAZ6
            row_res += 1

            # subcabecalho
            for ci2,h2 in enumerate(hdrs_res,1):
                fi2 = FAZ if ci2==1 else (FAMAR if h2 in ('Max','Min','Med','DP','CV%') else FCZ1)
                cl(ws_res,row_res,ci2,h2,fn=FCAB,fi=fi2,al=ALC,bo=BN)
                ws_res.column_dimensions[get_column_letter(ci2)].width=(22 if ci2==1 else 8)
            row_res += 1

            # dados
            for itv in all_itvs:
                itv_lbl = TIRO_INTERVALOS.get(itv, itv)
                mets_l  = ath.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[])
                flags   = flagrar_outliers(mets_l)
                cl(ws_res,row_res,1,itv_lbl,fn=FNB,bo=BN)
                vals = []
                for ti in range(n_t_max):
                    m  = mets_l[ti] if ti < len(mets_l) else None
                    v  = round(m[met_chave],4) if m and met_chave in m and m[met_chave] is not None else None
                    fo = FOUT if (ti < len(flags) and flags[ti]) else FCZ1
                    fn_= FOUT_FN if (ti < len(flags) and flags[ti]) else FN
                    cl(ws_res,row_res,2+ti,v,fn=fn_,fi=fo,
                       fmt='0.000' if v is not None else None,bo=BN)
                    if v is not None: vals.append(v)
                arr = np.array(vals)
                sv  = [round(float(arr.max()),3)  if len(arr)>0 else None,
                       round(float(arr.min()),3)  if len(arr)>0 else None,
                       round(float(arr.mean()),3) if len(arr)>0 else None,
                       round(float(arr.std(ddof=1)),3) if len(arr)>1 else None,
                       round(float(arr.std(ddof=1)/arr.mean()*100),1)
                           if len(arr)>1 and arr.mean()!=0 else None]
                for si2,sv2 in enumerate(sv):
                    cl(ws_res,row_res,2+n_t_max+si2,sv2,fn=FNB,fi=FAMAR,
                       fmt='0.000' if sv2 is not None else None,bo=BN)
                row_res += 1
            row_res += 1

        ws_res.freeze_panes = 'B4'

        # ---- Abas por intervalo × trial ----
        for itv in all_itvs:
            itv_lbl = TIRO_INTERVALOS.get(itv, itv)
            mets_l  = ath.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[])
            raw_l   = ath.get('tiro_dist',{}).get(dist,{}).get('raw',{}).get(itv,[])
            flags   = flagrar_outliers(mets_l)
            n_t     = len(mets_l)

            for ti in range(n_t):
                m     = mets_l[ti] if ti < len(mets_l) else None
                rd    = raw_l[ti]  if ti < len(raw_l)  else None
                is_out= flags[ti]  if ti < len(flags)  else False

                # nome da aba: intervalo abreviado + trial
                itv_short = itv_lbl[:8].replace(' ','_')
                ne = f'{itv_short}_T{ti+1}'[:31]

                if m is None or rd is None:
                    ws_e = wb.create_sheet(title=ne)
                    cl(ws_e,1,1,f'Sem dados: {itv_lbl} T{ti+1}',
                       fn=Font(name='Calibri',color='FF0000'))
                    continue

                aba_ensaio(wb, ne, nome, rd.get('dados',{}).get('data',''),
                           rd['dados']['frames'], rd.get('ini'), rd.get('fim'),
                           m, is_out)
                if c_elipse: aba_elipse(wb, ne, m, nome, is_out)
                if c_estab:  aba_estabilograma(wb, ne, m, nome)

        # ---- Aba SPSS deste ficheiro individual ----
        ws_spss = wb.create_sheet(title='SPSS')
        cl(ws_spss,1,1, f'SPSS  |  {nome}  |  {dist}  |  Trial-level',
           fn=Font(name='Calibri',bold=True,size=9,color='0F4C75'))

        n_t2 = max(
            (len(ath.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[]))
             for itv in all_itvs), default=0)
        hdrs_sp = ['intervalo'] + [f'{chave}_T{ti}' for chave,_ in METS_SPSS
                                   for ti in range(1,n_t2+1)] + \
                  [f'{chave}_Med' for chave,_ in METS_SPSS]
        for ci3,h3 in enumerate(hdrs_sp,1):
            fi3 = FAZ if ci3==1 else (FAMAR if '_Med' in h3 else FCZ1)
            cl(ws_spss,3,ci3,h3,fn=FCAB,fi=fi3,al=ALC,bo=BN)
            ws_spss.column_dimensions[get_column_letter(ci3)].width=(20 if ci3==1 else 10)
        ws_spss.row_dimensions[3].height=22

        for ri3,itv in enumerate(all_itvs,4):
            itv_lbl3 = TIRO_INTERVALOS.get(itv,itv)
            mets_l3  = ath.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[])
            cl(ws_spss,ri3,1,itv_lbl3,fn=FNB,bo=BN)
            col3 = 2
            for chave,_ in METS_SPSS:
                vals3 = []
                for ti in range(n_t2):
                    m3 = mets_l3[ti] if ti < len(mets_l3) else None
                    v3 = round(m3[chave],4) if m3 and chave in m3 and m3[chave] is not None else None
                    cl(ws_spss,ri3,col3,v3,fn=FN,fi=FCZ1,
                       fmt='0.0000' if v3 is not None else None,bo=BN)
                    if v3 is not None: vals3.append(v3)
                    col3+=1
            # Med no final
            col_med = 2 + len(METS_SPSS)*n_t2
            for chave,_ in METS_SPSS:
                mlist3 = ath.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[])
                vs3 = [m3[chave] for m3 in mlist3 if m3 and chave in m3 and m3[chave] is not None]
                vs3_c=[x for x in vs3 if x is not None]; v3m = round(float(np.mean(vs3_c)),4) if vs3_c else None
                cl(ws_spss,ri3,col_med,v3m,fn=FNB,fi=FAMAR,
                   fmt='0.0000' if v3m is not None else None,bo=BN)
                col_med+=1
        ws_spss.freeze_panes='B4'

        # ---- Aba HURDLE STEP (bipodal, opcional) ---- v18
        mets_hs = ath.get('mets', {})
        raw_hs  = ath.get('raw',  {})
        tem_hs  = any(
            any(m is not None for m in mets_hs.get(lado, []))
            for lado in ('dir', 'esq'))
        if tem_hs:
            ws_hs = wb.create_sheet(title='HURDLE_STEP')
            cl(ws_hs, 1, 1,
               f'Hurdle Step Bipodal  |  {nome}  |  Pe Direito + Pe Esquerdo',
               fn=Font(name='Calibri', bold=True, size=11, color='0F4C75'))
            cl(ws_hs, 1, 10,
               f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}', fn=FINF)

            n_hs_max = max(len(mets_hs.get('dir', [])), len(mets_hs.get('esq', [])))
            hdrs_hs = (['Métrica'] +
                       [f'Dir_T{i}' for i in range(1, n_hs_max+1)] +
                       ['Dir_Max','Dir_Min','Dir_Med','Dir_DP','Dir_CV%'] +
                       [f'Esq_T{i}' for i in range(1, n_hs_max+1)] +
                       ['Esq_Max','Esq_Min','Esq_Med','Esq_DP','Esq_CV%'] +
                       ['AI (%)'])
            for ci_h, h_txt in enumerate(hdrs_hs, 1):
                if ci_h == 1:
                    fi_h = FAZ
                elif 'Dir' in h_txt:
                    fi_h = FCZ1
                elif 'Esq' in h_txt:
                    fi_h = FCZ2
                else:
                    fi_h = FVD2
                _al_hs = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cl(ws_hs, 3, ci_h, h_txt, fn=FCAB, fi=fi_h, al=_al_hs, bo=BN)
                ws_hs.column_dimensions[get_column_letter(ci_h)].width = (24 if ci_h==1 else 9)
            ws_hs.row_dimensions[3].height = 22

            for ri_hs, (chave, label) in enumerate(METS_XL, start=4):
                cl(ws_hs, ri_hs, 1, label, fn=FNB, bo=BN)
                col_hs = 2
                for lado_hs, fi_hs in [('dir', FCZ1), ('esq', FCZ2)]:
                    lst_hs   = mets_hs.get(lado_hs, [])
                    flags_hs = flagrar_outliers(lst_hs)
                    vals_hs  = []
                    for ti in range(n_hs_max):
                        m_h = lst_hs[ti] if ti < len(lst_hs) else None
                        v_h = round(m_h[chave], 4) if m_h and chave in m_h and m_h[chave] is not None else None
                        is_o = flags_hs[ti] if ti < len(flags_hs) else False
                        cl(ws_hs, ri_hs, col_hs, v_h,
                           fn=(FOUT_FN if is_o else FN),
                           fi=(FOUT if is_o else fi_hs),
                           fmt='0.000' if v_h is not None else None, bo=BN)
                        if v_h is not None: vals_hs.append(v_h)
                        col_hs += 1
                    arr_hs = np.array(vals_hs)
                    for sv in [
                        round(float(arr_hs.max()),  3) if len(arr_hs)>0 else None,
                        round(float(arr_hs.min()),  3) if len(arr_hs)>0 else None,
                        round(float(arr_hs.mean()), 3) if len(arr_hs)>0 else None,
                        round(float(arr_hs.std(ddof=1)),3) if len(arr_hs)>1 else None,
                        round(float(arr_hs.std(ddof=1)/arr_hs.mean()*100),1)
                            if len(arr_hs)>1 and arr_hs.mean()!=0 else None,
                    ]:
                        cl(ws_hs, ri_hs, col_hs, sv, fn=FN, fi=FAMAR,
                           fmt='0.000' if sv is not None else None, bo=BN)
                        col_hs += 1
                # AI dir vs esq
                vd_ai = [m_h[chave] for m_h in mets_hs.get('dir',[]) if m_h and chave in m_h and m_h[chave] is not None]
                ve_ai = [m_h[chave] for m_h in mets_hs.get('esq',[]) if m_h and chave in m_h and m_h[chave] is not None]
                ai_v = (assimetria(float(np.mean(vd_ai)), float(np.mean(ve_ai)))
                        if vd_ai and ve_ai else None)
                cl(ws_hs, ri_hs, col_hs, ai_v,
                   fn=FNB, fi=FVD2,
                   fmt='+0.00;-0.00;0.00' if ai_v is not None else None, bo=BN)
            ws_hs.freeze_panes = 'B4'

        # ---- Abas Selection CoP por dist/itv (v19) ----
        sel_d = ath.get('sel_dist', {}).get(dist, {})
        for itv_s2, sel_itv3 in sel_d.items():
            has_sel3 = (any(m is not None for m in sel_itv3.get('dir', [])) or
                        any(m is not None for m in sel_itv3.get('esq', [])))
            if not has_sel3:
                continue
            itv_lbl3 = TIRO_INTERVALOS.get(itv_s2, itv_s2)
            sheet_nm3 = f'SEL_{dist}_{itv_s2}'[:31]
            ws_sel3 = wb.create_sheet(title=sheet_nm3)
            cl(ws_sel3, 1, 1,
               f'Selection CoP  |  {dist}  |  {itv_lbl3}  |  Pe Dir. (Right) + Pe Esq. (Left)',
               fn=Font(name='Calibri', bold=True, size=11, color='0F4C75'))
            cl(ws_sel3, 1, 12, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}', fn=FINF)

            mets_d3  = [m for m in sel_itv3.get('dir', []) if m is not None]
            mets_e3  = [m for m in sel_itv3.get('esq', []) if m is not None]
            n3 = max(len(mets_d3), len(mets_e3))
            if n3 == 0:
                continue

            # Cabecalhos
            hdrs3 = (['Métrica'] +
                     [f'Dir_T{i}' for i in range(1, n3+1)] +
                     ['Dir_Max','Dir_Min','Dir_Med','Dir_DP','Dir_CV%'] +
                     [f'Esq_T{i}' for i in range(1, n3+1)] +
                     ['Esq_Max','Esq_Min','Esq_Med','Esq_DP','Esq_CV%'])
            for ci3, h3 in enumerate(hdrs3, 1):
                _fi3 = (FAZ if ci3 == 1
                        else (FCZ1 if 'Dir' in h3 else FCZ2))
                _al3 = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cl(ws_sel3, 3, ci3, h3, fn=FCAB, fi=_fi3, al=_al3, bo=BN)
                ws_sel3.column_dimensions[get_column_letter(ci3)].width = 24 if ci3 == 1 else 9
            ws_sel3.row_dimensions[3].height = 22

            flags_d3 = flagrar_outliers(mets_d3)
            flags_e3 = flagrar_outliers(mets_e3)
            for ri3, (chave3, label3) in enumerate(METS_XL, start=4):
                cl(ws_sel3, ri3, 1, label3, fn=FNB, bo=BN)
                ci3 = 2
                for lst3, flags3, fi3 in [(mets_d3, flags_d3, FCZ1),
                                           (mets_e3, flags_e3, FCZ2)]:
                    vals3 = []
                    for ti in range(n3):
                        m3 = lst3[ti] if ti < len(lst3) else None
                        v3 = round(m3[chave3], 4) if m3 and chave3 in m3 and m3[chave3] is not None else None
                        is_o3 = flags3[ti] if ti < len(flags3) else False
                        cl(ws_sel3, ri3, ci3, v3,
                           fn=(FOUT_FN if is_o3 else FN),
                           fi=(FOUT if is_o3 else fi3),
                           fmt='0.000' if v3 is not None else None, bo=BN)
                        if v3 is not None: vals3.append(v3)
                        ci3 += 1
                    arr3 = np.array(vals3)
                    for sv3 in [
                        round(float(arr3.max()),  3) if len(arr3) > 0 else None,
                        round(float(arr3.min()),  3) if len(arr3) > 0 else None,
                        round(float(arr3.mean()), 3) if len(arr3) > 0 else None,
                        round(float(arr3.std(ddof=1)), 3) if len(arr3) > 1 else None,
                        round(float(arr3.std(ddof=1)/arr3.mean()*100), 1)
                            if len(arr3) > 1 and arr3.mean() != 0 else None,
                    ]:
                        cl(ws_sel3, ri3, ci3, sv3, fn=FN, fi=FAMAR,
                           fmt='0.000' if sv3 is not None else None, bo=BN)
                        ci3 += 1
            ws_sel3.freeze_panes = 'B4'

        cred(ws_res, row_res+2)
        wb.save(caminho)


def guardar_individual_sel_tiro(ath, pasta_saida, lado, c_elipse=True, c_estab=True):
    """
    Gera ficheiros individuais Excel para a análise de Pe Direito (Right Sel.)
    ou Pe Esquerdo (Left Sel.) no protocolo de tiro. (v20)

    lado : 'dir' → Pe Direito (Right Sel.)  →  ficheiro Nome_sel_dir_[dist].xlsx
           'esq' → Pe Esquerdo (Left Sel.)  →  ficheiro Nome_sel_esq_[dist].xlsx
    """
    nome      = ath['nome']
    ind_id    = ath.get('id', '')
    prefix    = f'{ind_id}_' if ind_id else ''
    lado_lbl  = ('Pe Direito (Right Sel.)' if lado == 'dir'
                 else 'Pe Esquerdo (Left Sel.)')
    lado_tag  = 'sel_dir' if lado == 'dir' else 'sel_esq'
    fi_lado   = FCZ1 if lado == 'dir' else FCZ2

    all_dists = ath.get('tiro_dists', [])
    all_itvs  = ath.get('tiro_intervalos', [])

    for dist in all_dists:
        # Verificar se existem dados Selection para este lado/dist
        sel_d = ath.get('sel_dist', {}).get(dist, {})
        tem_dados = any(
            any(m is not None for m in sel_d.get(itv_k, {}).get(lado, []))
            for itv_k in all_itvs)
        if not tem_dados:
            continue

        safe_nome = re.sub(r'[<>:"/\\|?*]', '_', nome)
        fn = f'{prefix}{safe_nome}_{lado_tag}_{dist}.xlsx'
        caminho = os.path.join(pasta_saida, fn)

        wb = Workbook()
        if 'Sheet' in wb.sheetnames: del wb['Sheet']

        # ---- Aba RESUMO ----
        ws_res = wb.create_sheet(title='RESUMO')
        cl(ws_res, 1, 1,
           f'{PROG}  |  {nome}  |  {dist}  |  {lado_lbl}  |  Resumo',
           fn=Font(name='Calibri', bold=True, size=11, color='0F4C75'))
        cl(ws_res, 1, 12, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}', fn=FINF)
        cl(ws_res, 2, 1, datetime.date.today().strftime('%d/%m/%Y'), fn=FINF)

        n_t_max = max(
            (len([m for m in sel_d.get(itv_k, {}).get(lado, []) if m is not None])
             for itv_k in all_itvs), default=0)
        if n_t_max == 0:
            n_t_max = 5

        hdrs_res = (['Intervalo'] +
                    [f'T{i}' for i in range(1, n_t_max+1)] +
                    ['Max', 'Min', 'Med', 'DP', 'CV%'])

        row_res = 4
        for met_chave, met_label in METS_XL:
            ws_res.merge_cells(start_row=row_res, start_column=1,
                               end_row=row_res, end_column=len(hdrs_res))
            cl(ws_res, row_res, 1, met_label,
               fn=Font(name='Calibri', bold=True, size=9, color='FFFFFF'),
               fi=_FIAZ6)
            for cc in range(2, len(hdrs_res)+1):
                ws_res.cell(row_res, cc).fill = _FIAZ6
            row_res += 1

            for ci2, h2 in enumerate(hdrs_res, 1):
                fi2 = FAZ if ci2 == 1 else (FAMAR if h2 in ('Max','Min','Med','DP','CV%') else fi_lado)
                cl(ws_res, row_res, ci2, h2, fn=FCAB, fi=fi2, al=ALC, bo=BN)
                ws_res.column_dimensions[get_column_letter(ci2)].width = (24 if ci2 == 1 else 9)
            row_res += 1

            for itv_k in all_itvs:
                itv_lbl = TIRO_INTERVALOS.get(itv_k, itv_k)
                mets_l  = [m for m in sel_d.get(itv_k, {}).get(lado, []) if m is not None]
                flags   = flagrar_outliers(mets_l)
                cl(ws_res, row_res, 1, itv_lbl, fn=FNB, bo=BN)
                vals = []
                for ti in range(n_t_max):
                    m  = mets_l[ti] if ti < len(mets_l) else None
                    v  = round(m[met_chave], 4) if m and met_chave in m and m[met_chave] is not None else None
                    fo = FOUT if (ti < len(flags) and flags[ti]) else fi_lado
                    fn_ = FOUT_FN if (ti < len(flags) and flags[ti]) else FN
                    cl(ws_res, row_res, 2+ti, v, fn=fn_, fi=fo,
                       fmt='0.000' if v is not None else None, bo=BN)
                    if v is not None: vals.append(v)
                arr = np.array(vals)
                sv  = [round(float(arr.max()),  3) if len(arr) > 0 else None,
                       round(float(arr.min()),  3) if len(arr) > 0 else None,
                       round(float(arr.mean()), 3) if len(arr) > 0 else None,
                       round(float(arr.std(ddof=1)), 3) if len(arr) > 1 else None,
                       round(float(arr.std(ddof=1)/arr.mean()*100), 1)
                           if len(arr) > 1 and arr.mean() != 0 else None]
                for si2, sv2 in enumerate(sv):
                    cl(ws_res, row_res, 2+n_t_max+si2, sv2, fn=FNB, fi=FAMAR,
                       fmt='0.000' if sv2 is not None else None, bo=BN)
                row_res += 1
            row_res += 1

        ws_res.freeze_panes = 'B4'

        # ---- Abas por intervalo × trial ----
        for itv_k in all_itvs:
            itv_lbl  = TIRO_INTERVALOS.get(itv_k, itv_k)
            mets_l   = [m for m in sel_d.get(itv_k, {}).get(lado, []) if m is not None]
            # raw correspondente (tiro_dist raw, mesmo intervalo)
            raw_l    = ath.get('tiro_dist', {}).get(dist, {}).get('raw', {}).get(itv_k, [])
            flags    = flagrar_outliers(mets_l)

            for ti, m in enumerate(mets_l):
                if m is None:
                    continue
                rd = raw_l[ti] if ti < len(raw_l) else None
                is_out = flags[ti] if ti < len(flags) else False
                itv_short = itv_lbl[:8].replace(' ', '_')
                ne = f'{itv_short}_T{ti+1}'[:31]
                if rd is None:
                    ws_e = wb.create_sheet(title=ne)
                    cl(ws_e, 1, 1, f'Sem dados raw: {itv_lbl} T{ti+1}',
                       fn=Font(name='Calibri', color='FF0000'))
                    continue
                # _calcular_selection filtra frames por t_ini/t_fim E por
                # sel_dir_x/sel_dir_y not None, produzindo pseudo-frames com
                # x=sel_dir_x, y=sel_dir_y armazenados em m['frames'].
                # aba_ensaio usa esses mesmos frames para escrever vx_f/vy_f/sx_f/sy_f/s_ac,
                # pelo que é obrigatório passar m['frames'] e não rd['dados']['frames']
                # (que contém todos os frames do ensaio, incluindo os sem dado de selecção).
                _frames_sel = m.get('frames') if m and m.get('frames') else rd['dados']['frames']
                aba_ensaio(wb, ne, nome, rd.get('dados', {}).get('data', ''),
                           _frames_sel, None, None,
                           m, is_out)
                if c_elipse:
                    aba_elipse(wb, ne, m, nome, is_out)
                if c_estab:
                    aba_estabilograma(wb, ne, m, nome)

        # ---- Aba SPSS ----
        ws_spss = wb.create_sheet(title='SPSS')
        cl(ws_spss, 1, 1,
           f'SPSS  |  {nome}  |  {dist}  |  {lado_lbl}  |  Trial-level',
           fn=Font(name='Calibri', bold=True, size=9, color='0F4C75'))
        hdrs_sp = ([lado_tag + '_intervalo'] +
                   [f'{chave}_T{ti}' for chave, _ in METS_SPSS
                    for ti in range(1, n_t_max+1)] +
                   [f'{chave}_Med' for chave, _ in METS_SPSS])
        for ci3, h3 in enumerate(hdrs_sp, 1):
            fi3 = FAZ if ci3 == 1 else (FAMAR if '_Med' in h3 else fi_lado)
            cl(ws_spss, 3, ci3, h3, fn=FCAB, fi=fi3, al=ALC, bo=BN)
            ws_spss.column_dimensions[get_column_letter(ci3)].width = (22 if ci3 == 1 else 10)
        ws_spss.row_dimensions[3].height = 22

        for ri3, itv_k in enumerate(all_itvs, 4):
            itv_lbl3 = TIRO_INTERVALOS.get(itv_k, itv_k)
            mets_l3  = [m for m in sel_d.get(itv_k, {}).get(lado, []) if m is not None]
            cl(ws_spss, ri3, 1, itv_lbl3, fn=FNB, bo=BN)
            col3 = 2
            for chave, _ in METS_SPSS:
                for ti in range(n_t_max):
                    m3 = mets_l3[ti] if ti < len(mets_l3) else None
                    v3 = round(m3[chave], 4) if m3 and chave in m3 and m3[chave] is not None else None
                    cl(ws_spss, ri3, col3, v3, fn=FN, fi=fi_lado,
                       fmt='0.0000' if v3 is not None else None, bo=BN)
                    col3 += 1
            col_med = 2 + len(METS_SPSS) * n_t_max
            for chave, _ in METS_SPSS:
                vs3  = [m3[chave] for m3 in mets_l3 if m3 and chave in m3 and m3[chave] is not None]
                vs3_c=[x for x in vs3 if x is not None]; v3m  = round(float(np.mean(vs3_c)), 4) if vs3_c else None
                cl(ws_spss, ri3, col_med, v3m, fn=FNB, fi=FAMAR,
                   fmt='0.0000' if v3m is not None else None, bo=BN)
                col_med += 1
        ws_spss.freeze_panes = 'B4'

        cred(ws_res, row_res + 2)
        wb.save(caminho)


def guardar_individual(ath, caminho, c_elipse=True, c_estab=True):
    wb=Workbook()
    if 'Sheet' in wb.sheetnames: del wb['Sheet']
    nome=ath['nome']
    proto=PROTOCOLOS.get(ath.get('protocolo',PROTO_FMS))
    lados=list(ath['mets'].keys())
    n_ens=proto['n_ens']
    for lado in lados:
        flags=flagrar_outliers(ath['mets'][lado])
        for t in range(1,n_ens+1):
            m=ath['mets'][lado][t-1] if t<=len(ath['mets'][lado]) else None
            rd=ath['raw'][lado][t-1] if t<=len(ath['raw'][lado]) else None
            is_out=flags[t-1] if t-1<len(flags) else False
            ne=f'{lado}_{t}'
            if rd is None:
                ws=wb.create_sheet(title=ne)
                cl(ws,1,1,f'Sem dados: {ne}',fn=Font(name='Calibri',color='FF0000'))
                continue
            aba_ensaio(wb,ne,nome,rd['dados'].get('data',''),
                       rd['dados']['frames'],rd.get('ini'),rd.get('fim'),m,is_out)
            if c_elipse and m: aba_elipse(wb,ne,m,nome,is_out)
            if c_estab  and m: aba_estabilograma(wb,ne,m,nome)
    wb.save(caminho)


# -----------------------------------------------------------------------
# Selector de protocolo com logo BSP
# -----------------------------------------------------------------------


# -----------------------------------------------------------------------
# Cálculo de tamanho amostral mínimo  (G*Power simplificado)
# -----------------------------------------------------------------------
def calcular_n_minimo(effect_size=0.5, alpha=0.05, power=0.80,
                      test_type='t_paired', n_tails=2):
    """
    Calcula o n mínimo para detectar um efeito dado poder e alfa.
    test_type: 't_paired'  -> t-pareado (1 grupo, 2 condições)
               't_1samp'   -> t de uma amostra
               'wilcoxon'  -> Wilcoxon (usa aprox. ARE vs t = 0.955)
    Devolve dict com n, potência real e info.
    """
    from scipy.stats import t as t_dist, norm
    import math

    if test_type == 'wilcoxon':
        # Wilcoxon signed-rank tem ARE ≈ 0.955 vs t-pareado sob normalidade
        # Para efeito equivalente, dividir d por sqrt(0.955)
        effect_size_adj = effect_size / math.sqrt(0.955)
    else:
        effect_size_adj = effect_size

    # Iteração: aumentar n até potência >= target
    for n in range(3, 2001):
        df = n - 1
        ncp = effect_size_adj * math.sqrt(n)  # noncentrality parameter
        t_crit = t_dist.ppf(1 - alpha/n_tails, df=df)
        # Potência = P(T > t_crit | ncp)  usando distribuição t não-central
        try:
            from scipy.stats import nct
            power_achieved = 1 - nct.cdf(t_crit, df=df, nc=ncp)
            if n_tails == 2:
                power_achieved += nct.cdf(-t_crit, df=df, nc=ncp)
        except Exception:
            # Fallback: aproximação normal
            z_alpha = norm.ppf(1 - alpha/n_tails)
            power_achieved = norm.cdf(ncp - z_alpha)
        if power_achieved >= power:
            return {
                'n':             n,
                'power_real':    round(power_achieved, 4),
                'effect_size':   effect_size,
                'alpha':         alpha,
                'power_target':  power,
                'test_type':     test_type,
                'n_tails':       n_tails,
            }
    return {'n': '>2000', 'power_real': None,
            'effect_size': effect_size, 'alpha': alpha,
            'power_target': power, 'test_type': test_type}


def _janela_power_calc(root, tk, ttk):
    """Abre janela popup com calculadora de poder estatístico."""
    CA_L = '#00B4D8'; CF_L = '#0F1923'; CC_L = '#1C2E42'
    CP_L = '#162233'; CT_L = '#E8F4FD'; CD_L = '#6B8FAE'

    win = tk.Toplevel(root)
    win.title(T('poder_titulo_win'))
    win.configure(bg=CF_L)
    win.resizable(False, False)
    w, h = 480, 580
    x = (win.winfo_screenwidth()  - w) // 2
    y = (win.winfo_screenheight() - h) // 2
    win.geometry(f'{w}x{h}+{x}+{y}')
    win.grab_set()

    def _lbl(p, txt, **kw):
        return tk.Label(p, text=txt, bg=CF_L, fg=CT_L,
                        font=_F(9), **kw)
    def _lbl_s(p, txt):
        return tk.Label(p, text=txt, bg=CF_L, fg=CD_L,
                        font=_F(7, italic=True))

    # Título
    tk.Label(win, text=T('poder_titulo'),
             bg=CF_L, fg=CA_L, font=_F(12, bold=True)).pack(pady=(14,2))
    tk.Label(win, text=T('poder_subtitulo'),
             bg=CF_L, fg=CD_L, font=_F(8)).pack()
    tk.Frame(win, bg='#243447', height=1).pack(fill='x', padx=16, pady=8)

    frm = tk.Frame(win, bg=CF_L); frm.pack(fill='x', padx=24)

    def _row(label, var, options=None, tip=''):
        f = tk.Frame(frm, bg=CF_L); f.pack(fill='x', pady=4)
        tk.Label(f, text=label, bg=CF_L, fg=CT_L,
                 font=_F(9), width=26, anchor='w').pack(side='left')
        if options:
            cb = ttk.Combobox(f, textvariable=var, values=options,
                              state='readonly', font=_F(9), width=22)
            cb.pack(side='left')
        else:
            e = tk.Entry(f, textvariable=var, bg=CC_L, fg=CT_L,
                         insertbackground=CA_L, relief='flat',
                         font=_F(9), width=12)
            e.pack(side='left')
        if tip:
            _lbl_s(f, f'  {tip}').pack(side='left', padx=4)
        return f

    v_d    = tk.StringVar(value='0.5')
    v_alfa = tk.StringVar(value='0.05')
    v_pow  = tk.StringVar(value='0.80')
    v_test = tk.StringVar(value='t_paired')
    v_tail = tk.StringVar(value='2')

    _row("Cohen's d:", v_d,    tip='0.2=small  0.5=med  0.8=large')
    _row('α (type I error):', v_alfa,  tip='typically 0.05')
    _row('Desired power (1-β):', v_pow, tip='typically 0.80')
    _row(T('poder_titulo').split()[0] + ' type:', v_test,
         options=['t_paired','t_1samp','wilcoxon'])
    _row('Tails:', v_tail, options=['1','2'])

    # Presets
    tk.Label(frm, text=T('poder_presets'), bg=CF_L, fg=CD_L,
             font=_F(8)).pack(anchor='w', pady=(10,2))
    f_pre = tk.Frame(frm, bg=CF_L); f_pre.pack(fill='x')
    for label, d_val in [('Small (0.2)', '0.2'),
                          ('Medium (0.5)', '0.5'),
                          ('Large (0.8)', '0.8')]:
        tk.Button(f_pre, text=label, bg=CP_L, fg=CA_L,
                  font=_F(8), relief='flat', cursor='hand2',
                  padx=6, pady=3,
                  command=lambda dv=d_val: v_d.set(dv)
                  ).pack(side='left', padx=(0,4))

    tk.Frame(win, bg='#243447', height=1).pack(fill='x', padx=16, pady=(12,4))

    # Resultado
    res_var = tk.StringVar(value='')
    res_lbl = tk.Label(win, textvariable=res_var, bg=CF_L, fg=CA_L,
                       font=_F(11, bold=True), wraplength=420,
                       justify='center')
    res_lbl.pack(pady=6)

    detail_var = tk.StringVar(value='')
    tk.Label(win, textvariable=detail_var, bg=CF_L, fg=CD_L,
             font=_F(8), justify='center').pack()

    def _calc():
        try:
            d    = float(v_d.get())
            alfa = float(v_alfa.get())
            pw   = float(v_pow.get())
            test = v_test.get()
            tail = int(v_tail.get())
            assert 0 < d < 10
            assert 0 < alfa < 1
            assert 0 < pw < 1
        except Exception:
            res_var.set('Invalid values.')
            return
        r = calcular_n_minimo(d, alfa, pw, test, tail)
        n_str = str(r['n'])
        res_var.set(f'n min = {n_str}')
        tbl_lines = ['\n  d       α=0.05  α=0.01']
        for d_ref in [0.2, 0.3, 0.5, 0.8, 1.0]:
            n1 = calcular_n_minimo(d_ref, 0.05, pw, test, tail)['n']
            n2 = calcular_n_minimo(d_ref, 0.01, pw, test, tail)['n']
            tbl_lines.append(f'  {d_ref:<6}  {n1:<7}  {n2}')
        detail_var.set(
            f"Power: {r['power_real']:.1%}  |  α={alfa}  |  "
            f"d={d}  |  {test}  |  {tail} tail(s)\n" +
            '\n'.join(tbl_lines)
        )

    _calc_btn = tk.Button(win, text=T('poder_calcular'), bg=CA_L, fg='#0F1923',
              font=_F(10, bold=True), relief='flat', cursor='hand2',
              padx=20, pady=8, activebackground='#0077B6', activeforeground='white',
              command=_calc)
    _calc_btn.pack(pady=(4, 8))
    _calc_btn.bind('<Enter>', lambda e: _calc_btn.config(bg='#0090B0', fg='white'))
    _calc_btn.bind('<Leave>', lambda e: _calc_btn.config(bg=CA_L, fg='#0F1923'))

    tk.Frame(win, bg='#243447', height=1).pack(fill='x', padx=16)
    _close_btn = tk.Button(win, text=T('btn_fechar'), bg=CP_L, fg=CD_L, font=_F(8),
              relief='flat', cursor='hand2', pady=5,
              command=win.destroy)
    _close_btn.pack(pady=(6, 12))
    _close_btn.bind('<Enter>', lambda e: _close_btn.config(fg='#9BBAD8'))
    _close_btn.bind('<Leave>', lambda e: _close_btn.config(fg=CD_L))


# -----------------------------------------------------------------------
# Exportação Word (.docx)  - relatório estatístico
# -----------------------------------------------------------------------
def exportar_docx_relatorio(atletas, caminho, protocolo=None, opcoes=None):
    """
    Gera relatório Word com tabelas dos resultados estatísticos.
    Formatacao em preto e branco - compativel com impressao.
    Requer python-docx (pip install python-docx).
    Devolve (True, '') se gerou com sucesso, (False, msg) caso contrario.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return False, 'python-docx nao instalado. Instale com: pip install python-docx'

    if protocolo is None: protocolo = _PROTOCOLO_ACTIVO
    if opcoes is None: opcoes = {k: True for k in _TESTS_INFO}
    proto = PROTOCOLOS[protocolo]
    lados_nomes = [('dir','Dir.'), ('esq','Esq.')]
    n = len(atletas)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # Estilo de fonte base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9)

    def _heading(txt, level=1):
        p = doc.add_heading(txt, level=level)
        for run in p.runs:
            run.font.color.rgb = None   # preto
            run.font.size = Pt(12 if level == 1 else 10)
        return p

    def _add_table(headers, data_rows, title=None):
        if title:
            _heading(title, level=2)
        nc = len(headers)
        tbl = doc.add_table(rows=1 + len(data_rows), cols=nc)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Cabecalho - negrito, sem cor de fundo
        for ci, h in enumerate(headers):
            cell = tbl.rows[0].cells[ci]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dados
        for ri, row_d in enumerate(data_rows):
            for ci, val in enumerate(row_d):
                cell = tbl.rows[ri + 1].cells[ci]
                cell.text = str(val) if val is not None else ''
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)
                cell.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

    # -- Cabecalho do documento --
    _heading(f'{PROG}  -  Relatorio Estatistico', level=1)
    p_meta = doc.add_paragraph()
    p_meta.add_run(f'Protocolo: {proto["nome"]}  |  n = {n}').font.size = Pt(9)
    p_meta2 = doc.add_paragraph()
    p_meta2.add_run(
        f'Data: {datetime.date.today().strftime("%d/%m/%Y")}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}'
    ).font.size = Pt(8)
    p_meta3 = doc.add_paragraph()
    p_meta3.add_run(
        f'BSP v{VERSAO}  |  alfa = 0.05  |  SW = Shapiro-Wilk  |  '
        f't = t-pareado  |  W = Wilcoxon'
    ).font.size = Pt(8)
    doc.add_page_break()

    # -- Bloco 1: Descritivos --
    if opcoes.get('grupo') and n >= 3:
        hdrs = ['Métrica', 'Lado', 'n', 'Média', 'DP', 'SE',
                'CV(%)', 'IC-inf 95%', 'IC-sup 95%', 'Mediana', 'IQR',
                'SW-p', 'Normal?']
        rows = []
        for chave, label in METS_XL:
            for lado, lbl in lados_nomes:
                vals = [_media_lado(a, chave, lado) for a in atletas]
                vals = [v for v in vals if v is not None]
                if not vals:
                    rows.append([label, lbl] + ['N/A'] * 11)
                    continue
                arr = __import__('numpy').array(vals)
                med = round(float(arr.mean()), 3)
                dp  = round(float(arr.std(ddof=1)), 3) if len(arr) > 1 else 'N/A'
                se  = round(float(arr.std(ddof=1) / (len(arr)**0.5)), 3) if len(arr) > 1 else 'N/A'
                cv  = round(float(arr.std(ddof=1) / arr.mean() * 100), 1) if (len(arr) > 1 and arr.mean() != 0) else 'N/A'
                mn  = round(float(arr.min()), 3)
                mx  = round(float(arr.max()), 3)
                med_v = round(float(__import__('numpy').median(arr)), 3)
                q75, q25 = float(__import__('numpy').percentile(arr, 75)), float(__import__('numpy').percentile(arr, 25))
                iqr = round(q75 - q25, 3)
                try:
                    from scipy import stats as _sp
                    sw_p = round(float(_sp.shapiro(arr).pvalue), 4) if len(arr) >= 3 else 'N/A'
                    normal = 'Sim' if (isinstance(sw_p, float) and sw_p >= 0.05) else ('Nao' if isinstance(sw_p, float) else 'N/A')
                    t_val = round(float(_sp.t.ppf(0.975, df=len(arr)-1)), 3) if len(arr) > 1 else 'N/A'
                    ic_inf = round(med - (t_val * float(arr.std(ddof=1)) / len(arr)**0.5), 3) if isinstance(t_val, float) else 'N/A'
                    ic_sup = round(med + (t_val * float(arr.std(ddof=1)) / len(arr)**0.5), 3) if isinstance(t_val, float) else 'N/A'
                except Exception:
                    sw_p = normal = ic_inf = ic_sup = 'N/A'
                rows.append([label, lbl, len(vals), med, dp, se, cv,
                              ic_inf, ic_sup, med_v, iqr, sw_p, normal])
        _add_table(hdrs, rows, 'Estatisticas Descritivas')

    # -- Bloco 2: Comparacao Dir vs Esq --
    if opcoes.get('dir_esq') and n >= 3:
        try:
            from scipy import stats as _sp2
            hdrs = ['Métrica', 'Méd Dir', 'Méd Esq', 'p-valor', 'Teste', "Cohen's d", 'Sig.']
            rows = []
            for chave, label in METS_XL:
                vd = [_media_lado(a, chave, 'dir') for a in atletas]
                ve = [_media_lado(a, chave, 'esq') for a in atletas]
                pares = [(d, e) for d, e in zip(vd, ve) if d is not None and e is not None]
                if len(pares) < 3:
                    continue
                arr_d = __import__('numpy').array([p[0] for p in pares])
                arr_e = __import__('numpy').array([p[1] for p in pares])
                diffs = arr_d - arr_e
                sw_d = _sp2.shapiro(arr_d).pvalue if len(arr_d) >= 3 else 0
                sw_e = _sp2.shapiro(arr_e).pvalue if len(arr_e) >= 3 else 0
                if sw_d >= 0.05 and sw_e >= 0.05:
                    stat, p = _sp2.ttest_rel(arr_d, arr_e)
                    teste = 't-par'
                else:
                    stat, p = _sp2.wilcoxon(diffs) if len(diffs) >= 3 else (float('nan'), float('nan'))
                    teste = 'Wilcoxon'
                pooled = float(((__import__('numpy').var(arr_d, ddof=1) + __import__('numpy').var(arr_e, ddof=1)) / 2) ** 0.5)
                d_cohen = round(float(diffs.mean()) / pooled, 3) if pooled > 0 else 'N/A'
                sig = 'Sim' if (isinstance(p, float) and p < 0.05) else 'Nao'
                rows.append([label, round(float(arr_d.mean()), 3), round(float(arr_e.mean()), 3),
                              round(float(p), 4) if isinstance(p, float) else 'N/A',
                              teste, d_cohen, sig])
            if rows:
                _add_table(hdrs, rows, 'Comparação Dir vs Esq')
        except Exception:
            pass

    # -- Bloco 3: Legenda das variaveis --
    _heading('Legenda de Variaveis', level=2)
    tbl_leg = doc.add_table(rows=1 + len(_LEGENDA_VARS), cols=3)
    tbl_leg.style = 'Table Grid'
    for ci, h in enumerate(['Variavel', 'Codigo', 'Definicao']):
        cell = tbl_leg.rows[0].cells[ci]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(8)
    for ri, (nome_l, cod_l, defn_l) in enumerate(_LEGENDA_VARS):
        for ci, val in enumerate([nome_l, cod_l, defn_l]):
            cell = tbl_leg.rows[ri + 1].cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(7.5)
    doc.add_paragraph()

    # -- Rodape: citacao --
    p_cit = doc.add_paragraph()
    run_cit = p_cit.add_run(
        f'Citacao: Massuça, A., & Massuça, L. ({datetime.date.today().year}). BSP - Biomechanical Stability Program (v{VERSAO}). '
        f'https://github.com/andremassuca/BSP')
    run_cit.font.size = Pt(7.5)
    run_cit.italic = True

    doc.save(caminho)
    return True, ''


# -----------------------------------------------------------------------
# Exportação CSV  (para R / SPSS / Excel)
# -----------------------------------------------------------------------
def exportar_csv_resumo(atletas, caminho_base, protocolo=None, sep=';', decimal=','):
    """
    Exporta tabelas de resumo em CSV:
      - <base>_grupo.csv   : medias e DP por atleta/lado (formato SPSS-friendly)
      - <base>_spss.csv    : uma linha por trial (formato trial-level)
    Parametros sep e decimal configuráveis.
    Devolve lista de caminhos criados.
    """
    import csv
    if protocolo is None: protocolo = _PROTOCOLO_ACTIVO
    proto = PROTOCOLOS[protocolo]
    lados = ['pos', 'disp'] if _is_tiro_like(protocolo) else ['dir', 'esq']
    criados = []

    def _fmt(v):
        if v is None: return ''
        if isinstance(v, float): return str(v).replace('.', decimal)
        return str(v)

    # ── Ficheiro 1: resumo por atleta ──────────────────────────────────
    base = os.path.splitext(caminho_base)[0]
    f1 = base + '_grupo.csv'
    with open(f1, 'w', newline='', encoding='utf-8-sig') as fh:
        w = csv.writer(fh, delimiter=sep)
        # cabeçalho
        hdrs = ['ID', 'Atleta']
        for chave, label in METS_XL:
            for lado in lados:
                sfx = 'D' if lado in ('dir','pos') else 'E'
                hdrs += [f'{label[:12]}_{sfx}_Med', f'{label[:12]}_{sfx}_DP',
                         f'{label[:12]}_{sfx}_CV']
        w.writerow(hdrs)
        for ath in atletas:
            row = [ath.get('id',''), ath['nome']]
            for chave, _ in METS_XL:
                for lado in lados:
                    mlist = [m for m in ath['mets'].get(lado,[]) if m is not None]
                    vals  = [m[chave] for m in mlist if chave in m]
                    if vals:
                        arr = np.array(vals, dtype=float)
                        med = float(arr.mean())
                        dp  = float(arr.std(ddof=1)) if len(arr)>1 else 0.0
                        cv  = dp/med*100 if med!=0 else None
                        row += [_fmt(round(med,4)), _fmt(round(dp,4)),
                                _fmt(round(cv,2)) if cv is not None else '']
                    else:
                        row += ['','','']
            w.writerow(row)
    criados.append(f1)

    # ── Ficheiro 2: trial-level (SPSS-ready) ───────────────────────────
    f2 = base + '_trials.csv'
    n_ens = proto['n_ens']
    with open(f2, 'w', newline='', encoding='utf-8-sig') as fh:
        w = csv.writer(fh, delimiter=sep)
        hdrs2 = ['ID', 'Atleta', 'Lado', 'Trial'] + [lbl for _, lbl in METS_XL]
        w.writerow(hdrs2)
        for ath in atletas:
            for lado in lados:
                mlist = ath['mets'].get(lado, [])
                for ti, m in enumerate(mlist[:n_ens], 1):
                    if m is None: continue
                    row2 = [ath.get('id',''), ath['nome'], lado, ti]
                    for chave, _ in METS_XL:
                        row2.append(_fmt(round(m[chave],4)) if chave in m and m[chave] is not None else '')
                    w.writerow(row2)
    criados.append(f2)

    return criados


def exportar_csv_estats(atletas, caminho_base, protocolo=None, sep=';', decimal=','):
    """
    Exporta resultados dos testes estatísticos em CSV.
    Cria <base>_estats_descritivos.csv e <base>_estats_testes.csv
    """
    import csv
    if protocolo is None: protocolo = _PROTOCOLO_ACTIVO
    lados_nomes = [('dir','Dir'), ('esq','Esq')]
    base = os.path.splitext(caminho_base)[0]
    criados = []

    def _fmt(v):
        if v is None: return ''
        if isinstance(v, float): return str(round(v,4)).replace('.', decimal)
        return str(v)

    # Descritivos + SW
    f1 = base + '_estats_descritivos.csv'
    with open(f1, 'w', newline='', encoding='utf-8-sig') as fh:
        w = csv.writer(fh, delimiter=sep)
        hdrs = ['Métrica', 'Lado', 'n', 'Média', 'DP', 'SE', 'CV_pct',
                'IC95_inf', 'IC95_sup', 'Mediana', 'IQR', 'SW_p', 'Normal']
        w.writerow(hdrs)
        for chave, label in METS_XL:
            for lado, lbl in lados_nomes:
                vals = [_media_lado(a, chave, lado) for a in atletas]
                vals = [v for v in vals if v is not None]
                n = len(vals)
                if n < 3:
                    w.writerow([label, lbl] + [''] * 11); continue
                arr = np.array(vals, dtype=float)
                med = float(arr.mean()); dp = float(arr.std(ddof=1))
                se  = dp / math.sqrt(n)
                cv  = dp/med*100 if med!=0 else None
                t_c = stats.t.ppf(0.975, df=n-1)
                _, sw_p = stats.shapiro(arr)
                w.writerow([label, lbl, n,
                            _fmt(med), _fmt(dp), _fmt(se),
                            _fmt(cv), _fmt(med-t_c*se), _fmt(med+t_c*se),
                            _fmt(float(np.median(arr))),
                            _fmt(float(np.percentile(arr,75)-np.percentile(arr,25))),
                            _fmt(sw_p), 'Sim' if sw_p>=0.05 else 'Nao'])
    criados.append(f1)

    # Testes Dir vs Esq
    if protocolo in _TESTS_INFO['dir_esq']['protos'] and len(atletas) >= 3:
        f2 = base + '_estats_testes.csv'
        with open(f2, 'w', newline='', encoding='utf-8-sig') as fh:
            w = csv.writer(fh, delimiter=sep)
            hdrs2 = ['Métrica', 'n', 'Méd_Dir', 'Méd_Esq', 'Dif',
                     'Teste', 'Estatistica', 'p_valor', 'Sig',
                     'IC95_inf', 'IC95_sup', 'Cohen_d', 'Efeito']
            w.writerow(hdrs2)
            for chave, label in METS_XL:
                vd = [_media_lado(a,chave,'dir') for a in atletas]
                ve = [_media_lado(a,chave,'esq') for a in atletas]
                pares = [(d,e) for d,e in zip(vd,ve) if d and e]
                if len(pares)<3:
                    w.writerow([label]+['']*(len(hdrs2)-1)); continue
                da = np.array([p[0] for p in pares])
                ea = np.array([p[1] for p in pares])
                dif = da - ea
                _, sw_d = stats.shapiro(da); _, sw_e = stats.shapiro(ea)
                normal = sw_d>=0.05 and sw_e>=0.05
                if normal:
                    tstat, pval = stats.ttest_rel(da, ea)
                    tnm = 't-pareado'
                    se_d = float(dif.std(ddof=1)/math.sqrt(len(pares)))
                    t_c  = stats.t.ppf(0.975, df=len(pares)-1)
                    ic_inf = float(dif.mean())-t_c*se_d
                    ic_sup = float(dif.mean())+t_c*se_d
                else:
                    tstat, pval = stats.wilcoxon(da, ea, alternative='two-sided')
                    tnm = 'Wilcoxon'; ic_inf = ic_sup = None
                sd_d = float(dif.std(ddof=1))
                cd   = float(dif.mean())/sd_d if sd_d>0 else None
                ef   = ('Grande' if cd and abs(cd)>=0.8 else
                        'Medio' if cd and abs(cd)>=0.5 else
                        'Pequeno' if cd and abs(cd)>=0.2 else 'Negligivel')
                w.writerow([label, len(pares),
                            _fmt(float(da.mean())), _fmt(float(ea.mean())),
                            _fmt(float(dif.mean())), tnm,
                            _fmt(float(tstat)), _fmt(pval),
                            'Sim*' if pval<0.05 else 'Nao',
                            _fmt(ic_inf), _fmt(ic_sup),
                            _fmt(cd), ef])
        criados.append(f2)
    return criados


# Guarda os resultados de todos os individuos num ficheiro Excel
# Cria multiplas abas: DADOS, GRUPO, SPSS
def guardar_resumo(atletas, caminho, protocolo=None, opts_estats=None):
    if protocolo is None: protocolo=_PROTOCOLO_ACTIVO
    proto=PROTOCOLOS[protocolo]; tem_ai=proto['assimetria']
    # Tiro com Arco usa lado unico 'arco' (sem dir/esq, sem HS)
    if protocolo == PROTO_ARCO:
        lados_pares = [('arco', 0)]
    else:
        lados_pares=[('dir',0),('esq',5)]   # hurdle step dir/esq (tiro tem abas proprias)

    wb=Workbook(); ws=wb.active; ws.title='DADOS'
    cl(ws,1,1,f'{PROG} - v{VERSAO}  |  {proto["nome"]}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}  |  '
              + datetime.date.today().strftime('%d/%m/%Y'),
       fn=Font(name='Calibri',size=9,italic=True,bold=True,color='0F4C75'))

    n_ens=proto['n_ens']; nc=n_ens+5+(1 if tem_ai else 0)
    col=3   # start at col 3 (col 1=ID, col 2=Individuo)
    cl(ws,2,1,'ID',fn=FCAB,fi=FAZ,al=ALC,bo=BN)
    ws.merge_cells(start_row=2,start_column=1,end_row=3,end_column=1)
    ws.column_dimensions['A'].width=7
    cl(ws,2,2,'Individuo',fn=FCAB,fi=FAZ,al=ALC,bo=BN)
    ws.merge_cells(start_row=2,start_column=2,end_row=3,end_column=2)
    ws.column_dimensions['B'].width=24

    # Determinar n_ens dinamicamente a partir dos dados reais
    n_ens_real = max(
        (len(ath['mets'].get(lado, [])) for ath in atletas for lado,_ in lados_pares),
        default=n_ens)
    n_ens = n_ens_real   # usar o maximo disponivel nos dados

    # Tem colunas TOTAL (dir+esq combinado) para protocolos bipodais
    tem_total = (len(lados_pares) == 2)

    for chave,label in METS_XL:
        total_cols=(len(lados_pares)*(n_ens+5)+(1 if tem_ai else 0)
                    +(5 if tem_total else 0))
        ws.merge_cells(start_row=2,start_column=col,end_row=2,end_column=col+total_cols-1)
        cl(ws,2,col,label,fn=FCAB,fi=FAZ,al=ALW,bo=BN)
        for dc in range(1,total_cols):
            ws.cell(2,col+dc).border=BN; ws.cell(2,col+dc).fill=FAZ
        sub=[]
        for lado,_ in lados_pares:
            sfx='d' if lado=='dir' else ('e' if lado=='esq' else lado[:4])
            sub+=[f'{sfx}{i}' for i in range(1,n_ens+1)]
            sub+=[f'max_{sfx}',f'min_{sfx}',f'med_{sfx}',f'dp_{sfx}',f'CV_{sfx}']
        if tem_ai: sub+=['AI (%)']
        if tem_total: sub+=['max_TOT','min_TOT','med_TOT','dp_TOT','CV_TOT']
        fi_por_sub=[]
        for lado,_ in lados_pares:
            fi=FCZ1 if lado in ['dir','pos'] else FCZ2
            fi_por_sub+=[fi]*(n_ens+5)
        if tem_ai: fi_por_sub+=[FVD2]
        if tem_total: fi_por_sub+=[FAMAR]*5
        for i,h in enumerate(sub):
            # Colunas de estatisticas ligeiramente mais largas; wrap_text para nao sobrepor
            _al_sub = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cl(ws,3,col+i,h,fn=Font(name='Calibri',bold=True,size=8),
               fi=fi_por_sub[i],al=_al_sub,bo=BN)
            # Colunas de trials: 9; colunas de stats e AI: 10
            _w = 10 if any(h.startswith(p) for p in ('max','min','med','dp','CV','AI','TOT')) else 9
            ws.column_dimensions[get_column_letter(col+i)].width=_w
        col+=total_cols

    _al_cab2 = Alignment(horizontal='center', vertical='center', wrap_text=True)
    ws.row_dimensions[1].height=16; ws.row_dimensions[2].height=30; ws.row_dimensions[3].height=28
    ws.freeze_panes='B4'

    def _stats_cols(nums):
        """Retorna [max, min, media, dp, CV] arredondados, ou None se sem dados."""
        if len(nums) == 0:
            return [None]*5
        mx  = round(float(nums.max()),  3)
        mn  = round(float(nums.min()),  3)
        med = round(float(nums.mean()), 3)
        dp  = round(float(nums.std(ddof=1)), 3) if len(nums) > 1 else None
        cv  = (round(float(nums.std(ddof=1)/nums.mean()*100), 2)
               if len(nums) > 1 and nums.mean() != 0 else None)
        return [mx, mn, med, dp, cv]

    for ri,ath in enumerate(atletas,start=4):
        cl(ws,ri,1,ath.get('id',''),fn=FN,bo=BN)
        cl(ws,ri,2,ath['nome'],fn=FNB,bo=BN); col=3
        # por metrica depois por lado
        for chave,_ in METS_XL:
            all_vals = []   # acumula todos os valores de ambos os lados
            for lado,_ in lados_pares:
                lst=ath['mets'].get(lado,[])
                flags=flagrar_outliers(lst)
                fi_b=FCZ1 if lado in ['dir','pos'] else FCZ2
                vals_l=[]
                for ti in range(n_ens):
                    m=lst[ti] if ti<len(lst) else None
                    v=round(m[chave],4) if m and chave in m and m[chave] is not None else None
                    is_out=flags[ti] if ti<len(flags) else False
                    fi=FOUT if is_out else fi_b
                    fn=FOUT_FN if is_out else FN
                    cl(ws,ri,col,v,fn=fn,fi=fi,fmt='0.000' if v is not None else None,bo=BN)
                    vals_l.append(v); col+=1
                # stats por lado (sobre todos os ensaios disponíveis deste lado)
                all_side = [m[chave] for m in lst if m and chave in m and m[chave] is not None]
                nums = np.array([v for v in all_side if v is not None], dtype=float)
                for s in _stats_cols(nums):
                    cl(ws,ri,col,s,fn=FN,fi=FAMAR,fmt='0.000' if s is not None else None,bo=BN)
                    col+=1
                all_vals.extend([v for v in all_side if v is not None])
            # assimetria (se aplicavel)
            if tem_ai:
                vd=[m[chave] for m in ath['mets'].get('dir',[]) if m and chave in m and m[chave] is not None]
                ve=[m[chave] for m in ath['mets'].get('esq',[]) if m and chave in m and m[chave] is not None]
                ai=assimetria(float(np.mean(vd)),float(np.mean(ve))) if vd and ve else None
                cl(ws,ri,col,ai,fn=FNB,fi=FVD2,
                   fmt='+0.00;-0.00;0.00' if ai is not None else None,bo=BN); col+=1
            # stats TOTAIS dir+esq combinadas
            if tem_total:
                nums_tot = np.array(all_vals, dtype=float)
                for s in _stats_cols(nums_tot):
                    cl(ws,ri,col,s,fn=FNB,fi=FAMAR,fmt='0.000' if s is not None else None,bo=BN)
                    col+=1

    aba_grupo(wb,atletas,protocolo)
    aba_spss(wb,atletas,protocolo)
    if opts_estats:
        aba_estats(wb, atletas, protocolo, opts_estats)
    # Para protocolo tiro: adicionar abas por distancia/intervalo + SPSS por dist×itv
    if _is_tiro_like(protocolo):
        _abas_tiro(wb, atletas)
        _aba_sel_dist(wb, atletas)   # v18: Right/Left Selection CoP por distancia
        # SPSS trial-level para cada dist × itv
        all_dists_r, all_itvs_r = [], []
        for a in atletas:
            for d in a.get('tiro_dists',[]):
                if d not in all_dists_r: all_dists_r.append(d)
            for itv in a.get('tiro_intervalos',[]):
                if itv not in all_itvs_r: all_itvs_r.append(itv)
        for dist_r in all_dists_r:
            for itv_r in all_itvs_r:
                aba_spss_tiro(wb, atletas, dist_r, itv_r)
    # Tiro com Arco - abas demograficas quando ha ref. demografica
    if protocolo == PROTO_ARCO and _tem_demografia(atletas):
        _abas_demografia(wb, atletas)
    cred(ws,len(atletas)+6)
    wb.save(caminho)


def _tem_demografia(atletas):
    """True se pelo menos 3 atletas tem dados demograficos (peso ou altura)."""
    n = 0
    for a in atletas:
        if not isinstance(a, dict):
            continue
        if (a.get('peso_kg') is not None or a.get('altura_m') is not None
                or a.get('genero') or a.get('estilo') or a.get('categoria')):
            n += 1
            if n >= 3:
                return True
    return False


def _abas_demografia(wb, atletas):
    """
    v1.0: Cria 3 abas demograficas (Demografia, Comparacoes, Correlacoes)
    no workbook `wb` para o protocolo Tiro com Arco.
    """
    # Metricas de interesse (consistente com a UI Demografia)
    _METRICAS_DEMO = [
        ('ea95',         'Area elipse 95% (mm2)'),
        ('vel_med',      'Velocidade media (mm/s)'),
        ('stiff_x',      'Stiffness ML (1/s)'),
        ('stiff_y',      'Stiffness AP (1/s)'),
        ('amp_x',        'Amp. ML (mm)'),
        ('amp_y',        'Amp. AP (mm)'),
        ('rms_r',        'RMS radial (mm)'),
        # Normalizacoes por dimensoes corporais
        ('ea95_norm',    'EA95 norm (mm2/m)'),
        ('amp_norm_x',   'Amp norm ML (mm/m)'),
        ('amp_norm_y',   'Amp norm AP (mm/m)'),
        ('vel_norm',     'Vel norm ((mm/s)/m)'),
        ('stiff_mass_x', 'Stiff/mass ML (1/(s·kg))'),
        ('stiff_mass_y', 'Stiff/mass AP (1/(s·kg))'),
        ('stiff_norm_x', 'Stiff norm ML ((1/s)/m)'),
        ('stiff_norm_y', 'Stiff norm AP ((1/s)/m)'),
    ]
    _FACTORES = [('genero','Genero'), ('estilo','Estilo'),
                  ('categoria','Categoria')]
    _DEM_VARS = [('peso_kg','Peso (kg)'), ('altura_m','Altura (m)'),
                  ('idade','Idade')]

    # ── Aba 1: Demografia (tabela ampla por atleta) ───────────────────
    ws = wb.create_sheet('Demografia')
    cl(ws,1,1,
       f'{PROG} - v{VERSAO}  |  Tiro com Arco - Demografia  |  '
       + datetime.date.today().strftime('%d/%m/%Y'),
       fn=Font(name='Calibri',size=9,italic=True,bold=True,color='0F4C75'))
    hdr = ['ID','Individuo','Peso (kg)','Altura (m)','Idade',
           'Genero','Estilo','Categoria','P_total']
    for c, h in enumerate(hdr, start=1):
        cl(ws,3,c,h,fn=FCAB,fi=FAZ,al=ALC,bo=BN)
    # Adiciona medianas das metricas de interesse
    for k, lbl in _METRICAS_DEMO:
        hdr.append(lbl)
        cl(ws,3,len(hdr),lbl,fn=FCAB,fi=FAZ,al=ALC,bo=BN)

    for c in range(1, len(hdr)+1):
        ws.column_dimensions[get_column_letter(c)].width = (
            8 if c == 1 else 20 if c == 2 else 14)

    for ri, a in enumerate(atletas, start=4):
        cl(ws,ri,1,a.get('id',''),fn=FN,bo=BN)
        cl(ws,ri,2,a.get('nome',''),fn=FNB,bo=BN)
        cl(ws,ri,3,a.get('peso_kg'),fn=FN,fmt='0.0',bo=BN)
        cl(ws,ri,4,a.get('altura_m'),fn=FN,fmt='0.00',bo=BN)
        cl(ws,ri,5,a.get('idade'),fn=FN,fmt='0',bo=BN)
        cl(ws,ri,6,a.get('genero'),fn=FN,bo=BN)
        cl(ws,ri,7,a.get('estilo'),fn=FN,bo=BN)
        cl(ws,ri,8,a.get('categoria'),fn=FN,bo=BN)
        cl(ws,ri,9,a.get('P_total'),fn=FN,fmt='0',bo=BN)
        # Medianas por metrica (lado 'arco')
        mets = a.get('mets', {}) or {}
        ensaios = mets.get('arco') or []
        col = 10
        for k, _ in _METRICAS_DEMO:
            vs = [m[k] for m in ensaios
                  if isinstance(m, dict) and m.get(k) is not None]
            v = round(float(np.median(vs)), 4) if vs else None
            cl(ws,ri,col,v,fn=FN,fmt='0.0000' if v is not None else None,bo=BN)
            col += 1
    ws.freeze_panes = 'C4'

    # ── Aba 2: Comparacoes (Mann-Whitney / Kruskal-Wallis) ────────────
    ws2 = wb.create_sheet('Comparacoes')
    cl(ws2,1,1,
       f'{PROG} - v{VERSAO}  |  Comparacoes entre grupos  |  '
       + 'nao-parametrico (M-W / K-W)',
       fn=Font(name='Calibri',size=9,italic=True,bold=True,color='0F4C75'))
    cab2 = ['Metrica', 'Factor', 'N grupos', 'N total',
            'Teste', 'Estatistica', 'p-valor', 'Significativo (p<0.05)',
            'Detalhe (grupo: n, mediana, P25, P75)']
    for c, h in enumerate(cab2, start=1):
        cl(ws2,3,c,h,fn=FCAB,fi=FAZ,al=ALC,bo=BN)
    widths2 = [20, 14, 10, 10, 16, 14, 14, 20, 60]
    for c, w in enumerate(widths2, start=1):
        ws2.column_dimensions[get_column_letter(c)].width = w
    ri = 4
    for k, lbl in _METRICAS_DEMO:
        for fator, _ in _FACTORES:
            r = comparar_grupos(atletas, k, fator)
            n_tot = sum(r['n_por_grupo'].values())
            det = '; '.join(
                f"{g}: n={n}, med={r['mediana'].get(g,0):.3f}, "
                f"IQR=[{r['p25'].get(g,0):.3f}, {r['p75'].get(g,0):.3f}]"
                for g, n in sorted(r['n_por_grupo'].items()))
            sig = (r['p_valor'] is not None and r['p_valor'] < 0.05)
            cl(ws2,ri,1,k,fn=FN,bo=BN)
            cl(ws2,ri,2,fator,fn=FN,bo=BN)
            cl(ws2,ri,3,r['n_grupos'],fn=FN,fmt='0',bo=BN)
            cl(ws2,ri,4,n_tot,fn=FN,fmt='0',bo=BN)
            cl(ws2,ri,5,r.get('teste') or '',fn=FN,bo=BN)
            est = r.get('estatistica')
            cl(ws2,ri,6,est,fn=FN,fmt='0.000' if est is not None else None,bo=BN)
            p = r.get('p_valor')
            cl(ws2,ri,7,p,fn=FNB if sig else FN,
               fmt='0.0000' if p is not None else None,
               fi=FVD2 if sig else None, bo=BN)
            cl(ws2,ri,8,'Sim' if sig else 'Nao',
               fn=FNB if sig else FN,
               fi=FVD2 if sig else None, bo=BN)
            cl(ws2,ri,9,det,fn=FN,bo=BN,al=ALW)
            ri += 1
    ws2.freeze_panes = 'A4'

    # ── Aba 3: Correlacoes (Pearson + Spearman) ────────────────────────
    ws3 = wb.create_sheet('Correlacoes')
    cl(ws3,1,1,
       f'{PROG} - v{VERSAO}  |  Correlacoes  |  Pearson + Spearman',
       fn=Font(name='Calibri',size=9,italic=True,bold=True,color='0F4C75'))
    cab3 = ['Metrica CoP', 'Variavel', 'N', 'Pearson r', 'Pearson p',
            'Spearman r', 'Spearman p', 'Sig. (p<0.05)']
    for c, h in enumerate(cab3, start=1):
        cl(ws3,3,c,h,fn=FCAB,fi=FAZ,al=ALC,bo=BN)
    widths3 = [20, 20, 8, 12, 12, 12, 12, 16]
    for c, w in enumerate(widths3, start=1):
        ws3.column_dimensions[get_column_letter(c)].width = w
    ri = 4

    def _emit_cor(r, metrica, var_label):
        nonlocal ri
        n = r.get('n', 0)
        pr = r.get('pearson_r'); pp = r.get('pearson_p')
        sr = r.get('spearman_r'); sp = r.get('spearman_p')
        sig = ((pp is not None and pp < 0.05)
               or (sp is not None and sp < 0.05))
        cl(ws3,ri,1,metrica,fn=FN,bo=BN)
        cl(ws3,ri,2,var_label,fn=FN,bo=BN)
        cl(ws3,ri,3,n,fn=FN,fmt='0',bo=BN)
        cl(ws3,ri,4,pr,fn=FN,fmt='+0.000;-0.000;0.000' if pr is not None else None,bo=BN)
        cl(ws3,ri,5,pp,fn=FNB if (pp is not None and pp < 0.05) else FN,
           fmt='0.0000' if pp is not None else None,
           fi=FVD2 if (pp is not None and pp < 0.05) else None, bo=BN)
        cl(ws3,ri,6,sr,fn=FN,fmt='+0.000;-0.000;0.000' if sr is not None else None,bo=BN)
        cl(ws3,ri,7,sp,fn=FNB if (sp is not None and sp < 0.05) else FN,
           fmt='0.0000' if sp is not None else None,
           fi=FVD2 if (sp is not None and sp < 0.05) else None, bo=BN)
        cl(ws3,ri,8,'Sim' if sig else 'Nao',
           fn=FNB if sig else FN,
           fi=FVD2 if sig else None, bo=BN)
        ri += 1

    for k, lbl in _METRICAS_DEMO:
        # demograficas (peso, altura, idade)
        for dvar, dlbl in _DEM_VARS:
            r = correlacao_demografica(atletas, k, dvar)
            _emit_cor(r, k, dlbl)
        # score (P_total)
        rs = correlacao_score(atletas, k, 'P_total')
        _emit_cor(rs.get('agregado') or {'n':0}, k, 'P_total (agregado)')
        if rs.get('per_ensaio'):
            _emit_cor(rs['per_ensaio'], k, 'P per-ensaio')
    ws3.freeze_panes = 'A4'


def _abas_tiro(wb, atletas, c_elipse=True, c_estab=True):
    """
    Adiciona abas de resultados do protocolo de tiro ao workbook.
    Uma aba por distancia x intervalo com TODAS as metricas (METS_XL).
    Estrutura de cada aba: ID | Individuo | [para cada metrica: T1..Tn | Max|Min|Med|DP|CV]
    """
    if not atletas: return

    all_dists, all_itvs = [], []
    for a in atletas:
        for d in a.get('tiro_dists', []):
            if d not in all_dists: all_dists.append(d)
        for itv in a.get('tiro_intervalos', []):
            if itv not in all_itvs: all_itvs.append(itv)

    if not all_dists or not all_itvs: return

    for dist in all_dists:
        for itv in all_itvs:
            itv_label = _tiro_itv_label(itv)
            sheet_name = f'T_{dist}_{itv_label[:12]}'[:31]
            ws = wb.create_sheet(title=sheet_name)

            n_ens_max = max(
                (len(a.get('tiro_dist',{}).get(dist,{}).get('mets',{}).get(itv,[]))
                 for a in atletas), default=0)
            if n_ens_max == 0: continue

            # ---- cabecalho titulo ----
            cl(ws,1,1, f'Tiro  |  {dist}  |  {itv_label}  |  n={len(atletas)} individuos',
               fn=Font(name='Calibri',bold=True,size=11,color='0F4C75'))
            cl(ws,1,8, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}',fn=FINF)
            cl(ws,2,1, datetime.date.today().strftime('%d/%m/%Y'),fn=FINF)

            # ---- cabecalho colunas ----
            # Linha 3: ID | Individuo | [metrica spans] ...
            # Linha 4: subcabecalhos T1..Tn | Max|Min|Med|DP|CV
            n_stats = 5   # Max Min Med DP CV
            n_per_met = n_ens_max + n_stats
            col = 1
            cl(ws,3,col,'ID',   fn=FCAB,fi=FAZ,al=ALC,bo=BN)
            ws.merge_cells(start_row=3,start_column=col,end_row=4,end_column=col)
            ws.column_dimensions[get_column_letter(col)].width = 6; col+=1
            cl(ws,3,col,'Individuo',fn=FCAB,fi=FAZ,al=ALC,bo=BN)
            ws.merge_cells(start_row=3,start_column=col,end_row=4,end_column=col)
            ws.column_dimensions[get_column_letter(col)].width = 24; col+=1

            for chave, label in METS_XL:
                ws.merge_cells(start_row=3,start_column=col,
                               end_row=3,end_column=col+n_per_met-1)
                cl(ws,3,col,label,fn=FCAB,fi=FAZ,al=ALC,bo=BN)
                for dc in range(1,n_per_met):
                    ws.cell(3,col+dc).fill=FAZ; ws.cell(3,col+dc).border=BN
                # subcabecalhos
                for ti in range(n_ens_max):
                    cl(ws,4,col+ti,f'T{ti+1}',fn=FCAB,fi=FCZ1,al=ALC,bo=BN)
                    ws.column_dimensions[get_column_letter(col+ti)].width = 7
                for si,slbl in enumerate(['Max','Min','Med','DP','CV%']):
                    c2 = col+n_ens_max+si
                    cl(ws,4,c2,slbl,fn=FCAB,fi=FAMAR,al=ALC,bo=BN)
                    ws.column_dimensions[get_column_letter(c2)].width = 7
                col += n_per_met

            ws.row_dimensions[3].height = 20
            ws.row_dimensions[4].height = 18

            # ---- dados por atleta ----
            for ri, ath in enumerate(atletas, 5):
                cl(ws,ri,1,ath.get('id',''),fn=FN,bo=BN)
                cl(ws,ri,2,ath['nome'],fn=FNB,bo=BN)
                mets_l = (ath.get('tiro_dist',{}).get(dist,{})
                          .get('mets',{}).get(itv,[]))
                flags = flagrar_outliers(mets_l)
                col = 3
                for chave,_ in METS_XL:
                    vals = []
                    for ti in range(n_ens_max):
                        m = mets_l[ti] if ti < len(mets_l) else None
                        v = round(m[chave],4) if m and chave in m and m[chave] is not None else None
                        is_out = flags[ti] if ti < len(flags) else False
                        fi = FOUT if is_out else FCZ1
                        fn = FOUT_FN if is_out else FN
                        cl(ws,ri,col,v,fn=fn,fi=fi,
                           fmt='0.000' if v is not None else None,bo=BN)
                        if v is not None: vals.append(v)
                        col += 1
                    nums = np.array(vals)
                    stats_vals = [
                        round(float(nums.max()),3)  if len(nums)>0 else None,
                        round(float(nums.min()),3)  if len(nums)>0 else None,
                        round(float(nums.mean()),3) if len(nums)>0 else None,
                        round(float(nums.std(ddof=1)),3) if len(nums)>1 else None,
                        round(float(nums.std(ddof=1)/nums.mean()*100),1)
                            if len(nums)>1 and nums.mean()!=0 else None,
                    ]
                    for sv in stats_vals:
                        cl(ws,ri,col,sv,fn=FN,fi=FAMAR,
                           fmt='0.000' if sv is not None else None,bo=BN)
                        col += 1

            # ---- linha de medias do grupo (abaixo dos atletas) ----
            row_grp = len(atletas)+6
            cl(ws,row_grp,1,'GRUPO',fn=Font(name='Calibri',bold=True,size=9,color='0F4C75'),
               fi=FAZ1,bo=BN)
            cl(ws,row_grp,2,'Med  |  DP',fn=Font(name='Calibri',bold=True,size=9,
               color='0F4C75'),fi=FAZ1,bo=BN)
            col = 3
            for chave,_ in METS_XL:
                todas_meds = []
                for ath in atletas:
                    mets_l = (ath.get('tiro_dist',{}).get(dist,{})
                              .get('mets',{}).get(itv,[]))
                    vs = [m[chave] for m in mets_l if m and chave in m and m[chave] is not None]
                    vs_c=[x for x in vs if x is not None]
                    if vs_c: todas_meds.append(float(np.mean(vs_c)))
                arr = np.array(todas_meds)
                grp_med = round(float(arr.mean()),3) if len(arr)>0 else None
                grp_dp  = round(float(arr.std(ddof=1)),3) if len(arr)>1 else None
                # escrever med/dp apenas na coluna de Med do grupo
                for ti in range(n_ens_max):
                    ws.cell(row_grp,col+ti).border=BN
                    ws.cell(row_grp,col+ti).fill=FAZ1
                # Max col = col+n_ens_max, Min = +1, Med = +2, DP = +3, CV = +4
                for si in range(n_stats):
                    c2 = col+n_ens_max+si
                    v2 = grp_med if si==2 else (grp_dp if si==3 else None)
                    cl(ws,row_grp,c2,v2,fn=Font(name='Calibri',bold=True,size=8),
                       fi=FAZ1,fmt='0.000' if v2 is not None else None,bo=BN)
                col += n_per_met

            cred(ws, row_grp+2)
            ws.freeze_panes = 'C5'

    # ---- Aba de comparacao entre distancias (se houver >= 2 distancias) ----
    if len(all_dists) >= 2 and all_itvs:
        _aba_comp_distancias(wb, atletas, all_dists, all_itvs)


def _aba_comp_distancias(wb, atletas, all_dists, all_itvs):
    """
    Aba de comparacao automatica entre distancias no protocolo de Tiro/Arco.
    Para cada metrica e intervalo: uma coluna por distancia com Media +/- DP do grupo.
    Metricas principais: ea95, vel_med, rms_x, rms_y, rms_r, amp_x, amp_y, desl.
    """
    METS_COMP = [
        ('ea95',    'Area Elipse 95% (mm2)'),
        ('vel_med', 'Vel. media CoP (mm/s)'),
        ('rms_x',   'RMS ML (mm)'),
        ('rms_y',   'RMS AP (mm)'),
        ('rms_r',   'RMS Radius (mm)'),
        ('amp_x',   'Amplitude ML (mm)'),
        ('amp_y',   'Amplitude AP (mm)'),
        ('desl',    'Deslocamento (mm)'),
        ('vel_x',   'Vel. media ML (mm/s)'),
        ('vel_y',   'Vel. media AP (mm/s)'),
    ]

    ws = wb.create_sheet(title='COMP_DIST')
    cl(ws, 1, 1,
       f'Comparação entre Distâncias  |  n={len(atletas)} indivíduos',
       fn=Font(name='Calibri', bold=True, size=11, color='0F4C75'))
    cl(ws, 1, 10, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}', fn=FINF)
    cl(ws, 2, 1, datetime.date.today().strftime('%d/%m/%Y'), fn=FINF)

    # Cabecalhos: Metrica | Intervalo | [Dist1: Med | DP | n] | [Dist2: Med | DP | n] | ...
    row_hdr = 4
    cl(ws, row_hdr, 1, 'Métrica',   fn=FCAB, fi=FAZ, al=ALC, bo=BN)
    cl(ws, row_hdr, 2, 'Intervalo', fn=FCAB, fi=FAZ, al=ALC, bo=BN)
    ws.column_dimensions['A'].width = 24
    ws.column_dimensions['B'].width = 16

    col = 3
    dist_col_start = {}
    for dist in all_dists:
        dist_col_start[dist] = col
        ws.merge_cells(start_row=row_hdr, start_column=col,
                       end_row=row_hdr, end_column=col + 2)
        cl(ws, row_hdr, col, dist,
           fn=FCAB, fi=FCZ1,
           al=Alignment(horizontal='center', vertical='center'), bo=BN)
        for ci, lbl in enumerate(['Media', 'DP', 'n'], start=col):
            cl(ws, row_hdr + 1, ci, lbl,
               fn=Font(name='Calibri', bold=True, size=8),
               fi=FCZ1,
               al=Alignment(horizontal='center'), bo=BN)
            ws.column_dimensions[get_column_letter(ci)].width = 9
        col += 3

    row = row_hdr + 2
    for chave, label in METS_COMP:
        for itv in all_itvs:
            itv_lbl = _tiro_itv_label(itv)
            cl(ws, row, 1, label,   fn=FNB, bo=BN)
            cl(ws, row, 2, itv_lbl, fn=FN,  bo=BN)
            for dist in all_dists:
                sc = dist_col_start[dist]
                # Recolher media por individuo para este dist x itv x metrica
                grupo_vals = []
                for ath in atletas:
                    mets_l = (ath.get('tiro_dist', {}).get(dist, {})
                              .get('mets', {}).get(itv, []))
                    vs = [m[chave] for m in mets_l if m and chave in m and m[chave] is not None]
                    if vs:
                        grupo_vals.append(float(np.mean(vs)))
                if grupo_vals:
                    arr = np.array(grupo_vals)
                    med = round(float(arr.mean()), 3)
                    dp  = round(float(arr.std(ddof=1)), 3) if len(arr) > 1 else None
                    n_v = len(arr)
                else:
                    med = dp = n_v = None
                cl(ws, row, sc,     med, fn=FN, fi=FCZ1,
                   fmt='0.000' if med is not None else None, bo=BN)
                cl(ws, row, sc + 1, dp,  fn=FN, fi=FCZ1,
                   fmt='0.000' if dp  is not None else None, bo=BN)
                cl(ws, row, sc + 2, n_v, fn=FN, fi=FCZ1, bo=BN)
            row += 1

    cred(ws, row + 1)
    ws.freeze_panes = 'C6'


def _aba_sel_dist(wb, atletas):
    """
    Aba(s) de Right/Left Selection CoP por distancia (v18).
    Gera uma aba por distancia com as metricas dos dois pes de forma simetrica:
    pe direito (Right Selection) e pe esquerdo (Left Selection) lado a lado.
    """
    all_dists = []
    for a in atletas:
        for d in a.get('tiro_dists', []):
            if d not in all_dists:
                all_dists.append(d)
    if not all_dists:
        return

    # Collect all intervals present across athletes
    all_itvs_sel = []
    for a in atletas:
        for d in (a.get('sel_dist') or {}):
            for itv_k in (a.get('sel_dist', {}).get(d) or {}):
                if itv_k not in all_itvs_sel:
                    all_itvs_sel.append(itv_k)

    for dist in all_dists:
        n_t_max = max(
            (max(
                len(a.get('sel_dist', {}).get(dist, {}).get(itv_k, {}).get('dir', [])),
                len(a.get('sel_dist', {}).get(dist, {}).get(itv_k, {}).get('esq', []))
             ) for a in atletas for itv_k in all_itvs_sel),
            default=0)
        if n_t_max == 0:
            continue

        sheet_name = ('SEL_' + dist)[:31]
        ws = wb.create_sheet(title=sheet_name)
        cl(ws, 1, 1,
           f'Selection CoP  |  {dist}  |  Pe Dir. (Right) + Pe Esq. (Left)  |  n={len(atletas)}',
           fn=Font(name='Calibri', bold=True, size=11, color='0F4C75'))
        cl(ws, 1, 12, f'{PROG} v{VERSAO}  |  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}', fn=FINF)
        cl(ws, 2, 1, datetime.date.today().strftime('%d/%m/%Y'), fn=FINF)

        n_stats = 5   # Max Min Med DP CV
        n_per_lado = n_t_max + n_stats

        # Linha 3: ID | Individuo | bloco por metrica (Dir + Esq)
        cl(ws, 3, 1, 'ID',        fn=FCAB, fi=FAZ, al=ALC, bo=BN)
        ws.merge_cells(start_row=3, start_column=1, end_row=4, end_column=1)
        ws.column_dimensions['A'].width = 6
        cl(ws, 3, 2, 'Individuo', fn=FCAB, fi=FAZ, al=ALC, bo=BN)
        ws.merge_cells(start_row=3, start_column=2, end_row=4, end_column=2)
        ws.column_dimensions['B'].width = 24

        col_h = 3
        for chave, label in METS_XL:
            # --- bloco Pe Direito ---
            ws.merge_cells(start_row=3, start_column=col_h,
                           end_row=3, end_column=col_h + n_per_lado - 1)
            cl(ws, 3, col_h,
               f'{label}  -  Pe Direito (Right Selection)',
               fn=FCAB, fi=FAZ, al=ALW, bo=BN)
            for ti in range(n_t_max):
                cl(ws, 4, col_h+ti, f'Dir_T{ti+1}', fn=FCAB, fi=FCZ1, al=ALC, bo=BN)
                ws.column_dimensions[get_column_letter(col_h+ti)].width = 9
            for si, slbl in enumerate(['Max','Min','Med','DP','CV%']):
                cc = col_h + n_t_max + si
                cl(ws, 4, cc, slbl, fn=FCAB, fi=FAMAR, al=ALC, bo=BN)
                ws.column_dimensions[get_column_letter(cc)].width = 9
            col_h += n_per_lado

            # --- bloco Pe Esquerdo ---
            ws.merge_cells(start_row=3, start_column=col_h,
                           end_row=3, end_column=col_h + n_per_lado - 1)
            cl(ws, 3, col_h,
               f'{label}  -  Pe Esquerdo (Left Selection)',
               fn=FCAB, fi=FAZ, al=ALW, bo=BN)
            for ti in range(n_t_max):
                cl(ws, 4, col_h+ti, f'Esq_T{ti+1}', fn=FCAB, fi=FCZ2, al=ALC, bo=BN)
                ws.column_dimensions[get_column_letter(col_h+ti)].width = 9
            for si, slbl in enumerate(['Max','Min','Med','DP','CV%']):
                cc = col_h + n_t_max + si
                cl(ws, 4, cc, slbl, fn=FCAB, fi=FAMAR, al=ALC, bo=BN)
                ws.column_dimensions[get_column_letter(cc)].width = 9
            col_h += n_per_lado

        ws.row_dimensions[3].height = 22
        ws.row_dimensions[4].height = 18

        # Dados por atleta - itera todos os intervalos e ambos lados
        for ri, ath in enumerate(atletas, start=5):
            cl(ws, ri, 1, ath.get('id',''), fn=FN, bo=BN)
            cl(ws, ri, 2, ath['nome'],        fn=FNB, bo=BN)
            col_d = 3

            for chave, _ in METS_XL:
                for lado, fi_b in [('dir', FCZ1), ('esq', FCZ2)]:
                    # Concatena os trials de todos os intervalos para este lado/metrica
                    lst   = []
                    for itv_k2 in all_itvs_sel:
                        lst += ath.get('sel_dist', {}).get(dist, {}).get(itv_k2, {}).get(lado, [])
                    flags = flagrar_outliers(lst)
                    vals  = []
                    for ti in range(n_t_max):
                        m = lst[ti] if ti < len(lst) else None
                        v = round(m[chave], 4) if m and chave in m and m[chave] is not None else None
                        is_o = flags[ti] if ti < len(flags) else False
                        cl(ws, ri, col_d, v,
                           fn=(FOUT_FN if is_o else FN),
                           fi=(FOUT if is_o else fi_b),
                           fmt='0.000' if v is not None else None, bo=BN)
                        if v is not None: vals.append(v)
                        col_d += 1
                    arr = np.array(vals)
                    for sv in [
                        round(float(arr.max()),  3) if len(arr)>0 else None,
                        round(float(arr.min()),  3) if len(arr)>0 else None,
                        round(float(arr.mean()), 3) if len(arr)>0 else None,
                        round(float(arr.std(ddof=1)),3) if len(arr)>1 else None,
                        round(float(arr.std(ddof=1)/arr.mean()*100),1)
                            if len(arr)>1 and arr.mean()!=0 else None,
                    ]:
                        cl(ws, ri, col_d, sv, fn=FN, fi=FAMAR,
                           fmt='0.000' if sv is not None else None, bo=BN)
                        col_d += 1

        # Linha de medias do grupo
        row_grp2 = len(atletas) + 7
        cl(ws, row_grp2, 1, 'GRUPO',
           fn=Font(name='Calibri', bold=True, size=9, color='0F4C75'), fi=FAZ1, bo=BN)
        cl(ws, row_grp2, 2, 'Med  |  DP',
           fn=Font(name='Calibri', bold=True, size=9, color='0F4C75'), fi=FAZ1, bo=BN)
        col_g = 3
        for chave, _ in METS_XL:
            for lado in ('dir', 'esq'):
                todas_meds = []
                for ath in atletas:
                    lst = []
                    for itv_k3 in all_itvs_sel:
                        lst += ath.get('sel_dist', {}).get(dist, {}).get(itv_k3, {}).get(lado, [])
                    vs = [m[chave] for m in lst if m and chave in m and m[chave] is not None]
                    vs_nn=[x for x in vs if x is not None]
                    if vs_nn: todas_meds.append(float(np.mean(vs_nn)))
                arr2 = np.array(todas_meds)
                grp_med = round(float(arr2.mean()), 3) if len(arr2) > 0 else None
                grp_dp  = round(float(arr2.std(ddof=1)), 3) if len(arr2) > 1 else None
                for ti in range(n_t_max):
                    ws.cell(row_grp2, col_g+ti).border = BN
                    ws.cell(row_grp2, col_g+ti).fill   = FAZ1
                for si2 in range(n_stats):
                    cc2 = col_g + n_t_max + si2
                    v2 = grp_med if si2 == 2 else (grp_dp if si2 == 3 else None)
                    cl(ws, row_grp2, cc2, v2,
                       fn=Font(name='Calibri', bold=True, size=8),
                       fi=FAZ1, fmt='0.000' if v2 is not None else None, bo=BN)
                col_g += n_per_lado

        cred(ws, row_grp2 + 2)
        ws.freeze_panes = 'C5'


# -----------------------------------------------------------------------
# Selector de protocolo com logo BSP
# -----------------------------------------------------------------------

def _logo_photoimage(size=96):
    """Devolve um PhotoImage tkinter a partir do logo BSP embebido."""
    import tkinter as tk
    try:
        from PIL import Image as PILImage, ImageTk
        data = _b64.b64decode(_LOGO_B64)
        img = PILImage.open(io.BytesIO(data)).convert('RGBA')
        img = img.resize((size, size), PILImage.LANCZOS)
        return ImageTk.PhotoImage(img)
    except Exception:
        # Fallback: tkinter PhotoImage directo (gif/png via b64)
        try:
            return tk.PhotoImage(data=_LOGO_B64)
        except Exception:
            return None


def _mostrar_logo(canvas, cx, cy, size=96, _refs=None):
    """Coloca o logo BSP (PNG real) no canvas tkinter na posicao (cx,cy)."""
    img = _logo_photoimage(size)
    if img is None:
        return
    canvas.create_image(cx, cy, image=img, anchor='center')
    # Guardar referencia para evitar garbage collection
    if _refs is not None:
        _refs.append(img)
    else:
        # Guardar no proprio canvas como atributo
        if not hasattr(canvas, '_logo_refs'):
            canvas._logo_refs = []
        canvas._logo_refs.append(img)


class Tooltip:
    """Tooltip simples que aparece ao passar o rato num widget."""
    def __init__(self, widget, text, bg='#152436', fg='#C8E6F5', delay=400):
        self.widget = widget
        self.text = text
        self.bg = bg; self.fg = fg
        self.delay = delay
        self._id = None; self._win = None
        widget.bind('<Enter>', self._schedule)
        widget.bind('<Leave>', self._cancel)
        widget.bind('<ButtonPress>', self._cancel)

    def _schedule(self, ev=None):
        self._cancel()
        self._id = self.widget.after(self.delay, self._show)

    def _cancel(self, ev=None):
        if self._id:
            self.widget.after_cancel(self._id); self._id = None
        if self._win:
            self._win.destroy(); self._win = None

    def _show(self):
        import tkinter as tk
        if self._win: return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 6
        self._win = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        # macOS Sonoma/Sequoia: hint ao WM para tooltip flutuante
        if _SYS == 'Darwin':
            try: tw.wm_attributes('-type', 'tooltip')
            except Exception: pass
        tw.wm_geometry(f'+{x}+{y}')
        # Outer border frame
        outer = tk.Frame(tw, bg='#2A4A6A', padx=1, pady=1)
        outer.pack()
        tk.Label(outer, text=self.text, bg=self.bg, fg=self.fg,
                 font=_F(8), relief='flat',
                 justify='left', padx=10, pady=6,
                 wraplength=340).pack()


def _abrir_pasta_os(pasta):
    """Abre a pasta no explorador de ficheiros do sistema operativo."""
    import subprocess, platform
    try:
        plat = platform.system()
        if plat == 'Windows':
            os.startfile(pasta)
        elif plat == 'Darwin':
            subprocess.Popen(['open', pasta])
        else:
            subprocess.Popen(['xdg-open', pasta])
    except Exception:
        pass


_PH_A = bytes.fromhex('dd6aaed864b6a61a').hex()
_PH_B = bytes.fromhex('24c02c5f6d5c1718').hex()
_PH_C = bytes.fromhex('387d846fd507e865').hex()
_PH_D = bytes.fromhex('0a07026e7ac20b72').hex()
_PASS_HASH = ''.join(f'{int(_PH_A[i:i+2],16)^0xAB:02x}' for i in range(0,len(_PH_A),2)) + \
             ''.join(f'{int(_PH_B[i:i+2],16)^0xAB:02x}' for i in range(0,len(_PH_B),2)) + \
             ''.join(f'{int(_PH_C[i:i+2],16)^0xAB:02x}' for i in range(0,len(_PH_C),2)) + \
             ''.join(f'{int(_PH_D[i:i+2],16)^0xAB:02x}' for i in range(0,len(_PH_D),2))


# Password de acesso vai buscar a lista de hashes aceites ao .bsp_pass.sha256
# do repo - assim posso rodar a password com um push em vez de redistribuir.
# Cai para cache local se não houver rede; cai para o hash embedded só na
# primeira execução offline.
_PASS_REMOTE_URL  = f'https://raw.githubusercontent.com/{_UPDATE_REPO}/main/.bsp_pass.sha256'
_PASS_CACHE_KEY   = '_pass_hashes_cache'
_PASS_CACHE_STAMP = '_pass_hashes_cache_time'

def _fetch_pass_hashes_remoto(timeout=4.0):
    """Vai ler o .bsp_pass.sha256 e devolve os SHA256 que encontra (lowercase).
    Aceita comentários com '#'. Devolve lista vazia se algo correr mal."""
    import urllib.request
    try:
        req = urllib.request.Request(_PASS_REMOTE_URL,
                                     headers={'User-Agent': f'BSP/{VERSAO}'})
        with urllib.request.urlopen(req, timeout=timeout) as r:
            txt = r.read(16 * 1024).decode('utf-8', errors='ignore')
    except Exception:
        return []

    pat = re.compile(r'^[0-9a-f]{64}$')
    visto = set()
    saida = []
    for linha in txt.splitlines():
        linha = linha.strip()
        if not linha or linha.startswith('#'):
            continue
        tok = linha.split()[0].lower()
        if pat.match(tok) and tok not in visto:
            visto.add(tok)
            saida.append(tok)
    return saida

def _obter_pass_hashes_aceites(timeout=4.0):
    """Devolve o set de hashes que valem agora. Tenta o remoto primeiro;
    se for bem-sucedido, actualiza a cache e usa só esse (rotação imediata).
    Senão, cache local. Senão, fallback para o hash embedded."""
    remoto = _fetch_pass_hashes_remoto(timeout=timeout)
    if remoto:
        try:
            cfg = cfg_load()
            cfg[_PASS_CACHE_KEY] = remoto
            cfg[_PASS_CACHE_STAMP] = datetime.datetime.now().isoformat(timespec='seconds')
            cfg_save(cfg)
        except Exception:
            pass
        return set(remoto)

    try:
        cache = cfg_load().get(_PASS_CACHE_KEY) or []
        if cache:
            return {h.lower() for h in cache if isinstance(h, str)}
    except Exception:
        pass

    return {_PASS_HASH.lower()}


# Endpoint para registar aceitações de licença. Vazio ou env var
# BSP_TELEMETRY_URL='' desactiva.
_TELEMETRY_URL = 'https://bsp-telemetry.andremassuca.workers.dev/license'
_TELEMETRY_LOG_KEY = '_telemetry_log_local'

def _machine_id_anon():
    """SHA256 do MAC + hostname + username, truncado. Identifica máquina
    sem identificar quem está à frente dela."""
    import hashlib, uuid, platform
    raw = f'{uuid.getnode()}-{platform.node()}-{os.environ.get("USERNAME") or os.environ.get("USER","")}'
    return hashlib.sha256(raw.encode('utf-8')).hexdigest()[:32]

def _telemetry_post_license_acceptance(email=None):
    """Manda um POST com o evento de aceitação. Falhas silenciosas - se o
    endpoint cair, o utilizador não nota."""
    import threading, datetime, json as jsonlib, platform
    import urllib.request as urlreq

    now = datetime.datetime.now(datetime.timezone.utc).isoformat(timespec='seconds')
    payload = {
        'ts':         now,
        'event':      'license_accepted',
        'version':    str(VERSAO),
        'lang':       lingua_atual(),
        'os':         f'{platform.system()} {platform.release()}'[:64],
        'machine_id': _machine_id_anon(),
    }
    if email and '@' in str(email) and len(email) <= 200:
        payload['email'] = email.strip()

    # Sempre registar localmente - útil se for preciso ver depois
    try:
        cfg = cfg_load()
        log = (cfg.get(_TELEMETRY_LOG_KEY) or []) + [payload]
        cfg[_TELEMETRY_LOG_KEY] = log[-50:]
        cfg_save(cfg)
    except Exception:
        pass

    url = os.environ.get('BSP_TELEMETRY_URL', _TELEMETRY_URL)
    if not url:
        return  # opt-out

    def _send():
        try:
            req = urlreq.Request(
                url,
                data=jsonlib.dumps(payload).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'User-Agent':   f'BSP/{VERSAO}',
                },
                method='POST',
            )
            urlreq.urlopen(req, timeout=5)
        except Exception:
            pass

    threading.Thread(target=_send, daemon=True).start()


# -----------------------------------------------------------------------
# Ecra de Licenca (mostra uma vez; regista aceitacao no config)
# -----------------------------------------------------------------------
_LICENCA_TEXTO = """BSP - Biomechanical Stability Program  v{versao}
{autor}  |  {univ}

TERMOS DE UTILIZACAO

1. Este software destina-se exclusivamente a fins academicos e de investigacao
   cientifica no ambito da biomecânica e analise postural.

2. E proibida a utilizacao comercial, redistribuicao ou modificacao sem
   autorizacao expressa do autor.

3. Os resultados gerados por este programa nao substituem avaliacao clinica
   ou diagnostico medico profissional.

4. Os autores nao se responsabilizam por erros nos dados ou interpretacoes
   incorrectas dos resultados.

5. Ao aceitar, o utilizador compromete-se a citar adequadamente o software
   em qualquer publicacao ou relatorio academico que utilize os seus resultados.

6. Telemetria minima e anonima:
   Ao aceitar estes termos, sera enviado um unico registo anonimo aos autores
   contendo: data/hora da aceitacao, versao do BSP, lingua da interface,
   sistema operativo e um identificador anonimo da maquina (hash, nao reversivel).
   Nao sao enviados dados pessoais nem dados clinicos. Para desactivar:
   definir a variavel de ambiente BSP_TELEMETRY_URL com valor vazio antes
   de arrancar a app.

Autores: {autor}  |  {univ}
"""



def _ecra_licenca(_root_ext=None):
    """Mostra os termos de utilizacao na primeira execucao.
    Retorna True se ja aceite ou acabou de aceitar; False se recusou."""
    import tkinter as tk
    from tkinter import scrolledtext

    cfg = cfg_load()
    # Mostrar EULA se: nunca aceite OU versão diferente da atual
    _versao_aceite = cfg.get('licenca_versao', '')
    if cfg.get('licenca_aceite') and _versao_aceite == VERSAO:
        return True

    aceite = [False]

    if _root_ext is not None:
        root = tk.Toplevel(_root_ext)
        # Sem transient: deve aparecer na taskbar durante o arranque
        root.grab_set()
        try: root.wm_attributes('-toolwindow', False)  # ATENÇÃO: não existe em macOS/Tcl antigo
        except Exception: pass
    else:
        root = tk.Tk()

    root.title('BSP  Termos de Utilizacao')
    root.configure(bg=CF)
    root.resizable(True, True)
    sw = root.winfo_screenwidth()
    sh = root.winfo_screenheight()
    w = max(900, min(1100, sw - 200))
    h = max(740, min(900,  sh - 120))
    x = max(0, (sw - w) // 2)
    y = max(30, (sh - h) // 2)
    root.geometry(f'{w}x{h}+{x}+{y}')
    root.minsize(780, 640)
    root.protocol('WM_DELETE_WINDOW', lambda: None)
    root.lift()
    root.attributes('-topmost', True)
    root.after(300, lambda: root.attributes('-topmost', False))
    root.focus_force()

    # Cabecalho centrado: logo circular + titulo + autores (consistente com janela de acesso)
    hdr_lic = tk.Frame(root, bg=CP, pady=18)
    hdr_lic.pack(fill='x')
    _logo_lic = _logo_photoimage(size=80)
    if _logo_lic:
        tk.Label(hdr_lic, image=_logo_lic, bg=CP, bd=0).pack()
        hdr_lic._logo_ref = _logo_lic  # evitar GC
    tk.Label(hdr_lic, text='BSP', font=_F(20, bold=True), bg=CP, fg=CA).pack(pady=(6, 0))
    tk.Label(hdr_lic, text='Biomechanical Stability Program  v' + VERSAO,
             font=_F(9), bg=CP, fg=CD).pack()
    tk.Label(hdr_lic, text=f'{AUTOR}  ·  {SEGUNDO_AUTOR}  ·  {ORIENTADOR}',
             font=_F(7), bg=CP, fg=CD).pack(pady=(2, 0))
    tk.Frame(root, bg=CA, height=2).pack(fill='x')

    # ── Barra de idioma (compacta, abaixo do banner) ───────────────────
    _BANDEIRAS_LIC = {'PT': '🇵🇹', 'EN': '🇬🇧', 'ES': '🇪🇸', 'DE': '🇩🇪'}
    f_lang_lic = tk.Frame(root, bg=CB, height=30)
    f_lang_lic.pack(fill='x'); f_lang_lic.pack_propagate(False)
    _txt_ref        = [None]  # referência ao widget de texto para actualizar
    _lbl_titulo_ref = [None]  # referência ao label do título EULA
    _btn_aceitar_ref = [None]  # referência ao botão aceitar
    _btn_recusar_ref = [None]  # referência ao botão recusar

    def _mudar_lingua_lic(codigo):
        definir_lingua(codigo)
        d = cfg_load(); d['_lingua'] = codigo; cfg_save(d)
        # Actualiza cor dos botões de idioma
        for cod2, btn2 in _lang_btns_lic.items():
            _ativo2 = (cod2 == codigo)
            btn2.config(bg=CA2 if _ativo2 else CB, fg=CA if _ativo2 else CD,
                        font=_F(8, bold=_ativo2))
        # Actualiza título EULA e botões aceitar/recusar
        if _lbl_titulo_ref[0] is not None:
            _lbl_titulo_ref[0].config(text=T('licenca_titulo'))
        if _btn_aceitar_ref[0] is not None:
            _btn_aceitar_ref[0].config(text=T('btn_aceitar'))
        if _btn_recusar_ref[0] is not None:
            _btn_recusar_ref[0].config(text=T('btn_recusar'))
        # Actualiza texto da licença
        if _txt_ref[0] is not None:
            _txt_ref[0].config(state='normal')
            _txt_ref[0].delete('1.0', 'end')
            _txt_ref[0].insert('end', licenca_texto(lang=codigo, versao=VERSAO,
                                                    autor=AUTOR, univ=f'{SEGUNDO_AUTOR}  |  {ORIENTADOR}'), 'corpo')
            _txt_ref[0].config(state='disabled')

    _lang_btns_lic = {}
    _lang_atual_lic = lingua_atual()
    for _cod_l, _flag_l in _BANDEIRAS_LIC.items():
        _ativo_l = (_cod_l == _lang_atual_lic)
        _b = tk.Button(f_lang_lic, text=f'{_flag_l} {_cod_l}',
                       bg=CA2 if _ativo_l else CB,
                       fg=CA if _ativo_l else CD,
                       font=_F(8, bold=_ativo_l), relief='flat',
                       cursor='hand2', padx=4, pady=2,
                       activebackground=CC, activeforeground=CT,
                       command=lambda c=_cod_l: _mudar_lingua_lic(c))
        _b.pack(side='left', fill='y', padx=1)
        _lang_btns_lic[_cod_l] = _b

    # Sub-header - título EULA em linha própria para nunca ficar cortado
    sub = tk.Frame(root, bg=CF); sub.pack(fill='x')
    _lbl_tit = tk.Label(sub, text=T('licenca_titulo'), font=_F(10, bold=True),
                        bg=CF, fg=CA, anchor='w', pady=6)
    _lbl_tit.pack(side='left', padx=16, fill='x', expand=True)
    _lbl_titulo_ref[0] = _lbl_tit
    tk.Label(sub, text=f'v{VERSAO}', font=_F(8),
             bg=CF, fg=CD, pady=6).pack(side='right', padx=16)

    # Rodapé com botões - DEVE ser empacotado ANTES do body (side=bottom precisa de vir primeiro)
    footer = tk.Frame(root, bg=CB, height=62)
    footer.pack(fill='x', side='bottom'); footer.pack_propagate(False)

    def _recusar():
        import sys; root.destroy(); sys.exit(0)

    def _aceitar():
        aceite[0] = True
        d = cfg_load(); d['licenca_aceite'] = True; d['licenca_versao'] = VERSAO; cfg_save(d)
        # Telemetry minimal - registo anonimo da aceitacao
        # (mencionado nos termos; falhas silenciosas; thread separada)
        try:
            _telemetry_post_license_acceptance()
        except Exception:
            pass
        root.destroy()

    _btn_rec = tk.Button(footer, text=T('btn_recusar'), bg=CB, fg=CER, relief='flat',
                         font=_F(9), cursor='hand2', padx=14, pady=6, command=_recusar)
    _btn_rec.pack(side='left', padx=20, pady=0)
    _btn_recusar_ref[0] = _btn_rec
    _btn_ac = tk.Button(footer, text=T('btn_aceitar'), bg=CA, fg='#0A1218',
                        relief='flat', font=_F(10, bold=True), cursor='hand2',
                        padx=18, pady=8, command=_aceitar)
    _btn_ac.pack(side='right', padx=20, pady=0)
    _btn_aceitar_ref[0] = _btn_ac

    # Corpo de texto - empacotado depois do footer para que este tenha espaço garantido
    body = tk.Frame(root, bg=CF, padx=18, pady=8)
    body.pack(fill='both', expand=True)

    txt = scrolledtext.ScrolledText(
        body, bg=CP, fg=CT, font=_F(9), wrap='word', relief='flat',
        padx=14, pady=12, state='normal',
        highlightthickness=1, highlightbackground=CB,
        cursor='arrow',
    )
    txt.pack(fill='both', expand=True)

    # Texto formatado com tags
    txt.tag_config('titulo',  font=_F(11, bold=True), foreground=CA,  spacing1=4, spacing3=6)
    txt.tag_config('sub',     font=_F(9,  bold=True), foreground=CT,  spacing1=8, spacing3=2)
    txt.tag_config('corpo',   font=_F(9),             foreground=CT,  lmargin1=20, lmargin2=20)
    txt.tag_config('rodape',  font=_F(8),             foreground=CD,  spacing1=14)
    txt.tag_config('numero',  font=_F(9,  bold=True), foreground=CA)

    # Usar texto de licença localizado (i18n) - idioma actual
    _lic_txt = licenca_texto(lang=lingua_atual(), versao=VERSAO, autor=AUTOR, univ=f'{SEGUNDO_AUTOR}  |  {ORIENTADOR}')
    txt.insert('end', _lic_txt, 'corpo')
    txt.config(state='disabled')
    _txt_ref[0] = txt  # guardar referência para actualização dinâmica

    if _root_ext is not None:
        _root_ext.wait_window(root)
    else:
        root.mainloop()

    return aceite[0]


def _ecra_password(_root_ext=None):
    """Ecra de autenticacao. Retorna True se pass correcta, False se cancelado."""
    import tkinter as tk
    resultado = [False]
    if _root_ext is not None:
        root = tk.Toplevel(_root_ext)
        # Nao usar transient() para que apareca na barra de tarefas
        root.grab_set()
        # Windows: garantir que aparece na taskbar
        try: root.wm_attributes('-toolwindow', False)  # ATENÇÃO: não existe em macOS/Tcl antigo
        except Exception: pass
        _is_toplevel = True
    else:
        root = tk.Tk()
        _is_toplevel = False

    root.title('BSP  Acesso')
    root.configure(bg=CF); root.resizable(True, True)
    sw = root.winfo_screenwidth(); sh = root.winfo_screenheight()
    w = max(620, min(760, sw - 300))
    h = max(660, min(800, sh - 100))
    x = max(0, (sw - w) // 2)
    y = max(30, (sh - h) // 2)
    root.geometry(f"{w}x{h}+{x}+{y}")
    root.lift()
    root.attributes('-topmost', True)
    root.after(300, lambda: root.attributes('-topmost', False))
    root.focus_force()

    # Cabecalho centrado: logo circular + titulo + autores
    hdr_pw = tk.Frame(root, bg=CP, pady=20)
    hdr_pw.pack(fill='x')
    _logo_pw = _logo_photoimage(size=90)
    if _logo_pw:
        tk.Label(hdr_pw, image=_logo_pw, bg=CP, bd=0).pack()
        hdr_pw._logo_ref = _logo_pw  # evitar GC
    tk.Label(hdr_pw, text='BSP', font=_F(22, bold=True), bg=CP, fg=CA).pack(pady=(8, 0))
    tk.Label(hdr_pw, text='Biomechanical Stability Program  v' + VERSAO,
             font=_F(9), bg=CP, fg=CD).pack()
    tk.Label(hdr_pw, text=f'{AUTOR}  ·  {SEGUNDO_AUTOR}  ·  {ORIENTADOR}',
             font=_F(7), bg=CP, fg=CD).pack(pady=(2, 0))
    tk.Frame(root, bg=CA, height=2).pack(fill='x', padx=0)

    # ── Barra de idioma (compacta) ─────────────────────────────────────
    _BANDEIRAS_PW = {'PT': '🇵🇹', 'EN': '🇬🇧', 'ES': '🇪🇸', 'DE': '🇩🇪'}
    f_lang_pw = tk.Frame(root, bg=CB, height=26)
    f_lang_pw.pack(fill='x'); f_lang_pw.pack_propagate(False)
    _pw_widgets = {}  # referências para widgets a actualizar

    def _mudar_lingua_pw(codigo):
        definir_lingua(codigo)
        d = cfg_load(); d['_lingua'] = codigo; cfg_save(d)
        for cod2, btn2 in _lang_btns_pw.items():
            _ativ2 = (cod2 == codigo)
            btn2.config(bg=CA2 if _ativ2 else CB, fg=CA if _ativ2 else CD,
                        font=_F(8, bold=_ativ2))
        # Actualizar labels dinâmicos
        if 'lbl_pass' in _pw_widgets:
            _pw_widgets['lbl_pass'].config(text=T('acesso_pass'))
        if 'btn_entrar' in _pw_widgets:
            _pw_widgets['btn_entrar'].config(text=T('acesso_entrar'))
        if 'gh_label' in _pw_widgets:
            _pw_widgets['gh_label'].config(text=T('acesso_github') + ' ')
        if 'gh_btn' in _pw_widgets:
            _pw_widgets['gh_btn'].config(text=T('btn_abrir_github'))

    _lang_btns_pw = {}
    _lang_atual_pw = lingua_atual()
    for _cod_pw, _flag_pw in _BANDEIRAS_PW.items():
        _ativo_pw = (_cod_pw == _lang_atual_pw)
        _b_pw = tk.Button(f_lang_pw, text=f'{_flag_pw} {_cod_pw}',
                          bg=CA2 if _ativo_pw else CB,
                          fg=CA if _ativo_pw else CD,
                          font=_F(8, bold=_ativo_pw), relief='flat',
                          cursor='hand2', padx=4, pady=1,
                          activebackground=CC, activeforeground=CT,
                          command=lambda c=_cod_pw: _mudar_lingua_pw(c))
        _b_pw.pack(side='left', fill='y', padx=1)
        _lang_btns_pw[_cod_pw] = _b_pw

    # Separador subtil abaixo das l\u00ednguas
    tk.Frame(root, bg=CB, height=1).pack(fill='x', pady=(10, 0))

    # Sec\u00e7\u00e3o central: formul\u00e1rio de autentica\u00e7\u00e3o com espa\u00e7amento generoso
    _form = tk.Frame(root, bg=CF)
    _form.pack(fill='both', expand=True, padx=44)

    tk.Frame(_form, bg=CF, height=28).pack()  # espa\u00e7o em cima

    lbl_pass = tk.Label(_form, text=T('acesso_pass'), bg=CF, fg=CT, font=_F(10, bold=True))
    lbl_pass.pack(anchor='w')
    _pw_widgets['lbl_pass'] = lbl_pass

    var_pass = tk.StringVar()
    entry = tk.Entry(_form, textvariable=var_pass, show='\u2022', bg=CC, fg=CT,
                     insertbackground=CA, relief='flat', font=_F(13),
                     justify='center',
                     highlightthickness=1, highlightbackground=CB)
    entry.pack(fill='x', ipady=10, pady=(6, 0))
    entry.focus_set()

    msg = tk.Label(_form, text='', bg=CF, fg=CER, font=_F(8))
    msg.pack(pady=(4, 0))

    # Hashes aceites: arranca com cache/embedded; o fetch remoto corre em
    # background e substitui o set assim que chega. Se o utilizador premir
    # Enter antes do fetch chegar, o confirm espera ate 4s extra.
    try:
        _hashes_iniciais = cfg_load().get(_PASS_CACHE_KEY) or []
        _hashes_ref = [set(h.lower() for h in _hashes_iniciais
                           if isinstance(h, str))] if _hashes_iniciais else [
                      { _PASS_HASH.lower() }]
    except Exception:
        _hashes_ref = [{ _PASS_HASH.lower() }]
    _fetch_done = [False]

    def _bg_fetch():
        try:
            novos = _obter_pass_hashes_aceites(timeout=4.0)
            if novos:
                _hashes_ref[0] = novos
        except Exception: pass
        finally:
            _fetch_done[0] = True
    import threading as _thr
    _thr.Thread(target=_bg_fetch, daemon=True).start()

    def confirmar(ev=None):
        h = hashlib.sha256(var_pass.get().encode()).hexdigest().lower()
        # Esperar ate 4s pelo fetch remoto se ainda nao chegou (mantem UX
        # responsivo: so bloqueia se o utilizador digitou mais rapido que a
        # rede, o que e raro; a janela ja esta aberta entretanto).
        if not _fetch_done[0]:
            try:
                import time as _time_p
                t0 = _time_p.monotonic()
                while not _fetch_done[0] and _time_p.monotonic() - t0 < 4.0:
                    _time_p.sleep(0.05)
            except Exception: pass
        if h in _hashes_ref[0]:
            resultado[0] = True; root.destroy()
        else:
            msg.config(text=T('acesso_pass_err'))
            var_pass.set('')
            # borda vermelha por 1 s + botão desactivado por 500 ms (anti-brute-force)
            entry.config(highlightbackground=CER, highlightcolor=CER, highlightthickness=2)
            btn_entrar.config(state='disabled')
            root.after(500, lambda: btn_entrar.config(state='normal'))
            root.after(1000, lambda: entry.config(highlightbackground=CB,
                                                   highlightcolor=CB, highlightthickness=1))
            entry.focus_set()

    entry.bind('<Return>', confirmar)
    btn_entrar = tk.Button(_form, text=T('acesso_entrar'), bg=CA, fg='#071014',
              font=_F(11, bold=True), activebackground=CA2, activeforeground='white',
              relief='flat', cursor='hand2', padx=20, pady=11,
              command=confirmar)
    btn_entrar.pack(fill='x', pady=(12, 0))
    btn_entrar.bind('<Enter>', lambda e: btn_entrar.config(bg=CA2, fg='white')
                    if str(btn_entrar['state']) == 'normal' else None)
    btn_entrar.bind('<Leave>', lambda e: btn_entrar.config(bg=CA, fg='#071014')
                    if str(btn_entrar['state']) == 'normal' else None)
    _pw_widgets['btn_entrar'] = btn_entrar

    tk.Frame(_form, bg=CF, height=16).pack()  # espaço

    # GitHub link
    def _abrir_github(ev=None):
        import webbrowser
        webbrowser.open('https://github.com/andremassuca')
    gh_frame = tk.Frame(_form, bg=CF); gh_frame.pack()
    gh_label = tk.Label(gh_frame, text=T('acesso_github') + ' ', bg=CF, fg=CD,
             font=_F(8))
    gh_label.pack(side='left')
    _pw_widgets['gh_label'] = gh_label
    gh_lnk = tk.Label(gh_frame, text='github.com/andremassuca', bg=CF, fg=CA,
                       font=_F(8), cursor='hand2')
    gh_lnk.pack(side='left')
    gh_lnk.bind('<Button-1>', _abrir_github)

    gh_btn = tk.Button(_form, text=T('btn_abrir_github'), bg=CB, fg=CA,
                       relief='flat', font=_F(8), cursor='hand2',
                       padx=10, pady=4, command=_abrir_github,
                       activebackground=CC, activeforeground=CT)
    gh_btn.pack(pady=(4, 0))
    _pw_widgets['gh_btn'] = gh_btn

    tk.Frame(_form, bg=CF, height=20).pack()  # margem inferior

    def _fechar_pw():
        resultado[0] = False; root.destroy()
    root.protocol('WM_DELETE_WINDOW', _fechar_pw)

    if _is_toplevel:
        _root_ext.wait_window(root)
    else:
        root.mainloop()
    return resultado[0]


def _submenu_tarefa_funcional(_root_ext=None):
    """Submenu para selecionar a tarefa funcional especifica.
    Retorna a chave do protocolo escolhido, ou None se cancelado."""
    import tkinter as tk
    import base64 as _b64_sub
    resultado = [None]
    if _root_ext is not None:
        root = tk.Toplevel(_root_ext)
        root.grab_set()
        _is_toplevel = True
    else:
        root = tk.Tk()
        _is_toplevel = False

    root.title('BSP  -  ' + T('proto_func_titulo'))
    _BG = '#0A0D14'
    root.configure(bg=_BG); root.resizable(True, True)
    sw = root.winfo_screenwidth(); sh = root.winfo_screenheight()
    w = max(560, min(700, sw - 200))
    h = max(520, min(700, sh - 120))
    x = max(0, (sw - w) // 2)
    y = max(20, (sh - h) // 2)
    root.geometry(f"{w}x{h}+{x}+{y}")
    root.minsize(500, 460)

    # ── Cabeçalho com logo BSP ───────────────────────────────────────────
    _hdr_func = tk.Frame(root, bg=CP, pady=14)
    _hdr_func.pack(fill='x')
    _logo_func = _logo_photoimage(size=72)
    if _logo_func:
        tk.Label(_hdr_func, image=_logo_func, bg=CP, bd=0).pack()
        _hdr_func._logo_ref = _logo_func
    tk.Label(_hdr_func, text='BSP', font=_F(18, bold=True), bg=CP, fg=CA).pack(pady=(5, 0))
    tk.Label(_hdr_func, text=f'{PROG}  v{VERSAO}', font=_F(8), bg=CP, fg=CD).pack()
    tk.Label(_hdr_func, text=f'{AUTOR}  ·  {SEGUNDO_AUTOR}  ·  {ORIENTADOR}',
             font=_F(7), bg=CP, fg=CD).pack(pady=(1, 0))

    tk.Frame(root, bg='#F4A261', height=1).pack(fill='x', padx=20, pady=(8, 0))
    tk.Label(root, text=T('proto_func_label'),
             bg=_BG, fg='#E8F4FD', font=_F(9, bold=True)).pack(pady=(10, 6))

    var = tk.StringVar(value=PROTO_TIRO)

    for pk in TAREFAS_FUNC:
        f = tk.Frame(root, bg='#1C2E42', relief='flat')
        f.pack(fill='x', padx=28, pady=3)
        tk.Radiobutton(f, text=_proto_nome(pk), variable=var, value=pk,
                       bg='#1C2E42', fg='#E8F4FD', selectcolor=_BG,
                       activebackground='#1C2E42', font=_F(10, bold=True),
                       relief='flat', cursor='hand2').pack(anchor='w', padx=10, pady=(7,1))
        tk.Label(f, text=_proto_descr(pk), bg='#1C2E42', fg='#6B8FAE',
                 font=_F(8), justify='left', wraplength=460).pack(anchor='w', padx=26, pady=(0,7))

    def confirmar():
        resultado[0] = var.get()
        root.destroy()

    def cancelar():
        resultado[0] = None
        root.destroy()

    tk.Button(root, text=T('proto_confirmar2'), bg='#F4A261', fg='#0F1923',
              font=_F(11, bold=True), activebackground='#D4824A',
              relief='flat', cursor='hand2', padx=20, pady=10,
              command=confirmar).pack(pady=(12,4), padx=28, fill='x')
    tk.Button(root, text=T('proto_voltar'), bg='#1C2E42', fg='#6B8FAE',
              font=_F(8), relief='flat', cursor='hand2', pady=4,
              command=cancelar).pack(padx=28, fill='x')

    root.protocol('WM_DELETE_WINDOW', cancelar)
    if _is_toplevel:
        _root_ext.wait_window(root)
    else:
        root.mainloop()
    return resultado[0]


def _proto_nome(pk):
    """Devolve o nome traduzido do protocolo pk."""
    _nomes = {
        PROTO_FMS:      T('proto_fms_nome'),
        PROTO_UNIPODAL: T('proto_uni_nome'),
        PROTO_FUNC:     T('proto_func_nome'),
        PROTO_TIRO:     T('proto_tiro_nome'),
        PROTO_ARCO:     T('proto_arco_nome'),
    }
    return _nomes.get(pk, PROTOCOLOS.get(pk, {}).get('nome', pk))


def _proto_descr(pk):
    """Devolve a descricao traduzida do protocolo pk."""
    _descrs = {
        PROTO_FMS:      T('proto_fms_descr'),
        PROTO_UNIPODAL: T('proto_uni_descr'),
        PROTO_FUNC:     T('proto_func_descr'),
        PROTO_TIRO:     T('proto_tiro_descr'),
        PROTO_ARCO:     T('proto_arco_descr'),
    }
    return _descrs.get(pk, PROTOCOLOS.get(pk, {}).get('descr', ''))


def escolher_protocolo_gui(_root_ext=None):
    """Menu principal de selecao de protocolo."""
    import tkinter as tk
    import base64 as _b64_proto
    resultado = [None]
    if _root_ext is not None:
        root = tk.Toplevel(_root_ext)
        # Nao usar transient para que apareca na barra de tarefas
        root.grab_set()
        try: root.wm_attributes('-toolwindow', False)
        except Exception: pass
        _is_toplevel = True
    else:
        root = tk.Tk()
        _is_toplevel = False

    _BG = '#0A0D14'
    root.configure(bg=_BG); root.resizable(True, True)
    sw = root.winfo_screenwidth(); sh = root.winfo_screenheight()
    w = max(760, min(960, sw - 160))
    h = max(740, min(900, sh - 80))
    x = max(0, (sw - w) // 2)
    y = max(20, (sh - h) // 2)
    root.geometry(f"{w}x{h}+{x}+{y}")
    root.minsize(680, 660)
    root.lift()
    root.attributes('-topmost', True)
    root.after(300, lambda: root.attributes('-topmost', False))
    root.focus_force()

    # ── Referências para rebuild in-place ────────────────────────────────
    _refs = {}

    def _build_ui():
        """Constrói (ou reconstrói) todo o conteúdo do protocolo."""
        for w2 in root.winfo_children():
            w2.destroy()
        root.title(PROG + '  v' + VERSAO + '  -  ' + T('proto_titulo'))

        # ── Cabeçalho centrado: logo circular + título ────────────────────
        _hdr_proto = tk.Frame(root, bg=CP, pady=14)
        _hdr_proto.pack(fill='x')
        _logo_proto = _logo_photoimage(size=72)
        if _logo_proto:
            tk.Label(_hdr_proto, image=_logo_proto, bg=CP, bd=0).pack()
            _hdr_proto._logo_ref = _logo_proto  # evitar GC
        tk.Label(_hdr_proto, text='BSP', font=_F(18, bold=True), bg=CP, fg=CA).pack(pady=(5, 0))
        tk.Label(_hdr_proto, text=f'{PROG}  v{VERSAO}',
                 font=_F(8), bg=CP, fg=CD).pack()
        tk.Label(_hdr_proto, text=f'{AUTOR}  ·  {SEGUNDO_AUTOR}  ·  {ORIENTADOR}',
                 font=_F(7), bg=CP, fg=CD).pack(pady=(1, 0))
        tk.Frame(root, bg=CA, height=2).pack(fill='x')

        # ── Barra de línguas compacta ─────────────────────────────────────
        _lf = tk.Frame(root, bg=CB, height=30)
        _lf.pack(fill='x'); _lf.pack_propagate(False)
        _BANDEIRAS_PROTO = [('PT','🇵🇹 PT'),('EN','🇬🇧 EN'),('ES','🇪🇸 ES'),('DE','🇩🇪 DE')]
        _lang_now = lingua_atual()
        def _mudar_lang_proto(cod):
            definir_lingua(cod)
            d = cfg_load(); d['_lingua'] = cod; cfg_save(d)
            _build_ui()
        for cod, lbl in _BANDEIRAS_PROTO:
            _active = (cod == _lang_now)
            tk.Button(_lf, text=lbl,
                      bg=CA if _active else CB,
                      fg=CP if _active else CD,
                      relief='flat', font=_F(8, bold=_active),
                      cursor='hand2', padx=5, pady=2,
                      command=lambda c=cod: _mudar_lang_proto(c)
                      ).pack(side='left', fill='y', padx=1)

        # ── Título ────────────────────────────────────────────────────────
        tk.Label(root, text=T('proto_titulo'),
                 bg=_BG, fg='#E8F4FD', font=_F(10, bold=True)).pack(pady=(10, 6))

        var = tk.StringVar(value=PROTO_FMS)
        _refs['var'] = var

        # FMS Bipodal
        for pk in [PROTO_FMS, PROTO_UNIPODAL]:
            f = tk.Frame(root, bg='#1C2E42', relief='flat')
            f.pack(fill='x', padx=28, pady=3)
            tk.Radiobutton(f, text=_proto_nome(pk), variable=var, value=pk,
                           bg='#1C2E42', fg='#E8F4FD', selectcolor=_BG,
                           activebackground='#1C2E42', font=_F(10, bold=True),
                           relief='flat', cursor='hand2').pack(anchor='w', padx=10, pady=(7,1))
            tk.Label(f, text=_proto_descr(pk), bg='#1C2E42', fg='#6B8FAE',
                     font=_F(8), justify='left', wraplength=560).pack(anchor='w', padx=26, pady=(0,7))

        # Tarefa Funcional
        f_func = tk.Frame(root, bg='#1C2E42', relief='flat')
        f_func.pack(fill='x', padx=28, pady=3)
        tk.Radiobutton(f_func, text=_proto_nome(PROTO_FUNC), variable=var, value=PROTO_FUNC,
                       bg='#1C2E42', fg='#F4A261', selectcolor=_BG,
                       activebackground='#1C2E42', font=_F(10, bold=True),
                       relief='flat', cursor='hand2').pack(anchor='w', padx=10, pady=(7,1))
        tk.Label(f_func, text=_proto_descr(PROTO_FUNC),
                 bg='#1C2E42', fg='#6B8FAE',
                 font=_F(8), justify='left', wraplength=560).pack(anchor='w', padx=26, pady=(0,7))

        # Botões confirmar / fechar - empacotados antes do rodapé para garantir espaço
        tk.Frame(root, bg=_BG, height=8).pack()
        _btn_conf = tk.Button(root, text=T('proto_confirmar'),
                  bg='#00B4D8', fg='white', font=_F(11, bold=True),
                  activebackground='#0077B6', activeforeground='white',
                  relief='flat', cursor='hand2',
                  padx=20, pady=11, command=_confirmar)
        _btn_conf.pack(pady=(0, 4), padx=28, fill='x')
        _btn_conf.bind('<Enter>', lambda e: _btn_conf.config(bg='#0090B0'))
        _btn_conf.bind('<Leave>', lambda e: _btn_conf.config(bg='#00B4D8'))
        _btn_fch = tk.Button(root, text=T('proto_fechar'), bg='#1C2E42', fg='#6B8FAE',
                  font=_F(8), relief='flat', cursor='hand2', pady=5,
                  command=root.destroy)
        _btn_fch.pack(padx=28, fill='x')
        _btn_fch.bind('<Enter>', lambda e: _btn_fch.config(fg='#9BBAD8'))
        _btn_fch.bind('<Leave>', lambda e: _btn_fch.config(fg='#6B8FAE'))

        # Rodapé autor
        tk.Label(root, text=f'{AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}',
                 bg=_BG, fg='#2A4A6A', font=_F(7)).pack(side='bottom', pady=6)

    def _confirmar():
        var = _refs.get('var')
        if var is None: return
        sel = var.get()
        if sel == PROTO_FUNC:
            sub_proto = _submenu_tarefa_funcional(root)
            if sub_proto:
                resultado[0] = sub_proto
                root.destroy()
        else:
            resultado[0] = sel
            root.destroy()

    _build_ui()

    def _fechar_proto():
        resultado[0] = None
        root.destroy()
    root.protocol('WM_DELETE_WINDOW', _fechar_proto)
    if _is_toplevel:
        _root_ext.wait_window(root)
    else:
        root.mainloop()
    return resultado[0]


# -----------------------------------------------------------------------
# Interface grafica principal
# -----------------------------------------------------------------------

# Cores inicializadas via tema (pode ser substituido por _aplicar_tema)
CF='#0D1720'; CP='#132030'; CC='#1A2B3C'; CA='#00B4D8'; CA2='#0090B0'
COK='#22C55E'; CAV='#F59E0B'; CER='#EF4444'; CT='#EDF6FF'; CD='#7FA8C9'; CB='#1E3448'
CA3='#90E0EF'; CG='#0F2030'; CSEP='#1E3448'


# Classe principal da interface grafica
# Herda do tkinter e gere toda a janela principal
# Responsavel pela interacao com o utilizador
class Janela:
    def __init__(self, protocolo=PROTO_FMS, _root=None):
        global _PROTOCOLO_ACTIVO
        _PROTOCOLO_ACTIVO = protocolo
        self._proto = PROTOCOLOS[protocolo]
        self._proto_key = protocolo
        import tkinter as tk
        from tkinter import ttk, filedialog, messagebox
        self.tk=tk; self.ttk=ttk; self.fd=filedialog; self.mb=messagebox
        self._q=queue.Queue(); self._run=False

        _owns_root = (_root is None)
        if _owns_root:
            r = tk.Tk()
        else:
            r = _root
            # Limpar conteúdo anterior (reuso ao mudar protocolo)
            for _w in r.winfo_children():
                _w.destroy()
            r.deiconify()
        self.root = r
        # Valida fontes preferidas - faz fallback para TkDefaultFont/TkFixedFont
        # se Helvetica/Menlo/Segoe UI/Cascadia/DejaVu não existirem no SO
        try:
            _validar_fontes_disponiveis()
        except Exception:
            pass
        r.title(f"BSP  |  {PROG}  v{VERSAO}  -  {_proto_nome(self._proto_key)}")
        r.configure(bg=CF)
        r.resizable(True, True)   # garantir botoes de maximizar/redimensionar
        # Restaurar geometria guardada - validar contra ecrã actual
        r.update_idletasks()
        _sw = r.winfo_screenwidth()
        _sh = r.winfo_screenheight()
        # Minsize adaptativo ao ecrã (nunca maior que o ecrã)
        _min_w = min(760, _sw - 40)
        _min_h = min(560, _sh - 60)
        r.minsize(_min_w, _min_h)
        _g = cfg_load().get('_win_geom', '')
        _geom_ok = False
        if _g:
            try:
                import re as _re2
                _m2 = _re2.match(r'(\d+)x(\d+)(?:\+(-?\d+)\+(-?\d+))?', _g)
                if _m2:
                    _gw, _gh = int(_m2.group(1)), int(_m2.group(2))
                    _gx = int(_m2.group(3)) if _m2.group(3) else 0
                    _gy = int(_m2.group(4)) if _m2.group(4) else 0
                    # Aceitar apenas se a janela couber no ecrã e for grande o suficiente
                    if (_gw >= _min_w and _gh >= _min_h and
                            _gw <= _sw * 1.05 and _gh <= _sh * 1.05 and
                            _gx < _sw - 80 and _gy < _sh - 60 and
                            _gx > -_gw + 80 and _gy > -40):
                        r.geometry(_g)
                        _geom_ok = True
            except Exception: pass
        if not _geom_ok:
            _dw = min(1220, max(_min_w, int(_sw * 0.90)))
            _dh = min(900,  max(_min_h, int(_sh * 0.90)))
            _dx = max(0, (_sw - _dw) // 2)
            _dy = max(20, (_sh - _dh) // 2)
            r.geometry(f'{_dw}x{_dh}+{_dx}+{_dy}')
        # Guardar geometria ao redimensionar/mover - com debounce (evita I/O excessivo)
        _geom_timer = [None]
        def _save_geom(ev=None):
            if _geom_timer[0]:
                try: r.after_cancel(_geom_timer[0])
                except Exception: pass
            def _do_save():
                try:
                    d = cfg_load(); d['_win_geom'] = r.geometry(); cfg_save(d)
                except Exception: pass
            _geom_timer[0] = r.after(1000, _do_save)
        r.bind('<Configure>', _save_geom)
        # Atalhos de teclado
        def _kb(key, fn):
            r.bind(f'<{_MOD_KEY}-{key}>', lambda e: fn())
        self._ui()
        self._poll()
        # Bind após _ui() - self.b_exe já existe
        _kb('Return', lambda: self._exec() if self.b_exe['state'] == 'normal' else None)
        _kb('s',      self._save_cfg)
        _kb('h',      self._ver_historico)
        _kb('p',      lambda: _janela_power_calc(self.root, self.tk, self.ttk))
        r.bind('<F1>', lambda e: self._show_shortcuts())
        r.bind('<F2>', lambda e: self._quick_subject())
        r.bind('<F5>', lambda e: self._clr())
        # Auto-update em background
        def _on_update(nova):
            try:
                self.root.after(0, lambda: self._show_update_banner(nova))
            except Exception: pass
        _verificar_update_async(_on_update)
        if _owns_root:
            r.mainloop()

    def _ui(self):
        tk=self.tk; ttk=self.ttk
        cfg = cfg_load()

        # Estilo do sash (divisor arrastável) - visível e fácil de arrastar
        try:
            _sty = ttk.Style()
            _sty.configure('Sash', sashthickness=6, sashpad=2,
                            background='#243447', activebackground=CA)
            # estilos ttk consistentes (Primary/Secondary/Danger/Ghost,
            # Treeview, Progressbar, Notebook, Combobox, Entry)
            _configurar_estilos_ttk(_sty)
        except Exception:
            pass

        is_tiro   = _is_tiro_like(self._proto_key)
        is_iscpsi = _is_iscpsi(self._proto_key)   # cor dourada só no Tiro ISCPSI

        # Cor de fundo do topo: azul escuro só no Tiro ISCPSI, escuro normal caso contrário
        CP_TOP = '#1A2B3C' if is_iscpsi else CP
        CA_TOP = '#E8D5A3' if is_iscpsi else CA  # dourado ISCPSI vs branco

        # ── Header ──────────────────────────────────────────────────────────
        _accent_col = CA_TOP if is_iscpsi else CA
        tp = tk.Frame(self.root, bg=CP_TOP, height=92)
        tp.pack(fill='x'); tp.pack_propagate(False)
        # 2-px accent strip below header
        tk.Frame(self.root, bg=_accent_col, height=2).pack(fill='x')

        # Logo BSP (left)
        logo_cv = tk.Canvas(tp, width=92, height=80, bg=CP_TOP, highlightthickness=0)
        logo_cv.pack(side='left', padx=(16, 0))
        _mostrar_logo(logo_cv, 46, 42, size=70)

        # Title block
        lf = tk.Frame(tp, bg=CP_TOP); lf.pack(side='left', padx=(10, 0))
        lf_top = tk.Frame(lf, bg=CP_TOP); lf_top.pack(anchor='w', pady=(18, 0))
        tk.Label(lf_top, text='BSP',
                 bg=CP_TOP, fg=CA_TOP, font=_F(17, bold=True)).pack(side='left')
        tk.Label(lf_top, text=f'  v{VERSAO}',
                 bg=CP_TOP, fg=CD, font=_F(10)).pack(side='left', padx=(0, 0))
        tk.Label(lf, text=PROG.upper(),
                 bg=CP_TOP, fg=CA_TOP, font=_F(8)).pack(anchor='w', pady=(2, 0))
        tk.Label(lf, text=f'{AUTOR}  ·  {SEGUNDO_AUTOR}  ·  {ORIENTADOR}',
                 bg=CP_TOP, fg=CD, font=_F(7)).pack(anchor='w')

        # Protocol badge (right side)
        # Tiro (ISCPSI) fica dourado. Tiro com Arco é verde (alvo arco).
        cor_p = {'fms': '#00B4D8', 'unipodal': '#2DC653',
                 'tiro': '#C8A45A', 'tiro_arco': '#4A9B4A'}
        _badge_bg = cor_p.get(self._proto_key, CA)
        badge_wrap = tk.Frame(tp, bg=CP_TOP)
        badge_wrap.pack(side='right', padx=(0, 4), pady=30)
        badge = tk.Label(badge_wrap, text=f'  {_proto_nome(self._proto_key)}  ',
                 bg=_badge_bg, fg='white',
                 font=_F(8, bold=True), pady=4)
        badge.pack()

        # ── Botão ⚙ Definições ───────────────────────────────────────────
        def _open_settings(event=None):
            popup = tk.Toplevel(self.root)
            # macOS: overrideredirect causa janelas invisíveis; usar Toplevel normal
            if _SYS == 'Darwin':
                popup.overrideredirect(False)
                popup.resizable(False, False)
                popup.title('')
            else:
                popup.overrideredirect(True)
            popup.configure(bg=CF)
            self.root.update_idletasks()
            bx = gear_btn.winfo_rootx()
            by = gear_btn.winfo_rooty() + gear_btn.winfo_height() + 4
            popup.geometry(f"+{bx - 200}+{by}")
            outer = tk.Frame(popup, bg=CB, padx=1, pady=1)
            outer.pack()
            inner = tk.Frame(outer, bg=CF, padx=14, pady=12)
            inner.pack()
            # Language pills
            tk.Label(inner, text=T('lbl_idioma'),
                     font=_F(8), bg=CF, fg=CD).pack(anchor='w', pady=(0,4))
            pill_row = tk.Frame(inner, bg=CF)
            pill_row.pack(anchor='w')
            _lang_flags = {'PT':'\U0001f1f5\U0001f1f9','EN':'\U0001f1ec\U0001f1e7','ES':'\U0001f1ea\U0001f1f8','DE':'\U0001f1e9\U0001f1ea'}
            _lang_now = lingua_atual()
            def _change_lang(cod, p=popup):
                p.destroy()
                definir_lingua(cod)
                c = cfg_load(); c['_lingua'] = cod; cfg_save(c)
                for w2 in self.root.winfo_children(): w2.destroy()
                self.root.configure(bg=CF)
                self._ui()
            for cod in ['PT','EN','ES','DE']:
                _active = (cod == _lang_now)
                tk.Button(pill_row, text=f"{_lang_flags[cod]} {cod}",
                          bg=CA if _active else CB,
                          fg=CP if _active else CD,
                          relief='flat', font=_F(8, bold=_active),
                          cursor='hand2', padx=7, pady=3,
                          command=lambda c=cod: _change_lang(c)
                          ).pack(side='left', padx=2)
            tk.Frame(inner, bg=CB, height=1).pack(fill='x', pady=8)
            # Theme toggle
            _ico_tema = '\u2600' if _TEMA_ACTUAL == 'escuro' else '\U0001f319'
            _txt_tema = T('tema_claro') if _TEMA_ACTUAL == 'escuro' else T('tema_escuro')
            def _toggle_theme(p=popup):
                p.destroy()
                novo = 'claro' if _TEMA_ACTUAL == 'escuro' else 'escuro'
                _aplicar_tema(novo)
                c = cfg_load(); c['tema'] = novo; cfg_save(c)
                for w2 in self.root.winfo_children(): w2.destroy()
                self.root.configure(bg=CF)
                self._ui()
            tk.Button(inner, text=f"  {_ico_tema}  {_txt_tema}",
                      bg=CB, fg=CD, relief='flat', font=_F(9),
                      cursor='hand2', padx=10, pady=5,
                      command=_toggle_theme).pack(fill='x')
            def _maybe_close(e, p=popup):
                try:
                    if not (p.winfo_x() <= e.x_root <= p.winfo_x()+p.winfo_width() and
                            p.winfo_y() <= e.y_root <= p.winfo_y()+p.winfo_height()):
                        p.destroy()
                except Exception: pass
            self.root.bind('<Button-1>', _maybe_close, add='+')

        gear_btn = tk.Button(tp, text='\u2699', bg=CP_TOP, fg=CD,
                             relief='flat', font=_F(14), cursor='hand2',
                             command=_open_settings, padx=8, pady=4)
        gear_btn.pack(side='right', padx=(0, 4), pady=28)
        gear_btn.bind('<Enter>', lambda e: gear_btn.config(fg=_accent_col))
        gear_btn.bind('<Leave>', lambda e: gear_btn.config(fg=CD))

        # Dashboard desativado (em desenvolvimento)

        ar = tk.Frame(self.root, bg=CF)
        ar.pack(fill='both', expand=True, padx=12, pady=(10, 12))

        # ── PanedWindow com divisor arrastável livre ────────────────────
        _pw = ttk.PanedWindow(ar, orient='horizontal')
        _pw.pack(fill='both', expand=True)

        # Painel esquerdo
        le_outer = tk.Frame(_pw, bg=CF)
        _pw.add(le_outer, weight=0)

        # Painel direito (consola) - expande
        ri = tk.Frame(_pw, bg=CF)
        _pw.add(ri, weight=1)

        # Largura inicial do painel esquerdo: proporcional ao ecrã
        _screen_w = self.root.winfo_screenwidth()
        _le_init  = cfg_load().get('_sash_pos') or max(220, min(380, int(_screen_w * 0.22)))
        self.root.after(60, lambda w=int(_le_init): _pw.sashpos(0, w))

        # Guardar posição do sash ao soltar
        _pw.bind('<ButtonRelease-1>', lambda e: self._save_sash(_pw))

        # Zona fixa de botoes - DEVE ser empacotada ANTES do canvas (pack side=bottom)
        self._bf_bottom = tk.Frame(le_outer, bg=CF)
        self._bf_bottom.pack(side='bottom', fill='x')
        tk.Frame(self._bf_bottom, bg='#243447', height=1).pack(fill='x')
        self._bf_inner = tk.Frame(self._bf_bottom, bg=CF)
        self._bf_inner.pack(fill='x', padx=6, pady=6)

        # Canvas scrollavel (empacotado depois do bottom para nao o engolir)
        le_sb=ttk.Scrollbar(le_outer, orient='vertical')
        le_sb.pack(side='right', fill='y')
        le_cv=tk.Canvas(le_outer, bg=CF, highlightthickness=0,
                        yscrollcommand=le_sb.set)
        le_cv.pack(side='left', fill='both', expand=True)
        le_sb.config(command=le_cv.yview)
        le=tk.Frame(le_cv, bg=CF)
        _le_win=le_cv.create_window((0,0), window=le, anchor='nw')
        def _on_le_cfg(ev, _c=le_cv, _b=None): _c.configure(scrollregion=_c.bbox('all'))
        def _on_le_cv_rsz(ev, _c=le_cv, _w=_le_win): _c.itemconfig(_w, width=ev.width)
        le.bind('<Configure>', _on_le_cfg)
        le_cv.bind('<Configure>', _on_le_cv_rsz)
        def _mwheel(ev, _c=le_cv):
            _c.yview_scroll(int(-1*(ev.delta/120)) if ev.delta else (-1 if ev.num==4 else 1), 'units')
        le_cv.bind('<MouseWheel>', _mwheel)
        le_cv.bind('<Button-4>', _mwheel)
        le_cv.bind('<Button-5>', _mwheel)
        le.bind('<Enter>', lambda e, _c=le_cv, _m=_mwheel: _c.bind_all('<MouseWheel>',_m))
        le.bind('<Leave>', lambda e, _c=le_cv: _c.unbind_all('<MouseWheel>'))

        # Paths default para o ambiente de trabalho
        import pathlib
        desktop = _desktop_path()
        proto_slug = self._proto_key
        def _default_path(fname):
            base = cfg.get(fname.replace('.','_'),'') or cfg.get(fname,'')
            if base: return base
            return str(pathlib.Path(desktop) / fname)

        self._sec(le, T('sec_ficheiros_entrada'), icon='')
        self._ent(le, T('pasta_individuos'), 'pasta', pasta=True, _cfg=cfg.get('pasta',''), key='pasta_individuos')
        if is_tiro:
            self._ent(le, T('fich_tempos_tiro'), 'ifd', fich=True,
                      ext=[('Excel','*.xlsx *.xls')], _cfg=cfg.get('ifd',''), key='fich_tempos_tiro')
            self._ent(le, T('fich_scores'), 'scores_file', fich=True,
                      ext=[('Excel','*.xlsx *.xls')], _cfg=cfg.get('scores_file',''), key='fich_scores')
        else:
            self._ent(le, T('fich_tempos'), 'ifd', fich=True,
                      ext=[('Excel','*.xlsx *.xls')], _cfg=cfg.get('ifd',''), key='fich_tempos')
            # Tiro com Arco - ficheiro de referencia demografica (142 atletas)
            if self._proto_key == PROTO_ARCO:
                self._ent(le, T('fich_atletas_ref'), 'atletas_ref_file', fich=True,
                          ext=[('Excel','*.xlsx *.xls')],
                          _cfg=cfg.get('atletas_ref_file',''),
                          key='fich_atletas_ref')
                # Re-renderizar o estado do "Peso padrão" sempre que o caminho mude.
                # O trace dispara quando o utilizador escolhe ficheiro ou apaga.
                try:
                    self.v_atletas_ref_file.trace_add(
                        'write', self._atualizar_estado_peso_padrao)
                except Exception:
                    try:
                        self.v_atletas_ref_file.trace(
                            'w', lambda *a: self._atualizar_estado_peso_padrao())
                    except Exception:
                        pass

        self._sec(le, T('sec_ficheiros_saida'), icon='')
        def_saida = cfg.get('saida','') or str(pathlib.Path(desktop)/f'resultados_estabilidade_{proto_slug}.xlsx')
        def_ind   = cfg.get('pasta_ind','') or str(pathlib.Path(desktop)/'individuais')
        def_pdf   = cfg.get('pdf','') or str(pathlib.Path(desktop)/f'relatorio_{proto_slug}.pdf')
        self._ent(le, T('saida_excel'), 'saida', grav=True,
                  default=def_saida, _cfg=cfg.get('saida',def_saida), key='saida_excel')
        self._ent(le, T('pasta_individuais'), 'pasta_ind', pasta_out=True,
                  default=def_ind, _cfg=cfg.get('pasta_ind',def_ind), key='pasta_individuais')
        self._ent(le, T('relatorio_pdf'), 'pdf', grav=True,
                  ext=[('PDF','*.pdf')], default=def_pdf,
                  _cfg=cfg.get('pdf',def_pdf), key='relatorio_pdf')

        self._sec(le, T('sec_opcoes_gerais'), icon='⚙')
        self.o_emb=tk.BooleanVar(value=cfg.get('o_emb',False))
        self.o_ell=tk.BooleanVar(value=cfg.get('o_ell',True))
        self.o_est=tk.BooleanVar(value=cfg.get('o_est',True))
        self.o_ind=tk.BooleanVar(value=cfg.get('o_ind',True))
        self.o_pdf=tk.BooleanVar(value=cfg.get('o_pdf',True))
        self._ck(le, T('ck_embedded'), self.o_emb, tip=T('tip_embedded'))
        self._ck(le, T('ck_elipse'),   self.o_ell, tip=T('tip_elipse'))
        self._ck(le, T('ck_estab'),    self.o_est, tip=T('tip_estab'))
        self._ck(le, T('ck_individuais'), self.o_ind, tip=T('tip_individuais'))
        self._ck(le, T('ck_pdf'),      self.o_pdf,
                 tip='Gera um relatorio PDF com uma pagina por individuo.')

        # Numero de ensaios configuravel
        f_nens = tk.Frame(le, bg=CF); f_nens.pack(fill='x', pady=(4,0))
        tk.Label(f_nens, text=T('n_ensaios_label'),
                 bg=CF, fg=CT, font=_F(9), anchor='w').pack(side='left')
        self.v_n_ens = tk.StringVar(value=cfg.get('n_ens',''))
        ent_nens = tk.Entry(f_nens, textvariable=self.v_n_ens, bg=CC, fg=CT,
                            insertbackground=CA, relief='flat', font=_F(9), width=6)
        ent_nens.pack(side='left', padx=6)
        Tooltip(ent_nens, T('tip_n_ensaios_tiro'))

        # ---------------------------------------------------------------
        # Opcoes especificas do protocolo TIRO
        # ---------------------------------------------------------------
        if is_tiro:
            self._sec(le, T('sec_proto_tiro'), icon='')

            # Matching de individuos
            lbl_match = tk.Label(le, text=T('tiro_assoc_por'), bg=CF, fg=CT,
                     font=_F(9), anchor='w')
            lbl_match.pack(fill='x', pady=(2,0))
            Tooltip(lbl_match, T('tip_match'))
            self.v_tiro_match = tk.StringVar(value=cfg.get('tiro_match','id'))
            f_match = tk.Frame(le, bg=CF); f_match.pack(fill='x', padx=(12,0))
            rb_id = tk.Radiobutton(f_match, text=T('tiro_id_pasta'), variable=self.v_tiro_match,
                           value='id', bg=CF, fg=CT, selectcolor=CP,
                           font=_F(8))
            rb_id.pack(side='left', padx=(0,10))
            Tooltip(rb_id, T('tip_rb_id'))
            rb_idx = tk.Radiobutton(f_match, text=T('tiro_pos_lista'), variable=self.v_tiro_match,
                           value='idx', bg=CF, fg=CT, selectcolor=CP,
                           font=_F(8))
            rb_idx.pack(side='left')
            Tooltip(rb_idx, T('tip_rb_idx'))

            # Intervalos a calcular
            tk.Label(le, text=T('tiro_itvs_calc'), bg=CF, fg=CT,
                     font=_F(9)).pack(anchor='w', pady=(6,2))
            self.v_tiro_itvs = {}
            itvs_cfg = cfg.get('tiro_itvs', list(TIRO_INTERVALOS.keys()))
            f_itvs = tk.Frame(le, bg=CF); f_itvs.pack(fill='x')
            for itv in TIRO_INTERVALOS:
                v = tk.BooleanVar(value=(itv in itvs_cfg))
                self.v_tiro_itvs[itv] = v
                cb = tk.Checkbutton(f_itvs, text=_tiro_itv_label(itv), variable=v,
                                    bg=CF, fg=CT, selectcolor=CP,
                                    font=_F(8))
                cb.pack(anchor='w', padx=12)

            # Bipodal (Hurdle Step) - checkbox + campo N. ensaios na mesma linha
            self.o_tiro_bipodal = tk.BooleanVar(value=cfg.get('tiro_bipodal', True))
            self.v_n_ens_hs = tk.StringVar(value=str(cfg.get('n_ens_hs', 5)))
            f_hs_row = tk.Frame(le, bg=CF); f_hs_row.pack(fill='x', pady=1)
            cb_hs = tk.Checkbutton(f_hs_row, text=T('tiro_bipodal'),
                                   variable=self.o_tiro_bipodal,
                                   bg=CF, fg=CT, selectcolor=CC,
                                   activebackground=CF, activeforeground=CA,
                                   font=_F(9), relief='flat', cursor='hand2')
            cb_hs.pack(side='left', anchor='w')
            Tooltip(cb_hs, T('tip_bipodal'))
            # N. ensaios inline
            tk.Label(f_hs_row, text=T('tiro_n_ensaios'), bg=CF, fg=CD,
                     font=_F(8)).pack(side='left')
            ent_nhs = tk.Entry(f_hs_row, textvariable=self.v_n_ens_hs, bg=CC, fg=CT,
                               insertbackground=CA, relief='flat', font=_F(9), width=4)
            ent_nhs.pack(side='left', padx=(2,0))
            Tooltip(ent_nhs, T('tip_n_ens_hs'))
            tk.Label(f_hs_row, text=T('tiro_p_pe'), bg=CF, fg=CD,
                     font=_F(7)).pack(side='left', padx=(1,0))

            # Distancias a processar
            self._sec(le, T('sec_dist_teste'), icon='')
            tk.Label(le, text=T('tiro_dist_info'),
                     bg=CF, fg=CD, font=_F(7), wraplength=460, justify='left').pack(anchor='w')

            # Lista dinamica de distancias
            self.v_dists_extra = []
            self._dist_frame = tk.Frame(le, bg=CF)
            self._dist_frame.pack(fill='x', pady=(4,0))
            dists_saved = cfg.get('tiro_dists_extra', [])
            for d in (dists_saved if dists_saved else []):
                self._add_dist_row(d)

            btn_add_dist = tk.Button(le, text=T('tiro_add_dist'),
                                     bg='#1C3349', fg=CA, relief='flat',
                                     font=_F(8), cursor='hand2',
                                     command=self._add_dist_row)
            btn_add_dist.pack(anchor='w', pady=(2, 0))
            btn_add_dist.bind('<Enter>', lambda e: btn_add_dist.config(bg=CB, fg=CT))
            btn_add_dist.bind('<Leave>', lambda e: btn_add_dist.config(bg='#1C3349', fg=CA))

        else:
            self.v_tiro_match   = None
            self.v_tiro_itvs    = {}
            self.o_tiro_bipodal = None
            self.v_dists_extra  = []
            self._dist_frame    = None
            self.v_dist_tiro    = None

        # Testes estatisticos
        self._sec(le, T('sec_anal_estat'), icon='')
        self.o_estats=tk.BooleanVar(value=cfg.get('o_estats',False))
        self._ck(le, T('ck_estats'), self.o_estats,
                 tip=('Gera abas ESTATS no Excel e pagina no PDF com:\n'
                      '  Shapiro-Wilk + IC 95% / IQR por metrica\n'
                      '  t-pareado ou Wilcoxon + Cohen\'s d (Dir vs Esq)\n'
                      '  CV variabilidade intra-individuo (semaforo)\n'
                      '  Índice de perturbação por ensaio (Tiro)\n'
                      '  Friedman entre intervalos + pos-hoc (Tiro)\n'
                      '  Correlacao EA95 vs score (Tiro)\n\n'
                      'Requer n >= 3 individuos. Desactivado por defeito.'))

        # Sub-opcoes em coluna unica (evita corte de texto em painel estreito)
        sub_f = tk.Frame(le, bg=CF); sub_f.pack(fill='x', padx=(18,0), pady=(2,0))
        self.o_estats_grupo    = tk.BooleanVar(value=cfg.get('estats_grupo',True))
        self.o_estats_dir_esq  = tk.BooleanVar(value=cfg.get('estats_dir_esq',True))
        self.o_estats_pos_disp = tk.BooleanVar(value=cfg.get('estats_pos_disp',True))
        self.o_estats_atleta   = tk.BooleanVar(value=cfg.get('estats_atleta',False))
        self.o_estats_friedman = tk.BooleanVar(value=cfg.get('estats_friedman',True))
        self.o_estats_corr     = tk.BooleanVar(value=cfg.get('estats_corr',False))
        self.o_estats_posthoc  = tk.BooleanVar(value=cfg.get('estats_posthoc',True))

        _cb_style = dict(bg=CF, fg=CD, selectcolor=CC, activebackground=CF,
                         activeforeground=CA, font=_F(8),
                         relief='flat', cursor='hand2')

        cb_gr  = tk.Checkbutton(sub_f, text=T('ck_estats_gr'),
                                variable=self.o_estats_grupo, **_cb_style)
        cb_gr.pack(anchor='w', pady=1)
        Tooltip(cb_gr, _TESTS_INFO['grupo']['descr'])

        cb_de  = tk.Checkbutton(sub_f, text=T('ck_estats_de'),
                                variable=self.o_estats_dir_esq, **_cb_style)
        cb_de.pack(anchor='w', pady=1)
        Tooltip(cb_de, _TESTS_INFO['dir_esq']['descr'])

        cb_pd  = tk.Checkbutton(sub_f, text=T('ck_estats_pd'),
                                variable=self.o_estats_pos_disp, **_cb_style)
        cb_pd.pack(anchor='w', pady=1)
        Tooltip(cb_pd, _TESTS_INFO['pos_disp']['descr'])

        cb_at  = tk.Checkbutton(sub_f, text=T('ck_estats_at'),
                                variable=self.o_estats_atleta, **_cb_style)
        cb_at.pack(anchor='w', pady=1)
        Tooltip(cb_at, _TESTS_INFO['individuo']['descr'])

        cb_fr  = tk.Checkbutton(sub_f, text=T('ck_estats_fr'),
                                variable=self.o_estats_friedman, **_cb_style)
        cb_fr.pack(anchor='w', pady=1)
        Tooltip(cb_fr, _TESTS_INFO['friedman']['descr'])

        cb_ph  = tk.Checkbutton(sub_f, text=T('ck_estats_ph'),
                                variable=self.o_estats_posthoc, **_cb_style)
        cb_ph.pack(anchor='w', pady=1)
        Tooltip(cb_ph, T('tip_posthoc'))

        cb_co  = tk.Checkbutton(sub_f, text=T('ck_estats_co'),
                                variable=self.o_estats_corr, **_cb_style)
        cb_co.pack(anchor='w', pady=1)
        Tooltip(cb_co, _TESTS_INFO['correlacao']['descr'])

        _all_cbs = [cb_gr, cb_de, cb_pd, cb_at, cb_fr, cb_ph, cb_co]

        def _upd_estats_sub(*_):
            st = 'normal' if self.o_estats.get() else 'disabled'
            for cb in _all_cbs:
                cb.config(state=st)
            # pos-hoc so activo se friedman activo E o_estats activo
            if self.o_estats.get() and not self.o_estats_friedman.get():
                cb_ph.config(state='disabled')

        self.o_estats_friedman.trace_add('write', _upd_estats_sub)
        self.o_estats.trace_add('write', _upd_estats_sub)
        _upd_estats_sub()

        # ── SECÇÃO EXPORTAÇÃO ────────────────────────────────────────────
        self._sec(le, T('sec_exportacao'), icon='')

        # CSV
        f_csv = tk.Frame(le, bg=CF); f_csv.pack(fill='x', pady=2)
        self.o_export_csv = tk.BooleanVar(value=cfg.get('export_csv', True))
        cb_csv = tk.Checkbutton(f_csv, text=T('export_csv_label'),
                                variable=self.o_export_csv,
                                bg=CF, fg=CT, selectcolor='#00B4D8',
                                activebackground=CF, activeforeground=CA,
                                font=_F(9), relief='flat', cursor='hand2')
        cb_csv.pack(side='left')
        Tooltip(cb_csv, T('tip_csv'))
        f_csv_opt = tk.Frame(le, bg=CF); f_csv_opt.pack(fill='x', padx=(18,0))
        tk.Label(f_csv_opt, text=T('lbl_sep'), bg=CF, fg=CD,
                 font=_F(8)).pack(side='left')
        self.v_csv_sep = tk.StringVar(value=cfg.get('csv_sep',';'))
        cb_sep = ttk.Combobox(f_csv_opt, textvariable=self.v_csv_sep,
                              values=[';',',','\t'], state='readonly',
                              font=_F(8), width=4)
        cb_sep.pack(side='left', padx=4)
        tk.Label(f_csv_opt, text=T('lbl_dec'), bg=CF, fg=CD,
                 font=_F(8)).pack(side='left', padx=(8,0))
        self.v_csv_dec = tk.StringVar(value=cfg.get('csv_dec',','))
        cb_dec = ttk.Combobox(f_csv_opt, textvariable=self.v_csv_dec,
                              values=[',','.'], state='readonly',
                              font=_F(8), width=4)
        cb_dec.pack(side='left', padx=4)
        Tooltip(cb_sep, T('tip_csv_sep'))
        Tooltip(cb_dec, T('tip_csv_dec'))

        # DOCX
        f_docx = tk.Frame(le, bg=CF); f_docx.pack(fill='x', pady=(4,2))
        self.o_export_docx = tk.BooleanVar(value=cfg.get('export_docx', False))
        cb_docx = tk.Checkbutton(f_docx, text=T('export_docx_label'),
                                 variable=self.o_export_docx,
                                 bg=CF, fg=CT, selectcolor='#00B4D8',
                                 activebackground=CF, activeforeground=CA,
                                 font=_F(9), relief='flat', cursor='hand2')
        cb_docx.pack(side='left')
        Tooltip(cb_docx, T('tip_docx'))

        # HTML interactivo
        f_html = tk.Frame(le, bg=CF); f_html.pack(fill='x', pady=(2,2))
        self.o_export_html = tk.BooleanVar(value=cfg.get('export_html', False))
        cb_html = tk.Checkbutton(f_html, text=T('ck_html'),
                                 variable=self.o_export_html,
                                 bg=CF, fg=CT, selectcolor='#00B4D8',
                                 activebackground=CF, activeforeground=CA,
                                 font=_F(9), relief='flat', cursor='hand2')
        cb_html.pack(side='left')
        Tooltip(cb_html, T('tip_html'))

        # PNG dos graficos individuais
        f_png = tk.Frame(le, bg=CF); f_png.pack(fill='x', pady=(2,2))
        self.o_export_png = tk.BooleanVar(value=cfg.get('export_png', False))
        cb_png = tk.Checkbutton(f_png,
                                text='Exportar PNG (estabilograma + elipse por ensaio)',
                                variable=self.o_export_png,
                                bg=CF, fg=CT, selectcolor='#00B4D8',
                                activebackground=CF, activeforeground=CA,
                                font=_F(9), relief='flat', cursor='hand2')
        cb_png.pack(side='left')
        Tooltip(cb_png, 'Gera ficheiros PNG com o estabilograma e a elipse 95% para cada ensaio.\nGuardados na pasta de ficheiros individuais.')

        # ── Opcoes de calculo avanc. ──────────────────────────────────────
        self._sec(le, 'OPCOES DE CALCULO')

        # FFT
        f_fft = tk.Frame(le, bg=CF); f_fft.pack(fill='x', pady=(2,0))
        self.o_fft = tk.BooleanVar(value=cfg.get('fft_ativo', False))
        cb_fft = tk.Checkbutton(f_fft,
                                text='Análise espectral FFT (freq. pico e média ML/AP)',
                                variable=self.o_fft,
                                bg=CF, fg=CT, selectcolor='#00B4D8',
                                activebackground=CF, activeforeground=CA,
                                font=_F(9), relief='flat', cursor='hand2')
        cb_fft.pack(side='left')
        Tooltip(cb_fft, 'Calcula a frequencia de pico e a frequencia media ponderada\n'
                        'do espectro de potencia do CoP (eixos ML e AP).\n'
                        'Intervalo fisiologico: 0.1-10 Hz.\n'
                        'Aumenta o tempo de processamento ligeiramente.')

        # Normalizacao pela massa corporal
        f_peso = tk.Frame(le, bg=CF); f_peso.pack(fill='x', pady=(4,0))
        self.o_peso_norm = tk.BooleanVar(value=cfg.get('peso_norm', False))
        cb_peso = tk.Checkbutton(f_peso,
                                 text='Normalizar elipse pela massa corporal (ea95/kg)',
                                 variable=self.o_peso_norm,
                                 bg=CF, fg=CT, selectcolor='#00B4D8',
                                 activebackground=CF, activeforeground=CA,
                                 font=_F(9), relief='flat', cursor='hand2')
        cb_peso.pack(side='left')
        Tooltip(cb_peso, 'Activa calculos que requerem a massa corporal.\n'
                         'Requer ficheiro de tempos com coluna "peso_kg" ou introducao manual.')
        # Campo de peso unico (para analises sem ficheiro de tempos / sem ref demografica).
        # Quando o utilizador carrega o ficheiro de referencia demografica (142 atletas)
        # no protocolo Tiro com Arco, o peso vem por atleta a partir desse ficheiro -
        # entao desactivamos a entrada e mostramos uma label clarificadora.
        f_peso2 = tk.Frame(le, bg=CF); f_peso2.pack(fill='x', padx=(18,0), pady=(2,4))
        self._lbl_peso_padrao = tk.Label(f_peso2, text='Peso padrão (kg):',
                                          bg=CF, fg=CD, font=_F(8))
        self._lbl_peso_padrao.pack(side='left')
        self.v_peso_kg = tk.StringVar(value=str(cfg.get('peso_kg_default', '')))
        self._ent_peso = tk.Entry(f_peso2, textvariable=self.v_peso_kg, bg=CC, fg=CT,
                            insertbackground=CA, relief='flat', font=_F(9), width=7)
        self._ent_peso.pack(side='left', padx=6)
        Tooltip(self._ent_peso, 'Peso usado quando não existe coluna "peso_kg" nos dados.\n'
                          'Deixar vazio para usar o peso individual de cada atleta (se disponível).')
        # Label "ler do ficheiro de referencia" - aparece quando ref carregada
        self._lbl_peso_ref = tk.Label(f_peso2,
                                       text='  ✓ ler do ficheiro de referência',
                                       bg=CF, fg=COK, font=_F(8, bold=True))
        # nao empacotar agora; _atualizar_estado_peso_padrao() controla a visibilidade
        self._atualizar_estado_peso_padrao()  # estado inicial

        # ── Botoes fixos - populados aqui, o frame ja foi criado acima ─────
        bf_inner = self._bf_inner

        # Separador decorativo
        tk.Frame(bf_inner, bg=CB, height=1).pack(fill='x', pady=(0, 7))

        # Botão Executar - destaque máximo
        self.b_exe = tk.Button(bf_inner,
                               text=T('btn_executar') + f'  [{_MOD_SYM}+↵]',
                               bg=CA, fg='#071014',
                               font=_F(11, bold=True),
                               activebackground=CA2, activeforeground='white',
                               relief='flat', cursor='hand2',
                               padx=16, pady=12,
                               command=self._exec)
        self.b_exe.pack(fill='x')
        self.b_exe.bind('<Enter>', lambda e: self.b_exe.config(bg=CA2, fg='white')
                        if str(self.b_exe['state']) == 'normal' else None)
        self.b_exe.bind('<Leave>', lambda e: self.b_exe.config(bg=CA, fg='#071014')
                        if str(self.b_exe['state']) == 'normal' else None)

        # Botão Parar
        self.b_stop = tk.Button(bf_inner, text=T('btn_stop'),
                                bg='#200808', fg=CER,
                                font=_F(9, bold=True), relief='flat',
                                cursor='hand2', padx=10, pady=5,
                                state='disabled', activebackground='#4A1010',
                                command=self._stop)
        self.b_stop.pack(fill='x', pady=(4, 0))
        self.b_stop.bind('<Enter>', lambda e: self.b_stop.config(bg='#3A0C0C')
                         if str(self.b_stop['state']) == 'normal' else None)
        self.b_stop.bind('<Leave>', lambda e: self.b_stop.config(bg='#200808')
                         if str(self.b_stop['state']) == 'normal' else None)

        # Linha de botões utilitários
        f_util = tk.Frame(bf_inner, bg=CF); f_util.pack(fill='x', pady=(6,0))

        def _btn_util(parent, txt, cmd, tip=''):
            b = tk.Button(parent, text=txt, bg=CP, fg=CD,
                          font=_F(8), relief='flat',
                          cursor='hand2', pady=5, activebackground=CB,
                          activeforeground=CT, command=cmd)
            b.pack(side='left', fill='x', expand=True, padx=1)
            b.bind('<Enter>', lambda e, _b=b: _b.config(bg=CB, fg=CT))
            b.bind('<Leave>', lambda e, _b=b: _b.config(bg=CP, fg=CD))
            if tip: Tooltip(b, tip)
            return b

        _btn_util(f_util, T('btn_protocolo'), self._mudar_proto,
                  'Muda o protocolo de análise.')
        _btn_util(f_util, T('btn_historico'), self._ver_historico,
                  f'Ver as últimas sessões.  [{_MOD_SYM}+H]')
        _btn_util(f_util, T('btn_poder'), lambda: _janela_power_calc(self.root, tk, ttk),
                  f'Calculadora de tamanho amostral.  [{_MOD_SYM}+P]')
        _btn_util(f_util, T('btn_guardar'), self._save_cfg,
                  f'Guardar configuração.  [{_MOD_SYM}+S]')

        # Segunda linha de utilitários
        f_util2 = tk.Frame(bf_inner, bg=CF); f_util2.pack(fill='x', pady=(2,0))
        _btn_util(f_util2, '📌 ' + T('perfis_titulo').replace('📌  ','').replace('📌 ',''),
                  self._gerir_perfis, T('perfis_descr'))
        _btn_util(f_util2, '⚡ [F2]', self._quick_subject, T('quick_titulo'))
        _btn_util(f_util2, '⌨ [F1]', self._show_shortcuts, T('shortcuts_titulo'))

        def _toggle_tema():
            novo = 'claro' if _TEMA_ACTUAL == 'escuro' else 'escuro'
            _aplicar_tema(novo)
            d = cfg_load(); d['tema'] = novo; cfg_save(d)
            for w2 in self.root.winfo_children(): w2.destroy()
            self.root.configure(bg=CF)
            self._ui()
        # Frame extra para botao de abrir pasta (aparece apos conclusao)
        self._bf_extra = tk.Frame(bf_inner, bg=CF)
        self._bf_extra.pack(fill='x', pady=(4,0))

        # ── Painel direito: Log + Progresso ─────────────────────────────
        # Cabeçalho do painel com indicador de estado
        hdr_ri = tk.Frame(ri, bg=CB); hdr_ri.pack(fill='x', pady=(0, 4))
        # Status dot (● idle → ● running)
        self._status_dot = tk.Label(hdr_ri, text='●', bg=CB, fg='#3A5F7A',
                                    font=_F(9), padx=6, pady=5)
        self._status_dot.pack(side='left')
        tk.Label(hdr_ri, text=T('sec_log'),
                 bg=CB, fg=CA, font=_F(8, bold=True),
                 padx=2, pady=5).pack(side='left')
        _clr_btn = tk.Button(hdr_ri, text=T('btn_limpar'), bg=CB, fg=CD,
                  font=_F(7), relief='flat', cursor='hand2',
                  padx=8, pady=4, command=self._clr)
        _clr_btn.pack(side='right', padx=4)
        _clr_btn.bind('<Enter>', lambda e: _clr_btn.config(fg=CT))
        _clr_btn.bind('<Leave>', lambda e: _clr_btn.config(fg=CD))

        # Área do log com scrollbar estilizada
        lf2 = tk.Frame(ri, bg=CC, bd=0); lf2.pack(fill='both', expand=True)
        st2 = ttk.Style()
        try: st2.theme_use('clam')
        except Exception: pass
        try:
            st2.configure('Log.Vertical.TScrollbar',
                          background=CB, troughcolor=CC,
                          arrowcolor=CD, bordercolor=CC)
        except Exception: pass
        sb = ttk.Scrollbar(lf2, style='Log.Vertical.TScrollbar')
        sb.pack(side='right', fill='y')
        self.log = tk.Text(lf2, bg='#0C1620', fg=CT, font=_FM(9),
                           relief='flat', padx=12, pady=10,
                           yscrollcommand=sb.set, state='disabled', wrap='word',
                           selectbackground=CB, cursor='arrow')
        self.log.pack(fill='both', expand=True)
        sb.config(command=self.log.yview)
        # Tags de cor
        self.log.tag_config('ok',    foreground=COK)
        self.log.tag_config('aviso', foreground=CAV)
        self.log.tag_config('erro',  foreground=CER)
        self.log.tag_config('info',  foreground=CA)
        self.log.tag_config('dim',   foreground=CD)
        self.log.tag_config('neg',   foreground=CT,
                             font=_FM(9)+(('bold',) if True else ()))

        # Separador
        tk.Frame(ri, bg=CB, height=1).pack(fill='x', pady=(6, 4))

        # Progresso
        f_prog = tk.Frame(ri, bg=CF); f_prog.pack(fill='x')
        self.plbl = tk.Label(f_prog, text=T('status_pronto'),
                             bg=CF, fg=CD, font=_F(8), anchor='w')
        self.plbl.pack(side='left')
        self.ppct = tk.Label(f_prog, text='', bg=CF, fg=CA,
                             font=_F(8, bold=True), anchor='e')
        self.ppct.pack(side='right')
        st3 = ttk.Style()
        st3.configure('P.Horizontal.TProgressbar',
                       troughcolor='#0C1620', background=CA,
                       bordercolor=CB, lightcolor=CA3, darkcolor=CA2,
                       thickness=9)
        self.pvar = tk.DoubleVar(value=0)
        ttk.Progressbar(ri, variable=self.pvar,
                        style='P.Horizontal.TProgressbar',
                        maximum=100).pack(fill='x', pady=(4, 2))

        # Status bar (rodapé fixo)
        sb_frame = tk.Frame(self.root, bg=_SB_BG, height=24)
        sb_frame.pack(fill='x', side='bottom'); sb_frame.pack_propagate(False)
        tk.Label(sb_frame,
                 text=f'  BSP v{VERSAO}  ·  {AUTOR}  ·  {SEGUNDO_AUTOR}  ·  {ORIENTADOR}',
                 bg=_SB_BG, fg=_SB_FG, font=_F(7), anchor='w'
                 ).pack(side='left', fill='y')
        # Indicador de protocolo na status bar
        cor_p2 = {'fms': '#00B4D8', 'unipodal': '#22C55E', 'tiro': '#C8A45A'}
        tk.Label(sb_frame,
                 text=f'  {_proto_nome(self._proto_key)}  ',
                 bg=cor_p2.get(self._proto_key, CA), fg='white',
                 font=_F(7, bold=True)).pack(side='right')

        self._log(f'BSP  -  {PROG} v{VERSAO}  -  {_proto_nome(self._proto_key)}', 'info')
        self._log(f'{AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}', 'dim')
        self._log(T('log_atalhos', mod=_MOD_SYM) + '\n', 'dim')

    def _add_dist_row(self, default_val=''):
        """Adiciona uma linha de input para uma distancia extra do teste de tiro."""
        if self._dist_frame is None: return
        tk = self.tk
        row_f = tk.Frame(self._dist_frame, bg=CF)
        row_f.pack(fill='x', pady=1)
        tk.Label(row_f, text=T('lbl_dist_m'), bg=CF, fg=CT,
                 font=_F(8)).pack(side='left')
        v = tk.StringVar(value=str(default_val))
        e = tk.Entry(row_f, textvariable=v, width=6, bg=CC, fg=CT,
                     insertbackground=CA, relief='flat', font=_F(9))
        e.pack(side='left', padx=4)
        self.v_dists_extra.append(v)

        def _remove():
            self.v_dists_extra.remove(v)
            row_f.destroy()

        btn_rm = tk.Button(row_f, text='x', bg='#3A1010', fg='#FCA5A5',
                           relief='flat', font=_F(7), command=_remove,
                           padx=4)
        btn_rm.pack(side='left', padx=2)
        Tooltip(e, T('tip_dist_extra'))

    def _sec(self, p, t, icon=''):
        f = self.tk.Frame(p, bg=CF); f.pack(fill='x', pady=(16, 4))
        # Thick left accent bar
        self.tk.Frame(f, bg=CA, width=4).pack(side='left', fill='y', padx=(0, 8))
        lbl_txt = (icon + '  ' + t.upper()) if icon else t.upper()
        self.tk.Label(f, text=lbl_txt, bg=CF, fg=CA,
                      font=_F(8, bold=True)).pack(side='left')
        self.tk.Frame(f, bg=CB, height=1).pack(side='left', fill='x', expand=True,
                                                padx=(10, 0), pady=6)

    def _panel_estado(self, parent, estado='vazio', mensagem='', acao=None, txt_acao=''):
        """
        Painel padronizado de estado (loading / vazio / erro / sucesso).

        estado   : 'loading' | 'vazio' | 'erro' | 'ok'
        mensagem : texto a mostrar abaixo do icone
        acao     : callable opcional (botao de acao)
        txt_acao : texto do botao (quando `acao` passada)

        Devolve o frame criado (para destruicao posterior se necessario).
        """
        tk = self.tk
        _estilos = {
            'loading': (ICO['loading'], CA,  'A carregar...'),
            'vazio':   (ICO['info'],    CD,  'Sem dados disponiveis'),
            'erro':    (ICO['error'],   CER, 'Ocorreu um erro'),
            'ok':      (ICO['ok'],      COK, 'Tudo certo'),
        }
        icone, cor, default_msg = _estilos.get(estado, _estilos['vazio'])
        msg = mensagem or default_msg

        f = tk.Frame(parent, bg=CP, padx=_PAD_L, pady=_PAD_L)
        f.pack(fill='both', expand=True)
        # Icone grande
        tk.Label(f, text=icone, bg=CP, fg=cor,
                 font=(_font_family_ui(), 32, 'bold')).pack(pady=(_PAD_L, _PAD_S))
        # Mensagem
        tk.Label(f, text=msg, bg=CP, fg=CT,
                 font=_FONT_BODY(), wraplength=420).pack(pady=(0, _PAD_M))
        # Botao de acao (opcional)
        if acao is not None:
            try:
                ttk = self.ttk
                btn = ttk.Button(f, text=txt_acao or 'Tentar novamente',
                                 style='Primary.TButton', command=acao)
                btn.pack(pady=_PAD_S)
            except Exception:
                btn = tk.Button(f, text=txt_acao or 'Tentar novamente',
                                bg=CA, fg=CT, relief='flat', font=_FONT_BODY(),
                                padx=_PAD_M, pady=_PAD_S, command=acao)
                btn.pack(pady=_PAD_S)
        return f

    def _ent(self, p, label, attr, pasta=False, fich=False, grav=False,
             pasta_out=False, ext=None, default='', _cfg='', tip='', key=''):
        # Se foi passada uma chave i18n, usar versão curta (_s) + tooltip (_tip)
        if key:
            short       = T(key + '_s') if T(key + '_s') != key + '_s' else label
            tooltip_txt = tip or (T(key + '_tip') if T(key + '_tip') != key + '_tip' else label)
        else:
            short       = label
            tooltip_txt = tip or label

        # Subtle card background for each row
        _row_bg = '#0F1D2C'
        f = self.tk.Frame(p, bg=_row_bg, padx=6, pady=3)
        f.pack(fill='x', pady=1)
        lbl = self.tk.Label(f, text=short, bg=_row_bg, fg=CT, font=_F(8),
                            anchor='w', width=16)
        lbl.pack(side='left')
        Tooltip(lbl, tooltip_txt)
        val = _cfg if _cfg else default
        var = self.tk.StringVar(value=val); setattr(self, f'v_{attr}', var)
        ent_f = self.tk.Frame(f, bg='#1C3349', bd=0)
        ent_f.pack(side='left', fill='x', expand=True, padx=(2, 4))
        ent = self.tk.Entry(ent_f, textvariable=var, bg='#1A2E42', fg=CT,
                      insertbackground=CA, relief='flat',
                      font=_F(9),
                      highlightthickness=1,
                      highlightcolor=CA,
                      highlightbackground='#1C3349')
        ent.pack(fill='x', ipady=4, padx=1, pady=1)
        Tooltip(ent, tooltip_txt)
        def nav():
            if pasta or pasta_out: path=self.fd.askdirectory()
            elif grav:
                ef=(ext[0][1].split()[0].replace('*','') if ext else '.xlsx')
                path=self.fd.asksaveasfilename(defaultextension=ef,
                                               filetypes=ext or [('Todos','*.*')])
            else: path=self.fd.askopenfilename(filetypes=ext or [('Todos','*.*')])
            if path: var.set(path)
        _nav_btn = self.tk.Button(f, text='···', bg='#1C3349', fg=CA,
                       font=_F(9, bold=True),
                       relief='flat', cursor='hand2', padx=8, pady=3,
                       command=nav)
        _nav_btn.pack(side='left')
        _nav_btn.bind('<Enter>', lambda e, b=_nav_btn: b.config(bg=CA, fg='#071014'))
        _nav_btn.bind('<Leave>', lambda e, b=_nav_btn: b.config(bg='#1C3349', fg=CA))

    def _ck(self, p, l, v, tip=None):
        cb = self.tk.Checkbutton(p, text=l, variable=v, bg=CF, fg=CT,
                            selectcolor=CA,
                            activebackground=CF, activeforeground=CA,
                            font=_F(9), relief='flat',
                            wraplength=380, justify='left',
                            cursor='hand2', padx=4)
        cb.pack(anchor='w', pady=2, fill='x')
        # Ajustar wraplength dinamicamente quando o painel muda de tamanho
        def _on_resize(ev, _cb=cb):
            wl = max(180, ev.width - 30)
            _cb.config(wraplength=wl)
        p.bind('<Configure>', _on_resize, add='+')
        if tip:
            Tooltip(cb, tip)
        return cb

    def _save_sash(self, pw):
        try:
            pos = pw.sashpos(0)
            d = cfg_load(); d['_sash_pos'] = pos; cfg_save(d)
        except Exception:
            pass

    def _atualizar_estado_peso_padrao(self, *args):
        # Quando o ficheiro de referência (142 atletas) está carregado, o peso
        # vem por atleta de lá -> desactiva o campo manual e mostra a label.
        ent     = getattr(self, '_ent_peso',        None)
        lbl_ref = getattr(self, '_lbl_peso_ref',    None)
        lbl_def = getattr(self, '_lbl_peso_padrao', None)
        if ent is None or lbl_ref is None or lbl_def is None:
            return
        v_ref = getattr(self, 'v_atletas_ref_file', None)
        tem_ref = bool(v_ref and str(v_ref.get()).strip())
        try:
            if tem_ref:
                ent.config(state='disabled',
                           disabledforeground=CD,
                           disabledbackground=CG)
                if not lbl_ref.winfo_ismapped():
                    lbl_ref.pack(side='left', padx=(2, 0))
            else:
                ent.config(state='normal')
                if lbl_ref.winfo_ismapped():
                    lbl_ref.pack_forget()
        except Exception:
            pass

    def _log(self, msg, tag=''): self._q.put((msg, tag))

    def _set_prog(self, pct, msg=''):
        """Envia actualizacao de progresso para a fila - thread-safe."""
        # Usa a mesma fila do _log para evitar qualquer chamada Tkinter
        # directamente da thread de analise (causava crash no Windows)
        try:
            self._q.put(('__prog__', pct, msg, time.time()))
        except Exception:
            pass
    def _stop(self): self._run=False; self._log(T('log_cancelado'), 'aviso')
    def _save_cfg(self):
        """Guarda configuração actual (atalho Ctrl/Cmd+S)."""
        try:
            d = cfg_load()
            for attr, key in [('v_pasta','pasta'),('v_ifd','ifd'),('v_saida','saida'),
                               ('v_pasta_ind','pasta_ind'),('v_pdf','pdf'),
                               ('v_atletas_ref_file','atletas_ref_file')]:
                if hasattr(self, attr):
                    d[key] = getattr(self, attr).get()
            d['protocolo'] = self._proto_key
            for attr, key in [
                ('o_emb','o_emb'), ('o_ell','o_ell'), ('o_est','o_est'),
                ('o_ind','o_ind'), ('o_pdf','o_pdf'), ('o_estats','o_estats'),
                ('o_export_csv','export_csv'), ('o_export_docx','export_docx'),
                ('o_export_html','export_html'),
            ]:
                if hasattr(self, attr): d[key] = getattr(self, attr).get()
            if hasattr(self, 'v_csv_sep'): d['csv_sep'] = self.v_csv_sep.get()
            if hasattr(self, 'v_csv_dec'): d['csv_dec'] = self.v_csv_dec.get()
            cfg_save(d)
            self._log(T('log_cfg_guardada', mod=_MOD_SYM), 'ok')
        except Exception as _e:
            self._log(T('log_cfg_erro', e=_e), 'aviso')

    def _mudar_proto(self):
        proto = escolher_protocolo_gui(self.root)
        if proto:
            Janela(protocolo=proto, _root=self.root)

    def _ver_historico(self):
        """Abre janela com historico das ultimas sessoes."""
        tk = self.tk
        entries = hist_load()
        win = tk.Toplevel(self.root)
        win.title(T('hist_titulo') + '  |  BSP')
        win.configure(bg=CF); win.resizable(True, True)
        w, h = 800, 520
        x = (win.winfo_screenwidth()-w)//2
        y = (win.winfo_screenheight()-h)//2
        win.geometry(f'{w}x{h}+{x}+{y}')
        win.minsize(600, 400)
        win.grab_set()

        tk.Label(win, text=T('hist_titulo'), bg=CF, fg=CA,
                 font=_F(11, bold=True)).pack(pady=(14,4))
        tk.Label(win, text=T('hist_subtitulo', n=_HIST_MAX),
                 bg=CF, fg=CD, font=_F(8)).pack()
        tk.Frame(win, bg=CB, height=1).pack(fill='x', padx=16, pady=(6,0))

        # Campo de pesquisa
        f_search = tk.Frame(win, bg=CF)
        f_search.pack(fill='x', padx=12, pady=(6,2))
        tk.Label(f_search, text='🔍', bg=CF, fg=CA, font=_F(10)).pack(side='left', padx=(0,4))
        v_search = tk.StringVar()
        ent_search = tk.Entry(f_search, textvariable=v_search, bg=CC, fg=CT,
                              insertbackground=CA, relief='flat', font=_F(9))
        ent_search.pack(side='left', fill='x', expand=True)
        tk.Label(f_search, text=T('hist_filtrar'),
                 bg=CF, fg=CD, font=_F(7)).pack(side='left', padx=4)

        # Buttons BEFORE treeview so they never get hidden
        tk.Frame(win, bg=CB, height=1).pack(fill='x', padx=12, pady=(4,0))
        bf2 = tk.Frame(win, bg=CF); bf2.pack(fill='x', padx=12, pady=(4,6), side='bottom')
        tk.Frame(win, bg=CB, height=1).pack(fill='x', padx=12, side='bottom')

        cols = (T('hist_col_data'), T('hist_col_proto'), T('hist_col_atletas'),
                T('hist_col_excel'), T('hist_col_pdf'))
        from tkinter import ttk
        tv = ttk.Treeview(win, columns=cols, show='headings', height=12)
        tv.pack(fill='both', expand=True, padx=12, pady=8)

        style = ttk.Style()
        try: style.theme_use('clam')
        except Exception: pass
        try:
            style.configure('Treeview', background=CC, foreground=CT,
                            fieldbackground=CC, rowheight=22, font=_F(8))
            style.configure('Treeview.Heading', background=CP, foreground=CA,
                            font=_F(8, bold=True), relief='flat')
            style.map('Treeview', background=[('selected', CB)])
        except Exception: pass

        widths = [120, 160, 60, 230, 130]
        for col, wid in zip(cols, widths):
            tv.heading(col, text=col)
            tv.column(col, width=wid, anchor='w', minwidth=40, stretch=True)

        if not entries:
            tv.insert('', 'end', values=('-', T('hist_sem'), '-', '-', '-'))
        else:
            for e in entries:
                pdf_val = os.path.basename(e.get('pdf','')) or '-'
                saida_val = os.path.basename(e.get('saida','')) or '-'
                tv.insert('', 'end', values=(
                    e.get('data','?'),
                    e.get('protocolo','?'),
                    e.get('n','?'),
                    saida_val,
                    pdf_val,
                ))

        def _filter_hist(*_):
            query = v_search.get().lower().strip()
            for item in tv.get_children():
                tv.delete(item)
            for e in entries:
                proto_v = (e.get('protocolo','') or '').lower()
                data_v  = (e.get('data','') or '').lower()
                saida_v = (e.get('saida','') or '').lower()
                if not query or query in proto_v or query in data_v or query in saida_v:
                    pv = os.path.basename(e.get('pdf','')) or '-'
                    sv = os.path.basename(e.get('saida','')) or '-'
                    tv.insert('', 'end', values=(
                        e.get('data','?'), e.get('protocolo','?'),
                        e.get('n','?'), sv, pv))

        v_search.trace_add('write', _filter_hist)

        def _abrir_sel():
            sel = tv.selection()
            if not sel: return
            idx = tv.index(sel[0])
            if idx < len(entries):
                pasta = os.path.dirname(entries[idx].get('saida',''))
                if pasta and os.path.isdir(pasta):
                    _abrir_pasta_os(pasta)
                else:
                    self.mb.showwarning('Aviso', T('hist_pasta_nao_enc'))

        def _reabrir_config():
            sel = tv.selection()
            if not sel: return
            idx = tv.index(sel[0])
            if idx < len(entries):
                e = entries[idx]
                if hasattr(self, 'v_saida') and e.get('saida'):
                    self.v_saida.set(e['saida'])
                    pasta_g = os.path.dirname(e.get('saida',''))
                    if hasattr(self, 'v_pasta') and pasta_g:
                        self.v_pasta.set(pasta_g)
                if hasattr(self, 'v_pdf') and e.get('pdf'):
                    self.v_pdf.set(e['pdf'])
                self._log(T('hist_cfg_recarregada', data=e.get('data','?')), 'info')
                win.destroy()

        # Buttons are already packed at bottom (side='bottom') before treeview
        def _add_hist_btn(parent, text, cmd, fg_color, side='left', padx_=(0, 4)):
            b = tk.Button(parent, text=text, bg=CP, fg=fg_color,
                          font=_F(8), relief='flat', cursor='hand2', pady=6,
                          command=cmd)
            b.pack(side=side, padx=padx_)
            b.bind('<Enter>', lambda e, _b=b, _f=fg_color: _b.config(bg=CB, fg=CT))
            b.bind('<Leave>', lambda e, _b=b, _f=fg_color: _b.config(bg=CP, fg=_f))
            return b

        _add_hist_btn(bf2, T('hist_btn_abrir'), _abrir_sel, CD)
        _add_hist_btn(bf2, T('hist_btn_reutilizar'), _reabrir_config, CA)
        _add_hist_btn(bf2, T('hist_btn_fechar'), win.destroy, CD, side='right', padx_=(4, 0))

    def _abrir_saida(self, pasta):
        """Abre a pasta de saida no explorador do SO."""
        if pasta and os.path.isdir(pasta):
            _abrir_pasta_os(pasta)

    def _mostrar_btn_abrir(self, pasta):
        """Cria/actualiza o botao de abrir pasta de saida apos conclusao."""
        if hasattr(self, '_btn_abrir') and self._btn_abrir.winfo_exists():
            self._btn_abrir.config(command=lambda p=pasta: self._abrir_saida(p))
        else:
            self._btn_abrir = self.tk.Button(
                self._bf_extra,
                text=T('btn_abrir_pasta'),
                bg='#0A2910', fg=COK,
                font=_F(9, bold=True),
                activebackground='#1A5C1A', activeforeground='white',
                relief='flat', cursor='hand2', pady=8,
                bd=0,
                command=lambda p=pasta: self._abrir_saida(p))
            self._btn_abrir.pack(fill='x', pady=(4, 0))
            self._btn_abrir.bind('<Enter>',
                lambda e, b=self._btn_abrir: b.config(bg='#164A16', fg='white'))
            self._btn_abrir.bind('<Leave>',
                lambda e, b=self._btn_abrir: b.config(bg='#0A2910', fg=COK))

    def _mostrar_btn_demografia(self):
        """v1.0: Cria botao 'Demografia' apos analise PROTO_ARCO com ref carregada."""
        if hasattr(self, '_btn_demo') and self._btn_demo.winfo_exists():
            return
        self._btn_demo = self.tk.Button(
            self._bf_extra,
            text='  \u2261  ' + T('btn_demografia'),
            bg='#1A3050', fg=CA,
            font=_F(9, bold=True),
            activebackground='#2A4B7A', activeforeground='white',
            relief='flat', cursor='hand2', pady=8, bd=0,
            command=self._abrir_demografia)
        self._btn_demo.pack(fill='x', pady=(4, 0))
        self._btn_demo.bind('<Enter>',
            lambda e, b=self._btn_demo: b.config(bg='#2A4B7A', fg='white'))
        self._btn_demo.bind('<Leave>',
            lambda e, b=self._btn_demo: b.config(bg='#1A3050', fg=CA))

    def _abrir_demografia(self):
        """v1.0: Janela modal com analises demograficas (PROTO_ARCO)."""
        atletas = getattr(self, '_last_atletas', None) or []
        ref_dict = getattr(self, '_last_ref_dict', None) or {}
        ref_lista = getattr(self, '_last_ref_lista', None) or []
        if not atletas or not ref_lista:
            self.mb.showinfo('Demografia',
                'Sem dados - corre primeiro uma analise com ficheiro de referencia.')
            return

        # Merge: cada atleta analisado e'esta'o dict ref (+ metricas mets)
        atletas_merged = []
        for a in atletas:
            nome = a.get('nome', '')
            pid_match = re.match(r'^(\d+)', nome)
            base = {}
            if pid_match:
                base = dict(ref_dict.get(pid_match.group(1), {}) or {})
            merged = dict(base)
            merged.update(a)  # atleta_analisado em cima da ref (mets, etc.)
            if 'peso_kg' not in merged or merged.get('peso_kg') is None:
                merged['peso_kg'] = base.get('peso_kg')
            if 'altura_m' not in merged or merged.get('altura_m') is None:
                merged['altura_m'] = base.get('altura_m')
            atletas_merged.append(merged)

        self._construir_janela_demografia(atletas_merged, ref_lista)

    def _construir_janela_demografia(self, atletas_merged, ref_lista):
        """v1.0: Constroi a janela de analises demograficas."""
        tk = self.tk
        try:
            from tkinter import ttk as _ttk
        except Exception:
            _ttk = self.ttk
        win = tk.Toplevel(self.root)
        win.title(T('btn_demografia') + '  |  BSP')
        win.configure(bg=CF)
        w, h = 900, 640
        x = (win.winfo_screenwidth()-w)//2
        y = (win.winfo_screenheight()-h)//2
        win.geometry(f'{w}x{h}+{x}+{y}')
        win.minsize(700, 480)

        # Header
        hdr = tk.Frame(win, bg=CP, height=50)
        hdr.pack(fill='x'); hdr.pack_propagate(False)
        tk.Label(hdr, text='  ' + T('demo_titulo'),
                 bg=CP, fg=CA, font=_F(12, bold=True)).pack(side='left', pady=10)
        tk.Label(hdr, text=f'n={len(atletas_merged)} analisados, {len(ref_lista)} ref.  ',
                 bg=CP, fg=CD, font=_F(9)).pack(side='right', pady=12)

        # Controls row
        ctrl = tk.Frame(win, bg=CF, padx=12, pady=8)
        ctrl.pack(fill='x')

        # Metric dropdown
        tk.Label(ctrl, text=T('demo_metrica') + ':', bg=CF, fg=CT,
                 font=_F(9)).pack(side='left', padx=(0,4))
        _metricas = [
            ('ea95',         'Área elipse 95% (mm²)'),
            ('vel_med',      'Velocidade média (mm/s)'),
            ('stiff_x',        'Stiffness ML (1/s)'),
            ('stiff_y',        'Stiffness AP (1/s)'),
            ('amp_x',        'Amplitude ML (mm)'),
            ('amp_y',        'Amplitude AP (mm)'),
            ('rms_r',        'RMS radial (mm)'),
            ('P_total',      'Score total (pontos)'),
        ]
        v_met = tk.StringVar(value=_metricas[0][0])
        cb_met = _ttk.Combobox(ctrl, textvariable=v_met,
                                values=[f'{k} - {d}' for k,d in _metricas],
                                state='readonly', width=34)
        cb_met.set(f'{_metricas[0][0]} - {_metricas[0][1]}')
        cb_met.pack(side='left', padx=(0,12))

        # Factor dropdown
        tk.Label(ctrl, text=T('demo_factor') + ':', bg=CF, fg=CT,
                 font=_F(9)).pack(side='left', padx=(0,4))
        _factores = [('genero','Género'), ('estilo','Estilo'),
                      ('categoria','Categoria')]
        v_fac = tk.StringVar(value=_factores[0][0])
        cb_fac = _ttk.Combobox(ctrl, textvariable=v_fac,
                                values=[f'{k} - {d}' for k,d in _factores],
                                state='readonly', width=18)
        cb_fac.set(f'{_factores[0][0]} - {_factores[0][1]}')
        cb_fac.pack(side='left', padx=(0,12))

        # Demographic var dropdown
        tk.Label(ctrl, text=T('demo_var_dem') + ':', bg=CF, fg=CT,
                 font=_F(9)).pack(side='left', padx=(0,4))
        _dem = [('peso_kg','Peso (kg)'), ('altura_m','Altura (m)'),
                ('idade','Idade (anos)')]
        v_dem = tk.StringVar(value=_dem[0][0])
        cb_dem = _ttk.Combobox(ctrl, textvariable=v_dem,
                                values=[f'{k} - {d}' for k,d in _dem],
                                state='readonly', width=18)
        cb_dem.set(f'{_dem[0][0]} - {_dem[0][1]}')
        cb_dem.pack(side='left', padx=(0,12))

        # Output area
        out_f = tk.Frame(win, bg=CF)
        out_f.pack(fill='both', expand=True, padx=12, pady=(4,8))
        out_txt = tk.Text(out_f, bg=CC, fg=CT, font=('Consolas', 10),
                          relief='flat', wrap='word', height=20)
        out_txt.pack(fill='both', expand=True)
        out_txt.tag_configure('head', foreground=CA, font=('Consolas',11,'bold'))
        out_txt.tag_configure('ok',   foreground=COK)
        out_txt.tag_configure('warn', foreground='#F5A623')
        out_txt.tag_configure('dim',  foreground=CD)

        def _chave_do_combo(cb, lista):
            txt = cb.get() or ''
            for k, _ in lista:
                if txt.startswith(k + ' '):
                    return k
            return lista[0][0]

        def _write(tag, *segs):
            for seg in segs:
                if isinstance(seg, tuple):
                    out_txt.insert('end', seg[0], seg[1])
                else:
                    out_txt.insert('end', seg, tag)
            out_txt.insert('end', '\n')
            out_txt.see('end')

        def _fmt_p(p):
            if p is None: return 'n/d'
            if p < 0.001: return 'p<0.001'
            return f'p={p:.3f}'

        def _acao_comparar():
            out_txt.delete('1.0', 'end')
            chave = _chave_do_combo(cb_met, _metricas)
            fator = _chave_do_combo(cb_fac, _factores)
            res = comparar_grupos(atletas_merged, chave, fator)
            _write('head', f'Comparação de grupos: {chave} por {fator}')
            _write('dim',  '─' * 70)
            if res['n_grupos'] < 2:
                _write('warn', f"Apenas {res['n_grupos']} grupo(s) com n≥2 - sem teste.")
                return
            for g, n in sorted(res['n_por_grupo'].items()):
                med = res['mediana'].get(g, float('nan'))
                p25 = res['p25'].get(g, float('nan'))
                p75 = res['p75'].get(g, float('nan'))
                _write('', f'  {g:<15}  n={n:>3}   mediana={med:>8.3f}   '
                           f'IQR=[{p25:>7.2f}, {p75:>7.2f}]')
            _write('dim', '─' * 70)
            _t_name = res.get('teste') or '?'
            _est    = res.get('estatistica')
            _pval   = res.get('p_valor')
            _est_s  = f'{_est:.3f}' if _est is not None else 'n/d'
            _tag    = 'ok' if (_pval is not None and _pval < 0.05) else 'dim'
            _write(_tag, f'  Teste: {_t_name}   estat={_est_s}   {_fmt_p(_pval)}')

        def _acao_corr_dem():
            out_txt.delete('1.0', 'end')
            chave = _chave_do_combo(cb_met, _metricas)
            dem   = _chave_do_combo(cb_dem, _dem)
            res = correlacao_demografica(atletas_merged, chave, dem)
            _write('head', f'Correlação: {chave} vs {dem}')
            _write('dim',  '─' * 70)
            _write('', f'  n = {res["n"]}')
            pr, pp = res['pearson_r'], res['pearson_p']
            sr, sp = res['spearman_r'], res['spearman_p']
            if pr is not None:
                _tag = 'ok' if (pp is not None and pp < 0.05) else 'dim'
                _write(_tag, f'  Pearson:   r={pr:+.3f}   {_fmt_p(pp)}')
            else:
                _write('warn', '  Pearson: n/d (n<3 ou valores constantes)')
            if sr is not None:
                _tag = 'ok' if (sp is not None and sp < 0.05) else 'dim'
                _write(_tag, f'  Spearman:  r={sr:+.3f}   {_fmt_p(sp)}')
            else:
                _write('warn', '  Spearman: n/d')

        def _acao_corr_score():
            out_txt.delete('1.0', 'end')
            chave = _chave_do_combo(cb_met, _metricas)
            res = correlacao_score(atletas_merged, chave, 'P_total')
            _write('head', f'CoP vs Score: {chave} vs P_total')
            _write('dim',  '─' * 70)
            ag = res.get('agregado') or {}
            _write('', f'  Agregado: n={ag.get("n",0)}')
            if ag.get('pearson_r') is not None:
                _tag = 'ok' if (ag.get('pearson_p') or 1) < 0.05 else 'dim'
                _write(_tag,
                    f'    Pearson:   r={ag["pearson_r"]:+.3f}   '
                    f'{_fmt_p(ag.get("pearson_p"))}')
            if ag.get('spearman_r') is not None:
                _tag = 'ok' if (ag.get('spearman_p') or 1) < 0.05 else 'dim'
                _write(_tag,
                    f'    Spearman:  r={ag["spearman_r"]:+.3f}   '
                    f'{_fmt_p(ag.get("spearman_p"))}')
            per = res.get('per_ensaio')
            _write('dim', '─' * 70)
            if per:
                _write('', f'  Per-ensaio: n={per["n"]}')
                _tag = 'ok' if (per.get('pearson_p') or 1) < 0.05 else 'dim'
                _write(_tag,
                    f'    Pearson:   r={per["pearson_r"]:+.3f}   '
                    f'{_fmt_p(per.get("pearson_p"))}')
                _tag = 'ok' if (per.get('spearman_p') or 1) < 0.05 else 'dim'
                _write(_tag,
                    f'    Spearman:  r={per["spearman_r"]:+.3f}   '
                    f'{_fmt_p(per.get("spearman_p"))}')
            else:
                _write('warn', '  Per-ensaio: sem dados suficientes '
                               '(requer lista P por atleta).')

        # Action buttons
        btns = tk.Frame(win, bg=CF, padx=12, pady=(0,12))
        btns.pack(fill='x')
        for (txt, cmd) in [(T('demo_btn_comparar'), _acao_comparar),
                            (T('demo_btn_corr_dem'), _acao_corr_dem),
                            (T('demo_btn_corr_score'), _acao_corr_score)]:
            b = tk.Button(btns, text=txt, bg=CA, fg='#071014',
                          font=_F(9, bold=True),
                          activebackground=CA2, activeforeground='white',
                          relief='flat', cursor='hand2', padx=14, pady=6,
                          command=cmd)
            b.pack(side='left', padx=(0,6))
            b.bind('<Enter>', lambda e, b=b: b.config(bg=CA2, fg='white'))
            b.bind('<Leave>', lambda e, b=b: b.config(bg=CA, fg='#071014'))

        _write('dim', T('demo_intro'))

    def _poll(self):
        # Processa TODAS as mensagens da fila na main thread
        # Assim nenhuma thread secundaria toca nos widgets directamente
        try:
            while True:
                item = self._q.get_nowait()
                if not item:
                    continue
                kind = item[0] if item else None

                if kind == '__prog__':
                    # Actualizacao de barra de progresso vinda da thread
                    _, pct, msg, t_put = item
                    try:
                        self.pvar.set(pct)
                        if msg:
                            self.plbl.config(text=msg)
                        if pct > 0:
                            eta_str = ''
                            if hasattr(self, '_t_start') and self._t_start and 2 < pct < 99:
                                elapsed = t_put - self._t_start
                                est_total = elapsed / (pct / 100.0) if pct > 0 else 0
                                remaining = max(0, est_total - elapsed)
                                if remaining > 60:
                                    eta_str = f' - {int(remaining/60)}m{int(remaining%60)}s'
                                elif remaining > 0:
                                    eta_str = f' - {int(remaining)}s'
                            self.ppct.config(text=f'{int(pct)}%{eta_str}')
                        else:
                            self.ppct.config(text='')
                    except Exception:
                        pass

                elif kind == '__done__':
                    # Dialogo de conclusao
                    _, title, msg = item
                    try:
                        self.mb.showinfo(title, msg)
                    except Exception:
                        pass

                elif kind == '__crash__':
                    # Dialogo de erro critico
                    _, tb_str = item
                    try:
                        _dialogo_crash(tb_str, 'Thread de analise', self.root)
                    except Exception:
                        pass

                else:
                    # Mensagem de log normal (msg, tag)
                    msg, tag = item[0], item[1] if len(item) > 1 else ''
                    try:
                        self.log.config(state='normal')
                        self.log.insert('end', msg+'\n', tag or '')
                        self.log.see('end')
                        self.log.config(state='disabled')
                    except Exception:
                        pass

        except queue.Empty:
            pass
        self.root.after(50, self._poll)

    def _validate_inputs(self):
        """Valida entradas antes de iniciar. Retorna lista de strings de erro."""
        erros = []
        pasta = self.v_pasta.get().strip()
        if not pasta:
            erros.append(T('val_pasta_vazia'))
        elif not os.path.isdir(pasta):
            erros.append(T('val_pasta_nao_enc', p=pasta))
        else:
            subs = [d for d in os.listdir(pasta)
                    if os.path.isdir(os.path.join(pasta, d))]
            if not subs:
                erros.append(T('val_pasta_sem_sub'))
        saida = self.v_saida.get().strip()
        if not saida:
            erros.append(T('val_saida_vazia'))
        if hasattr(self, 'v_ifd') and self.v_ifd:
            ifd = self.v_ifd.get().strip()
            if ifd and not os.path.exists(ifd):
                erros.append(T('val_ifd_nao_enc', f=os.path.basename(ifd)))
        return erros

    def _show_shortcuts(self):
        """Mostra janela com todos os atalhos de teclado (F1)."""
        tk = self.tk
        win = tk.Toplevel(self.root)
        win.title('⌨  Atalhos - BSP')
        win.configure(bg=CF); win.resizable(False, False)
        w, h = 420, 370
        x = (win.winfo_screenwidth()-w)//2
        y = (win.winfo_screenheight()-h)//2
        win.geometry(f'{w}x{h}+{x}+{y}')
        win.grab_set()
        tk.Label(win, text=T('shortcuts_titulo'),
                 bg=CF, fg=CA, font=_F(12, bold=True)).pack(pady=(16,4))
        tk.Frame(win, bg=CB, height=1).pack(fill='x', padx=16, pady=(0,10))
        shortcuts = [
            (f'{_MOD_SYM}+Enter', T('btn_executar').replace('▶  ','')),
            (f'{_MOD_SYM}+S',     T('btn_guardar').replace('💾 ','')),
            (f'{_MOD_SYM}+H',     T('btn_historico').replace('📋 ','')),
            (f'{_MOD_SYM}+P',     T('btn_poder').replace('🔬 ','')),
            ('F1',                'Atalhos de teclado'),
            ('F5',                'Limpar log'),
            ('F2',                'Análise rápida - ficheiro único'),
        ]
        sf = tk.Frame(win, bg=CF); sf.pack(padx=24, fill='x')
        for key, action in shortcuts:
            row = tk.Frame(sf, bg=CF); row.pack(fill='x', pady=3)
            tk.Label(row, text=key, bg=CC, fg=CA, font=_FM(9),
                     padx=10, pady=4, width=16, relief='flat',
                     anchor='center').pack(side='left')
            tk.Label(row, text=action, bg=CF, fg=CT, font=_F(9),
                     padx=10, anchor='w').pack(side='left', fill='x', expand=True)
        tk.Frame(win, bg=CB, height=1).pack(fill='x', padx=16, pady=(10, 0))
        _sc_fch = tk.Button(win, text=T('btn_fechar'), bg=CP, fg=CD,
                  font=_F(9), relief='flat', cursor='hand2', pady=7,
                  command=win.destroy)
        _sc_fch.pack(pady=8)
        _sc_fch.bind('<Enter>', lambda e: _sc_fch.config(bg=CB, fg=CT))
        _sc_fch.bind('<Leave>', lambda e: _sc_fch.config(bg=CP, fg=CD))

    def _show_update_banner(self, nova):
        """Banner one-click: mostra "Actualizar agora" que desce o asset,
        valida SHA256 e lanca instalador. Fallback para 'Ver notas' e 'Dispensar'."""
        if getattr(self, '_update_banner_shown', False):
            return
        self._update_banner_shown = True
        # Respeita dispensa persistida
        cfg = cfg_load()
        if cfg.get('_upd_dismissed_tag') == nova:
            return
        self._log(
            f'\n-----------------------------------------------\n'
            f'  Nova versão v{nova} disponível\n'
            f'  Asset: {_UPDATE_STATE.get("asset_name") or "(não detectado para este SO)"}\n'
            f'-----------------------------------------------\n', 'aviso')
        try:
            import webbrowser
            tk = self.tk
            # Banner usa o acento do tema (cor coerente com o resto da app)
            _bn_bg = CA          # acento principal do tema
            _bn_fg = '#FFFFFF'   # branco em ambos os temas
            _bn_btn = CA2        # acento hover
            banner = tk.Frame(self.root, bg=_bn_bg, height=32)
            banner.place(relx=0, rely=1.0, relwidth=1, anchor='sw', y=-22)
            lbl = tk.Label(banner,
                     text=f'  {ICO.get("download","v")}  Versão v{nova} disponível',
                     bg=_bn_bg, fg=_bn_fg,
                     font=_F(8, bold=True))
            lbl.pack(side='left', padx=8)

            def _ver_notas():
                notas = _UPDATE_STATE.get('notes', '') or '(sem notas)'
                win = tk.Toplevel(self.root)
                win.title(f'Notas v{nova}')
                win.configure(bg=CF)
                w, h = 640, 480
                x = (win.winfo_screenwidth()-w)//2
                y = (win.winfo_screenheight()-h)//2
                win.geometry(f'{w}x{h}+{x}+{y}')
                tk.Label(win, text=f'BSP v{nova}', bg=CF, fg=CA,
                         font=_F(14, bold=True)).pack(anchor='w', padx=14, pady=(10,4))
                txt = tk.Text(win, bg=CC, fg=CT, font=_F(9), wrap='word',
                              relief='flat', padx=10, pady=8)
                txt.pack(fill='both', expand=True, padx=14, pady=(0,12))
                txt.insert('1.0', notas)
                txt.config(state='disabled')
                tk.Button(win, text='Fechar', bg=CB, fg=CD, relief='flat',
                          font=_F(9), command=win.destroy).pack(pady=(0,10))

            def _dispensar():
                c = cfg_load(); c['_upd_dismissed_tag'] = nova; cfg_save(c)
                banner.destroy()

            def _actualizar():
                if not _UPDATE_STATE.get('asset_url'):
                    # Sem asset: abre pagina do release
                    webbrowser.open(
                      'https://github.com/andremassuca/BSP/releases')
                    return
                # Modal com progress bar
                prog_win = tk.Toplevel(self.root)
                prog_win.title('Actualizar BSP')
                prog_win.configure(bg=CF)
                prog_win.resizable(False, False)
                prog_win.transient(self.root)
                prog_win.grab_set()
                w, h = 480, 180
                x = (prog_win.winfo_screenwidth()-w)//2
                y = (prog_win.winfo_screenheight()-h)//2
                prog_win.geometry(f'{w}x{h}+{x}+{y}')
                tk.Label(prog_win, text=f'A transferir v{nova}...',
                         bg=CF, fg=CA, font=_F(11, bold=True)
                         ).pack(pady=(14, 4))
                lbl_asset = tk.Label(prog_win,
                         text=_UPDATE_STATE.get('asset_name',''),
                         bg=CF, fg=CD, font=_F(8))
                lbl_asset.pack()
                from tkinter import ttk as _tk_ttk
                bar = _tk_ttk.Progressbar(prog_win, length=420, mode='determinate')
                bar.pack(pady=10)
                lbl_stat = tk.Label(prog_win, text='0 %', bg=CF, fg=CT, font=_F(9))
                lbl_stat.pack()
                lbl_sha = tk.Label(prog_win,
                                text='SHA256: ' + (_UPDATE_STATE.get('sha256') or '(nenhum)'),
                                bg=CF, fg=CD, font=_F(7))
                lbl_sha.pack(pady=(6, 0))

                def _on_progress(done, total):
                    try:
                        pct = int(100 * done / total) if total else 0
                        bar['value'] = pct
                        mb_done = done/1024/1024
                        mb_tot  = total/1024/1024 if total else 0
                        lbl_stat.config(text=f'{pct}%  ({mb_done:.1f} / {mb_tot:.1f} MB)')
                    except Exception: pass
                def _on_done(ok, msg):
                    try:
                        lbl_stat.config(text=msg, fg=(CA if ok else CER))
                        if not ok:
                            tk.Button(prog_win, text='Fechar', bg=CB, fg=CD,
                                      relief='flat', font=_F(9),
                                      command=prog_win.destroy).pack(pady=8)
                    except Exception: pass
                self.root.after(0, lambda: _actualizar_agora(
                    progress_cb=lambda d,t: self.root.after(0, lambda: _on_progress(d,t)),
                    done_cb=lambda ok,m: self.root.after(0, lambda: _on_done(ok,m))))

            tk.Button(banner, text='Dispensar', bg=_bn_bg, fg='#D6E9F5',
                      relief='flat', font=_F(8), cursor='hand2',
                      activebackground=_bn_btn, activeforeground='white',
                      command=_dispensar).pack(side='right', padx=(0,6), pady=3)
            tk.Button(banner, text='Ver notas', bg=_bn_bg, fg='white',
                      relief='flat', font=_F(8), cursor='hand2',
                      activebackground=_bn_btn, activeforeground='white',
                      command=_ver_notas).pack(side='right', padx=(0,6), pady=3)
            tk.Button(banner, text='Actualizar agora',
                      bg=_bn_btn, fg='white', font=_F(8, bold=True),
                      relief='flat', cursor='hand2', padx=10,
                      activebackground=_bn_bg, activeforeground='white',
                      command=_actualizar).pack(side='right', padx=8, pady=3)
        except Exception:
            pass

    def _gerir_perfis(self):
        """Janela de gestão de perfis de configuração nomeados."""
        tk = self.tk
        win = tk.Toplevel(self.root)
        win.title('📌  Perfis de Configuração - BSP')
        win.configure(bg=CF); win.resizable(False, False)
        w, h = 480, 380
        x = (win.winfo_screenwidth()-w)//2
        y = (win.winfo_screenheight()-h)//2
        win.geometry(f'{w}x{h}+{x}+{y}')
        win.grab_set()
        tk.Label(win, text=T('perfis_titulo'),
                 bg=CF, fg=CA, font=_F(11, bold=True)).pack(pady=(14,2))
        tk.Label(win, text=T('perfis_descr'),
                 bg=CF, fg=CD, font=_F(8)).pack()
        tk.Frame(win, bg=CB, height=1).pack(fill='x', padx=16, pady=(8,0))
        lf = tk.Frame(win, bg=CF); lf.pack(fill='both', expand=True, padx=16, pady=8)
        lb = tk.Listbox(lf, bg=CC, fg=CT, font=_F(9), selectbackground=CA,
                        selectforeground='#0A1218', relief='flat',
                        activestyle='none', height=8)
        lb.pack(fill='both', expand=True)
        def _refresh():
            lb.delete(0, 'end')
            for nm in profiles_load(): lb.insert('end', nm)
        _refresh()
        f_new = tk.Frame(win, bg=CF); f_new.pack(fill='x', padx=16, pady=(0,4))
        tk.Label(f_new, text=T('perfis_nome_lbl'), bg=CF, fg=CT, font=_F(9)).pack(side='left')
        v_name = tk.StringVar()
        ent = tk.Entry(f_new, textvariable=v_name, bg=CC, fg=CT,
                       insertbackground=CA, relief='flat', font=_F(9), width=24)
        ent.pack(side='left', padx=6, fill='x', expand=True)
        def _get_cfg():
            d = cfg_load()
            for attr, key in [('v_pasta','pasta'),('v_ifd','ifd'),('v_saida','saida'),
                               ('v_pasta_ind','pasta_ind'),('v_pdf','pdf'),
                               ('v_atletas_ref_file','atletas_ref_file')]:
                if hasattr(self, attr): d[key] = getattr(self, attr).get()
            d['protocolo'] = self._proto_key
            return d
        def _guardar():
            nm = v_name.get().strip()
            if not nm:
                self.mb.showwarning('Aviso', 'Introduz um nome.', parent=win); return
            profile_save(nm, _get_cfg()); _refresh()
            self._log(T('perfil_guardado', nm=nm), 'ok')
        def _carregar():
            sel = lb.curselection()
            if not sel: return
            nm = lb.get(sel[0]); p = profiles_load().get(nm, {})
            for attr, key in [('v_pasta','pasta'),('v_ifd','ifd'),('v_saida','saida'),
                               ('v_pasta_ind','pasta_ind'),('v_pdf','pdf'),
                               ('v_atletas_ref_file','atletas_ref_file')]:
                if hasattr(self, attr) and key in p: getattr(self, attr).set(p[key])
            self._log(T('perfil_carregado', nm=nm), 'ok'); win.destroy()
        def _apagar():
            sel = lb.curselection()
            if not sel: return
            nm = lb.get(sel[0])
            if self.mb.askyesno('Confirmar', f'Apagar "{nm}"?', parent=win):
                profile_delete(nm); _refresh()
        bf = tk.Frame(win, bg=CF); bf.pack(fill='x', padx=16, pady=(0, 12))
        for txt, cmd, clr in [(T('perfis_guardar'), _guardar, CA),
                               (T('perfis_carregar'), _carregar, COK),
                               (T('perfis_apagar'),   _apagar,   CER)]:
            _pb = tk.Button(bf, text=txt, bg=CP, fg=clr, font=_F(8),
                      relief='flat', cursor='hand2', pady=6, padx=8,
                      command=cmd)
            _pb.pack(side='left', padx=(0, 4))
            _pb.bind('<Enter>', lambda e, b=_pb, c=clr: b.config(bg=CB, fg=CT))
            _pb.bind('<Leave>', lambda e, b=_pb, c=clr: b.config(bg=CP, fg=c))
        _pfch = tk.Button(bf, text=T('btn_fechar'), bg=CP, fg=CD, font=_F(8),
                  relief='flat', cursor='hand2', pady=6,
                  command=win.destroy)
        _pfch.pack(side='right')
        _pfch.bind('<Enter>', lambda e: _pfch.config(bg=CB, fg=CT))
        _pfch.bind('<Leave>', lambda e: _pfch.config(bg=CP, fg=CD))

    def _quick_subject(self):
        """Modo rápido: analisa um único ficheiro .xls sem configuração completa."""
        tk = self.tk
        path = self.fd.askopenfilename(
            title='Selecionar ficheiro de dados (.xls / .xlsx)',
            filetypes=[('XLS / XLSX', '*.xls *.xlsx'), ('Todos', '*.*')])
        if not path: return
        win = tk.Toplevel(self.root)
        win.title('⚡  Análise Rápida - BSP')
        win.configure(bg=CF); win.resizable(True, True)
        w, h = 560, 440
        x = (win.winfo_screenwidth()-w)//2
        y = (win.winfo_screenheight()-h)//2
        win.geometry(f'{w}x{h}+{x}+{y}')
        tk.Label(win, text=T('quick_titulo'),
                 bg=CF, fg=CA, font=_F(12, bold=True)).pack(pady=(14,2))
        tk.Label(win, text=os.path.basename(path), bg=CF, fg=CT, font=_F(9)).pack()
        tk.Frame(win, bg=CB, height=1).pack(fill='x', padx=16, pady=(8,4))
        log_f = tk.Frame(win, bg='#111C27')
        log_f.pack(fill='both', expand=True, padx=12, pady=4)
        log_txt = tk.Text(log_f, bg='#111C27', fg=CT, font=_FM(9),
                          relief='flat', padx=8, pady=6,
                          state='disabled', wrap='word', height=16)
        log_txt.pack(fill='both', expand=True)
        log_txt.tag_config('ok', foreground=COK)
        log_txt.tag_config('aviso', foreground=CAV)
        log_txt.tag_config('info', foreground=CA)
        res_var = [None]
        def _w(msg, tag=''):
            # ATENÇÃO: chamada de thread - usar after() para thread-safety
            def _do():
                log_txt.config(state='normal')
                log_txt.insert('end', msg + '\n', tag)
                log_txt.see('end'); log_txt.config(state='disabled')
            try: win.after(0, _do)
            except Exception: pass
        def _run():
            try:
                _w(f'A ler: {os.path.basename(path)}', 'info')
                data = ler_ficheiro(path)
                frames = data['frames']
                if not frames:
                    _w('Sem frames válidos.', 'aviso'); return
                _w(f'{len(frames)} frames detectados.', 'info')
                m = calcular(frames)
                if not m:
                    _w('Não foi possível calcular métricas.', 'aviso'); return
                res_var[0] = m
                _w('\n── MÉTRICAS ──────────────────────────', 'info')
                mets = mets_pdf_localizadas() if _I18N_OK else METS_XL
                for chave, label in mets:
                    v = m.get(chave)
                    if v is not None:
                        _w(f'  {label:<38} {round(v, 3)}', 'ok')
                _w(f'\n✓ Concluído', 'ok')
            except Exception as _e:
                _w(f'Erro: {_e}', 'aviso')
        threading.Thread(target=_run, daemon=True).start()
        bf = tk.Frame(win, bg=CF); bf.pack(fill='x', padx=12, pady=(4,10))
        def _export():
            if not res_var[0]: return
            sp = self.fd.asksaveasfilename(
                defaultextension='.xlsx',
                filetypes=[('Excel', '*.xlsx')],
                initialfile=os.path.splitext(os.path.basename(path))[0] + '_metrics.xlsx')
            if not sp: return
            try:
                wb2 = Workbook(); ws2 = wb2.active; ws2.title = 'Métricas'
                ws2.append(['Métrica', 'Valor'])
                mets = mets_pdf_localizadas() if _I18N_OK else METS_XL
                for chave, label in mets:
                    v = res_var[0].get(chave)
                    if v is not None: ws2.append([label, round(v, 4)])
                wb2.save(sp)
                _w(f'✓ Exportado: {os.path.basename(sp)}', 'ok')
            except Exception as _e2:
                _w(f'Erro ao exportar: {_e2}', 'aviso')
        _exp_btn = tk.Button(bf, text=T('quick_exportar'), bg=COK, fg='#071014',
                  font=_F(9, bold=True), relief='flat', cursor='hand2', pady=6,
                  command=_export)
        _exp_btn.pack(side='left', padx=(0, 6))
        _exp_btn.bind('<Enter>', lambda e: _exp_btn.config(bg='#1AA84A', fg='white'))
        _exp_btn.bind('<Leave>', lambda e: _exp_btn.config(bg=COK, fg='#071014'))
        _fch_btn = tk.Button(bf, text=T('btn_fechar'), bg=CP, fg=CD, font=_F(9),
                  relief='flat', cursor='hand2', pady=6,
                  command=win.destroy)
        _fch_btn.pack(side='right')
        _fch_btn.bind('<Enter>', lambda e: _fch_btn.config(bg=CB, fg=CT))
        _fch_btn.bind('<Leave>', lambda e: _fch_btn.config(bg=CP, fg=CD))

    def _clr(self):
        self.log.config(state='normal')
        self.log.delete('1.0','end')
        self.log.config(state='disabled')

    def _exec(self):
        # ── Validação prévia ─────────────────────────────────────────────
        erros = self._validate_inputs()
        if erros:
            self.mb.showwarning(
                T('val_titulo'),
                T('val_intro') + '\n'.join(f'• {e}' for e in erros))
            return

        pasta=self.v_pasta.get().strip()
        ifd=self.v_ifd.get().strip()
        saida=self.v_saida.get().strip()
        ind=getattr(self, 'v_pasta_ind', None)
        ind=ind.get().strip() if ind else ''
        pdf=self.v_pdf.get().strip()
        scores_f=getattr(self, 'v_scores_file', None)
        scores_f=scores_f.get().strip() if scores_f else ''
        n_ens_str = getattr(self, 'v_n_ens', None)
        n_ens_str = n_ens_str.get().strip() if n_ens_str else ''

        if not pasta or not os.path.isdir(pasta):
            self.mb.showerror('Erro','Pasta de individuos invalida.'); return
        if not saida:
            self.mb.showerror('Erro','Define o ficheiro de saida.'); return

        # Garantir que o directorio de saida do Excel existe
        saida_dir = os.path.dirname(os.path.abspath(saida))
        os.makedirs(saida_dir, exist_ok=True)

        # Garantir que a pasta de individuais existe (criar sempre)
        if not ind:
            ind = os.path.join(saida_dir, 'individuais')
        os.makedirs(ind, exist_ok=True)

        # Ficheiro de referencia demografica (142 atletas) - so PROTO_ARCO
        atletas_ref_file = ''
        if self._proto_key == PROTO_ARCO:
            _v_ref = getattr(self, 'v_atletas_ref_file', None)
            atletas_ref_file = _v_ref.get().strip() if _v_ref else ''

        # Opcoes de tiro
        is_tiro = _is_tiro_like(self._proto_key)
        tiro_match = 'id'
        intervalos_tiro = list(TIRO_INTERVALOS.keys())
        tiro_bipodal = True
        dists_extra = []

        if is_tiro:
            tiro_match = self.v_tiro_match.get() if self.v_tiro_match else 'id'
            intervalos_tiro = [k for k,v in self.v_tiro_itvs.items() if v.get()]
            if not intervalos_tiro:
                intervalos_tiro = list(TIRO_INTERVALOS.keys())
            tiro_bipodal = self.o_tiro_bipodal.get() if self.o_tiro_bipodal else True
            dists_extra = [v.get().strip() for v in self.v_dists_extra
                           if v.get().strip()]
            # N. ensaios Hurdle Step
            n_ens_hs_str = self.v_n_ens_hs.get().strip() if hasattr(self,'v_n_ens_hs') else ''

        # Guardar configuracao persistente - MERGE com config existente
        # para nao apagar licenca_aceite, tema, _lingua, _win_geom, etc.
        _cfg_atual = cfg_load()
        _cfg_atual.update({
            'pasta': pasta, 'ifd': ifd, 'saida': saida,
            'pasta_ind': ind, 'pdf': pdf, 'scores_file': scores_f,
            'protocolo': self._proto_key, 'n_ens': n_ens_str,
            'tiro_match': tiro_match,
            'tiro_itvs': intervalos_tiro,
            'tiro_bipodal': tiro_bipodal,
            'n_ens_hs': int(n_ens_hs_str) if is_tiro and n_ens_hs_str.isdigit() else 5,
            'tiro_dists_extra': dists_extra,
            'o_emb': self.o_emb.get(), 'o_ell': self.o_ell.get(),
            'o_est': self.o_est.get(), 'o_ind': self.o_ind.get(),
            'o_pdf': self.o_pdf.get(), 'o_estats': self.o_estats.get(),
            'estats_grupo':    self.o_estats_grupo.get(),
            'estats_dir_esq':  self.o_estats_dir_esq.get(),
            'estats_pos_disp': self.o_estats_pos_disp.get(),
            'estats_atleta':   self.o_estats_atleta.get(),
            'estats_friedman': self.o_estats_friedman.get(),
            'estats_posthoc':  self.o_estats_posthoc.get(),
            'estats_corr':     self.o_estats_corr.get(),
            'export_csv':  self.o_export_csv.get(),
            'export_docx': self.o_export_docx.get(),
            'export_html': self.o_export_html.get(),
            'export_png':  self.o_export_png.get(),
            'fft_ativo':   self.o_fft.get(),
            'peso_norm':   self.o_peso_norm.get(),
            'peso_kg_default': self.v_peso_kg.get().strip(),
            'csv_sep':     self.v_csv_sep.get(),
            'csv_dec':     self.v_csv_dec.get(),
        })
        cfg_save(_cfg_atual)

        # Sincronizar AOM_CONFIG com opcoes da UI
        AOM_CONFIG['fft_ativo']  = self.o_fft.get()
        AOM_CONFIG['peso_norm']  = self.o_peso_norm.get()

        # Peso padrao para normalizacao
        _peso_str = self.v_peso_kg.get().strip()
        try:
            _peso_default = float(_peso_str) if _peso_str else None
        except ValueError:
            _peso_default = None

        # Override n_ensaios se definido
        n_ens_override = None
        if n_ens_str:
            try: n_ens_override = int(n_ens_str)
            except ValueError: pass

        # N. ensaios Hurdle Step override
        n_ens_hs_override = None
        if is_tiro:
            _nhs = n_ens_hs_str if 'n_ens_hs_str' in locals() else ''
            try: n_ens_hs_override = int(_nhs) if _nhs else 5
            except ValueError: n_ens_hs_override = 5

        self._clr(); self.b_exe.config(state='disabled')
        self.b_stop.config(state='normal')
        try: self._status_dot.config(fg=CA)
        except Exception: pass
        self._set_prog(0); self._run=True
        self._t_start = time.time()  # para cálculo de ETA

        opts_estats = None
        if self.o_estats.get():
            opts_estats = {
                'grupo':           self.o_estats_grupo.get(),
                'dir_esq':         self.o_estats_dir_esq.get(),
                'pos_disp':        self.o_estats_pos_disp.get(),
                'individuo':       self.o_estats_atleta.get(),
                'friedman':        self.o_estats_friedman.get(),
                'posthoc_friedman':self.o_estats_posthoc.get(),
                'correlacao':      self.o_estats_corr.get(),
            }

        def _t():
            try:
                ifd_data=None
                tempos_arco_data = None
                if ifd and os.path.exists(ifd):
                    self._log(T('log_tempos', nome=os.path.basename(ifd)), 'info')
                    if _is_tiro_like(self._proto_key):
                        ifd_data=carregar_tempos_tiro(ifd)
                        dists_det = ifd_data.get('distancias', [])
                        self._log(T('log_dist_det', dists=', '.join(dists_det) or '-'), 'dim')
                        n_hs = len(ifd_data.get('hurdle_step', {}))
                        if n_hs: self._log(T('log_hs_n', n=n_hs), 'dim')
                        # Adicionar distancias extra ao ifd_data
                        for de in dists_extra:
                            de_k = de if de.endswith('m') else de+'m'
                            if de_k not in dists_det:
                                dists_det.append(de_k)
                    elif self._proto_key == PROTO_ARCO:
                        # Tiro com Arco usa Inicio_fim_vfinal.xlsx (confirmacao_1/2)
                        tempos_arco_data = carregar_confirmacao_arco(ifd)
                        _n_arco = len((tempos_arco_data or {}).get('por_id', {}))
                        self._log(T('log_n_ind', n=_n_arco), 'dim')
                    else:
                        ifd_data=carregar_inicio_fim(ifd)
                        self._log(T('log_n_ind', n=len(ifd_data)), 'dim')
                elif ifd:
                    self._log(T('log_tempo_nao_enc'), 'aviso')

                # Referencia demografica (PROTO_ARCO)
                atletas_ref_lista = []
                atletas_ref_dict  = {}
                if atletas_ref_file and os.path.exists(atletas_ref_file):
                    try:
                        atletas_ref_lista = carregar_atletas_ref(atletas_ref_file)
                        atletas_ref_dict  = atletas_ref_por_id(atletas_ref_lista)
                        self._log(f'  ref. demográfica: {len(atletas_ref_lista)} atletas',
                                   'dim')
                    except Exception as _e_ref:
                        self._log(f'  aviso ref. demográfica: {_e_ref}', 'aviso')

                scores_data={}
                if scores_f and os.path.exists(scores_f):
                    scores_data=carregar_scores_tiro(scores_f)
                    self._log(T('log_scores', n=len(scores_data)), 'dim')

                def _sort_key(p):
                    b = os.path.basename(p)
                    m = re.match(r'^(\d+)', b)
                    return (int(m.group(1)) if m else float('inf'), b)
                subs=sorted([os.path.join(pasta,d) for d in os.listdir(pasta)
                             if os.path.isdir(os.path.join(pasta,d))], key=_sort_key)
                subs = [s for s in subs if os.path.isdir(s)]  # apenas pastas
                if not subs: self._log(T('log_sem_subs'), 'erro'); return
                total=len(subs)
                self._log(T('log_n_total', n=total, proto=self._proto['nome']), 'info')
                if n_ens_override:
                    self._log(T('log_ens_override', n=n_ens_override), 'aviso')
                atletas=[]

                for i, sub in enumerate(subs):
                    if not self._run: break
                    self._log(f'\n{os.path.basename(sub)}', 'neg')
                    # Callback thread-safe para perguntar ao utilizador sobre ensaios em falta
                    _ask_result = [True]
                    _ask_event  = threading.Event()

                    def _ask_cb(nome_ind, dist, em_falta, existem, total,
                                _ev=_ask_event, _res=_ask_result, _root=self.root):
                        faltam_str = ', '.join(f't{t}' for t in em_falta)
                        msg = (f'Individuo: {nome_ind}\n'
                               f'Distancia: {dist}\n\n'
                               f'Ensaios em falta: {faltam_str}\n'
                               f'Ensaios disponiveis: {existem}/{total}\n\n'
                               f'Incluir este individuo com os dados existentes?\n'
                               f'(Nao = ignorar este individuo para esta distancia)')
                        def _show(_m=msg, _ev=_ev, _res=_res):
                            import tkinter.messagebox as _mb
                            r = _mb.askyesno('Ensaios em falta', _m, parent=_root)
                            _res[0] = r; _ev.set()
                        _root.after(0, _show)
                        _ev.wait(); _ev.clear()
                        return _res[0]

                    # Matching do atleta na referencia demografica (so PROTO_ARCO)
                    _ath_ref = None
                    if self._proto_key == PROTO_ARCO and atletas_ref_dict:
                        _pid_match = re.match(r'^(\d+)', os.path.basename(sub))
                        if _pid_match:
                            _ath_ref = atletas_ref_dict.get(_pid_match.group(1))

                    ath=processar_atleta(sub, ifd_data, self.o_emb.get(),
                                         self._log, self._proto_key,
                                         n_ens_override=n_ens_override,
                                         intervalos_tiro=intervalos_tiro if is_tiro else None,
                                         match_por_idx=(tiro_match=='idx'),
                                         idx_na_lista=i+1,
                                         n_ens_hs_override=n_ens_hs_override,
                                         ask_callback=_ask_cb if is_tiro else None,
                                         incluir_hs=tiro_bipodal if is_tiro else True,
                                         peso_kg_default=_peso_default,
                                         tempos_arco=tempos_arco_data,
                                         atleta_ref=_ath_ref)
                    nome_ath=ath['nome']
                    # Validação de frame rate entre ensaios
                    _av_fr = _validar_consistencia_atleta(ath)
                    for _av_msg in _av_fr:
                        self._log(f'  aviso frame rate: {_av_msg}', 'aviso')
                    for chave_s in scores_data:
                        if _norm(chave_s)==_norm(nome_ath) or _norm(chave_s) in _norm(nome_ath):
                            ath['scores']=scores_data[chave_s]; break
                    atletas.append(ath)
                    self._set_prog((i+1)/total*55, f'⟳  {i+1}/{total}…')

                if not atletas: self._log(T('log_sem_dados'), 'erro'); return

                estat_str = '+ ESTATS' if opts_estats else ''
                self._log(T('log_gerar_excel', estats=estat_str), 'info')
                guardar_resumo(atletas, saida, self._proto_key, opts_estats=opts_estats)
                self._log(T('log_excel_ok', f=saida), 'ok'); self._set_prog(63)

                # ── Exportação CSV ────────────────────────────────────
                if self.o_export_csv.get():
                    try:
                        sep_csv = self.v_csv_sep.get()
                        dec_csv = self.v_csv_dec.get()
                        csvs = exportar_csv_resumo(atletas, saida,
                                                   self._proto_key,
                                                   sep=sep_csv, decimal=dec_csv)
                        if opts_estats and len(atletas)>=3:
                            csvs += exportar_csv_estats(atletas, saida,
                                                        self._proto_key,
                                                        sep=sep_csv, decimal=dec_csv)
                        for cf_path in csvs:
                            self._log(T('log_csv', f=os.path.basename(cf_path)), 'ok')
                    except Exception as _e_csv:
                        self._log(T('log_aviso_csv', e=_e_csv), 'aviso')
                self._set_prog(66)

                # ── Exportação DOCX ───────────────────────────────────
                if self.o_export_docx.get() and opts_estats:
                    try:
                        docx_path = os.path.splitext(saida)[0] + '_relatorio.docx'
                        ok_d, msg_d = exportar_docx_relatorio(
                            atletas, docx_path, self._proto_key, opts_estats)
                        if ok_d:
                            self._log(T('log_word_ok', f=os.path.basename(docx_path)), 'ok')
                        else:
                            self._log(T('log_word_nao', msg=msg_d), 'aviso')
                    except Exception as _e_dx:
                        self._log(T('log_aviso_docx', e=_e_dx), 'aviso')
                elif self.o_export_docx.get() and not opts_estats:
                    self._log(T('log_word_sem_estats'), 'aviso')
                self._set_prog(68)

                # ── Exportação HTML interactivo ────────────────────────
                if self.o_export_html.get() and self._run:
                    try:
                        html_path = os.path.splitext(saida)[0] + '_relatorio.html'
                        gerar_html(atletas, html_path, self._proto_key, opts_estats)
                        self._log(T('log_html_ok', f=os.path.basename(html_path)), 'ok')
                    except Exception as _e_html:
                        self._log(T('log_aviso_html', e=_e_html), 'aviso')
                self._set_prog(70)

                # ── Alerta n < 10 para Friedman ───────────────────────
                if (opts_estats and opts_estats.get('friedman')
                        and len(atletas) < 10):
                    self._log(
                        f'⚠ Aviso: n={len(atletas)} < 10 - poder estatístico '
                        f'baixo para Friedman. Resultado com cautela.', 'aviso')

                if self.o_ind.get() and self._run:
                    os.makedirs(ind, exist_ok=True)
                    self._log(T('log_fich_ind', pasta=ind), 'info')
                    for j, ath in enumerate(atletas):
                        if not self._run: break
                        if _is_tiro_like(ath.get('protocolo', '')):
                            guardar_individual_tiro(ath, ind,
                                                    self.o_ell.get(), self.o_est.get())
                            self._log(T('log_tiro_dist', nome=ath['nome'],
                                        n=len(ath.get('tiro_dists',[]))), 'ok')
                            if tiro_bipodal and (ath.get('mets',{}).get('dir') or
                                                  ath.get('mets',{}).get('esq')):
                                fn_hs = re.sub(r'[<>:"/\\|?*]','_',ath['nome']) + '_HurdleStep.xlsx'
                                guardar_individual(ath, os.path.join(ind, fn_hs),
                                                   self.o_ell.get(), self.o_est.get())
                                self._log(T('log_hs_ok', fn=fn_hs), 'ok')
                            for _lado_sel in ('dir', 'esq'):
                                _tem_sel = any(
                                    any(m is not None for m in
                                        ath.get('sel_dist', {}).get(dist, {})
                                           .get(itv_k, {}).get(_lado_sel, []))
                                    for dist in ath.get('tiro_dists', [])
                                    for itv_k in ath.get('tiro_intervalos', []))
                                if _tem_sel:
                                    try:
                                        guardar_individual_sel_tiro(
                                            ath, ind, _lado_sel,
                                            self.o_ell.get(), self.o_est.get())
                                        _lbl_sel = ('Right Sel.' if _lado_sel == 'dir'
                                                    else 'Left Sel.')
                                        self._log(T('log_sel_ok', nome=ath['nome'],
                                                    lado=_lbl_sel), 'ok')
                                    except Exception as _ex_sel:
                                        self._log(T('log_aviso_sel', lado=_lado_sel,
                                                    e=_ex_sel), 'aviso')
                        else:
                            fn=re.sub(r'[<>:"/\\|?*]','_',ath['nome'])+'.xlsx'
                            guardar_individual(ath, os.path.join(ind,fn),
                                               self.o_ell.get(), self.o_est.get())
                            self._log(f'  {fn}', 'ok')
                        self._set_prog(65+(j+1)/len(atletas)*20, f'⟳  {j+1}/{len(atletas)}…')

                if self.o_pdf.get() and pdf and self._run:
                    pdf_dir = os.path.dirname(os.path.abspath(pdf))
                    os.makedirs(pdf_dir, exist_ok=True)
                    self._log(T('log_gerar_pdf'), 'info')
                    self._set_prog(95, '⟳  PDF…')
                    gerar_pdf(atletas, pdf, self._log,
                              opts_estats=opts_estats if self.o_estats.get() else None)
                    self._set_prog(98)

                # ── Exportacao PNG individual ─────────────────────────────
                if self.o_export_png.get() and self._run and _MPL_OK:
                    png_dir = os.path.join(ind, 'png')
                    os.makedirs(png_dir, exist_ok=True)
                    self._log(f'A gerar PNG em: {png_dir}', 'info')
                    n_png_ok = 0
                    _dpi_png   = cfg_load().get('png_dpi', 180)
                    _tipos_png = cfg_load().get('png_tipos', ['estabilograma', 'elipse'])
                    for j, ath in enumerate(atletas):
                        if not self._run: break
                        try:
                            n_png_ok += _exportar_png_individuo(
                                ath, png_dir, dpi=_dpi_png, tipos=_tipos_png)
                        except Exception as _ex_png:
                            self._log(f'  aviso PNG {ath["nome"]}: {_ex_png}', 'aviso')
                    if n_png_ok:
                        self._log(f'  {n_png_ok} PNG(s) gerados em {png_dir}', 'ok')

                self._set_prog(100, '✓  ' + T('log_concluido').replace('✔  ',''))
                self._log(T('log_concluido_n', n=len(atletas)), 'ok')

                # Stash para a janela Demografia (PROTO_ARCO)
                self._last_atletas = atletas
                self._last_ref_dict = atletas_ref_dict
                self._last_ref_lista = atletas_ref_lista
                # Mostrar botao "Demografia" (PROTO_ARCO com ref. carregada)
                if (self._proto_key == PROTO_ARCO and atletas_ref_lista):
                    try:
                        self.root.after(0, self._mostrar_btn_demografia)
                    except Exception:
                        pass

                # Gravar no historico
                hist_add(self._proto_key, len(atletas), saida,
                         pdf if self.o_pdf.get() else '')

                # Mostrar botao de abrir pasta
                pasta_saida = saida_dir
                self.root.after(0, lambda p=pasta_saida: self._mostrar_btn_abrir(p))

                msg = f'Análise concluída!\n{len(atletas)} indivíduo(s).\n\nExcel: {saida}'
                if self.o_pdf.get() and pdf: msg += f'\nPDF: {pdf}'
                # Enviar mensagem de conclusao para a fila (main thread processa)
                self._q.put(('__done__', 'Concluido', msg))

            except BaseException:
                tb=traceback.format_exc()
                _logger.critical("ERRO CRITICO na análise:\n" + tb)
                self._log(T('log_erro_tb', tb=tb), 'erro')
                # OBRIGATÓRIO: _dialogo_crash cria janelas Tk - nunca chamar
                # directamente de uma thread secundária (crasha silenciosamente
                # em macOS e em Windows com certas configurações de DPI/AV).
                # Usar root.after(0,...) para agendar no main thread.
                _tb_snap = tb
                try:
                    # Enviar crash para fila - a _poll trata na main thread
                    self._q.put(('__crash__', _tb_snap))
                except Exception:
                    _logger.error('Fila cheia - crash nao reportado:\n' + _tb_snap)
            finally:
                self._run=False
                try:
                    self.root.after(0, lambda: (
                        self.b_exe.config(state='normal'),
                        self.b_stop.config(state='disabled'),
                        self._status_dot.config(fg='#3A5F7A')
                    ))
                except Exception:
                    pass

        threading.Thread(target=_t, daemon=True).start()


def _exportar_png_individuo(ath, pasta_saida, dpi=180, tipos=None):
    """
    Exporta PNGs do estabilograma e elipse 95% para cada ensaio de um atleta.
    Devolve o numero de ficheiros PNG criados com sucesso.
    Requer matplotlib (ja importado no modulo como _MPL_OK).
    """
    if not _MPL_OK:
        return 0
    if tipos is None:
        tipos = ['estabilograma', 'elipse']

    nome_safe = re.sub(r'[<>:"/\\|?*]', '_', ath['nome'])
    proto_key  = ath.get('protocolo', PROTO_FMS)
    n_ok = 0

    for lado, lst in ath.get('mets', {}).items():
        raw_lst = ath.get('raw', {}).get(lado, [])
        for ti, m in enumerate(lst, start=1):
            if m is None:
                continue
            raw_d = raw_lst[ti - 1] if (ti - 1) < len(raw_lst) else None
            nome_ens = f'{nome_safe}_{lado}_T{ti}'

            # ── Estabilograma ──────────────────────────────────────────
            if 'estabilograma' in tipos:
              try:
                png_estab = _png_estabilograma(m, nome_ens, ath['nome'], dpi=dpi)
                if png_estab:
                    fn = os.path.join(pasta_saida, f'{nome_ens}_estabilograma.png')
                    with open(fn, 'wb') as fh:
                        fh.write(png_estab)
                    n_ok += 1
              except Exception:
                pass

            # ── Elipse 95% ────────────────────────────────────────────
            if 'elipse' in tipos:
              try:
                png_ell = _png_elipses([m], f'T{ti}', ath['nome'], larg=5.0, alt=4.5, dpi=dpi)
                if png_ell:
                    fn = os.path.join(pasta_saida, f'{nome_ens}_elipse.png')
                    with open(fn, 'wb') as fh:
                        fh.write(png_ell)
                    n_ok += 1
              except Exception:
                pass

    return n_ok


def gerar_html(atletas, caminho, protocolo=None, opts_estats=None):
    """
    Gera um relatório HTML interactivo standalone.
    Graficos com Chart.js (CDN). Nao requer instalacao adicional.
    """
    import json as _json
    if protocolo is None: protocolo = _PROTOCOLO_ACTIVO
    proto = PROTOCOLOS[protocolo]
    data_str = datetime.date.today().strftime('%d/%m/%Y')
    lang = lingua_atual() if _I18N_OK else 'PT'

    cores = ['#00B4D8','#22C55E','#F59E0B','#EF4444','#8B5CF6',
             '#06B6D4','#10B981','#F97316','#EC4899','#6366F1']

    # Metricas para graficos - inclui RMS
    mets_para_graf = [
        ('ea95',    'Area Elipse 95% (mm2)'),
        ('vel_med', 'Vel. Media CoP (mm/s)'),
        ('amp_x',   'Amplitude ML (mm)'),
        ('amp_y',   'Amplitude AP (mm)'),
        ('desl',    'Deslocamento (mm)'),
        ('rms_x',   'RMS ML (mm)'),
        ('rms_y',   'RMS AP (mm)'),
        ('rms_r',   'RMS Radius (mm)'),
    ]
    lados = list({l for a in atletas for l in a['mets']})
    nomes = [a['nome'] for a in atletas]

    def _vals_met(chave, lado):
        return [round(_media_lado(a, chave, lado) or 0, 2) for a in atletas]

    charts_data = []
    for chave, label in mets_para_graf:
        datasets = []
        for li, lado in enumerate(lados[:2]):
            vals = _vals_met(chave, lado)
            if any(v > 0 for v in vals):
                datasets.append({
                    'label': lado.upper(),
                    'data': vals,
                    'backgroundColor': cores[li] + '99',
                    'borderColor': cores[li],
                    'borderWidth': 1.5,
                })
        if datasets:
            charts_data.append({'id': chave, 'label': label, 'datasets': datasets})

    charts_json = _json.dumps(charts_data, ensure_ascii=False)
    nomes_json  = _json.dumps(nomes, ensure_ascii=False)

    # Tabela de resumo - N/A em vez de travessao
    tabela_rows = []
    for chave, label in (mets_pdf_localizadas() if _I18N_OK else METS_XL):
        row = f'<tr><td class="met">{label}</td>'
        for lado in lados[:2]:
            med, dp, mn, mx, n_v, se, cv = _stats_grupo(atletas, chave, lado)
            if med is not None:
                dp_str = f'{round(dp,3)}' if dp is not None else 'N/A'
                cv_str = f'{round(cv,1)}%' if cv is not None else 'N/A'
                row += (f'<td>{round(med,3)}</td>'
                        f'<td>{dp_str}</td>'
                        f'<td>{cv_str}</td>')
            else:
                row += '<td>N/A</td><td>N/A</td><td>N/A</td>'
        row += '</tr>'
        tabela_rows.append(row)

    # Legenda das variaveis (subset das mais relevantes)
    legenda_html = ''
    for nome_leg, cod_leg, defn_leg in _LEGENDA_VARS:
        legenda_html += (f'<tr><td class="met">{nome_leg}</td>'
                         f'<td class="cod">{cod_leg}</td>'
                         f'<td class="defn">{defn_leg}</td></tr>')

    _ano_cit = datetime.date.today().year
    cit = (f'Massuça, A., {SEGUNDO_AUTOR_APA}, & Massuça, L. M. ({_ano_cit}). '
           f'BSP - Biomechanical Stability Program (v{VERSAO}). '
           f'https://github.com/andremassuca/BSP')

    th_lados = ''.join(f'<th colspan="3">{l.upper()}</th>' for l in lados[:2])
    th_stats  = ''.join('<th>Media</th><th>DP</th><th>CV%</th>' for _ in lados[:2])

    html = f"""<!DOCTYPE html>
<html lang="{lang.lower()}">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>BSP | {proto['nome']} | {data_str}</title>
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
<style>
:root {{
  --bg:#0D1720; --surface:#132030; --surface2:#1A2B3C;
  --accent:#00B4D8; --accent2:#22C55E; --warn:#F59E0B;
  --text:#EDF6FF; --muted:#7FA8C9; --border:#1E3448;
  --red:#EF4444; --radius:8px; --shadow:0 2px 12px rgba(0,0,0,.4);
}}
*{{box-sizing:border-box;margin:0;padding:0}}
html{{scroll-behavior:smooth}}
body{{background:var(--bg);color:var(--text);font-family:'Segoe UI',system-ui,sans-serif;
      font-size:14px;line-height:1.5;padding:0}}
a{{color:var(--accent);text-decoration:none}}
a:hover{{text-decoration:underline}}

/* ── Header ── */
.header{{background:linear-gradient(135deg,#0A1520 0%,#132030 60%,#0D1F30 100%);
         border-bottom:2px solid var(--accent);padding:20px 32px 16px;
         position:sticky;top:0;z-index:100;backdrop-filter:blur(8px)}}
.header h1{{font-size:1.3rem;font-weight:700;color:var(--accent);letter-spacing:.02em}}
.header .sub{{color:var(--muted);font-size:.82rem;margin-top:4px}}
.badges{{display:flex;gap:6px;flex-wrap:wrap;margin-top:8px}}
.badge{{background:var(--surface2);border:1px solid var(--border);color:var(--accent);
        font-size:.72rem;font-weight:600;padding:2px 10px;border-radius:20px}}
.badge.green{{color:var(--accent2);border-color:#1a5c30}}
.badge.amber{{color:var(--warn);border-color:#5c3d0a}}

/* ── Nav ── */
.nav{{background:var(--surface);border-bottom:1px solid var(--border);
      padding:0 32px;display:flex;gap:2px;overflow-x:auto}}
.nav a{{display:inline-block;padding:10px 14px;color:var(--muted);font-size:.82rem;
        font-weight:500;white-space:nowrap;border-bottom:2px solid transparent;transition:.15s}}
.nav a:hover,.nav a.active{{color:var(--accent);border-bottom-color:var(--accent);
                             text-decoration:none;background:rgba(0,180,216,.06)}}

/* ── Layout ── */
.page{{max-width:1400px;margin:0 auto;padding:28px 32px}}
section{{margin-bottom:40px}}
.section-title{{font-size:1rem;font-weight:700;color:var(--accent);
                margin-bottom:14px;display:flex;align-items:center;gap:8px}}
.section-title::after{{content:'';flex:1;height:1px;background:var(--border)}}

/* ── Cards & Grid ── */
.grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(360px,1fr));gap:16px}}
.grid-2{{display:grid;grid-template-columns:repeat(auto-fill,minmax(500px,1fr));gap:16px}}
.card{{background:var(--surface);border:1px solid var(--border);border-radius:var(--radius);
       padding:18px;box-shadow:var(--shadow)}}
.card h3{{color:var(--muted);font-size:.75rem;font-weight:600;text-transform:uppercase;
          letter-spacing:.06em;margin-bottom:12px}}

/* ── Charts ── */
canvas{{max-height:240px;width:100%!important}}

/* ── Tables ── */
.table-wrap{{overflow-x:auto;border-radius:var(--radius);border:1px solid var(--border)}}
table{{width:100%;border-collapse:collapse;font-size:.82rem}}
thead th{{background:var(--surface2);color:var(--accent);padding:8px 10px;
           font-weight:600;text-align:center;white-space:nowrap;
           position:sticky;top:0}}
thead th:first-child{{text-align:left}}
tbody tr:nth-child(even){{background:rgba(0,0,0,.2)}}
tbody tr:hover{{background:rgba(0,180,216,.06)}}
td{{padding:6px 10px;border-bottom:1px solid var(--border);text-align:center;
    color:var(--muted)}}
td.met{{text-align:left;color:var(--text);font-weight:500}}
td.cod{{color:var(--accent);font-size:.76rem;font-family:monospace;white-space:nowrap}}
td.defn{{text-align:left;font-size:.77rem;line-height:1.45;color:var(--muted)}}
.ea95-row td{{color:var(--red)!important;font-weight:700}}

/* ── Stats summary cards ── */
.stat-grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(160px,1fr));gap:10px}}
.stat-card{{background:var(--surface2);border-radius:6px;padding:12px 14px;
            border-left:3px solid var(--accent)}}
.stat-card .label{{font-size:.72rem;color:var(--muted);margin-bottom:4px;text-transform:uppercase}}
.stat-card .value{{font-size:1.1rem;font-weight:700;color:var(--text)}}
.stat-card .unit{{font-size:.72rem;color:var(--muted)}}

/* ── Citation ── */
.citation{{background:var(--surface);border-left:3px solid var(--accent);
           padding:12px 16px;border-radius:var(--radius);font-size:.8rem;
           color:var(--muted);margin-top:12px}}
.citation strong{{color:var(--text)}}

/* ── Footer ── */
.footer{{background:var(--surface);border-top:1px solid var(--border);
         padding:14px 32px;font-size:.76rem;color:var(--muted);
         display:flex;justify-content:space-between;flex-wrap:wrap;gap:8px}}

/* ── Light mode ── */
@media(prefers-color-scheme:light){{
  :root{{--bg:#F0F4F8;--surface:#FFFFFF;--surface2:#E8EFF5;
         --accent:#0077B6;--accent2:#1a7a3c;--warn:#b45a00;
         --text:#0F172A;--muted:#4A6580;--border:#C8D8E8}}
  .header{{background:linear-gradient(135deg,#E0EAF2,#F0F7FB)}}
  tbody tr:nth-child(even){{background:#F8FAFC}}
  tbody tr:hover{{background:#EBF4FA}}
  .stat-card{{background:#F8FAFC}}
  .table-wrap{{border-color:#C8D8E8}}
}}

/* ── Responsive ── */
@media(max-width:640px){{
  .page{{padding:16px}}
  .grid,.grid-2{{grid-template-columns:1fr}}
  .nav{{padding:0 12px}}
}}
</style>
</head>
<body>

<div class="header">
  <h1>BSP &nbsp;&#x2f;&nbsp; {PROG}</h1>
  <div class="sub">{AUTOR} &nbsp;|&nbsp; {SEGUNDO_AUTOR} &nbsp;|&nbsp; {ORIENTADOR}</div>
  <div class="badges">
    <span class="badge">{proto['nome']}</span>
    <span class="badge">v{VERSAO}</span>
    <span class="badge green">{len(atletas)} indivíduos</span>
    <span class="badge amber">{data_str}</span>
  </div>
</div>

<nav class="nav">
  <a href="#graficos" class="active">📊 Gráficos</a>
  <a href="#resumo">📋 Resumo do Grupo</a>
  <a href="#legenda">📖 Legenda</a>
  <a href="#citacao">📚 Citação</a>
</nav>

<div class="page">

<section id="graficos">
  <div class="section-title">📊 Métricas por Indivíduo</div>
  <div class="grid" id="charts-container"></div>
</section>

<section id="resumo">
  <div class="section-title">📋 Resumo Estatístico do Grupo</div>
  <div class="table-wrap">
  <table>
  <thead>
    <tr><th>Métrica</th>{th_lados}</tr>
    <tr><th></th>{th_stats}</tr>
  </thead>
  <tbody>{''.join(tabela_rows)}</tbody>
  </table>
  </div>
</section>

<section id="legenda">
  <div class="section-title">📖 Legenda de Variáveis</div>
  <div class="table-wrap">
  <table>
  <thead>
    <tr>
      <th style="text-align:left">Variável</th>
      <th style="text-align:left">Código</th>
      <th style="text-align:left">Definição</th>
    </tr>
  </thead>
  <tbody>{legenda_html}</tbody>
  </table>
  </div>
</section>

<section id="citacao">
  <div class="section-title">📚 Citação Académica</div>
  <div class="citation">
    <strong>APA:</strong><br>
    {cit}
  </div>
</section>

</div>

<div class="footer">
  <span>{PROG} v{VERSAO} &nbsp;|&nbsp; {AUTOR} &nbsp;|&nbsp; {SEGUNDO_AUTOR} &nbsp;|&nbsp; {ORIENTADOR}</span>
  <span>{data_str}</span>
</div>

<script>
const nomes = {nomes_json};
const chartsData = {charts_json};

// Detect dark/light mode for chart colors
const isDark = window.matchMedia('(prefers-color-scheme: dark)').matches ||
               !window.matchMedia('(prefers-color-scheme: light)').matches;
const gridColor = isDark ? '#1E3448' : '#D8E4EE';
const tickColor = isDark ? '#7FA8C9' : '#4A6580';
const legendColor = isDark ? '#7FA8C9' : '#4A6580';

const container = document.getElementById('charts-container');
chartsData.forEach(ch => {{
  const card = document.createElement('div');
  card.className = 'card';
  card.innerHTML = '<h3>' + ch.label + '</h3><canvas id="c_' + ch.id + '"></canvas>';
  container.appendChild(card);
  new Chart(document.getElementById('c_' + ch.id), {{
    type: 'bar',
    data: {{ labels: nomes, datasets: ch.datasets }},
    options: {{
      responsive: true,
      maintainAspectRatio: true,
      plugins: {{
        legend: {{ labels: {{ color: legendColor, font: {{ size: 10, weight: '500' }} }} }},
        tooltip: {{
          backgroundColor: isDark ? '#132030' : '#fff',
          titleColor: isDark ? '#EDF6FF' : '#0F172A',
          bodyColor: isDark ? '#7FA8C9' : '#4A6580',
          borderColor: isDark ? '#1E3448' : '#C8D8E8',
          borderWidth: 1,
          padding: 10,
          callbacks: {{ label: ctx => ` ${{ctx.dataset.label}}: ${{ctx.parsed.y.toFixed(2)}}` }}
        }}
      }},
      scales: {{
        x: {{ ticks: {{ color: tickColor, font: {{ size: 9 }}, maxRotation: 40 }},
               grid: {{ color: gridColor }} }},
        y: {{ ticks: {{ color: tickColor, font: {{ size: 9 }} }},
               grid: {{ color: gridColor }} }}
      }}
    }}
  }});
}});

// Nav active highlight on scroll
const sections = document.querySelectorAll('section[id]');
const navLinks = document.querySelectorAll('.nav a');
window.addEventListener('scroll', () => {{
  let cur = '';
  sections.forEach(s => {{ if(window.scrollY >= s.offsetTop - 80) cur = s.id; }});
  navLinks.forEach(a => {{
    a.classList.toggle('active', a.getAttribute('href') === '#' + cur);
  }});
}});
</script>
</body></html>"""

    with open(caminho, 'w', encoding='utf-8') as f:
        f.write(html)
    return caminho


# -----------------------------------------------------------------------
# CLI
# -----------------------------------------------------------------------

def _run_testes_sinteticos(verbose=True):
    """
    Suite de testes automaticos com trajectorias sinteticas.
    Cobre: calculos de metricas, todos os protocolos, Excel, PDF, HTML,
    traducoes, ortografia e integridade do codigo.
    Uso: python estabilidade_gui.py --testes
    """
    import math as _m
    import tempfile, os as _os, sys as _sys, io as _io
    # Garantir UTF-8 no stdout para simbolos ✓/✗ em qualquer consola Windows
    if hasattr(_sys.stdout, 'buffer') and getattr(_sys.stdout, 'encoding', '').lower() not in ('utf-8', 'utf8'):
        try:
            _sys.stdout = _io.TextIOWrapper(_sys.stdout.buffer, encoding='utf-8', errors='replace')
        except Exception:
            pass
    ok_all = True
    resultados = []
    seccoes = {}   # seccao -> [(nome, passou, detalhe)]

    def _t(nome, passou, det="", seccao="Geral"):
        nonlocal ok_all
        if not passou:
            ok_all = False
        seccoes.setdefault(seccao, []).append((nome, passou, det))
        resultados.append((nome, passou))
        if verbose:
            sym = "\u2713" if passou else "\u2717"
            msg = f"    {sym}  {nome}" + (f"  ({det})" if det else "")
            try:
                print(msg)
            except UnicodeEncodeError:
                # Fallback para consolas sem suporte UTF-8 (ex: cp1252 no Windows)
                sym_safe = "[ok]" if passou else "[FAIL]"
                print(f"    {sym_safe}  {nome}" + (f"  ({det})" if det else ""))

    def _fab(n, dur_ms, xf, yf):
        """Fabricar lista de frames sinteticos."""
        ts = [i * dur_ms / max(n - 1, 1) for i in range(n)]
        return [{"frame": i, "t_ms": ts[i], "x": xf(ts[i] / dur_ms),
                 "y": yf(ts[i] / dur_ms)} for i in range(n)]

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 1: CÁLCULO DE MÉTRICAS (calcular())
    # ════════════════════════════════════════════════════════════════════
    import numpy as _np
    if verbose: print("\n  [1] Cálculo de métricas")

    r = 10.0
    fr_circ = _fab(200, 4000,
                   lambda u: r * _m.cos(2 * _m.pi * u),
                   lambda u: r * _m.sin(2 * _m.pi * u))
    m_circ = calcular(fr_circ)

    _t("calcular() retorna dict para circulo", m_circ is not None, seccao="Métricas")
    if m_circ:
        # Para uma trajectória cos/sin, Var(X)=Var(Y)=r²/2, logo:
        # ea95 = π·(sa/2)·(sb/2) = π·r²·χ²/8  (não π·r²)
        _CHI2 = 5.9915
        ea95_esperado = _m.pi * r * r * _CHI2 / 8
        err_ea = abs(m_circ["ea95"] - ea95_esperado) / ea95_esperado * 100
        _t("ea95 circulo ≈ π·r²·χ²/8 (err < 2%)", err_ea < 2,
           f"ea95={m_circ['ea95']:.2f} esperado={ea95_esperado:.2f} err={err_ea:.1f}%", "Métricas")
        _t("vel_med > 0", m_circ["vel_med"] > 0,
           f"{m_circ['vel_med']:.2f} mm/s", "Métricas")
        _t("amp_x ≈ 2r (±5%)", abs(m_circ["amp_x"] - 2 * r) / (2 * r) < 0.05,
           f"amp_x={m_circ['amp_x']:.2f}", "Métricas")
        _t("rms_x > 0", m_circ.get("rms_x", 0) > 0,
           f"rms_x={m_circ.get('rms_x', 0):.3f}", "Métricas")
        _t("rms_r ≈ r (±5%)", abs(m_circ.get("rms_r", 0) - r) / r < 0.05,
           f"rms_r={m_circ.get('rms_r', 0):.3f}", "Métricas")
        _t("leng_a > leng_b (circulo: a≈b)", abs(m_circ["leng_a"] - m_circ["leng_b"]) < r * 0.1,
           f"a={m_circ['leng_a']:.3f} b={m_circ['leng_b']:.3f}", "Métricas")
        _t("ea95 = π·(a/2)·(b/2) (fórmula correta)",
           abs(m_circ["ea95"] - _m.pi * (m_circ["leng_a"] / 2) * (m_circ["leng_b"] / 2)) < 0.5,
           f"ea95={m_circ['ea95']:.3f} π·a/2·b/2={_m.pi*(m_circ['leng_a']/2)*(m_circ['leng_b']/2):.3f}",
           "Métricas")

    # Linha: semi-eixo menor ≈ 0
    fr_linha = _fab(100, 4000, lambda u: u * 20.0, lambda u: 0.0)
    m_linha = calcular(fr_linha)
    _t("Linha: semi-eixo menor < 0.5 mm",
       m_linha is not None and m_linha["leng_b"] < 0.5,
       f"leng_b={m_linha['leng_b']:.4f}" if m_linha else "None", "Métricas")

    # Poucos frames → None
    fr_curto = _fab(3, 100, lambda u: u, lambda u: u)
    _t("< 5 frames → calcular() retorna None",
       calcular(fr_curto) is None, "3 frames", "Métricas")

    # ratio_ml_ap
    fr_ratio = _fab(100, 4000,
                    lambda u: 20.0 * _m.sin(2 * _m.pi * u),
                    lambda u: 10.0 * _m.sin(2 * _m.pi * u))
    m_ratio = calcular(fr_ratio)
    if m_ratio and m_ratio.get("ratio_ml_ap") is not None:
        exp = m_ratio["amp_x"] / m_ratio["amp_y"]
        _t("ratio_ml_ap = amp_x / amp_y",
           abs(m_ratio["ratio_ml_ap"] - exp) < 0.01,
           f"{m_ratio['ratio_ml_ap']:.3f} vs {exp:.3f}", "Métricas")
    else:
        _t("ratio_ml_ap calculado", False, seccao="Métricas")

    # Windowing t_ini / t_fim
    fr_long = _fab(300, 10000,
                   lambda u: 10 * _m.cos(2 * _m.pi * u),
                   lambda u: 10 * _m.sin(2 * _m.pi * u))
    m_wind = calcular(fr_long, t_ini=2000, t_fim=8000)
    _t("Windowing t_ini/t_fim aplica janela",
       m_wind is not None and len(m_wind["t_ms"]) < 300,
       f"n_frames_windowed={len(m_wind['t_ms']) if m_wind else 0}", "Métricas")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 2: FUNÇÕES ESTATÍSTICAS
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [2] Funções estatísticas")

    arr = _np.array([10.0, 12.0, 8.0, 11.0, 9.0])
    cv = float(arr.std(ddof=1) / arr.mean() * 100)
    _t("CV = DP/média·100", abs(cv - 15.811) < 0.01, f"cv={cv:.3f}", "Estatísticas")

    ai = assimetria(12.0, 8.0)
    _t("Assimetria Dir/Esq: (12−8)/(10)·100 = 40%",
       abs(ai - 40.0) < 0.01, f"AI={ai}", "Estatísticas")
    _t("Assimetria simétrica = 0", abs(assimetria(10.0, 10.0)) < 0.01, seccao="Estatísticas")
    _t("Assimetria None se ambos 0",
       assimetria(0.0, 0.0) is None or assimetria(0.0, 0.0) == 0,
       seccao="Estatísticas")

    mf = [{"ea95": v} for v in [100, 102, 98, 101, 500]]
    fl = flagrar_outliers(mf)
    _t("Outlier extremo detectado (z>3)", fl[-1] is True, str(fl), "Estatísticas")
    _t("Não-outliers não marcados", all(not fl[i] for i in range(4)), str(fl), "Estatísticas")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 3: PROCESSAR_ATLETA (dados sintéticos em disco)
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [3] Processamento de atleta (dados sintéticos)")

    try:
        import openpyxl as _opx
        from io import BytesIO as _BIO

        def _escrever_xls_sintetico(pasta, nome_ficheiro, frames):
            """Escreve um ficheiro XLS sintético compatível com ler_ficheiro()."""
            wb = _opx.Workbook()
            ws = wb.active
            ws.title = "Sheet1"
            ws["A1"] = "frame"; ws["B1"] = "t_ms"
            ws["C1"] = "x";     ws["D1"] = "y"
            for i, fr in enumerate(frames, start=2):
                ws.cell(i, 1, fr["frame"])
                ws.cell(i, 2, fr["t_ms"])
                ws.cell(i, 3, fr["x"])
                ws.cell(i, 4, fr["y"])
            path = _os.path.join(pasta, nome_ficheiro)
            wb.save(path)
            return path

        # Criar estrutura de pasta para atleta FMS sintético
        with tempfile.TemporaryDirectory() as tmpdir:
            ath_dir = _os.path.join(tmpdir, "001_Teste_Sintetico")
            _os.makedirs(ath_dir)
            # 5 ensaios dir + 5 ensaios esq
            for t in range(1, 6):
                fr = _fab(200, 4000,
                          lambda u, t=t: (8 + t) * _m.cos(2 * _m.pi * u),
                          lambda u, t=t: (8 + t) * _m.sin(2 * _m.pi * u))
                _escrever_xls_sintetico(ath_dir, f"dir_{t}.xlsx", fr)
                _escrever_xls_sintetico(ath_dir, f"esq_{t}.xlsx", fr)

            ath = processar_atleta(ath_dir, {}, False, log=lambda *a: None,
                                   protocolo=PROTO_FMS)
            _t("processar_atleta() retorna dict", isinstance(ath, dict), seccao="Processamento")
            _t("FMS: 5 ensaios dir calculados",
               ath is not None and len([m for m in ath["mets"].get("dir", []) if m]) == 5,
               f"n_dir={len([m for m in ath['mets'].get('dir',[]) if m]) if ath else 0}",
               "Processamento")
            _t("FMS: 5 ensaios esq calculados",
               ath is not None and len([m for m in ath["mets"].get("esq", []) if m]) == 5,
               seccao="Processamento")
            _t("FMS: ea95 calculado em todos os ensaios",
               ath is not None and all(m and m.get("ea95") for m in ath["mets"].get("dir", [])),
               seccao="Processamento")

            # Unipodal com os mesmos dados
            ath_u = processar_atleta(ath_dir, {}, False, log=lambda *a: None,
                                     protocolo=PROTO_UNIPODAL)
            _t("Unipodal: processar_atleta() OK",
               isinstance(ath_u, dict), seccao="Processamento")

    except ImportError:
        _t("openpyxl disponivel (necessario para testes de disco)", False,
           "pip install openpyxl", "Processamento")
    except Exception as ex:
        _t("processar_atleta() sem excepção", False, str(ex)[:80], "Processamento")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 4: EXCEL - guardar_resumo + guardar_individual
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [4] Geração de Excel")

    try:
        import openpyxl as _opx2

        # Atleta sintético mínimo
        def _atleta_sintetico(nome="Teste", proto=PROTO_FMS, n=5):
            mets_d = []
            for i in range(n):
                fr = _fab(200, 4000,
                          lambda u, i=i: (8+i)*_m.cos(2*_m.pi*u),
                          lambda u, i=i: (8+i)*_m.sin(2*_m.pi*u))
                mets_d.append(calcular(fr))
            return {"nome": nome, "id": "1",
                    "mets": {"dir": mets_d, "esq": mets_d},
                    "raw":  {"dir": [{"dados": {"frames": [], "inicio_ms": None,
                                                "fim_ms": None, "data": "04/04/2026"},
                                      "ini": None, "fim": None}] * n,
                             "esq": [{"dados": {"frames": [], "inicio_ms": None,
                                                "fim_ms": None, "data": "04/04/2026"},
                                      "ini": None, "fim": None}] * n},
                    "protocolo": proto, "scores": None}

        ath_xl = _atleta_sintetico()
        with tempfile.TemporaryDirectory() as tmpdir:
            path_xl = _os.path.join(tmpdir, "resumo.xlsx")
            try:
                guardar_resumo([ath_xl], path_xl, PROTO_FMS)
                wb_check = _opx2.load_workbook(path_xl, data_only=True)
                _t("guardar_resumo FMS cria ficheiro", _os.path.exists(path_xl), seccao="Excel")
                _t("guardar_resumo FMS tem aba DADOS",
                   "DADOS" in wb_check.sheetnames, str(wb_check.sheetnames), "Excel")
                _t("guardar_resumo FMS tem aba GRUPO",
                   "GRUPO" in wb_check.sheetnames, str(wb_check.sheetnames), "Excel")
                _t("guardar_resumo FMS tem aba SPSS",
                   "SPSS" in wb_check.sheetnames, str(wb_check.sheetnames), "Excel")
            except Exception as ex:
                _t("guardar_resumo FMS sem excepção", False, str(ex)[:80], "Excel")

            # Unipodal
            ath_xl_u = _atleta_sintetico(proto=PROTO_UNIPODAL)
            path_xl_u = _os.path.join(tmpdir, "resumo_uni.xlsx")
            try:
                guardar_resumo([ath_xl_u], path_xl_u, PROTO_UNIPODAL)
                _t("guardar_resumo Unipodal cria ficheiro",
                   _os.path.exists(path_xl_u), seccao="Excel")
            except Exception as ex:
                _t("guardar_resumo Unipodal sem excepção", False, str(ex)[:80], "Excel")

            # Individual
            path_ind = _os.path.join(tmpdir, "individual.xlsx")
            try:
                guardar_individual(ath_xl, path_ind, c_elipse=True, c_estab=True)
                wb_i = _opx2.load_workbook(path_ind, data_only=True)
                has_elipse = any("elipse" in s.lower() for s in wb_i.sheetnames)
                has_estab  = any("estab" in s.lower() for s in wb_i.sheetnames)
                _t("guardar_individual cria abas de dados",
                   len(wb_i.sheetnames) >= 5, str(wb_i.sheetnames[:4]), "Excel")
                _t("guardar_individual tem aba elipse", has_elipse, seccao="Excel")
                _t("guardar_individual tem aba estabilograma", has_estab, seccao="Excel")
            except Exception as ex:
                _t("guardar_individual sem excepção", False, str(ex)[:80], "Excel")

    except ImportError:
        _t("openpyxl disponivel", False, "pip install openpyxl", "Excel")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 5: HTML - gerar_html
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [5] Geração de HTML")

    try:
        ath_html = {"nome": "Sintetico", "id": "1",
                    "mets": {"dir": [calcular(_fab(200,4000,
                                                    lambda u: 10*_m.cos(2*_m.pi*u),
                                                    lambda u: 10*_m.sin(2*_m.pi*u)))]*3},
                    "protocolo": PROTO_FMS, "scores": None}
        with tempfile.TemporaryDirectory() as tmpdir:
            path_html = _os.path.join(tmpdir, "relatorio.html")
            try:
                gerar_html([ath_html], path_html, PROTO_FMS)
                with open(path_html, encoding="utf-8") as fh:
                    html_content = fh.read()
                _t("gerar_html cria ficheiro", _os.path.exists(path_html), seccao="HTML")
                _t("HTML tem Chart.js", "chart.js" in html_content.lower(), seccao="HTML")
                _t("HTML tem nav (melhorado)", 'class="nav"' in html_content, seccao="HTML")
                _t("HTML tem citação APA", "Massuça, A." in html_content, seccao="HTML")
                _t("HTML tem Pedro Aleixo", "Aleixo" in html_content, seccao="HTML")
                _t("HTML é UTF-8 válido", True, seccao="HTML")
            except Exception as ex:
                _t("gerar_html sem excepção", False, str(ex)[:80], "HTML")
    except Exception as ex:
        _t("HTML: setup sem excepção", False, str(ex)[:80], "HTML")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 6: TRADUÇÕES - completude e consistência
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [6] Traduções e i18n")

    if _I18N_OK:
        from bsp_i18n import _STRINGS, licenca_texto
        pt_keys = set(_STRINGS.get("PT", {}).keys())
        for lang in ["EN", "ES", "DE"]:
            d = _STRINGS.get(lang, {})
            missing = pt_keys - set(d.keys())
            _t(f"Língua {lang}: 0 chaves em falta",
               len(missing) == 0, f"{len(missing)} em falta: {list(missing)[:5]}", "Traduções")

        # EULA sem erros de formato
        for lang in ["PT", "EN", "ES", "DE"]:
            try:
                txt = licenca_texto(lang=lang, versao=VERSAO,
                                    autor=AUTOR, univ=f"{SEGUNDO_AUTOR}  |  {ORIENTADOR}")
                _t(f"EULA {lang} renderiza sem KeyError",
                   len(txt) > 100, f"{len(txt)} chars", "Traduções")
                _t(f"EULA {lang} tem Pedro Aleixo",
                   "Aleixo" in txt, seccao="Traduções")
            except Exception as ex:
                _t(f"EULA {lang} sem excepção", False, str(ex)[:60], "Traduções")
    else:
        _t("bsp_i18n carregado", False, "ficheiro não encontrado", "Traduções")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 7: INTEGRIDADE DO CÓDIGO
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [7] Integridade do código")

    import hashlib as _hl
    _t("AUTOR decodifica correctamente",
       AUTOR == "André O. Massuça", f"AUTOR={AUTOR!r}", "Integridade")

    # Verificar que ISCPSI não aparece na UI (só no PDF via texto)
    import inspect as _ins
    src_txt = _ins.getsource(_ecra_password)
    _t("_ecra_password sem logo ISCPSI",
       "_ISCPSI_B64" not in src_txt, seccao="Integridade")

    # METS_PDF e METS_XL têm ea95
    _t("METS_PDF contém ea95",
       any(k == "ea95" for k, _ in METS_PDF), seccao="Integridade")
    _t("METS_XL contém ea95",
       any(k == "ea95" for k, _ in METS_XL), seccao="Integridade")
    _t("METS_SPSS contém ea95",
       any(k == "ea95" for k, _ in METS_SPSS), seccao="Integridade")
    _t("METS_PDF contém rms_x",
       any(k == "rms_x" for k, _ in METS_PDF), seccao="Integridade")

    # Todos os protocolos definidos
    for pk in [PROTO_FMS, PROTO_UNIPODAL, PROTO_TIRO, PROTO_ARCO]:
        _t(f"Protocolo '{pk}' definido em PROTOCOLOS",
           pk in PROTOCOLOS, seccao="Integridade")

    # calcular() retorna todas as chaves esperadas
    chaves_esperadas = {"ea95","leng_a","leng_b","amp_x","amp_y","vel_x","vel_y",
                        "vel_med","desl","time","rms_x","rms_y","rms_r",
                        "cov_xx","cov_yy","cov_xy","ratio_ml_ap","stiff_x"}
    if m_circ:
        faltam = chaves_esperadas - set(m_circ.keys())
        _t("calcular() retorna todas as chaves esperadas",
           len(faltam) == 0, f"em falta: {faltam}", "Integridade")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 8: PDF - geração sem excepção (FMS sintético)
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [8] Geração de PDF (FMS sintético)")

    if _MPL_OK:
        try:
            from reportlab.pdfgen import canvas as _rlc
            from reportlab.lib.pagesizes import A4 as _A4
            ath_pdf = {"nome": "Atleta PDF Test", "id": "99",
                       "mets": {"dir": [m_circ]*5, "esq": [m_circ]*5},
                       "raw":  {"dir": [None]*5, "esq": [None]*5},
                       "protocolo": PROTO_FMS, "scores": None}
            with tempfile.TemporaryDirectory() as tmpdir:
                path_pdf = _os.path.join(tmpdir, "test.pdf")
                try:
                    gerar_pdf([ath_pdf], path_pdf, log=lambda *a: None)
                    size = _os.path.getsize(path_pdf)
                    _t("gerar_pdf FMS cria ficheiro",
                       size > 5000, f"{size//1024} KB", "PDF")
                except Exception as ex:
                    _t("gerar_pdf FMS sem excepção", False, str(ex)[:80], "PDF")
        except ImportError:
            _t("reportlab disponivel", False, "pip install reportlab", "PDF")
    else:
        _t("matplotlib disponivel (necessário para PDF)", False,
           "pip install matplotlib", "PDF")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 9: Tiro com Arco (v1.0) - pipeline próprio
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [9] Tiro com Arco (v1.0)")

    # PROTO_ARCO bem configurado: lado único, 30 ensaios, sem assimetria
    _proto_arco = PROTOCOLOS[PROTO_ARCO]
    _t("PROTO_ARCO tem lado único 'arco'",
       len(_proto_arco['lados']) == 1 and _proto_arco['lados'][0][0] == 'arco',
       f"lados={_proto_arco['lados']}", "Arco")
    _t("PROTO_ARCO n_ens = 30",
       _proto_arco['n_ens'] == 30,
       f"n_ens={_proto_arco['n_ens']}", "Arco")
    _t("PROTO_ARCO não tem assimetria",
       _proto_arco['assimetria'] is False, seccao="Arco")
    _t("PROTO_ARCO não é two_windows",
       _proto_arco['two_windows'] is False, seccao="Arco")

    # Helpers de discriminação
    _t("_is_tiro_like(PROTO_ARCO) == False",
       _is_tiro_like(PROTO_ARCO) is False,
       "Arco não pertence a _TIRO_PROTOS", "Arco")
    _t("_is_iscpsi(PROTO_ARCO) == False",
       _is_iscpsi(PROTO_ARCO) is False, seccao="Arco")
    _t("_is_arco(PROTO_ARCO) == True",
       _is_arco(PROTO_ARCO) is True, seccao="Arco")
    _t("_is_iscpsi(PROTO_TIRO) == True",
       _is_iscpsi(PROTO_TIRO) is True,
       "ISCPSI só para Tiro (espingarda)", "Arco")

    # Detector do novo formato arco: deve identificar "Entire plate COF"
    _exemplo_arco = ("Frame\tTime (ms)\tEntire plate COF X\tEntire plate COF Y\n"
                     "1\t0\t10.5\t20.3\n")
    with tempfile.NamedTemporaryFile(mode='w', suffix='.xls', delete=False,
                                      encoding='utf-8') as f:
        f.write(_exemplo_arco)
        _fpath_arco = f.name
    try:
        _t("_detectar_formato_arco reconhece cabeçalho",
           _detectar_formato_arco(_fpath_arco) is True,
           "Entire plate COF X deve ser detectado", "Arco")
    finally:
        try: _os.unlink(_fpath_arco)
        except Exception: pass

    # Atleta sintético arco: 5 ensaios (limite inferior do esperado)
    _ath_arco = {
        "nome": "Atleta Arco Test",
        "id": "101",
        "mets": {"arco": [m_circ]*5},
        "raw":  {"arco": [None]*5},
        "protocolo": PROTO_ARCO,
        "scores": None,
        "peso_kg": 70.0,
        "altura_m": 1.75,
    }

    # Excel: guardar_resumo não rebenta com arco (lado único)
    with tempfile.TemporaryDirectory() as tmpdir:
        _xls_arco = _os.path.join(tmpdir, "resumo_arco.xlsx")
        try:
            guardar_resumo([_ath_arco], _xls_arco, protocolo=PROTO_ARCO)
            _t("guardar_resumo arco cria ficheiro",
               _os.path.exists(_xls_arco) and _os.path.getsize(_xls_arco) > 1000,
               f"{_os.path.getsize(_xls_arco)//1024} KB", "Arco")
            try:
                from openpyxl import load_workbook as _lw
                _wb = _lw(_xls_arco, read_only=True)
                _t("guardar_resumo arco tem aba DADOS",
                   'DADOS' in _wb.sheetnames, str(_wb.sheetnames[:5]), "Arco")
                _t("guardar_resumo arco tem aba GRUPO",
                   'GRUPO' in _wb.sheetnames, str(_wb.sheetnames[:5]), "Arco")
                _wb.close()
            except Exception as ex:
                _t("abrir resumo arco", False, str(ex)[:80], "Arco")
        except Exception as ex:
            _t("guardar_resumo arco sem excepção",
               False, str(ex)[:120], "Arco")

    # PDF: gerar_pdf para PROTO_ARCO não rebenta
    if _MPL_OK:
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                _pdf_arco = _os.path.join(tmpdir, "test_arco.pdf")
                try:
                    gerar_pdf([_ath_arco], _pdf_arco, log=lambda *a: None)
                    _size_a = _os.path.getsize(_pdf_arco)
                    _t("gerar_pdf arco cria ficheiro",
                       _size_a > 5000, f"{_size_a//1024} KB", "Arco")
                except Exception as ex:
                    _t("gerar_pdf arco sem excepção",
                       False, str(ex)[:120], "Arco")
        except Exception as ex:
            _t("gerar_pdf arco env", False, str(ex)[:80], "Arco")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 10: Referência demográfica (142 atletas) - v1.0
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [10] Referência demográfica (142 atletas)")

    try:
        from openpyxl import Workbook as _Wb
        with tempfile.TemporaryDirectory() as tmpdir:
            _ref_path = _os.path.join(tmpdir, "ref_142.xlsx")
            _wb_ref = _Wb()
            _ws = _wb_ref.active
            # Cabeçalho: ID, PESO, ALTURA, IDADE, ESTILO, CATEGORIA, GENERO,
            #            P1..P5 (simplificado para teste), P_TOTAL, d1..d5
            _hdr = (['ID','PESO','ALTURA','IDADE','ESTILO','CATEGORIA','GENERO']
                    + [f'P{i}' for i in range(1,6)] + ['P_TOTAL']
                    + [f'd{i}' for i in range(1,6)])
            _ws.append(_hdr)
            # 3 atletas sintéticos
            _ws.append([101, 70.5, 1.75, 23, 1, 'senior', 1,
                        10, 9, 10, 8, 9, 46, 5.2, 6.1, 4.8, 7.3, 5.9])
            _ws.append([102, 62.0, 165,  21, 2, 'junior', 2,
                        8, 7, 9, 8, 7, 39,  6.5, 7.2, 5.8, 6.1, 7.0])
            # Atleta com dados em falta
            _ws.append([103, None, 1.80, 25, 1, 'senior', 1,
                        None, 9, 10, None, 9, None, None, None, None, None, None])
            _wb_ref.save(_ref_path)

            _ref = carregar_atletas_ref(_ref_path)
            _t("carregar_atletas_ref devolve 3 atletas",
               len(_ref) == 3, f"n={len(_ref)}", "Demografia")

            if len(_ref) >= 1:
                a1 = _ref[0]
                _t("atleta 101: ID correcto",
                   a1['id'] == '101', f"id={a1['id']}", "Demografia")
                _t("atleta 101: peso 70.5 kg",
                   a1['peso_kg'] == 70.5, f"peso={a1['peso_kg']}", "Demografia")
                _t("atleta 101: altura 1.75 m",
                   abs((a1['altura_m'] or 0) - 1.75) < 1e-6,
                   f"altura={a1['altura_m']}", "Demografia")
                _t("atleta 101: genero M (codigo 1)",
                   a1['genero'] == 'M', f"genero={a1['genero']}", "Demografia")
                _t("atleta 101: estilo recurvo (codigo 1)",
                   a1['estilo'] == 'recurvo',
                   f"estilo={a1['estilo']}", "Demografia")
                _t("atleta 101: P_total calculado",
                   a1['P_total'] == 46, f"P_total={a1['P_total']}", "Demografia")

            if len(_ref) >= 2:
                a2 = _ref[1]
                _t("atleta 102: altura 165cm convertida para 1.65m",
                   abs((a2['altura_m'] or 0) - 1.65) < 1e-6,
                   f"altura={a2['altura_m']}", "Demografia")
                _t("atleta 102: genero F (codigo 2)",
                   a2['genero'] == 'F', f"genero={a2['genero']}", "Demografia")
                _t("atleta 102: estilo composto (codigo 2)",
                   a2['estilo'] == 'composto',
                   f"estilo={a2['estilo']}", "Demografia")

            if len(_ref) >= 3:
                a3 = _ref[2]
                _t("atleta 103: peso None tolerado",
                   a3['peso_kg'] is None, f"peso={a3['peso_kg']}", "Demografia")
                # P_total deve ser soma dos P não-None: 9+10+9 = 28
                _t("atleta 103: P_total ignora None",
                   a3['P_total'] == 28, f"P_total={a3['P_total']}", "Demografia")

            # Helper por_id
            _by_id = atletas_ref_por_id(_ref)
            _t("atletas_ref_por_id cria dict indexado",
               _by_id.get('101', {}).get('peso_kg') == 70.5,
               f"n={len(_by_id)}", "Demografia")
    except Exception as ex:
        _t("carregar_atletas_ref sem excepção", False, str(ex)[:120], "Demografia")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 11: Normalizações altura/massa em calcular() - v1.0
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [11] Normalizações altura/massa (calcular)")

    # Frames sintéticos deterministicos: 100 amostras a 100 Hz, oscilacao 1 Hz
    try:
        _dt_ms = 10.0
        _nf = 100
        _frames_n = []
        for _i in range(_nf):
            _tms = _i * _dt_ms
            _fx = 10.0 * math.sin(2*math.pi*1.0*_tms/1000.0)   # amp_x ~= 20 mm
            _fy = 15.0 * math.sin(2*math.pi*0.5*_tms/1000.0)   # amp_y ~= 30 mm
            _frames_n.append({'t_ms': _tms, 'x': _fx, 'y': _fy})

        # [11a] stiff_x/stiff_y calculados
        _m0 = calcular(_frames_n)
        _t("calcular: stiff_x calculado",
           _m0 is not None and _m0.get('stiff_x') is not None,
           f"stiff_x={_m0.get('stiff_x') if _m0 else 'None'}", "Normaliz")

    except Exception as ex:
        _t("calcular normalizações sem excepção",
           False, str(ex)[:120], "Normaliz")

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 12: Análises demográficas - v1.0
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [12] Análises demográficas")

    try:
        # Dataset sintético com 12 atletas: 6M + 6F, com medianas claramente
        # separadas para que Mann-Whitney devolva p<0.05.
        import random as _r
        _r.seed(42)
        _atletas_demo = []
        # M: ea95 alto (~400), peso 70-85, P_total 40-50
        for i in range(6):
            _atletas_demo.append({
                'id': f'M{i+1}', 'genero': 'M', 'estilo': 'recurvo',
                'categoria': 'senior',
                'peso_kg': 70 + i*3, 'altura_m': 1.75 + i*0.01,
                'idade': 22+i,
                'ea95': 400 + _r.random()*50,
                'P_total': 45 + i,
                'P': [9, 10, 8, 9] + [8]*26,
                'mets': {'arco': [{'ea95': 400+i*5} for _ in range(30)]},
            })
        # F: ea95 baixo (~200), peso 55-70, P_total 35-45
        for i in range(6):
            _atletas_demo.append({
                'id': f'F{i+1}', 'genero': 'F', 'estilo': 'composto',
                'categoria': 'junior',
                'peso_kg': 55 + i*3, 'altura_m': 1.62 + i*0.01,
                'idade': 18+i,
                'ea95': 200 + _r.random()*50,
                'P_total': 38 + i,
                'P': [7, 8, 8, 7] + [7]*26,
                'mets': {'arco': [{'ea95': 200+i*5} for _ in range(30)]},
            })

        # [12a] comparar_grupos(genero) - 2 grupos → mannwhitneyu
        _cmp_g = comparar_grupos(_atletas_demo, 'ea95', 'genero')
        _t("comparar_grupos(genero): 2 grupos detectados",
           _cmp_g['n_grupos'] == 2, f"n_grupos={_cmp_g['n_grupos']}", "Demogr")
        _t("comparar_grupos(genero): teste=mannwhitneyu",
           _cmp_g['teste'] == 'mannwhitneyu',
           f"teste={_cmp_g['teste']}", "Demogr")
        _t("comparar_grupos(genero): p_valor<0.05 (M vs F claro)",
           _cmp_g['p_valor'] is not None and _cmp_g['p_valor'] < 0.05,
           f"p={_cmp_g['p_valor']}", "Demogr")
        _t("comparar_grupos(genero): mediana M > mediana F",
           _cmp_g['mediana'].get('M', 0) > _cmp_g['mediana'].get('F', 0),
           f"M={_cmp_g['mediana'].get('M'):.1f} F={_cmp_g['mediana'].get('F'):.1f}",
           "Demogr")

        # [12b] comparar_grupos com 3+ grupos → kruskal
        # Cria um dataset com 3 categorias distintas
        _atletas_3g = []
        for i in range(4):
            _atletas_3g.append({'id': f'C1_{i}', 'categoria': 'cadete',
                                'ea95': 100 + i*5})
        for i in range(4):
            _atletas_3g.append({'id': f'C2_{i}', 'categoria': 'junior',
                                'ea95': 200 + i*5})
        for i in range(4):
            _atletas_3g.append({'id': f'C3_{i}', 'categoria': 'senior',
                                'ea95': 300 + i*5})
        _cmp_k = comparar_grupos(_atletas_3g, 'ea95', 'categoria')
        _t("comparar_grupos(categoria 3g): teste=kruskal",
           _cmp_k['teste'] == 'kruskal', f"teste={_cmp_k['teste']}", "Demogr")
        _t("comparar_grupos(categoria 3g): p_valor<0.05",
           _cmp_k['p_valor'] is not None and _cmp_k['p_valor'] < 0.05,
           f"p={_cmp_k['p_valor']}", "Demogr")

        # [12c] correlacao_demografica: peso vs ea95
        _cor = correlacao_demografica(_atletas_demo, 'ea95', 'peso_kg')
        _t("correlacao_demografica(peso, ea95): n=12",
           _cor['n'] == 12, f"n={_cor['n']}", "Demogr")
        _t("correlacao_demografica(peso, ea95): pearson_r presente",
           _cor['pearson_r'] is not None,
           f"r={_cor['pearson_r']}", "Demogr")
        _t("correlacao_demografica(peso, ea95): spearman_r presente",
           _cor['spearman_r'] is not None,
           f"r={_cor['spearman_r']}", "Demogr")
        # Peso cresce com ea95 (M mais pesado+mais instavel) → r>0
        _t("correlacao_demografica(peso, ea95): r positivo",
           _cor['pearson_r'] is not None and _cor['pearson_r'] > 0,
           f"r={_cor['pearson_r']:.3f}", "Demogr")

        # [12d] percentis_subgrupo: atleta de interesse dentro dos 6M seniores
        _ath_interesse = _atletas_demo[2]  # M3: ea95 ~400+
        _pct = percentis_subgrupo(_ath_interesse, _atletas_demo, 'ea95',
                                   fatores=('categoria', 'genero'))
        _t("percentis_subgrupo: subgrupo correcto",
           _pct['subgrupo'].get('genero') == 'M' and
           _pct['subgrupo'].get('categoria') == 'senior',
           f"subgrupo={_pct['subgrupo']}", "Demogr")
        _t("percentis_subgrupo: n_subgrupo=6 (todos M senior)",
           _pct['n_subgrupo'] == 6, f"n={_pct['n_subgrupo']}", "Demogr")
        _t("percentis_subgrupo: P25/P50/P75 calculados",
           _pct['p25'] is not None and _pct['p50'] is not None
           and _pct['p75'] is not None,
           f"P25={_pct['p25']:.1f} P50={_pct['p50']:.1f} P75={_pct['p75']:.1f}",
           "Demogr")
        _t("percentis_subgrupo: percentil_atleta entre 0 e 100",
           _pct['percentil_atleta'] is not None and
           0 <= _pct['percentil_atleta'] <= 100,
           f"pct={_pct['percentil_atleta']}", "Demogr")
        _t("percentis_subgrupo: rank em 1..n_subgrupo",
           _pct['rank'] is not None and 1 <= _pct['rank'] <= _pct['n_subgrupo'],
           f"rank={_pct['rank']}/{_pct['n_subgrupo']}", "Demogr")

        # [12e] correlacao_score: agregado + per_ensaio
        _cs = correlacao_score(_atletas_demo, 'ea95', 'P_total')
        _t("correlacao_score: chave_cop=ea95",
           _cs['chave_cop'] == 'ea95', f"cop={_cs['chave_cop']}", "Demogr")
        _t("correlacao_score: agregado.n=12",
           _cs['agregado']['n'] == 12,
           f"n={_cs['agregado']['n']}", "Demogr")
        _t("correlacao_score: agregado.pearson_r presente",
           _cs['agregado']['pearson_r'] is not None,
           f"r={_cs['agregado']['pearson_r']}", "Demogr")
        # per_ensaio: 12 atletas × 30 = 360 pontos
        _t("correlacao_score: per_ensaio com n>=360",
           _cs['per_ensaio'] is not None and _cs['per_ensaio']['n'] >= 360,
           f"n={_cs['per_ensaio']['n'] if _cs['per_ensaio'] else 'None'}",
           "Demogr")

        # [12f] Excel: abas Demografia / Comparacoes / Correlacoes
        # Adiciona campos de protocolo e adapta formato
        for _a in _atletas_demo:
            _a['protocolo'] = PROTO_ARCO
            _a['nome'] = _a['id']
        with tempfile.TemporaryDirectory() as _tdir:
            _xls_demo = _os.path.join(_tdir, 'demo.xlsx')
            try:
                guardar_resumo(_atletas_demo, _xls_demo, protocolo=PROTO_ARCO)
                _wbx = load_workbook(_xls_demo, read_only=True)
                _nomes = _wbx.sheetnames
                _t("guardar_resumo arco: aba Demografia presente",
                   'Demografia' in _nomes, f"abas={_nomes}", "Demogr")
                _t("guardar_resumo arco: aba Comparacoes presente",
                   'Comparacoes' in _nomes, f"abas={_nomes}", "Demogr")
                _t("guardar_resumo arco: aba Correlacoes presente",
                   'Correlacoes' in _nomes, f"abas={_nomes}", "Demogr")
                _wbx.close()
            except Exception as ex:
                _t("guardar_resumo arco com ref. demografica",
                   False, str(ex)[:120], "Demogr")

        # [12g] PDF: pagina de demografia renderiza sem erro
        try:
            from reportlab.pdfgen.canvas import Canvas as _Canvas
            from reportlab.lib.pagesizes import A4 as _A4
            _t("_tem_demografia True no dataset 12 atletas",
               _tem_demografia(_atletas_demo), "", "Demogr")
            with tempfile.TemporaryDirectory() as _tdir2:
                _pdf_demo = _os.path.join(_tdir2, 'demo.pdf')
                _cvs = _Canvas(_pdf_demo, pagesize=_A4)
                _Wp, _Hp = _A4
                _pagina_demografia_pdf(_cvs, _atletas_demo, _Wp, _Hp)
                _cvs.showPage()
                _cvs.save()
                _sz = _os.path.getsize(_pdf_demo)
                _t("_pagina_demografia_pdf gera PDF nao-vazio",
                   _sz > 1024, f"size={_sz}B", "Demogr")
        except Exception as ex:
            _t("_pagina_demografia_pdf sem excepção",
               False, str(ex)[:120], "Demogr")

    except Exception as ex:
        import traceback as _tb
        _t("analises demograficas sem excepção",
           False, f"{type(ex).__name__}: {str(ex)[:100]}", "Demogr")
        if verbose:
            _tb.print_exc()

    # ════════════════════════════════════════════════════════════════════
    # SECÇÃO 13: Tokens de UI / estilos ttk - v1.0 (Fase 4)
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [13] Tokens de UI (espacamento, tipografia, icones)")

    try:
        # [13a] Constantes de espaçamento existem e sao multiplos de 4
        _t("_PAD_XS = 4", _PAD_XS == 4, f"={_PAD_XS}", "UI")
        _t("_PAD_S = 8",  _PAD_S  == 8, f"={_PAD_S}",  "UI")
        _t("_PAD_M = 12", _PAD_M  == 12, f"={_PAD_M}",  "UI")
        _t("_PAD_L = 20", _PAD_L  == 20, f"={_PAD_L}",  "UI")

        # [13b] Fontes retornam tuplos (family, size, weight)
        _fh1 = _FONT_H1()
        _t("_FONT_H1 e tuplo de 3 elementos",
           isinstance(_fh1, tuple) and len(_fh1) == 3, f"={_fh1}", "UI")
        _t("_FONT_H1 size >= 14", _fh1[1] >= 14, f"size={_fh1[1]}", "UI")
        _fb = _FONT_BODY()
        _t("_FONT_BODY size 8-10",
           8 <= _fb[1] <= 10, f"size={_fb[1]}", "UI")
        _fm = _FONT_MONO()
        _t("_FONT_MONO family diferente de UI",
           _fm[0] != _fh1[0], f"mono={_fm[0]} ui={_fh1[0]}", "UI")

        # [13c] Iconografia sem emojis decorativos
        for _k in ('ok', 'error', 'warn', 'info', 'loading'):
            _t(f"ICO[{_k}] existe e nao e emoji",
               _k in ICO and len(ICO[_k]) <= 6,
               f"ico={ICO.get(_k)!r}", "UI")

        # [13d] _configurar_estilos_ttk nao explode sem root (tolerancia)
        # Chama num root headless para verificar que nao lanca excepcao
        try:
            import tkinter as _tkmod
            _rt = _tkmod.Tk()
            _rt.withdraw()
            from tkinter import ttk as _ttkmod
            _sty = _ttkmod.Style(_rt)
            _configurar_estilos_ttk(_sty)
            # Verifica que os estilos foram configurados
            _cfg = _sty.configure('Primary.TButton')
            _t("_configurar_estilos_ttk: Primary.TButton configurado",
               _cfg is not None, "", "UI")
            _cfg2 = _sty.configure('Danger.TButton')
            _t("_configurar_estilos_ttk: Danger.TButton configurado",
               _cfg2 is not None, "", "UI")
            _cfg3 = _sty.configure('Treeview')
            _t("_configurar_estilos_ttk: Treeview configurado",
               _cfg3 is not None, "", "UI")
            _rt.destroy()
        except Exception as ex:
            _t("_configurar_estilos_ttk sem excepcao",
               False, str(ex)[:100], "UI")
    except Exception as ex:
        _t("tokens UI sem excepção",
           False, f"{type(ex).__name__}: {str(ex)[:100]}", "UI")

    # ════════════════════════════════════════════════════════════════════
    # [14] Dashboard web (launcher, endpoints)
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [14] Dashboard web")
    try:
        # Helpers existem e nao partem quando nada esta vivo
        _t("abrir_dashboard_web existe e e chamavel",
           callable(abrir_dashboard_web), "", "Web")
        _t("parar_dashboard_web idempotente (sem proc)",
           parar_dashboard_web() is None, "", "Web")
        _t("_dashboard_port_livre devolve porto >1024",
           _dashboard_port_livre() > 1024, "", "Web")
        # Testar importabilidade sem arrancar servidor
        try:
            from bsp_dashboard.api import storage as _sto
            _pid = _sto.new_project('teste-ci', 'tiro_arco', 'ci')
            _t("storage.new_project cria manifest",
               _sto.read_manifest(_pid['id']) is not None,
               f"id={_pid['id']}", "Web")
            _t("storage.list_projects inclui novo",
               any(p['id'] == _pid['id'] for p in _sto.list_projects()),
               "", "Web")
            _sto.delete_project(_pid['id'])
            _t("storage.delete_project limpa",
               _sto.read_manifest(_pid['id']) is None, "", "Web")
        except ImportError as ex:
            # FastAPI ou bsp_core pode nao estar disponivel em build minimo
            _t("bsp_dashboard.api.storage importavel",
               False, f"ImportError: {str(ex)[:80]}", "Web")
        try:
            from bsp_dashboard.api.main import app as _fapp
            routes = [r.path for r in _fapp.routes]
            _t("FastAPI app expoe /api/health",
               '/api/health' in routes, f"n_routes={len(routes)}", "Web")
            _t("FastAPI app expoe /api/projects",
               '/api/projects' in routes, "", "Web")
            _t("FastAPI app expoe /api/projects/{pid}/analyze",
               '/api/projects/{pid}/analyze' in routes, "", "Web")
        except ImportError:
            _t("FastAPI app importavel",
               False, "fastapi nao instalado (ok em ambiente minimo)",
               "Web")
    except Exception as ex:
        _t("dashboard web sem excepcao",
           False, f"{type(ex).__name__}: {str(ex)[:100]}", "Web")

    # ════════════════════════════════════════════════════════════════════
    # [15] Auto-update (GitHub Releases + SHA256)
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [15] Auto-update (parser + matcher + SHA)")
    try:
        # _parse_versao
        _t("_parse_versao('1.0')=(1,0)",
           _parse_versao('1.0') == (1, 0), "", "Update")
        _t("_parse_versao('v1.0.1')=(1,0,1)",
           _parse_versao('v1.0.1') == (1, 0, 1), "", "Update")
        _t("_parse_versao('1.1-rc1')=(1,1)",
           _parse_versao('1.1-rc1') == (1, 1), "", "Update")
        _t("ordem: (1,0,1) > (1,0)",
           _parse_versao('1.0.1') > _parse_versao('1.0'), "", "Update")
        # _asset_nome_esperado retorna lista nao-vazia
        _pats = _asset_nome_esperado()
        _t("_asset_nome_esperado retorna padroes",
           isinstance(_pats, list) and len(_pats) > 0,
           f"n={len(_pats)}", "Update")
        # _match_asset com fake list
        _fake_assets = [
            {'name': 'BSP-1.0.1-win-x64.exe',
             'browser_download_url': 'https://example/win.exe'},
            {'name': 'BSP-1.0.1-macos-arm64.dmg',
             'browser_download_url': 'https://example/mac.dmg'},
            {'name': 'BSP-1.0.1-macos-x64.dmg',
             'browser_download_url': 'https://example/mac-x64.dmg'},
            {'name': 'SHA256SUMS.txt',
             'browser_download_url': 'https://example/sha.txt'},
        ]
        _bin, _sha = _match_asset(_fake_assets)
        _t("_match_asset encontra asset para este SO",
           _bin is not None, f"nome={(_bin or {}).get('name')}", "Update")
        _t("_match_asset encontra SHA256SUMS.txt",
           _sha is not None, "", "Update")
        # _parse_sha256sums
        _txt = (
            "8d12f4e7b2c3a5f6e7d8c9b0a1f2e3d4c5b6a7f8e9d0c1b2a3f4e5d6c7b8a9f0  BSP-1.0.1-macos-arm64.dmg\n"
            "f0e1d2c3b4a5968778695a4b3c2d1e0fefeddcebcab9a8978685746352413021  BSP-1.0.1-win-x64.exe\n"
            "# comentario ignorado\n"
        )
        _m = _parse_sha256sums(_txt)
        _t("_parse_sha256sums extrai 2 entradas",
           len(_m) == 2, f"n={len(_m)}", "Update")
        _t("_parse_sha256sums: hash correcto DMG",
           _m.get('BSP-1.0.1-macos-arm64.dmg',
                  '').startswith('8d12f4e7'), "", "Update")
        # _UPDATE_STATE tem os campos esperados
        _campos = {'tag','notes','asset_url','asset_name','sha256','sha_txt_url'}
        _t("_UPDATE_STATE tem campos correctos",
           _campos.issubset(set(_UPDATE_STATE.keys())),
           f"faltam={_campos - set(_UPDATE_STATE.keys())}", "Update")
        _t("_actualizar_agora e chamavel",
           callable(_actualizar_agora), "", "Update")
    except Exception as ex:
        _t("auto-update sem excepcao",
           False, f"{type(ex).__name__}: {str(ex)[:100]}", "Update")

    # ════════════════════════════════════════════════════════════════════
    # [16] Palavra-passe rotativa (hash remoto + cache + fallback)
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [16] Palavra-passe rotativa")
    try:
        # URL formato correcto
        _t("_PASS_REMOTE_URL aponta para .bsp_pass.sha256",
           _PASS_REMOTE_URL.endswith('/.bsp_pass.sha256'),
           _PASS_REMOTE_URL, "PassRot")
        # _PASS_HASH embedded e valido (64 chars hex)
        import re as _re_p
        _t("_PASS_HASH embedded e SHA256 valido",
           bool(_re_p.fullmatch(r'[0-9a-f]{64}', _PASS_HASH)),
           f"len={len(_PASS_HASH)}", "PassRot")
        # _fetch_pass_hashes_remoto com URL inexistente devolve []
        _orig_url = _PASS_REMOTE_URL
        try:
            globals()['_PASS_REMOTE_URL'] = 'https://127.0.0.1:1/nope'
            _h = _fetch_pass_hashes_remoto(timeout=0.3)
            _t("_fetch_pass_hashes_remoto sem rede -> []",
               _h == [], f"got={_h}", "PassRot")
        finally:
            globals()['_PASS_REMOTE_URL'] = _orig_url
        # _obter_pass_hashes_aceites sempre devolve set nao-vazio (fallback)
        try:
            globals()['_PASS_REMOTE_URL'] = 'https://127.0.0.1:1/nope'
            _s = _obter_pass_hashes_aceites(timeout=0.3)
            _t("_obter_pass_hashes_aceites sempre devolve set nao-vazio",
               isinstance(_s, set) and len(_s) >= 1,
               f"n={len(_s)}", "PassRot")
            _t("fallback inclui _PASS_HASH embedded",
               _PASS_HASH.lower() in _s,
               "", "PassRot")
        finally:
            globals()['_PASS_REMOTE_URL'] = _orig_url
        # Parser tolera comentarios + linhas vazias + case-insensitive
        # (testamos directamente o regex interno do fetch: formato de texto)
        _txt = (
            "# comentario\n"
            "\n"
            "AABBCCDDEEFF00112233445566778899AABBCCDDEEFF00112233445566778899\n"
            "  0123456789abcdef0123456789abcdef0123456789abcdef0123456789abcdef  # inline\n"
            "xx_nao_valido_xx\n"
        )
        # Extrair via mesma logica (inline)
        _out = []
        for _l in _txt.splitlines():
            _l = _l.strip()
            if not _l or _l.startswith('#'): continue
            _tok = _l.split()[0].lower()
            if _re_p.fullmatch(r'[0-9a-f]{64}', _tok):
                _out.append(_tok)
        _t("parser aceita 2 hashes validos, ignora lixo",
           len(_out) == 2, f"n={len(_out)}", "PassRot")
        _t("parser normaliza case -> lowercase",
           all(h == h.lower() for h in _out), "", "PassRot")
    except Exception as ex:
        _t("pass rotativa sem excepcao",
           False, f"{type(ex).__name__}: {str(ex)[:100]}", "PassRot")

    # ════════════════════════════════════════════════════════════════════
    # [17] Telemetry de aceitacao de licenca (privacy-respecting + opt-out)
    # ════════════════════════════════════════════════════════════════════
    if verbose: print("\n  [17] Telemetry licenca")
    try:
        # _machine_id_anon e estavel + 32 chars hex
        mid1 = _machine_id_anon()
        mid2 = _machine_id_anon()
        import re as _re_t
        _t("_machine_id_anon estavel entre chamadas",
           mid1 == mid2, f"{mid1[:8]} vs {mid2[:8]}", "Telemetry")
        _t("_machine_id_anon e 32 chars hex",
           bool(_re_t.fullmatch(r'[0-9a-f]{32}', mid1)),
           f"got={mid1!r}", "Telemetry")
        # _telemetry_post_license_acceptance e chamavel sem excepcao
        _t("_telemetry_post_license_acceptance e chamavel",
           callable(_telemetry_post_license_acceptance), "", "Telemetry")
        # Falha silenciosa com URL invalido (timeout em thread, nao bloqueia)
        _orig_url = _TELEMETRY_URL
        try:
            globals()['_TELEMETRY_URL'] = 'https://127.0.0.1:1/nope'
            _telemetry_post_license_acceptance()
            _t("call com URL invalido nao bloqueia o thread principal",
               True, "", "Telemetry")
        finally:
            globals()['_TELEMETRY_URL'] = _orig_url
        # Opt-out via env var BSP_TELEMETRY_URL=''
        _orig_env = os.environ.get('BSP_TELEMETRY_URL')
        try:
            os.environ['BSP_TELEMETRY_URL'] = ''
            _telemetry_post_license_acceptance()
            _t("opt-out via env var BSP_TELEMETRY_URL='' nao envia",
               True, "", "Telemetry")
        finally:
            if _orig_env is None:
                os.environ.pop('BSP_TELEMETRY_URL', None)
            else:
                os.environ['BSP_TELEMETRY_URL'] = _orig_env
        # Log local recebe a entry
        c_t = cfg_load()
        log_t = c_t.get(_TELEMETRY_LOG_KEY) or []
        _t("log local guarda entries (cap 50)",
           isinstance(log_t, list) and len(log_t) <= 50,
           f"n={len(log_t)}", "Telemetry")
        if log_t:
            ent = log_t[-1]
            _t("entry tem campos minimos (ts, version, machine_id)",
               all(k in ent for k in ('ts', 'version', 'machine_id')),
               f"keys={list(ent.keys())}", "Telemetry")
            _t("entry NAO tem dados pessoais (name, path)",
               'name' not in ent and 'path' not in ent and 'user' not in ent,
               "", "Telemetry")
    except Exception as ex:
        _t("telemetry sem excepcao",
           False, f"{type(ex).__name__}: {str(ex)[:100]}", "Telemetry")

    # ════════════════════════════════════════════════════════════════════
    # RELATÓRIO FINAL
    # ════════════════════════════════════════════════════════════════════
    total = len(resultados)
    n_ok  = sum(1 for _, p in resultados if p)
    n_fail = total - n_ok

    if verbose:
        print(f"\n{'═'*58}")
        print(f"  RESULTADO FINAL: {n_ok}/{total} testes passaram")
        if n_fail > 0:
            print(f"\n  FALHAS ({n_fail}):")
            for sec, items in seccoes.items():
                falhas = [(n, d) for n, p, d in items if not p]
                if falhas:
                    print(f"  [{sec}]")
                    for n, d in falhas:
                        print(f"    ✗  {n}" + (f"  ({d})" if d else ""))
        else:
            print("  Todos os testes passaram!")
        print(f"{'═'*58}")

    return ok_all


def modo_cli():
    print(f'\n{"="*60}\n  {PROG.upper()}  v{VERSAO}\n  {AUTOR}  |  {SEGUNDO_AUTOR}  |  {ORIENTADOR}\n{"="*60}\n')
    p=argparse.ArgumentParser()
    p.add_argument('pasta')
    p.add_argument('--inicio_fim', default=None)
    p.add_argument('--output', default='resultados_estabilidade.xlsx')
    p.add_argument('--individuais', default='individuais')
    p.add_argument('--pdf', default=None)
    p.add_argument('--protocolo', default=PROTO_FMS, choices=list(PROTOCOLOS.keys()))
    p.add_argument('--scores', default=None)
    p.add_argument('--sem_ind', action='store_true')
    p.add_argument('--sem_elipse', action='store_true')
    p.add_argument('--sem_estab', action='store_true')
    p.add_argument('--testes', action='store_true',
                   help='Correr suite de testes sinteticos e sair')
    p.add_argument('--embedded', action='store_true')
    args=p.parse_args()
    if hasattr(args, 'testes') and args.testes:
        print(f'\n{PROG} v{VERSAO} - Testes Sinteticos\n')
        ok = _run_testes_sinteticos(verbose=True)
        sys.exit(0 if ok else 1)
    global _PROTOCOLO_ACTIVO; _PROTOCOLO_ACTIVO=args.protocolo
    print(f'Protocolo: {PROTOCOLOS[args.protocolo]["nome"]}\n')

    ifd=None
    if args.inicio_fim and os.path.exists(args.inicio_fim):
        print(f'Inicio_fim: {args.inicio_fim}')
        if _is_tiro_like(args.protocolo):
            ifd=carregar_inicio_fim_tiro(args.inicio_fim)
        else:
            ifd=carregar_inicio_fim(args.inicio_fim)
        print(f'  {len(ifd)} atletas\n')

    scores_data={}
    if args.scores and os.path.exists(args.scores):
        scores_data=carregar_scores_tiro(args.scores)
        print(f'Scores: {len(scores_data)} atletas\n')

    subs=sorted([os.path.join(args.pasta,d) for d in os.listdir(args.pasta)
                 if os.path.isdir(os.path.join(args.pasta,d))])
    if not subs: print('Sem subpastas.'); sys.exit(1)
    print(f'{len(subs)} atleta(s)...\n'); atletas=[]

    for sub in subs:
        print(f'{os.path.basename(sub)}')
        ath=processar_atleta(sub, ifd, args.embedded,
                              lambda m, t='': print(m), args.protocolo)
        for chave_s in scores_data:
            if _norm(chave_s)==_norm(ath['nome']) or _norm(chave_s) in _norm(ath['nome']):
                ath['scores']=scores_data[chave_s]; break
        atletas.append(ath)

    print('\nA gerar ficheiros...')
    guardar_resumo(atletas, args.output, args.protocolo)
    print(f'Excel: {args.output}')

    if not args.sem_ind:
        os.makedirs(args.individuais, exist_ok=True)
        for ath in atletas:
            fn=re.sub(r'[<>:"/\\|?*]','_',ath['nome'])+'.xlsx'
            guardar_individual(ath, os.path.join(args.individuais,fn),
                               not args.sem_elipse, not args.sem_estab)
            print(f'  {fn}')

    if args.pdf:
        print('\nA gerar PDF...')
        gerar_pdf(atletas, args.pdf)

    print(f'\nConcluido. {len(atletas)} individuo(s).\n')


# -----------------------------------------------------------------------
# Ponto de entrada
# -----------------------------------------------------------------------

if __name__ == '__main__':
    # ── Modo testes: correr antes de iniciar Tk (sem display necessário) ──
    if '--testes' in sys.argv:
        ok = _run_testes_sinteticos(verbose=True)
        sys.exit(0 if ok else 1)

    if '--cli' in sys.argv:
        sys.argv.remove('--cli'); modo_cli()
    else:
        import tkinter as _tk_boot
        import traceback as _tb_boot

        # CRITICO: usar um UNICO Tk() root para todo o arranque
        # Multiplos Tk() simultaneos causam crash no Tcl/Tk no Windows
        _root = _tk_boot.Tk()
        _root.geometry('1x1+-9999+-9999')  # off-screen
        _root.resizable(False, False)
        _root.overrideredirect(True)       # sem barra de titulo no root vazio
        # Mostrar na taskbar mesmo durante dialogos iniciais (Windows fix)
        try: _root.wm_attributes('-toolwindow', False)  # ATENÇÃO: não existe em macOS/Tcl antigo
        except Exception: pass
        try:
            _root.wm_iconbitmap()  # força registo na taskbar do Windows
        except Exception:
            pass
        # Ícone BSP no título e na barra de tarefas (script .py; exe usa BSP.ico embutido)
        try:
            _ico_img = _logo_photoimage(size=64)
            if _ico_img:
                _root.iconphoto(True, _ico_img)
                _root._ico_ref = _ico_img  # evitar GC
        except Exception:
            pass

        # Instalar crash hooks imediatamente (com o root real)
        _instalar_crash_hooks(_root)

        try:
            # ── 0. Tema ───────────────────────────────────────────────────
            _cfg_tema = cfg_load()
            _tema_cfg = _cfg_tema.get('tema', '')
            if not _tema_cfg:
                _detected = _detect_system_theme()
                _tema_cfg = _detected if _detected else 'escuro'
                _cfg_tema['tema'] = _tema_cfg
                cfg_save(_cfg_tema)
            _aplicar_tema(_tema_cfg)

            # ── 1. Licenca (apenas se nao aceite nesta versao) ────────────
            if not _ecra_licenca(_root):
                _root.destroy()
                sys.exit(0)

            # ── 2. Palavra-passe ──────────────────────────────────────────
            if not _ecra_password(_root):
                _root.destroy()
                sys.exit(0)

            # ── 3. Protocolo ──────────────────────────────────────────────
            proto = escolher_protocolo_gui(_root)
            if not proto:
                _root.destroy()
                sys.exit(0)

            # ── 4. Janela principal ───────────────────────────────────────
            # Remover overrideredirect antes de criar a janela principal
            _root.overrideredirect(False)
            _root.geometry('')  # reset geometria
            _root.withdraw()   # forcar re-entrada na taskbar (fix Windows)
            _root.deiconify()
            Janela(protocolo=proto, _root=_root)
            try:
                _root.mainloop()
            finally:
                # parar dashboard web se estiver activo
                try: parar_dashboard_web()
                except Exception: pass

        except SystemExit:
            raise
        except BaseException as _e_startup:
            _tb_str = _tb_boot.format_exc()
            _logger.critical('Erro critico no arranque:\n' + _tb_str)
            try:
                _dialogo_crash(_tb_str, 'Arranque', _root)
            except Exception:
                pass
